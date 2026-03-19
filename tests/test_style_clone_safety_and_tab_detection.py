import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.docx_io.style_clone import (
    _build_style_config_from_paragraph,
    _infer_heading_pattern,
    clone_scene_style_from_docx,
)
from src.scene.manager import load_scene_from_data
from src.scene.schema import StyleConfig


def _rewrite_zip_entry(path: Path, entry_name: str, transform) -> None:
    tmp_path = path.with_suffix(".rewrite.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == entry_name:
                data = transform(data)
            zout.writestr(item, data)
    tmp_path.replace(path)


def test_infer_heading_pattern_allows_explicit_tab_separator():
    level1 = _infer_heading_pattern("1\t研究目的")
    assert level1 is not None
    assert level1[0] == "heading1"
    assert level1[1] == "arabic"
    assert level1[3] == "\t"

    level2 = _infer_heading_pattern("1.1\t研究方法")
    assert level2 is not None
    assert level2[0] == "heading2"
    assert level2[3] == "\t"


def test_infer_heading_pattern_still_rejects_toc_like_tab_page_suffix():
    assert _infer_heading_pattern("1\t研究目的\t12") is None


def test_clone_scene_style_from_docx_does_not_modify_source_file(tmp_path):
    source_path = tmp_path / "clone_source.docx"
    doc = Document()
    doc.add_heading("测试标题", level=1)
    doc.save(source_path)

    _rewrite_zip_entry(
        source_path,
        "word/_rels/document.xml.rels",
        lambda data: data.replace(
            b"</Relationships>",
            (
                b'<Relationship Id="rIdNULL" '
                b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                b'Target="NULL"/></Relationships>'
            ),
        ),
    )
    before_bytes = source_path.read_bytes()

    cfg = load_scene_from_data({})
    summary = clone_scene_style_from_docx(cfg, source_path)

    assert summary["styles_updated"]
    assert source_path.read_bytes() == before_bytes
    with zipfile.ZipFile(source_path, "r") as zf:
        rels = zf.read("word/_rels/document.xml.rels")
    assert b'Target="NULL"' in rels


def test_build_style_config_from_paragraph_preserves_hanging_indent():
    doc = Document()
    para = doc.add_paragraph("参考文献条目")
    para.paragraph_format.left_indent = Pt(24)
    para.paragraph_format.first_line_indent = Pt(-24)

    sc = _build_style_config_from_paragraph(para, StyleConfig())

    assert sc.left_indent_chars == 24.0
    assert sc.left_indent_unit == "pt"
    assert sc.first_line_indent_chars == 0.0
    assert sc.hanging_indent_chars == 24.0
    assert sc.hanging_indent_unit == "pt"

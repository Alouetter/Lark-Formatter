import json
import zipfile
from pathlib import Path

from docx import Document

from src.engine.pipeline import Pipeline
from src.scene.schema import SceneConfig, StyleConfig


def _stub_word_page_spans(monkeypatch, spans):
    monkeypatch.setattr(
        "src.engine.page_scope._resolve_word_paragraph_page_spans",
        lambda doc, source_doc_path=None, timeout_sec=None: list(spans),
    )


def _inject_null_relationship(doc_path: Path) -> None:
    temp_path = doc_path.with_name(f"{doc_path.stem}_rewritten{doc_path.suffix}")
    with zipfile.ZipFile(doc_path, "r") as zin:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/_rels/document.xml.rels":
                    marker = b"</Relationships>"
                    bogus_rel = (
                        b'<Relationship Id="rId999" '
                        b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                        b'Target="media/NULL"/>'
                    )
                    data = data.replace(marker, bogus_rel + marker, 1)
                zout.writestr(item, data)
    temp_path.replace(doc_path)


def test_pipeline_removes_stale_disabled_outputs(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "output_options.docx"
    doc = Document()
    doc.add_paragraph("hello")
    doc.save(str(doc_path))

    cfg = SceneConfig()
    first_result = Pipeline(cfg).run(str(doc_path))
    first_paths = {
        key: Path(value)
        for key, value in first_result.output_paths.items()
        if key in {"final", "compare", "report_json", "report_md"}
    }
    assert first_paths
    assert all(path.exists() for path in first_paths.values())

    cfg.output.final_docx = False
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False

    second_result = Pipeline(cfg).run(str(doc_path))
    assert "final" not in second_result.output_paths
    assert "compare" not in second_result.output_paths
    assert "report_json" not in second_result.output_paths
    assert "report_md" not in second_result.output_paths
    assert all(not path.exists() for path in first_paths.values())


def test_pipeline_sanitize_uses_working_copy_without_mutating_source(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "sanitize_source.docx"
    doc = Document()
    doc.add_paragraph("hello")
    doc.save(str(doc_path))
    _inject_null_relationship(doc_path)
    before_bytes = doc_path.read_bytes()

    cfg = SceneConfig()
    cfg.pipeline = []
    cfg.output.final_docx = False
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False

    result = Pipeline(cfg).run(str(doc_path))

    assert result.success is True
    assert doc_path.read_bytes() == before_bytes
    assert not any(".pipeline." in p.name for p in tmp_path.iterdir())


def test_pipeline_range_mode_skips_global_page_and_style_rules(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")
    _stub_word_page_spans(monkeypatch, [(1, 1)])

    doc_path = tmp_path / "range_global_skip.docx"
    doc = Document()
    doc.add_paragraph("hello")
    doc.save(str(doc_path))

    original_doc = Document(str(doc_path))
    before_top_margin = round(float(original_doc.sections[0].top_margin.cm), 2)
    before_font_name = original_doc.styles["Normal"].font.name
    before_bold = original_doc.styles["Normal"].font.bold

    cfg = SceneConfig()
    cfg.pipeline = ["page_setup", "style_manager"]
    cfg.page_setup.margin.top_cm = 1.23
    cfg.styles["normal"] = StyleConfig(
        font_cn="黑体",
        font_en="Courier New",
        size_pt=16,
        bold=True,
    )
    cfg.format_scope.mode = "manual"
    cfg.format_scope.page_ranges_text = "1-1"
    cfg.output.final_docx = False
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False

    result = Pipeline(cfg).run(str(doc_path))

    assert result.success is True
    skipped_rules = {
        rec.rule_name
        for rec in result.tracker.records
        if rec.target == "range_scope" and rec.change_type == "skip"
    }
    assert {"page_setup", "style_manager"}.issubset(skipped_rules)
    assert round(float(result.doc.sections[0].top_margin.cm), 2) == before_top_margin
    assert result.doc.styles["Normal"].font.name == before_font_name
    assert result.doc.styles["Normal"].font.bold == before_bold


def test_pipeline_range_mode_validation_ignores_out_of_scope_errors(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")
    _stub_word_page_spans(monkeypatch, [(1, 1), (1, 1), (2, 2)])

    doc_path = tmp_path / "range_validation.docx"
    doc = Document()
    doc.add_paragraph("第一页")
    doc.add_paragraph("第一页补充")
    heading = doc.add_paragraph("1\t第二页标题")
    heading.style = doc.styles["Heading 1"]
    doc.save(str(doc_path))

    cfg = SceneConfig()
    cfg.pipeline = ["validation"]
    cfg.format_scope.mode = "manual"
    cfg.format_scope.page_ranges_text = "1-1"
    cfg.output.final_docx = False
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False

    result = Pipeline(cfg).run(str(doc_path))

    assert result.success is True
    assert not any(rec.failure_reason for rec in result.tracker.records if rec.rule_name == "validation")


def test_pipeline_report_records_input_file(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "report_input.docx"
    doc = Document()
    doc.add_paragraph("hello")
    doc.save(str(doc_path))

    cfg = SceneConfig()
    cfg.pipeline = []
    cfg.output.final_docx = False
    cfg.output.compare_docx = False
    cfg.output.report_json = True
    cfg.output.report_markdown = False

    result = Pipeline(cfg).run(str(doc_path))

    report_path = Path(result.output_paths["report_json"])
    payload = json.loads(report_path.read_text(encoding="utf-8"))
    assert payload["input_file"] == str(doc_path)

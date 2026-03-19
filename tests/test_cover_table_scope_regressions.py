from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from src.engine.change_tracker import ChangeTracker
from src.engine.doc_tree import DocSection, DocTree
from src.engine.rules.equation_table_format import EquationTableFormatRule
from src.engine.rules.formula_style import FormulaStyleRule
from src.engine.rules.section_format import SectionFormatRule
from src.engine.rules.style_manager import StyleManagerRule
from src.engine.rules.table_format import TableFormatRule
from src.scene.manager import load_default_scene

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _doc_tree(doc: Document) -> DocTree:
    tree = DocTree()
    tree.build(doc)
    return tree


def _table_alignment_values(tbl) -> list[str]:
    return tbl._tbl.xpath(
        "./*[namespace-uri()='%s' and local-name()='tblPr']"
        "/*[namespace-uri()='%s' and local-name()='jc']"
        "/@*[local-name()='val']"
        % (_W_NS, _W_NS)
    )


def _cell_has_bold(cell) -> bool:
    return bool(
        cell._tc.xpath(
            ".//*[namespace-uri()='%s' and local-name()='b']" % _W_NS
        )
    )


def _cell_border_values(cell, side: str) -> list[str]:
    return cell._tc.xpath(
        "./*[namespace-uri()='%s' and local-name()='tcPr']"
        "/*[namespace-uri()='%s' and local-name()='tcBorders']"
        f"/*[namespace-uri()='%s' and local-name()='{side}']"
        "/@*[local-name()='val']"
        % (_W_NS, _W_NS, _W_NS)
    )


def _table_border_values(tbl, side: str) -> list[str]:
    return tbl._tbl.xpath(
        "./*[namespace-uri()='%s' and local-name()='tblPr']"
        "/*[namespace-uri()='%s' and local-name()='tblBorders']"
        f"/*[namespace-uri()='%s' and local-name()='{side}']"
        "/@*[local-name()='val']"
        % (_W_NS, _W_NS, _W_NS)
    )



def _paragraph_indent_values(para) -> list[str]:
    return para._p.xpath(
        "./*[namespace-uri()='%s' and local-name()='pPr']"
        "/*[namespace-uri()='%s' and local-name()='ind']"
        "/@*[local-name()='firstLineChars' or local-name()='firstLine']"
        % (_W_NS, _W_NS)
    )


def test_table_format_skips_leading_cover_table_but_formats_body_table() -> None:
    doc = Document()
    cover_tbl = doc.add_table(rows=1, cols=2)
    cover_tbl.cell(0, 0).text = "姓名"
    cover_tbl.cell(0, 1).text = "张三"

    doc.add_paragraph("硕士学位论文")
    doc.add_paragraph("第一章 绪论")

    body_tbl = doc.add_table(rows=1, cols=2)
    body_tbl.cell(0, 0).text = "项目"
    body_tbl.cell(0, 1).text = "数值"

    cfg = load_default_scene()
    TableFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {"doc_tree": _doc_tree(doc)},
    )

    assert _cell_has_bold(cover_tbl.cell(0, 0)) is False
    assert _cell_has_bold(body_tbl.cell(0, 0)) is True


def test_equation_table_format_skips_leading_cover_equation_table() -> None:
    doc = Document()
    cover_tbl = doc.add_table(rows=1, cols=2)
    cover_tbl.cell(0, 0).text = "x^2+y^2=z^2"
    cover_tbl.cell(0, 1).text = ""

    doc.add_paragraph("硕士学位论文")
    doc.add_paragraph("第一章 绪论")

    body_tbl = doc.add_table(rows=1, cols=2)
    body_tbl.cell(0, 0).text = "a^2+b^2=c^2"
    body_tbl.cell(0, 1).text = ""

    cfg = load_default_scene()
    cfg.equation_table_format.enabled = True
    tracker = ChangeTracker()
    context = {"doc_tree": _doc_tree(doc)}

    EquationTableFormatRule().apply(doc, cfg, tracker, context)

    assert cover_tbl.cell(0, 1).text.strip() == ""
    assert body_tbl.cell(0, 1).text.strip() == "(1.1)"


def test_formula_style_skips_leading_cover_equation_table() -> None:
    doc = Document()
    cover_tbl = doc.add_table(rows=1, cols=2)
    cover_tbl.cell(0, 0).text = "a^2+b^2=c^2"
    cover_tbl.cell(0, 1).text = "(1.1)"

    doc.add_paragraph("硕士学位论文")
    doc.add_paragraph("第一章 绪论")

    body_tbl = doc.add_table(rows=1, cols=2)
    body_tbl.cell(0, 0).text = "x_i^2+y_i^2"
    body_tbl.cell(0, 1).text = "(1.2)"

    cfg = load_default_scene()
    cfg.formula_style.enabled = True
    cfg.formula_table.table_alignment = "right"
    cfg.formula_table.formula_font_name = "Cambria Math"

    FormulaStyleRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {"doc_tree": _doc_tree(doc)},
    )

    assert cover_tbl.cell(0, 0).paragraphs[0].runs[0].font.name is None
    assert body_tbl.cell(0, 0).paragraphs[0].runs[0].font.name == "Cambria Math"
    assert _table_alignment_values(cover_tbl) != ["right"]
    assert _table_alignment_values(body_tbl) == ["right"]


def test_section_format_does_not_center_leading_cover_table() -> None:
    doc = Document()
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    doc.add_paragraph("硕士学位论文")
    doc.add_paragraph("第一章 绪论")

    body_tbl = doc.add_table(rows=1, cols=1)
    body_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    cfg = load_default_scene()
    context = {
        "doc_tree": _doc_tree(doc),
        "format_scope": cfg.format_scope,
    }

    SectionFormatRule().apply(doc, cfg, ChangeTracker(), context)

    assert cover_tbl.alignment == WD_TABLE_ALIGNMENT.LEFT
    assert body_tbl.alignment == WD_TABLE_ALIGNMENT.CENTER



def test_style_manager_keeps_cover_content_unmodified_even_if_cover_scope_is_forced_on() -> None:
    doc = Document()
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.cell(0, 0).text = "????"
    cover_tbl.cell(0, 0).paragraphs[0].paragraph_format.first_line_indent = Pt(9)

    cover_para = doc.add_paragraph("??????")
    doc.add_paragraph("??? ??")
    body_para = doc.add_paragraph("???????")

    cfg = load_default_scene()
    cfg.format_scope.sections["cover"] = True
    tree = DocTree()
    tree.sections = [
        DocSection("cover", 0, 0, confidence=12.0),
        DocSection("body", 1, 2, confidence=12.0),
    ]
    context = {
        "doc_tree": tree,
        "format_scope": cfg.format_scope,
    }

    cover_table_indent_before = _paragraph_indent_values(cover_tbl.cell(0, 0).paragraphs[0])
    assert doc.styles["Normal"].paragraph_format.first_line_indent is None

    StyleManagerRule().apply(doc, cfg, ChangeTracker(), context)

    assert doc.styles["Normal"].paragraph_format.first_line_indent is None

    SectionFormatRule().apply(doc, cfg, ChangeTracker(), context)

    assert _paragraph_indent_values(cover_tbl.cell(0, 0).paragraphs[0]) == cover_table_indent_before
    assert _paragraph_indent_values(cover_para) == []
    assert _paragraph_indent_values(body_para)


def test_table_format_applies_three_line_even_to_md_cleanup_marked_tables() -> None:
    doc = Document()
    doc.add_paragraph("硕士学位论文")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("这是正文内容。")

    body_tbl = doc.add_table(rows=3, cols=2)
    body_tbl.cell(0, 0).text = "项目"
    body_tbl.cell(0, 1).text = "数值"
    body_tbl.cell(1, 0).text = "A"
    body_tbl.cell(1, 1).text = "1"
    body_tbl.cell(2, 0).text = "B"
    body_tbl.cell(2, 1).text = "2"

    cfg = load_default_scene()
    cfg.normal_table_border_mode = "three_line"

    TableFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": _doc_tree(doc),
            "md_cleanup_table_element_ids": {id(body_tbl._tbl)},
        },
    )

    assert _table_border_values(body_tbl, "insideH") == ["none"]
    assert _table_border_values(body_tbl, "insideV") == ["none"]
    assert _cell_border_values(body_tbl.cell(0, 0), "top") == ["single"]
    assert _cell_border_values(body_tbl.cell(0, 0), "bottom") == ["single"]
    assert _cell_border_values(body_tbl.cell(2, 0), "bottom") == ["single"]

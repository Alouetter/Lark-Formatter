from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from types import SimpleNamespace

from src.engine.rules.toc_format import (
    TocFormatRule,
    _build_toc_entries,
    _native_toc_field_instruction,
    _normalize_toc_styles_in_doc,
)
from src.scene.manager import load_scene_from_data
from src.scene.schema import HeadingLevelConfig

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _ensure_toc_heading_style(doc: Document):
    try:
        return doc.styles["TOC Heading"]
    except KeyError:
        return doc.styles.add_style("TOC Heading", WD_STYLE_TYPE.PARAGRAPH)


def test_normalize_toc_heading_style_detaches_heading_numbering_inheritance():
    doc = Document()
    toc_heading = _ensure_toc_heading_style(doc)
    toc_heading.base_style = doc.styles["Heading 1"]

    cfg = load_scene_from_data({})
    _normalize_toc_styles_in_doc(doc, cfg)

    toc_heading = doc.styles["TOC Heading"]
    assert toc_heading.base_style is not None
    assert toc_heading.base_style.name == "Normal"

    ppr = toc_heading.element.find(_w("pPr"))
    assert ppr is not None
    assert ppr.find(_w("numPr")) is None
    outline = ppr.find(_w("outlineLvl"))
    assert outline is not None
    assert outline.get(_w("val")) == "9"


def test_format_existing_toc_title_para_keeps_title_out_of_heading_outline():
    doc = Document()
    toc_heading = _ensure_toc_heading_style(doc)
    toc_heading.base_style = doc.styles["Heading 1"]
    para = doc.add_paragraph("目录")

    cfg = load_scene_from_data({})
    _normalize_toc_styles_in_doc(doc, cfg)
    TocFormatRule._format_existing_toc_paras(doc, cfg)

    assert para.style.name == "TOC Heading"
    ppr = para._element.find(_w("pPr"))
    assert ppr is not None
    assert ppr.find(_w("numPr")) is None
    outline = ppr.find(_w("outlineLvl"))
    assert outline is not None
    assert outline.get(_w("val")) == "9"


def test_native_toc_field_instruction_tracks_v2_include_in_toc_levels():
    cfg = load_scene_from_data({})
    cfg.heading_numbering_v2.level_bindings["heading2"].include_in_toc = False
    cfg.heading_numbering_v2.level_bindings["heading4"].include_in_toc = True

    instruction = _native_toc_field_instruction(cfg)

    assert '\\o "1-4"' in instruction
    assert '\\h' in instruction
    assert '\\z' in instruction
    assert '\\u' in instruction
    assert '\\t "' in instruction
    assert 'Front Matter Heading Unnumbered,1' in instruction
    assert 'Heading 1 Unnumbered,1' in instruction


def test_build_toc_entries_respects_hidden_levels_but_keeps_numbering_counters():
    doc = Document()
    doc.add_paragraph("3 Title A")
    doc.add_paragraph("3.5 Title B")
    doc.add_paragraph("3.5.7 Title C")
    doc.add_paragraph("3.5.7.1 Title D")

    cfg = load_scene_from_data({})
    cfg.heading_numbering.levels = {
        "heading1": HeadingLevelConfig(format="arabic", template="{current}", separator=" "),
        "heading2": HeadingLevelConfig(format="arabic_dotted", template="{parent}.{current}", separator=" "),
        "heading3": HeadingLevelConfig(format="arabic_dotted", template="{parent}.{current}", separator=" "),
        "heading4": HeadingLevelConfig(format="arabic_dotted", template="{parent}.{current}", separator=" "),
    }
    cfg.heading_numbering_v2.level_bindings["heading1"].start_at = 3
    cfg.heading_numbering_v2.level_bindings["heading2"].start_at = 5
    cfg.heading_numbering_v2.level_bindings["heading2"].include_in_toc = False
    cfg.heading_numbering_v2.level_bindings["heading2"].restart_on = None
    cfg.heading_numbering_v2.level_bindings["heading3"].start_at = 7
    cfg.heading_numbering_v2.level_bindings["heading3"].restart_on = "heading1"
    cfg.heading_numbering_v2.level_bindings["heading4"].enabled = True
    cfg.heading_numbering_v2.level_bindings["heading4"].include_in_toc = True
    cfg.heading_numbering_v2.level_bindings["heading4"].restart_on = "heading3"

    headings = [
        SimpleNamespace(level="heading1", text="3 Title A", para_index=0),
        SimpleNamespace(level="heading2", text="3.5 Title B", para_index=1),
        SimpleNamespace(level="heading3", text="3.5.7 Title C", para_index=2),
        SimpleNamespace(level="heading4", text="3.5.7.1 Title D", para_index=3),
    ]

    entries = _build_toc_entries(
        doc,
        headings,
        cfg.heading_numbering.levels,
        cfg.heading_numbering_v2.level_bindings,
    )

    assert [entry["level"] for entry in entries] == ["heading1", "heading3", "heading4"]
    assert [entry["numbering"] for entry in entries] == ["3", "3.5.7", "3.5.7.1"]


def test_build_toc_entries_renders_tab_separator_as_spaces():
    doc = Document()
    doc.add_paragraph("绪论")

    cfg = load_scene_from_data({})
    cfg.heading_numbering.levels = {
        "heading1": HeadingLevelConfig(format="arabic", template="{current}", separator="\t"),
    }
    cfg.heading_numbering_v2.level_bindings["heading1"].enabled = True

    headings = [
        SimpleNamespace(level="heading1", text="绪论", para_index=0),
    ]

    entries = _build_toc_entries(
        doc,
        headings,
        cfg.heading_numbering.levels,
        cfg.heading_numbering_v2.level_bindings,
    )

    assert len(entries) == 1
    assert entries[0]["sep"] == "   "
    assert "\t" not in entries[0]["sep"]


def test_normalize_toc_styles_uses_level2_style_for_deeper_toc_levels():
    doc = Document()
    cfg = load_scene_from_data({})
    cfg.heading_numbering_v2.level_bindings["heading4"].include_in_toc = True
    cfg.styles["toc_level2"].alignment = "right"

    _normalize_toc_styles_in_doc(doc, cfg)

    toc4 = doc.styles["TOC 4"]
    assert toc4.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_build_toc_entries_keeps_full_title_when_body_prefix_already_stripped():
    doc = Document()
    doc.add_paragraph("研究内容及意义")

    cfg = load_scene_from_data({})
    cfg.heading_numbering.levels = {
        "heading2": HeadingLevelConfig(
            format="chinese_section",
            template="第{current}节",
            separator="\u3000",
        ),
    }
    cfg.heading_numbering_v2.level_bindings["heading2"].enabled = True
    cfg.heading_numbering_v2.level_bindings["heading2"].start_at = 6
    cfg.heading_numbering_v2.level_bindings["heading2"].restart_on = None

    headings = [
        SimpleNamespace(level="heading2", text="第六节　研究内容及意义", para_index=0),
    ]

    entries = _build_toc_entries(
        doc,
        headings,
        cfg.heading_numbering.levels,
        cfg.heading_numbering_v2.level_bindings,
    )

    assert len(entries) == 1
    assert entries[0]["title"] == "研究内容及意义"
    assert "".join(part["text"] for part in entries[0]["title_runs"]) == "研究内容及意义"

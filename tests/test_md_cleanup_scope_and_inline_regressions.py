from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.md_cleanup import MdCleanupRule
from src.scene.schema import SceneConfig
from src.markdown.word_render import apply_num_pr, register_list_numbering


class _DummyDocTree:
    def __init__(self, paragraph_count: int):
        self.paragraph_count = paragraph_count

    def get_section(self, section_type: str):
        if section_type != "body":
            return None
        return SimpleNamespace(start_index=0, end_index=max(0, self.paragraph_count - 1))

    def build(self, doc: Document, body_start_index=None):
        self.paragraph_count = len(doc.paragraphs)


def _apply_md_cleanup(
    doc: Document,
    cfg: SceneConfig,
    *,
    target_indices: set[int] | None = None,
) -> None:
    context = {"doc_tree": _DummyDocTree(len(doc.paragraphs))}
    if target_indices is not None:
        context["target_paragraph_indices"] = set(target_indices)
    MdCleanupRule().apply(doc, cfg, ChangeTracker(), context)


def _has_num_pr(para) -> bool:
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        return False
    return ppr.find(qn("w:numPr")) is not None


def test_md_cleanup_range_scope_splits_manual_break_markdown_lists_without_dropping_lines():
    doc = Document()
    doc.add_paragraph("1. first\n2. second")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg, target_indices={0})

    assert [p.text for p in doc.paragraphs] == ["first", "second"]
    assert all(_has_num_pr(p) for p in doc.paragraphs)


def test_md_cleanup_orphan_native_list_fix_respects_target_scope():
    doc = Document()
    doc.add_paragraph("in scope")
    orphan = doc.add_paragraph("")
    doc.add_paragraph("outside text")
    num_id = register_list_numbering(doc, "bullet", marker_separator="tab")
    apply_num_pr(orphan, num_id, 0)

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg, target_indices={0})

    assert [p.text for p in doc.paragraphs] == ["in scope", "", "outside text"]
    assert _has_num_pr(doc.paragraphs[1]) is True
    assert _has_num_pr(doc.paragraphs[2]) is False


def test_md_cleanup_preserves_windows_paths_while_formatting_inline_markdown():
    doc = Document()
    doc.add_paragraph(r"Use **bold** around C:\Temp")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    para = doc.paragraphs[0]
    assert para.text == r"Use bold around C:\Temp"
    assert any((run.text == "bold") and bool(run.bold) for run in para.runs)


def test_md_cleanup_keeps_plain_blank_paragraphs_outside_markdown_paste_context():
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("")
    doc.add_paragraph("Use **bold**")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert len(doc.paragraphs) == 3
    assert [p.text for p in doc.paragraphs] == ["alpha", "", "Use bold"]
    assert any(bool(run.bold) for run in doc.paragraphs[2].runs)


def test_md_cleanup_converts_markdown_footnotes_after_formula_protection_narrowing():
    doc = Document()
    doc.add_paragraph("Body[^1]")
    doc.add_paragraph("[^1]: footnote content")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "Body"
    assert "footnoteReference" in doc.paragraphs[0]._p.xml


def test_md_cleanup_protects_plain_text_formula_explanation_with_markdown_sensitive_markers():
    original = "其中，*代表CQDs，E(*O2/*H+)、E(*)、E(O2/H+)分别表示CQDs表面吸附O2或H+的能量。"

    doc = Document()
    doc.add_paragraph(original)

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    cfg.md_cleanup.formula_copy_noise_cleanup = False
    cfg.md_cleanup.suppress_formula_fake_lists = False
    _apply_md_cleanup(doc, cfg)

    para = doc.paragraphs[0]
    assert para.text == original
    assert "*" in para.text
    assert all(not bool(run.italic) for run in para.runs if run.text)

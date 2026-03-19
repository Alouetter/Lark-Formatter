from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.formula_convert import FormulaConvertRule
from src.engine.rules.formula_to_table import FormulaToTableRule
from src.formula_core.parse import parse_document_formulas
from src.scene.schema import SceneConfig


def test_parse_document_formulas_groups_multiline_display_block():
    doc = Document()
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\mathcal{Z}(\beta,\mu)")
    doc.add_paragraph("=")
    doc.add_paragraph(r"\sum_{N=0}^{\infty}")
    doc.add_paragraph(r"\]")

    result = parse_document_formulas(doc)

    assert len(result.occurrences) == 1
    occ = result.occurrences[0]
    assert occ.source_type == "latex"
    assert occ.is_block is True
    assert occ.is_formula_only is True
    assert occ.source_paragraph_indices == [0, 1, 2, 3, 4]
    assert r"\mathcal{Z}(\beta,\mu)" in occ.source_text
    assert r"\sum_{N=0}^{\infty}" in occ.source_text
    assert r"\[" not in occ.source_text
    assert r"\]" not in occ.source_text


def test_parse_document_formulas_ignores_latex_document_wrapper_lines():
    doc = Document()
    doc.add_paragraph(r"\documentclass{article}")
    doc.add_paragraph(r"\usepackage{amsmath}")
    doc.add_paragraph(r"\begin{document}")
    doc.add_paragraph(r"\end{document}")

    result = parse_document_formulas(doc)

    assert result.occurrences == []


def test_formula_convert_collapses_multiline_display_block_before_table_conversion():
    doc = Document()
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\mathcal{Z}(\beta,\mu)")
    doc.add_paragraph("=")
    doc.add_paragraph(r"\sum_{N=0}^{\infty}")
    doc.add_paragraph(r"\]")

    cfg = SceneConfig()
    cfg.formula_convert.enabled = True
    cfg.formula_to_table.enabled = True

    tracker = ChangeTracker()
    context = {}
    FormulaConvertRule().apply(doc, cfg, tracker, context)

    assert len(doc.paragraphs) == 1
    assert "oMath" in doc.paragraphs[0]._p.xml

    FormulaToTableRule().apply(doc, cfg, tracker, context)

    assert len(doc.tables) == 1

def test_formula_convert_removes_latex_document_wrapper_lines_after_grouped_conversion():
    doc = Document()
    doc.add_paragraph(r"\documentclass{article}")
    doc.add_paragraph(r"\usepackage{amsmath}")
    doc.add_paragraph(r"\begin{document}")
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\mathcal{Z}(\beta,\mu)")
    doc.add_paragraph("=")
    doc.add_paragraph(r"\sum_{N=0}^{\infty}")
    doc.add_paragraph(r"\]")
    doc.add_paragraph(r"\end{document}")

    cfg = SceneConfig()
    cfg.formula_convert.enabled = True
    cfg.formula_to_table.enabled = True

    tracker = ChangeTracker()
    context = {}
    FormulaConvertRule().apply(doc, cfg, tracker, context)

    assert len(doc.paragraphs) == 1
    assert 'oMath' in doc.paragraphs[0]._p.xml
    assert all(not p.text.strip() for p in doc.paragraphs if 'oMath' in p._p.xml)

    FormulaToTableRule().apply(doc, cfg, tracker, context)

    assert len(doc.paragraphs) == 0
    assert len(doc.tables) == 1


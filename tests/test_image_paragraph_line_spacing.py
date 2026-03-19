import base64
from pathlib import Path

from docx import Document

from src.engine.pipeline import Pipeline
from src.engine.rules.style_manager import _apply_style_config
from src.markdown.word_render import insert_image
from src.scene.manager import load_scene_from_data
from src.scene.schema import StyleConfig

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+a9+gAAAAASUVORK5CYII="
)


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _spacing_attrs(spacing_el) -> dict[str, str]:
    attrs: dict[str, str] = {}
    if spacing_el is None:
        return attrs
    for name in ("before", "after", "line", "lineRule", "beforeAutospacing", "afterAutospacing"):
        value = spacing_el.get(_w(name))
        if value is not None:
            attrs[name] = value
    return attrs


def _write_tiny_png(tmp_path: Path) -> Path:
    img_path = tmp_path / "tiny.png"
    img_path.write_bytes(_TINY_PNG)
    return img_path


def _exact_normal_style() -> StyleConfig:
    return StyleConfig(
        font_cn="SimSun",
        font_en="Times New Roman",
        size_pt=12.0,
        alignment="justify",
        line_spacing_type="exact",
        line_spacing_pt=20.0,
        space_before_pt=0.0,
        space_after_pt=0.0,
    )


def test_insert_image_forces_safe_single_line_spacing_on_picture_paragraph(tmp_path):
    img_path = _write_tiny_png(tmp_path)

    doc = Document()
    _apply_style_config(doc.styles["Normal"], _exact_normal_style())

    para = doc.add_paragraph()
    assert insert_image(para, "tiny", str(img_path)) is True

    out_path = tmp_path / "insert_image_spacing.docx"
    doc.save(str(out_path))
    out_doc = Document(str(out_path))

    para = out_doc.paragraphs[0]
    ppr = para._element.find(_w("pPr"))
    spacing = ppr.find(_w("spacing")) if ppr is not None else None
    attrs = _spacing_attrs(spacing)

    normal_ppr = out_doc.styles["Normal"].element.find(_w("pPr"))
    normal_spacing = normal_ppr.find(_w("spacing")) if normal_ppr is not None else None
    normal_attrs = _spacing_attrs(normal_spacing)

    assert attrs["line"] == "240"
    assert attrs["lineRule"] == "auto"
    assert normal_attrs["line"] == "400"
    assert normal_attrs["lineRule"] == "exact"


def test_pipeline_section_format_fixes_existing_picture_paragraph_line_spacing(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")
    img_path = _write_tiny_png(tmp_path)

    doc_path = tmp_path / "existing_picture.docx"
    doc = Document()
    _apply_style_config(doc.styles["Normal"], _exact_normal_style())
    para = doc.add_paragraph()
    para.add_run().add_picture(str(img_path))
    doc.save(str(doc_path))

    cfg = load_scene_from_data({})
    cfg.pipeline = ["style_manager", "section_format"]
    cfg.pipeline_critical_rules = list(cfg.pipeline)
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False
    cfg.styles["normal"] = _exact_normal_style()

    result = Pipeline(cfg).run(str(doc_path))
    assert result.success is True

    out_doc = Document(result.output_paths["final"])
    para = out_doc.paragraphs[0]
    ppr = para._element.find(_w("pPr"))
    spacing = ppr.find(_w("spacing")) if ppr is not None else None
    attrs = _spacing_attrs(spacing)

    normal_ppr = out_doc.styles["Normal"].element.find(_w("pPr"))
    normal_spacing = normal_ppr.find(_w("spacing")) if normal_ppr is not None else None
    normal_attrs = _spacing_attrs(normal_spacing)

    assert attrs["line"] == "240"
    assert attrs["lineRule"] == "auto"
    assert normal_attrs["line"] == "400"
    assert normal_attrs["lineRule"] == "exact"

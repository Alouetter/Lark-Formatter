from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from src.engine.pipeline import Pipeline
from src.engine.rules.section_format import SectionFormatRule
from src.engine.rules.style_manager import _apply_style_config
from src.scene.manager import load_scene_from_data
from src.scene.schema import StyleConfig

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _spacing_attrs(spacing_el) -> dict[str, str]:
    attrs: dict[str, str] = {}
    if spacing_el is None:
        return attrs
    for name in ("before", "after", "line", "lineRule", "beforeAutospacing", "afterAutospacing"):
        value = spacing_el.get(_w(name))
        if value is not None:
            attrs[name] = value
    return attrs


def _seed_auto_spacing(container_element) -> None:
    ppr = container_element.find(_w("pPr"))
    if ppr is None:
        from lxml import etree

        ppr = etree.SubElement(container_element, _w("pPr"))
    spacing = ppr.find(_w("spacing"))
    if spacing is None:
        from lxml import etree

        spacing = etree.SubElement(ppr, _w("spacing"))
    spacing.set(_w("before"), "240")
    spacing.set(_w("after"), "120")
    spacing.set(_w("beforeAutospacing"), "1")
    spacing.set(_w("afterAutospacing"), "1")
    if ppr.find(_w("contextualSpacing")) is None:
        from lxml import etree

        etree.SubElement(ppr, _w("contextualSpacing"))


def test_heading_style_spacing_explicitly_overrides_inherited_normal_spacing(tmp_path):
    doc = Document()
    _apply_style_config(
        doc.styles["Normal"],
        StyleConfig(
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=12.0,
            alignment="justify",
            first_line_indent_chars=2.0,
            first_line_indent_unit="chars",
            line_spacing_type="exact",
            line_spacing_pt=20.0,
            space_before_pt=12.0,
            space_after_pt=8.0,
        ),
    )
    heading_style = doc.styles["Heading 1"]
    _seed_auto_spacing(heading_style.element)
    _apply_style_config(
        heading_style,
        StyleConfig(
            font_cn="黑体",
            font_en="Arial",
            size_pt=16.0,
            bold=True,
            alignment="left",
            first_line_indent_chars=0.0,
            first_line_indent_unit="chars",
            line_spacing_type="multiple",
            line_spacing_pt=1.5,
            space_before_pt=0.0,
            space_after_pt=0.0,
        ),
    )

    doc_path = tmp_path / "heading_style_spacing.docx"
    doc.save(str(doc_path))
    out_doc = Document(str(doc_path))

    style = out_doc.styles["Heading 1"]
    ppr = style.element.find(_w("pPr"))
    spacing = ppr.find(_w("spacing")) if ppr is not None else None
    attrs = _spacing_attrs(spacing)

    assert attrs["before"] == "0"
    assert attrs["after"] == "0"
    assert attrs["line"] == "360"
    assert attrs["lineRule"] == "auto"
    assert attrs.get("beforeAutospacing") is None
    assert attrs.get("afterAutospacing") is None
    assert ppr.find(_w("contextualSpacing")) is None


def test_unnumbered_heading_paragraph_spacing_overrides_inherited_normal_spacing(tmp_path):
    doc = Document()
    _apply_style_config(
        doc.styles["Normal"],
        StyleConfig(
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=12.0,
            alignment="justify",
            first_line_indent_chars=2.0,
            first_line_indent_unit="chars",
            line_spacing_type="exact",
            line_spacing_pt=20.0,
            space_before_pt=12.0,
            space_after_pt=8.0,
        ),
    )

    para = doc.add_paragraph("摘要")
    _seed_auto_spacing(para._element)
    rule = SectionFormatRule()
    cfg = load_scene_from_data({})
    rule._active_config = cfg
    title_sc = cfg.styles["heading1"]
    title_sc.space_before_pt = 0.0
    title_sc.space_after_pt = 0.0
    title_sc.line_spacing_type = "multiple"
    title_sc.line_spacing_pt = 1.5

    rule._make_para_unnumbered(doc, para)
    rule._format_paragraph(para, title_sc, force_plain=True, alignment_override="center")

    doc_path = tmp_path / "unnumbered_heading_spacing.docx"
    doc.save(str(doc_path))
    out_doc = Document(str(doc_path))

    para = out_doc.paragraphs[0]
    ppr = para._element.find(_w("pPr"))
    spacing = ppr.find(_w("spacing")) if ppr is not None else None
    attrs = _spacing_attrs(spacing)

    assert attrs["before"] == "0"
    assert attrs["after"] == "0"
    assert attrs["line"] == "360"
    assert attrs["lineRule"] == "auto"
    assert attrs.get("beforeAutospacing") is None
    assert attrs.get("afterAutospacing") is None
    assert ppr.find(_w("contextualSpacing")) is None


def test_section_format_paragraph_spacing_overrides_existing_direct_spacing(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "body_spacing_override.docx"
    doc = Document()
    para = doc.add_paragraph("正文第一段")
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(18)
    para.paragraph_format.line_spacing = Pt(30)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    _seed_auto_spacing(para._element)
    doc.save(str(doc_path))

    cfg = load_scene_from_data({})
    cfg.pipeline = ["style_manager", "section_format"]
    cfg.pipeline_critical_rules = list(cfg.pipeline)
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False
    cfg.styles["normal"].line_spacing_type = "exact"
    cfg.styles["normal"].line_spacing_pt = 20.0
    cfg.styles["normal"].space_before_pt = 0.0
    cfg.styles["normal"].space_after_pt = 0.0

    result = Pipeline(cfg).run(str(doc_path))
    assert result.success is True

    out_doc = Document(result.output_paths["final"])
    para = out_doc.paragraphs[0]
    ppr = para._element.find(_w("pPr"))
    spacing = ppr.find(_w("spacing")) if ppr is not None else None
    attrs = _spacing_attrs(spacing)

    assert attrs["before"] == "0"
    assert attrs["after"] == "0"
    assert attrs["line"] == "400"
    assert attrs["lineRule"] == "exact"
    assert attrs.get("beforeAutospacing") is None
    assert attrs.get("afterAutospacing") is None
    assert ppr.find(_w("contextualSpacing")) is None

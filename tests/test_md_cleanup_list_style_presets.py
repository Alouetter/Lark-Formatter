from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn

import src.markdown.word_render as word_render
from src.engine.change_tracker import ChangeTracker
from src.engine.rules.md_cleanup import MdCleanupRule
from src.markdown.word_render import apply_num_pr, register_list_numbering
from src.scene.schema import SceneConfig


class _DummyDocTree:
    def __init__(self, paragraph_count: int):
        self.paragraph_count = paragraph_count

    def get_section(self, section_type: str):
        if section_type != "body":
            return None
        return SimpleNamespace(start_index=0, end_index=max(0, self.paragraph_count - 1))

    def build(self, doc: Document, body_start_index=None):
        self.paragraph_count = len(doc.paragraphs)


def _apply_md_cleanup(doc: Document, cfg: SceneConfig):
    context = {"doc_tree": _DummyDocTree(len(doc.paragraphs))}
    MdCleanupRule().apply(doc, cfg, ChangeTracker(), context)


def _get_para_num_id(para) -> str | None:
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        return None
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    return num_id_el.get(qn("w:val"))


def _get_num_level_style(doc: Document, para, ilvl: int = 0) -> tuple[str | None, str | None, str | None]:
    num_id = _get_para_num_id(para)
    if not num_id:
        return None, None, None
    numbering = doc.part.numbering_part.element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    num_el = numbering.find(f".//w:num[@w:numId='{num_id}']", ns)
    if num_el is None:
        return None, None, None
    abs_ref = num_el.find("w:abstractNumId", ns)
    if abs_ref is None:
        return None, None, None
    abs_id = abs_ref.get(qn("w:val"))
    if not abs_id:
        return None, None, None
    lvl = numbering.find(
        f".//w:abstractNum[@w:abstractNumId='{abs_id}']/w:lvl[@w:ilvl='{ilvl}']",
        ns,
    )
    if lvl is None:
        return None, None, None
    fmt = lvl.find("w:numFmt", ns)
    txt = lvl.find("w:lvlText", ns)
    suff = lvl.find("w:suff", ns)
    return (
        fmt.get(qn("w:val")) if fmt is not None else None,
        txt.get(qn("w:val")) if txt is not None else None,
        suff.get(qn("w:val")) if suff is not None else None,
    )


def test_md_cleanup_ordered_style_decimal_cn_dun_applies_to_markdown_lists():
    doc = Document()
    doc.add_paragraph("1. 项目一")
    doc.add_paragraph("2. 项目二")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    cfg.md_cleanup.ordered_list_style = "decimal_cn_dun"
    _apply_md_cleanup(doc, cfg)

    fmt, lvl_text, suff = _get_num_level_style(doc, doc.paragraphs[0])
    assert fmt == "decimal"
    assert lvl_text == "%1、"
    assert suff == "tab"


def test_md_cleanup_ordered_style_decimal_full_paren_applies_to_markdown_lists():
    doc = Document()
    doc.add_paragraph("1. 项目一")
    doc.add_paragraph("2. 项目二")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    cfg.md_cleanup.ordered_list_style = "decimal_full_paren"
    _apply_md_cleanup(doc, cfg)

    fmt, lvl_text, _ = _get_num_level_style(doc, doc.paragraphs[0])
    assert fmt == "decimal"
    assert lvl_text == "（%1）"


def test_md_cleanup_unordered_style_square_applies_to_markdown_lists():
    doc = Document()
    doc.add_paragraph("- item1")
    doc.add_paragraph("- item2")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    cfg.md_cleanup.unordered_list_style = "bullet_square"
    _apply_md_cleanup(doc, cfg)

    fmt, lvl_text, _ = _get_num_level_style(doc, doc.paragraphs[0])
    assert fmt == "bullet"
    assert lvl_text == "■"


def test_register_list_numbering_ignores_stale_cache_and_preserves_bullet_style(monkeypatch):
    doc = Document()
    numbering = doc.part.numbering_part.element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    existing_num_ids = {
        int(num_el.get(qn("w:numId")))
        for num_el in numbering.findall(".//w:num", ns)
        if (num_el.get(qn("w:numId")) or "").isdigit()
    }

    monkeypatch.setattr(
        word_render,
        "_numbering_id_cache",
        {
            (id(numbering), "abstractNum", "abstractNumId"): 1,
            (id(numbering), "num", "numId"): 1,
        },
    )

    para = doc.add_paragraph("item1")
    num_id = register_list_numbering(doc, "bullet", unordered_style="bullet_square")
    apply_num_pr(para, num_id, 0)

    assert num_id > max(existing_num_ids)
    fmt, lvl_text, _ = _get_num_level_style(doc, para)
    assert fmt == "bullet"
    assert lvl_text == "■"


def test_md_cleanup_updates_existing_native_ordered_list_style():
    doc = Document()
    para1 = doc.add_paragraph("item1")
    para2 = doc.add_paragraph("item2")
    num_id = register_list_numbering(doc, "decimal")
    apply_num_pr(para1, num_id, 0)
    apply_num_pr(para2, num_id, 0)

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    cfg.md_cleanup.ordered_list_style = "decimal_cn_dun"
    cfg.md_cleanup.list_marker_separator = "half_space"
    _apply_md_cleanup(doc, cfg)

    fmt, lvl_text, suff = _get_num_level_style(doc, doc.paragraphs[0])
    assert fmt == "decimal"
    assert lvl_text == "%1、"
    assert suff == "space"

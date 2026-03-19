from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.engine.pipeline import Pipeline
from src.scene.manager import load_scene

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def test_default_preset_heading_conversion_clears_body_indent_override(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    cfg = load_scene(Path("src/scene/presets/default_format.json"))
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False
    cfg.pipeline = ["style_manager", "heading_detect", "heading_numbering"]
    cfg.pipeline_critical_rules = list(cfg.pipeline)
    cfg.styles["normal"].first_line_indent_chars = 2.0
    cfg.styles["normal"].first_line_indent_unit = "chars"

    doc_path = tmp_path / "heading_with_body_indent.docx"
    doc = Document()
    heading_para = doc.add_paragraph("1 Intro")
    heading_para.paragraph_format.first_line_indent = Pt(24)
    doc.add_paragraph("This is body text.")
    doc.save(str(doc_path))

    result = Pipeline(cfg).run(str(doc_path))
    assert result.success is True
    assert "final" in result.output_paths

    out_doc = Document(result.output_paths["final"])
    para = out_doc.paragraphs[0]
    assert para.style.name == "Heading 1"

    ppr = para._element.find(_w("pPr"))
    ind = ppr.find(_w("ind")) if ppr is not None else None
    assert ind is None or ind.get(_w("firstLine")) in {None, "0"}
    assert ind is None or ind.get(_w("firstLineChars")) in {None, "0"}

    heading_style = out_doc.styles["Heading 1"]
    style_ppr = heading_style.element.find(_w("pPr"))
    style_ind = style_ppr.find(_w("ind")) if style_ppr is not None else None
    assert style_ind is not None
    assert style_ind.get(_w("firstLine")) == "0"

    normal_style = out_doc.styles["Normal"]
    normal_ppr = normal_style.element.find(_w("pPr"))
    normal_ind = normal_ppr.find(_w("ind")) if normal_ppr is not None else None
    assert normal_ind is not None
    assert normal_ind.get(_w("firstLineChars")) == "200"

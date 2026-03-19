from __future__ import annotations

from types import SimpleNamespace

from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.section_format import SectionFormatRule, _run_vert_align
from src.scene.manager import load_default_scene


class _DummyDocTree:
    def __init__(self, sections):
        self.sections = sections

    def get_section(self, section_type: str):
        for section in self.sections:
            if section.section_type == section_type:
                return section
        return None

    def get_section_for_paragraph(self, para_index: int) -> str:
        for section in self.sections:
            if section.start_index <= para_index <= section.end_index:
                return section.section_type
        return "unknown"


def _has_super_or_sub(runs) -> bool:
    return any(style in {"superscript", "subscript"} for _, style in runs)


def test_caption_scope_restores_chem_typography_for_caption_paragraphs() -> None:
    doc = Document()
    caption = doc.add_paragraph("图1 H2O2")
    body_section = SimpleNamespace(section_type="body", paragraph_range=range(0, 1), start_index=0, end_index=0)

    cfg = load_default_scene()
    cfg.chem_typography.enabled = True
    cfg.chem_typography.scopes = {
        "references": False,
        "body": False,
        "headings": False,
        "abstract_cn": False,
        "abstract_en": False,
        "captions": True,
        "tables": False,
    }

    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": _DummyDocTree([body_section]),
            "format_scope": cfg.format_scope,
            "caption_indices": {0},
            "target_paragraph_indices": {0},
        },
    )

    runs = [(run.text, _run_vert_align(run)) for run in caption.runs]
    assert _has_super_or_sub(runs) is True


def test_caption_scope_off_keeps_caption_plain_even_when_body_scope_off() -> None:
    doc = Document()
    caption = doc.add_paragraph("图1 H2O2")
    body_section = SimpleNamespace(section_type="body", paragraph_range=range(0, 1), start_index=0, end_index=0)

    cfg = load_default_scene()
    cfg.chem_typography.enabled = True
    cfg.chem_typography.scopes = {
        "references": False,
        "body": False,
        "headings": False,
        "abstract_cn": False,
        "abstract_en": False,
        "captions": False,
        "tables": False,
    }

    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": _DummyDocTree([body_section]),
            "format_scope": cfg.format_scope,
            "caption_indices": {0},
            "target_paragraph_indices": {0},
        },
    )

    runs = [(run.text, _run_vert_align(run)) for run in caption.runs]
    assert _has_super_or_sub(runs) is False


def test_table_scope_restores_chem_typography_inside_table_cells() -> None:
    doc = Document()
    doc.add_paragraph("正文前导")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "H2O2"
    body_section = SimpleNamespace(section_type="body", paragraph_range=range(0, 1), start_index=0, end_index=0)

    cfg = load_default_scene()
    cfg.chem_typography.enabled = True
    cfg.chem_typography.scopes = {
        "references": False,
        "body": False,
        "headings": False,
        "abstract_cn": False,
        "abstract_en": False,
        "captions": False,
        "tables": True,
    }

    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": _DummyDocTree([body_section]),
            "format_scope": cfg.format_scope,
            "target_paragraph_indices": {0},
        },
    )

    para = table.cell(0, 0).paragraphs[0]
    runs = [(run.text, _run_vert_align(run)) for run in para.runs]
    assert _has_super_or_sub(runs) is True


def test_table_scope_off_keeps_table_cell_plain() -> None:
    doc = Document()
    doc.add_paragraph("正文前导")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "H2O2"
    body_section = SimpleNamespace(section_type="body", paragraph_range=range(0, 1), start_index=0, end_index=0)

    cfg = load_default_scene()
    cfg.chem_typography.enabled = True
    cfg.chem_typography.scopes = {
        "references": False,
        "body": False,
        "headings": False,
        "abstract_cn": False,
        "abstract_en": False,
        "captions": False,
        "tables": False,
    }

    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": _DummyDocTree([body_section]),
            "format_scope": cfg.format_scope,
            "target_paragraph_indices": {0},
        },
    )

    para = table.cell(0, 0).paragraphs[0]
    runs = [(run.text, _run_vert_align(run)) for run in para.runs]
    assert _has_super_or_sub(runs) is False

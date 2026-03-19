import json
import os
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PySide6.QtWidgets import QApplication

from src.engine.pipeline import Pipeline
from src.scene.manager import load_scene
from src.scene.manager import load_scene_from_data
from src.ui.main_window import FormatConfigDialog

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _app():
    return QApplication.instance() or QApplication([])


def _combo_select_by_data(combo, value) -> None:
    idx = combo.findData(value)
    assert idx >= 0
    combo.setCurrentIndex(idx)


def _style_row(dialog: FormatConfigDialog, style_key: str) -> int:
    return dialog._style_keys.index(style_key)


def _find_numbering_binding(numbering_el, *, num_id: str):
    num_el = None
    for candidate in numbering_el.findall(_w("num")):
        if str(candidate.get(_w("numId"), "")).strip() == str(num_id).strip():
            num_el = candidate
            break
    assert num_el is not None

    abs_ref = num_el.find(_w("abstractNumId"))
    assert abs_ref is not None
    abs_id = str(abs_ref.get(_w("val"), "")).strip()
    assert abs_id

    abs_num = None
    for candidate in numbering_el.findall(_w("abstractNum")):
        if str(candidate.get(_w("abstractNumId"), "")).strip() == abs_id:
            abs_num = candidate
            break
    assert abs_num is not None
    return abs_num


def _find_lvl(abs_num, ilvl: int):
    for lvl in abs_num.findall(_w("lvl")):
        if str(lvl.get(_w("ilvl"), "")).strip() == str(ilvl):
            return lvl
    raise AssertionError(f"missing numbering level {ilvl}")


def test_format_config_style_values_round_trip_to_json_and_final_docx(tmp_path, monkeypatch):
    _app()
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    cfg = load_scene_from_data({})
    dlg = FormatConfigDialog(cfg, scene_path=str(tmp_path / "seed.json"))

    row = _style_row(dlg, "normal")

    size_combo = dlg._style_cell_widget(row, dlg._style_table_col_index("size_pt"))
    size_combo.setCurrentText("14")

    alignment_combo = dlg._style_cell_widget(row, dlg._style_table_col_index("alignment"))
    _combo_select_by_data(alignment_combo, "center")

    line_spacing_widget = dlg._style_cell_widget(row, dlg._style_table_col_index("line_spacing_pt"))
    line_spacing_widget.setSpacingType("exact", 20)

    special_indent_widget = dlg._style_cell_widget(row, dlg._style_table_col_index("special_indent_value"))
    special_indent_widget.setMode("first_line")
    special_indent_widget.setUnit("pt")
    special_indent_widget._value_widget._edit.setText("18")

    left_indent_widget = dlg._style_cell_widget(row, dlg._style_table_col_index("left_indent_chars"))
    left_indent_widget.setUnit("pt")
    left_indent_widget._edit.setText("12")

    right_indent_widget = dlg._style_cell_widget(row, dlg._style_table_col_index("right_indent_chars"))
    right_indent_widget.setUnit("pt")
    right_indent_widget._edit.setText("6")

    dlg._output_checks["report_json"].setChecked(False)
    dlg._prompt_format_name = lambda *args, **kwargs: ("style_e2e", True)
    dlg._save_format()

    saved_path = Path(dlg._scene_path)
    assert saved_path.exists()

    raw = json.loads(saved_path.read_text(encoding="utf-8"))
    normal = raw["styles"]["normal"]
    assert normal["size_pt"] == 14.0
    assert normal["size_display"] == "14"
    assert normal["alignment"] == "center"
    assert normal["line_spacing_type"] == "exact"
    assert normal["line_spacing_pt"] == 20.0
    assert normal["special_indent_mode"] == "first_line"
    assert normal["special_indent_value"] == 18.0
    assert normal["special_indent_unit"] == "pt"
    assert normal["first_line_indent_chars"] == 18.0
    assert normal["first_line_indent_unit"] == "pt"
    assert normal["left_indent_chars"] == 12.0
    assert normal["left_indent_unit"] == "pt"
    assert normal["right_indent_chars"] == 6.0
    assert normal["right_indent_unit"] == "pt"
    assert raw["output"]["report_json"] is False

    reloaded = load_scene(saved_path)
    assert reloaded.styles["normal"].size_pt == 14.0
    assert reloaded.styles["normal"].size_display == "14"
    assert reloaded.styles["normal"].alignment == "center"
    assert reloaded.styles["normal"].line_spacing_type == "exact"
    assert reloaded.styles["normal"].line_spacing_pt == 20.0
    assert reloaded.styles["normal"].special_indent_mode == "first_line"
    assert reloaded.styles["normal"].special_indent_value == 18.0
    assert reloaded.styles["normal"].special_indent_unit == "pt"
    assert reloaded.styles["normal"].first_line_indent_chars == 18.0
    assert reloaded.styles["normal"].first_line_indent_unit == "pt"
    assert reloaded.styles["normal"].left_indent_chars == 12.0
    assert reloaded.styles["normal"].left_indent_unit == "pt"
    assert reloaded.styles["normal"].right_indent_chars == 6.0
    assert reloaded.styles["normal"].right_indent_unit == "pt"
    assert reloaded.output.report_json is False

    doc_path = tmp_path / "style_input.docx"
    doc = Document()
    doc.add_paragraph("Style pipeline sample.")
    doc.save(str(doc_path))

    reloaded.output.compare_docx = False
    reloaded.output.report_json = False
    reloaded.output.report_markdown = False

    result = Pipeline(reloaded).run(str(doc_path))
    assert result.success is True
    assert "final" in result.output_paths

    out_doc = Document(result.output_paths["final"])
    normal_style = out_doc.styles["Normal"]
    assert round(normal_style.font.size.pt, 2) == 14.0

    ppr = normal_style.element.find(_w("pPr"))
    assert ppr is not None
    jc = ppr.find(_w("jc"))
    ind = ppr.find(_w("ind"))
    spacing = ppr.find(_w("spacing"))
    assert jc is not None and jc.get(_w("val")) == "center"
    assert ind is not None and ind.get(_w("firstLine")) == "360"
    assert ind.get(_w("left")) == "240"
    assert ind.get(_w("right")) == "120"
    assert spacing is not None and spacing.get(_w("lineRule")) == "exact"
    assert spacing.get(_w("line")) == "400"


def test_format_config_heading_numbering_round_trips_to_json_and_final_docx(tmp_path, monkeypatch):
    _app()
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    cfg = load_scene_from_data({})
    cfg._heading_numbering_v2_source = "payload"
    dlg = FormatConfigDialog(cfg, scene_path=str(tmp_path / "seed.json"))

    _combo_select_by_data(dlg._num_mode_combo, "custom")

    dlg._num_cell_widget(0, "display_shell").setCurrentText("[{}]")
    _combo_select_by_data(dlg._num_cell_widget(0, "display_core_style"), "arabic")
    dlg._num_cell_widget(0, "title_separator").setRawText("_")
    dlg._num_cell_widget(0, "start_at").setValue(3)

    dlg._num_cell_widget(2, "start_at").setValue(7)
    _combo_select_by_data(dlg._num_cell_widget(2, "restart_on"), "heading1")
    dlg._num_cell_widget(2, "include_in_toc").setChecked(False)

    dlg._prompt_format_name = lambda *args, **kwargs: ("heading_e2e", True)
    dlg._save_format()

    saved_path = Path(dlg._scene_path)
    assert saved_path.exists()

    raw = json.loads(saved_path.read_text(encoding="utf-8"))
    v2 = raw["heading_numbering_v2"]["level_bindings"]
    legacy = raw["heading_numbering"]["levels"]

    h1_binding = v2["heading1"]
    h3_binding = v2["heading3"]
    assert h1_binding["start_at"] == 3
    assert h3_binding["start_at"] == 7
    assert h3_binding["restart_on"] == "heading1"
    assert h3_binding["include_in_toc"] is False

    h1_legacy = legacy["heading1"]
    assert h1_legacy["format"] == "arabic"
    assert h1_legacy["template"] == "[{current}]"
    assert h1_legacy["separator"] == "_"

    reloaded = load_scene(saved_path)
    assert reloaded.heading_numbering_v2.level_bindings["heading1"].start_at == 3
    assert reloaded.heading_numbering_v2.level_bindings["heading3"].start_at == 7
    assert reloaded.heading_numbering_v2.level_bindings["heading3"].restart_on == "heading1"
    assert reloaded.heading_numbering_v2.level_bindings["heading3"].include_in_toc is False
    assert reloaded.heading_numbering.levels["heading1"].template == "[{current}]"
    assert reloaded.heading_numbering.levels["heading1"].separator == "_"

    doc_path = tmp_path / "heading_input.docx"
    doc = Document()
    doc.add_paragraph("Heading one sample.", style="Heading 1")
    doc.add_paragraph("Heading two sample.", style="Heading 2")
    doc.add_paragraph("Heading three sample.", style="Heading 3")
    doc.save(str(doc_path))

    reloaded.output.compare_docx = False
    reloaded.output.report_json = False
    reloaded.output.report_markdown = False

    result = Pipeline(reloaded).run(str(doc_path))
    assert result.success is True
    assert "final" in result.output_paths

    out_doc = Document(result.output_paths["final"])
    heading1_style = out_doc.styles["Heading 1"]
    ppr = heading1_style.element.find(_w("pPr"))
    assert ppr is not None
    num_pr = ppr.find(_w("numPr"))
    assert num_pr is not None
    num_id = num_pr.find(_w("numId")).get(_w("val"))
    assert num_id is not None

    abs_num = _find_numbering_binding(out_doc.part.numbering_part.element, num_id=num_id)
    lvl0 = _find_lvl(abs_num, 0)
    lvl2 = _find_lvl(abs_num, 2)

    assert lvl0.find(_w("start")).get(_w("val")) == "3"
    assert lvl0.find(_w("lvlText")).get(_w("val")) == "[%1]_"
    assert lvl2.find(_w("start")).get(_w("val")) == "7"

    lvl2_restart = lvl2.find(_w("lvlRestart"))
    assert lvl2_restart is not None
    assert lvl2_restart.get(_w("val")) == "1"


def test_format_config_heading_numbering_preset_mode_persists_after_reopen(tmp_path):
    _app()

    cfg = load_scene_from_data({})
    cfg._heading_numbering_v2_source = "payload"
    dlg = FormatConfigDialog(cfg, scene_path=str(tmp_path / "seed.json"))

    _combo_select_by_data(dlg._num_mode_combo, "preset_1")
    dlg._prompt_format_name = lambda *args, **kwargs: ("preset_roundtrip", True)
    dlg._save_format()

    saved_path = Path(dlg._scene_path)
    reloaded = load_scene(saved_path)
    reopened = FormatConfigDialog(reloaded, scene_path=str(saved_path))

    assert reopened._num_mode_combo.currentData() == "preset_1"
    assert reopened._num_detail_container.isEnabled() is False
    assert reloaded.heading_numbering_v2.level_bindings["heading1"].display_shell == "chapter_cn"
    assert reloaded.heading_numbering_v2.level_bindings["heading2"].display_shell == "section_cn"


def test_reference_hanging_indent_round_trips_to_json(tmp_path):
    _app()

    cfg = load_scene_from_data({})
    dlg = FormatConfigDialog(cfg, scene_path=str(tmp_path / "seed.json"))

    row = _style_row(dlg, "references_body")
    left_widget = dlg._style_cell_widget(row, dlg._style_table_col_index("left_indent_chars"))
    left_widget.setUnit("chars")
    left_widget._edit.setText("2")

    special_widget = dlg._style_cell_widget(row, dlg._style_table_col_index("special_indent_value"))
    special_widget.setMode("hanging")
    special_widget.setUnit("chars")
    special_widget._value_widget._edit.setText("2")

    dlg._prompt_format_name = lambda *args, **kwargs: ("references_hanging", True)
    dlg._save_format()

    saved_path = Path(dlg._scene_path)
    raw = json.loads(saved_path.read_text(encoding="utf-8"))
    ref_style = raw["styles"]["references_body"]

    assert ref_style["left_indent_chars"] == 2.0
    assert ref_style["left_indent_unit"] == "chars"
    assert ref_style["special_indent_mode"] == "hanging"
    assert ref_style["special_indent_value"] == 2.0
    assert ref_style["special_indent_unit"] == "chars"
    assert ref_style["hanging_indent_chars"] == 2.0
    assert ref_style["hanging_indent_unit"] == "chars"

    reloaded = load_scene(saved_path)
    assert reloaded.styles["references_body"].left_indent_chars == 2.0
    assert reloaded.styles["references_body"].left_indent_unit == "chars"
    assert reloaded.styles["references_body"].special_indent_mode == "hanging"
    assert reloaded.styles["references_body"].special_indent_value == 2.0
    assert reloaded.styles["references_body"].special_indent_unit == "chars"
    assert reloaded.styles["references_body"].hanging_indent_chars == 2.0
    assert reloaded.styles["references_body"].hanging_indent_unit == "chars"


def test_format_config_misc_values_round_trip_to_json_and_reload(tmp_path):
    _app()

    cfg = load_scene_from_data({})
    dlg = FormatConfigDialog(cfg, scene_path=str(tmp_path / "seed.json"))

    dlg._enforcement_checks["ban_tab"].setChecked(False)
    dlg._enforcement_checks["auto_fix"].setChecked(False)
    dlg._risk_guard_controls["enabled"].setChecked(False)
    dlg._risk_guard_controls["no_body_min_candidates"].setValue(5)
    dlg._risk_guard_controls["tiny_body_max_ratio"].setValue(0.125)
    dlg._risk_guard_controls["keep_after_first_chapter"].setChecked(False)

    dlg._cap_edits["figure_prefix"].setText("Fig")
    dlg._cap_edits["table_prefix"].setText("Tbl")
    dlg._cap_edits["placeholder"].setText("[PENDING]")
    _combo_select_by_data(dlg._cap_separator_mode_combo, "custom")
    dlg._cap_custom_separator_edit.setText("::")
    _combo_select_by_data(dlg._numbering_format_combo, "chapter-seq")
    _combo_select_by_data(dlg._equation_numbering_format_combo, "seq")
    dlg._cap_enabled.setChecked(False)

    dlg._formula_table_unify_font_check.setChecked(False)
    dlg._formula_table_unify_size_check.setChecked(False)
    dlg._formula_table_unify_spacing_check.setChecked(True)
    dlg._formula_table_font_combo.setCurrentText("XITS Math")
    dlg._formula_table_font_size_combo.setCurrentText("11.5pt")
    dlg._formula_table_line_spacing_spin.setValue(1.25)
    dlg._formula_table_space_before_spin.setValue(2.0)
    dlg._formula_table_space_after_spin.setValue(3.0)
    _combo_select_by_data(dlg._formula_table_block_alignment_combo, "right")
    _combo_select_by_data(dlg._formula_table_table_alignment_combo, "left")
    _combo_select_by_data(dlg._formula_table_cell_alignment_combo, "left")
    _combo_select_by_data(dlg._formula_table_number_alignment_combo, "center")
    dlg._formula_table_number_font_combo.setCurrentText("Arial")
    dlg._formula_table_number_size_combo.setCurrentText("9pt")
    dlg._formula_table_auto_shrink_check.setChecked(False)

    _combo_select_by_data(dlg._paper_size_combo, "Letter")
    dlg._page_spins["top_cm"].setValue(2.6)
    dlg._page_spins["bottom_cm"].setValue(2.4)
    dlg._page_spins["left_cm"].setValue(2.8)
    dlg._page_spins["right_cm"].setValue(2.2)
    dlg._page_spins["gutter_cm"].setValue(0.4)
    dlg._page_spins["header_distance_cm"].setValue(1.5)
    dlg._page_spins["footer_distance_cm"].setValue(1.2)

    dlg._output_checks["compare_docx"].setChecked(False)
    dlg._output_checks["report_json"].setChecked(False)
    dlg._output_checks["report_markdown"].setChecked(False)

    dlg._pipeline_checks["caption_format"].setChecked(False)
    dlg._pipeline_checks["validation"].setChecked(False)

    dlg._prompt_format_name = lambda *args, **kwargs: ("misc_roundtrip", True)
    dlg._save_format()

    saved_path = Path(dlg._scene_path)
    raw = json.loads(saved_path.read_text(encoding="utf-8"))

    assert raw["heading_numbering"]["enforcement"]["ban_tab"] is False
    assert raw["heading_numbering"]["enforcement"]["auto_fix"] is False
    assert raw["heading_numbering"]["risk_guard"]["enabled"] is False
    assert raw["heading_numbering"]["risk_guard"]["no_body_min_candidates"] == 5
    assert raw["heading_numbering"]["risk_guard"]["tiny_body_max_ratio"] == 0.125
    assert raw["heading_numbering"]["risk_guard"]["keep_after_first_chapter"] is False

    assert raw["caption"]["enabled"] is False
    assert raw["caption"]["figure_prefix"] == "Fig"
    assert raw["caption"]["table_prefix"] == "Tbl"
    assert raw["caption"]["placeholder"] == "[PENDING]"
    assert raw["caption"]["separator"] == "::"
    assert raw["caption"]["numbering_format"] == "chapter-seq"
    assert raw["equation_table_format"]["numbering_format"] == "seq"

    assert raw["formula_style"]["unify_font"] is False
    assert raw["formula_style"]["unify_size"] is False
    assert raw["formula_style"]["unify_spacing"] is True
    assert raw["formula_table"]["formula_font_name"] == "XITS Math"
    assert raw["formula_table"]["formula_font_size_pt"] == 11.5
    assert raw["formula_table"]["formula_font_size_display"] == "11.5pt"
    assert raw["formula_table"]["formula_line_spacing"] == 1.25
    assert raw["formula_table"]["formula_space_before_pt"] == 2.0
    assert raw["formula_table"]["formula_space_after_pt"] == 3.0
    assert raw["formula_table"]["block_alignment"] == "right"
    assert raw["formula_table"]["table_alignment"] == "left"
    assert raw["formula_table"]["formula_cell_alignment"] == "left"
    assert raw["formula_table"]["number_alignment"] == "center"
    assert raw["formula_table"]["number_font_name"] == "Arial"
    assert raw["formula_table"]["number_font_size_pt"] == 9.0
    assert raw["formula_table"]["number_font_size_display"] == "9pt"
    assert raw["formula_table"]["auto_shrink_number_column"] is False

    assert raw["page_setup"]["paper_size"] == "Letter"
    assert raw["page_setup"]["margin"]["top_cm"] == 2.6
    assert raw["page_setup"]["margin"]["bottom_cm"] == 2.4
    assert raw["page_setup"]["margin"]["left_cm"] == 2.8
    assert raw["page_setup"]["margin"]["right_cm"] == 2.2
    assert raw["page_setup"]["gutter_cm"] == 0.4
    assert raw["page_setup"]["header_distance_cm"] == 1.5
    assert raw["page_setup"]["footer_distance_cm"] == 1.2

    assert raw["output"]["compare_docx"] is False
    assert raw["output"]["report_json"] is False
    assert raw["output"]["report_markdown"] is False
    assert "caption_format" not in raw["pipeline"]
    assert "validation" not in raw["pipeline"]
    assert raw["pipeline_critical_rules"] == raw["pipeline"]

    reloaded = load_scene(saved_path)
    assert reloaded.heading_numbering.enforcement.ban_tab is False
    assert reloaded.heading_numbering.enforcement.auto_fix is False
    assert reloaded.heading_numbering.risk_guard.enabled is False
    assert reloaded.heading_numbering.risk_guard.no_body_min_candidates == 5
    assert reloaded.heading_numbering.risk_guard.tiny_body_max_ratio == 0.125
    assert reloaded.heading_numbering.risk_guard.keep_after_first_chapter is False

    assert reloaded.caption.enabled is False
    assert reloaded.caption.figure_prefix == "Fig"
    assert reloaded.caption.table_prefix == "Tbl"
    assert reloaded.caption.placeholder == "[PENDING]"
    assert reloaded.caption.separator == "::"
    assert reloaded.caption.numbering_format == "chapter-seq"
    assert reloaded.equation_table_format.numbering_format == "seq"

    assert reloaded.formula_style.unify_font is False
    assert reloaded.formula_style.unify_size is False
    assert reloaded.formula_style.unify_spacing is True
    assert reloaded.formula_table.formula_font_name == "XITS Math"
    assert reloaded.formula_table.formula_font_size_pt == 11.5
    assert reloaded.formula_table.formula_font_size_display == "11.5pt"
    assert reloaded.formula_table.formula_line_spacing == 1.25
    assert reloaded.formula_table.formula_space_before_pt == 2.0
    assert reloaded.formula_table.formula_space_after_pt == 3.0
    assert reloaded.formula_table.block_alignment == "right"
    assert reloaded.formula_table.table_alignment == "left"
    assert reloaded.formula_table.formula_cell_alignment == "left"
    assert reloaded.formula_table.number_alignment == "center"
    assert reloaded.formula_table.number_font_name == "Arial"
    assert reloaded.formula_table.number_font_size_pt == 9.0
    assert reloaded.formula_table.number_font_size_display == "9pt"
    assert reloaded.formula_table.auto_shrink_number_column is False

    assert reloaded.page_setup.paper_size == "Letter"
    assert reloaded.page_setup.margin.top_cm == 2.6
    assert reloaded.page_setup.margin.bottom_cm == 2.4
    assert reloaded.page_setup.margin.left_cm == 2.8
    assert reloaded.page_setup.margin.right_cm == 2.2
    assert reloaded.page_setup.gutter_cm == 0.4
    assert reloaded.page_setup.header_distance_cm == 1.5
    assert reloaded.page_setup.footer_distance_cm == 1.2

    assert reloaded.output.compare_docx is False
    assert reloaded.output.report_json is False
    assert reloaded.output.report_markdown is False
    assert "caption_format" not in reloaded.pipeline
    assert "validation" not in reloaded.pipeline
    assert reloaded.pipeline_critical_rules == reloaded.pipeline


def test_format_config_caption_and_page_values_apply_to_final_docx(tmp_path, monkeypatch):
    _app()
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    cfg = load_scene_from_data({})
    dlg = FormatConfigDialog(cfg, scene_path=str(tmp_path / "seed.json"))

    dlg._cap_edits["table_prefix"].setText("Tbl")
    dlg._cap_edits["placeholder"].setText("[PENDING]")
    _combo_select_by_data(dlg._cap_separator_mode_combo, "custom")
    dlg._cap_custom_separator_edit.setText("::")
    _combo_select_by_data(dlg._numbering_format_combo, "chapter-seq")

    _combo_select_by_data(dlg._paper_size_combo, "Letter")
    dlg._page_spins["top_cm"].setValue(2.6)
    dlg._page_spins["bottom_cm"].setValue(2.4)
    dlg._page_spins["left_cm"].setValue(2.8)
    dlg._page_spins["right_cm"].setValue(2.2)
    dlg._page_spins["gutter_cm"].setValue(0.4)
    dlg._page_spins["header_distance_cm"].setValue(1.5)
    dlg._page_spins["footer_distance_cm"].setValue(1.2)

    dlg._output_checks["compare_docx"].setChecked(False)
    dlg._output_checks["report_json"].setChecked(False)
    dlg._output_checks["report_markdown"].setChecked(False)
    dlg._prompt_format_name = lambda *args, **kwargs: ("caption_page_docx", True)
    dlg._save_format()

    reloaded = load_scene(Path(dlg._scene_path))
    reloaded.output.compare_docx = False
    reloaded.output.report_json = False
    reloaded.output.report_markdown = False

    doc_path = tmp_path / "caption_page_input.docx"
    doc = Document()
    doc.add_paragraph("Chapter One", style="Heading 1")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(str(doc_path))

    result = Pipeline(reloaded).run(str(doc_path))
    assert result.success is True
    assert "final" in result.output_paths

    out_doc = Document(result.output_paths["final"])
    section = out_doc.sections[0]
    assert round(section.page_width.cm, 2) == 21.59
    assert round(section.page_height.cm, 2) == 27.94
    assert round(section.top_margin.cm, 1) == 2.6
    assert round(section.bottom_margin.cm, 1) == 2.4
    assert round(section.left_margin.cm, 1) == 2.8
    assert round(section.right_margin.cm, 1) == 2.2
    assert round(section.gutter.cm, 1) == 0.4
    assert round(section.header_distance.cm, 1) == 1.5
    assert round(section.footer_distance.cm, 1) == 1.2

    caption_texts = [para.text.strip() for para in out_doc.paragraphs if para.text.strip().startswith("Tbl")]
    assert "Tbl1-1::[PENDING]" in caption_texts

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.engine.pipeline import Pipeline
from src.scene.schema import SceneConfig


def _build_whitespace_regression_doc(doc_path: Path) -> None:
    doc = Document()

    doc.add_paragraph("  Alpha\u00A0\u3000Beta\u200b  Gamma  ")
    doc.add_paragraph("中文,测试. 这是(中文)ＡＢＣ１２３")
    doc.add_paragraph("English，test（demo）ＡＢＣ１２３")
    doc.add_paragraph("参考文献 [ 1 , 2 ] 保留")
    doc.add_paragraph(r"路径 C:\Temp\Foo Bar\a.txt 不应乱改 ＡＢ")
    doc.add_paragraph("字段\t值")

    heading = doc.add_paragraph("  标题ＡＢＣ  ")
    heading.style = doc.styles["Heading 1"]

    doc.save(str(doc_path))


def _paragraph_has_tab_element(para) -> bool:
    return para._p.find(f".//{qn('w:tab')}") is not None


def test_whitespace_normalize_regression_on_generated_docx(tmp_path):
    doc_path = tmp_path / "文本空白清洗回归.docx"
    _build_whitespace_regression_doc(doc_path)

    cfg = SceneConfig()
    cfg.pipeline = ["whitespace_normalize"]
    cfg.whitespace_normalize.enabled = True
    cfg.output.final_docx = False
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False

    result = Pipeline(cfg).run(str(doc_path))

    assert result.success is True
    assert result.status == "success"
    assert result.doc is not None

    assert [p.text for p in result.doc.paragraphs] == [
        "Alpha Beta Gamma",
        "中文，测试。 这是（中文）ABC123",
        "English,test(demo)ABC123",
        "参考文献 [ 1 , 2 ] 保留",
        r"路径 C:\Temp\Foo Bar\a.txt 不应乱改 AB",
        "字段 值",
        "  标题ＡＢＣ  ",
    ]
    assert _paragraph_has_tab_element(result.doc.paragraphs[5]) is False

    changed_indices = sorted(
        {
            record.paragraph_index
            for record in result.tracker.get_by_rule("whitespace_normalize")
            if record.paragraph_index >= 0
        }
    )
    assert changed_indices == [0, 1, 2, 4, 5]

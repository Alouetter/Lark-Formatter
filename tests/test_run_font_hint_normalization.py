from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from src.engine.rules.section_format import (
    _apply_run_font,
    _normalize_explicit_run_font_hints,
    _normalize_related_story_part_font_hints,
)
from src.markdown.word_render import set_no_proof


def _ensure_rpr(run):
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run._element, qn("w:rPr"))
    return rpr


def _rfonts_attrs(run) -> dict[str, str]:
    rpr = _ensure_rpr(run)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return {}
    out = {}
    for key in ("ascii", "hAnsi", "eastAsia", "cs", "hint"):
        val = rfonts.get(qn(f"w:{key}"))
        if val is not None:
            out[key] = val
    return out


def test_apply_run_font_normalizes_hint_for_mixed_cjk_run() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("C=N-NH-C=O")

    rpr = _ensure_rpr(run)
    rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:hint"), "eastAsia")

    _apply_run_font(run, "宋体", "Times New Roman", 12)

    attrs = _rfonts_attrs(run)
    assert attrs["ascii"] == "Times New Roman"
    assert attrs["hAnsi"] == "Times New Roman"
    assert attrs["eastAsia"] == "宋体"
    assert attrs["cs"] == "Times New Roman"
    assert attrs["hint"] == "default"


def test_apply_run_font_formats_noproof_run_without_explicit_fonts() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("Urea-based carbon quantum dots")

    rpr = _ensure_rpr(run)
    etree.SubElement(rpr, qn("w:noProof"))

    _apply_run_font(run, "宋体", "Times New Roman", 10.5)

    attrs = _rfonts_attrs(run)
    assert attrs["ascii"] == "Times New Roman"
    assert attrs["eastAsia"] == "宋体"
    assert attrs["hint"] == "default"


def test_apply_run_font_preserves_explicit_code_font_on_noproof_run() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("code-span")
    run.font.name = "Consolas"
    set_no_proof(run)

    _apply_run_font(run, "宋体", "Times New Roman", 12)

    attrs = _rfonts_attrs(run)
    assert attrs["ascii"] == "Consolas"
    assert attrs["hAnsi"] == "Consolas"
    assert attrs.get("eastAsia") is None
    assert attrs.get("hint") is None


def test_normalize_explicit_run_font_hints_repairs_leftover_eastasia_hint() -> None:
    doc = Document()
    run = doc.add_paragraph().add_run("BD-CQDs")
    rpr = _ensure_rpr(run)
    rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:hint"), "eastAsia")

    assert _normalize_explicit_run_font_hints(doc) == 1
    assert _rfonts_attrs(run)["hint"] == "default"


def test_normalize_related_story_part_font_hints_repairs_comment_blob() -> None:
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:comment w:id="0">'
        '<w:p><w:r><w:rPr><w:rFonts '
        'w:ascii="Times New Roman" '
        'w:hAnsi="Times New Roman" '
        'w:eastAsia="瀹嬩綋" '
        'w:cs="Times New Roman" '
        'w:hint="eastAsia"/></w:rPr><w:t>Fe-N</w:t></w:r></w:p>'
        '</w:comment></w:comments>'
    ).encode("utf-8")

    class _DummyPart:
        def __init__(self, partname: str, blob: bytes):
            self.partname = partname
            self._blob = blob

        @property
        def blob(self) -> bytes:
            return self._blob

    class _DummyRel:
        def __init__(self, target_part):
            self.target_part = target_part

    class _DummyDoc:
        def __init__(self, part):
            self.part = type("DocPart", (), {"rels": {"rId1": _DummyRel(part)}})()

    part = _DummyPart("/word/comments.xml", comments_xml)
    doc = _DummyDoc(part)

    assert _normalize_related_story_part_font_hints(doc) == {"/word/comments.xml": 1}

    root = etree.fromstring(part.blob)
    rfonts = root.find(".//" + qn("w:rFonts"))
    assert rfonts is not None
    assert rfonts.get(qn("w:hint")) == "default"

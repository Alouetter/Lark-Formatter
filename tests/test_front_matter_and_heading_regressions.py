from __future__ import annotations

from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.doc_tree import DocSection, DocTree
from src.engine.rules.equation_table_format import EquationTableFormatRule
from src.engine.rules.header_footer import HeaderFooterRule
from src.engine.rules.heading_detect import HeadingDetectRule
from src.engine.rules.heading_numbering import HeadingNumberingRule
from src.engine.rules.section_format import SectionFormatRule
from src.engine.rules.toc_format import TocFormatRule
from src.scene.heading_model import get_non_numbered_heading_style_name
from src.scene.manager import load_default_scene

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class _DummyDocTree:
    def __init__(self, sections: list[DocSection]):
        self.sections = sections

    def get_section(self, section_type: str):
        for section in self.sections:
            if section.section_type == section_type:
                return section
        return None

    def get_section_for_paragraph(self, para_index: int) -> str:
        for section in self.sections:
            if section.start_index <= para_index <= section.end_index:
                return section.section_type
        return "unknown"


def _build_front_matter_excerpt_doc() -> Document:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 绪论\t1")
    doc.add_paragraph("Abstract\t5")
    doc.add_paragraph("摘要")
    doc.add_paragraph(
        "这是中文摘要部分的说明文字，用于验证目录之后的摘要内容不会被错误识别为正文。"
    )
    doc.add_paragraph(
        "（1）这是摘要中的第一点说明，内容较长且具有完整叙述，不应被识别为参考文献。"
    )
    doc.add_paragraph("关键词：智能编译；统一IR")
    doc.add_paragraph("Abstract")
    doc.add_paragraph(
        "This is the English abstract body used to verify that abstract pages are "
        "not treated as the first chapter body."
    )
    doc.add_paragraph(
        "(1) This is a long narrative abstract bullet item and it should not be "
        "treated as a reference entry."
    )
    doc.add_paragraph("Key Words: intelligent compilation; unified IR")
    return doc


def _header_texts(doc: Document) -> list[str]:
    texts: list[str] = []
    for sec in doc.sections:
        text = "".join(
            (para.text or "").strip()
            for para in sec.header.paragraphs
            if (para.text or "").strip()
        )
        texts.append(text)
    return texts


def _outline_level(para) -> str | None:
    ppr = para._element.find(f"{{{_W_NS}}}pPr")
    if ppr is None:
        return None
    outline = ppr.find(f"{{{_W_NS}}}outlineLvl")
    if outline is None:
        return None
    return outline.get(f"{{{_W_NS}}}val")


def _has_numpr(target) -> bool:
    element = getattr(target, "element", None)
    if element is None:
        element = getattr(target, "_element", None)
    if element is None:
        return False
    ppr = element.find(f"{{{_W_NS}}}pPr")
    if ppr is None:
        return False
    return ppr.find(f"{{{_W_NS}}}numPr") is not None


def test_doc_tree_keeps_abstract_excerpt_as_front_matter() -> None:
    doc = _build_front_matter_excerpt_doc()

    tree = DocTree()
    tree.build(doc)

    assert [(sec.section_type, sec.start_index, sec.end_index) for sec in tree.sections] == [
        ("toc", 0, 2),
        ("abstract_cn", 3, 6),
        ("abstract_en", 7, 10),
    ]
    assert tree.get_section("body") is None
    assert tree.get_section("references") is None


def test_header_footer_uses_static_front_headers_for_abstract_excerpt() -> None:
    doc = _build_front_matter_excerpt_doc()
    cfg = load_default_scene()

    tree = DocTree()
    tree.build(doc)

    HeaderFooterRule().apply(doc, cfg, ChangeTracker(), {"doc_tree": tree})

    assert len(doc.sections) == 3
    assert _header_texts(doc) == ["目录", "摘要", "Abstract"]


def test_section_format_keeps_abstract_titles_out_of_back_matter_styleref_pool() -> None:
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("中文摘要内容")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("English abstract content")
    doc.add_paragraph("绪论")
    doc.add_paragraph("正文内容")
    doc.add_paragraph("致谢")
    doc.add_paragraph("感谢内容")

    sections = [
        DocSection("abstract_cn", 0, 1),
        DocSection("abstract_en", 2, 3),
        DocSection("body", 4, 5),
        DocSection("acknowledgment", 6, 7),
    ]

    cfg = load_default_scene()
    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {"doc_tree": _DummyDocTree(sections), "format_scope": cfg.format_scope},
    )

    non_numbered_style = get_non_numbered_heading_style_name(cfg)
    assert doc.paragraphs[0].style.name != non_numbered_style
    assert doc.paragraphs[2].style.name != non_numbered_style
    assert doc.paragraphs[6].style.name == non_numbered_style
    assert _outline_level(doc.paragraphs[0]) == "0"
    assert _outline_level(doc.paragraphs[2]) == "0"
    assert _has_numpr(doc.styles[doc.paragraphs[0].style.name]) is False
    assert _has_numpr(doc.styles[doc.paragraphs[2].style.name]) is False


def test_front_matter_titles_do_not_inherit_heading1_style_numbering() -> None:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 绪论\t1")
    abstract_cn = doc.add_paragraph("摘要", style="Heading 1")
    doc.add_paragraph("中文摘要内容")
    abstract_en = doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph("English abstract content")
    chapter = doc.add_paragraph("绪论", style="Heading 1")
    doc.add_paragraph("正文内容")
    acknowledgement = doc.add_paragraph("致谢", style="Heading 1")
    doc.add_paragraph("感谢内容")
    references = doc.add_paragraph("参考文献", style="Heading 1")
    doc.add_paragraph("[1] 示例参考文献")

    cfg = load_default_scene()
    tree = DocTree()
    tree.build(doc)
    context = {"doc_tree": tree, "format_scope": cfg.format_scope}

    HeadingDetectRule().apply(doc, cfg, ChangeTracker(), context)
    HeadingNumberingRule().apply(doc, cfg, ChangeTracker(), context)
    SectionFormatRule().apply(doc, cfg, ChangeTracker(), context)

    non_numbered_style = get_non_numbered_heading_style_name(cfg)

    assert _has_numpr(doc.styles["Heading 1"]) is True
    assert abstract_cn.style.name not in {"Heading 1", non_numbered_style}
    assert abstract_en.style.name not in {"Heading 1", non_numbered_style}
    assert _has_numpr(doc.styles[abstract_cn.style.name]) is False
    assert _has_numpr(doc.styles[abstract_en.style.name]) is False
    assert _outline_level(abstract_cn) == "0"
    assert _outline_level(abstract_en) == "0"
    assert chapter.style.name == "Heading 1"
    assert acknowledgement.style.name == non_numbered_style
    assert references.style.name == non_numbered_style


def test_back_matter_titles_stay_unnumbered_after_heading1_numbering() -> None:
    doc = Document()
    chapter = doc.add_paragraph("绪论", style="Heading 1")
    doc.add_paragraph("正文内容")
    errata = doc.add_paragraph("勘误页", style="Heading 1")
    doc.add_paragraph("勘误说明")
    acknowledgement = doc.add_paragraph("致谢", style="Heading 1")
    doc.add_paragraph("感谢内容")
    appendix_a = doc.add_paragraph("附录A 公式推导", style="Heading 1")
    doc.add_paragraph("附录A 内容")
    appendix_b = doc.add_paragraph("附录B 补充实验", style="Heading 1")
    doc.add_paragraph("附录B 内容")
    resume = doc.add_paragraph("个人简历", style="Heading 1")
    doc.add_paragraph("个人简历内容")
    resume_subtitle = doc.add_paragraph("在学期间发表的学术论文与研究成果", style="Heading 1")
    doc.add_paragraph("成果列表")
    references = doc.add_paragraph("参考文献", style="Heading 1")
    doc.add_paragraph("[1] 示例参考文献")

    cfg = load_default_scene()
    tree = DocTree()
    tree.build(doc)
    assert [sec.section_type for sec in tree.sections] == [
        "body",
        "errata",
        "acknowledgment",
        "appendix",
        "resume",
        "references",
    ]

    context = {"doc_tree": tree, "format_scope": cfg.format_scope}
    HeadingDetectRule().apply(doc, cfg, ChangeTracker(), context)
    HeadingNumberingRule().apply(doc, cfg, ChangeTracker(), context)
    SectionFormatRule().apply(doc, cfg, ChangeTracker(), context)

    non_numbered_style = get_non_numbered_heading_style_name(cfg)

    assert _has_numpr(doc.styles["Heading 1"]) is True
    assert chapter.style.name == "Heading 1"
    for para in (
        errata,
        acknowledgement,
        appendix_a,
        appendix_b,
        resume,
        resume_subtitle,
        references,
    ):
        assert para.style.name == non_numbered_style
        assert _outline_level(para) == "0"


def test_native_toc_keeps_front_and_back_matter_entries_without_numbering() -> None:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("旧目录条目\t1")
    doc.add_paragraph("旧目录条目\t2")
    doc.add_paragraph("摘要", style="Heading 1")
    doc.add_paragraph("中文摘要内容")
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph("English abstract content")
    doc.add_paragraph("绪论", style="Heading 1")
    doc.add_paragraph("正文内容")
    doc.add_paragraph("致谢", style="Heading 1")
    doc.add_paragraph("感谢内容")
    doc.add_paragraph("参考文献", style="Heading 1")
    doc.add_paragraph("[1] 示例参考文献")

    cfg = load_default_scene()
    tree = DocTree()
    tree.build(doc)
    context = {"doc_tree": tree, "format_scope": cfg.format_scope}

    HeadingDetectRule().apply(doc, cfg, ChangeTracker(), context)
    HeadingNumberingRule().apply(doc, cfg, ChangeTracker(), context)
    SectionFormatRule().apply(doc, cfg, ChangeTracker(), context)
    TocFormatRule().apply(doc, cfg, ChangeTracker(), context)

    toc = tree.get_section("toc")
    assert toc is not None
    toc_texts = [
        (doc.paragraphs[i].text or "").strip()
        for i in range(toc.start_index, toc.end_index + 1)
        if (doc.paragraphs[i].text or "").strip()
    ]
    labels = [text.split("\t", 1)[0] for text in toc_texts]

    assert "摘要" in labels
    assert "Abstract" in labels
    assert "致谢" in labels
    assert "参考文献" in labels
    assert any("绪论" in label for label in labels)
    assert all("第一章" not in text for text in toc_texts if text.split("\t", 1)[0] in {"摘要", "Abstract", "致谢", "参考文献"})


def test_heading_detect_skips_tabular_measurement_rows() -> None:
    doc = Document()
    doc.add_paragraph("表 2.8 BD-CQDs与其余CQDs基复合催化剂的光催化产H2O2性能对比")
    doc.add_paragraph("41.9 mW cm-2\t1562\tthis work")
    doc.add_paragraph("34.8 mW cm-2\t1776\t[68]")
    doc.add_paragraph("22.3 mW cm−2\t1036\t[124]")

    tree = DocTree()
    tree.build(doc)
    context = {"doc_tree": tree}

    HeadingDetectRule().apply(doc, load_default_scene(), ChangeTracker(), context)

    assert context.get("headings", []) == []


def test_equation_table_refresh_keeps_explicit_front_matter_sections() -> None:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 绪论\t1")
    doc.add_paragraph("摘要")
    doc.add_paragraph("中文摘要内容")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("English abstract content")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("正文内容")

    cfg = load_default_scene()
    tree = DocTree()
    tree.build(doc)
    assert [sec.section_type for sec in tree.sections] == [
        "toc",
        "abstract_cn",
        "abstract_en",
        "body",
    ]

    context = {"doc_tree": tree}
    EquationTableFormatRule().apply(doc, cfg, ChangeTracker(), context)

    refreshed = context["doc_tree"]
    assert [sec.section_type for sec in refreshed.sections] == [
        "toc",
        "abstract_cn",
        "abstract_en",
        "body",
    ]

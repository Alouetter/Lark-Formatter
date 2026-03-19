from types import SimpleNamespace

from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.md_cleanup import MdCleanupRule
from src.formula_core.parse import parse_document_formulas
from src.scene.schema import SceneConfig


class _DummyDocTree:
    def __init__(self, paragraph_count: int):
        self.paragraph_count = paragraph_count

    def get_section(self, section_type: str):
        if section_type != "body":
            return None
        return SimpleNamespace(start_index=0, end_index=max(0, self.paragraph_count - 1))

    def build(self, doc: Document, body_start_index=None):
        self.paragraph_count = len(doc.paragraphs)


def _apply_md_cleanup(doc: Document, cfg: SceneConfig) -> None:
    context = {"doc_tree": _DummyDocTree(len(doc.paragraphs))}
    MdCleanupRule().apply(doc, cfg, ChangeTracker(), context)


def test_md_cleanup_normalizes_formula_copy_noise_and_suppresses_fake_formula_lists():
    doc = Document()
    doc.add_paragraph("a+b=ca+b=ca+b=c")
    doc.add_paragraph("1.")
    doc.add_paragraph("2.")
    doc.add_paragraph("a2+b2=c2a^2+b^2=c^2a2+b2=c2")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert [p.text for p in doc.paragraphs] == ["a+b=c", "a^2+b^2=c^2"]


def test_md_cleanup_dedup_unicode_minus_formula():
    """Formula with U+2212 MINUS SIGN repeated 3 times should be deduped."""
    doc = Document()
    doc.add_paragraph(
        "(a\u2212b)2=a2\u22122ab+b2"
        "(a\u2212b)2=a2\u22122ab+b2"
        "(a\u2212b)2=a2\u22122ab+b2"
    )

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    result = doc.paragraphs[0].text
    # Should be deduplicated to a single instance (with ASCII minus after compact)
    assert result.count("=") == 1
    assert "a2" in result or "a^{2}" in result  # dedup happened


def test_md_cleanup_formula_copy_noise_subswitches_can_be_disabled():
    doc = Document()
    doc.add_paragraph("a+b=ca+b=ca+b=c")
    doc.add_paragraph("1.")
    doc.add_paragraph("2.")
    doc.add_paragraph("a2+b2=c2a^2+b^2=c^2a2+b2=c2")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    cfg.md_cleanup.formula_copy_noise_cleanup = False
    cfg.md_cleanup.suppress_formula_fake_lists = False
    _apply_md_cleanup(doc, cfg)

    assert [p.text for p in doc.paragraphs] == [
        "a+b=ca+b=ca+b=c",
        "1.",
        "2.",
        "a2+b2=c2a^2+b^2=c^2a2+b2=c2",
    ]


def test_md_cleanup_recovers_escape_fraction_noise():
    doc = Document()
    doc.add_paragraph("11−xESC_cb14872crac{1}{1-x}1−x1")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert doc.paragraphs[0].text == r"\frac{1}{1-x}"


def test_md_cleanup_recovers_escape_sum_noise():
    doc = Document()
    doc.add_paragraph("∑i=1ni2ESC_ff676190um_{i=1}^{n}i^2i=1∑n\u200bi2")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert doc.paragraphs[0].text == r"\sum_{i=1}^{n}"


def test_md_cleanup_dedups_repeated_standalone_expression():
    doc = Document()
    doc.add_paragraph("(x+y+z)2(x+y+z)2(x+y+z)2")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert doc.paragraphs[0].text == "(x+y+z)2"


def test_md_cleanup_recovers_rendered_simple_fraction_noise():
    doc = Document()
    doc.add_paragraph("11−x11−x1−x1")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert doc.paragraphs[0].text == r"\frac{1}{1-x}"


def test_md_cleanup_dedup_function_notation_formula():
    """Test E[X]=xp(x) repeated 3 times should be deduped."""
    doc = Document()
    doc.add_paragraph("E[X]=xp(x)E[X]=xp(x)E[X]=xp(x)")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    result = doc.paragraphs[0].text
    # Should be deduplicated to a single instance
    assert result.count("=") == 1
    assert "E[X]" in result
    assert "xp(x)" in result


def test_md_cleanup_dedup_limit_formula_with_arrow():
    """Test limit formula with Unicode arrow repeated 3 times should be deduped."""
    doc = Document()
    doc.add_paragraph("limx→0tanx/x=1limx→0tanx/x=1limx→0tanx/x=1")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    result = doc.paragraphs[0].text
    # Should be deduplicated to a single instance (arrow normalized to ->)
    assert result.count("=") == 1
    assert "lim" in result
    assert "tan" in result

def test_md_cleanup_preserves_multiline_latex_display_blocks_for_later_formula_conversion():
    doc = Document()
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\mathcal{Z}(\beta,\mu)")
    doc.add_paragraph("=")
    doc.add_paragraph(r"\sum_{N=0}^{\infty}")
    doc.add_paragraph(r"\]")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    _apply_md_cleanup(doc, cfg)

    assert [p.text for p in doc.paragraphs] == [
        r"\[",
        r"\mathcal{Z}(\beta,\mu)",
        "=",
        r"\sum_{N=0}^{\infty}",
        r"\]",
    ]

    result = parse_document_formulas(doc)
    assert len(result.occurrences) == 1
    assert result.occurrences[0].source_paragraph_indices == [0, 1, 2, 3, 4]


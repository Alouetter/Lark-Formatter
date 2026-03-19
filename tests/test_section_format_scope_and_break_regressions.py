from __future__ import annotations

from copy import deepcopy
from types import SimpleNamespace

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from lxml import etree

from src.engine.change_tracker import ChangeTracker
from src.engine.doc_tree import DocSection
from src.engine.rules.section_format import SectionFormatRule, _looks_like_reference_entry
from src.scene.manager import load_default_scene

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class _DummyDocTree:
    def __init__(self, sections):
        self.sections = sections

    def get_section(self, section_type: str):
        for section in self.sections:
            if section.section_type == section_type:
                return section
        return None

    def get_section_for_paragraph(self, para_index: int) -> str:
        for section in self.sections:
            if section.start_index <= para_index <= section.end_index:
                return section.section_type
        return "unknown"


def _para_has_sectpr(para) -> bool:
    ppr = para._element.find(f"{{{W_NS}}}pPr")
    return ppr is not None and ppr.find(f"{{{W_NS}}}sectPr") is not None


def _append_para_sectpr(doc: Document, para_index: int) -> None:
    para = doc.paragraphs[para_index]
    ppr = para._element.find(f"{{{W_NS}}}pPr")
    if ppr is None:
        ppr = etree.SubElement(para._element, f"{{{W_NS}}}pPr")
        para._element.insert(0, ppr)

    body_sectpr = doc.element.body.find(f"{{{W_NS}}}sectPr")
    if body_sectpr is not None:
        ppr.append(deepcopy(body_sectpr))
    else:
        ppr.append(etree.Element(f"{{{W_NS}}}sectPr"))


def test_disabled_appendix_scope_does_not_center_appendix_table() -> None:
    doc = Document()
    doc.add_paragraph("附录")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cfg = load_default_scene()
    cfg.format_scope.sections["appendix"] = False

    doc_tree = _DummyDocTree([DocSection("appendix", 0, 0)])
    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": doc_tree,
            "format_scope": cfg.format_scope,
        },
    )

    assert table.alignment == WD_TABLE_ALIGNMENT.LEFT


def test_disabled_references_scope_keeps_reference_entry_indent() -> None:
    doc = Document()
    doc.add_paragraph("参考文献")
    para = doc.add_paragraph("[1] reference item")
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.left_indent = Pt(24)

    cfg = load_default_scene()
    cfg.format_scope.sections["references"] = False

    doc_tree = _DummyDocTree([DocSection("references", 0, 1)])
    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": doc_tree,
            "format_scope": cfg.format_scope,
        },
    )

    assert round(para.paragraph_format.first_line_indent.pt, 2) == 24.0
    assert round(para.paragraph_format.left_indent.pt, 2) == 24.0


def test_reference_scope_applies_hanging_indent_and_blocks_body_indent_spillover() -> None:
    doc = Document()
    doc.styles["Normal"].paragraph_format.first_line_indent = Pt(24)
    doc.styles["Normal"].paragraph_format.left_indent = Pt(0)
    doc.add_paragraph("这是正文中的一段普通文字，用于验证正文首行缩进。")
    doc.add_paragraph("参考文献")
    ref_para = doc.add_paragraph("[1] reference item")

    cfg = load_default_scene()
    cfg.styles["normal"].first_line_indent_chars = 2.0
    cfg.styles["normal"].first_line_indent_unit = "chars"
    cfg.styles["normal"].left_indent_chars = 0.0
    cfg.styles["normal"].left_indent_unit = "chars"
    cfg.styles["references_body"].first_line_indent_chars = 0.0
    cfg.styles["references_body"].first_line_indent_unit = "chars"
    cfg.styles["references_body"].left_indent_chars = 2.0
    cfg.styles["references_body"].left_indent_unit = "chars"
    cfg.styles["references_body"].hanging_indent_chars = 2.0
    cfg.styles["references_body"].hanging_indent_unit = "chars"

    doc_tree = _DummyDocTree([DocSection("body", 0, 0), DocSection("references", 1, 2)])
    SectionFormatRule().apply(
        doc,
        cfg,
        ChangeTracker(),
        {
            "doc_tree": doc_tree,
            "format_scope": cfg.format_scope,
        },
    )

    ppr = ref_para._element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind") if ppr is not None else None
    assert ind is not None
    assert ind.get(f"{{{W_NS}}}leftChars") == "200"
    assert ind.get(f"{{{W_NS}}}hangingChars") == "200"
    assert ind.get(f"{{{W_NS}}}left") is None
    assert ind.get(f"{{{W_NS}}}firstLineChars") is None
    assert ind.get(f"{{{W_NS}}}firstLine") is None


def test_reference_entry_detection_accepts_only_square_bracket_numbers() -> None:
    assert _looks_like_reference_entry("[1] reference item") is True
    assert _looks_like_reference_entry("［2］ reference item") is True
    assert _looks_like_reference_entry("(1) 表面接枝CQDs用于光催化合成H2O2：") is False
    assert _looks_like_reference_entry("（1）表面接枝CQDs用于光催化合成H2O2：") is False
    assert _looks_like_reference_entry("1. reference item") is False


def test_chapter_section_break_ignores_unrelated_nearby_sectpr() -> None:
    doc = Document()
    doc.add_paragraph("正文段1")
    doc.add_paragraph("")
    doc.add_paragraph("正文段2")
    doc.add_paragraph("第一章 绪论")
    _append_para_sectpr(doc, 1)

    cfg = load_default_scene()
    rule = SectionFormatRule()
    rule._active_config = cfg
    rule._detected_heading1_indices = {3}

    rule._ensure_chapter_section_breaks(
        doc,
        _DummyDocTree([DocSection("body", 0, 3)]),
        ChangeTracker(),
    )

    assert _para_has_sectpr(doc.paragraphs[2]) is True


def test_post_title_section_break_treats_trailing_table_as_content() -> None:
    doc = Document()
    doc.add_paragraph("正文结尾")
    doc.add_paragraph("附录")
    doc.add_table(rows=1, cols=1)

    cfg = load_default_scene()
    rule = SectionFormatRule()
    rule._active_config = cfg

    doc_tree = SimpleNamespace(
        sections=[DocSection("body", 0, 0), DocSection("appendix", 1, 1)],
        get_section=lambda section_type: None,
    )
    rule._ensure_post_title_section_breaks(doc, doc_tree, ChangeTracker())

    assert _para_has_sectpr(doc.paragraphs[0]) is True

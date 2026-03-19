from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.md_cleanup import MdCleanupRule
from src.markdown.block_parser import parse_docx_paragraphs, parse_markdown_text
from src.markdown.ir import BlockType
from src.scene.schema import SceneConfig


class _DummyDocTree:
    def __init__(self, paragraph_count: int):
        self.paragraph_count = paragraph_count

    def get_section(self, section_type: str):
        if section_type != "body":
            return None
        return SimpleNamespace(start_index=0, end_index=max(0, self.paragraph_count - 1))

    def build(self, doc: Document, body_start_index=None):
        self.paragraph_count = len(doc.paragraphs)


def _has_num_pr(para) -> bool:
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        return False
    return ppr.find(qn("w:numPr")) is not None


def test_parse_markdown_text_supports_cn_ordered_markers():
    blocks = parse_markdown_text("（1）一级\n1、二级\n(2)三级")
    markers = [b.list_marker for b in blocks if b.type == BlockType.LIST_ITEM]
    assert markers == ["（1）", "1、", "(2)"]


def test_parse_docx_cn_ordered_markers_are_ignored_without_markdown_context():
    para_texts = [
        (0, "（1）总则"),
        (1, "1、范围"),
        (2, "2、术语"),
        (3, "（2）方法"),
    ]
    blocks = parse_docx_paragraphs(para_texts, markdown_paste_hint=False)
    list_blocks = [b for b in blocks if b.type == BlockType.LIST_ITEM]
    assert list_blocks == []


def test_parse_docx_cn_ordered_markers_enabled_with_markdown_hint():
    para_texts = [
        (0, "（1）总则"),
        (1, "1、范围"),
        (2, "2、术语"),
        (3, "（2）方法"),
    ]
    blocks = parse_docx_paragraphs(para_texts, markdown_paste_hint=True)
    markers = [b.list_marker for b in blocks if b.type == BlockType.LIST_ITEM]
    assert markers == ["（1）", "1、", "2、", "（2）"]


def test_md_cleanup_skips_cn_ordered_outline_when_no_markdown_signature():
    doc = Document()
    doc.add_paragraph("（1）总则")
    doc.add_paragraph("1、范围")
    doc.add_paragraph("2、术语")
    doc.add_paragraph("（2）方法")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    context = {"doc_tree": _DummyDocTree(len(doc.paragraphs))}
    MdCleanupRule().apply(doc, cfg, ChangeTracker(), context)

    assert [p.text for p in doc.paragraphs] == ["（1）总则", "1、范围", "2、术语", "（2）方法"]
    assert all(not _has_num_pr(p) for p in doc.paragraphs)


def test_md_cleanup_enables_cn_ordered_list_when_manual_break_signature_exists():
    doc = Document()
    doc.add_paragraph("（1）总则\n1、范围\n2、术语")

    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    context = {"doc_tree": _DummyDocTree(len(doc.paragraphs))}
    MdCleanupRule().apply(doc, cfg, ChangeTracker(), context)

    assert len(doc.paragraphs) == 3
    assert [p.text for p in doc.paragraphs] == ["总则", "范围", "术语"]
    assert all(_has_num_pr(p) for p in doc.paragraphs)

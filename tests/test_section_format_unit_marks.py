from __future__ import annotations

import pytest
from docx import Document

from src.engine.rules.section_format import (
    _build_chem_style_marks,
    _restore_reference_chem_typography,
    _run_vert_align,
)


def _superscript_positions(text: str) -> list[int]:
    marks = _build_chem_style_marks(text)
    return [idx for idx, style in enumerate(marks) if style == "superscript"]


def _subscript_positions(text: str) -> list[int]:
    marks = _build_chem_style_marks(text)
    return [idx for idx, style in enumerate(marks) if style == "subscript"]


def _marked_positions(text: str) -> list[int]:
    marks = _build_chem_style_marks(text)
    return [idx for idx, style in enumerate(marks) if style in {"superscript", "subscript"}]


def test_unit_exponents_preserved_with_cjk_tail() -> None:
    text = "mmol h-1 g-1这个参数"
    assert _superscript_positions(text) == [6, 7, 10, 11]

    text_single = "cm-2这个"
    assert _superscript_positions(text_single) == [2, 3]


@pytest.mark.parametrize("text", ["(cm-2)", "（cm-2）", "[cm-2]", "{cm-2}"])
def test_wrapped_unit_exponents_are_detected(text: str) -> None:
    assert _superscript_positions(text) == [3, 4]


def test_multiple_unit_exponents_are_detected() -> None:
    assert _superscript_positions("mol m-2 s-1") == [5, 6, 9, 10]


@pytest.mark.parametrize(
    "text",
    [
        "D1-D4",
        "200k-1.1M",
        "Buck2 Documentation",
        "(D1-D4)",
        "[200k-1.1M]",
        "(A-1)",
    ],
)
def test_non_unit_tokens_remain_plain_text(text: str) -> None:
    assert _superscript_positions(text) == []


@pytest.mark.parametrize(
    "text",
    [
        "A-1",
        "D-1",
        "Appendix A-1",
        "Appendix D-1",
        "AppendixD-1",
        "Fig. D-1",
        "Fig.D-1",
        "Table D-1",
        "TableD-1",
        "表A-1",
        "表 D-1",
        "图D-1",
        "图A-1",
        "图 D-1",
        "见表A-1",
        "p53",
        "P53",
        "H1",
        "F1",
        "H1 标题",
        "h3c",
        "h3c.com",
    ],
)
def test_appendix_labels_and_gene_tokens_remain_plain_text(text: str) -> None:
    assert _marked_positions(text) == []


@pytest.mark.parametrize("text", ["H2B", "B2B", "P2P", "C3PO", "H3K27"])
def test_ambiguous_compact_identifiers_remain_plain_text(text: str) -> None:
    assert _marked_positions(text) == []


@pytest.mark.parametrize(
    "text",
    [
        "300 W??",
        "300 W ??",
        "220 V??",
        "40 K???",
        "300 N??",
    ],
)
def test_spaced_single_letter_measurements_remain_plain_text(text: str) -> None:
    assert _marked_positions(text) == []


@pytest.mark.parametrize(
    "text",
    [
        "300W氙灯",
        "220V电源",
        "40K条件下",
        "300N载荷",
        "220mV偏压",
        "500mW激光",
        "100Pa压力",
        "100MPa强度",
    ],
)
def test_compact_measurement_like_tokens_remain_plain_text(text: str) -> None:
    assert _marked_positions(text) == []


@pytest.mark.parametrize("text", ["14 C??", "15 N??", "35 S??"])
def test_plausible_spaced_single_letter_isotopes_still_format(text: str) -> None:
    assert _superscript_positions(text) == [0, 1]


@pytest.mark.parametrize(
    ("text", "superscripts", "subscripts"),
    [
        ("13C标记", [0, 1], []),
        ("14C示踪", [0, 1], []),
        ("15N标记", [0, 1], []),
        ("18O标记", [0, 1], []),
        ("19F信号", [0, 1], []),
        ("29Si前驱体", [0, 1], []),
        ("31P谱峰", [0, 1], []),
        ("35S标记", [0, 1], []),
        ("57Fe位点", [0, 1], []),
        ("60Co辐照", [0, 1], []),
        ("99Tc示踪", [0, 1], []),
        ("99mTc显像", [0, 1, 2], []),
        ("D2O溶剂", [], [1]),
        ("T2O示踪", [], [1]),
    ],
)
def test_paper_scene_allowlisted_isotopes_are_preserved(
    text: str,
    superscripts: list[int],
    subscripts: list[int],
) -> None:
    assert _superscript_positions(text) == superscripts
    assert _subscript_positions(text) == subscripts


@pytest.mark.parametrize(
    ("text", "superscripts", "subscripts"),
    [
        ("1H", [0], []),
        ("2H", [0], []),
        ("3H", [0], []),
        ("3He", [0], []),
        ("4He", [0], []),
        ("2H2O", [0], [2]),
        ("3H2O", [0], [2]),
        ("1HNMR", [0], []),
        ("13CNMR", [0, 1], []),
        ("15NNMR", [0, 1], []),
        ("19FNMR", [0, 1], []),
        ("29SiNMR", [0, 1], []),
        ("31PNMR", [0, 1], []),
        ("1H-NMR", [0], []),
        ("13C-NMR", [0, 1], []),
        ("15N-NMR", [0, 1], []),
        ("19F-NMR", [0, 1], []),
        ("29Si-NMR", [0, 1], []),
        ("31P-NMR", [0, 1], []),
        ("1H NMR", [0], []),
        ("13C NMR", [0, 1], []),
    ],
)
def test_paper_scene_manual_overrides_cover_nuclei_and_nmr_shorthand(
    text: str,
    superscripts: list[int],
    subscripts: list[int],
) -> None:
    assert _superscript_positions(text) == superscripts
    assert _subscript_positions(text) == subscripts


def test_long_gap_index_like_span_remains_plain_text() -> None:
    assert _marked_positions("初始浓度；11,24                                 H") == []


@pytest.mark.parametrize(
    "text",
    [
        "表A-1",
        "Fig. D-1",
        "Table D-1",
        "p53",
        "H2B",
        "h3c.com",
        "初始浓度；11,24                                 H",
        "300 W??",
        "220 V??",
        "40 K???",
        "300 N??",
    ],
)
def test_restore_reference_chem_typography_keeps_false_positive_tokens_plain(text: str) -> None:
    doc = Document()
    para = doc.add_paragraph(text)

    font_changed, mark_changed = _restore_reference_chem_typography(para)

    assert font_changed == 0
    assert mark_changed == 0
    assert [(run.text, _run_vert_align(run)) for run in para.runs] == [(text, None)]


def test_restore_reference_chem_typography_still_formats_real_formula() -> None:
    doc = Document()
    para = doc.add_paragraph("SO42-")

    font_changed, mark_changed = _restore_reference_chem_typography(para)

    assert font_changed > 0
    assert mark_changed > 0
    assert [(run.text, _run_vert_align(run)) for run in para.runs] == [
        ("SO", None),
        ("4", "subscript"),
        ("2-", "superscript"),
    ]

from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.formula_convert import FormulaConvertRule
from src.engine.rules.formula_to_table import FormulaToTableRule
from src.formula_core.ast import FormulaNode
from src.formula_core.convert import convert_formula_node
from src.formula_core.parse import parse_document_formulas
from src.scene.schema import SceneConfig


def _formula_stats(context: dict, rule_name: str) -> dict:
    return context.get("formula_runtime", {}).get("stats", {}).get(rule_name, {})


def _set_native_formula(paragraph, latex: str, *, block: bool = True) -> None:
    outcome = convert_formula_node(
        FormulaNode(
            kind="equation",
            payload={"latex": latex},
            source_type="latex",
            confidence=0.95,
        ),
        "word_native",
        block=block,
    )
    assert outcome.success is True
    assert outcome.omml_element is not None

    p = paragraph._p
    for child in list(p):
        p.remove(child)

    omml = deepcopy(outcome.omml_element)
    local_name = omml.tag.split("}")[-1] if "}" in omml.tag else omml.tag
    if local_name == "oMathPara":
        p.append(omml)
        return

    run = OxmlElement("w:r")
    run.append(omml)
    p.append(run)


def _add_mathtype_formula(paragraph, text: str) -> None:
    paragraph.add_run(text)
    run = OxmlElement("w:r")
    obj = OxmlElement("w:object")
    control = OxmlElement("w:control")
    control.set(qn("w:name"), "MathType")
    obj.append(control)
    run.append(obj)
    paragraph._p.append(run)


def _body_texts(doc: Document) -> list[str]:
    return [para.text for para in doc.paragraphs if para.text]


def _assert_equation_tables(doc: Document, expected_count: int) -> None:
    assert len(doc.tables) == expected_count
    for table in doc.tables:
        assert "oMath" in table.cell(0, 0)._tc.xml
        assert table.cell(0, 1).text == ""


def test_formula_convert_preserves_body_between_two_extreme_multiline_blocks():
    doc = Document()
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\sum_{k=1}^{n} k = \frac{n(n+1)}{2}")
    doc.add_paragraph(r"\]")
    doc.add_paragraph("BODY_KEEP")
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\int_0^1 x^2 \, dx = \frac{1}{3}")
    doc.add_paragraph(r"\]")

    cfg = SceneConfig()
    cfg.formula_convert.enabled = True
    cfg.formula_to_table.enabled = True

    tracker = ChangeTracker()
    context = {}
    FormulaConvertRule().apply(doc, cfg, tracker, context)

    assert len(doc.paragraphs) == 3
    assert "oMath" in doc.paragraphs[0]._p.xml
    assert doc.paragraphs[1].text == "BODY_KEEP"
    assert "oMath" in doc.paragraphs[2]._p.xml

    FormulaToTableRule().apply(doc, cfg, tracker, context)

    assert len(doc.tables) == 2
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "BODY_KEEP"

    convert_stats = _formula_stats(context, "formula_convert")
    to_table_stats = _formula_stats(context, "formula_to_table")
    assert convert_stats.get("converted") == 2
    assert to_table_stats.get("converted") == 2


def test_formula_convert_handles_aligned_extreme_block_and_round_trips_to_table():
    doc = Document()
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\begin{aligned}")
    doc.add_paragraph(r"f(x)&=x^2+2x+1\\")
    doc.add_paragraph(r"&=(x+1)^2")
    doc.add_paragraph(r"\end{aligned}")
    doc.add_paragraph(r"\]")

    cfg = SceneConfig()
    cfg.formula_convert.enabled = True
    cfg.formula_to_table.enabled = True

    tracker = ChangeTracker()
    context = {}
    FormulaConvertRule().apply(doc, cfg, tracker, context)

    assert len(doc.paragraphs) == 1
    para_xml = doc.paragraphs[0]._p.xml
    assert "oMath" in para_xml
    assert r"\begin{aligned}" not in para_xml
    assert r"\end{aligned}" not in para_xml

    FormulaToTableRule().apply(doc, cfg, tracker, context)

    assert len(doc.tables) == 1
    assert len(doc.paragraphs) == 0

    convert_stats = _formula_stats(context, "formula_convert")
    to_table_stats = _formula_stats(context, "formula_to_table")
    assert convert_stats.get("converted") == 1
    assert to_table_stats.get("converted") == 1


def test_formula_convert_mixed_sources_word_native_round_trip_to_tables():
    doc = Document()
    doc.add_paragraph("BODY_KEEP_TOP")
    doc.add_paragraph("sum_(i=1)^n i^2 = n(n+1)(2n+1)/6")
    doc.add_paragraph("BODY_KEEP_MID")

    native_para = doc.add_paragraph()
    _set_native_formula(native_para, r"\int_0^1 \frac{x^2+1}{x+1} \, dx = 1")

    doc.add_paragraph("BODY_KEEP_AFTER_NATIVE")
    doc.add_paragraph(r"\[")
    doc.add_paragraph(r"\begin{aligned}")
    doc.add_paragraph(r"f(x)&=\sum_{k=1}^{n} k^2\\")
    doc.add_paragraph(r"&=\frac{n(n+1)(2n+1)}{6}")
    doc.add_paragraph(r"\end{aligned}")
    doc.add_paragraph(r"\]")
    doc.add_paragraph("BODY_KEEP_AFTER_LATEX")

    mathtype_para = doc.add_paragraph()
    _add_mathtype_formula(mathtype_para, "lim x->0 (sin x)/x = 1")
    doc.add_paragraph("BODY_KEEP_TAIL")

    parsed = parse_document_formulas(doc)
    assert parsed.total == 4
    assert parsed.by_source == {
        "plain_text": 1,
        "word_native": 1,
        "latex": 1,
        "mathtype": 1,
    }

    cfg = SceneConfig()
    cfg.formula_convert.enabled = True
    cfg.formula_to_table.enabled = True

    tracker = ChangeTracker()
    context = {}
    FormulaConvertRule().apply(doc, cfg, tracker, context)

    assert len(doc.paragraphs) == 9
    assert _body_texts(doc) == [
        "BODY_KEEP_TOP",
        "BODY_KEEP_MID",
        "BODY_KEEP_AFTER_NATIVE",
        "BODY_KEEP_AFTER_LATEX",
        "BODY_KEEP_TAIL",
    ]
    assert sum("oMath" in para._p.xml for para in doc.paragraphs) == 4

    convert_stats = _formula_stats(context, "formula_convert")
    assert convert_stats.get("matched") == 4
    assert convert_stats.get("converted") == 3
    assert convert_stats.get("removed_wrappers") == 0
    assert context.get("formula_runtime", {}).get("low_confidence", []) == []

    FormulaToTableRule().apply(doc, cfg, tracker, context)

    assert _body_texts(doc) == [
        "BODY_KEEP_TOP",
        "BODY_KEEP_MID",
        "BODY_KEEP_AFTER_NATIVE",
        "BODY_KEEP_AFTER_LATEX",
        "BODY_KEEP_TAIL",
    ]
    _assert_equation_tables(doc, expected_count=4)

    to_table_stats = _formula_stats(context, "formula_to_table")
    assert to_table_stats.get("matched") == 4
    assert to_table_stats.get("converted") == 4


def test_formula_convert_mixed_sources_latex_mode_round_trip_to_tables():
    doc = Document()
    doc.add_paragraph("BODY_KEEP_HEAD")
    doc.add_paragraph("E=mc^2")

    native_para = doc.add_paragraph()
    _set_native_formula(native_para, r"\begin{bmatrix}a & b \\ c & d\end{bmatrix}")

    doc.add_paragraph(r"\left(\frac{a+b}{c+d}\right)^2=\frac{(a+b)^2}{(c+d)^2}")

    mathtype_para = doc.add_paragraph()
    _add_mathtype_formula(mathtype_para, "x_i^2+y_i^2=z_i^2")
    doc.add_paragraph("BODY_KEEP_TAIL")

    parsed = parse_document_formulas(doc)
    assert parsed.total == 4
    assert parsed.by_source == {
        "plain_text": 1,
        "word_native": 1,
        "latex": 1,
        "mathtype": 1,
    }

    native_xml_before = native_para._p.xml

    cfg = SceneConfig()
    cfg.formula_convert.enabled = True
    cfg.formula_convert.output_mode = "latex"
    cfg.formula_to_table.enabled = True

    tracker = ChangeTracker()
    context = {}
    FormulaConvertRule().apply(doc, cfg, tracker, context)

    assert native_para._p.xml == native_xml_before
    assert _body_texts(doc) == ["BODY_KEEP_HEAD", "BODY_KEEP_TAIL"]
    assert sum("oMath" in para._p.xml for para in doc.paragraphs) == 4

    runtime = context.get("formula_runtime", {})
    convert_stats = _formula_stats(context, "formula_convert")
    assert runtime.get("output_mode") == "latex"
    assert runtime.get("latex_logic_editable_output") is True
    assert convert_stats.get("matched") == 4
    assert convert_stats.get("converted") == 3

    latex_exchange = runtime.get("latex_exchange", [])
    assert len(latex_exchange) == 3
    latex_texts = [item.get("latex", "") for item in latex_exchange]
    assert any("mc^{2}" in text for text in latex_texts)
    assert any(r"\frac{a+b}{c+d}" in text for text in latex_texts)
    assert any(r"x_{i}^{2}+y_{i}^{2}=z_{i}^{2}" in text for text in latex_texts)

    FormulaToTableRule().apply(doc, cfg, tracker, context)

    assert _body_texts(doc) == ["BODY_KEEP_HEAD", "BODY_KEEP_TAIL"]
    _assert_equation_tables(doc, expected_count=4)

    to_table_stats = _formula_stats(context, "formula_to_table")
    assert to_table_stats.get("matched") == 4
    assert to_table_stats.get("converted") == 4

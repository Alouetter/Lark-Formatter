from pathlib import Path

from docx import Document
from lxml import etree

from src.engine.pipeline import Pipeline
from src.engine.rules.section_format import SectionFormatRule
from src.engine.rules.style_manager import _apply_style_config
from src.scene.manager import load_scene
from src.scene.manager import load_scene_from_data
from src.scene.schema import StyleConfig

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_normal_style_char_indents(doc: Document, *, first_chars: float = 0.0, left_chars: float = 0.0) -> None:
    style = doc.styles["Normal"]
    ppr = style.element.find(f"{{{W_NS}}}pPr")
    if ppr is None:
        ppr = etree.SubElement(style.element, f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind")
    if ind is None:
        ind = etree.SubElement(ppr, f"{{{W_NS}}}ind")
    if first_chars > 0:
        ind.set(f"{{{W_NS}}}firstLineChars", str(int(round(first_chars * 100))))
    if left_chars > 0:
        ind.set(f"{{{W_NS}}}leftChars", str(int(round(left_chars * 100))))


def _ind_attrs(ind) -> dict[str, str]:
    attrs: dict[str, str] = {}
    if ind is None:
        return attrs
    for name in (
        "left",
        "leftChars",
        "firstLine",
        "firstLineChars",
        "hanging",
        "hangingChars",
        "right",
        "rightChars",
    ):
        value = ind.get(f"{{{W_NS}}}{name}")
        if value is not None:
            attrs[name] = value
    return attrs


def test_pipeline_applies_explicit_paragraph_indents_over_legacy_char_indents(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "indent_demo.docx"
    doc = Document()
    _set_normal_style_char_indents(doc, first_chars=2.0, left_chars=1.0)
    doc.add_paragraph("第一段正文")
    doc.save(str(doc_path))

    cfg = load_scene(Path("src/scene/presets/default_format.json"))
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False
    cfg.pipeline = ["style_manager", "section_format"]
    cfg.pipeline_critical_rules = list(cfg.pipeline)
    cfg.styles["normal"].size_pt = 12.0
    cfg.styles["normal"].first_line_indent_chars = 20.0
    cfg.styles["normal"].first_line_indent_unit = "pt"
    cfg.styles["normal"].left_indent_chars = 18.0
    cfg.styles["normal"].left_indent_unit = "pt"

    result = Pipeline(cfg).run(str(doc_path))
    out_doc = Document(result.output_paths["final"])
    para = out_doc.paragraphs[0]

    assert round(para.paragraph_format.first_line_indent.pt, 2) == 20.0
    assert round(para.paragraph_format.left_indent.pt, 2) == 18.0

    ppr = para._element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind")
    assert ind.get(f"{{{W_NS}}}firstLine") == "400"
    assert ind.get(f"{{{W_NS}}}left") == "360"
    assert ind.get(f"{{{W_NS}}}firstLineChars") is None
    assert ind.get(f"{{{W_NS}}}leftChars") is None


def test_pipeline_applies_large_pt_indent_exactly(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "indent_large_pt.docx"
    doc = Document()
    doc.add_paragraph("缩进测试段落")
    doc.save(str(doc_path))

    cfg = load_scene(Path("src/scene/presets/default_format.json"))
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False
    cfg.pipeline = ["style_manager", "section_format"]
    cfg.pipeline_critical_rules = list(cfg.pipeline)
    cfg.styles["normal"].size_pt = 12.0
    cfg.styles["normal"].first_line_indent_chars = 200.0
    cfg.styles["normal"].first_line_indent_unit = "pt"
    cfg.styles["normal"].left_indent_chars = 200.0
    cfg.styles["normal"].left_indent_unit = "pt"

    result = Pipeline(cfg).run(str(doc_path))
    out_doc = Document(result.output_paths["final"])
    para = out_doc.paragraphs[0]

    assert round(para.paragraph_format.first_line_indent.pt, 2) == 200.0
    assert round(para.paragraph_format.left_indent.pt, 2) == 200.0

    ppr = para._element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind")
    assert ind.get(f"{{{W_NS}}}firstLine") == "4000"
    assert ind.get(f"{{{W_NS}}}left") == "4000"
    assert ind.get(f"{{{W_NS}}}firstLineChars") is None
    assert ind.get(f"{{{W_NS}}}leftChars") is None


def test_pipeline_updates_normal_style_indent_when_pt_mode_is_configured(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "indent_normal_style_legacy.docx"
    doc = Document()
    _set_normal_style_char_indents(doc, first_chars=2.0, left_chars=0.0)
    doc.add_paragraph("缩进测试段落")
    doc.save(str(doc_path))

    cfg = load_scene(Path("src/scene/presets/default_format.json"))
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False
    cfg.pipeline = ["style_manager", "section_format"]
    cfg.pipeline_critical_rules = list(cfg.pipeline)
    cfg.styles["normal"].size_pt = 12.0
    cfg.styles["normal"].first_line_indent_chars = 40.0
    cfg.styles["normal"].first_line_indent_unit = "pt"
    cfg.styles["normal"].left_indent_chars = 0.0
    cfg.styles["normal"].left_indent_unit = "chars"

    result = Pipeline(cfg).run(str(doc_path))
    out_doc = Document(result.output_paths["final"])

    style = out_doc.styles["Normal"]
    ppr = style.element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind")

    assert ind.get(f"{{{W_NS}}}firstLine") == "800"
    assert ind.get(f"{{{W_NS}}}firstLineChars") is None


def test_pipeline_preserves_char_unit_indents_as_word_char_indents(tmp_path, monkeypatch):
    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")

    doc_path = tmp_path / "indent_chars.docx"
    doc = Document()
    doc.add_paragraph("第一段正文")
    doc.save(str(doc_path))

    cfg = load_scene(Path("src/scene/presets/default_format.json"))
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False
    cfg.pipeline = ["style_manager", "section_format"]
    cfg.pipeline_critical_rules = list(cfg.pipeline)
    cfg.styles["normal"].size_pt = 12.0
    cfg.styles["normal"].first_line_indent_chars = 3.0
    cfg.styles["normal"].first_line_indent_unit = "chars"
    cfg.styles["normal"].left_indent_chars = 1.5
    cfg.styles["normal"].left_indent_unit = "chars"

    result = Pipeline(cfg).run(str(doc_path))
    out_doc = Document(result.output_paths["final"])
    para = out_doc.paragraphs[0]

    ppr = para._element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind")
    assert ind.get(f"{{{W_NS}}}firstLineChars") == "300"
    assert ind.get(f"{{{W_NS}}}leftChars") == "150"
    assert ind.get(f"{{{W_NS}}}firstLine") is None
    assert ind.get(f"{{{W_NS}}}left") is None


def test_zero_indent_on_heading_style_overrides_inherited_char_indent(tmp_path):
    doc = Document()
    _set_normal_style_char_indents(doc, first_chars=2.0, left_chars=1.0)

    _apply_style_config(
        doc.styles["Heading 1"],
        StyleConfig(
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=12.0,
            alignment="left",
            first_line_indent_chars=0.0,
            first_line_indent_unit="chars",
            left_indent_chars=0.0,
            left_indent_unit="chars",
            line_spacing_type="multiple",
            line_spacing_pt=1.5,
        ),
    )

    doc_path = tmp_path / "heading_style_zero_indent.docx"
    doc.save(str(doc_path))
    out_doc = Document(str(doc_path))

    style = out_doc.styles["Heading 1"]
    ppr = style.element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind") if ppr is not None else None
    attrs = _ind_attrs(ind)

    assert attrs["left"] == "0"
    assert attrs["leftChars"] == "0"
    assert attrs["firstLine"] == "0"
    assert attrs["firstLineChars"] == "0"


def test_unnumbered_heading_style_explicitly_clears_inherited_char_indent(tmp_path):
    doc = Document()
    _set_normal_style_char_indents(doc, first_chars=2.0, left_chars=1.0)

    rule = SectionFormatRule()
    rule._active_config = load_scene_from_data({})
    rule._ensure_non_numbered_heading_style(doc)

    doc_path = tmp_path / "unnumbered_heading_style_zero_indent.docx"
    doc.save(str(doc_path))
    out_doc = Document(str(doc_path))

    style = out_doc.styles["Heading 1 Unnumbered"]
    ppr = style.element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind") if ppr is not None else None
    attrs = _ind_attrs(ind)

    assert attrs["left"] == "0"
    assert attrs["leftChars"] == "0"
    assert attrs["firstLine"] == "0"
    assert attrs["firstLineChars"] == "0"


def test_apply_style_config_writes_char_based_hanging_indent(tmp_path):
    doc = Document()

    _apply_style_config(
        doc.styles["Normal"],
        StyleConfig(
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=12.0,
            alignment="justify",
            first_line_indent_chars=0.0,
            first_line_indent_unit="chars",
            left_indent_chars=2.0,
            left_indent_unit="chars",
            hanging_indent_chars=2.0,
            hanging_indent_unit="chars",
            line_spacing_type="exact",
            line_spacing_pt=20.0,
        ),
    )

    doc_path = tmp_path / "style_hanging_indent.docx"
    doc.save(str(doc_path))
    out_doc = Document(str(doc_path))

    style = out_doc.styles["Normal"]
    ppr = style.element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind") if ppr is not None else None
    attrs = _ind_attrs(ind)

    assert attrs["leftChars"] == "200"
    assert attrs["hangingChars"] == "200"
    assert "firstLine" not in attrs
    assert "firstLineChars" not in attrs


def test_apply_style_config_writes_cm_based_first_line_indent(tmp_path):
    doc = Document()

    _apply_style_config(
        doc.styles["Normal"],
        StyleConfig(
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=12.0,
            alignment="justify",
            left_indent_chars=0.0,
            left_indent_unit="chars",
            special_indent_mode="first_line",
            special_indent_value=1.0,
            special_indent_unit="cm",
            line_spacing_type="exact",
            line_spacing_pt=20.0,
        ),
    )

    doc_path = tmp_path / "style_cm_first_indent.docx"
    doc.save(str(doc_path))
    out_doc = Document(str(doc_path))

    style = out_doc.styles["Normal"]
    ppr = style.element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind") if ppr is not None else None
    attrs = _ind_attrs(ind)

    expected_twips = str(int(round((72.0 / 2.54) * 20)))
    assert attrs["firstLine"] == expected_twips
    assert "firstLineChars" not in attrs


def test_apply_style_config_writes_right_indent(tmp_path):
    doc = Document()

    _apply_style_config(
        doc.styles["Normal"],
        StyleConfig(
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=12.0,
            alignment="justify",
            right_indent_chars=1.5,
            right_indent_unit="chars",
            line_spacing_type="exact",
            line_spacing_pt=20.0,
        ),
    )

    doc_path = tmp_path / "style_right_indent.docx"
    doc.save(str(doc_path))
    out_doc = Document(str(doc_path))

    style = out_doc.styles["Normal"]
    ppr = style.element.find(f"{{{W_NS}}}pPr")
    ind = ppr.find(f"{{{W_NS}}}ind") if ppr is not None else None
    attrs = _ind_attrs(ind)

    assert attrs["rightChars"] == "150"

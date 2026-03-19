from pathlib import Path

from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.citation_link import CitationLinkRule
from src.scene.manager import load_scene
from src.scene.schema import SceneConfig


def _apply_citation_link(doc: Document, *, auto_number: bool = True) -> ChangeTracker:
    cfg = SceneConfig()
    cfg.citation_link.enabled = True
    cfg.citation_link.auto_number_reference_entries = auto_number
    tracker = ChangeTracker()
    CitationLinkRule().apply(doc, cfg, tracker, {"doc_tree": None})
    return tracker


def test_citation_link_repairs_existing_field_paragraph_and_links_new_plain_text_tail():
    doc = Document()
    body = doc.add_paragraph("Body [1].")
    doc.add_paragraph("References")
    doc.add_paragraph("[1] Ref A")
    doc.add_paragraph("[2] Ref B")

    _apply_citation_link(doc)

    body.add_run(" and [2].")
    tracker = _apply_citation_link(doc)

    xml = body._element.xml
    assert "REF _RefNum_RefEntry_1" in xml
    assert "REF _RefNum_RefEntry_2" in xml
    assert any(
        rec.target == "paragraph #0"
        and "normalized existing citation fields" in rec.after
        for rec in tracker.records
    )
    assert any(
        rec.target == "paragraph #0"
        and "converted to linked field" in rec.after
        for rec in tracker.records
    )


def test_citation_link_links_only_square_bracket_numeric_citations_in_body():
    doc = Document()
    body = doc.add_paragraph("Body ［1］, [2,3], [1-3], (2) and （3）.")

    doc.add_paragraph("References")
    doc.add_paragraph("［1］ Ref A")
    doc.add_paragraph("[2] Ref B")
    doc.add_paragraph("[3] Ref C")

    _apply_citation_link(doc)

    xml = body._element.xml
    assert xml.count("REF _RefNum_") == 5
    assert "(2)" in body.text
    assert "（3）" in body.text


def test_citation_link_ignores_non_square_bracket_list_items_inside_references():
    doc = Document()
    body = doc.add_paragraph("Body [1] and [2].")
    doc.add_paragraph("References")
    pseudo_item = doc.add_paragraph("(1) 表面接枝CQDs用于光催化合成H2O2：")
    doc.add_paragraph("[2] Ref B")

    _apply_citation_link(doc)

    body_xml = body._element.xml
    assert body_xml.count("REF _RefNum_") == 1
    assert "REF _RefNum_" not in pseudo_item._element.xml
    assert "SEQ RefEntry" not in pseudo_item._element.xml


def test_default_scene_preset_enables_citation_auto_numbering():
    cfg = load_scene(Path("src/scene/presets/default_format.json"))

    assert cfg.citation_link.enabled is True
    assert cfg.citation_link.auto_number_reference_entries is True
    assert "citation_link" in (cfg.pipeline or [])


def test_duplicated_reference_numbers_are_left_unlinked_to_avoid_ambiguity():
    doc = Document()
    body = doc.add_paragraph("Body [1].")
    doc.add_paragraph("References")
    doc.add_paragraph("[1] Ref A")
    doc.add_paragraph("[1] Ref B")

    tracker = _apply_citation_link(doc)

    assert " REF " not in body._element.xml
    assert body.text == "Body [1]."
    assert any(
        rec.target == "references"
        and "duplicated numbering detected" in rec.before
        and rec.after == "duplicate references left unlinked to avoid ambiguity"
        for rec in tracker.records
    )

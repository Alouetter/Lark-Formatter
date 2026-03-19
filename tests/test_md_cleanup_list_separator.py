from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.md_cleanup import MdCleanupRule
from src.markdown.word_render import apply_num_pr, register_list_numbering
from src.scene.schema import SceneConfig


class _DummyDocTree:
    def __init__(self, paragraph_count: int):
        self.paragraph_count = paragraph_count

    def get_section(self, section_type: str):
        if section_type != "body":
            return None
        return SimpleNamespace(start_index=0, end_index=max(0, self.paragraph_count - 1))

    def build(self, doc: Document, body_start_index=None):
        self.paragraph_count = len(doc.paragraphs)


def _make_native_bullet_doc() -> Document:
    doc = Document()
    para1 = doc.add_paragraph("item1")
    para2 = doc.add_paragraph("item2")
    num_id = register_list_numbering(doc, "bullet", marker_separator="tab")
    apply_num_pr(para1, num_id, 0)
    apply_num_pr(para2, num_id, 0)
    return doc


def _get_para_num_id(para) -> str | None:
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        return None
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        return None
    return num_id.get(qn("w:val"))


def _get_para_suff(doc: Document, para) -> str | None:
    num_id = _get_para_num_id(para)
    if not num_id:
        return None
    numbering = doc.part.numbering_part.element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    num_el = numbering.find(f".//w:num[@w:numId='{num_id}']", ns)
    if num_el is None:
        return None
    abstract_ref = num_el.find("w:abstractNumId", ns)
    if abstract_ref is None:
        return None
    abstract_id = abstract_ref.get(qn("w:val"))
    if not abstract_id:
        return None
    lvl0 = numbering.find(
        f".//w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='0']",
        ns,
    )
    if lvl0 is None:
        return None
    suff = lvl0.find("w:suff", ns)
    if suff is None:
        return None
    return suff.get(qn("w:val"))


def _run_md_cleanup(doc: Document, separator: str) -> None:
    cfg = SceneConfig()
    cfg.md_cleanup.enabled = True
    cfg.md_cleanup.list_marker_separator = separator
    context = {"doc_tree": _DummyDocTree(len(doc.paragraphs))}
    MdCleanupRule().apply(doc, cfg, ChangeTracker(), context)


def test_md_cleanup_can_convert_existing_native_list_separator_to_half_space():
    doc = _make_native_bullet_doc()
    _run_md_cleanup(doc, "half_space")

    assert _get_para_suff(doc, doc.paragraphs[0]) == "space"


def test_md_cleanup_can_convert_existing_native_list_separator_to_full_space():
    doc = _make_native_bullet_doc()
    _run_md_cleanup(doc, "full_space")

    assert _get_para_suff(doc, doc.paragraphs[0]) == "nothing"
    assert doc.paragraphs[0].text.startswith("\u3000")


def test_md_cleanup_tab_mode_removes_full_space_prefix_on_existing_lists():
    doc = _make_native_bullet_doc()
    _run_md_cleanup(doc, "full_space")
    _run_md_cleanup(doc, "tab")

    assert _get_para_suff(doc, doc.paragraphs[0]) == "tab"
    assert not doc.paragraphs[0].text.startswith("\u3000")

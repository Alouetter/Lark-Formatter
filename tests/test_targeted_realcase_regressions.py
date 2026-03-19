from __future__ import annotations

from pathlib import Path

from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.doc_tree import DocTree
from src.engine.pipeline import Pipeline
from src.engine.rules.header_footer import HeaderFooterRule
from src.engine.rules.heading_detect import HeadingDetectRule
from src.scene.manager import load_default_scene


def _header_texts(doc: Document) -> list[str]:
    texts: list[str] = []
    for sec in doc.sections:
        text = "".join(
            (para.text or "").strip()
            for para in sec.header.paragraphs
            if (para.text or "").strip()
        )
        texts.append(text)
    return texts


def test_catalog_abstract_front_matter_headers_regression() -> None:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 绪论\t1")
    doc.add_paragraph("Abstract\t4")

    doc.add_paragraph("摘要")
    doc.add_paragraph(
        "这是中文摘要的主体内容，用于回归测试目录、摘要、Abstract 的分节与页眉。"
    )
    doc.add_paragraph(
        "（1）这里故意放一个较长的摘要枚举段，验证它不会被误判为参考文献或正文。"
    )
    doc.add_paragraph("关键词：智能编译；统一IR")

    doc.add_paragraph("Abstract")
    doc.add_paragraph(
        "This is the English abstract body used to verify front-matter section "
        "breaks and header assignment."
    )
    doc.add_paragraph(
        "(1) This long abstract bullet should stay inside the English abstract "
        "instead of becoming references or body content."
    )
    doc.add_paragraph("Key Words: intelligent compilation; unified IR")

    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("这里是正文内容。")

    tree = DocTree()
    tree.build(doc)
    cfg = load_default_scene()

    HeaderFooterRule().apply(doc, cfg, ChangeTracker(), {"doc_tree": tree})

    headers = _header_texts(doc)
    assert headers[:3] == ["目录", "摘要", "Abstract"]

    assert tree.get_section("toc") is not None
    assert tree.get_section("abstract_cn") is not None
    assert tree.get_section("abstract_en") is not None
    assert tree.get_section("body") is not None


def test_table2_measurement_rows_are_not_detected_as_headings() -> None:
    doc = Document()
    doc.add_paragraph("表 2.8 BD-CQDs与其余CQDs基复合催化剂的光催化产H2O2性能对比")
    doc.add_paragraph("41.9 mW cm-2\t1562\tthis work")
    doc.add_paragraph("34.8 mW cm-2\t1776\t[68]")
    doc.add_paragraph("22.3 mW cm−2\t1036\t[124]")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("这里才是真正的正文。")

    cfg = load_default_scene()
    tree = DocTree()
    tree.build(doc)
    context = {"doc_tree": tree}

    HeadingDetectRule().apply(doc, cfg, ChangeTracker(), context)

    headings = context.get("headings", [])
    assert [(item.para_index, item.level) for item in headings] == [(4, "heading1")]


def test_pipeline_keeps_cover_and_front_headers_for_zhongtu_case(
    tmp_path, monkeypatch
) -> None:
    source = Path(__file__).with_name("中图分类号.docx")
    work_doc = tmp_path / source.name
    work_doc.write_bytes(source.read_bytes())

    cfg = load_default_scene()
    cfg.output.final_docx = True
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False

    monkeypatch.setenv("DOCX_DISABLE_FIELD_REFRESH", "1")
    result = Pipeline(cfg).run(str(work_doc))

    assert result.success is True
    output_doc = Document(str(tmp_path / "中图分类号_new.docx"))

    assert len(output_doc.sections) == 5
    assert _header_texts(output_doc)[:4] == ["", "目录", "摘要", "Abstract"]

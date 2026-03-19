from docx import Document

from src.engine.doc_tree import DocTree
from src.formula_core.normalize import looks_like_formula_text
from src.formula_core.parse import parse_document_formulas
from src.utils.toc_entry import (
    looks_like_bibliographic_reference_line,
    looks_like_reference_entry_line,
)


_REF_WITHOUT_SPACE = (
    "[1]SHARMAS,CHOWDHURYP.Urea\u2212basedcarbonquantumdotsforrecognitionofFe3+"
    "byfluorescencequenchingandPb2+byfluorescenceenhancement[J]."
    "Chempluschem,2025,90(6):e202400715."
)

_REF_WITH_SPACE = (
    "[1] SHARMA S, CHOWDHURY P. Urea-based carbon quantum dots for recognition "
    "of Fe3+ by fluorescence quenching and Pb2+ by fluorescence enhancement [J]. "
    "ChemPlusChem, 2025, 90(6): e202400715."
)

_UNNUMBERED_REF = (
    "SHARMA S, CHOWDHURY P. Urea-based carbon quantum dots for recognition of "
    "Fe3+ by fluorescence quenching and Pb2+ by fluorescence enhancement [J]. "
    "ChemPlusChem, 2025, 90(6): e202400715."
)


def test_reference_entry_heuristics_accept_missing_space_after_index():
    assert looks_like_reference_entry_line(_REF_WITHOUT_SPACE) is True
    assert looks_like_bibliographic_reference_line(_REF_WITHOUT_SPACE) is True


def test_formula_detection_skips_reference_entries_with_chemical_ions():
    assert looks_like_formula_text(_REF_WITHOUT_SPACE) == (False, 0.0, "plain_text")
    assert looks_like_formula_text(_REF_WITH_SPACE) == (False, 0.0, "plain_text")
    assert looks_like_formula_text(_UNNUMBERED_REF) == (False, 0.0, "plain_text")


def test_parse_document_formulas_skips_reference_entries_with_chemical_ions():
    doc = Document()
    doc.add_paragraph(_REF_WITHOUT_SPACE)
    doc.add_paragraph(_UNNUMBERED_REF)

    parsed = parse_document_formulas(doc)

    assert parsed.total == 0


def test_formula_detection_keeps_real_chemical_equation():
    text = "Fe3+ + e- = Fe2+"

    matched, confidence, source = looks_like_formula_text(text)

    assert matched is True
    assert confidence >= 0.66
    assert source == "plain_text"

    doc = Document()
    doc.add_paragraph(text)
    parsed = parse_document_formulas(doc)
    assert parsed.total == 1
    assert parsed.occurrences[0].source_type == "plain_text"


def test_doc_tree_detects_reference_cluster_even_without_space_after_index():
    doc = Document()
    doc.add_paragraph("\u6b63\u6587\u6bb5\u843d")
    doc.add_paragraph(_REF_WITHOUT_SPACE)
    doc.add_paragraph(
        "[2]WANGX,LIY.CarbondotsforCu2+detection[J].Analyticalletters,2024,58(3):123-130."
    )

    tree = DocTree()
    tree.build(doc)

    ref = tree.get_section("references")
    assert ref is not None
    assert ref.start_index == 1

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.caption_format import _format_caption_para
from src.engine.rules.header_footer import HeaderFooterRule
from src.engine.rules.page_setup import PageSetupRule
from src.scene.manager import load_scene_from_data

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def test_caption_style_fields_apply_to_docx_paragraph_and_runs():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("图1 测试题注")

    cfg = load_scene_from_data({})
    sc = cfg.styles["figure_caption"]
    sc.bold = True
    sc.italic = True
    sc.alignment = "right"
    sc.first_line_indent_chars = 18.0
    sc.first_line_indent_unit = "pt"
    sc.left_indent_chars = 24.0
    sc.left_indent_unit = "pt"
    sc.line_spacing_type = "exact"
    sc.line_spacing_pt = 18.0
    sc.space_before_pt = 6.0
    sc.space_after_pt = 8.0

    _format_caption_para(para, sc)

    assert para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert round(para.paragraph_format.first_line_indent.pt, 2) == 18.0
    assert round(para.paragraph_format.left_indent.pt, 2) == 24.0
    assert round(para.paragraph_format.space_before.pt, 2) == 6.0
    assert round(para.paragraph_format.space_after.pt, 2) == 8.0
    assert round(para.paragraph_format.line_spacing.pt, 2) == 18.0

    assert run.font.bold is True
    assert run.font.italic is True
    assert round(run.font.size.pt, 2) == round(sc.size_pt, 2)

    ppr = para._element.find(_w("pPr"))
    ind = ppr.find(_w("ind"))
    assert ind.get(_w("left")) == "480"
    assert ind.get(_w("firstLine")) == "360"
    assert ind.get(_w("right")) == "0"

    rpr = run._element.find(_w("rPr"))
    assert rpr.find(_w("b")) is not None
    assert rpr.find(_w("i")) is not None
    rfonts = rpr.find(_w("rFonts"))
    assert rfonts.get(_w("ascii")) == sc.font_en
    assert rfonts.get(_w("eastAsia")) == sc.font_cn
    assert rfonts.get(_w("hint")) == "default"


def test_header_footer_rule_builders_apply_full_style_config():
    doc = Document()
    section = doc.sections[0]

    cfg = load_scene_from_data({})
    header_sc = cfg.styles["header_cn"]
    header_sc.bold = True
    header_sc.italic = True
    header_sc.alignment = "right"
    header_sc.left_indent_chars = 18.0
    header_sc.left_indent_unit = "pt"
    header_sc.first_line_indent_chars = 6.0
    header_sc.first_line_indent_unit = "pt"
    header_sc.line_spacing_type = "exact"
    header_sc.line_spacing_pt = 17.0
    header_sc.space_before_pt = 3.0
    header_sc.space_after_pt = 5.0

    page_sc = cfg.styles["page_number"]
    page_sc.bold = False
    page_sc.italic = True
    page_sc.alignment = "left"
    page_sc.left_indent_chars = 12.0
    page_sc.left_indent_unit = "pt"
    page_sc.first_line_indent_chars = 3.0
    page_sc.first_line_indent_unit = "pt"
    page_sc.line_spacing_type = "double"
    page_sc.space_before_pt = 2.0
    page_sc.space_after_pt = 4.0

    rule = HeaderFooterRule()
    rule._apply_header(
        section,
        "body",
        header_sc,
        header_sc,
        show_line=False,
        front_header_text={},
        non_numbered_back_types=set(),
    )
    rule._apply_footer_body(section, page_sc, start=1)

    header_p = section.header._element.find(_w("p"))
    header_ppr = header_p.find(_w("pPr"))
    assert header_ppr.find(_w("jc")).get(_w("val")) == "right"
    assert header_ppr.find(_w("ind")).get(_w("left")) == "360"
    assert header_ppr.find(_w("ind")).get(_w("firstLine")) == "120"
    assert header_ppr.find(_w("spacing")).get(_w("lineRule")) == "exact"
    assert header_ppr.find(_w("spacing")).get(_w("line")) == "340"

    header_rpr = section.header._element.find(f".//{_w('r')}/{_w('rPr')}")
    assert header_rpr.find(_w("b")) is not None
    assert header_rpr.find(_w("i")) is not None
    header_rfonts = header_rpr.find(_w("rFonts"))
    assert header_rfonts.get(_w("eastAsia")) == header_sc.font_cn
    assert header_rfonts.get(_w("hint")) == "default"

    footer_p = section.footer._element.find(_w("p"))
    footer_ppr = footer_p.find(_w("pPr"))
    assert footer_ppr.find(_w("jc")).get(_w("val")) == "left"
    assert footer_ppr.find(_w("ind")).get(_w("left")) == "240"
    assert footer_ppr.find(_w("ind")).get(_w("firstLine")) == "60"
    assert footer_ppr.find(_w("spacing")).get(_w("lineRule")) == "auto"
    assert footer_ppr.find(_w("spacing")).get(_w("line")) == "480"

    footer_rpr = section.footer._element.find(f".//{_w('r')}/{_w('rPr')}")
    assert footer_rpr.find(_w("i")) is not None
    assert footer_rpr.find(_w("b")).get(_w("val")) == "0"

    pg_num_type = section._sectPr.find(_w("pgNumType"))
    assert pg_num_type.get(_w("fmt")) == "decimal"
    assert pg_num_type.get(_w("start")) == "1"


def test_header_footer_unknown_section_is_not_treated_as_cover() -> None:
    doc = Document()
    section = doc.sections[0]

    cfg = load_scene_from_data({})
    header_sc = cfg.styles["header_cn"]

    rule = HeaderFooterRule()
    rule._apply_header(
        section,
        "unknown",
        header_sc,
        header_sc,
        show_line=False,
        front_header_text={},
        non_numbered_back_types=set(),
    )

    header_text = "".join(p.text for p in section.header.paragraphs).strip()
    assert header_text != ""


def test_page_setup_formats_existing_header_footer_with_full_style():
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "已有页眉"
    section.footer.paragraphs[0].text = "已有页脚"

    cfg = load_scene_from_data({})
    header_sc = cfg.styles["header_cn"]
    header_sc.bold = True
    header_sc.italic = True
    header_sc.alignment = "right"
    header_sc.left_indent_chars = 24.0
    header_sc.left_indent_unit = "pt"
    header_sc.first_line_indent_chars = 12.0
    header_sc.first_line_indent_unit = "pt"
    header_sc.line_spacing_type = "exact"
    header_sc.line_spacing_pt = 18.0
    header_sc.space_before_pt = 4.0
    header_sc.space_after_pt = 6.0

    page_sc = cfg.styles["page_number"]
    page_sc.bold = True
    page_sc.italic = False
    page_sc.alignment = "left"
    page_sc.left_indent_chars = 12.0
    page_sc.left_indent_unit = "pt"
    page_sc.first_line_indent_chars = 6.0
    page_sc.first_line_indent_unit = "pt"
    page_sc.line_spacing_type = "double"
    page_sc.space_before_pt = 2.0
    page_sc.space_after_pt = 3.0

    PageSetupRule().apply(doc, cfg, ChangeTracker(), {})

    header_ppr = section.header._element.find(f"{_w('p')}/{_w('pPr')}")
    assert header_ppr.find(_w("jc")).get(_w("val")) == "right"
    assert header_ppr.find(_w("ind")).get(_w("left")) == "480"
    assert header_ppr.find(_w("ind")).get(_w("firstLine")) == "240"
    assert header_ppr.find(_w("spacing")).get(_w("lineRule")) == "exact"
    assert header_ppr.find(_w("spacing")).get(_w("line")) == "360"

    header_rpr = section.header._element.find(f".//{_w('r')}/{_w('rPr')}")
    assert header_rpr.find(_w("b")) is not None
    assert header_rpr.find(_w("i")) is not None

    footer_ppr = section.footer._element.find(f"{_w('p')}/{_w('pPr')}")
    assert footer_ppr.find(_w("jc")).get(_w("val")) == "left"
    assert footer_ppr.find(_w("spacing")).get(_w("line")) == "480"
    assert footer_ppr.find(_w("spacing")).get(_w("lineRule")) == "auto"

    footer_rpr = section.footer._element.find(f".//{_w('r')}/{_w('rPr')}")
    assert footer_rpr.find(_w("b")) is not None
    assert footer_rpr.find(_w("i")).get(_w("val")) == "0"

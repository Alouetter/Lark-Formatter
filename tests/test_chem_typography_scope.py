from __future__ import annotations

from types import SimpleNamespace

import pytest
from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.section_format import SectionFormatRule, _run_vert_align
from src.scene.manager import load_default_scene


class _DummyDocTree:
    def __init__(self, sections):
        self.sections = sections

    def get_section(self, section_type: str):
        for section in self.sections:
            if section.section_type == section_type:
                return section
        return None


def _body_runs_after_apply(
    section_type: str,
    *,
    section_enabled: bool,
    chem_scope_enabled: bool,
    target_indices: set[int] | None = None,
):
    doc = Document()
    title = "摘要" if section_type == "abstract_cn" else "Abstract"
    doc.add_paragraph(title)
    body = doc.add_paragraph("H2O and SO42-")
    section = SimpleNamespace(section_type=section_type, paragraph_range=range(0, 2), start_index=0)

    cfg = load_default_scene()
    cfg.chem_typography.enabled = True
    cfg.chem_typography.scopes["headings"] = False
    cfg.chem_typography.scopes[section_type] = chem_scope_enabled
    cfg.format_scope.sections[section_type] = section_enabled

    context = {
        "doc_tree": _DummyDocTree([section]),
        "format_scope": cfg.format_scope,
    }
    if target_indices is not None:
        context["target_paragraph_indices"] = set(target_indices)

    SectionFormatRule().apply(doc, cfg, ChangeTracker(), context)
    return [(run.text, _run_vert_align(run)) for run in body.runs]


def _has_super_or_sub(runs) -> bool:
    return any(style in {"superscript", "subscript"} for _, style in runs)


@pytest.mark.parametrize("section_type", ["abstract_cn", "abstract_en"])
def test_abstract_chem_restore_obeys_disabled_section_scope(section_type: str) -> None:
    runs = _body_runs_after_apply(
        section_type,
        section_enabled=False,
        chem_scope_enabled=True,
        target_indices={0, 1},
    )
    assert _has_super_or_sub(runs) is False


@pytest.mark.parametrize("section_type", ["abstract_cn", "abstract_en"])
def test_abstract_chem_restore_runs_when_section_and_scope_are_enabled(section_type: str) -> None:
    runs = _body_runs_after_apply(
        section_type,
        section_enabled=True,
        chem_scope_enabled=True,
        target_indices={0, 1},
    )
    assert _has_super_or_sub(runs) is True


@pytest.mark.parametrize("section_type", ["abstract_cn", "abstract_en"])
def test_abstract_chem_restore_stays_off_when_chem_scope_disabled(section_type: str) -> None:
    runs = _body_runs_after_apply(
        section_type,
        section_enabled=True,
        chem_scope_enabled=False,
        target_indices={0, 1},
    )
    assert _has_super_or_sub(runs) is False

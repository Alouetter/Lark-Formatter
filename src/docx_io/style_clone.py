"""Clone style/page setup settings from a reference DOCX into SceneConfig."""

from __future__ import annotations

from collections import Counter, defaultdict
import os
import re
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from src.scene.manager import _derive_heading_numbering_v2_from_legacy, load_default_scene
from src.scene.schema import HeadingLevelConfig, SceneConfig, StyleConfig
from src.utils.heading_numbering_template import default_heading_numbering_template
from src.utils.heading_numbering_v2 import legacy_levels_from_v2
from src.utils.indent import sync_style_config_indent_fields
from src.ui.font_sizes import display_font_size_with_name
from src.utils.toc_entry import (
    looks_like_numbered_toc_entry_with_page_suffix,
    looks_like_toc_entry_line,
)


STYLE_CANDIDATES: dict[str, list[str]] = {
    "normal": ["Normal", "正文"],
    "heading1": ["Heading 1", "标题 1", "章标题", "一级标题"],
    "heading2": ["Heading 2", "标题 2", "二级标题"],
    "heading3": ["Heading 3", "标题 3", "三级标题"],
    "heading4": ["Heading 4", "标题 4", "四级标题", "4", "附录章标题", "附录一级条标题"],
    "heading5": ["Heading 5", "标题 5", "五级标题", "5", "附录二级条标题"],
    "heading6": ["Heading 6", "标题 6", "六级标题", "6", "附录三级条标题"],
    "heading7": ["Heading 7", "标题 7", "七级标题", "7", "附录四级条标题"],
    "heading8": ["Heading 8", "标题 8", "八级标题", "8", "附录五级条标题"],
    "abstract_title_cn": ["中文摘要标题", "摘要标题", "Abstract Title"],
    "abstract_title_en": ["英文摘要标题", "Abstract Title"],
    "abstract_body": ["中文摘要正文", "摘要正文", "Normal"],
    "abstract_body_en": ["英文摘要正文", "Normal"],
    "toc_title": ["TOC Heading", "目录标题"],
    "toc_chapter": ["TOC 1", "目录 1"],
    "toc_level1": ["TOC 2", "目录 2"],
    "toc_level2": ["TOC 3", "目录 3"],
    "references_body": ["参考文献", "Bibliography", "Normal"],
    "acknowledgment_body": ["致谢正文", "Normal"],
    "appendix_body": ["附录正文", "Normal"],
    "resume_body": ["个人简历正文", "Normal"],
    "symbol_table_body": ["符号注释表", "Normal"],
    "code_block": ["Code", "代码块", "Normal"],
    "figure_caption": ["Caption", "图题注"],
    "table_caption": ["Caption", "表题注"],
    "header_cn": ["Header", "页眉"],
    "header_en": ["Header", "页眉"],
    "page_number": ["Footer", "页脚"],
}

ALIGNMENT_REVERSE = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

_RE_TABLE_CAPTION = re.compile(r"^\s*表\s*\d+(?:[.\-—–:：]\d+)*\s*\S")
_RE_FIGURE_CAPTION = re.compile(r"^\s*(?:图|Fig\.?)\s*\d+(?:[.\-—–:：]\d+)*\s*\S", re.IGNORECASE)
_PARAGRAPH_SAMPLE_RESET_INDENT_KEYS = {
    "abstract_title_cn",
    "abstract_title_en",
    "toc_title",
    "table_caption",
    "figure_caption",
}
_MIN_STYLE_SLOTS_FOR_CLONE = ("header_cn", "header_en", "page_number")
_HEADING_LEVEL_KEYS = [f"heading{i}" for i in range(1, 9)]
_TOC_TITLE_NORMS = {"目录", "目次", "contents", "tableofcontents"}
_RE_TOC_STYLE_LEVEL = re.compile(r"^(?:toc|目录)\s*(\d+)$", re.IGNORECASE)
_RE_TOC_LEVEL2_PREFIX = re.compile(
    r"^(?:\d+\.\d+\.\d+(?:\.\d+)*|[一二三四五六七八九十百零两]+、|[（(][一二三四五六七八九十百零两]+[）)])"
)
_RE_TOC_LEVEL1_PREFIX = re.compile(
    r"^(?:\d+\.\d+(?:\.\d+)*|第[一二三四五六七八九十百零两0-9]+节)"
)

_ARABIC_HEADING_RE = re.compile(
    r"^\s*(?P<num>\d+(?:[.\uFF0E\u3002\uFF61\uFE52]\d+)*)"
    r"(?P<sep>[ \t\u3000]+)"
    r"(?P<title>\S.+?)\s*$"
)
_CHAPTER_HEADING_RE = re.compile(
    r"^\s*(?P<num>第[一二三四五六七八九十百千万零〇两0-9]+章)"
    r"(?P<sep>[ \t\u3000]*)"
    r"(?P<title>\S.+?)\s*$"
)
_SECTION_HEADING_RE = re.compile(
    r"^\s*(?P<num>第[一二三四五六七八九十百千万零〇两0-9]+节)"
    r"(?P<sep>[ \t\u3000]*)"
    r"(?P<title>\S.+?)\s*$"
)
_CN_ORDINAL_HEADING_RE = re.compile(
    r"^\s*(?P<num>[一二三四五六七八九十百千万零〇两]+、)"
    r"(?P<sep>[ \t\u3000]*)"
    r"(?P<title>\S.+?)\s*$"
)
_CN_ORDINAL_PAREN_RE = re.compile(
    r"^\s*(?P<num>[（(][一二三四五六七八九十百千万零〇两]+[）)])"
    r"(?P<sep>[ \t\u3000]*)"
    r"(?P<title>\S.+?)\s*$"
)
_TOC_LEADER_TAIL_RE = re.compile(r"[\.．。…·]{3,}\s*\d+\s*$")
_RE_OOXML_LEVEL_REF = re.compile(r"%([1-9])")
_OOXML_CN_NUMFMTS = {
    "aiueo",
    "chineseCounting",
    "chineseCountingThousand",
    "chineseLegalSimplified",
    "ideographDigital",
    "ideographTraditional",
    "iroha",
}
_OOXML_CORE_STYLE_BY_NUMFMT = {
    "decimal": "arabic",
    "decimalZero": "arabic_pad2",
    "decimalEnclosedCircle": "circled",
    "ideographEnclosedCircle": "circled",
    "upperRoman": "roman_upper",
}


def _iter_style_chain(style):
    cur = style
    while cur is not None:
        yield cur
        cur = cur.base_style


def _first_non_none(style, getter):
    for st in _iter_style_chain(style):
        try:
            value = getter(st)
        except Exception:
            value = None
        if value is not None:
            return value
    return None


def _rfonts_value(style, attr: str) -> str | None:
    key = qn(attr)
    for st in _iter_style_chain(style):
        rpr = st.element.find(qn("w:rPr"))
        if rpr is None:
            continue
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            continue
        val = rfonts.get(key)
        if val:
            return val
    return None


def _size_pt(style) -> float | None:
    size_len = _first_non_none(style, lambda st: st.font.size)
    if size_len is not None:
        try:
            return float(size_len.pt)
        except Exception:
            pass

    for st in _iter_style_chain(style):
        rpr = st.element.find(qn("w:rPr"))
        if rpr is None:
            continue
        sz = rpr.find(qn("w:sz"))
        if sz is None:
            continue
        raw = sz.get(qn("w:val"))
        if not raw:
            continue
        try:
            return float(raw) / 2.0
        except Exception:
            continue
    return None


def _length_pt_or_none(value) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        try:
            return float(value.pt)
        except Exception:
            return None
    return None


def _style_direct_paragraph_length_pt(style, attr_name: str) -> float | None:
    try:
        pf = style.paragraph_format
    except Exception:
        return None
    if pf is None:
        return None
    try:
        value = getattr(pf, attr_name)
    except Exception:
        return None
    return _length_pt_or_none(value)


def _assign_style_indent_from_pt(sc: StyleConfig, attr_prefix: str, value_pt: float | None) -> None:
    if value_pt is None:
        return
    setattr(sc, f"{attr_prefix}_chars", round(max(0.0, value_pt), 2))
    setattr(sc, f"{attr_prefix}_unit", "pt")


def _read_xml_indent_info(style) -> dict:
    """Read character-based indent and hanging-indent from Word XML directly.

    Word stores indents in two forms:
      - EMU/twip values: ``w:ind/@w:firstLine``, ``w:ind/@w:left``
      - Character counts (×100): ``w:ind/@w:firstLineChars``, ``w:ind/@w:leftChars``
      - Hanging indent: ``w:ind/@w:hanging`` / ``w:ind/@w:hangingChars``

    python-docx only exposes EMU values, missing the chars attributes.
    Returns a dict with char-based and pt-based values when available.
    """
    result: dict = {
        "first_line_chars": None,
        "left_chars": None,
        "right_chars": None,
        "hanging_chars": None,
        "first_line_pt": None,
        "left_pt": None,
        "right_pt": None,
        "hanging_pt": None,
        "has_hanging": False,
    }
    try:
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            return result
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            return result
    except Exception:
        return result

    fl_chars = ind.get(qn("w:firstLineChars"))
    if fl_chars is not None:
        try:
            result["first_line_chars"] = int(fl_chars) / 100.0
        except (ValueError, TypeError):
            pass

    l_chars = ind.get(qn("w:leftChars"))
    if l_chars is not None:
        try:
            result["left_chars"] = int(l_chars) / 100.0
        except (ValueError, TypeError):
            pass

    r_chars = ind.get(qn("w:rightChars"))
    if r_chars is not None:
        try:
            result["right_chars"] = int(r_chars) / 100.0
        except (ValueError, TypeError):
            pass

    hanging_chars = ind.get(qn("w:hangingChars"))
    if hanging_chars is not None:
        try:
            result["hanging_chars"] = int(hanging_chars) / 100.0
            if result["hanging_chars"] > 0:
                result["has_hanging"] = True
        except (ValueError, TypeError):
            pass

    first_line_twip = ind.get(qn("w:firstLine"))
    if first_line_twip is not None:
        try:
            result["first_line_pt"] = int(first_line_twip) / 20.0
        except (ValueError, TypeError):
            pass

    left_twip = ind.get(qn("w:left"))
    if left_twip is not None:
        try:
            result["left_pt"] = int(left_twip) / 20.0
        except (ValueError, TypeError):
            pass

    right_twip = ind.get(qn("w:right"))
    if right_twip is not None:
        try:
            result["right_pt"] = int(right_twip) / 20.0
        except (ValueError, TypeError):
            pass

    hanging = ind.get(qn("w:hanging"))
    if hanging is not None:
        try:
            h_val = int(hanging)
            if h_val > 0:
                result["has_hanging"] = True
                result["hanging_pt"] = h_val / 20.0
        except (ValueError, TypeError):
            pass

    return result


def _container_element(obj):
    element = getattr(obj, "element", None)
    if element is not None:
        return element
    return getattr(obj, "_element", None)


def _read_xml_spacing_info(obj) -> dict:
    """Read explicit paragraph spacing from XML, including explicit zero-line settings.

    python-docx may return ``None`` for spacing when OOXML uses ``beforeLines=0`` /
    ``afterLines=0`` with auto-spacing disabled. For clone fidelity, that still means
    the source explicitly requested zero spacing and should override preset defaults.
    """
    result = {
        "before_pt": None,
        "after_pt": None,
        "before_explicit": False,
        "after_explicit": False,
    }
    try:
        element = _container_element(obj)
        if element is None:
            return result
        ppr = element.find(qn("w:pPr"))
        if ppr is None:
            return result
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            return result
    except Exception:
        return result

    def _read_side(name: str) -> tuple[float | None, bool]:
        twip_attr = spacing.get(qn(f"w:{name}"))
        if twip_attr is not None:
            try:
                return max(0.0, int(str(twip_attr).strip()) / 20.0), True
            except (TypeError, ValueError):
                pass

        lines_attr = spacing.get(qn(f"w:{name}Lines"))
        auto_attr = str(spacing.get(qn(f"w:{name}Autospacing"), "")).strip().lower()
        if lines_attr is not None:
            try:
                lines_value = int(str(lines_attr).strip())
            except (TypeError, ValueError):
                lines_value = None
            if lines_value == 0 and auto_attr in {"", "0", "false", "off"}:
                return 0.0, True
        return None, False

    before_pt, before_explicit = _read_side("before")
    after_pt, after_explicit = _read_side("after")
    result["before_pt"] = before_pt
    result["after_pt"] = after_pt
    result["before_explicit"] = before_explicit
    result["after_explicit"] = after_explicit
    return result


def _build_style_config(
    src_style,
    base: StyleConfig,
    *,
    inherit_indents_from_style_chain: bool = True,
) -> StyleConfig:
    sc = StyleConfig(**vars(base))

    font_cn = _rfonts_value(src_style, "w:eastAsia")
    if font_cn:
        sc.font_cn = font_cn

    font_en = (
        _rfonts_value(src_style, "w:ascii")
        or _rfonts_value(src_style, "w:hAnsi")
        or _first_non_none(src_style, lambda st: st.font.name)
    )
    if font_en:
        sc.font_en = font_en
        if not font_cn:
            sc.font_cn = sc.font_cn or font_en

    size_pt = _size_pt(src_style)
    if size_pt and size_pt > 0:
        sc.size_pt = round(size_pt, 2)
        sc.size_display = display_font_size_with_name(sc.size_pt)

    bold = _first_non_none(src_style, lambda st: st.font.bold)
    if bold is not None:
        sc.bold = bool(bold)
    italic = _first_non_none(src_style, lambda st: st.font.italic)
    if italic is not None:
        sc.italic = bool(italic)

    align = _first_non_none(src_style, lambda st: st.paragraph_format.alignment)
    if align in ALIGNMENT_REVERSE:
        sc.alignment = ALIGNMENT_REVERSE[align]

    xml_spacing = _read_xml_spacing_info(src_style)
    before_pt = xml_spacing["before_pt"]
    if not xml_spacing["before_explicit"]:
        before_pt = _length_pt_or_none(_first_non_none(src_style, lambda st: st.paragraph_format.space_before))
    if before_pt is not None:
        sc.space_before_pt = max(0.0, round(before_pt, 2))
    after_pt = xml_spacing["after_pt"]
    if not xml_spacing["after_explicit"]:
        after_pt = _length_pt_or_none(_first_non_none(src_style, lambda st: st.paragraph_format.space_after))
    if after_pt is not None:
        sc.space_after_pt = max(0.0, round(after_pt, 2))

    # --- Indentation ---
    # First, try reading character-based indent from XML (preferred).
    # Word stores both EMU/twip values and character counts; python-docx
    # only exposes EMU, so we read character counts from XML directly.
    xml_indent = _read_xml_indent_info(src_style)

    if xml_indent["has_hanging"]:
        sc.first_line_indent_chars = 0.0
        sc.first_line_indent_unit = "chars"
        if xml_indent["left_chars"] is not None:
            sc.left_indent_chars = round(max(0.0, xml_indent["left_chars"]), 2)
            sc.left_indent_unit = "chars"
        elif xml_indent["left_pt"] is not None:
            _assign_style_indent_from_pt(sc, "left_indent", xml_indent["left_pt"])
        if xml_indent["right_chars"] is not None:
            sc.right_indent_chars = round(max(0.0, xml_indent["right_chars"]), 2)
            sc.right_indent_unit = "chars"
        elif xml_indent["right_pt"] is not None:
            _assign_style_indent_from_pt(sc, "right_indent", xml_indent["right_pt"])

        if xml_indent["hanging_chars"] is not None:
            sc.hanging_indent_chars = round(max(0.0, xml_indent["hanging_chars"]), 2)
            sc.hanging_indent_unit = "chars"
        elif xml_indent["hanging_pt"] is not None:
            _assign_style_indent_from_pt(sc, "hanging_indent", xml_indent["hanging_pt"])

        if getattr(sc, "left_indent_chars", 0.0) <= 0 and getattr(sc, "hanging_indent_chars", 0.0) > 0:
            sc.left_indent_chars = sc.hanging_indent_chars
            sc.left_indent_unit = sc.hanging_indent_unit
    elif xml_indent["first_line_chars"] is not None or xml_indent["left_chars"] is not None:
        # Use character-based values from XML when available.
        if xml_indent["first_line_chars"] is not None:
            sc.first_line_indent_chars = round(max(0.0, xml_indent["first_line_chars"]), 2)
            sc.first_line_indent_unit = "chars"
        if xml_indent["left_chars"] is not None:
            sc.left_indent_chars = round(max(0.0, xml_indent["left_chars"]), 2)
            sc.left_indent_unit = "chars"
        elif xml_indent["left_pt"] is not None:
            _assign_style_indent_from_pt(sc, "left_indent", xml_indent["left_pt"])
        if xml_indent["right_chars"] is not None:
            sc.right_indent_chars = round(max(0.0, xml_indent["right_chars"]), 2)
            sc.right_indent_unit = "chars"
        elif xml_indent["right_pt"] is not None:
            _assign_style_indent_from_pt(sc, "right_indent", xml_indent["right_pt"])
        if xml_indent["right_chars"] is not None:
            sc.right_indent_chars = round(max(0.0, xml_indent["right_chars"]), 2)
            sc.right_indent_unit = "chars"
    else:
        # Fall back to EMU/pt-based values from python-docx.
        if inherit_indents_from_style_chain:
            first_indent_pt = _length_pt_or_none(
                _first_non_none(src_style, lambda st: st.paragraph_format.first_line_indent)
            )
        else:
            first_indent_pt = _style_direct_paragraph_length_pt(src_style, "first_line_indent")
        _assign_style_indent_from_pt(sc, "first_line_indent", first_indent_pt)

        if inherit_indents_from_style_chain:
            left_indent_pt = _length_pt_or_none(
                _first_non_none(src_style, lambda st: st.paragraph_format.left_indent)
            )
        else:
            left_indent_pt = _style_direct_paragraph_length_pt(src_style, "left_indent")
        _assign_style_indent_from_pt(sc, "left_indent", left_indent_pt)
        if inherit_indents_from_style_chain:
            right_indent_pt = _length_pt_or_none(
                _first_non_none(src_style, lambda st: st.paragraph_format.right_indent)
            )
        else:
            right_indent_pt = _style_direct_paragraph_length_pt(src_style, "right_indent")
        _assign_style_indent_from_pt(sc, "right_indent", right_indent_pt)

    # Walk the style chain once to get a *coupled* (rule, value) pair from the
    # same hierarchy level.  Using two independent ``_first_non_none`` calls can
    # return a rule from one level and a value from a different level, leading to
    # mismatches such as EXACTLY + 1.12 (multiple float from the base style).
    line_rule = None
    line_spacing = None
    for st in _iter_style_chain(src_style):
        try:
            _rule = st.paragraph_format.line_spacing_rule
            _val = st.paragraph_format.line_spacing
        except Exception:
            continue
        if _rule is not None or _val is not None:
            line_rule = _rule
            line_spacing = _val
            break
    line_spacing_pt = _length_pt_or_none(line_spacing)
    if line_rule in {
        WD_LINE_SPACING.EXACTLY,
        WD_LINE_SPACING.AT_LEAST,
    }:
        sc.line_spacing_type = "exact"
        if line_spacing_pt is not None and line_spacing_pt > 0:
            sc.line_spacing_pt = round(line_spacing_pt, 2)
        elif hasattr(line_spacing, "pt"):
            try:
                sc.line_spacing_pt = round(float(line_spacing.pt), 2)
            except Exception:
                pass
    elif line_rule in {
        WD_LINE_SPACING.SINGLE,
        WD_LINE_SPACING.ONE_POINT_FIVE,
        WD_LINE_SPACING.DOUBLE,
        WD_LINE_SPACING.MULTIPLE,
    }:
        sc.line_spacing_type = "multiple"
        if isinstance(line_spacing, (int, float)) and not hasattr(line_spacing, "pt"):
            sc.line_spacing_pt = round(max(0.1, float(line_spacing)), 2)
        elif line_rule == WD_LINE_SPACING.SINGLE:
            sc.line_spacing_pt = 1.0
        elif line_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
            sc.line_spacing_pt = 1.5
        elif line_rule == WD_LINE_SPACING.DOUBLE:
            sc.line_spacing_pt = 2.0
    elif line_spacing_pt is not None and line_spacing_pt > 0:
        sc.line_spacing_type = "exact"
        sc.line_spacing_pt = round(line_spacing_pt, 2)
    elif isinstance(line_spacing, (int, float)) and not hasattr(line_spacing, "pt") and float(line_spacing) > 0:
        sc.line_spacing_type = "multiple"
        sc.line_spacing_pt = round(max(0.1, float(line_spacing)), 2)

    sync_style_config_indent_fields(sc)
    return sc


def _norm_no_space(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _run_rfonts_value(run, attr: str) -> str | None:
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn(attr))


def _build_style_config_from_paragraph(para, base: StyleConfig) -> StyleConfig:
    if para.style is not None:
        sc = _build_style_config(para.style, base)
    else:
        sc = StyleConfig(**vars(base))

    pf = para.paragraph_format
    if pf is not None:
        if pf.alignment in ALIGNMENT_REVERSE:
            sc.alignment = ALIGNMENT_REVERSE[pf.alignment]

        xml_spacing = _read_xml_spacing_info(para)
        before_pt = xml_spacing["before_pt"]
        if not xml_spacing["before_explicit"]:
            before_pt = _length_pt_or_none(pf.space_before)
        if before_pt is not None:
            sc.space_before_pt = max(0.0, round(before_pt, 2))
        after_pt = xml_spacing["after_pt"]
        if not xml_spacing["after_explicit"]:
            after_pt = _length_pt_or_none(pf.space_after)
        if after_pt is not None:
            sc.space_after_pt = max(0.0, round(after_pt, 2))

        left_indent_pt = _length_pt_or_none(pf.left_indent)
        _assign_style_indent_from_pt(sc, "left_indent", left_indent_pt)
        right_indent_pt = _length_pt_or_none(pf.right_indent)
        _assign_style_indent_from_pt(sc, "right_indent", right_indent_pt)
        first_indent_pt = _length_pt_or_none(pf.first_line_indent)
        if first_indent_pt is not None and first_indent_pt < 0:
            _assign_style_indent_from_pt(sc, "hanging_indent", abs(first_indent_pt))
            sc.first_line_indent_chars = 0.0
            sc.first_line_indent_unit = "chars"
            if left_indent_pt is None or left_indent_pt <= 0:
                sc.left_indent_chars = sc.hanging_indent_chars
                sc.left_indent_unit = sc.hanging_indent_unit
        else:
            if first_indent_pt is not None:
                sc.hanging_indent_chars = 0.0
                sc.hanging_indent_unit = "chars"
            _assign_style_indent_from_pt(sc, "first_line_indent", first_indent_pt)

        line_rule = pf.line_spacing_rule
        line_spacing = pf.line_spacing
        line_spacing_pt = _length_pt_or_none(line_spacing)
        if line_rule in {
            WD_LINE_SPACING.EXACTLY,
            WD_LINE_SPACING.AT_LEAST,
        }:
            sc.line_spacing_type = "exact"
            if line_spacing_pt is not None and line_spacing_pt > 0:
                sc.line_spacing_pt = round(line_spacing_pt, 2)
            elif hasattr(line_spacing, "pt"):
                try:
                    sc.line_spacing_pt = round(float(line_spacing.pt), 2)
                except Exception:
                    pass
        elif line_rule in {
            WD_LINE_SPACING.SINGLE,
            WD_LINE_SPACING.ONE_POINT_FIVE,
            WD_LINE_SPACING.DOUBLE,
            WD_LINE_SPACING.MULTIPLE,
        }:
            sc.line_spacing_type = "multiple"
            if isinstance(line_spacing, (int, float)) and not hasattr(line_spacing, "pt"):
                sc.line_spacing_pt = round(max(0.1, float(line_spacing)), 2)
            elif line_rule == WD_LINE_SPACING.SINGLE:
                sc.line_spacing_pt = 1.0
            elif line_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                sc.line_spacing_pt = 1.5
            elif line_rule == WD_LINE_SPACING.DOUBLE:
                sc.line_spacing_pt = 2.0
        elif line_spacing_pt is not None and line_spacing_pt > 0:
            sc.line_spacing_type = "exact"
            sc.line_spacing_pt = round(line_spacing_pt, 2)

    runs = [run for run in para.runs if (run.text or "").strip()]
    if runs:
        run0 = runs[0]
        font_cn = _run_rfonts_value(run0, "w:eastAsia")
        font_en = (
            _run_rfonts_value(run0, "w:ascii")
            or _run_rfonts_value(run0, "w:hAnsi")
            or run0.font.name
        )
        if font_cn:
            sc.font_cn = font_cn
        if font_en:
            sc.font_en = font_en
            if not font_cn:
                sc.font_cn = sc.font_cn or font_en

        size_len = run0.font.size
        if size_len is not None:
            try:
                size_pt = float(size_len.pt)
            except Exception:
                size_pt = None
            if size_pt and size_pt > 0:
                sc.size_pt = round(size_pt, 2)
                sc.size_display = display_font_size_with_name(sc.size_pt)

        bold = run0.bold if run0.bold is not None else run0.font.bold
        if bold is not None:
            sc.bold = bool(bold)
        italic = run0.italic if run0.italic is not None else run0.font.italic
        if italic is not None:
            sc.italic = bool(italic)

    sync_style_config_indent_fields(sc)
    return sc


def _collect_paragraph_samples(doc: Document) -> dict[str, object]:
    samples: dict[str, object] = {}
    table_caption_para = None
    figure_caption_para = None
    toc_title_idx = None

    for idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue

        normalized = _norm_no_space(text)
        normalized_lower = normalized.lower()

        if "abstract_title_cn" not in samples and normalized == "摘要":
            samples["abstract_title_cn"] = para
        if "abstract_title_en" not in samples and normalized_lower == "abstract":
            samples["abstract_title_en"] = para
        if "toc_title" not in samples and normalized_lower in _TOC_TITLE_NORMS:
            samples["toc_title"] = para
            toc_title_idx = idx

        if table_caption_para is None and _RE_TABLE_CAPTION.match(text):
            table_caption_para = para
        if figure_caption_para is None and _RE_FIGURE_CAPTION.match(text):
            figure_caption_para = para

    if table_caption_para is not None:
        samples["table_caption"] = table_caption_para
    if figure_caption_para is not None:
        samples["figure_caption"] = figure_caption_para
    elif table_caption_para is not None:
        # Many规范文档 only include table-caption examples. Reuse for figure caption.
        samples["figure_caption"] = table_caption_para

    if toc_title_idx is not None:
        found_entry = False
        blank_streak = 0
        for para in doc.paragraphs[toc_title_idx + 1:]:
            text = (para.text or "").strip()
            style_name = ((para.style.name if para.style else "") or "").strip()
            style_id = (
                (getattr(para.style, "style_id", "") if para.style else "") or ""
            ).strip()
            if not text:
                if found_entry:
                    blank_streak += 1
                    if blank_streak >= 3:
                        break
                continue

            blank_streak = 0
            level_style_key = _infer_toc_style_key_from_para(text, style_name, style_id)
            if level_style_key:
                if level_style_key not in samples:
                    samples[level_style_key] = para
                found_entry = True
                continue

            if found_entry:
                break

    return samples


def _ensure_clone_style_slots(config: SceneConfig) -> list[str]:
    added: list[str] = []
    for key in (
        "abstract_title_cn",
        "abstract_title_en",
        "abstract_body",
        "abstract_body_en",
        "toc_title",
        "toc_chapter",
        "toc_level1",
        "toc_level2",
        *_MIN_STYLE_SLOTS_FOR_CLONE,
    ):
        if key in config.styles:
            continue
        config.styles[key] = StyleConfig()
        added.append(key)
    return added


def _clone_style_seed_map(config: SceneConfig) -> dict[str, StyleConfig]:
    try:
        seed_config = load_default_scene()
        source_styles = getattr(seed_config, "styles", {}) or {}
    except Exception:
        source_styles = getattr(config, "styles", {}) or {}

    seeds: dict[str, StyleConfig] = {}
    for key, style_cfg in source_styles.items():
        if isinstance(style_cfg, StyleConfig):
            seeds[key] = StyleConfig(**vars(style_cfg))
    return seeds


def _fresh_style_seed(
    style_key: str,
    seed_styles: dict[str, StyleConfig],
    fallback: StyleConfig | None = None,
) -> StyleConfig:
    seed = seed_styles.get(style_key)
    if isinstance(seed, StyleConfig):
        return StyleConfig(**vars(seed))
    if isinstance(fallback, StyleConfig):
        return StyleConfig(**vars(fallback))
    return StyleConfig()


def _infer_toc_style_key_from_para(text: str, style_name: str, style_id: str) -> str | None:
    style_levels = []
    for s in (style_name, style_id):
        m = _RE_TOC_STYLE_LEVEL.match((s or "").strip())
        if not m:
            continue
        try:
            style_levels.append(int(m.group(1)))
        except Exception:
            continue
    if style_levels:
        level = min(style_levels)
        if level <= 1:
            return "toc_chapter"
        if level == 2:
            return "toc_level1"
        return "toc_level2"

    if not (
        looks_like_toc_entry_line(text)
        or looks_like_numbered_toc_entry_with_page_suffix(text)
    ):
        return None

    raw = (text or "").strip()
    if _RE_TOC_LEVEL2_PREFIX.match(raw):
        return "toc_level2"
    if _RE_TOC_LEVEL1_PREFIX.match(raw):
        return "toc_level1"
    return "toc_chapter"


def _normalize_separator(raw: str) -> str:
    if not raw:
        return ""
    if "\u3000" in raw:
        return "\u3000"
    if "\t" in raw:
        return "\t"
    if " " in raw:
        return " "
    return raw


def _looks_like_toc_tail(title: str) -> bool:
    tail = (title or "").strip()
    if not tail:
        return False
    return bool(_TOC_LEADER_TAIL_RE.search(tail))


def _infer_heading_pattern(text: str):
    text = (text or "").strip()
    if not text:
        return None
    if looks_like_toc_entry_line(text) or looks_like_numbered_toc_entry_with_page_suffix(text):
        return None

    m = _CHAPTER_HEADING_RE.match(text)
    if m:
        title = m.group("title")
        if _looks_like_toc_tail(title):
            return None
        return ("heading1", "chinese_chapter", "第{current}章", _normalize_separator(m.group("sep")), title)

    m = _SECTION_HEADING_RE.match(text)
    if m:
        title = m.group("title")
        if _looks_like_toc_tail(title):
            return None
        return ("heading2", "chinese_section", "第{current}节", _normalize_separator(m.group("sep")), title)

    m = _CN_ORDINAL_HEADING_RE.match(text)
    if m:
        title = m.group("title")
        if _looks_like_toc_tail(title):
            return None
        return ("heading3", "chinese_ordinal", "{current}、", _normalize_separator(m.group("sep")), title)

    m = _CN_ORDINAL_PAREN_RE.match(text)
    if m:
        num = m.group("num")
        template = "({current})" if num.startswith("(") else "（{current}）"
        title = m.group("title")
        if _looks_like_toc_tail(title):
            return None
        return ("heading4", "chinese_ordinal_paren", template, _normalize_separator(m.group("sep")), title)

    m = _ARABIC_HEADING_RE.match(text)
    if not m:
        return None

    number_text = m.group("num")
    parts = [p for p in re.split(r"[.\uFF0E\u3002\uFF61\uFE52]", number_text) if p]
    if not parts:
        return None
    if len(parts[0]) >= 4:
        try:
            if int(parts[0]) >= 1900:
                return None
        except ValueError:
            pass

    level_idx = min(len(parts), 8)
    level_key = f"heading{level_idx}"
    format_name = "arabic" if level_idx == 1 else "arabic_dotted"
    template = default_heading_numbering_template(level_key)
    title = m.group("title")
    if _looks_like_toc_tail(title):
        return None
    if title.startswith(("年", "月", "日")):
        return None
    return (level_key, format_name, template, _normalize_separator(m.group("sep")), title)


def _guess_heading_indent_chars(para, style_cfg: StyleConfig | None) -> float | None:
    left_pt = _length_pt_or_none(para.paragraph_format.left_indent)
    if left_pt is None or left_pt <= 0:
        return None
    size_pt = 0.0
    if style_cfg is not None:
        try:
            size_pt = float(getattr(style_cfg, "size_pt", 0) or 0)
        except Exception:
            size_pt = 0.0
    if size_pt <= 0:
        size_pt = 12.0
    return round(max(0.0, left_pt) / size_pt, 2)


def _heading_level_index(level_key: str) -> int:
    raw = str(level_key or "").strip().lower()
    if raw.startswith("heading"):
        try:
            return int(raw[7:])
        except ValueError:
            return 0
    return 0


def _parse_numpr(num_pr, *, fallback_ilvl: int | None = None) -> tuple[str, int] | None:
    if num_pr is None:
        return None

    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        return None

    num_id = str(num_id_el.get(qn("w:val"), "")).strip()
    if not num_id:
        return None

    ilvl = fallback_ilvl
    ilvl_el = num_pr.find(qn("w:ilvl"))
    if ilvl_el is not None:
        ilvl_raw = str(ilvl_el.get(qn("w:val"), "")).strip()
        if ilvl_raw:
            try:
                ilvl = int(ilvl_raw)
            except ValueError:
                return None

    if ilvl is None or ilvl < 0:
        return None
    return num_id, ilvl


def _container_numpr(container, *, fallback_ilvl: int | None = None) -> tuple[str, int] | None:
    element = _container_element(container)
    if element is None:
        return None

    try:
        ppr = element.find(qn("w:pPr"))
    except Exception:
        ppr = None
    if ppr is None:
        return None

    try:
        num_pr = ppr.find(qn("w:numPr"))
    except Exception:
        num_pr = None
    return _parse_numpr(num_pr, fallback_ilvl=fallback_ilvl)


def _style_numpr(style) -> tuple[str, int] | None:
    for st in _iter_style_chain(style):
        value = _container_numpr(st)
        if value is not None:
            return value
    return None


def _paragraph_numpr(para, *, fallback_ilvl: int | None = None) -> tuple[str, int] | None:
    return _container_numpr(para, fallback_ilvl=fallback_ilvl)


def _numbering_root(doc: Document):
    numbering_part = getattr(doc.part, "numbering_part", None)
    if numbering_part is None:
        return None
    root = getattr(numbering_part, "element", None)
    if root is not None:
        return root
    return getattr(numbering_part, "_element", None)


def _build_numbering_maps(doc: Document) -> tuple[dict[str, dict[str, object]], dict[str, object]]:
    root = _numbering_root(doc)
    if root is None:
        return {}, {}

    abstract_map: dict[str, object] = {}
    for abs_num in root.findall(qn("w:abstractNum")):
        abs_id = str(abs_num.get(qn("w:abstractNumId"), "")).strip()
        if abs_id:
            abstract_map[abs_id] = abs_num

    num_map: dict[str, dict[str, object]] = {}
    for num in root.findall(qn("w:num")):
        num_id = str(num.get(qn("w:numId"), "")).strip()
        if not num_id:
            continue

        abs_ref = num.find(qn("w:abstractNumId"))
        abstract_id = str(abs_ref.get(qn("w:val"), "")).strip() if abs_ref is not None else ""
        overrides: dict[int, dict[str, object]] = {}

        for override in num.findall(qn("w:lvlOverride")):
            ilvl_raw = str(override.get(qn("w:ilvl"), "")).strip()
            if not ilvl_raw:
                continue
            try:
                ilvl = int(ilvl_raw)
            except ValueError:
                continue

            start_override_el = override.find(qn("w:startOverride"))
            start_override = None
            if start_override_el is not None:
                try:
                    start_override = int(str(start_override_el.get(qn("w:val"), "")).strip())
                except ValueError:
                    start_override = None

            overrides[ilvl] = {
                "lvl": override.find(qn("w:lvl")),
                "start_override": start_override,
            }

        num_map[num_id] = {
            "abstract_id": abstract_id,
            "overrides": overrides,
        }

    return num_map, abstract_map


def _find_numbering_lvl(num_map: dict[str, dict[str, object]], abstract_map: dict[str, object], *, num_id: str, ilvl: int):
    num_info = num_map.get(str(num_id or "").strip())
    if not num_info:
        return None, None

    overrides = num_info.get("overrides", {}) or {}
    override_info = overrides.get(ilvl, {}) if isinstance(overrides, dict) else {}
    override_lvl = override_info.get("lvl") if isinstance(override_info, dict) else None
    abstract_id = str(num_info.get("abstract_id", "")).strip()
    abstract_lvl = None
    abstract_num = abstract_map.get(abstract_id)
    if abstract_num is not None:
        for candidate in abstract_num.findall(qn("w:lvl")):
            if str(candidate.get(qn("w:ilvl"), "")).strip() == str(ilvl):
                abstract_lvl = candidate
                break

    return (override_lvl or abstract_lvl), abstract_lvl


def _numbering_child(*parents, tag: str):
    child_tag = qn(f"w:{tag}")
    for parent in parents:
        if parent is None:
            continue
        try:
            child = parent.find(child_tag)
        except Exception:
            child = None
        if child is not None:
            return child
    return None


def _numbering_ppr_child(*parents, tag: str):
    child_tag = qn(f"w:{tag}")
    for parent in parents:
        if parent is None:
            continue
        try:
            ppr = parent.find(qn("w:pPr"))
        except Exception:
            ppr = None
        if ppr is None:
            continue
        try:
            child = ppr.find(child_tag)
        except Exception:
            child = None
        if child is not None:
            return child
    return None


_OOXML_IND_ATTR_NAMES = (
    "left",
    "leftChars",
    "firstLine",
    "firstLineChars",
    "hanging",
    "hangingChars",
    "right",
    "rightChars",
)


def _extract_ind_attrs(ind_el) -> dict[str, str]:
    attrs: dict[str, str] = {}
    if ind_el is None:
        return attrs
    for attr_name in _OOXML_IND_ATTR_NAMES:
        raw = ind_el.get(qn(f"w:{attr_name}"))
        if raw is None:
            continue
        value = str(raw).strip()
        if value:
            attrs[attr_name] = value
    return attrs


def _numbering_lvl_ind_attrs(resolved_lvl, abstract_lvl, src_style) -> dict[str, str]:
    ind_attrs = _extract_ind_attrs(_numbering_ppr_child(resolved_lvl, abstract_lvl, tag="ind"))
    if ind_attrs:
        return ind_attrs

    for style in _iter_style_chain(src_style):
        try:
            ppr = style.element.find(qn("w:pPr"))
        except Exception:
            ppr = None
        if ppr is None:
            continue
        ind = ppr.find(qn("w:ind"))
        ind_attrs = _extract_ind_attrs(ind)
        if ind_attrs:
            return ind_attrs
    return {}


def _numbering_separator_from_lvl_text(raw: str) -> tuple[str, str | None]:
    text = str(raw or "")
    if not text:
        return "", None

    match = re.search(r"([ \t\u3000]+)$", text)
    if match and _RE_OOXML_LEVEL_REF.search(text[: match.start()]):
        return text[: match.start()], _normalize_separator(match.group(1))
    return text, None


def _numbering_suff_value(*lvls) -> str | None:
    suff = _numbering_child(*lvls, tag="suff")
    if suff is None:
        return None

    value = str(suff.get(qn("w:val"), "")).strip().lower()
    if value in {"tab", "space", "nothing"}:
        return value
    return None


def _numbering_separator_from_suff_value(value: str | None) -> str | None:
    if value is None:
        return None
    if value == "tab":
        return "\t"
    if value == "space":
        return " "
    if value == "nothing":
        return ""
    return None


def _ooxml_lvl_text_to_template(level_key: str, lvl_text: str) -> str:
    level_idx = _heading_level_index(level_key) or 1

    def _replace(match: re.Match[str]) -> str:
        try:
            ref_idx = int(match.group(1))
        except ValueError:
            return match.group(0)
        if ref_idx == level_idx:
            return "{current}"
        return f"{{level{ref_idx}}}"

    return _RE_OOXML_LEVEL_REF.sub(_replace, str(lvl_text or ""))


def _ooxml_numfmt_to_format(level_key: str, *, num_fmt: str, template: str) -> str:
    normalized_num_fmt = str(num_fmt or "").strip()
    level_idx = _heading_level_index(level_key)
    core_style = _OOXML_CORE_STYLE_BY_NUMFMT.get(normalized_num_fmt)
    if core_style is None and normalized_num_fmt in _OOXML_CN_NUMFMTS:
        core_style = "cn_lower"
    if core_style is None:
        core_style = "arabic"

    if core_style == "cn_lower":
        if level_idx == 1 and template == "第{current}章":
            return "chinese_chapter"
        if level_idx == 2 and template == "第{current}节":
            return "chinese_section"
        if template == "{current}、":
            return "chinese_ordinal"
        if template in {"（{current}）", "({current})"}:
            return "chinese_ordinal_paren"

    if (
        core_style == "arabic"
        and level_idx > 1
        and any(f"{{level{idx}}}" in template for idx in range(1, level_idx))
    ):
        return "arabic_dotted"

    return core_style


def _restart_on_from_lvl(level_key: str, *lvls) -> str | None:
    level_idx = _heading_level_index(level_key)
    if level_idx <= 1:
        return None

    restart_el = _numbering_child(*lvls, tag="lvlRestart")
    if restart_el is None:
        return f"heading{level_idx - 1}"

    raw = str(restart_el.get(qn("w:val"), "")).strip()
    try:
        restart_idx = int(raw)
    except ValueError:
        return f"heading{level_idx - 1}"

    if restart_idx <= 0:
        return ""
    if restart_idx >= level_idx:
        return f"heading{level_idx - 1}"
    return f"heading{restart_idx}"


def _extract_heading_numbering_level_from_numpr(
    existing_levels: dict[str, HeadingLevelConfig],
    *,
    level_key: str,
    src_style,
    num_map: dict[str, dict[str, object]],
    abstract_map: dict[str, object],
    num_id: str,
    ilvl: int,
) -> tuple[HeadingLevelConfig, dict[str, object]] | None:
    resolved_lvl, abstract_lvl = _find_numbering_lvl(num_map, abstract_map, num_id=num_id, ilvl=ilvl)
    if resolved_lvl is None:
        return None

    lvl_text_el = _numbering_child(resolved_lvl, abstract_lvl, tag="lvlText")
    num_fmt_el = _numbering_child(resolved_lvl, abstract_lvl, tag="numFmt")
    if lvl_text_el is None or num_fmt_el is None:
        return None

    raw_lvl_text = str(lvl_text_el.get(qn("w:val"), "") or "")
    if not raw_lvl_text:
        return None

    lvl_text_body, separator = _numbering_separator_from_lvl_text(raw_lvl_text)
    suff_value = _numbering_suff_value(resolved_lvl, abstract_lvl)
    separator_mode = "inline"
    if separator is None:
        separator = _numbering_separator_from_suff_value(suff_value)
        if suff_value is not None:
            separator_mode = "suff"
        else:
            # OOXML 17.9.28: missing <w:suff> defaults to tab.
            separator = "\t"
            suff_value = "tab"
            separator_mode = "suff"

    template = _ooxml_lvl_text_to_template(level_key, lvl_text_body)
    if not template:
        return None

    base = existing_levels.get(level_key, HeadingLevelConfig())
    level_cfg = HeadingLevelConfig(**vars(base))
    level_cfg.format = _ooxml_numfmt_to_format(
        level_key,
        num_fmt=str(num_fmt_el.get(qn("w:val"), "")).strip(),
        template=template,
    )
    level_cfg.template = template
    level_cfg.separator = str(separator or "")
    level_cfg.custom_separator = None

    start_el = _numbering_child(resolved_lvl, abstract_lvl, tag="start")
    start_override = None
    num_info = num_map.get(str(num_id or "").strip(), {}) or {}
    overrides = num_info.get("overrides", {}) or {}
    override_info = overrides.get(ilvl, {}) if isinstance(overrides, dict) else {}
    if isinstance(override_info, dict):
        start_override = override_info.get("start_override")

    start_at = 1
    if start_override is not None:
        try:
            start_at = max(1, int(start_override))
        except (TypeError, ValueError):
            start_at = 1
    elif start_el is not None:
        try:
            start_at = max(1, int(str(start_el.get(qn("w:val"), "")).strip()))
        except ValueError:
            start_at = 1

    is_lgl = _numbering_child(resolved_lvl, abstract_lvl, tag="isLgl") is not None
    binding_override = {
        "start_at": start_at,
        "restart_on": _restart_on_from_lvl(level_key, resolved_lvl, abstract_lvl),
        "reference_core_style": "arabic" if is_lgl else None,
        "ooxml_separator_mode": separator_mode,
        "ooxml_suff": suff_value,
        "ooxml_lvl_ind": _numbering_lvl_ind_attrs(resolved_lvl, abstract_lvl, src_style),
    }
    return level_cfg, binding_override


def _heading_numbering_candidate_signature(
    level_cfg: HeadingLevelConfig,
    binding_override: dict[str, object],
) -> tuple[object, ...]:
    return (
        str(getattr(level_cfg, "format", "") or ""),
        str(getattr(level_cfg, "template", "") or ""),
        str(getattr(level_cfg, "effective_separator", "") or ""),
        str(binding_override.get("restart_on") or ""),
        str(binding_override.get("reference_core_style") or ""),
        str(binding_override.get("ooxml_separator_mode") or ""),
        str(binding_override.get("ooxml_suff") or ""),
        tuple(sorted(dict(binding_override.get("ooxml_lvl_ind", {}) or {}).items())),
    )


def _select_heading_numbering_candidate(candidates: list[dict[str, object]]) -> tuple[HeadingLevelConfig, dict[str, object]] | None:
    if not candidates:
        return None

    signature_counter = Counter(
        _heading_numbering_candidate_signature(c["level_cfg"], c["binding_override"])
        for c in candidates
        if isinstance(c.get("level_cfg"), HeadingLevelConfig) and isinstance(c.get("binding_override"), dict)
    )
    if not signature_counter:
        return None

    best_signature, _ = signature_counter.most_common(1)[0]
    chosen_matches = [
        c for c in candidates
        if isinstance(c.get("level_cfg"), HeadingLevelConfig)
        and isinstance(c.get("binding_override"), dict)
        and _heading_numbering_candidate_signature(c["level_cfg"], c["binding_override"]) == best_signature
    ]
    if not chosen_matches:
        return None

    first = chosen_matches[0]
    level_cfg = HeadingLevelConfig(**vars(first["level_cfg"]))
    binding_override = dict(first["binding_override"])

    start_values = []
    for match in chosen_matches:
        try:
            start_values.append(max(1, int(match["binding_override"].get("start_at", 1) or 1)))
        except (TypeError, ValueError, AttributeError):
            continue
    if start_values:
        binding_override["start_at"] = min(start_values)

    return level_cfg, binding_override


def _clone_heading_numbering_from_styles(
    config: SceneConfig,
    doc: Document,
) -> tuple[dict[str, HeadingLevelConfig], dict[str, dict[str, object]]]:
    num_map, abstract_map = _build_numbering_maps(doc)
    if not num_map or not abstract_map:
        return {}, {}

    heading_style_ids = set()
    for level_key in _HEADING_LEVEL_KEYS:
        src_style = _find_style(doc, level_key)
        if src_style is None:
            continue
        style_id = str(getattr(src_style, "style_id", "") or "").strip()
        if style_id:
            heading_style_ids.add(style_id)

    if heading_style_ids:
        has_heading_usage = any(
            str(getattr(para.style, "style_id", "") or "").strip() in heading_style_ids
            for para in doc.paragraphs
            if getattr(para, "style", None) is not None
        )
        if not has_heading_usage:
            return {}, {}

    existing_levels = getattr(config.heading_numbering, "levels", {}) or {}
    extracted_levels: dict[str, HeadingLevelConfig] = {}
    binding_overrides: dict[str, dict[str, object]] = {}

    for level_key in _HEADING_LEVEL_KEYS:
        src_style = _find_style(doc, level_key)
        if src_style is None:
            continue

        num_pr = _style_numpr(src_style)
        if num_pr is None:
            continue

        num_id, ilvl = num_pr
        extracted = _extract_heading_numbering_level_from_numpr(
            existing_levels,
            level_key=level_key,
            src_style=src_style,
            num_map=num_map,
            abstract_map=abstract_map,
            num_id=num_id,
            ilvl=ilvl,
        )
        if extracted is None:
            continue
        level_cfg, binding_override = extracted
        extracted_levels[level_key] = level_cfg
        binding_overrides[level_key] = binding_override

    return extracted_levels, binding_overrides


def _heading_level_from_paragraph(para, heading_styles: dict[str, object]) -> str | None:
    para_style = getattr(para, "style", None)
    if para_style is None:
        return None

    para_style_id = str(getattr(para_style, "style_id", "") or "").strip()
    para_style_name = str(getattr(para_style, "name", "") or "").strip().lower()
    for level_key in _HEADING_LEVEL_KEYS:
        heading_style = heading_styles.get(level_key)
        if heading_style is None:
            continue
        if para_style == heading_style:
            return level_key

        heading_style_id = str(getattr(heading_style, "style_id", "") or "").strip()
        if para_style_id and heading_style_id and para_style_id == heading_style_id:
            return level_key

        heading_style_name = str(getattr(heading_style, "name", "") or "").strip().lower()
        if para_style_name and heading_style_name and para_style_name == heading_style_name:
            return level_key
    return None


def _clone_heading_numbering_from_paragraphs(
    config: SceneConfig,
    doc: Document,
) -> tuple[dict[str, HeadingLevelConfig], dict[str, dict[str, object]]]:
    num_map, abstract_map = _build_numbering_maps(doc)
    if not num_map or not abstract_map:
        return {}, {}

    heading_styles = {
        level_key: _find_style(doc, level_key)
        for level_key in _HEADING_LEVEL_KEYS
    }
    if not any(style is not None for style in heading_styles.values()):
        return {}, {}

    existing_levels = getattr(config.heading_numbering, "levels", {}) or {}
    candidates_by_level: dict[str, list[dict[str, object]]] = defaultdict(list)

    for para in doc.paragraphs:
        level_key = _heading_level_from_paragraph(para, heading_styles)
        if level_key is None:
            continue

        fallback_ilvl = max(0, _heading_level_index(level_key) - 1)
        num_pr = _paragraph_numpr(para, fallback_ilvl=fallback_ilvl)
        if num_pr is None:
            continue

        num_id, ilvl = num_pr
        extracted = _extract_heading_numbering_level_from_numpr(
            existing_levels,
            level_key=level_key,
            src_style=getattr(para, "style", None),
            num_map=num_map,
            abstract_map=abstract_map,
            num_id=num_id,
            ilvl=ilvl,
        )
        if extracted is None:
            continue

        level_cfg, binding_override = extracted
        candidates_by_level[level_key].append(
            {
                "level_cfg": level_cfg,
                "binding_override": binding_override,
            }
        )

    extracted_levels: dict[str, HeadingLevelConfig] = {}
    binding_overrides: dict[str, dict[str, object]] = {}
    for level_key in _HEADING_LEVEL_KEYS:
        selected = _select_heading_numbering_candidate(candidates_by_level.get(level_key, []))
        if selected is None:
            continue
        level_cfg, binding_override = selected
        extracted_levels[level_key] = level_cfg
        binding_overrides[level_key] = binding_override

    return extracted_levels, binding_overrides


def _build_default_heading_numbering_levels(config: SceneConfig) -> dict[str, HeadingLevelConfig]:
    defaults: dict[str, HeadingLevelConfig] = {}
    existing = getattr(config.heading_numbering, "levels", {}) or {}
    for idx in range(1, 5):
        key = f"heading{idx}"
        base = existing.get(key, HeadingLevelConfig())
        lc = HeadingLevelConfig(**vars(base))
        lc.format = "arabic" if idx == 1 else "arabic_dotted"
        lc.template = default_heading_numbering_template(key)
        lc.separator = lc.separator or "\u3000"
        lc.custom_separator = None
        defaults[key] = lc
    return defaults


def _clone_heading_numbering(config: SceneConfig, doc: Document) -> dict:
    samples_by_level: dict[str, list[dict[str, object]]] = defaultdict(list)
    sample_count = 0

    for para in doc.paragraphs:
        style_name = ((para.style.name if para.style is not None else "") or "").strip().lower()
        if "toc" in style_name or "目录" in style_name:
            continue

        text = (para.text or "").strip()
        inferred = _infer_heading_pattern(text)
        if inferred is None:
            continue

        level_key, fmt, template, separator, _title = inferred
        if level_key not in _HEADING_LEVEL_KEYS:
            continue

        sample = {
            "format": fmt,
            "template": template,
            "separator": separator,
        }
        align = ALIGNMENT_REVERSE.get(para.paragraph_format.alignment)
        if align:
            sample["alignment"] = align

        style_cfg = config.styles.get(level_key) or config.styles.get("normal")
        indent_chars = _guess_heading_indent_chars(para, style_cfg)
        if indent_chars is not None:
            sample["left_indent_chars"] = indent_chars

        samples_by_level[level_key].append(sample)
        sample_count += 1

    inferred_levels: dict[str, HeadingLevelConfig] = {}
    existing_levels = getattr(config.heading_numbering, "levels", {}) or {}

    for level_key in _HEADING_LEVEL_KEYS:
        level_samples = samples_by_level.get(level_key)
        if not level_samples:
            continue

        base = existing_levels.get(level_key, HeadingLevelConfig())
        lc = HeadingLevelConfig(**vars(base))

        fmt_template_counter = Counter(
            (str(s["format"]), str(s["template"])) for s in level_samples
        )
        (fmt, template), _ = fmt_template_counter.most_common(1)[0]
        lc.format = fmt
        lc.template = template

        separator_counter = Counter(str(s["separator"]) for s in level_samples)
        lc.separator = separator_counter.most_common(1)[0][0]
        lc.custom_separator = None

        align_counter = Counter(
            str(s["alignment"]) for s in level_samples if s.get("alignment")
        )
        if align_counter:
            align_value, align_count = align_counter.most_common(1)[0]
            if align_count >= 2:
                lc.alignment = align_value

        indent_values = [
            float(s["left_indent_chars"])
            for s in level_samples
            if s.get("left_indent_chars") is not None
        ]
        if len(indent_values) >= 2:
            lc.left_indent_chars = round(sum(indent_values) / len(indent_values), 2)

        inferred_levels[level_key] = lc

    use_inferred = bool(inferred_levels) and (
        sample_count >= 3 or (sample_count >= 2 and len(inferred_levels) >= 2)
    )
    style_bound_levels, style_bound_binding_overrides = _clone_heading_numbering_from_styles(config, doc)
    paragraph_bound_levels, paragraph_bound_binding_overrides = _clone_heading_numbering_from_paragraphs(config, doc)
    direct_bound_levels = dict(style_bound_levels)
    direct_bound_levels.update(paragraph_bound_levels)
    direct_bound_binding_overrides = dict(style_bound_binding_overrides)
    direct_bound_binding_overrides.update(paragraph_bound_binding_overrides)

    levels_updated: dict[str, HeadingLevelConfig] | None = None
    inferred_used = False
    fallback_used = False
    before_levels = getattr(config.heading_numbering, "levels", {}) or {}

    if direct_bound_levels:
        levels_updated = dict(inferred_levels) if use_inferred else {}
        levels_updated.update(direct_bound_levels)
        inferred_used = use_inferred and any(
            level_key not in direct_bound_levels for level_key in inferred_levels
        )
    elif use_inferred:
        levels_updated = inferred_levels
        inferred_used = True
    elif before_levels:
        levels_updated = None
    else:
        levels_updated = _build_default_heading_numbering_levels(config)
        fallback_used = True

    if levels_updated is not None:
        config.heading_numbering.levels = dict(levels_updated)
        scheme_id = str(getattr(config.heading_numbering, "scheme", "") or "").strip()
        if scheme_id:
            schemes = getattr(config.heading_numbering, "schemes", {}) or {}
            schemes[scheme_id] = dict(levels_updated)
            config.heading_numbering.schemes = schemes
        config.heading_numbering_v2 = _derive_heading_numbering_v2_from_legacy(
            config.heading_numbering,
            enabled=bool(config.capabilities.get("heading_numbering", True)),
        )
        for level_key, override in direct_bound_binding_overrides.items():
            binding = config.heading_numbering_v2.level_bindings.get(level_key)
            level_cfg = levels_updated.get(level_key)
            if binding is None or level_cfg is None:
                continue

            binding.enabled = True
            binding.start_at = int(override.get("start_at", 1) or 1)
            binding.restart_on = override.get("restart_on")
            binding.title_separator = str(level_cfg.effective_separator or "")
            binding.ooxml_separator_mode = str(
                override.get("ooxml_separator_mode", "inline") or "inline"
            )
            binding.ooxml_suff = override.get("ooxml_suff")
            binding.ooxml_lvl_ind = dict(override.get("ooxml_lvl_ind", {}) or {})

            reference_core_style = override.get("reference_core_style")
            if reference_core_style:
                binding.reference_core_style = str(reference_core_style)

        config.heading_numbering.levels = legacy_levels_from_v2(
            config.heading_numbering_v2,
            existing_levels=config.heading_numbering.levels,
        )
        config._heading_numbering_v2_source = "payload"  # type: ignore[attr-defined]

    return {
        "numbering_updated": levels_updated is not None,
        "numbering_levels_updated": list(levels_updated.keys()) if levels_updated else [],
        "numbering_inferred": inferred_used,
        "numbering_fallback_used": fallback_used,
        "numbering_samples": sample_count,
    }


def _find_style(doc: Document, style_key: str):
    candidates = [style_key]
    candidates.extend(STYLE_CANDIDATES.get(style_key, []))

    styles = list(doc.styles)
    lowered = {(st.name or "").strip().lower(): st for st in styles if st.name}
    for name in candidates:
        key = (name or "").strip().lower()
        if not key:
            continue
        st = lowered.get(key)
        if st is not None:
            return st
    return None


def _clone_page_setup(config: SceneConfig, doc: Document) -> bool:
    if not doc.sections:
        return False
    sec = doc.sections[0]
    ps = config.page_setup

    changed = False

    def _set_cm(attr: str, value):
        nonlocal changed
        old = getattr(ps.margin, attr)
        new = round(float(value.cm), 2)
        if abs(old - new) > 1e-6:
            setattr(ps.margin, attr, new)
            changed = True

    _set_cm("top_cm", sec.top_margin)
    _set_cm("bottom_cm", sec.bottom_margin)
    _set_cm("left_cm", sec.left_margin)
    _set_cm("right_cm", sec.right_margin)

    header_new = round(float(sec.header_distance.cm), 2)
    footer_new = round(float(sec.footer_distance.cm), 2)
    if abs(ps.header_distance_cm - header_new) > 1e-6:
        ps.header_distance_cm = header_new
        changed = True
    if abs(ps.footer_distance_cm - footer_new) > 1e-6:
        ps.footer_distance_cm = footer_new
        changed = True

    w_cm = round(float(sec.page_width.cm), 2)
    h_cm = round(float(sec.page_height.cm), 2)
    detected_paper_size = None
    paper_specs = {
        "A4": (21.0, 29.7),
        "Letter": (21.59, 27.94),
        "A3": (29.7, 42.0),
    }
    for paper_name, (pw, ph) in paper_specs.items():
        if (
            abs(w_cm - pw) <= 0.3 and abs(h_cm - ph) <= 0.3
        ) or (
            abs(w_cm - ph) <= 0.3 and abs(h_cm - pw) <= 0.3
        ):
            detected_paper_size = paper_name
            break

    if detected_paper_size and ps.paper_size != detected_paper_size:
        ps.paper_size = detected_paper_size
        changed = True

    return changed


def clone_scene_style_from_docx(config: SceneConfig, template_docx_path: str | Path) -> dict:
    """Clone style/page settings from template docx into scene config.

    Returns:
        {
            "styles_updated": [...style_keys...],
            "styles_missing": [...style_keys...],
            "page_setup_updated": bool
        }
    """
    from src.docx_io.sanitize import sanitize_docx

    src_docx_path = str(template_docx_path)
    suffix = Path(src_docx_path).suffix or ".docx"
    temp_docx_path: str | None = None
    fd, temp_docx_path = tempfile.mkstemp(suffix=suffix, prefix=".clone-style-")
    os.close(fd)
    try:
        shutil.copy2(src_docx_path, temp_docx_path)
        try:
            sanitize_docx(temp_docx_path)
        except Exception:
            pass  # best-effort sanitize on temp copy only

        try:
            doc = Document(temp_docx_path)
        except Exception as exc:
            # Fallback for templates with parser-blocking oversized XML attributes.
            sanitize_docx(temp_docx_path, aggressive=True)
            try:
                doc = Document(temp_docx_path)
            except Exception as inner_exc:
                raise RuntimeError(
                    "无法读取模板 DOCX：检测到异常 XML 内容（可能存在超长属性值）。"
                ) from inner_exc
    finally:
        if temp_docx_path:
            try:
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
            except Exception:
                pass
    style_slots_added = _ensure_clone_style_slots(config)
    style_seeds = _clone_style_seed_map(config)
    updated: list[str] = []
    missing: list[str] = []

    for key, sc in config.styles.items():
        src_style = _find_style(doc, key)
        if src_style is None:
            missing.append(key)
            continue
        clean_base = _fresh_style_seed(key, style_seeds, fallback=sc)
        config.styles[key] = _build_style_config(
            src_style,
            clean_base,
            inherit_indents_from_style_chain=key not in _HEADING_LEVEL_KEYS,
        )
        updated.append(key)

    para_samples = _collect_paragraph_samples(doc)
    for key, para in para_samples.items():
        current = config.styles.get(key)
        if key in updated and isinstance(current, StyleConfig):
            base = StyleConfig(**vars(current))
        else:
            base = _fresh_style_seed(key, style_seeds, fallback=current)
        sc = _build_style_config_from_paragraph(para, base)
        if key in _PARAGRAPH_SAMPLE_RESET_INDENT_KEYS:
            sc.first_line_indent_chars = 0.0
            sc.left_indent_chars = 0.0
            sc.right_indent_chars = 0.0
            sync_style_config_indent_fields(sc)
        config.styles[key] = sc
        if key not in updated:
            updated.append(key)
        if key in missing:
            missing.remove(key)

    page_setup_updated = _clone_page_setup(config, doc)
    numbering_summary = _clone_heading_numbering(config, doc)

    return {
        "styles_updated": updated,
        "styles_missing": missing,
        "styles_added": style_slots_added,
        "page_setup_updated": page_setup_updated,
        **numbering_summary,
    }

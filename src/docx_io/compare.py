"""对比稿生成：基于 ChangeTracker 在原始文档上注入 OOXML 修订标记"""

import copy
from difflib import SequenceMatcher
from lxml import etree
from docx import Document

from src.utils.constants import OOXML_NS, REVISION_AUTHOR
from src.utils.ooxml import next_revision_id, reset_revision_counter, revision_date
from src.engine.change_tracker import ChangeTracker


W_NS = OOXML_NS["w"]


def _w(tag: str) -> str:
    """生成带 w: 命名空间的完整标签名"""
    return f"{{{W_NS}}}{tag}"


def _make_run_props(source_run_elem):
    """从源 run 元素复制 rPr（运行属性）"""
    rpr = source_run_elem.find(_w("rPr"))
    if rpr is not None:
        return copy.deepcopy(rpr)
    return None


def _build_del_run(text: str, rpr_elem, rev_id: int,
                   author: str, date: str):
    """构建 <w:del> 包裹的删除 run"""
    del_elem = etree.Element(_w("del"))
    del_elem.set(_w("id"), str(rev_id))
    del_elem.set(_w("author"), author)
    del_elem.set(_w("date"), date)

    r = etree.SubElement(del_elem, _w("r"))
    if rpr_elem is not None:
        r.append(copy.deepcopy(rpr_elem))
    dt = etree.SubElement(r, _w("delText"))
    dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    dt.text = text
    return del_elem


def _build_ins_run(text: str, rpr_elem, rev_id: int,
                   author: str, date: str):
    """构建 <w:ins> 包裹的插入 run"""
    ins_elem = etree.Element(_w("ins"))
    ins_elem.set(_w("id"), str(rev_id))
    ins_elem.set(_w("author"), author)
    ins_elem.set(_w("date"), date)

    r = etree.SubElement(ins_elem, _w("r"))
    if rpr_elem is not None:
        r.append(copy.deepcopy(rpr_elem))
    t = etree.SubElement(r, _w("t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return ins_elem


def _build_ins_container(rev_id: int, author: str, date: str):
    """构建空的 <w:ins> 容器，用于包裹原段落内容。"""
    ins_elem = etree.Element(_w("ins"))
    ins_elem.set(_w("id"), str(rev_id))
    ins_elem.set(_w("author"), author)
    ins_elem.set(_w("date"), date)
    return ins_elem


def _paragraph_has_complex_content(para_elem) -> bool:
    """Detect whether paragraph contains non-trivial OOXML that should be preserved."""
    allowed_direct = {_w("pPr"), _w("r")}
    for child in list(para_elem):
        if child.tag not in allowed_direct:
            return True
    complex_tags = {
        _w("drawing"),
        _w("object"),
        _w("pict"),
        _w("fldSimple"),
        _w("instrText"),
        _w("fldChar"),
        "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath",
        "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara",
    }
    for node in para_elem.iter():
        if node.tag in complex_tags:
            return True
    return False


def _wrap_paragraph_content_as_inserted(para_elem, author: str, date: str) -> bool:
    """将段落现有内容整体标为插入修订（不改写内容本身）。"""
    content_children = [child for child in list(para_elem) if child.tag != _w("pPr")]
    if not content_children:
        return False

    ins_elem = _build_ins_container(next_revision_id(), author, date)
    for child in content_children:
        para_elem.remove(child)
        ins_elem.append(child)

    ppr = para_elem.find(_w("pPr"))
    insert_idx = list(para_elem).index(ppr) + 1 if ppr is not None else 0
    para_elem.insert(insert_idx, ins_elem)
    return True


def _inject_text_revision(para_elem, old_text: str, new_text: str,
                          author: str, date: str):
    """在段落 XML 中注入文本修订标记：删除旧文本 + 插入新文本"""
    old_text = old_text or ""
    new_text = new_text or ""
    if old_text == "" and new_text == "":
        return False

    # 获取第一个 run（可在 hyperlink/fld 容器内）的格式属性作为模板
    first_run = None
    for node in para_elem.iter(_w("r")):
        first_run = node
        break
    rpr_elem = None
    if first_run is not None:
        rpr_elem = _make_run_props(first_run)

    # 找到插入位置（在 pPr 之后）
    ppr = para_elem.find(_w("pPr"))
    insert_idx = list(para_elem).index(ppr) + 1 if ppr is not None else 0

    # Preserve complex paragraph structures (e.g., hyperlinks/fields/objects/math).
    # In this branch we only prepend revision marks and keep existing nodes untouched.
    if _paragraph_has_complex_content(para_elem):
        cur_idx = insert_idx
        if old_text:
            del_id = next_revision_id()
            del_elem = _build_del_run(old_text, rpr_elem, del_id, author, date)
            para_elem.insert(cur_idx, del_elem)
            cur_idx += 1
        if new_text:
            ins_id = next_revision_id()
            ins_elem = _build_ins_run(new_text, rpr_elem, ins_id, author, date)
            para_elem.insert(cur_idx, ins_elem)
        return True

    # 清除段落中所有内容节点，保留段落属性 pPr。
    for child in list(para_elem):
        if child.tag != _w("pPr"):
            para_elem.remove(child)

    cur_idx = insert_idx
    # 注入 <w:del> 旧文本
    if old_text:
        del_id = next_revision_id()
        del_elem = _build_del_run(old_text, rpr_elem, del_id, author, date)
        para_elem.insert(cur_idx, del_elem)
        cur_idx += 1

    # 注入 <w:ins> 新文本
    if new_text:
        ins_id = next_revision_id()
        ins_elem = _build_ins_run(new_text, rpr_elem, ins_id, author, date)
        para_elem.insert(cur_idx, ins_elem)

    return True


def _insert_revision_paragraph(compare_doc: Document, insert_at: int,
                               old_text: str, new_text: str,
                               author: str, date: str) -> bool:
    """在 compare_doc 指定段落索引处插入一段修订段落。"""
    new_p = etree.Element(_w("p"))
    injected = _inject_text_revision(
        new_p, old_text=old_text, new_text=new_text, author=author, date=date
    )
    if not injected:
        return False

    paragraphs = compare_doc.paragraphs
    body = compare_doc.element.body

    if paragraphs:
        if insert_at <= 0:
            paragraphs[0]._element.addprevious(new_p)
        elif insert_at >= len(paragraphs):
            paragraphs[-1]._element.addnext(new_p)
        else:
            paragraphs[insert_at]._element.addprevious(new_p)
    else:
        sect_pr = body.find(_w("sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(new_p)
        else:
            body.append(new_p)
    return True


def _node_xml_bytes(node) -> bytes:
    if node is None:
        return b""
    return etree.tostring(node, with_tail=False)


def _build_style_format_signature_map(doc: Document) -> dict[str, bytes]:
    """Build recursive style-format signatures keyed by style_id."""
    style_elems: dict[str, etree._Element] = {}
    for style in doc.styles:
        style_id = getattr(style, "style_id", "") or ""
        if not style_id:
            continue
        style_elems[style_id] = style.element

    memo: dict[str, bytes] = {}
    visiting: set[str] = set()

    def _sig(style_id: str) -> bytes:
        if not style_id:
            return b""
        if style_id in memo:
            return memo[style_id]
        if style_id in visiting:
            return b""
        visiting.add(style_id)
        el = style_elems.get(style_id)
        if el is None:
            visiting.remove(style_id)
            memo[style_id] = b""
            return b""

        ppr_sig = _node_xml_bytes(el.find(_w("pPr")))
        rpr_sig = _node_xml_bytes(el.find(_w("rPr")))
        base_id = ""
        based_on = el.find(_w("basedOn"))
        if based_on is not None:
            base_id = (
                based_on.get(_w("val"))
                or based_on.get("val")
                or ""
            )
        base_sig = _sig(base_id)
        visiting.remove(style_id)
        out = b"|PPR|" + ppr_sig + b"|RPR|" + rpr_sig + b"|BASE|" + base_sig
        memo[style_id] = out
        return out

    for sid in list(style_elems.keys()):
        _sig(sid)
    return memo


def _paragraph_style_id(para) -> str:
    try:
        style = para.style
    except Exception:
        style = None
    if style is None:
        return ""
    return getattr(style, "style_id", "") or ""


def _run_style_id(run_el) -> str:
    rpr = run_el.find(_w("rPr"))
    if rpr is None:
        return ""
    rstyle = rpr.find(_w("rStyle"))
    if rstyle is None:
        return ""
    return (
        rstyle.get(_w("val"))
        or rstyle.get("val")
        or ""
    )


def _paragraph_format_fingerprint(para, style_sig_map: dict[str, bytes]) -> tuple[bytes, bytes, tuple[tuple[bytes, bytes], ...]]:
    """Return text-insensitive paragraph formatting fingerprint."""
    ppr = para._element.find(_w("pPr"))
    ppr_sig = _node_xml_bytes(ppr)
    para_style_sig = style_sig_map.get(_paragraph_style_id(para), b"")
    run_sigs: list[tuple[bytes, bytes]] = []
    for run_el in para._element.iter(_w("r")):
        rpr_sig = _node_xml_bytes(run_el.find(_w("rPr")))
        rstyle_sig = style_sig_map.get(_run_style_id(run_el), b"")
        run_sigs.append((rpr_sig, rstyle_sig))
    return ppr_sig, para_style_sig, tuple(run_sigs)


def _collect_format_changed_final_indices(
    original_doc: Document,
    final_doc: Document,
) -> set[int]:
    """Find final-doc paragraph indices whose text is unchanged but format differs."""
    original_texts = [p.text for p in original_doc.paragraphs]
    final_texts = [p.text for p in final_doc.paragraphs]
    sm = SequenceMatcher(a=original_texts, b=final_texts, autojunk=False)
    original_style_sig = _build_style_format_signature_map(original_doc)
    final_style_sig = _build_style_format_signature_map(final_doc)

    changed: set[int] = set()
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag != "equal":
            continue
        span = min(i2 - i1, j2 - j1)
        for k in range(span):
            oi = i1 + k
            fj = j1 + k
            if oi >= len(original_doc.paragraphs) or fj >= len(final_doc.paragraphs):
                continue
            if _paragraph_format_fingerprint(original_doc.paragraphs[oi], original_style_sig) != _paragraph_format_fingerprint(final_doc.paragraphs[fj], final_style_sig):
                changed.add(fj)
    return changed


def _build_final_to_compare_index_map(
    original_texts: list[str],
    final_texts: list[str],
) -> dict[int, int]:
    """Map final_doc paragraph index -> compare_doc index after text-revision injection."""
    sm = SequenceMatcher(a=original_texts, b=final_texts, autojunk=False)
    idx_map: dict[int, int] = {}
    offset = 0
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag in {"equal", "insert"}:
            for j in range(j1, j2):
                idx_map[j] = j + offset
            continue

        if tag == "replace":
            for j in range(j1, j2):
                idx_map[j] = j + offset
            old_count = i2 - i1
            new_count = j2 - j1
            pair_count = min(old_count, new_count)
            offset += (old_count - pair_count)
            continue

        if tag == "delete":
            offset += (i2 - i1)
            continue
    return idx_map


def _inject_revisions_by_doc_diff_on_original_base(
        compare_doc: Document, final_doc: Document,
        author: str, date: str) -> tuple[int, set[int]]:
    """以原稿为底稿注入修订。"""
    original_texts = [p.text for p in compare_doc.paragraphs]
    final_texts = [p.text for p in final_doc.paragraphs]

    sm = SequenceMatcher(a=original_texts, b=final_texts, autojunk=False)
    revised_para_indices = set()
    injected_count = 0
    offset = 0

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue

        if tag == "replace":
            old_count = i2 - i1
            new_count = j2 - j1
            pair_count = min(old_count, new_count)

            for k in range(pair_count):
                idx = i1 + k + offset
                if idx < len(compare_doc.paragraphs):
                    ok = _inject_text_revision(
                        compare_doc.paragraphs[idx]._element,
                        old_text=original_texts[i1 + k],
                        new_text=final_texts[j1 + k],
                        author=author,
                        date=date,
                    )
                    if ok:
                        injected_count += 1
                        revised_para_indices.add(idx)

            for k in range(pair_count, old_count):
                idx = i1 + k + offset
                if idx < len(compare_doc.paragraphs):
                    ok = _inject_text_revision(
                        compare_doc.paragraphs[idx]._element,
                        old_text=original_texts[i1 + k],
                        new_text="",
                        author=author,
                        date=date,
                    )
                    if ok:
                        injected_count += 1
                        revised_para_indices.add(idx)

            insert_at = i1 + pair_count + offset
            for k in range(pair_count, new_count):
                ok = _insert_revision_paragraph(
                    compare_doc,
                    insert_at=insert_at,
                    old_text="",
                    new_text=final_texts[j1 + k],
                    author=author,
                    date=date,
                )
                if ok:
                    injected_count += 1
                    insert_at += 1
                    offset += 1

        elif tag == "delete":
            for k in range(i1, i2):
                idx = k + offset
                if idx < len(compare_doc.paragraphs):
                    ok = _inject_text_revision(
                        compare_doc.paragraphs[idx]._element,
                        old_text=original_texts[k],
                        new_text="",
                        author=author,
                        date=date,
                    )
                    if ok:
                        injected_count += 1
                        revised_para_indices.add(idx)

        elif tag == "insert":
            insert_at = i1 + offset
            for k in range(j1, j2):
                ok = _insert_revision_paragraph(
                    compare_doc,
                    insert_at=insert_at,
                    old_text="",
                    new_text=final_texts[k],
                    author=author,
                    date=date,
                )
                if ok:
                    injected_count += 1
                    insert_at += 1
                    offset += 1

    return injected_count, revised_para_indices


def _inject_revisions_by_doc_diff_on_final_base(
        compare_doc: Document, original_doc: Document, final_doc: Document,
        author: str, date: str) -> tuple[int, set[int]]:
    """以终稿为底稿注入修订，保留终稿样式。"""
    original_texts = [p.text for p in original_doc.paragraphs]
    final_texts = [p.text for p in final_doc.paragraphs]

    sm = SequenceMatcher(a=original_texts, b=final_texts, autojunk=False)
    revised_para_indices = set()
    injected_count = 0
    offset = 0

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue

        if tag == "replace":
            old_count = i2 - i1
            new_count = j2 - j1
            pair_count = min(old_count, new_count)

            # 同位替换：在终稿段落上注入 del+ins，保留终稿段落样式
            for k in range(pair_count):
                idx = j1 + k + offset
                if idx < len(compare_doc.paragraphs):
                    ok = _inject_text_revision(
                        compare_doc.paragraphs[idx]._element,
                        old_text=original_texts[i1 + k],
                        new_text=final_texts[j1 + k],
                        author=author,
                        date=date,
                    )
                    if ok:
                        injected_count += 1
                        revised_para_indices.add(idx)

            # 原稿有、终稿无：插入删除段落
            insert_at = j1 + pair_count + offset
            for k in range(pair_count, old_count):
                ok = _insert_revision_paragraph(
                    compare_doc,
                    insert_at=insert_at,
                    old_text=original_texts[i1 + k],
                    new_text="",
                    author=author,
                    date=date,
                )
                if ok:
                    injected_count += 1
                    revised_para_indices.add(insert_at)
                    insert_at += 1
                    offset += 1

            # 终稿有、原稿无：在终稿现有段落上标注插入
            for k in range(pair_count, new_count):
                idx = j1 + k + offset
                if idx < len(compare_doc.paragraphs):
                    ok = _wrap_paragraph_content_as_inserted(
                        compare_doc.paragraphs[idx]._element,
                        author=author,
                        date=date,
                    )
                    if ok:
                        injected_count += 1
                        revised_para_indices.add(idx)

        elif tag == "delete":
            # 原稿删除段落插入到终稿对应位置
            insert_at = j1 + offset
            for k in range(i1, i2):
                ok = _insert_revision_paragraph(
                    compare_doc,
                    insert_at=insert_at,
                    old_text=original_texts[k],
                    new_text="",
                    author=author,
                    date=date,
                )
                if ok:
                    injected_count += 1
                    revised_para_indices.add(insert_at)
                    insert_at += 1
                    offset += 1

        elif tag == "insert":
            # 终稿新增段落已存在 compare_doc，对其打插入修订
            for k in range(j1, j2):
                idx = k + offset
                if idx < len(compare_doc.paragraphs):
                    ok = _wrap_paragraph_content_as_inserted(
                        compare_doc.paragraphs[idx]._element,
                        author=author,
                        date=date,
                    )
                    if ok:
                        injected_count += 1
                        revised_para_indices.add(idx)

    return injected_count, revised_para_indices


def generate_compare_doc(original_doc: Document,
                         tracker: ChangeTracker,
                         output_path: str,
                         final_doc=None,
                         include_text: bool = True,
                         include_formatting: bool = True) -> None:
    """基于原始文档和变更记录生成带修订标记的对比稿。

    优先使用 original_doc 与 final_doc 的段落文本差异生成修订；
    若 final_doc 缺失，则回退到 tracker.text 记录。
    """
    reset_revision_counter()
    # Use save/reload clone to avoid deepcopy corruption on python-docx/lxml trees.
    from io import BytesIO
    source_doc = final_doc if final_doc is not None else original_doc
    buf = BytesIO()
    source_doc.save(buf)
    buf.seek(0)
    compare_doc = Document(buf)
    author = REVISION_AUTHOR
    date = revision_date()

    injected_count = 0
    revised_para_indices: set[int] = set()

    if include_text:
        if final_doc is not None:
            injected_count, revised_para_indices = _inject_revisions_by_doc_diff_on_final_base(
                compare_doc, original_doc, final_doc, author, date
            )
        else:
            # 回退模式：仅使用 tracker 中 text 变更
            text_changes = {}
            for rec in tracker.records:
                if rec.change_type == "text" and rec.paragraph_index >= 0:
                    text_changes[rec.paragraph_index] = rec

            paragraphs = compare_doc.paragraphs
            for para_idx, rec in text_changes.items():
                if para_idx >= len(paragraphs):
                    continue
                ok = _inject_text_revision(
                    paragraphs[para_idx]._element,
                    old_text=rec.before,
                    new_text=rec.after,
                    author=author,
                    date=date,
                )
                if ok:
                    injected_count += 1
                    revised_para_indices.add(para_idx)

    # Formatting revisions:
    # 1) With final_doc: detect unchanged-text paragraphs whose formatting differs.
    # 2) Without final_doc (fallback mode): keep tracker paragraph-index behavior.
    if include_formatting:
        if final_doc is not None:
            format_changed_final = _collect_format_changed_final_indices(original_doc, final_doc)
            if format_changed_final:
                original_texts = [p.text for p in original_doc.paragraphs]
                final_texts = [p.text for p in final_doc.paragraphs]
                if include_text:
                    final_to_compare = _build_final_to_compare_index_map(original_texts, final_texts)
                else:
                    final_to_compare = {i: i for i in range(len(final_texts))}
                for final_idx in sorted(format_changed_final):
                    idx = final_to_compare.get(final_idx)
                    if idx is None or idx in revised_para_indices:
                        continue
                    if idx >= len(compare_doc.paragraphs):
                        continue
                    ok = _wrap_paragraph_content_as_inserted(
                        compare_doc.paragraphs[idx]._element,
                        author=author,
                        date=date,
                    )
                    if ok:
                        revised_para_indices.add(idx)
                        injected_count += 1
        elif final_doc is None:
            for rec in tracker.records:
                idx = rec.paragraph_index
                if idx < 0 or idx in revised_para_indices or rec.change_type in {"error", "text"}:
                    continue
                if idx >= len(compare_doc.paragraphs):
                    continue
                ok = _wrap_paragraph_content_as_inserted(
                    compare_doc.paragraphs[idx]._element,
                    author=author,
                    date=date,
                )
                if ok:
                    revised_para_indices.add(idx)
                    injected_count += 1

    # 最终兜底：若仍无修订标记，但确有格式变更记录，则插入一段修订说明
    if include_formatting and injected_count == 0:
        effective = [
            rec for rec in tracker.records
            if rec.success and rec.change_type not in {"error", "text"}
        ]
        if effective:
            preview = "；".join(
                f"[{rec.rule_name}] {rec.target}" for rec in effective[:5]
            )
            summary = f"自动排版共 {len(effective)} 项变更：{preview}"
            _insert_revision_paragraph(
                compare_doc,
                insert_at=len(compare_doc.paragraphs),
                old_text="",
                new_text=summary,
                author=author,
                date=date,
            )

    compare_doc.save(output_path)

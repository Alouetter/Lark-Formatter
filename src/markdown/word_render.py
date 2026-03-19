"""共享 Word 渲染工具函数：Markdown IR → Word OOXML"""

import os
import logging
import re
import unicodedata
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.markdown.ir import InlineSpan, InlineType
from src.markdown.inline_parser import parse_inline
from src.utils.line_spacing import apply_safe_picture_line_spacing

logger = logging.getLogger(__name__)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# 缓存编号 ID，避免重复遍历
_numbering_id_cache = {}


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


# ── 自动更正防护 ──

def set_no_proof(run) -> None:
    """为 run 添加 <w:noProof/> 防止 Word 自动更正。

    防止 -- → 破折号、... → 省略号、" → 智能引号等。
    """
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn('w:noProof')) is None:
        etree.SubElement(rpr, qn('w:noProof'))


# ── 真超链接 ──

def add_hyperlink(para, text: str, url: str) -> None:
    """在段落中创建 Word 真超链接 <w:hyperlink>。

    通过 part.relate_to() 注册外部关系，生成 r:id，
    构建 <w:hyperlink r:id="..."> 包含蓝色下划线 run。
    """
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = etree.SubElement(para._element, qn('w:hyperlink'))
    hyperlink.set(qn('r:id'), r_id)

    run_el = etree.SubElement(hyperlink, qn('w:r'))
    rpr = etree.SubElement(run_el, qn('w:rPr'))

    # 蓝色
    color = etree.SubElement(rpr, qn('w:color'))
    color.set(qn('w:val'), '0563C1')
    # 下划线
    u = etree.SubElement(rpr, qn('w:u'))
    u.set(qn('w:val'), 'single')
    # 使用 Hyperlink 样式（如果存在）
    rstyle = etree.SubElement(rpr, qn('w:rStyle'))
    rstyle.set(qn('w:val'), 'Hyperlink')

    t = etree.SubElement(run_el, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')


# ── 列表编号 ──

# 无序列表各层级：(字符, ascii字体, hAnsi字体)
# Word 标准 bullet 定义：Symbol \uF0B7, Courier New "o", Wingdings \uF0A7
_BULLET_DEFS = [
    ("\uF0B7", "Symbol", "Symbol"),
    ("o", "Courier New", "Courier New"),
    ("\uF0A7", "Wingdings", "Wingdings"),
    ("\uF0B7", "Symbol", "Symbol"),
    ("o", "Courier New", "Courier New"),
    ("\uF0A7", "Wingdings", "Wingdings"),
    ("\uF0B7", "Symbol", "Symbol"),
    ("o", "Courier New", "Courier New"),
    ("\uF0A7", "Wingdings", "Wingdings"),
]
# 有序列表各层级格式
_DECIMAL_FMTS = [
    ("decimal", "%1."),
    ("lowerLetter", "%2)"),
    ("lowerRoman", "%3."),
    ("decimal", "(%4)"),
    ("lowerLetter", "%5."),
    ("lowerRoman", "%6."),
    ("decimal", "%7."),
    ("lowerLetter", "%8)"),
    ("lowerRoman", "%9."),
]


def normalize_ordered_list_style(style: str | None) -> str:
    raw = str(style or "").strip().lower()
    aliases = {
        "legacy": "mixed",
        "legacy_mixed": "mixed",
        "default": "mixed",
        "dot": "decimal_dot",
        "paren_right": "decimal_paren_right",
        "cn_dun": "decimal_cn_dun",
        "full_paren": "decimal_full_paren",
    }
    raw = aliases.get(raw, raw)
    if raw in {
        "mixed",
        "decimal_dot",
        "decimal_paren_right",
        "decimal_cn_dun",
        "decimal_full_paren",
    }:
        return raw
    return "mixed"


def normalize_unordered_list_style(style: str | None) -> str:
    raw = str(style or "").strip().lower()
    aliases = {
        "legacy": "word_default",
        "default": "word_default",
        "dot": "bullet_dot",
        "circle": "bullet_circle",
        "square": "bullet_square",
        "dash": "bullet_dash",
    }
    raw = aliases.get(raw, raw)
    if raw in {
        "word_default",
        "bullet_dot",
        "bullet_circle",
        "bullet_square",
        "bullet_dash",
    }:
        return raw
    return "word_default"


def get_ordered_level_format(style: str | None, ilvl: int) -> tuple[str, str]:
    ilvl = max(0, min(int(ilvl), 8))
    style_key = normalize_ordered_list_style(style)
    token = f"%{ilvl + 1}"
    if style_key == "decimal_dot":
        return "decimal", f"{token}."
    if style_key == "decimal_paren_right":
        return "decimal", f"{token})"
    if style_key == "decimal_cn_dun":
        return "decimal", f"{token}、"
    if style_key == "decimal_full_paren":
        return "decimal", f"（{token}）"
    return _DECIMAL_FMTS[ilvl]


def get_unordered_level_definition(style: str | None, ilvl: int) -> tuple[str, str, str]:
    ilvl = max(0, min(int(ilvl), 8))
    style_key = normalize_unordered_list_style(style)
    if style_key == "bullet_dot":
        return "•", "Times New Roman", "Times New Roman"
    if style_key == "bullet_circle":
        return "○", "宋体", "宋体"
    if style_key == "bullet_square":
        return "■", "宋体", "宋体"
    if style_key == "bullet_dash":
        return "-", "Times New Roman", "Times New Roman"
    return _BULLET_DEFS[ilvl]


def register_list_numbering(
        doc,
        list_type: str = "bullet",
        *,
        marker_separator: str = "tab",
        ordered_style: str = "mixed",
        unordered_style: str = "word_default") -> int:
    """注册 9 级多级列表编号定义，返回 numId。

    list_type: "bullet" 或 "decimal"
    """
    numbering_part = doc.part.numbering_part
    numbering_el = numbering_part.element

    abs_id = _next_num_id(numbering_el, "abstractNum", "abstractNumId")
    num_id = _next_num_id(numbering_el, "num", "numId")

    sep_key = str(marker_separator or "tab").strip().lower()
    if sep_key in {"half_space", "space", "halfwidth_space"}:
        suff_val = "space"
    elif sep_key in {"full_space", "fullwidth_space"}:
        # Word 编号后缀不支持全角空格，使用 nothing 并由调用方补入 U+3000。
        suff_val = "nothing"
    else:
        suff_val = "tab"
    ordered_style = normalize_ordered_list_style(ordered_style)
    unordered_style = normalize_unordered_list_style(unordered_style)

    # 构建 abstractNum
    abs_num = etree.SubElement(numbering_el, _w("abstractNum"))
    abs_num.set(_w("abstractNumId"), str(abs_id))
    ml_type = etree.SubElement(abs_num, _w("multiLevelType"))
    ml_type.set(_w("val"), "multilevel")

    for ilvl in range(9):
        lvl = etree.SubElement(abs_num, _w("lvl"))
        lvl.set(_w("ilvl"), str(ilvl))

        start = etree.SubElement(lvl, _w("start"))
        start.set(_w("val"), "1")

        if list_type == "bullet":
            bullet_char, ascii_font, hansi_font = get_unordered_level_definition(
                unordered_style, ilvl
            )
            fmt = etree.SubElement(lvl, _w("numFmt"))
            fmt.set(_w("val"), "bullet")
            txt = etree.SubElement(lvl, _w("lvlText"))
            txt.set(_w("val"), bullet_char)
        else:
            num_fmt_val, lvl_text_val = get_ordered_level_format(
                ordered_style, ilvl
            )
            fmt = etree.SubElement(lvl, _w("numFmt"))
            fmt.set(_w("val"), num_fmt_val)
            txt = etree.SubElement(lvl, _w("lvlText"))
            txt.set(_w("val"), lvl_text_val)

        jc = etree.SubElement(lvl, _w("lvlJc"))
        jc.set(_w("val"), "left")

        # 列表序号后缀分隔符（tab / space / nothing）
        suff = etree.SubElement(lvl, _w("suff"))
        suff.set(_w("val"), suff_val)

        # 缩进：每层增加 0.75cm
        ppr = etree.SubElement(lvl, _w("pPr"))
        indent = etree.SubElement(ppr, _w("ind"))
        left_twips = 425 * (ilvl + 1)  # ~0.75cm per level
        hanging = 425
        indent.set(_w("left"), str(left_twips))
        indent.set(_w("hanging"), str(hanging))

        # 字体格式 rPr（bullet 需要对应字体才能正确渲染符号）
        rpr = etree.SubElement(lvl, _w("rPr"))
        rfonts = etree.SubElement(rpr, _w("rFonts"))
        if list_type == "bullet":
            rfonts.set(_w("ascii"), ascii_font)
            rfonts.set(_w("hAnsi"), hansi_font)
            rfonts.set(_w("hint"), "default")
        else:
            rfonts.set(_w("hint"), "default")

    # 移动 abstractNum 到 num 之前
    first_num = numbering_el.find(_w("num"))
    if first_num is not None:
        first_num.addprevious(abs_num)

    # 构建 num 引用
    num_el = etree.SubElement(numbering_el, _w("num"))
    num_el.set(_w("numId"), str(num_id))
    abs_ref = etree.SubElement(num_el, _w("abstractNumId"))
    abs_ref.set(_w("val"), str(abs_id))

    return num_id


def _next_num_id(numbering_el, tag: str, attr: str) -> int:
    """获取下一个可用的编号 ID（使用缓存优化性能）"""
    cache_key = (id(numbering_el), tag, attr)

    max_id = 0
    for el in numbering_el.findall(_w(tag)):
        val = el.get(_w(attr))
        if val and val.isdigit():
            max_id = max(max_id, int(val))

    # `id(numbering_el)` may be reused by another Document, so stale cache
    # entries must not allocate an ID below the real max in the current file.
    cached_max = int(_numbering_id_cache.get(cache_key, 0) or 0)
    next_id = max(max_id, cached_max) + 1
    _numbering_id_cache[cache_key] = next_id
    return next_id


def apply_num_pr(para, num_id: int, ilvl: int) -> None:
    """为段落设置 <w:numPr> 列表编号属性"""
    ppr = para._element.get_or_add_pPr()
    # 移除已有的 numPr
    old = ppr.find(qn('w:numPr'))
    if old is not None:
        ppr.remove(old)

    num_pr = etree.SubElement(ppr, qn('w:numPr'))
    ilvl_el = etree.SubElement(num_pr, qn('w:ilvl'))
    ilvl_el.set(qn('w:val'), str(ilvl))
    num_id_el = etree.SubElement(num_pr, qn('w:numId'))
    num_id_el.set(qn('w:val'), str(num_id))


# ── 引用块样式 ──

def apply_blockquote_border(para, level: int = 1) -> None:
    """为引用块段落添加左边框 + 浅灰底色。

    嵌套层级越深，边框越粗 (sz = 6 * level)。
    """
    ppr = para._element.get_or_add_pPr()

    # 左边框
    pbdr = ppr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = etree.SubElement(ppr, qn('w:pBdr'))
    left = etree.SubElement(pbdr, qn('w:left'))
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(6 * level))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), 'CCCCCC')

    # 浅灰底色
    shd = ppr.find(qn('w:shd'))
    if shd is None:
        shd = etree.SubElement(ppr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F5F5F5')

    # 左缩进
    para.paragraph_format.left_indent = Cm(0.5 * level)


# ── 统一 InlineSpan → Word runs ──

def write_spans(para, spans: list[InlineSpan], *,
                base_font_name: str | None = None,
                base_font_size=None,
                doc=None,
                footnote_defs: dict[str, str] | None = None,
                clear_existing: bool = True,
                image_base_path: str | None = None,
                prefer_real_image: bool = False) -> None:
    """将 InlineSpan 列表渲染为段落 runs。

    支持所有行内类型：TEXT, BOLD, ITALIC, CODE, HYPERLINK,
    IMAGE, FOOTNOTE_REF, LINE_BREAK, LATEX_INLINE 等。
    对 CODE / LATEX_INLINE 自动添加 noProof。
    HYPERLINK 使用真超链接。
    """
    if not spans:
        return

    if clear_existing:
        for run in para.runs:
            run._element.getparent().remove(run._element)

    for span in spans:
        if span.type == InlineType.LINE_BREAK:
            run = para.add_run()
            br = run._element.makeelement(qn('w:br'), {})
            run._element.append(br)
            continue

        if span.type == InlineType.HYPERLINK:
            add_hyperlink(para, span.text, span.url)
            continue

        if span.type == InlineType.IMAGE:
            if prefer_real_image:
                inserted = insert_image(
                    para, span.text, span.url, base_path=image_base_path
                )
                if inserted:
                    continue
            # 回退：斜体占位文本
            run = para.add_run(f"[{span.text}]")
            run.italic = True
            if base_font_name:
                run.font.name = base_font_name
            if base_font_size:
                run.font.size = base_font_size
            continue

        if span.type == InlineType.FOOTNOTE_REF:
            fn_id = span.footnote_id
            if doc is not None and footnote_defs:
                content = footnote_defs.get(fn_id)
                if content:
                    try:
                        add_footnote(doc, para, fn_id, content)
                        continue
                    except Exception as e:
                        logger.warning(f"Failed to add footnote {fn_id}: {e}")
                        # 在文档中标记失败
                        run = para.add_run(f"[脚注{fn_id}]")
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.superscript = True
                        continue

            run = para.add_run(fn_id or span.text or "")
            run.font.superscript = True
            if base_font_size:
                run.font.size = base_font_size
            continue

        # 普通 run
        run = para.add_run(span.text)
        if base_font_name:
            run.font.name = base_font_name
        if base_font_size:
            run.font.size = base_font_size

        _apply_inline_format(run, span)


def _apply_inline_format(run, span: InlineSpan) -> None:
    """根据 InlineType 设置 run 格式"""
    if span.type == InlineType.BOLD:
        run.bold = True
    elif span.type == InlineType.ITALIC:
        run.italic = True
    elif span.type == InlineType.BOLD_ITALIC:
        run.bold = True
        run.italic = True
    elif span.type == InlineType.STRIKETHROUGH:
        run.font.strike = True
    elif span.type == InlineType.CODE:
        run.font.name = "Consolas"
        set_no_proof(run)
    elif span.type == InlineType.LATEX_INLINE:
        run.font.name = "Consolas"
        set_no_proof(run)


# ── 表格单元格渲染 ──

_RE_BR_SPLIT = re.compile(r"<br\s*/?\s*>", re.IGNORECASE)


def _display_text_width(text: str) -> int:
    """Estimate visual width for markdown table column distribution."""
    raw = _RE_BR_SPLIT.sub("\n", str(text or ""))
    lines = raw.splitlines() or [raw]
    best = 1
    for line in lines:
        width = 0
        for ch in line.strip():
            if ch == "\t":
                width += 2
            elif unicodedata.east_asian_width(ch) in {"F", "W"}:
                width += 2
            else:
                width += 1
        best = max(best, width)
    return best


def _available_page_width_twips(doc) -> int:
    emu_per_twip = 635
    try:
        section = doc.sections[0]
        page_w = int(section.page_width or 0) // emu_per_twip
        left = int(section.left_margin or 0) // emu_per_twip
        right = int(section.right_margin or 0) // emu_per_twip
        avail = page_w - left - right
        if avail > 0:
            return avail
    except Exception:
        pass
    return 9000


def _table_col_widths(headers: list[str], rows: list[list[str]], total_twips: int) -> list[int]:
    col_count = len(headers)
    if col_count <= 0:
        return []
    if col_count == 1:
        return [max(1440, total_twips)]

    weights = [max(2.0, float(_display_text_width(h) + 2)) for h in headers]
    for row in rows:
        for idx in range(min(col_count, len(row))):
            weights[idx] = max(weights[idx], float(_display_text_width(row[idx]) + 2))

    sorted_weights = sorted(weights)
    median = sorted_weights[len(sorted_weights) // 2]
    cap = max(12.0, median * 3.0)
    weights = [min(w, cap) for w in weights]

    min_col = 640
    min_total = min_col * col_count
    if total_twips <= min_total:
        base = total_twips // col_count
        widths = [base] * col_count
        widths[0] += total_twips - sum(widths)
        return widths

    extra = total_twips - min_total
    weight_sum = sum(weights) or float(col_count)
    widths = [min_col + int(extra * (w / weight_sum)) for w in weights]
    widths[0] += total_twips - sum(widths)
    return widths


def _set_native_table_style(tbl_el) -> None:
    tbl_pr = tbl_el.find(_w("tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl_el, _w("tblPr"))
        tbl_el.insert(0, tbl_pr)
    tbl_style = tbl_pr.find(_w("tblStyle"))
    if tbl_style is not None:
        tbl_pr.remove(tbl_style)
    tbl_look = tbl_pr.find(_w("tblLook"))
    if tbl_look is not None:
        tbl_pr.remove(tbl_look)


def _clear_table_cell_spacing(tbl_el) -> None:
    tbl_pr = tbl_el.find(_w("tblPr"))
    if tbl_pr is not None:
        tbl_cell_spacing = tbl_pr.find(_w("tblCellSpacing"))
        if tbl_cell_spacing is not None:
            tbl_pr.remove(tbl_cell_spacing)

    # Some pasted Markdown tables store cell spacing at row level.
    for tr in tbl_el.findall(_w("tr")):
        tr_pr = tr.find(_w("trPr"))
        if tr_pr is None:
            continue
        tr_spacing = tr_pr.find(_w("tblCellSpacing"))
        if tr_spacing is not None:
            tr_pr.remove(tr_spacing)


def _clear_table_row_property_exceptions(tbl_el) -> None:
    # Remove per-row table property exceptions (tblPrEx), which can override
    # table-level borders and create dashed/double-outline artifacts.
    for tr in tbl_el.findall(_w("tr")):
        tbl_pr_ex = tr.find(_w("tblPrEx"))
        if tbl_pr_ex is not None:
            tr.remove(tbl_pr_ex)


def _clear_cell_border_overrides(tbl_el) -> None:
    for tc in tbl_el.iter(_w("tc")):
        tc_pr = tc.find(_w("tcPr"))
        if tc_pr is None:
            continue
        tc_borders = tc_pr.find(_w("tcBorders"))
        if tc_borders is not None:
            tc_pr.remove(tc_borders)


def _set_table_borders_single(tbl_el, size_eighth_pt: int = 4) -> None:
    tbl_pr = tbl_el.find(_w("tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl_el, _w("tblPr"))
        tbl_el.insert(0, tbl_pr)
    old = tbl_pr.find(_w("tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = etree.SubElement(tbl_pr, _w("tblBorders"))
    sz = str(max(1, int(size_eighth_pt)))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = etree.SubElement(borders, _w(side))
        border.set(_w("val"), "single")
        border.set(_w("sz"), sz)
        border.set(_w("space"), "0")
        border.set(_w("color"), "000000")


def _clear_table_paragraph_indents(tbl_el) -> None:
    for tc in tbl_el.iter(_w("tc")):
        for p in tc.iter(_w("p")):
            p_pr = p.find(_w("pPr"))
            if p_pr is None:
                p_pr = etree.SubElement(p, _w("pPr"))
                p.insert(0, p_pr)
            p_style = p_pr.find(_w("pStyle"))
            if p_style is not None:
                p_pr.remove(p_style)
            ind = p_pr.find(_w("ind"))
            if ind is None:
                ind = etree.SubElement(p_pr, _w("ind"))
            for attr in (
                    "left", "leftChars", "right", "rightChars",
                    "firstLine", "firstLineChars", "hanging", "hangingChars"):
                ind.set(_w(attr), "0")
            spacing = p_pr.find(_w("spacing"))
            if spacing is None:
                spacing = etree.SubElement(p_pr, _w("spacing"))
            spacing.set(_w("before"), "0")
            spacing.set(_w("after"), "0")
            spacing.set(_w("line"), "240")
            spacing.set(_w("lineRule"), "auto")


def _center_table_cells(tbl_el) -> None:
    for tc in tbl_el.iter(_w("tc")):
        tc_pr = tc.find(_w("tcPr"))
        if tc_pr is None:
            tc_pr = etree.SubElement(tc, _w("tcPr"))
            tc.insert(0, tc_pr)
        v_align = tc_pr.find(_w("vAlign"))
        if v_align is None:
            v_align = etree.SubElement(tc_pr, _w("vAlign"))
        v_align.set(_w("val"), "center")

        for p in tc.iter(_w("p")):
            p_pr = p.find(_w("pPr"))
            if p_pr is None:
                p_pr = etree.SubElement(p, _w("pPr"))
                p.insert(0, p_pr)
            jc = p_pr.find(_w("jc"))
            if jc is None:
                jc = etree.SubElement(p_pr, _w("jc"))
            jc.set(_w("val"), "center")


def _set_table_alignment_center(tbl_el) -> None:
    tbl_pr = tbl_el.find(_w("tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl_el, _w("tblPr"))
        tbl_el.insert(0, tbl_pr)
    jc = tbl_pr.find(_w("jc"))
    if jc is None:
        jc = etree.SubElement(tbl_pr, _w("jc"))
    jc.set(_w("val"), "center")
    tbl_ind = tbl_pr.find(_w("tblInd"))
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)


def _set_table_widths(tbl_el, widths: list[int]) -> None:
    if not widths:
        return
    total = max(1, int(sum(widths)))
    col_count = len(widths)

    tbl_pr = tbl_el.find(_w("tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl_el, _w("tblPr"))
        tbl_el.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(_w("tblW"))
    if tbl_w is None:
        tbl_w = etree.SubElement(tbl_pr, _w("tblW"))
    tbl_w.set(_w("type"), "dxa")
    tbl_w.set(_w("w"), str(total))

    layout = tbl_pr.find(_w("tblLayout"))
    if layout is None:
        layout = etree.SubElement(tbl_pr, _w("tblLayout"))
    layout.set(_w("type"), "fixed")

    tbl_grid = tbl_el.find(_w("tblGrid"))
    if tbl_grid is None:
        tbl_grid = etree.SubElement(tbl_el, _w("tblGrid"))
        tbl_pr.addnext(tbl_grid)
    else:
        for col in list(tbl_grid):
            tbl_grid.remove(col)
    for w in widths:
        grid_col = etree.SubElement(tbl_grid, _w("gridCol"))
        grid_col.set(_w("w"), str(max(1, int(w))))

    for tr in tbl_el.findall(_w("tr")):
        tcs = tr.findall(_w("tc"))
        for idx, tc in enumerate(tcs):
            width = widths[min(idx, col_count - 1)]
            tc_pr = tc.find(_w("tcPr"))
            if tc_pr is None:
                tc_pr = etree.SubElement(tc, _w("tcPr"))
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(_w("tcW"))
            if tc_w is None:
                tc_w = etree.SubElement(tc_pr, _w("tcW"))
            tc_w.set(_w("type"), "dxa")
            tc_w.set(_w("w"), str(max(1, int(width))))


def normalize_markdown_table(doc, table, headers: list[str], rows: list[list[str]]) -> None:
    """Apply Word-native table style + content-aware width allocation."""
    try:
        _set_native_table_style(table._tbl)
        _clear_table_cell_spacing(table._tbl)
        _clear_table_row_property_exceptions(table._tbl)
        _clear_cell_border_overrides(table._tbl)
        _set_table_borders_single(table._tbl, size_eighth_pt=4)
        _set_table_alignment_center(table._tbl)
        total_w = _available_page_width_twips(doc)
        widths = _table_col_widths(headers, rows, total_w)
        _set_table_widths(table._tbl, widths)
        _clear_table_paragraph_indents(table._tbl)
        _center_table_cells(table._tbl)
    except Exception as exc:
        logger.debug("normalize_markdown_table skipped: %s", exc)


def normalize_markdown_table_visual_only(table) -> None:
    """Apply markdown table border/alignment normalization without resizing columns."""
    try:
        _set_native_table_style(table._tbl)
        _clear_table_cell_spacing(table._tbl)
        _clear_table_row_property_exceptions(table._tbl)
        _clear_cell_border_overrides(table._tbl)
        _set_table_borders_single(table._tbl, size_eighth_pt=4)
        _set_table_alignment_center(table._tbl)
        _clear_table_paragraph_indents(table._tbl)
        _center_table_cells(table._tbl)
    except Exception as exc:
        logger.debug("normalize_markdown_table_visual_only skipped: %s", exc)


def write_table_cell(
        cell,
        text: str,
        *,
        image_base_path: str | None = None,
        prefer_real_image: bool = False) -> None:
    """渲染表格单元格：支持 <br> 换行和行内格式。

    将 cell.text = ... 替换为带格式的渲染。
    <br> 在单元格内产生新段落而非软换行。
    """
    # 清除默认段落文本
    cell.text = ""
    para = cell.paragraphs[0]

    # 按 <br> 分割为多段
    parts = _RE_BR_SPLIT.split(text)

    for idx, part in enumerate(parts):
        if idx > 0:
            para = cell.add_paragraph()
        spans = parse_inline(part.strip())
        write_spans(
            para,
            spans,
            clear_existing=False,
            image_base_path=image_base_path,
            prefer_real_image=prefer_real_image,
        )


# ── 图片插入 ──

def insert_image(para, alt: str, src: str,
                 base_path: str | None = None,
                 max_width_cm: float = 14.0) -> bool:
    """插入本地图片到段落，不存在则回退为斜体文本。

    Returns: True 如果成功插入图片，False 如果回退文本。
    """
    # 解析图片路径并验证安全性
    try:
        if base_path:
            base = Path(base_path).resolve()
            img_path = (base / src).resolve()
            # 验证路径在允许范围内
            try:
                img_path.relative_to(base)
            except ValueError:
                logger.warning(f"Image path outside base directory: {src}")
                run = para.add_run(f"[图片路径不安全: {alt or src}]")
                run.italic = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                return False
        else:
            img_path = Path(src).resolve()
    except Exception as e:
        logger.warning(f"Invalid image path {src}: {e}")
        run = para.add_run(f"[图片路径无效: {alt or src}]")
        run.italic = True
        return False

    if not img_path.exists():
        logger.debug(f"Image not found: {img_path}")
        run = para.add_run(f"[图片: {alt or src}]")
        run.italic = True
        return False

    try:
        run = para.add_run()
        run.add_picture(str(img_path), width=Cm(max_width_cm))
        apply_safe_picture_line_spacing(para)
        return True
    except Exception as e:
        logger.warning(f"Failed to insert image {img_path}: {e}")
        run = para.add_run(f"[图片加载失败: {alt or src}]")
        run.italic = True
        return False


# ── 脚注 ──

def add_footnote(doc, para, fn_id: str, content: str) -> None:
    """在段落当前位置插入 Word 脚注。

    fn_id: 脚注标识符（如 "1", "note"）
    content: 脚注正文内容
    """
    # 确保 footnotes.xml part 存在（python-docx 无高层 API）
    _get_footnotes_element(doc)

    # 分配脚注 ID
    footnote_id = _next_footnote_id(doc)

    # 在 footnotes.xml 中添加脚注内容
    _add_footnote_content(doc, footnote_id, content)

    # 在段落中插入脚注引用标记
    _add_footnote_reference(para, footnote_id)


def _get_footnotes_element(doc):
    """获取 footnotes.xml 的根元素，不存在则创建"""
    FN_URI = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    package = doc.part.package
    # 在 document part 的关系中查找 footnotes part
    for rel in doc.part.rels.values():
        if rel.reltype != FN_URI:
            continue
        target = rel.target_part
        if hasattr(target, "element"):
            return target.element
        if hasattr(target, "_element"):
            return target._element
        if hasattr(target, "blob"):
            el = etree.fromstring(target.blob)
            target._element = el
            return el

    # 不存在，创建 footnotes part
    from docx.opc.packuri import PackURI
    from docx.opc.part import XmlPart

    fn_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:footnote w:type="separator" w:id="-1">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>'
        '</w:footnotes>'
    )
    fn_part = XmlPart.load(
        PackURI("/word/footnotes.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        fn_xml.encode('utf-8'),
        package,
    )
    doc.part.relate_to(fn_part, FN_URI)
    return fn_part.element


def _next_footnote_id(doc) -> int:
    """获取下一个可用的脚注 ID"""
    try:
        fn_el = _get_footnotes_element(doc)
    except Exception:
        return 1
    max_id = 0
    for fn in fn_el.findall(_w("footnote")):
        fid = fn.get(_w("id"))
        if fid and fid.lstrip('-').isdigit():
            max_id = max(max_id, int(fid))
    return max_id + 1


def _add_footnote_content(doc, footnote_id: int, content: str) -> None:
    """在 footnotes.xml 中添加脚注内容"""
    fn_el = _get_footnotes_element(doc)

    footnote = etree.SubElement(fn_el, _w("footnote"))
    footnote.set(_w("id"), str(footnote_id))

    # 脚注段落
    p = etree.SubElement(footnote, _w("p"))
    ppr = etree.SubElement(p, _w("pPr"))
    pstyle = etree.SubElement(ppr, _w("pStyle"))
    pstyle.set(_w("val"), "FootnoteText")

    # 脚注引用标记（脚注区域内的上标数字）
    ref_run = etree.SubElement(p, _w("r"))
    ref_rpr = etree.SubElement(ref_run, _w("rPr"))
    ref_style = etree.SubElement(ref_rpr, _w("rStyle"))
    ref_style.set(_w("val"), "FootnoteReference")
    etree.SubElement(ref_run, _w("footnoteRef"))

    # 空格分隔
    sep_run = etree.SubElement(p, _w("r"))
    sep_t = etree.SubElement(sep_run, _w("t"))
    sep_t.text = " "
    sep_t.set(qn('xml:space'), 'preserve')

    # 脚注正文
    text_run = etree.SubElement(p, _w("r"))
    text_t = etree.SubElement(text_run, _w("t"))
    text_t.text = content
    text_t.set(qn('xml:space'), 'preserve')


def _add_footnote_reference(para, footnote_id: int) -> None:
    """在段落中插入脚注引用标记（正文中的上标数字）"""
    run_el = etree.SubElement(para._element, _w("r"))
    rpr = etree.SubElement(run_el, _w("rPr"))
    rstyle = etree.SubElement(rpr, _w("rStyle"))
    rstyle.set(_w("val"), "FootnoteReference")

    fn_ref = etree.SubElement(run_el, _w("footnoteReference"))
    fn_ref.set(_w("id"), str(footnote_id))

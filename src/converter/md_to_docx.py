"""独立 Markdown → DOCX 转换器"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.markdown.block_parser import parse_markdown_text
from src.markdown.ir import MarkdownBlock, BlockType, InlineSpan, InlineType
from src.markdown.word_render import (
    write_spans, set_no_proof, add_hyperlink, write_table_cell,
    normalize_ordered_list_style, normalize_unordered_list_style,
    register_list_numbering, apply_num_pr, apply_blockquote_border,
    insert_image, add_footnote,
)
from src.scene.schema import SceneConfig, StyleConfig
from src.engine.change_tracker import ChangeTracker


HEADING_STYLE_MAP = {
    1: "Heading 1", 2: "Heading 2",
    3: "Heading 3", 4: "Heading 4",
    5: "Heading 4", 6: "Heading 4",
}


def _resolve_list_marker_separator(config: SceneConfig) -> str:
    md_cfg = getattr(config, "md_cleanup", None)
    raw = str(getattr(md_cfg, "list_marker_separator", "tab") or "").strip().lower()
    if raw in {"half_space", "space", "halfwidth_space"}:
        return "half_space"
    if raw in {"full_space", "fullwidth_space"}:
        return "full_space"
    return "tab"


def _list_separator_text(separator: str) -> str:
    if separator == "full_space":
        return "\u3000"
    if separator == "half_space":
        return " "
    return "\t"


def _resolve_ordered_list_style(config: SceneConfig) -> str:
    md_cfg = getattr(config, "md_cleanup", None)
    raw = str(getattr(md_cfg, "ordered_list_style", "mixed") or "").strip().lower()
    return normalize_ordered_list_style(raw)


def _resolve_unordered_list_style(config: SceneConfig) -> str:
    md_cfg = getattr(config, "md_cleanup", None)
    raw = str(getattr(md_cfg, "unordered_list_style", "word_default") or "").strip().lower()
    return normalize_unordered_list_style(raw)


def _is_ordered_list_marker(marker: str) -> bool:
    s = str(marker or "").strip()
    if not s:
        return False
    return bool(re.match(r"^(?:\d+[.)、]|[（(]\d+[)）])$", s))


def convert_md_to_docx(md_path: str, output_path: str,
                       config: SceneConfig) -> dict:
    """读取 .md 文件，创建格式化的 .docx

    Returns: {"success": bool, "error": str|None, "output": str}
    """
    md_path = Path(md_path)
    if not md_path.exists():
        return {"success": False, "error": f"文件不存在: {md_path}"}

    try:
        text = md_path.read_text(encoding="utf-8")
    except Exception as e:
        return {"success": False, "error": f"读取失败: {e}"}

    blocks = parse_markdown_text(text)
    doc = Document()
    tracker = ChangeTracker()

    # 预注册列表编号
    marker_separator = _resolve_list_marker_separator(config)
    ordered_list_style = _resolve_ordered_list_style(config)
    unordered_list_style = _resolve_unordered_list_style(config)
    bullet_num_id = register_list_numbering(
        doc,
        "bullet",
        marker_separator=marker_separator,
        unordered_style=unordered_list_style,
    )
    decimal_num_id = register_list_numbering(
        doc,
        "decimal",
        marker_separator=marker_separator,
        ordered_style=ordered_list_style,
    )

    # 收集脚注定义
    footnote_defs: dict[str, str] = {}
    for block in blocks:
        if block.type == BlockType.FOOTNOTE_DEF:
            footnote_defs[block.list_marker] = block.raw_text

    ctx = {
        "bullet_num_id": bullet_num_id,
        "decimal_num_id": decimal_num_id,
        "footnote_defs": footnote_defs,
        "md_base_path": str(md_path.parent),
        "list_marker_separator": marker_separator,
        "ordered_list_style": ordered_list_style,
        "unordered_list_style": unordered_list_style,
    }

    for block in blocks:
        if block.type == BlockType.FOOTNOTE_DEF:
            continue  # 脚注定义不输出为段落
        _write_block(doc, block, config, tracker, ctx)

    out = Path(output_path)
    doc.save(str(out))
    return {"success": True, "error": None, "output": str(out)}


def _write_block(doc: Document, block: MarkdownBlock,
                 config: SceneConfig, tracker: ChangeTracker, ctx: dict):
    """将单个 block 写入 Document"""
    if block.type == BlockType.BLANK:
        return
    elif block.type == BlockType.HEADING:
        _write_heading(doc, block, config)
    elif block.type == BlockType.CODE_BLOCK:
        _write_code_block(doc, block, config)
    elif block.type == BlockType.TABLE:
        _write_table(doc, block)
    elif block.type == BlockType.BLOCKQUOTE:
        _write_blockquote(doc, block, config)
    elif block.type == BlockType.HORIZONTAL_RULE:
        _write_hr(doc)
    elif block.type == BlockType.LIST_ITEM:
        _write_list_item(doc, block, config, ctx)
    elif block.type == BlockType.TASK_LIST_ITEM:
        _write_task_list(doc, block, config)
    elif block.type == BlockType.PARAGRAPH:
        _write_paragraph(doc, block, config, ctx)
    elif block.type == BlockType.LATEX_BLOCK:
        _write_latex_block(doc, block)


def _write_heading(doc, block, config):
    """写入标题"""
    style_name = HEADING_STYLE_MAP.get(block.level, "Heading 4")
    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    write_spans(para, block.spans, clear_existing=False)


def _write_code_block(doc, block, config):
    """写入代码块"""
    code_sc = config.styles.get("code_block")
    font = code_sc.font_en if code_sc else "Consolas"
    size = code_sc.size_pt if code_sc else 10
    for line in block.code_lines:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.left_indent = Cm(1.0)
        run = para.add_run(line)
        run.font.name = font
        run.font.size = Pt(size)
        set_no_proof(run)


def _write_table(doc, block):
    """写入表格"""
    num_cols = len(block.table_headers)
    num_rows = 1 + len(block.table_rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    for j, header in enumerate(block.table_headers):
        if j < num_cols:
            cell = table.rows[0].cells[j]
            write_table_cell(cell, header)
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for i, row_data in enumerate(block.table_rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                write_table_cell(table.rows[i + 1].cells[j], cell_text)

    # 应用列对齐
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }
    if block.table_alignments:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(block.table_alignments):
                    align = align_map.get(block.table_alignments[j])
                    if align and cell.paragraphs:
                        cell.paragraphs[0].alignment = align

def _write_blockquote(doc, block, config):
    """写入引用块（支持嵌套层级，带左边框+底色）"""
    para = doc.add_paragraph()
    apply_blockquote_border(para, block.quote_level)
    write_spans(para, block.spans, clear_existing=False)


def _write_hr(doc):
    """写入水平线"""
    para = doc.add_paragraph()
    para.add_run("─" * 40)


def _write_list_item(doc, block, config, ctx):
    """写入列表项（Word 原生编号）"""
    para = doc.add_paragraph()
    # 判断有序/无序
    is_ordered = _is_ordered_list_marker(block.list_marker)
    num_id = ctx["decimal_num_id"] if is_ordered else ctx["bullet_num_id"]
    apply_num_pr(para, num_id, block.list_level)
    spans = list(block.spans or [])
    if ctx.get("list_marker_separator") == "full_space":
        if not (
            spans
            and spans[0].type == InlineType.TEXT
            and str(spans[0].text or "").startswith("\u3000")
        ):
            spans = [InlineSpan(InlineType.TEXT, _list_separator_text("full_space"))] + spans
    write_spans(para, spans, clear_existing=False)


def _write_task_list(doc, block, config):
    """写入任务清单项"""
    para = doc.add_paragraph()
    marker_separator = _resolve_list_marker_separator(config)
    prefix = f"{'☑' if block.checked else '☐'}{_list_separator_text(marker_separator)}"
    para.add_run(prefix)
    write_spans(para, block.spans, clear_existing=False)


def _write_paragraph(doc, block, config, ctx):
    """写入普通段落（支持图片和脚注）"""
    # 检查是否整段是图片
    if (len(block.spans) == 1
            and block.spans[0].type == InlineType.IMAGE):
        span = block.spans[0]
        para = doc.add_paragraph()
        base = ctx.get("md_base_path")
        insert_image(para, span.text, span.url, base_path=base)
        return

    para = doc.add_paragraph()
    # 处理脚注引用
    fn_defs = ctx.get("footnote_defs", {})
    has_footnotes = any(
        s.type == InlineType.FOOTNOTE_REF for s in block.spans)
    if has_footnotes and fn_defs:
        _write_spans_with_footnotes(doc, para, block.spans, fn_defs)
    else:
        write_spans(para, block.spans, clear_existing=False)


def _write_latex_block(doc, block):
    """写入 LaTeX 块（保留原文，防自动更正）"""
    para = doc.add_paragraph()
    run = para.add_run(block.raw_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    set_no_proof(run)


def _write_spans_with_footnotes(doc, para, spans, fn_defs: dict):
    """渲染 spans，遇到 FOOTNOTE_REF 时插入 Word 脚注"""
    from src.markdown.word_render import _apply_inline_format

    for span in spans:
        if span.type == InlineType.FOOTNOTE_REF:
            content = fn_defs.get(span.footnote_id, "")
            if content:
                add_footnote(doc, para, span.footnote_id, content)
            else:
                run = para.add_run(f"[^{span.footnote_id}]")
                run.font.superscript = True
        elif span.type == InlineType.LINE_BREAK:
            run = para.add_run()
            br = run._element.makeelement(qn('w:br'), {})
            run._element.append(br)
        elif span.type == InlineType.HYPERLINK:
            add_hyperlink(para, span.text, span.url)
        else:
            run = para.add_run(span.text)
            _apply_inline_format(run, span)

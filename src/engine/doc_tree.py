"""文档结构树：将 docx 段落划分为逻辑分区（多信号加权识别 + 顺序校验）"""

import re
from dataclasses import dataclass
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from src.utils.toc_entry import (
    looks_like_toc_entry_line,
    looks_like_date_placeholder_line,
    looks_like_reference_entry_line,
    looks_like_numbered_toc_entry_with_page_suffix,
)


# 分区锚点关键字
SECTION_ANCHORS: dict[str, list[str]] = {
    "cover": ["学位论文", "博士学位", "硕士学位"],
    "abstract_cn": ["摘要", "摘 要"],
    "abstract_en": ["Abstract", "ABSTRACT"],
    "toc": ["目录", "目 录"],
    "references": ["参考文献"],
    "errata": ["勘误页", "勘误"],
    "appendix": ["附录"],
    "acknowledgment": ["致谢", "致 谢"],
    "resume": ["简历", "个人简历", "在学期间"],
}

# 分区出现的典型顺序（用于顺序校验）
SECTION_ORDER = [
    "cover", "abstract_cn", "abstract_en", "toc",
    "body", "references", "errata", "appendix",
    "acknowledgment", "resume",
]

# 识别阈值：低于此分数的候选不采纳
MIN_SCORE = 4

_PRE_BODY_TYPES = {"cover", "toc", "abstract_cn", "abstract_en"}
_POST_BODY_TYPES = {"references", "errata", "appendix", "acknowledgment", "resume"}

# 分区最小可信分（避免“弱匹配 + 位置侥幸”误切分）
_SECTION_MIN_ACCEPT_SCORE: dict[str, float] = {
    "cover": 8.0,
    "abstract_cn": 8.0,
    "abstract_en": 8.0,
    "toc": 8.0,
    "references": 8.0,
    "errata": 8.0,
    "appendix": 8.0,
    "acknowledgment": 8.0,
    "resume": 8.0,
}

# 标题行“可直接按标题处理”的最低置信度
_SECTION_TITLE_CONFIDENCE: dict[str, float] = {
    "cover": 10.0,
    "abstract_cn": 10.0,
    "abstract_en": 10.0,
    "toc": 10.0,
    "references": 10.0,
    "errata": 10.0,
    "appendix": 10.0,
    "acknowledgment": 10.0,
    "resume": 10.0,
}

_POST_BODY_MIN_RATIO = 0.35
_PRE_BODY_MAX_RATIO = 0.70
_HIGH_CONF_OVERRIDE = 12.0

_RE_REFERENCE_ENTRY = [
    re.compile(r"^\s*\[\d{1,4}\]\s*\S"),       # [1] ...
    re.compile(r"^\s*[（(]\d{1,4}[）)]\s*\S"),  # (1) ...
    re.compile(r"^\s*\d{1,4}\.\s+\S"),          # 1. ...
]
_REFERENCE_ENTRY_CLUSTER_MAX_LEN = 240

_FRONT_TITLE_NORMS = {
    "\u6458\u8981",
    "\u6458\u8981\u3002",
    "\u6458\u8981\uff1a",
    "abstract",
    "\u76ee\u5f55",
    "\u76ee\u9304",
    "contents",
    "tableofcontents",
}
_RE_FRONT_TITLE_TAIL_MARKS = re.compile(r"[：:;；·•\-—_~\.。…]+$")
_RE_FRONT_TITLE_PATTERNS = [
    re.compile(r"^\u6458\u8981(?:[（(][^()（）]{0,8}[)）])?$"),
    re.compile(r"^abstract(?:[（(][^()（）]{0,16}[)）])?$", re.IGNORECASE),
    re.compile(
        r"^(?:\u76ee\u5f55|\u76ee\u9304|contents|tableofcontents)(?:[（(][^()（）]{0,8}[)）])?$",
        re.IGNORECASE,
    ),
]

@dataclass
class DocSection:
    section_type: str
    start_index: int
    end_index: int = -1
    confidence: float = 0.0  # 识别置信度得分
    title_confident: bool = True

    @property
    def paragraph_range(self) -> range:
        return range(self.start_index, self.end_index + 1) if self.end_index >= 0 else range(0)


@dataclass
class _Candidate:
    """分区候选"""
    para_index: int
    section_type: str
    score: float


def _is_reference_entry_text(text: str) -> bool:
    return looks_like_reference_entry_line(text)


def _is_reference_entry_content_text(text: str) -> bool:
    raw = (text or "").strip()
    if not raw:
        return False
    if len(raw) > _REFERENCE_ENTRY_CLUSTER_MAX_LEN:
        return False
    return _is_reference_entry_text(raw)


def _norm_no_space_lower(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip().lower()


def _norm_front_title_token(text: str) -> str:
    norm = _norm_no_space_lower(text)
    if not norm:
        return ""
    return _RE_FRONT_TITLE_TAIL_MARKS.sub("", norm)


def _is_front_matter_title_text(text: str) -> bool:
    norm = _norm_front_title_token(text)
    if not norm:
        return False
    if norm in _FRONT_TITLE_NORMS:
        return True
    for pat in _RE_FRONT_TITLE_PATTERNS:
        if pat.match(norm):
            return True
    return False


def _is_section_title_confident(section_type: str, score: float) -> bool:
    return score >= _SECTION_TITLE_CONFIDENCE.get(section_type, 10.0)


def _is_candidate_credible(c: _Candidate, total: int) -> bool:
    min_score = _SECTION_MIN_ACCEPT_SCORE.get(c.section_type, MIN_SCORE)
    if c.score < min_score:
        return False

    ratio = c.para_index / max(total, 1)
    if (c.section_type in _POST_BODY_TYPES
            and ratio < _POST_BODY_MIN_RATIO
            and c.score < _HIGH_CONF_OVERRIDE):
        return False
    if (c.section_type in _PRE_BODY_TYPES
            and ratio > _PRE_BODY_MAX_RATIO
            and c.score < _HIGH_CONF_OVERRIDE):
        return False
    return True


def _order_compatible(prev_type: str, cur_type: str, order_rank: dict[str, int]) -> bool:
    """Return True when two section types can appear in this order.

    Pre-body blocks (cover/abstract/toc) are treated as an unordered group
    before body, so internal order can vary by template.
    """
    if prev_type == cur_type:
        return False

    prev_is_pre = prev_type in _PRE_BODY_TYPES
    cur_is_pre = cur_type in _PRE_BODY_TYPES
    if prev_is_pre and cur_is_pre:
        return True
    if prev_is_pre and not cur_is_pre:
        return True
    if not prev_is_pre and cur_is_pre:
        return False

    # Post-body blocks can appear in different template-specific orders
    # (e.g. appendix before references, or acknowledgment before appendix).
    prev_is_post = prev_type in _POST_BODY_TYPES
    cur_is_post = cur_type in _POST_BODY_TYPES
    if prev_is_post and cur_is_post:
        return True
    if prev_is_post and not cur_is_post:
        return False
    if not prev_is_post and cur_is_post:
        return True

    return order_rank.get(prev_type, 10_000) < order_rank.get(cur_type, 10_000)


def _score_paragraph(para: Paragraph, sec_type: str,
                     anchors: list[str], total: int,
                     para_index: int) -> float:
    """对段落进行多信号加权评分，判断其作为某分区标题的可能性。

    信号:
      - 精确匹配（去空格后 == 关键字）: +10
      - 前缀匹配（以关键字开头，后面可跟空格/标点）: +6
      - 子串包含: +2
      - Heading 样式加分: +5
      - 加粗 + 大字号加分: +3
      - 独占段落（文本长度 ≤ 关键字长度+4）: +3
      - 位置加分（符合预期位置区间）: +2
    """
    text = para.text.strip()
    if not text or len(text) > 80:
        return 0.0
    if sec_type != "toc" and (
        looks_like_toc_entry_line(text)
        or looks_like_numbered_toc_entry_with_page_suffix(text)
    ):
        return 0.0

    score = 0.0
    matched_anchor = None

    # ── 文本匹配 ──
    text_no_space = re.sub(r'\s+', '', text)
    best_text_score = 0.0
    for anchor in anchors:
        anchor_no_space = re.sub(r'\s+', '', anchor)
        local_score = 0.0
        if text_no_space == anchor_no_space:
            local_score = 10.0
        elif text_no_space.startswith(anchor_no_space):
            local_score = 6.0
            if (
                sec_type == "appendix"
                and re.match(r"^附录[A-Za-zＡ-Ｚａ-ｚ0-9一二三四五六七八九十]", text_no_space)
            ):
                # Treat "附录A/附录B/附录1" title forms as strong appendix anchors.
                local_score = max(local_score, 10.0)
        elif anchor_no_space and anchor_no_space in text_no_space:
            local_score = 2.0
        if local_score > best_text_score:
            best_text_score = local_score
            matched_anchor = anchor
    score += best_text_score

    if not matched_anchor:
        return 0.0

    # ── 样式加分 ──
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading") or style_name.startswith("heading"):
        score += 5

    # ── 格式加分（加粗 + 大字号）──
    if para.runs:
        first_run = para.runs[0]
        is_bold = first_run.bold or (first_run.font.bold is True)
        font_size = first_run.font.size
        if is_bold and font_size and font_size >= Pt(14):
            score += 3

    # ── 独占段落加分 ──
    if len(text) <= len(matched_anchor) + 4:
        score += 3

    # ── 位置加分（仅当文本匹配较强时才给位置加分）──
    # 子串包含(+2)太弱，不应仅靠位置凑到阈值
    if score >= 6:
        ratio = para_index / max(total, 1)
        if sec_type in ("cover",) and ratio < 0.15:
            score += 2
        elif sec_type in ("abstract_cn", "abstract_en") and ratio < 0.35:
            score += 2
        elif sec_type == "toc" and ratio < 0.60:
            # TOC often appears after CN/EN abstracts in thesis templates.
            score += 2
        elif sec_type in ("references", "errata", "acknowledgment", "resume") and ratio > 0.5:
            score += 2

    return score


class DocTree:
    def __init__(self):
        self.sections: list[DocSection] = []
        self._detection_log: list[str] = []  # 识别日志，供调试/报告
        self._doc_ref = None  # 供 _infer_body 回退扫描

    def build(self, doc: Document, *,
              body_start_index: int | None = None) -> None:
        """扫描文档段落，按多信号加权识别分区。

        Args:
            body_start_index: 手动指定正文起始段落索引。
                若提供，则跳过自动识别 body 边界，直接使用该值。
        """
        self.sections.clear()
        self._detection_log.clear()
        self._doc_ref = doc
        total = len(doc.paragraphs)

        if body_start_index is not None:
            # ── 手动模式：用户指定正文起始位置 ──
            self._build_manual(doc, total, body_start_index)
        else:
            # ── 自动模式：多信号加权识别 ──
            self._build_auto(doc, total)

    def _build_manual(self, doc: Document, total: int,
                      body_start: int) -> None:
        """手动模式：以用户指定的段落索引为正文起点"""
        body_start = max(0, min(body_start, total - 1))

        # 正文之前的部分标记为 pre_body（不排版）
        if body_start > 0:
            self.sections.append(DocSection("pre_body", 0, body_start - 1))

        # 从 body_start 开始，仍然尝试识别后半部分的分区锚点
        post_anchors = self._scan_candidates(doc, total,
                                             start_from=body_start)
        # 只保留后半部分的分区（references, errata, appendix, acknowledgment, resume）
        post_best = self._pick_best_candidates(post_anchors, _POST_BODY_TYPES, total)

        # 后置分区补偿：即使“参考文献”标题缺失，也尽量通过内容识别参考文献起点
        if "references" not in post_best:
            ref_by_content = self._detect_references_by_content(
                doc, total, scan_start=body_start
            )
            if ref_by_content:
                post_best["references"] = ref_by_content
                self._detection_log.append(
                    f"手动模式: 通过参考文献条目识别 references @#{ref_by_content.start_index}"
                )

        # 按顺序校验并构建后半部分分区
        ordered = self._order_validate(post_best)

        if ordered:
            # body: body_start → 第一个后分区之前
            body_end = ordered[0].start_index - 1
        else:
            body_end = total - 1

        if body_start <= body_end:
            self.sections.append(DocSection("body", body_start, body_end))

        # 添加后半部分分区
        for i, sec in enumerate(ordered):
            if i + 1 < len(ordered):
                sec.end_index = ordered[i + 1].start_index - 1
            else:
                sec.end_index = total - 1
            self.sections.append(sec)

        self.sections.sort(key=lambda s: s.start_index)
        self._detection_log.append(
            f"手动模式: body 从段落 #{body_start} 开始")

    def _build_auto(self, doc: Document, total: int) -> None:
        """自动模式：多信号加权识别所有分区"""
        candidates = self._scan_candidates(doc, total)

        if not candidates:
            # 无文本锚点时，回退到 TOC 样式检测 + 参考文献条目检测
            toc_sec = self._detect_toc_by_style(doc, total)
            ref_sec = self._detect_references_by_content(doc, total)
            if toc_sec is not None:
                self.sections.append(toc_sec)
            if ref_sec is not None:
                self.sections.append(ref_sec)
            if self.sections:
                self._infer_body(total)
                self.sections.sort(key=lambda s: s.start_index)
                self._detection_log.append("无标题锚点：通过 TOC/参考文献内容识别构建分区")
                return

            self.sections.append(DocSection("body", 0, total - 1))
            self._detection_log.append("未找到任何分区锚点，整个文档视为 body")
            return

        all_types = set(SECTION_ANCHORS.keys())
        best = self._pick_best_candidates(candidates, all_types, total)

        # 样式补充：如果文本锚点未检测到 toc，尝试用 toc 样式检测
        if "toc" not in best:
            toc_sec = self._detect_toc_by_style(doc, total)
            if toc_sec:
                best["toc"] = toc_sec

        # 内容补偿：若 references 缺失/可疑，使用“连续参考条目”兜底
        ref_by_content = self._detect_references_by_content(doc, total)
        if ref_by_content is not None:
            old_ref = best.get("references")
            if old_ref is None:
                best["references"] = ref_by_content
            else:
                old_ratio = old_ref.start_index / max(total, 1)
                if (old_ratio < _POST_BODY_MIN_RATIO and ref_by_content.start_index > old_ref.start_index
                        or not old_ref.title_confident):
                    best["references"] = ref_by_content
                    self._detection_log.append(
                        f"references 纠偏: 由 #{old_ref.start_index} -> #{ref_by_content.start_index}"
                    )

        ordered = self._order_validate(best)

        # 构建分区边界
        for i, sec in enumerate(ordered):
            if i + 1 < len(ordered):
                sec.end_index = ordered[i + 1].start_index - 1
            else:
                sec.end_index = total - 1
            self.sections.append(sec)

        # 推断 body 分区
        self._infer_body(total)
        self.sections.sort(key=lambda s: s.start_index)

    def _detect_toc_by_style(self, doc: Document, total: int) -> DocSection | None:
        """通过 Word 样式检测目录区域（toc 1/2/3 等样式）"""
        toc_start = None
        toc_end = None
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            style_id = getattr(para.style, "style_id", "") if para.style else ""
            style_text = f"{style_name} {style_id}".strip()
            if re.search(r'(toc|目录)\s*\d+', style_text, re.IGNORECASE):
                if toc_start is None:
                    toc_start = i
                toc_end = i
        if toc_start is not None:
            self._detection_log.append(
                f"toc 样式检测: [{toc_start}, {toc_end}]")
            return DocSection("toc", toc_start, confidence=10.0, title_confident=True)
        return None

    def _scan_candidates(self, doc: Document, total: int,
                         *, start_from: int = 0) -> list[_Candidate]:
        """扫描段落，为每个段落对每种分区类型评分"""
        candidates: list[_Candidate] = []
        for i, para in enumerate(doc.paragraphs):
            if i < start_from:
                continue
            text = para.text.strip()
            if not text:
                continue
            for sec_type, anchors in SECTION_ANCHORS.items():
                score = _score_paragraph(para, sec_type, anchors,
                                         total, i)
                if score >= MIN_SCORE:
                    candidates.append(_Candidate(i, sec_type, score))
        return candidates

    def _pick_best_candidates(
            self,
            candidates: list[_Candidate],
            allowed_types: set[str],
            total: int) -> dict[str, DocSection]:
        """对每种分区类型，选取“可信且最优”的候选。"""
        best: dict[str, _Candidate] = {}
        for c in candidates:
            if c.section_type not in allowed_types:
                continue
            if not _is_candidate_credible(c, total):
                self._detection_log.append(
                    f"候选丢弃: {c.section_type} @#{c.para_index}, score={c.score:.1f} (低可信)"
                )
                continue

            if c.section_type not in best:
                best[c.section_type] = c
                continue

            cur = best[c.section_type]
            if c.score > cur.score:
                best[c.section_type] = c
                continue

            # 同分时：前置分区更偏向文档前部，后置分区更偏向文档尾部
            if c.score == cur.score:
                if c.section_type in _PRE_BODY_TYPES and c.para_index < cur.para_index:
                    best[c.section_type] = c
                elif c.section_type == "appendix":
                    # Appendix section should start from the first appendix title.
                    if c.para_index < cur.para_index:
                        best[c.section_type] = c
                elif c.section_type in _POST_BODY_TYPES and c.para_index > cur.para_index:
                    best[c.section_type] = c

        return {
            st: DocSection(
                st,
                c.para_index,
                confidence=c.score,
                title_confident=_is_section_title_confident(st, c.score),
            )
            for st, c in best.items()
        }

    def _order_validate(self, best: dict[str, DocSection]) -> list[DocSection]:
        """按“文中位置 + 标准分区顺序”选择高置信一致子序列。"""
        # cover 始终从段落 0 开始
        if "cover" in best:
            best["cover"].start_index = 0

        sorted_secs = sorted(best.values(), key=lambda s: s.start_index)
        if not sorted_secs:
            return []

        order_rank = {name: idx for idx, name in enumerate(SECTION_ORDER)}
        ranks = [order_rank.get(sec.section_type, 10_000) for sec in sorted_secs]

        # DP: 在保持段落索引递增的前提下，选取“顺序不回退”的最高置信子序列
        n = len(sorted_secs)
        dp = [sorted_secs[i].confidence for i in range(n)]
        prev = [-1] * n
        for i in range(n):
            for j in range(i):
                if not _order_compatible(
                    sorted_secs[j].section_type,
                    sorted_secs[i].section_type,
                    order_rank,
                ):
                    continue
                cand = dp[j] + sorted_secs[i].confidence
                if cand > dp[i]:
                    dp[i] = cand
                    prev[i] = j

        end_idx = max(range(n), key=lambda i: dp[i])
        picked: list[int] = []
        cur = end_idx
        while cur != -1:
            picked.append(cur)
            cur = prev[cur]
        picked.reverse()

        validated = [sorted_secs[i] for i in picked]
        picked_types = {s.section_type for s in validated}
        for sec in sorted_secs:
            if sec.section_type not in picked_types:
                self._detection_log.append(
                    f"顺序校验: 丢弃 {sec.section_type}"
                    f"(段落 #{sec.start_index}, conf={sec.confidence:.1f})"
                )
        return validated

    def _infer_body(self, total: int) -> None:
        """推断正文分区：所有前置分区之后，后置分区之前。

        前置分区: cover, toc, abstract_cn, abstract_en
        后置分区: references, errata, appendix, acknowledgment, resume
        """
        pre_secs = [s for s in self.sections
                    if s.section_type in _PRE_BODY_TYPES]
        post_secs = [s for s in self.sections
                     if s.section_type in _POST_BODY_TYPES]
        post_secs.sort(key=lambda s: s.start_index)
        last_pre = max(pre_secs, key=lambda s: s.start_index) if pre_secs else None
        if last_pre:
            pre_tail = (
                last_pre.end_index
                if last_pre.end_index >= last_pre.start_index
                else last_pre.start_index
            )
            scan_start = max(0, min(total, pre_tail + 1))
        else:
            scan_start = 0

        # 优先取 scan_start 之后的首个后置分区；若出现“后置分区早于正文起点”，先忽略
        first_post = None
        for sec in post_secs:
            if sec.start_index >= scan_start:
                first_post = sec
                break

        body_end = (first_post.start_index - 1) if first_post else total - 1
        if body_end < scan_start:
            # first_post aligns exactly with scan_start (common when TOC overcovers body
            # and the next detected anchor is references). Keep the empty range so that
            # recovery paths can pull body_start back from TOC/pre-body overreach.
            if first_post is not None and first_post.start_index == scan_start:
                pass
            elif first_post is None:
                body_end = total - 1
            # Keep empty [scan_start, body_end] when first_post starts exactly
            # at scan_start so recovery logic can still move body start earlier.

        # body 起点：样式/标题信号 -> 稳定正文簇回退
        body_start = self._scan_body_start_range(scan_start, body_end + 1)
        if body_start is None:
            body_start = self._scan_stable_body_cluster(scan_start, body_end + 1)
        if body_start is None:
            body_start = scan_start

        if body_start > body_end:
            toc_sec = self.get_section("toc")
            recovered_start = self._recover_body_start_from_toc_overreach(toc_sec, total)
            recover_reason = "TOC 过覆盖"
            if recovered_start is None:
                recovered_start = self._recover_body_start_from_pre_overreach(last_pre, total)
                recover_reason = "pre-body 过覆盖"
            if recovered_start is not None:
                body_start = recovered_start
                if first_post and first_post.start_index > body_start:
                    body_end = first_post.start_index - 1
                else:
                    body_end = total - 1
                self._detection_log.append(
                    f"body 纠偏: {recover_reason}，正文起点回退到段落 #{body_start}"
                )

        # Pre-body blocks should never start at/after body start.
        cleaned_sections: list[DocSection] = []
        dropped_pre: list[DocSection] = []
        for sec in self.sections:
            if sec.section_type in _PRE_BODY_TYPES and sec.start_index >= body_start:
                dropped_pre.append(sec)
                continue
            cleaned_sections.append(sec)
        if dropped_pre:
            self._detection_log.append(
                "body 纠偏: 清理过覆盖 pre-body -> "
                + ", ".join(f"{s.section_type}@#{s.start_index}" for s in dropped_pre)
            )
        self.sections = cleaned_sections

        valid_pre_secs = [
            s for s in self.sections
            if s.section_type in _PRE_BODY_TYPES and s.start_index < body_start
        ]
        if valid_pre_secs:
            valid_last_pre = max(valid_pre_secs, key=lambda s: s.start_index)
            valid_last_pre.end_index = max(valid_last_pre.start_index, body_start - 1)

        if body_start <= body_end:
            self.sections.append(DocSection("body", body_start, body_end, confidence=10.0))
        elif not self.sections:
            self.sections.append(DocSection("body", 0, total - 1, confidence=10.0))

    def _recover_body_start_from_toc_overreach(
            self, toc_sec: DocSection | None, total: int) -> int | None:
        """When TOC swallows tail content, try to split out body start from TOC range."""
        if not toc_sec:
            return None
        if total <= 0:
            return None

        start = max(0, toc_sec.start_index + 1)
        end = min(total, max(toc_sec.end_index + 1, start))
        if start >= end:
            return None

        recovered = self._scan_body_start_range(start, end)
        if recovered is None:
            recovered = self._scan_stable_body_cluster(start, end)
        if recovered is None:
            return None

        # Keep TOC-only paragraphs out of body even in recovery path.
        text = (self._doc_ref.paragraphs[recovered].text or "").strip() if self._doc_ref else ""
        if looks_like_toc_entry_line(text):
            return None
        return recovered

    def _recover_body_start_from_pre_overreach(
            self, last_pre: DocSection | None, total: int) -> int | None:
        """Recover body start when a pre-body section overcovers to document tail."""
        if not last_pre:
            return None
        if total <= 0:
            return None

        start = max(0, last_pre.start_index + 1)
        if last_pre.start_index >= int(total * 0.55):
            toc_sec = self.get_section("toc")
            if toc_sec and toc_sec.start_index < last_pre.start_index:
                start = max(0, toc_sec.start_index + 1)
            else:
                pre_secs = [s for s in self.sections if s.section_type in _PRE_BODY_TYPES]
                if pre_secs:
                    earliest = min(pre_secs, key=lambda s: s.start_index)
                    if earliest.start_index >= int(total * 0.30):
                        start = 0
                    else:
                        start = max(0, earliest.start_index + 1)
        end = min(total, max(last_pre.end_index + 1, start))
        if start >= end:
            return None

        recovered = self._scan_body_start_range(start, end)
        recovered_by_cluster = False
        if recovered is None:
            recovered = self._scan_stable_body_cluster(start, end)
            recovered_by_cluster = recovered is not None
            if recovered is not None and recovered <= last_pre.start_index:
                recovered = None
                recovered_by_cluster = False
        if recovered is None and end < total:
            # last_pre end may already be over-extended; allow full-tail scan once.
            recovered = self._scan_body_start_range(start, total)
            if recovered is None:
                recovered = self._scan_stable_body_cluster(start, total)
                recovered_by_cluster = recovered is not None
                if recovered is not None and recovered <= last_pre.start_index:
                    recovered = None
                    recovered_by_cluster = False
        if recovered is None:
            return None

        if recovered_by_cluster and last_pre.section_type in {"abstract_cn", "abstract_en"}:
            # Abstract bodies are narrative prose by design; without a stronger
            # heading/body-start signal, do not reclassify them as正文.
            return None

        text = (self._doc_ref.paragraphs[recovered].text or "").strip() if self._doc_ref else ""
        if (
            looks_like_toc_entry_line(text)
            or looks_like_numbered_toc_entry_with_page_suffix(text)
            or looks_like_date_placeholder_line(text)
            or _is_reference_entry_text(text)
        ):
            return None
        return recovered

    def _detect_references_by_content(
            self, doc: Document, total: int, *, scan_start: int = 0) -> DocSection | None:
        """通过连续参考文献条目识别 references 起点（用于无“参考文献”标题场景）。"""
        if total <= 0:
            return None

        start = max(scan_start, int(total * 0.35))
        best_idx = None
        best_score = 0.0

        for i in range(start, total):
            text = (doc.paragraphs[i].text or "").strip()
            if not _is_reference_entry_content_text(text):
                continue

            # 观察后续窗口内是否形成“连续参考条目簇”
            window_end = min(total, i + 12)
            ref_count = 0
            nonempty_count = 0
            for j in range(i, window_end):
                t = (doc.paragraphs[j].text or "").strip()
                if not t:
                    continue
                nonempty_count += 1
                if _is_reference_entry_content_text(t):
                    ref_count += 1

            if ref_count < 2:
                continue
            density = ref_count / max(nonempty_count, 1)
            if density < 0.5:
                continue

            score = 8.0 + min(4.0, ref_count * 0.8) + density
            if score > best_score:
                best_score = score
                best_idx = i

        if best_idx is None:
            return None

        self._detection_log.append(
            f"内容识别: references @#{best_idx}, score={best_score:.1f}"
        )
        return DocSection(
            "references",
            best_idx,
            confidence=best_score,
            title_confident=False,
        )

    def _scan_stable_body_cluster(self, start: int, end: int) -> int | None:
        """在 [start, end) 中寻找稳定正文簇（至少两段连续正文样式文本）。"""
        doc = self._doc_ref
        if not doc:
            return None

        cluster_start = None
        cluster_size = 0
        for idx in range(start, min(end, len(doc.paragraphs))):
            p = doc.paragraphs[idx]
            text = (p.text or "").strip()
            if not text:
                continue

            style = p.style.name if p.style else ""
            if re.search(r'^(toc|目录)\s*\d+', style, re.IGNORECASE):
                cluster_start = None
                cluster_size = 0
                continue
            if looks_like_toc_entry_line(text):
                cluster_start = None
                cluster_size = 0
                continue
            if looks_like_numbered_toc_entry_with_page_suffix(text):
                cluster_start = None
                cluster_size = 0
                continue
            if _is_front_matter_title_text(text):
                cluster_start = None
                cluster_size = 0
                continue
            if _is_reference_entry_text(text):
                cluster_start = None
                cluster_size = 0
                continue
            if looks_like_date_placeholder_line(text):
                cluster_start = None
                cluster_size = 0
                continue
            if len(text) < 8:
                cluster_start = None
                cluster_size = 0
                continue

            if cluster_start is None:
                cluster_start = idx
                cluster_size = 1
            else:
                # 允许中间夹一个空段，仍视为连续正文簇
                if idx - (cluster_start + cluster_size - 1) <= 2:
                    cluster_size += 1
                else:
                    cluster_start = idx
                    cluster_size = 1

            if cluster_size >= 2:
                self._detection_log.append(
                    f"body 回退: 稳定正文簇 @#{cluster_start}"
                )
                return cluster_start
        return None

    # 正文起点检测模式（按优先级排列）
    _RE_BODY_START = [
        re.compile(r'^#{1,6}\s+'),                    # Markdown 标题
        re.compile(r'^第[一二三四五六七八九十\d]+[章节篇]'),  # 第X章/节/篇
        re.compile(r'^\d+[\.、]\S'),                   # 1.绪论 / 1、绪论
        re.compile(r'^\d+\s+\S'),                      # 1 绪论
    ]

    def _scan_body_start_range(self, start: int, end: int) -> int | None:
        """在 [start, end) 范围内扫描，寻找正文起点段落索引。

        检测信号（按优先级）：
          1. Heading / 一级标题 等样式
          2. Markdown # 标题
          3. 中文章节标题（第X章/节）
          4. 数字编号标题（1.绪论、1 绪论）
        """
        doc = self._doc_ref
        if not doc:
            return None

        for idx in range(start, end):
            if idx >= len(doc.paragraphs):
                break
            p = doc.paragraphs[idx]
            style = p.style.name if p.style else ""
            text = p.text.strip()
            if not text:
                continue

            # 跳过 toc 样式段落（避免将目录条目误判为正文起点）
            if re.search(r'^(toc|目录)\s*\d+', style, re.IGNORECASE):
                continue
            if looks_like_toc_entry_line(text):
                continue
            if looks_like_numbered_toc_entry_with_page_suffix(text):
                continue
            if _is_front_matter_title_text(text):
                continue
            if looks_like_date_placeholder_line(text):
                continue
            if _is_reference_entry_text(text):
                continue

            # Heading 样式或自定义一级标题样式
            if (style.startswith("Heading")
                    or "标题" in style):
                self._detection_log.append(
                    f"body 扫描: 样式 '{style}'，"
                    f"段落 #{idx} ({text[:30]})")
                return idx

            # 正则模式匹配
            for pat in self._RE_BODY_START:
                if pat.match(text):
                    self._detection_log.append(
                        f"body 扫描: 模式匹配，"
                        f"段落 #{idx} ({text[:30]})")
                    return idx

        return None

    def get_section(self, section_type: str) -> DocSection | None:
        for s in self.sections:
            if s.section_type == section_type:
                return s
        return None

    def get_section_for_paragraph(self, para_index: int) -> str:
        """返回段落所属的分区类型"""
        for s in self.sections:
            if s.start_index <= para_index <= s.end_index:
                return s.section_type
        return "unknown"



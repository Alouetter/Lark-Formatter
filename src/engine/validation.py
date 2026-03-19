"""Validation rule: check whether formatted result matches scene constraints."""

from docx import Document

from src.engine.rules.base import BaseRule, ValidationIssue
from src.engine.change_tracker import ChangeTracker
from src.scene.heading_model import get_level_to_word_style
from src.scene.schema import SceneConfig
from src.utils.toc_entry import (
    looks_like_numbered_toc_entry_with_page_suffix,
    looks_like_toc_entry_line,
)


_BROKEN_TOC_BOOKMARK_HINTS = (
    "error! bookmark not defined",
    "bookmark not defined",
    "\u672a\u5b9a\u4e49\u4e66\u7b7e",
)
_EXPLICIT_FRONT_MATTER_TYPES = {"cover", "toc", "abstract_cn", "abstract_en"}


def _preserved_body_start_index(old_tree, total_paragraphs: int) -> int | None:
    """Keep manual pre_body/body scopes stable without collapsing explicit front matter."""
    if old_tree is None or total_paragraphs <= 0:
        return None

    try:
        if any(
            old_tree.get_section(sec_type) is not None
            for sec_type in _EXPLICIT_FRONT_MATTER_TYPES
        ):
            return None
        pre_body = old_tree.get_section("pre_body")
        body_sec = old_tree.get_section("body")
    except Exception:
        return None

    if pre_body is None or body_sec is None or not isinstance(body_sec.start_index, int):
        return None
    return max(0, min(body_sec.start_index, total_paragraphs - 1))


class ValidationRule(BaseRule):
    name = "validation"
    description = "校验排版结果"

    def apply(
        self,
        doc: Document,
        config: SceneConfig,
        tracker: ChangeTracker,
        context: dict,
    ) -> None:
        issues = self.validate(doc, config, context)
        context["validation_issues"] = issues
        for issue in issues:
            tracker.record(
                rule_name=self.name,
                target=issue.location,
                section="global",
                change_type="format",
                before=issue.message,
                after=f"[{issue.level}]",
                paragraph_index=-1,
                success=(issue.level != "error"),
                failure_reason=issue.message if issue.level == "error" else None,
            )

    def validate(
        self,
        doc: Document,
        config: SceneConfig,
        context: dict | None = None,
    ) -> list[ValidationIssue]:
        issues: list[ValidationIssue] = []
        ctx = self._prepare_validation_context(doc, config, context)
        target_indices = self._target_paragraph_indices(ctx)
        if not target_indices:
            issues.extend(self._check_page_setup(doc, config))
        issues.extend(self._check_heading_styles(doc, config, ctx))
        issues.extend(self._check_separator_violations(doc, config, ctx))
        return issues

    @staticmethod
    def _target_paragraph_indices(context: dict | None) -> set[int]:
        if not isinstance(context, dict):
            return set()
        raw = context.get("target_paragraph_indices")
        if raw is None:
            return set()
        return {idx for idx in raw if isinstance(idx, int) and idx >= 0}

    def _prepare_validation_context(
        self,
        doc: Document,
        config: SceneConfig,
        context: dict | None,
    ) -> dict:
        """Strict mode: refresh heading context on the final document snapshot."""
        ctx = context if isinstance(context, dict) else {}
        refreshed = dict(ctx)

        try:
            from src.engine.doc_tree import DocTree
            from src.engine.rules.heading_detect import HeadingDetectRule
        except Exception:
            return refreshed

        total_paragraphs = len(doc.paragraphs)
        old_tree = ctx.get("doc_tree")
        body_start_index = _preserved_body_start_index(old_tree, total_paragraphs)

        doc_tree = DocTree()
        if total_paragraphs > 0:
            if body_start_index is None:
                doc_tree.build(doc)
            else:
                doc_tree.build(doc, body_start_index=body_start_index)
        refreshed["doc_tree"] = doc_tree

        heading_context = {"doc_tree": doc_tree}
        target_indices = self._target_paragraph_indices(ctx)
        if target_indices:
            heading_context["target_paragraph_indices"] = set(target_indices)
        HeadingDetectRule().apply(doc, config, ChangeTracker(), heading_context)
        refreshed["headings"] = heading_context.get("headings", [])
        if target_indices:
            refreshed["target_paragraph_indices"] = set(target_indices)

        if isinstance(context, dict):
            context["doc_tree"] = doc_tree
            context["headings"] = refreshed["headings"]
            if target_indices:
                context["target_paragraph_indices"] = set(target_indices)
        return refreshed

    def _check_page_setup(self, doc, config):
        issues: list[ValidationIssue] = []
        ps = config.page_setup
        for i, section in enumerate(doc.sections):
            expected_top = round(ps.margin.top_cm * 360000)
            actual_top = section.top_margin
            if actual_top and abs(actual_top - expected_top) > 5000:
                issues.append(
                    ValidationIssue(
                        level="warning",
                        rule_name=self.name,
                        message=f"页边距(上)不符: 期望 {ps.margin.top_cm}cm",
                        location=f"Section {i + 1}",
                    )
                )
        return issues

    def _check_heading_styles(self, doc, config, context):
        issues: list[ValidationIssue] = []
        headings = context.get("headings", [])
        level_style = get_level_to_word_style(config)
        skipped_invalid_index = 0
        for h in headings:
            para_idx = getattr(h, "para_index", None)
            if not isinstance(para_idx, int) or para_idx < 0 or para_idx >= len(doc.paragraphs):
                skipped_invalid_index += 1
                continue
            para = doc.paragraphs[para_idx]
            expected = level_style.get(h.level, "")
            actual = para.style.name if para.style else ""
            if actual != expected:
                issues.append(
                    ValidationIssue(
                        level="warning",
                        rule_name=self.name,
                        message=f"标题样式不符: 期望 '{expected}', 实际 '{actual}'",
                        location=f"段落 #{para_idx}",
                    )
                )
        if skipped_invalid_index:
            issues.append(
                ValidationIssue(
                    level="error",
                    rule_name=self.name,
                    message=f"标题样式校验命中 {skipped_invalid_index} 个失效标题索引",
                    location="validation:heading_styles",
                )
            )
        return issues

    def _is_toc_placeholder_heading(self, text: str) -> bool:
        raw = (text or "").strip()
        if "\t" not in raw:
            return False
        low = raw.lower()
        if looks_like_toc_entry_line(raw) or looks_like_numbered_toc_entry_with_page_suffix(raw):
            return True
        if any(hint in low for hint in _BROKEN_TOC_BOOKMARK_HINTS):
            return True
        return ("\u4e66\u7b7e" in low and "\u672a\u5b9a\u4e49" in low)

    def _check_separator_violations(self, doc, config, context):
        issues: list[ValidationIssue] = []
        headings = context.get("headings", [])
        enforcement = config.heading_numbering.enforcement
        skipped_invalid_index = 0
        for h in headings:
            para_idx = getattr(h, "para_index", None)
            if not isinstance(para_idx, int) or para_idx < 0 or para_idx >= len(doc.paragraphs):
                skipped_invalid_index += 1
                continue
            text = doc.paragraphs[para_idx].text or ""
            if self._is_toc_placeholder_heading(text):
                continue
            if enforcement.ban_tab and "\t" in text:
                issues.append(
                    ValidationIssue(
                        level="error",
                        rule_name=self.name,
                        message="标题中仍存在 Tab 字符",
                        location=f"段落 #{para_idx}",
                    )
                )
            if enforcement.ban_double_halfwidth_space and "  " in text:
                issues.append(
                    ValidationIssue(
                        level="error",
                        rule_name=self.name,
                        message="标题中仍存在连续两个半角空格",
                        location=f"段落 #{para_idx}",
                    )
                )
        if skipped_invalid_index:
            issues.append(
                ValidationIssue(
                    level="error",
                    rule_name=self.name,
                    message=f"标题分隔符校验命中 {skipped_invalid_index} 个失效标题索引",
                    location="validation:separator",
                )
            )
        return issues

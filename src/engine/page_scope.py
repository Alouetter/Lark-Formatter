"""Helpers for mapping user page ranges to paragraph ranges."""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree
from src.utils.runtime_features import word_page_scope_forced_disabled

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_PAGE_BREAK = f"{{{_W_NS}}}br"
_PPR = f"{{{_W_NS}}}pPr"
_SECTPR = f"{{{_W_NS}}}sectPr"
_SECT_TYPE = f"{{{_W_NS}}}type"
_SECT_VAL = f"{{{_W_NS}}}val"
_BR_TYPE = f"{{{_W_NS}}}type"
_BOOKMARK_START = f"{{{_W_NS}}}bookmarkStart"
_BOOKMARK_END = f"{{{_W_NS}}}bookmarkEnd"
_BOOKMARK_ID = f"{{{_W_NS}}}id"
_BOOKMARK_NAME = f"{{{_W_NS}}}name"

_RANGE_TOKEN_RE = re.compile(r"^(\d+)\s*[-~\u2013\u2014\uFF0D\uFF5E]\s*(\d+)$")
_SINGLE_TOKEN_RE = re.compile(r"^\d+$")
_SPLIT_RE = re.compile(r"[,\uFF0C;\uFF1B\u3001]+")

_PAGE_SCOPE_BOOKMARK_PREFIX = "_LF_PAGE_SCOPE_"
_DEFAULT_TIMEOUT_SEC = 30
_WORD_PAGE_NUMBER_INFO = 3

_PYWIN32_PAGE_SPAN_CHILD = r"""
import json
import os
from pathlib import Path

doc_path = (os.environ.get("DOCX_PAGE_SCOPE_PATH") or "").strip()
prefix = os.environ.get("DOCX_PAGE_SCOPE_PREFIX") or ""
if not doc_path:
    raise RuntimeError("DOCX_PAGE_SCOPE_PATH is empty.")

import pythoncom
import win32com.client

WD_ACTIVE_END_PAGE_NUMBER = 3

word = None
doc = None
pythoncom.CoInitialize()
try:
    p = str(Path(doc_path).resolve())
    word = win32com.client.DispatchEx("Word.Application")
    try:
        word.Visible = False
    except Exception:
        pass
    try:
        word.DisplayAlerts = 0
    except Exception:
        pass

    doc = word.Documents.Open(
        p,
        False, True, False, "", "", False, "", "", 0, 0, False, True
    )
    doc.Repaginate()

    spans = {}
    for bookmark in list(doc.Bookmarks):
        try:
            name = str(bookmark.Name)
        except Exception:
            continue
        if not name.startswith(prefix):
            continue
        try:
            para_range = bookmark.Range.Paragraphs(1).Range
            start_pos = max(0, int(para_range.Start))
            end_pos = max(start_pos, int(para_range.End) - 1)
            start_page = int(doc.Range(start_pos, start_pos).Information(WD_ACTIVE_END_PAGE_NUMBER))
            end_page = int(doc.Range(end_pos, end_pos).Information(WD_ACTIVE_END_PAGE_NUMBER))
            spans[name] = [start_page, end_page]
        except Exception:
            continue

    print(json.dumps(spans, ensure_ascii=False))
finally:
    if doc is not None:
        try:
            doc.Close(False)
        except Exception:
            pass
    if word is not None:
        try:
            word.Quit()
        except Exception:
            pass
    pythoncom.CoUninitialize()
"""

_POWERSHELL_PAGE_SPAN_SCRIPT = r"""
$ErrorActionPreference = 'Stop'

$docPath = $env:DOCX_PAGE_SCOPE_PATH
$prefix = $env:DOCX_PAGE_SCOPE_PREFIX
if (-not $docPath) {
    throw "DOCX_PAGE_SCOPE_PATH is empty."
}
if (-not (Test-Path -LiteralPath $docPath)) {
    throw "File not found: $docPath"
}

$word = $null
$doc = $null

try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open($docPath, $false, $true)
    $doc.Repaginate() | Out-Null

    $spans = @{}
    foreach ($bookmark in @($doc.Bookmarks)) {
        try {
            $name = [string]$bookmark.Name
        } catch {
            continue
        }
        if (-not $name.StartsWith($prefix)) {
            continue
        }
        try {
            $paraRange = $bookmark.Range.Paragraphs.Item(1).Range
            $startPos = [int]$paraRange.Start
            $endPos = [int]$paraRange.End - 1
            if ($endPos -lt $startPos) {
                $endPos = $startPos
            }
            $startPage = [int]$doc.Range($startPos, $startPos).Information(3)
            $endPage = [int]$doc.Range($endPos, $endPos).Information(3)
            $spans[$name] = @($startPage, $endPage)
        } catch {
        }
    }

    $spans | ConvertTo-Json -Compress
}
finally {
    if ($doc -ne $null) {
        try { $doc.Close($false) } catch {}
        try { [void][Runtime.InteropServices.Marshal]::FinalReleaseComObject($doc) } catch {}
    }
    if ($word -ne $null) {
        try { $word.Quit() } catch {}
        try { [void][Runtime.InteropServices.Marshal]::FinalReleaseComObject($word) } catch {}
    }
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
}
"""


def merge_ranges(ranges: list[tuple[int, int]]) -> list[tuple[int, int]]:
    if not ranges:
        return []
    ordered = sorted((int(s), int(e)) for s, e in ranges)
    merged: list[list[int]] = []
    for start, end in ordered:
        if not merged:
            merged.append([start, end])
            continue
        prev = merged[-1]
        if start <= prev[1] + 1:
            prev[1] = max(prev[1], end)
        else:
            merged.append([start, end])
    return [(s, e) for s, e in merged]


def parse_page_ranges_text(text: str) -> list[tuple[int, int]]:
    """Parse page-range text like ``27-40,44-56`` into merged ranges."""
    raw = (text or "").strip()
    if not raw:
        return []

    tokens = [token.strip() for token in _SPLIT_RE.split(raw) if token.strip()]
    if not tokens:
        return []

    parsed: list[tuple[int, int]] = []
    for token in tokens:
        m = _RANGE_TOKEN_RE.match(token)
        if m:
            start = int(m.group(1))
            end = int(m.group(2))
        elif _SINGLE_TOKEN_RE.match(token):
            start = end = int(token)
        else:
            raise ValueError(f"invalid page token: {token}")
        if start < 1 or end < 1:
            raise ValueError("page number must be >= 1")
        if start > end:
            raise ValueError(f"page range start/end order is invalid: {token}")
        parsed.append((start, end))
    return merge_ranges(parsed)


def format_page_ranges_text(ranges: list[tuple[int, int]]) -> str:
    chunks: list[str] = []
    for start, end in merge_ranges(ranges):
        if start == end:
            chunks.append(str(start))
        else:
            chunks.append(f"{start}-{end}")
    return ", ".join(chunks)


def paragraph_ranges_to_index_set(ranges: list[tuple[int, int]]) -> set[int]:
    indices: set[int] = set()
    for start, end in ranges:
        if end < start:
            continue
        indices.update(range(start, end + 1))
    return indices


def _probe_timeout_sec() -> int:
    raw = os.environ.get("DOCX_PAGE_SCOPE_TIMEOUT_SEC", str(_DEFAULT_TIMEOUT_SEC))
    try:
        timeout = int(raw)
    except ValueError:
        timeout = _DEFAULT_TIMEOUT_SEC
    return max(10, min(timeout, 600))


def _append_page_start(page_starts: list[int], para_index: int, total_paragraphs: int) -> None:
    if total_paragraphs <= 0:
        return
    if para_index <= 0 or para_index >= total_paragraphs:
        return
    if para_index > page_starts[-1]:
        page_starts.append(para_index)


def _collect_explicit_page_start_indices(doc: Document) -> list[int]:
    total = len(doc.paragraphs)
    if total <= 0:
        return []

    page_starts = [0]
    para_idx = -1
    body = doc.element.body

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "p":
            continue

        para_idx += 1
        for br in child.iter(_PAGE_BREAK):
            if br.get(_BR_TYPE, "") == "page":
                _append_page_start(page_starts, para_idx + 1, total)

        ppr = child.find(_PPR)
        if ppr is None:
            continue
        sect = ppr.find(_SECTPR)
        if sect is None:
            continue
        sect_type = ""
        sect_type_el = sect.find(_SECT_TYPE)
        if sect_type_el is not None:
            sect_type = sect_type_el.get(_SECT_VAL, "")
        if sect_type in ("nextPage", "oddPage", "evenPage", ""):
            _append_page_start(page_starts, para_idx + 1, total)

    return page_starts


def _build_explicit_paragraph_page_spans(doc: Document) -> list[tuple[int, int]]:
    total = len(doc.paragraphs)
    if total <= 0:
        return []

    page_starts = _collect_explicit_page_start_indices(doc)
    if not page_starts:
        return [(1, 1)] * total

    spans = [(1, 1)] * total
    page_count = len(page_starts)
    for page_idx, start_para in enumerate(page_starts, start=1):
        if page_idx < page_count:
            end_para = page_starts[page_idx] - 1
        else:
            end_para = total - 1
        for para_idx in range(start_para, end_para + 1):
            spans[para_idx] = (page_idx, page_idx)
    return spans


def _page_scope_bookmark_name(para_index: int) -> str:
    return f"{_PAGE_SCOPE_BOOKMARK_PREFIX}{para_index}"


def _strip_generated_probe_bookmarks(body_el) -> None:
    stale_ids: set[str] = set()
    for b_start in list(body_el.iter(_BOOKMARK_START)):
        name = (b_start.get(_BOOKMARK_NAME) or "").strip()
        if not name.startswith(_PAGE_SCOPE_BOOKMARK_PREFIX):
            continue
        bid = (b_start.get(_BOOKMARK_ID) or "").strip()
        if bid:
            stale_ids.add(bid)
        parent = b_start.getparent()
        if parent is not None:
            parent.remove(b_start)

    if not stale_ids:
        return

    for b_end in list(body_el.iter(_BOOKMARK_END)):
        bid = (b_end.get(_BOOKMARK_ID) or "").strip()
        if bid not in stale_ids:
            continue
        parent = b_end.getparent()
        if parent is not None:
            parent.remove(b_end)


def _inject_page_scope_bookmarks(docx_path: Path) -> int:
    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    with zipfile.ZipFile(docx_path, "r") as zin:
        try:
            xml_bytes = zin.read("word/document.xml")
        except KeyError as exc:
            raise RuntimeError("DOCX is missing word/document.xml") from exc

    root = etree.fromstring(xml_bytes)
    body = root.find(f"{{{_W_NS}}}body")
    if body is None:
        raise RuntimeError("DOCX document.xml is missing w:body")

    _strip_generated_probe_bookmarks(body)

    max_id = 0
    for b in body.iter(_BOOKMARK_START):
        try:
            max_id = max(max_id, int(b.get(_BOOKMARK_ID, "0")))
        except Exception:
            continue

    para_idx = 0
    next_id = max_id + 1
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "p":
            continue
        ppr = child.find(_PPR)
        insert_idx = list(child).index(ppr) + 1 if ppr is not None else 0

        b_start = etree.Element(_BOOKMARK_START)
        b_start.set(_BOOKMARK_ID, str(next_id))
        b_start.set(_BOOKMARK_NAME, _page_scope_bookmark_name(para_idx))

        b_end = etree.Element(_BOOKMARK_END)
        b_end.set(_BOOKMARK_ID, str(next_id))

        child.insert(insert_idx, b_start)
        child.insert(insert_idx + 1, b_end)

        next_id += 1
        para_idx += 1

    updated_xml = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    tmp_zip = docx_path.with_name(f".{docx_path.name}.{uuid.uuid4().hex}.tmp")
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            with zipfile.ZipFile(tmp_zip, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = updated_xml if item.filename == "word/document.xml" else zin.read(item.filename)
                    zout.writestr(item, data)
        os.replace(tmp_zip, docx_path)
    finally:
        try:
            if tmp_zip.exists():
                tmp_zip.unlink()
        except OSError:
            pass

    return para_idx


def _make_page_scope_probe_doc(doc: Document, source_doc_path: str | Path | None) -> tuple[Path, int]:
    tmp_dir = Path(tempfile.gettempdir()) / "lark_formatter_page_scope"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    probe_path = tmp_dir / f"{uuid.uuid4().hex}.docx"

    source_path = Path(source_doc_path) if source_doc_path else None
    if source_path is not None and source_path.exists():
        shutil.copy2(source_path, probe_path)
    else:
        doc.save(str(probe_path))

    probe_paragraphs = _inject_page_scope_bookmarks(probe_path)
    return probe_path, probe_paragraphs


def _parse_word_span_payload(payload: str, paragraph_count: int) -> list[tuple[int, int]]:
    text = (payload or "").strip()
    if not text:
        raise RuntimeError("Word page probe returned empty output.")
    try:
        raw = json.loads(text)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Word page probe returned invalid JSON: {text}") from exc

    if not isinstance(raw, dict):
        raise RuntimeError("Word page probe returned invalid payload type.")

    spans: list[tuple[int, int]] = []
    missing: list[int] = []
    for para_idx in range(paragraph_count):
        entry = raw.get(_page_scope_bookmark_name(para_idx))
        if (
            not isinstance(entry, list)
            or len(entry) != 2
            or not isinstance(entry[0], int)
            or not isinstance(entry[1], int)
        ):
            missing.append(para_idx)
            continue
        start_page = max(1, int(entry[0]))
        end_page = max(start_page, int(entry[1]))
        spans.append((start_page, end_page))

    if missing:
        raise RuntimeError(
            f"Word page probe missing {len(missing)} paragraph spans."
        )
    return spans


def _run_pywin32_page_probe(probe_path: Path, timeout_sec: int) -> tuple[bool, str]:
    if getattr(sys, "frozen", False):
        return False, "pywin32 page probe skipped in frozen executable."
    exe_name = Path(sys.executable).name.lower()
    if "python" not in exe_name:
        return False, "pywin32 page probe skipped: current runtime is not a Python interpreter."

    env = os.environ.copy()
    env["DOCX_PAGE_SCOPE_PATH"] = str(probe_path.resolve())
    env["DOCX_PAGE_SCOPE_PREFIX"] = _PAGE_SCOPE_BOOKMARK_PREFIX
    cmd = [sys.executable, "-c", _PYWIN32_PAGE_SPAN_CHILD]
    try:
        proc = subprocess.run(
            cmd,
            env=env,
            capture_output=True,
            text=True,
            timeout=timeout_sec,
            check=False,
        )
    except subprocess.TimeoutExpired:
        return False, f"pywin32 page probe timed out after {timeout_sec}s."
    except Exception as exc:
        return False, f"pywin32 page probe failed to start: {exc}"

    if proc.returncode == 0:
        return True, (proc.stdout or "").strip()

    stderr = (proc.stderr or "").strip()
    stdout = (proc.stdout or "").strip()
    detail = stderr if stderr else stdout
    if not detail:
        detail = f"pywin32 page probe exited with code {proc.returncode}"
    return False, detail


def _run_powershell_page_probe(probe_path: Path, timeout_sec: int) -> tuple[bool, str]:
    env = os.environ.copy()
    env["DOCX_PAGE_SCOPE_PATH"] = str(probe_path.resolve())
    env["DOCX_PAGE_SCOPE_PREFIX"] = _PAGE_SCOPE_BOOKMARK_PREFIX
    cmd = [
        "powershell",
        "-NoProfile",
        "-NonInteractive",
        "-Command",
        _POWERSHELL_PAGE_SPAN_SCRIPT,
    ]
    try:
        proc = subprocess.run(
            cmd,
            env=env,
            capture_output=True,
            text=True,
            timeout=timeout_sec,
            check=False,
        )
    except subprocess.TimeoutExpired:
        return False, f"Word page probe timed out after {timeout_sec}s."
    except Exception as exc:
        return False, f"Word page probe failed to start: {exc}"

    if proc.returncode == 0:
        return True, (proc.stdout or "").strip()

    stderr = (proc.stderr or "").strip()
    stdout = (proc.stdout or "").strip()
    detail = stderr if stderr else stdout
    if not detail:
        detail = f"powershell page probe exited with code {proc.returncode}"
    return False, detail


def _resolve_word_paragraph_page_spans(
    doc: Document,
    *,
    source_doc_path: str | Path | None = None,
    timeout_sec: int | None = None,
) -> list[tuple[int, int]]:
    if word_page_scope_forced_disabled():
        raise RuntimeError("当前运行环境不支持“指定修正范围”；请改用自动识别分区。")
    if os.name != "nt":
        raise RuntimeError("真实页码解析仅支持 Windows + Microsoft Word。")

    paragraph_count = len(doc.paragraphs)
    if paragraph_count <= 0:
        return []

    probe_path, probe_paragraphs = _make_page_scope_probe_doc(doc, source_doc_path)
    try:
        if probe_paragraphs != paragraph_count:
            raise RuntimeError(
                f"Word page probe paragraph count mismatch: expected {paragraph_count}, got {probe_paragraphs}."
            )

        timeout = timeout_sec if timeout_sec is not None else _probe_timeout_sec()
        ok, payload = _run_pywin32_page_probe(probe_path, timeout)
        if ok:
            return _parse_word_span_payload(payload, paragraph_count)

        ok_ps, payload_ps = _run_powershell_page_probe(probe_path, timeout)
        if ok_ps:
            return _parse_word_span_payload(payload_ps, paragraph_count)

        raise RuntimeError(f"{payload}; fallback failed: {payload_ps}")
    finally:
        try:
            probe_path.unlink(missing_ok=True)
        except Exception:
            pass


def resolve_paragraph_page_spans(
    doc: Document,
    *,
    source_doc_path: str | Path | None = None,
    require_word: bool = False,
) -> list[tuple[int, int]]:
    paragraph_count = len(doc.paragraphs)
    if paragraph_count <= 0:
        return []

    try:
        return _resolve_word_paragraph_page_spans(doc, source_doc_path=source_doc_path)
    except Exception as exc:
        if require_word:
            raise RuntimeError(str(exc)) from exc
        return _build_explicit_paragraph_page_spans(doc)


def page_ranges_to_paragraph_ranges(
    doc: Document,
    page_ranges: list[tuple[int, int]],
    *,
    source_doc_path: str | Path | None = None,
    paragraph_page_spans: list[tuple[int, int]] | None = None,
    require_word: bool = False,
) -> list[tuple[int, int]]:
    """Map physical page ranges to paragraph index ranges."""
    total = len(doc.paragraphs)
    if total <= 0:
        return []

    normalized = merge_ranges(page_ranges)
    if not normalized:
        return []

    spans = (
        paragraph_page_spans
        if paragraph_page_spans is not None
        else resolve_paragraph_page_spans(
            doc,
            source_doc_path=source_doc_path,
            require_word=require_word,
        )
    )
    if not spans:
        return []

    para_ranges: list[tuple[int, int]] = []
    current_start: int | None = None
    current_end: int | None = None
    for para_idx, (start_page, end_page) in enumerate(spans):
        in_scope = any(
            not (end_page < range_start or start_page > range_end)
            for range_start, range_end in normalized
        )
        if in_scope:
            if current_start is None:
                current_start = para_idx
            current_end = para_idx
            continue
        if current_start is not None and current_end is not None:
            para_ranges.append((current_start, current_end))
            current_start = None
            current_end = None

    if current_start is not None and current_end is not None:
        para_ranges.append((current_start, current_end))
    return merge_ranges(para_ranges)


def page_number_to_start_paragraph_index(
    doc: Document,
    target_page: int,
    *,
    source_doc_path: str | Path | None = None,
    paragraph_page_spans: list[tuple[int, int]] | None = None,
    require_word: bool = False,
) -> int | None:
    if target_page is None or target_page < 1:
        return None

    spans = (
        paragraph_page_spans
        if paragraph_page_spans is not None
        else resolve_paragraph_page_spans(
            doc,
            source_doc_path=source_doc_path,
            require_word=require_word,
        )
    )
    for para_idx, (start_page, end_page) in enumerate(spans):
        if start_page <= target_page <= end_page:
            return para_idx
    return None

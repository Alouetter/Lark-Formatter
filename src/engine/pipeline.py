"""Pipeline 调度器：按顺序执行排版规则"""

import copy
import os
import shutil
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from src.engine.change_tracker import ChangeTracker, ChangeRecord
from src.engine.doc_tree import DocTree
from src.engine.rules.page_setup import PageSetupRule
from src.engine.rules.style_manager import StyleManagerRule
from src.engine.rules.heading_detect import HeadingDetectRule
from src.engine.rules.heading_numbering import HeadingNumberingRule
from src.engine.rules.caption_format import CaptionFormatRule
from src.engine.rules.section_format import SectionFormatRule
from src.engine.rules.table_format import TableFormatRule
from src.engine.rules.formula_convert import FormulaConvertRule
from src.engine.rules.formula_to_table import FormulaToTableRule
from src.engine.rules.formula_style import FormulaStyleRule
from src.engine.rules.equation_table_format import EquationTableFormatRule
from src.engine.rules.whitespace_normalize import WhitespaceNormalizeRule
from src.engine.rules.citation_link import CitationLinkRule
from src.engine.rules.md_cleanup import MdCleanupRule
from src.engine.rules.toc_format import TocFormatRule, _normalize_toc_styles_in_doc
from src.engine.rules.header_footer import HeaderFooterRule
from src.engine.validation import ValidationRule
from src.scene.manager import (
    _sync_citation_link_pipeline_switch,
    _sync_formula_pipeline_switch,
    _sync_pipeline_critical_rules,
    _sync_whitespace_pipeline_switch,
)
from src.scene.schema import SceneConfig, DEFAULT_PIPELINE_STEPS
from src.engine.page_scope import (
    parse_page_ranges_text,
    page_ranges_to_paragraph_ranges,
    page_number_to_start_paragraph_index,
    paragraph_ranges_to_index_set,
)
from src.docx_io.sanitize import sanitize_docx, docx_needs_sanitization
from src.docx_io.compare import generate_compare_doc
from src.docx_io.field_refresh import refresh_doc_fields_with_word
from src.docx_io.mathtype_office_fallback import apply_mathtype_office_fallback


@dataclass
class PipelineResult:
    """Pipeline 执行结果"""
    success: bool
    status: str = "success"  # success | partial_success | failed | cancelled
    doc: object | None = None
    original_doc: object | None = None
    tracker: ChangeTracker = field(default_factory=ChangeTracker)
    output_paths: dict = field(default_factory=dict)
    failed_items: list[dict] = field(default_factory=list)
    error: str | None = None
    cancelled: bool = False

    def __post_init__(self) -> None:
        if self.cancelled:
            self.status = "cancelled"
            return
        if not self.success and self.status == "success":
            self.status = "failed"


# V0 可用规则注册表
RULE_REGISTRY = {
    "page_setup": PageSetupRule,
    "md_cleanup": MdCleanupRule,
    "style_manager": StyleManagerRule,
    "heading_detect": HeadingDetectRule,
    "heading_numbering": HeadingNumberingRule,
    "toc_format": TocFormatRule,
    "caption_format": CaptionFormatRule,
    "table_format": TableFormatRule,
    "formula_convert": FormulaConvertRule,
    "formula_to_table": FormulaToTableRule,
    "formula_style": FormulaStyleRule,
    "equation_table_format": EquationTableFormatRule,
    "whitespace_normalize": WhitespaceNormalizeRule,
    "citation_link": CitationLinkRule,
    "section_format": SectionFormatRule,
    "header_footer": HeaderFooterRule,
    "validation": ValidationRule,
}

# V0 默认 pipeline 顺序（引用 schema.py 中的单一来源定义）
DEFAULT_PIPELINE = list(DEFAULT_PIPELINE_STEPS)


class Pipeline:
    """排版 Pipeline：加载文档 → 构建结构树 → 依次执行规则 → 输出结果"""

    def __init__(self, config: SceneConfig,
                 progress_callback=None,
                 cancel_requested=None):
        self.config = config
        self.tracker = ChangeTracker()
        self.progress_callback = progress_callback
        self.cancel_requested = cancel_requested
        self._last_input_file = ""

        # Keep config flags and pipeline steps synchronized even for callers
        # that construct SceneConfig directly instead of loading via scene.manager.
        _sync_whitespace_pipeline_switch(config)
        _sync_formula_pipeline_switch(config)
        _sync_citation_link_pipeline_switch(config)
        _sync_pipeline_critical_rules(config)

        # 按场景配置构建规则列表；仅在 None 时回退默认顺序。
        # 空列表表示显式禁用规则步骤，应被保留。
        step_names = DEFAULT_PIPELINE if config.pipeline is None else config.pipeline
        self.steps = []
        for name in step_names:
            rule_cls = RULE_REGISTRY.get(name)
            if rule_cls:
                self.steps.append(rule_cls())

    def _is_strict_mode_enabled(self) -> bool:
        """Strict mode can be configured by env var or scene config."""
        env_val = os.environ.get("DOCX_PIPELINE_STRICT_MODE")
        if env_val is not None and env_val.strip() != "":
            return env_val.strip().lower() in {"1", "true", "yes", "on"}
        return bool(getattr(self.config, "pipeline_strict_mode", True))

    def _critical_rules(self) -> set[str]:
        rules = getattr(self.config, "pipeline_critical_rules", None)
        if not isinstance(rules, list) or not rules:
            return {rule.name for rule in self.steps if getattr(rule, "name", "")}
        return {str(x).strip() for x in rules if str(x).strip()}

    def _collect_failed_items(self) -> list[dict]:
        items: list[dict] = []
        for rec in self.tracker.get_failures():
            items.append({
                "rule_name": rec.rule_name,
                "target": rec.target,
                "section": rec.section,
                "change_type": rec.change_type,
                "paragraph_index": rec.paragraph_index,
                "reason": rec.failure_reason or "",
            })
        return items

    def _finalize_result(self, *, doc, original_doc, output_paths: dict) -> PipelineResult:
        """Build structured status for success/partial/failure states."""
        failed_items = self._collect_failed_items()
        if not failed_items:
            return PipelineResult(
                success=True,
                status="success",
                doc=doc,
                original_doc=original_doc,
                tracker=self.tracker,
                output_paths=output_paths,
                failed_items=[],
            )

        strict_mode = self._is_strict_mode_enabled()
        critical_rules = self._critical_rules()
        has_critical_failure = any(
            item.get("rule_name") in critical_rules for item in failed_items
        )

        if strict_mode and has_critical_failure:
            return PipelineResult(
                success=False,
                status="failed",
                doc=doc,
                original_doc=original_doc,
                tracker=self.tracker,
                output_paths=output_paths,
                failed_items=failed_items,
                error=f"关键步骤失败 {len(failed_items)} 项（严格模式）",
            )

        return PipelineResult(
            success=True,
            status="partial_success",
            doc=doc,
            original_doc=original_doc,
            tracker=self.tracker,
            output_paths=output_paths,
            failed_items=failed_items,
            error=f"存在未成功项 {len(failed_items)} 项",
        )

    def _emit_progress(self, current: int, total: int, message: str):
        if self.progress_callback:
            self.progress_callback(current, total, message)

    def _is_cancel_requested(self) -> bool:
        if self.cancel_requested is None:
            return False
        try:
            return bool(self.cancel_requested())
        except Exception:
            return False

    def _cancel_result(self, error: str = "用户已取消") -> PipelineResult:
        return PipelineResult(
            success=False,
            status="cancelled",
            tracker=self.tracker,
            failed_items=self._collect_failed_items(),
            error=error,
            cancelled=True,
        )

    @staticmethod
    def _remove_temp_doc(path: Path | None) -> None:
        if path is None:
            return
        try:
            if path.exists():
                path.unlink()
        except OSError:
            pass

    @staticmethod
    def _make_temp_doc_copy(src_path: Path) -> Path:
        suffix = src_path.suffix or ".docx"
        fd, temp_path = tempfile.mkstemp(
            suffix=suffix,
            prefix=f".{src_path.stem}.pipeline.",
            dir=str(src_path.parent),
        )
        os.close(fd)
        tmp = Path(temp_path)
        shutil.copy2(str(src_path), str(tmp))
        return tmp

    def _open_processing_doc(self, doc_path: Path) -> tuple[Document, Document, Path | None]:
        """Open a working docx without mutating the user's source file."""
        baseline_doc = None
        try:
            baseline_doc = Document(str(doc_path))
        except Exception:
            baseline_doc = None

        try:
            needs_sanitize = docx_needs_sanitization(str(doc_path))
        except Exception:
            needs_sanitize = False

        if baseline_doc is not None and not needs_sanitize:
            return baseline_doc, copy.deepcopy(baseline_doc), None

        working_doc_path = self._make_temp_doc_copy(doc_path)
        sanitized = False
        if needs_sanitize:
            try:
                sanitize_docx(str(working_doc_path))
                sanitized = True
            except Exception:
                pass

        try:
            doc = Document(str(working_doc_path))
        except Exception as first_exc:
            try:
                sanitize_docx(str(working_doc_path), aggressive=True)
                sanitized = True
                doc = Document(str(working_doc_path))
            except Exception as second_exc:
                if sanitized:
                    raise RuntimeError(f"无法打开文档: {second_exc}") from second_exc
                raise RuntimeError(f"无法打开文档: {first_exc}") from first_exc

        original_doc = copy.deepcopy(baseline_doc if baseline_doc is not None else doc)
        return doc, original_doc, working_doc_path

    def _resolve_scope_targets(
        self,
        doc: Document,
        scope,
        *,
        source_doc_path: Path | None = None,
    ) -> tuple[int | None, list[tuple[int, int]], list[tuple[int, int]], set[int]]:
        body_start = None
        target_page_ranges: list[tuple[int, int]] = []
        target_paragraph_ranges: list[tuple[int, int]] = []
        target_paragraph_indices: set[int] = set()

        if scope.mode == "manual":
            page_ranges_text = str(getattr(scope, "page_ranges_text", "") or "").strip()
            if page_ranges_text:
                target_page_ranges = parse_page_ranges_text(page_ranges_text)
                if not target_page_ranges:
                    raise ValueError("修正范围为空，请填写例如 27-40,44-56")
            else:
                if scope.body_start_index is not None:
                    body_start = scope.body_start_index
                elif scope.body_start_page is not None:
                    body_start = page_number_to_start_paragraph_index(
                        doc,
                        scope.body_start_page,
                        source_doc_path=source_doc_path,
                        require_word=True,
                    )
                    if body_start is None:
                        raise ValueError("正文起始页未命中文档中的有效页码，请检查输入的物理页码。")
                elif scope.body_start_keyword:
                    body_start = self._find_keyword_index(doc, scope.body_start_keyword)

        if target_page_ranges:
            target_paragraph_ranges = page_ranges_to_paragraph_ranges(
                doc,
                target_page_ranges,
                source_doc_path=source_doc_path,
                require_word=True,
            )
            if not target_paragraph_ranges:
                raise ValueError("修正范围未命中文档中的有效页码，请检查输入的物理页码范围。")
            target_paragraph_indices = paragraph_ranges_to_index_set(target_paragraph_ranges)

        return body_start, target_page_ranges, target_paragraph_ranges, target_paragraph_indices

    def _run_legacy(self, doc_path: str) -> PipelineResult:
        """执行完整排版流程"""
        if self._is_cancel_requested():
            return self._cancel_result()
        doc_path = Path(doc_path)
        if not doc_path.exists():
            return PipelineResult(success=False,
                                  status="failed",
                                  error=f"文件不存在: {doc_path}")

        disable_refresh = os.environ.get("DOCX_DISABLE_FIELD_REFRESH", "").strip() == "1"
        planned_refresh_step = bool(self.config.output.final_docx and not disable_refresh)
        planned_fallback_step = bool(
            self.config.output.final_docx
            and getattr(getattr(self.config, "formula_convert", None), "office_fallback_enabled", False)
        )
        # 阶段：加载、结构分析、规则执行、保存、(可选)MathType 兜底、(可选)域刷新、报告
        total_steps = len(self.steps) + 4 + (1 if planned_refresh_step else 0) + (1 if planned_fallback_step else 0)
        step_idx = 0

        # 1. 预处理 & 加载文档
        self._emit_progress(step_idx, total_steps, "正在加载文档…")
        if self._is_cancel_requested():
            return self._cancel_result()
        try:
            sanitize_docx(str(doc_path))
        except Exception:
            pass  # 预处理失败不阻断，继续尝试加载
        if self._is_cancel_requested():
            return self._cancel_result()
        try:
            doc = Document(str(doc_path))
        except Exception as e:
            return PipelineResult(success=False,
                                  status="failed",
                                  error=f"无法打开文档: {e}")

        # 深拷贝原始文档用于对比稿
        original_doc = copy.deepcopy(doc)
        self._last_input_file = str(doc_path)
        step_idx += 1
        if self._is_cancel_requested():
            return self._cancel_result()

        # 2. 构建文档结构树（根据 format_scope 决定自动/手动模式）
        self._emit_progress(step_idx, total_steps, "正在分析文档结构…")
        if self._is_cancel_requested():
            return self._cancel_result()
        doc_tree = DocTree()
        scope = self.config.format_scope
        try:
            body_start, target_page_ranges, target_paragraph_ranges, target_paragraph_indices = (
                self._resolve_scope_targets(doc, scope, source_doc_path=doc_path)
            )
        except (ValueError, RuntimeError) as exc:
            return PipelineResult(
                success=False,
                status="failed",
                error=str(exc),
            )

        if target_page_ranges:
            # Range mode keeps section auto-detection, then restricts edits by paragraph indices.
            doc_tree.build(doc)
        else:
            doc_tree.build(doc, body_start_index=body_start)

        context = {
            "doc_tree": doc_tree,
            "format_scope": scope,
            "source_doc_path": str(doc_path),
            "source_doc_dir": str(doc_path.parent),
        }
        if target_paragraph_ranges:
            context["target_page_ranges"] = target_page_ranges
            context["target_paragraph_ranges"] = target_paragraph_ranges
            context["target_paragraph_indices"] = target_paragraph_indices
        step_idx += 1
        if self._is_cancel_requested():
            return self._cancel_result()

        # 3. 依次执行规则
        range_mode_skip_rules = {
            "toc_format",
            "caption_format",
            "citation_link",
            "header_footer",
            "formula_convert",
            "formula_to_table",
            "equation_table_format",
            "formula_style",
        }
        for rule in self.steps:
            if self._is_cancel_requested():
                return self._cancel_result()
            if target_paragraph_ranges and rule.name in range_mode_skip_rules:
                self._emit_progress(
                    step_idx,
                    total_steps,
                    f"修正范围模式：跳过规则 {rule.description}",
                )
                self.tracker.record(
                    rule_name=rule.name,
                    target="range_scope",
                    section="global",
                    change_type="skip",
                    before="full document operation",
                    after="skipped in page-range mode",
                    paragraph_index=-1,
                )
                step_idx += 1
                continue
            self._emit_progress(step_idx, total_steps,
                                f"正在执行规则：{rule.description}")
            try:
                rule.apply(doc, self.config, self.tracker, context)
            except Exception as e:
                self.tracker.record(
                    rule_name=rule.name,
                    target="pipeline",
                    section="global",
                    change_type="error",
                    before="",
                    after="",
                    paragraph_index=-1,
                    success=False,
                    failure_reason=str(e),
                )
            step_idx += 1
            if self._is_cancel_requested():
                return self._cancel_result()

        # 4. 保存输出文件
        if self._is_cancel_requested():
            return self._cancel_result()
        self._emit_progress(step_idx, total_steps, "正在保存输出文件…")
        output_paths = self._save_outputs(doc, doc_path)
        step_idx += 1

        if self._is_cancel_requested():
            return PipelineResult(
                success=False,
                status="cancelled",
                doc=doc,
                original_doc=original_doc,
                tracker=self.tracker,
                output_paths=output_paths,
                failed_items=self._collect_failed_items(),
                error="用户已取消",
                cancelled=True,
            )

        final_path = output_paths.get("final")
        formula_convert_cfg = getattr(self.config, "formula_convert", None)
        office_fallback_enabled = bool(
            getattr(formula_convert_cfg, "office_fallback_enabled", False)
        )
        office_fallback_timeout = int(
            getattr(formula_convert_cfg, "office_fallback_timeout_sec", 30) or 30
        )
        office_fallback_timeout = max(10, min(office_fallback_timeout, 600))
        if final_path and office_fallback_enabled:
            self._emit_progress(step_idx, total_steps, "正在尝试 MathType Office 兜底转换…")
            if target_paragraph_ranges:
                self.tracker.record(
                    rule_name="pipeline",
                    target="mathtype_office_fallback",
                    section="formula",
                    change_type="skip",
                    before="page-range mode",
                    after="MathType Office 兜底已跳过：page-range mode",
                    paragraph_index=-1,
                    success=True,
                    failure_reason=None,
                )
            else:
                changed, detail, stats = apply_mathtype_office_fallback(
                    final_path,
                    timeout_sec=office_fallback_timeout,
                )
                found_count = int(stats.get("found", 0) or 0)
                replaced_count = int(stats.get("replaced", 0) or 0)
                extracted_count = int(stats.get("extracted", 0) or 0)
                self.tracker.record(
                    rule_name="pipeline",
                    target="mathtype_office_fallback",
                    section="formula",
                    change_type="fallback",
                    before=f"found={found_count}",
                    after=(
                        f"MathType Office 兜底已应用：extracted={extracted_count}, replaced={replaced_count}"
                        if changed
                        else f"MathType Office 兜底未应用：{detail}"
                    ),
                    paragraph_index=-1,
                    success=True,
                    failure_reason=None,
                )
            step_idx += 1
        if final_path and not disable_refresh:
            self._emit_progress(step_idx, total_steps, "正在更新目录与页码…")
            try:
                refresh_timeout = int(
                    os.environ.get("DOCX_FIELD_REFRESH_TIMEOUT_SEC", "10")
                )
            except ValueError:
                refresh_timeout = 10
            refresh_timeout = max(10, min(refresh_timeout, 600))
            ok, detail = refresh_doc_fields_with_word(
                final_path, timeout_sec=refresh_timeout
            )
            self.tracker.record(
                rule_name="pipeline",
                target="final_doc_fields",
                section="global",
                change_type="refresh",
                before="TOC/fields pending",
                after="TOC/fields refreshed" if ok else "TOC/fields refresh skipped",
                paragraph_index=-1,
                success=ok,
                failure_reason=None if ok else detail,
            )
            if ok:
                toc_ok, toc_detail = self._normalize_toc_after_field_refresh(final_path)
                self.tracker.record(
                    rule_name="pipeline",
                    target="post_refresh_toc_format",
                    section="global",
                    change_type="format",
                    before="Word refreshed TOC fields",
                    after="TOC styles normalized" if toc_ok else "TOC style normalization skipped",
                    paragraph_index=-1,
                    success=toc_ok,
                    failure_reason=None if toc_ok else toc_detail,
                )
            step_idx += 1
        elif final_path and disable_refresh:
            self.tracker.record(
                rule_name="pipeline",
                target="final_doc_fields",
                section="global",
                change_type="refresh",
                before="TOC/fields pending",
                after="TOC/fields refresh disabled by env",
                paragraph_index=-1,
                success=True,
                failure_reason=None,
            )

        if self._is_cancel_requested():
            return PipelineResult(
                success=False,
                status="cancelled",
                doc=doc,
                original_doc=original_doc,
                tracker=self.tracker,
                output_paths=output_paths,
                failed_items=self._collect_failed_items(),
                error="用户已取消",
                cancelled=True,
            )

        compare_path = output_paths.get("compare")
        if compare_path and original_doc is not None:
            try:
                if final_path and Path(final_path).exists():
                    # Compare should be based on the final on-disk deliverable
                    # (after optional field refresh) to avoid report drift.
                    final_doc_for_compare = Document(str(final_path))
                else:
                    # Fallback: clone in-memory final doc safely.
                    from io import BytesIO
                    buf = BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    final_doc_for_compare = Document(buf)

                generate_compare_doc(
                    original_doc=original_doc,
                    tracker=self.tracker,
                    output_path=compare_path,
                    final_doc=final_doc_for_compare,
                    include_text=bool(self.config.output.compare_text),
                    include_formatting=bool(self.config.output.compare_formatting),
                )
            except Exception as e:
                self.tracker.record(
                    rule_name="pipeline",
                    target="compare_doc",
                    section="global",
                    change_type="error",
                    before="",
                    after="",
                    paragraph_index=-1,
                    success=False,
                    failure_reason=f"对比稿生成失败: {e}",
                )

        # 5. 生成报告
        self._emit_progress(step_idx, total_steps, "正在生成处理报告…")
        self._generate_reports(output_paths, input_file=str(doc_path))

        return self._finalize_result(
            doc=doc,
            original_doc=original_doc,
            output_paths=output_paths,
        )

    def run(self, doc_path: str) -> PipelineResult:
        """Execute the full formatting pipeline."""
        if self._is_cancel_requested():
            return self._cancel_result()

        doc_path = Path(doc_path)
        if not doc_path.exists():
            return PipelineResult(
                success=False,
                status="failed",
                error=f"文件不存在: {doc_path}",
            )

        disable_refresh = os.environ.get("DOCX_DISABLE_FIELD_REFRESH", "").strip() == "1"
        planned_refresh_step = bool(self.config.output.final_docx and not disable_refresh)
        planned_fallback_step = bool(
            self.config.output.final_docx
            and getattr(getattr(self.config, "formula_convert", None), "office_fallback_enabled", False)
        )
        total_steps = len(self.steps) + 4 + (1 if planned_refresh_step else 0) + (1 if planned_fallback_step else 0)
        step_idx = 0
        working_doc_path: Path | None = None

        try:
            self._emit_progress(step_idx, total_steps, "正在加载文档…")
            if self._is_cancel_requested():
                return self._cancel_result()
            try:
                doc, original_doc, working_doc_path = self._open_processing_doc(doc_path)
            except RuntimeError as exc:
                return PipelineResult(
                    success=False,
                    status="failed",
                    error=str(exc),
                )
            self._last_input_file = str(doc_path)

            step_idx += 1
            if self._is_cancel_requested():
                return self._cancel_result()

            self._emit_progress(step_idx, total_steps, "正在分析文档结构…")
            if self._is_cancel_requested():
                return self._cancel_result()

            doc_tree = DocTree()
            scope = self.config.format_scope
            source_doc_for_page_scope = working_doc_path if working_doc_path is not None else doc_path
            try:
                body_start, target_page_ranges, target_paragraph_ranges, target_paragraph_indices = (
                    self._resolve_scope_targets(
                        doc,
                        scope,
                        source_doc_path=source_doc_for_page_scope,
                    )
                )
            except (ValueError, RuntimeError) as exc:
                return PipelineResult(
                    success=False,
                    status="failed",
                    error=str(exc),
                )

            if target_page_ranges:
                doc_tree.build(doc)
            else:
                doc_tree.build(doc, body_start_index=body_start)

            context = {
                "doc_tree": doc_tree,
                "format_scope": scope,
                "source_doc_path": str(doc_path),
                "source_doc_dir": str(doc_path.parent),
            }
            if target_paragraph_ranges:
                context["target_page_ranges"] = target_page_ranges
                context["target_paragraph_ranges"] = target_paragraph_ranges
                context["target_paragraph_indices"] = target_paragraph_indices

            step_idx += 1
            if self._is_cancel_requested():
                return self._cancel_result()

            range_mode_skip_rules = {
                "page_setup",
                "style_manager",
                "toc_format",
                "caption_format",
                "citation_link",
                "header_footer",
                "formula_convert",
                "formula_to_table",
                "equation_table_format",
                "formula_style",
            }
            for rule in self.steps:
                if self._is_cancel_requested():
                    return self._cancel_result()
                if target_paragraph_ranges and rule.name in range_mode_skip_rules:
                    self._emit_progress(
                        step_idx,
                        total_steps,
                        f"修正范围模式：跳过规则 {rule.description}",
                    )
                    self.tracker.record(
                        rule_name=rule.name,
                        target="range_scope",
                        section="global",
                        change_type="skip",
                        before="full document operation",
                        after="skipped in page-range mode",
                        paragraph_index=-1,
                    )
                    step_idx += 1
                    continue

                self._emit_progress(step_idx, total_steps, f"正在执行规则：{rule.description}")
                try:
                    rule.apply(doc, self.config, self.tracker, context)
                except Exception as exc:
                    self.tracker.record(
                        rule_name=rule.name,
                        target="pipeline",
                        section="global",
                        change_type="error",
                        before="",
                        after="",
                        paragraph_index=-1,
                        success=False,
                        failure_reason=str(exc),
                    )
                step_idx += 1
                if self._is_cancel_requested():
                    return self._cancel_result()

            if self._is_cancel_requested():
                return self._cancel_result()

            self._emit_progress(step_idx, total_steps, "正在保存输出文件…")
            output_paths = self._save_outputs(doc, doc_path)
            step_idx += 1

            if self._is_cancel_requested():
                return PipelineResult(
                    success=False,
                    status="cancelled",
                    doc=doc,
                    original_doc=original_doc,
                    tracker=self.tracker,
                    output_paths=output_paths,
                    failed_items=self._collect_failed_items(),
                    error="用户已取消",
                    cancelled=True,
                )

            final_path = output_paths.get("final")
            formula_convert_cfg = getattr(self.config, "formula_convert", None)
            office_fallback_enabled = bool(
                getattr(formula_convert_cfg, "office_fallback_enabled", False)
            )
            office_fallback_timeout = int(
                getattr(formula_convert_cfg, "office_fallback_timeout_sec", 30) or 30
            )
            office_fallback_timeout = max(10, min(office_fallback_timeout, 600))
            if final_path and office_fallback_enabled:
                self._emit_progress(step_idx, total_steps, "正在尝试 MathType Office 兜底转换…")
                if target_paragraph_ranges:
                    self.tracker.record(
                        rule_name="pipeline",
                        target="mathtype_office_fallback",
                        section="formula",
                        change_type="skip",
                        before="page-range mode",
                        after="MathType Office 兜底已跳过：page-range mode",
                        paragraph_index=-1,
                        success=True,
                        failure_reason=None,
                    )
                else:
                    changed, detail, stats = apply_mathtype_office_fallback(
                        final_path,
                        timeout_sec=office_fallback_timeout,
                    )
                    found_count = int(stats.get("found", 0) or 0)
                    replaced_count = int(stats.get("replaced", 0) or 0)
                    extracted_count = int(stats.get("extracted", 0) or 0)
                    self.tracker.record(
                        rule_name="pipeline",
                        target="mathtype_office_fallback",
                        section="formula",
                        change_type="fallback",
                        before=f"found={found_count}",
                        after=(
                            f"MathType Office 兜底已应用：extracted={extracted_count}, replaced={replaced_count}"
                            if changed
                            else f"MathType Office 兜底未应用：{detail}"
                        ),
                        paragraph_index=-1,
                        success=True,
                        failure_reason=None,
                    )
                step_idx += 1
            if final_path and not disable_refresh:
                self._emit_progress(step_idx, total_steps, "正在更新目录与页码…")
                try:
                    refresh_timeout = int(os.environ.get("DOCX_FIELD_REFRESH_TIMEOUT_SEC", "10"))
                except ValueError:
                    refresh_timeout = 10
                refresh_timeout = max(10, min(refresh_timeout, 600))
                ok, detail = refresh_doc_fields_with_word(final_path, timeout_sec=refresh_timeout)
                self.tracker.record(
                    rule_name="pipeline",
                    target="final_doc_fields",
                    section="global",
                    change_type="refresh",
                    before="TOC/fields pending",
                    after="TOC/fields refreshed" if ok else "TOC/fields refresh skipped",
                    paragraph_index=-1,
                    success=ok,
                    failure_reason=None if ok else detail,
                )
                if ok:
                    toc_ok, toc_detail = self._normalize_toc_after_field_refresh(final_path)
                    self.tracker.record(
                        rule_name="pipeline",
                        target="post_refresh_toc_format",
                        section="global",
                        change_type="format",
                        before="Word refreshed TOC fields",
                        after="TOC styles normalized" if toc_ok else "TOC style normalization skipped",
                        paragraph_index=-1,
                        success=toc_ok,
                        failure_reason=None if toc_ok else toc_detail,
                    )
                step_idx += 1
            elif final_path and disable_refresh:
                self.tracker.record(
                    rule_name="pipeline",
                    target="final_doc_fields",
                    section="global",
                    change_type="refresh",
                    before="TOC/fields pending",
                    after="TOC/fields refresh disabled by env",
                    paragraph_index=-1,
                    success=True,
                    failure_reason=None,
                )

            if self._is_cancel_requested():
                return PipelineResult(
                    success=False,
                    status="cancelled",
                    doc=doc,
                    original_doc=original_doc,
                    tracker=self.tracker,
                    output_paths=output_paths,
                    failed_items=self._collect_failed_items(),
                    error="用户已取消",
                    cancelled=True,
                )

            compare_path = output_paths.get("compare")
            if compare_path and original_doc is not None:
                try:
                    if final_path and Path(final_path).exists():
                        final_doc_for_compare = Document(str(final_path))
                    else:
                        from io import BytesIO

                        buf = BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        final_doc_for_compare = Document(buf)

                    generate_compare_doc(
                        original_doc=original_doc,
                        tracker=self.tracker,
                        output_path=compare_path,
                        final_doc=final_doc_for_compare,
                        include_text=bool(self.config.output.compare_text),
                        include_formatting=bool(self.config.output.compare_formatting),
                    )
                except Exception as exc:
                    self.tracker.record(
                        rule_name="pipeline",
                        target="compare_doc",
                        section="global",
                        change_type="error",
                        before="",
                        after="",
                        paragraph_index=-1,
                        success=False,
                        failure_reason=f"对比稿生成失败: {exc}",
                    )

            self._emit_progress(step_idx, total_steps, "正在生成处理报告…")
            self._generate_reports(output_paths, input_file=str(doc_path))

            return self._finalize_result(
                doc=doc,
                original_doc=original_doc,
                output_paths=output_paths,
            )
        finally:
            self._remove_temp_doc(working_doc_path)

    def _normalize_toc_after_field_refresh(self, doc_path: str) -> tuple[bool, str]:
        """Normalize TOC styles after Word refresh introduces locale-specific TOC styles."""
        try:
            refreshed_doc = Document(str(doc_path))
        except Exception as exc:
            return False, f"unable to reopen refreshed docx: {exc}"

        try:
            normalized_styles = _normalize_toc_styles_in_doc(refreshed_doc, self.config)
            toc_rule = TocFormatRule()
            doc_tree = DocTree()
            doc_tree.build(refreshed_doc)
            toc_section = doc_tree.get_section("toc")
            if toc_section is None:
                toc_section = toc_rule._fallback_toc_section(refreshed_doc, self.config)

            if toc_section is not None:
                toc_rule._format_existing_toc_paras(
                    refreshed_doc,
                    self.config,
                    toc_section=toc_section,
                )
            elif normalized_styles <= 0:
                return True, "no toc section detected"

            refreshed_doc.save(str(doc_path))
            return True, f"normalized_styles={normalized_styles}"
        except Exception as exc:
            return False, f"toc normalization failed: {exc}"

    def _generate_reports(self, output_paths: dict, *, input_file: str = "") -> None:
        """生成 JSON / Markdown 报告文件（如果路径存在于 output_paths）。"""
        from src.report.collector import collect_report
        from src.report.json_report import generate_json_report
        from src.report.markdown_report import generate_markdown_report
        from src.engine.rules.base import ValidationIssue

        context_issues: list[ValidationIssue] = []
        for rec in self.tracker.records:
            if rec.rule_name == "validation":
                context_issues.append(ValidationIssue(
                    level="warning" if rec.success else "error",
                    rule_name=rec.rule_name,
                    message=rec.before,
                    location=rec.target,
                ))

        report_data = collect_report(
            self.tracker,
            scene_name=getattr(self.config, "name", ""),
            input_file=input_file or self._last_input_file,
            validation_issues=context_issues,
        )

        json_path = output_paths.get("report_json")
        if json_path:
            try:
                generate_json_report(report_data, json_path)
            except Exception:
                pass

        md_path = output_paths.get("report_md")
        if md_path:
            try:
                generate_markdown_report(report_data, md_path)
            except Exception:
                pass

    @staticmethod
    def _find_keyword_index(doc: Document, keyword: str) -> int | None:
        """在文档中查找包含关键字的段落索引"""
        import re
        kw = re.sub(r'\s+', '', keyword)
        for i, para in enumerate(doc.paragraphs):
            text = re.sub(r'\s+', '', para.text.strip())
            if kw in text:
                return i
        return None

    @staticmethod
    def _find_page_start_index(doc: Document, target_page: int) -> int | None:
        """根据目标页码找到该页出现的第一个段落索引。"""
        return page_number_to_start_paragraph_index(doc, target_page)

    def _save_outputs(self, doc: Document, src_path: Path) -> dict:
        """保存输出文件，返回路径字典

        策略：原始文件保持不变，新生成文件以 _new 后缀保存在同目录。
        """
        stem = src_path.stem
        out_dir = src_path.parent
        sub_dir = out_dir / f"{stem}_排版附件"
        sub_dir.mkdir(exist_ok=True)
        self._cleanup_disabled_outputs(src_path=src_path, sub_dir=sub_dir)
        paths = {"sub_dir": str(sub_dir)}

        # 原始文件保持不变，不做备份/重命名
        paths["original"] = str(src_path)

        # 最终稿：保存为 {stem}_new.docx（不覆盖原文件）
        if self.config.output.final_docx:
            primary_final_path = out_dir / f"{stem}_new.docx"
            final_path = self._save_doc_with_fallback(doc, primary_final_path)
            paths["final"] = str(final_path)
            if final_path != primary_final_path:
                paths["final_primary"] = str(primary_final_path)

        # 对比稿 → 子文件夹
        if self.config.output.compare_docx:
            paths["compare"] = str(sub_dir / f"{stem}_对比稿.docx")

        # 报告 → 子文件夹
        if self.config.output.report_json:
            paths["report_json"] = str(sub_dir / f"{stem}_报告.json")
        if self.config.output.report_markdown:
            paths["report_md"] = str(sub_dir / f"{stem}_报告.md")

        return paths

    def _cleanup_disabled_outputs(self, *, src_path: Path, sub_dir: Path) -> None:
        """Remove stale artifacts for outputs disabled in the current run."""
        stem = src_path.stem

        stale_paths: list[Path] = []
        if not self.config.output.final_docx:
            stale_paths.append(src_path.parent / f"{stem}_new.docx")
            stale_paths.extend(sorted(src_path.parent.glob(f"{stem}_new_*.docx")))

        if not self.config.output.compare_docx:
            stale_paths.append(sub_dir / f"{stem}_对比稿.docx")

        if not self.config.output.report_json:
            stale_paths.append(sub_dir / f"{stem}_报告.json")

        if not self.config.output.report_markdown:
            stale_paths.append(sub_dir / f"{stem}_报告.md")

        for path in stale_paths:
            try:
                if path.exists():
                    path.unlink()
            except OSError:
                continue

    @staticmethod
    def _save_doc_with_fallback(doc: Document, primary_path: Path) -> Path:
        """保存文档，若目标文件被占用则自动尝试递增后缀路径。"""
        try:
            doc.save(str(primary_path))
            return primary_path
        except PermissionError:
            pass

        base_dir = primary_path.parent
        stem = primary_path.stem
        suffix = primary_path.suffix
        for idx in range(2, 100):
            alt = base_dir / f"{stem}_{idx}{suffix}"
            try:
                doc.save(str(alt))
                return alt
            except PermissionError:
                continue

        # 连续失败时抛出原始异常语义
        doc.save(str(primary_path))
        return primary_path

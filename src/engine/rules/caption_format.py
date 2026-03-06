"""图表题注自动编号与格式修正规则"""

import re
from dataclasses import dataclass, field
from lxml import etree
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.engine.rules.base import BaseRule
from src.engine.change_tracker import ChangeTracker
from src.scene.schema import SceneConfig, StyleConfig
from src.utils.line_spacing import apply_line_spacing

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


# ── 题注识别正则 ──

_NUM_PART = (
    r'[一二三四五六七八九十百零]+'
    r'|（\d+(?:[.\-]\d+)*）'
    r'|\(\d+(?:[.\-]\d+)*\)'
    r'|\d+(?:[.\-]\d+)*'
)

_RE_FIG_CAPTION = re.compile(
    r'^(?P<prefix>图|Figure|Fig\.?)\s*'
    r'(?P<number>' + _NUM_PART + r')'
    r'(?P<sep>[\s\u3000\t]*)'
    r'(?P<title>.*)',
    re.IGNORECASE,
)

_RE_TBL_CAPTION = re.compile(
    r'^(?P<prefix>表|Table)\s*'
    r'(?P<number>' + _NUM_PART + r')'
    r'(?P<sep>[\s\u3000\t]*)'
    r'(?P<title>.*)',
    re.IGNORECASE,
)

# 图注续行：仅识别 "(a)/(b)/(c)" 这类子图标记开头，尽量避免误判正文。
_RE_FIG_CAPTION_CONT_LINE = re.compile(
    r'^\s*[（(]\s*[A-Za-z]\s*[）)]'
)


# ── 数据结构 ──

@dataclass
class _FigureInfo:
    para_index: int
    element: object
    caption_para_index: int | None = None
    caption_para_indices: list[int] = field(default_factory=list)

@dataclass
class _TableInfo:
    table_index: int
    element: object
    para_index: int          # 虚拟段落索引（用于章号映射）
    caption_para_index: int | None = None


# ── 工具函数 ──

_RE_EQ_NUM = re.compile(
    r'^\s*(?:[\.．…·•]{2,}\s*)?[\(（]\s*\d+(?:[\-\.]\d+)+\s*[\)）]\s*$'
)
_RE_EQ_NUM_SUFFIX = re.compile(
    r'^(?P<prefix>.*?)(?:[\.．…·•]{2,}\s*)?[\(（]\s*\d+(?:[\-\.]\d+)+\s*[\)）]\s*$'
)
_RE_EQ_ARROW = re.compile(r'(→|←|↔|⇌|->|<-|<->)')
_MAX_EQ_TABLE_ROWS = 20


def _get_cell_text(tc) -> str:
    """提取单元格纯文本"""
    parts = []
    for p in tc.findall(_w("p")):
        for r in p.findall(_w("r")):
            t = r.find(_w("t"))
            if t is not None and t.text:
                parts.append(t.text)
    return "".join(parts)


def _looks_like_formula_prefix(prefix: str) -> bool:
    """判断编号前缀是否像公式表达式（覆盖“表达式+(2.8)”同格场景）。"""
    if not prefix:
        return False
    if re.search(r'[=→←↔±×÷∑∫≤≥<>]', prefix):
        return True
    # 仅有一个连字符常见于普通文本（如 A-1），至少两个基础运算符再认定为公式
    op_count = len(re.findall(r'[+\-*/]', prefix))
    return op_count >= 2


def _looks_like_equation_text(text: str) -> bool:
    """判断文本是否像公式/反应式表达。"""
    s = (text or "").strip()
    if not s:
        return False
    if _RE_EQ_ARROW.search(s):
        return True
    if re.search(r'[A-Za-z0-9\)\]]\s*=\s*[A-Za-z0-9\(\[]', s):
        return True
    # 化学式/反应式常见：含字母且含多个运算符（兼容 Unicode 负号）
    op_count = len(re.findall(r'[+\-−*/]', s))
    alpha_count = len(re.findall(r'[A-Za-zΑ-Ωα-ω]', s))
    if op_count >= 2 and alpha_count >= 2:
        return True
    return False


def _is_equation_table(tbl_el) -> bool:
    """检测是否为公式表格。

    判据（满足任一即可）：
    1. 含 oMath/oMathPara 或 MathType OLE 对象
    2. 小表格（行数受限）且右侧单元格命中公式编号：
       - 纯编号：(x-y)/(x.y)/……(x.y)
       - 末尾编号：表达式\t(x.y) / 表达式……(x.y)
    3. 无编号时，若多行内容高度像反应式/公式，也视为公式表格
    """
    if tbl_el.findall(f".//{{{_M_NS}}}oMathPara"):
        return True
    if tbl_el.findall(f".//{{{_M_NS}}}oMath"):
        return True
    if tbl_el.findall(f".//{{{_W_NS}}}r/{{{_W_NS}}}object"):
        return True

    rows = tbl_el.findall(_w("tr"))
    if not rows or len(rows) > _MAX_EQ_TABLE_ROWS:
        return False

    formula_like_rows = 0
    non_empty_rows = 0
    max_cols = 0
    for tr in rows:
        cells = tr.findall(_w("tc"))
        if not cells:
            continue
        max_cols = max(max_cols, len(cells))

        cell_texts = [(_get_cell_text(tc) or "").strip() for tc in cells]
        non_empty = [txt for txt in cell_texts if txt]
        if non_empty:
            non_empty_rows += 1

        # 从右向左优先匹配“纯编号”单元格，兼容末列空白的情况
        for txt in reversed(cell_texts):
            if txt and _RE_EQ_NUM.match(txt):
                return True

        # 兼容“表达式 + 编号”同单元格写法（Word 可能丢失 \t 为 <w:tab/>）
        right_non_empty = next((txt for txt in reversed(cell_texts) if txt), "")
        if right_non_empty:
            m = _RE_EQ_NUM_SUFFIX.match(right_non_empty)
            if m and _looks_like_formula_prefix(m.group("prefix")):
                return True

        # 无编号行：按反应式/公式特征计数
        if non_empty:
            merged = " ".join(
                txt for txt in non_empty if not _RE_EQ_NUM.match(txt)
            )
            if _looks_like_equation_text(merged):
                formula_like_rows += 1

    # 无明确编号时的兜底识别：
    # 小列数表格（通常 2~3 列）且大多数非空行是公式/反应式
    if non_empty_rows and max_cols <= 4 and formula_like_rows >= max(1, non_empty_rows - 1):
        return True
    return False


def _para_has_image(para) -> bool:
    """检查段落是否包含图片"""
    el = para._element
    if el.findall(f"{{{_W_NS}}}r/{{{_W_NS}}}drawing"):
        return True
    if el.findall(f"{{{_W_NS}}}r/{{{_W_NS}}}pict"):
        return True
    if el.findall(f".//{{{_WP_NS}}}inline"):
        return True
    if el.findall(f".//{{{_WP_NS}}}anchor"):
        return True
    return False


def _build_chapter_ranges(headings, body_end: int):
    """从标题列表构建章号范围表 [(start, end, chapter_num), ...]"""
    chapters = []
    num = 0
    for h in headings:
        if h.level == "heading1":
            num += 1
            chapters.append((h.para_index, num))
    ranges = []
    for i, (start, n) in enumerate(chapters):
        end = chapters[i + 1][0] - 1 if i + 1 < len(chapters) else body_end
        ranges.append((start, end, n))
    return ranges


def _get_chapter_num(para_index: int, chapter_ranges) -> int:
    """查询段落所属章号，章前返回 0"""
    for start, end, num in chapter_ranges:
        if start <= para_index <= end:
            return num
    return 0


def _build_body_map(doc, body_start, body_end):
    """遍历 w:body 子节点，构建段落/表格交错序列（仅 body 范围内）。

    返回 [("para", para_idx, element), ("table", tbl_idx, element), ...]
    """
    body_el = doc.element.body
    items = []
    para_idx = 0
    tbl_idx = 0
    last_para_idx = -1
    for child in body_el:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            last_para_idx = para_idx
            if body_start <= para_idx <= body_end:
                items.append(("para", para_idx, child))
            para_idx += 1
        elif tag == "tbl":
            # 表格属于 body 当且仅当其前一个段落索引落在 body 范围内。
            # 必须同时满足 <= body_end，避免把正文之后的表格误并入正文。
            if body_start <= last_para_idx <= body_end:
                items.append(("table", tbl_idx, child))
            tbl_idx += 1
    return items


def _detect_figures(doc, body_items):
    """扫描 body 元素序列，找出图片段落并配对题注"""
    figures = []
    for i, (typ, idx, elem) in enumerate(body_items):
        if typ != "para":
            continue
        para = doc.paragraphs[idx]
        if not _para_has_image(para):
            continue
        # 向后查找第一个非空段落
        cap_idx = None
        cap_indices: list[int] = []
        for j in range(i + 1, len(body_items)):
            nt, ni, ne = body_items[j]
            if nt != "para":
                break
            text = doc.paragraphs[ni].text.strip()
            if not text:
                continue
            if _RE_FIG_CAPTION.match(text):
                cap_idx = ni
                cap_indices = [ni]
                cap_indices.extend(
                    _collect_figure_caption_continuations(doc, body_items, j)
                )
            break
        figures.append(_FigureInfo(idx, elem, cap_idx, cap_indices))
    return figures


def _is_figure_caption_continuation(text: str) -> bool:
    """Whether a paragraph can be treated as figure-caption continuation.

    Conservative strategy:
    - Must start with a sub-figure marker like "(c)" / "（d）"
    - Keep a reasonable max length to avoid swallowing normal paragraphs
    """
    s = (text or "").strip()
    if not s or len(s) > 180:
        return False
    return bool(_RE_FIG_CAPTION_CONT_LINE.match(s))


def _collect_figure_caption_continuations(doc, body_items, caption_pos: int) -> list[int]:
    """Collect immediate continuation lines after a detected figure caption head."""
    indices: list[int] = []
    for k in range(caption_pos + 1, len(body_items)):
        typ, para_idx, _ = body_items[k]
        if typ != "para":
            break
        text = (doc.paragraphs[para_idx].text or "").strip()
        if not text:
            break
        if not _is_figure_caption_continuation(text):
            break
        indices.append(para_idx)
    return indices


def _detect_tables(doc, body_items):
    """扫描 body 元素序列，找出表格并配对题注"""
    # 预计算公式表格在 body_items 中的位置，供向前搜索时跳过
    eq_table_positions = set()
    for k, (t, _, e) in enumerate(body_items):
        if t == "table" and _is_equation_table(e):
            eq_table_positions.add(k)

    tables = []
    for i, (typ, idx, elem) in enumerate(body_items):
        if typ != "table":
            continue
        # 公式表格跳过题注处理
        if i in eq_table_positions:
            continue
        # 虚拟段落索引：取表格前最近的段落索引
        virt = 0
        for j in range(i - 1, -1, -1):
            if body_items[j][0] == "para":
                virt = body_items[j][1]
                break
        # 向前查找第一个非空段落（跳过公式表格，但遇到数据表格仍停止）
        cap_idx = None
        for j in range(i - 1, -1, -1):
            pt, pi, pe = body_items[j]
            if pt == "table":
                if j in eq_table_positions:
                    continue  # 跳过公式表格，继续向前搜索
                break         # 遇到数据表格，停止搜索
            if pt != "para":
                break
            text = doc.paragraphs[pi].text.strip()
            if not text:
                continue
            if _RE_TBL_CAPTION.match(text):
                cap_idx = pi
            break
        tables.append(_TableInfo(idx, elem, virt, cap_idx))
    return tables


# ── 域代码构建 ──

def _make_run(parent, text=None, rpr_src=None):
    """创建 w:r 元素，可选复制 rPr 格式"""
    r = etree.SubElement(parent, _w("r"))
    if rpr_src is not None:
        import copy
        r.append(copy.deepcopy(rpr_src))
    if text is not None:
        t = etree.SubElement(r, _w("t"))
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _make_field(parent, instr_text, result_text, rpr_src=None):
    """在 parent 中追加一组完整的域代码 runs（begin/instrText/separate/result/end）"""
    # begin
    r_begin = _make_run(parent, rpr_src=rpr_src)
    fc_begin = etree.SubElement(r_begin, _w("fldChar"))
    fc_begin.set(_w("fldCharType"), "begin")

    # instrText
    r_instr = _make_run(parent, rpr_src=rpr_src)
    it = etree.SubElement(r_instr, _w("instrText"))
    it.text = instr_text
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # separate
    r_sep = _make_run(parent, rpr_src=rpr_src)
    fc_sep = etree.SubElement(r_sep, _w("fldChar"))
    fc_sep.set(_w("fldCharType"), "separate")

    # result (显示值，打开文档后 Word 会自动更新)
    _make_run(parent, text=result_text, rpr_src=rpr_src)

    # end
    r_end = _make_run(parent, rpr_src=rpr_src)
    fc_end = etree.SubElement(r_end, _w("fldChar"))
    fc_end.set(_w("fldCharType"), "end")


def _make_run_element(text=None, rpr_src=None):
    """创建独立的 w:r 元素（不挂载到 parent），用于 insert 操作"""
    import copy
    r = etree.Element(_w("r"))
    if rpr_src is not None:
        r.append(copy.deepcopy(rpr_src))
    if text is not None:
        t = etree.SubElement(r, _w("t"))
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _make_field_elements(instr_text, result_text, rpr_src=None):
    """创建域代码的 5 个独立 run 元素列表（begin/instr/separate/result/end）"""
    import copy
    runs = []
    # begin
    r_begin = _make_run_element(rpr_src=rpr_src)
    fc = etree.SubElement(r_begin, _w("fldChar"))
    fc.set(_w("fldCharType"), "begin")
    runs.append(r_begin)
    # instrText
    r_instr = _make_run_element(rpr_src=rpr_src)
    it = etree.SubElement(r_instr, _w("instrText"))
    it.text = instr_text
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    runs.append(r_instr)
    # separate
    r_sep = _make_run_element(rpr_src=rpr_src)
    fc2 = etree.SubElement(r_sep, _w("fldChar"))
    fc2.set(_w("fldCharType"), "separate")
    runs.append(r_sep)
    # result
    runs.append(_make_run_element(text=result_text, rpr_src=rpr_src))
    # end
    r_end = _make_run_element(rpr_src=rpr_src)
    fc3 = etree.SubElement(r_end, _w("fldChar"))
    fc3.set(_w("fldCharType"), "end")
    runs.append(r_end)
    return runs


def _parse_caption_numbering_format(numbering_format: str | None) -> tuple[bool, str]:
    """Parse caption numbering mode to (include_chapter, chapter_seq_separator).

    Supports Word-compatible separators:
      .  句点    -  连字符    :  冒号    —  长划线(em dash)    –  短划线(en dash)
    """
    raw = (numbering_format or "chapter.seq").strip().lower().replace(" ", "")
    _CHAPTER_FORMATS = {
        "chapter.seq": ".",
        "chapter-seq": "-",
        "chapter_seq": "_",
        "chapter/seq": "/",
        "chapterseq": "",
        "chapter:seq": ":",
        "chapter\u2014seq": "\u2014",   # em dash
        "chapter\u2013seq": "\u2013",   # en dash
    }
    if raw in _CHAPTER_FORMATS:
        return True, _CHAPTER_FORMATS[raw]
    if raw in {"seq", "arabic", "number"}:
        return False, ""
    m = re.match(r"^chapter([.\-_/::\u2014\u2013]?)seq$", raw)
    if m:
        return True, m.group(1)
    return True, "."


def _compose_caption_number_text(chapter_num, seq_num, numbering_format: str | None) -> str:
    include_chapter, chapter_sep = _parse_caption_numbering_format(numbering_format)
    if include_chapter:
        return f"{chapter_num}{chapter_sep}{seq_num}"
    return str(seq_num)


def _append_caption_number_field_runs(target_runs: list, prefix, chapter_num, seq_num,
                                      sep, numbering_format, rpr_src=None):
    include_chapter, chapter_sep = _parse_caption_numbering_format(numbering_format)
    target_runs.append(_make_run_element(text=prefix, rpr_src=rpr_src))
    if include_chapter:
        target_runs.extend(_make_field_elements(
            " STYLEREF 1 \\s ", str(chapter_num), rpr_src=rpr_src))
        if chapter_sep:
            target_runs.append(_make_run_element(text=chapter_sep, rpr_src=rpr_src))
        target_runs.extend(_make_field_elements(
            f" SEQ {prefix} \\* ARABIC \\s 1 ", str(seq_num), rpr_src=rpr_src))
    else:
        target_runs.extend(_make_field_elements(
            f" SEQ {prefix} \\* ARABIC ", str(seq_num), rpr_src=rpr_src))
    target_runs.append(_make_run_element(text=sep, rpr_src=rpr_src))


def _add_caption_number_fields(parent, prefix, chapter_num, seq_num,
                               sep, numbering_format, rpr_src=None):
    include_chapter, chapter_sep = _parse_caption_numbering_format(numbering_format)
    _make_run(parent, text=prefix, rpr_src=rpr_src)
    if include_chapter:
        _make_field(parent, " STYLEREF 1 \\s ", str(chapter_num), rpr_src=rpr_src)
        if chapter_sep:
            _make_run(parent, text=chapter_sep, rpr_src=rpr_src)
        _make_field(parent, f" SEQ {prefix} \\* ARABIC \\s 1 ",
                    str(seq_num), rpr_src=rpr_src)
    else:
        _make_field(parent, f" SEQ {prefix} \\* ARABIC ",
                    str(seq_num), rpr_src=rpr_src)
    _make_run(parent, text=sep, rpr_src=rpr_src)


def _first_text_run(el):
    """Return the first run containing visible text."""
    for r in el.findall(_w("r")):
        t_el = r.find(_w("t"))
        if t_el is not None and t_el.text:
            return r
    return None


def _extract_caption_rpr(el):
    """Extract a reusable run style template from existing runs."""
    import copy as _copy
    rpr_src = None
    for r in el.findall(_w("r")):
        rpr = r.find(_w("rPr"))
        if rpr is not None:
            rpr_src = _copy.deepcopy(rpr)
            # Caption numbering should not inherit italic style.
            for tag in ("i", "iCs"):
                it = rpr_src.find(_w(tag))
                if it is not None:
                    rpr_src.remove(it)
            break
    return rpr_src


def _strip_caption_number_prefix(el, full_text: str):
    """Strip caption prefix/number from text runs and keep title runs."""
    fig_m = _RE_FIG_CAPTION.match(full_text)
    tbl_m = _RE_TBL_CAPTION.match(full_text)
    m = fig_m or tbl_m
    strip_chars = m.start("title") if m else 0

    remaining_strip = strip_chars
    for r in el.findall(_w("r")):
        t2 = r.find(_w("t"))
        if t2 is None or not t2.text:
            continue
        if remaining_strip >= len(t2.text):
            remaining_strip -= len(t2.text)
            t2.text = ""
        else:
            t2.text = t2.text[remaining_strip:]
            remaining_strip = 0
            break


def _rebuild_para_as_caption_field(para, prefix, chapter_num, seq_num, sep, title,
                                   numbering_format="chapter.seq"):
    """Rebuild caption numbering as field codes while preserving title runs."""
    el = para._element
    rpr_src = _extract_caption_rpr(el)
    first_text_run = _first_text_run(el)

    if first_text_run is None:
        _rebuild_para_full(el, prefix, chapter_num, seq_num, sep, title, rpr_src,
                           numbering_format=numbering_format)
        return

    _strip_caption_number_prefix(el, para.text or "")

    insert_pos = list(el).index(first_text_run)
    new_runs = []
    _append_caption_number_field_runs(
        new_runs, prefix, chapter_num, seq_num, sep,
        numbering_format=numbering_format, rpr_src=rpr_src
    )

    for i, run_el in enumerate(new_runs):
        el.insert(insert_pos + i, run_el)


def _rebuild_para_as_caption_text(para, prefix, chapter_num, seq_num, sep, title,
                                  numbering_format="chapter.seq"):
    """Rebuild caption numbering as plain text while preserving title runs."""
    el = para._element
    rpr_src = _extract_caption_rpr(el)
    first_text_run = _first_text_run(el)

    if first_text_run is None:
        _rebuild_para_text_full(el, prefix, chapter_num, seq_num, sep, title, rpr_src,
                                numbering_format=numbering_format)
        return

    _strip_caption_number_prefix(el, para.text or "")

    insert_pos = list(el).index(first_text_run)
    lead_text = f"{prefix}{_compose_caption_number_text(chapter_num, seq_num, numbering_format)}{sep}"
    el.insert(insert_pos, _make_run_element(text=lead_text, rpr_src=rpr_src))


def _rebuild_para_full(el, prefix, chapter_num, seq_num, sep, title, rpr_src,
                       numbering_format="chapter.seq"):
    """Fallback: rebuild caption paragraph as field-code numbering."""
    for r in el.findall(_w("r")):
        el.remove(r)
    _add_caption_number_fields(
        el, prefix, chapter_num, seq_num, sep,
        numbering_format=numbering_format, rpr_src=rpr_src
    )
    _make_run(el, text=title, rpr_src=rpr_src)


def _rebuild_para_text_full(el, prefix, chapter_num, seq_num, sep, title, rpr_src,
                            numbering_format="chapter.seq"):
    """Fallback: rebuild caption paragraph as plain-text numbering."""
    for r in el.findall(_w("r")):
        el.remove(r)
    number_text = _compose_caption_number_text(chapter_num, seq_num, numbering_format)
    _make_run(el, text=f"{prefix}{number_text}{sep}{title}", rpr_src=rpr_src)

def _format_caption_para(para, sc):
    """对题注段落应用格式（居中、字体、间距）"""
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = None
    pf.left_indent = None
    pf.space_before = Pt(sc.space_before_pt)
    pf.space_after = Pt(sc.space_after_pt)
    apply_line_spacing(pf, sc.line_spacing_type, sc.line_spacing_pt)
    for run in para.runs:
        run.font.name = sc.font_en
        run.font.size = Pt(sc.size_pt)
        rpr = run._element.find(_w("rPr"))
        if rpr is None:
            rpr = etree.SubElement(run._element, _w("rPr"))
        rfonts = rpr.find(_w("rFonts"))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, _w("rFonts"))
        rfonts.set(_w("eastAsia"), sc.font_cn)
        rfonts.set(_w("ascii"), sc.font_en)
        rfonts.set(_w("hAnsi"), sc.font_en)
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            rfonts.attrib.pop(_w(attr), None)


def _insert_caption_para(doc, anchor_elem, position,
                         prefix, chapter, seq, sep, title, sc,
                         use_field=True, numbering_format="chapter.seq"):
    """在锚点元素前/后插入题注段落。

    Args:
        use_field: True 用域代码（STYLEREF+SEQ），False 用纯文本
    """
    new_p = OxmlElement('w:p')
    ppr = etree.SubElement(new_p, _w("pPr"))
    jc = etree.SubElement(ppr, _w("jc"))
    jc.set(_w("val"), "center")

    if use_field:
        _add_caption_number_fields(
            new_p, prefix, chapter, seq, sep, numbering_format=numbering_format
        )
        _make_run(new_p, text=title)
    else:
        number_text = _compose_caption_number_text(chapter, seq, numbering_format)
        _make_run(new_p, text=f"{prefix}{number_text}{sep}{title}")

    if position == "after":
        anchor_elem.addnext(new_p)
    else:
        anchor_elem.addprevious(new_p)

    # 通过 doc.paragraphs 找到新段落并应用完整格式
    if sc:
        for para in doc.paragraphs:
            if para._element is new_p:
                _format_caption_para(para, sc)
                break


def _fix_indices_after_insert(doc_tree, body_end_old, count, headings):
    """题注插入后，修正 body 之后所有分区和标题的段落索引。"""
    for sec in doc_tree.sections:
        if sec.start_index > body_end_old:
            sec.start_index += count
        if sec.end_index > body_end_old:
            sec.end_index += count
        if sec.section_type == "body":
            sec.end_index += count
    for h in headings:
        if h.para_index > body_end_old:
            h.para_index += count


class CaptionFormatRule(BaseRule):
    name = "caption_format"
    description = "图表题注自动编号与格式修正"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        caps = getattr(config, "capabilities", {}) or {}
        if not bool(caps.get("caption", True)):
            return

        cap_cfg = config.caption
        if not cap_cfg.enabled:
            return

        doc_tree = context.get("doc_tree")
        if not doc_tree:
            return

        headings = context.get("headings", [])

        # 确定 body 范围
        body_section = None
        for sec in doc_tree.sections:
            if sec.section_type == "body":
                body_section = sec
                break
        if not body_section:
            return

        body_start = body_section.start_index
        body_end = body_section.end_index

        # 构建章号范围表
        chapter_ranges = _build_chapter_ranges(headings, body_end)

        # 构建 body 元素交错序列
        body_items = _build_body_map(doc, body_start, body_end)

        # 检测图片和表格
        figures = _detect_figures(doc, body_items)
        tables = _detect_tables(doc, body_items)

        # 获取样式配置
        fig_style = config.styles.get("figure_caption")
        tbl_style = config.styles.get("table_caption")

        # 按文档顺序编号并收集操作
        caption_indices = set()
        operations = []

        self._process_figures(
            doc, figures, chapter_ranges, cap_cfg,
            fig_style, tracker, caption_indices, operations,
        )
        self._process_tables(
            doc, tables, chapter_ranges, cap_cfg,
            tbl_style, tracker, caption_indices, operations,
        )

        # 从后往前执行插入操作（避免索引偏移）
        operations.sort(key=lambda op: op["sort_key"], reverse=True)
        for op in operations:
            _insert_caption_para(
                doc, op["anchor"], op["position"],
                op["prefix"], op["heading1"], op["seq"],
                op["sep"], op["title"], op["style"],
                use_field=op.get("use_field", True),
                numbering_format=op.get("numbering_format", "chapter.seq"),
            )

        # 插入题注后修正 doc_tree 后续分区索引
        if operations:
            _fix_indices_after_insert(doc_tree, body_end, len(operations),
                                     headings)

        # 存入 context 供 section_format 跳过
        context["caption_indices"] = caption_indices

    def _process_figures(self, doc, figures, chapter_ranges, cap_cfg,
                         fig_style, tracker, caption_indices, operations):
        """处理所有图片：用域代码重建题注编号、收集缺失题注插入操作"""
        use_chapter_number, _ = _parse_caption_numbering_format(cap_cfg.numbering_format)
        chapter_fig_counts = {}
        global_fig_count = 0

        for fig in figures:
            chap = _get_chapter_num(fig.para_index, chapter_ranges)
            if use_chapter_number:
                chapter_fig_counts[chap] = chapter_fig_counts.get(chap, 0) + 1
                seq = chapter_fig_counts[chap]
            else:
                global_fig_count += 1
                seq = global_fig_count
            prefix = cap_cfg.figure_prefix
            sep = cap_cfg.separator

            if fig.caption_para_index is not None:
                # 已有题注 → 提取标题文本，用域代码重建
                para = doc.paragraphs[fig.caption_para_index]
                old_text = para.text.strip()
                regex = _RE_FIG_CAPTION
                m = regex.match(old_text)
                title = m.group("title").strip() if m else old_text
                cap_block_indices = fig.caption_para_indices or [fig.caption_para_index]

                if cap_cfg.format_inserted:
                    _rebuild_para_as_caption_field(
                        para, prefix, chap, seq, sep, title,
                        numbering_format=cap_cfg.numbering_format)
                    after_text = f"{prefix}STYLEREF.SEQ{sep}{title[:30]}"
                else:
                    _rebuild_para_as_caption_text(
                        para, prefix, chap, seq, sep, title,
                        numbering_format=cap_cfg.numbering_format)
                    number_text = _compose_caption_number_text(
                        chap, seq, cap_cfg.numbering_format
                    )
                    after_text = f"{prefix}{number_text}{sep}{title[:30]}"
                if fig_style:
                    for cap_i in cap_block_indices:
                        if 0 <= cap_i < len(doc.paragraphs):
                            _format_caption_para(doc.paragraphs[cap_i], fig_style)
                for cap_i in cap_block_indices:
                    caption_indices.add(cap_i)

                tracker.record(
                    rule_name=self.name,
                    target=f"段落 #{fig.caption_para_index}"
                           if len(cap_block_indices) <= 1
                           else f"段落 #{fig.caption_para_index} (+{len(cap_block_indices)-1} 行续行)",
                    section="body", change_type="caption",
                    before=old_text[:60],
                    after=after_text,
                    paragraph_index=fig.caption_para_index,
                )
            else:
                # 缺失题注 → 收集插入操作
                if cap_cfg.auto_insert:
                    operations.append({
                        "sort_key": fig.para_index,
                        "anchor": fig.element,
                        "position": "after",
                        "prefix": prefix,
                        "heading1": chap,
                        "seq": seq,
                        "sep": sep,
                        "title": cap_cfg.placeholder,
                        "style": fig_style,
                        "use_field": cap_cfg.format_inserted,
                        "numbering_format": cap_cfg.numbering_format,
                    })
                    number_text = _compose_caption_number_text(
                        chap, seq, cap_cfg.numbering_format
                    )
                    tracker.record(
                        rule_name=self.name,
                        target=f"图片段落 #{fig.para_index} 后",
                        section="body", change_type="insert_caption",
                        before="(无题注)",
                        after=f"{prefix}{number_text}{sep}{cap_cfg.placeholder}",
                        paragraph_index=fig.para_index,
                    )

    def _process_tables(self, doc, tables, chapter_ranges, cap_cfg,
                        tbl_style, tracker, caption_indices, operations):
        """处理所有表格：编号、修正已有题注、收集缺失题注插入操作"""
        use_chapter_number, _ = _parse_caption_numbering_format(cap_cfg.numbering_format)
        chapter_tbl_counts = {}
        global_tbl_count = 0

        for tbl in tables:
            chap = _get_chapter_num(tbl.para_index, chapter_ranges)
            if use_chapter_number:
                chapter_tbl_counts[chap] = chapter_tbl_counts.get(chap, 0) + 1
                seq = chapter_tbl_counts[chap]
            else:
                global_tbl_count += 1
                seq = global_tbl_count
            prefix = cap_cfg.table_prefix
            sep = cap_cfg.separator

            if tbl.caption_para_index is not None:
                # 已有题注 → 提取标题文本，用域代码重建
                para = doc.paragraphs[tbl.caption_para_index]
                old_text = para.text.strip()
                regex = _RE_TBL_CAPTION
                m = regex.match(old_text)
                title = m.group("title").strip() if m else old_text

                if cap_cfg.format_inserted:
                    _rebuild_para_as_caption_field(
                        para, prefix, chap, seq, sep, title,
                        numbering_format=cap_cfg.numbering_format)
                    after_text = f"{prefix}STYLEREF.SEQ{sep}{title[:30]}"
                else:
                    _rebuild_para_as_caption_text(
                        para, prefix, chap, seq, sep, title,
                        numbering_format=cap_cfg.numbering_format)
                    number_text = _compose_caption_number_text(
                        chap, seq, cap_cfg.numbering_format
                    )
                    after_text = f"{prefix}{number_text}{sep}{title[:30]}"
                if tbl_style:
                    _format_caption_para(para, tbl_style)
                caption_indices.add(tbl.caption_para_index)

                tracker.record(
                    rule_name=self.name,
                    target=f"段落 #{tbl.caption_para_index}",
                    section="body", change_type="caption",
                    before=old_text[:60],
                    after=after_text,
                    paragraph_index=tbl.caption_para_index,
                )
            else:
                # 缺失题注 → 收集插入操作
                if cap_cfg.auto_insert:
                    operations.append({
                        "sort_key": tbl.para_index,
                        "anchor": tbl.element,
                        "position": "before",
                        "prefix": prefix,
                        "heading1": chap,
                        "seq": seq,
                        "sep": sep,
                        "title": cap_cfg.placeholder,
                        "style": tbl_style,
                        "use_field": cap_cfg.format_inserted,
                        "numbering_format": cap_cfg.numbering_format,
                    })
                    number_text = _compose_caption_number_text(
                        chap, seq, cap_cfg.numbering_format
                    )
                    tracker.record(
                        rule_name=self.name,
                        target=f"表格 #{tbl.table_index} 前",
                        section="body", change_type="insert_caption",
                        before="(无题注)",
                        after=f"{prefix}{number_text}{sep}{cap_cfg.placeholder}",
                        paragraph_index=tbl.para_index,
                    )

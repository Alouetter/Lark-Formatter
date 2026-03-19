"""Formula-to-table conversion rule."""

from __future__ import annotations

from copy import deepcopy

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.engine.change_tracker import ChangeTracker
from src.engine.rules import table_format as table_helpers
from src.engine.rules.base import BaseRule
from src.formula_core.parse import parse_document_formulas
from src.formula_core.runtime import FormulaRuleStats
from src.scene.schema import SceneConfig

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TABLE_ALIGNMENT_MAP = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}


def _is_body_paragraph(paragraph) -> bool:
    p = paragraph._p
    if p is None or p.getparent() is None:
        return False
    parent_tag = p.getparent().tag
    local = parent_tag.split("}")[-1] if "}" in parent_tag else parent_tag
    return local == "body"


def _copy_paragraph_payload(src_paragraph, dst_paragraph) -> None:
    src_p = src_paragraph._p
    dst_p = dst_paragraph._p
    ppr_tag = f"{{{_W_NS}}}pPr"
    for child in list(dst_p):
        if child.tag != ppr_tag:
            dst_p.remove(child)
    for child in list(src_p):
        if child.tag == ppr_tag:
            continue
        dst_p.append(deepcopy(child))


class FormulaToTableRule(BaseRule):
    name = "formula_to_table"
    description = "Formula to table conversion"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        cfg = getattr(config, "formula_to_table", None)
        if cfg is not None and not bool(getattr(cfg, "enabled", False)):
            return

        block_only = bool(getattr(cfg, "block_only", True))
        visual_cfg = getattr(config, "formula_table", None)
        table_alignment = str(
            getattr(visual_cfg, "table_alignment", "center")
        ).strip().lower()
        table_alignment = table_alignment if table_alignment in _TABLE_ALIGNMENT_MAP else "center"
        auto_shrink_number_column = bool(
            getattr(visual_cfg, "auto_shrink_number_column", True)
        )
        avail_w = table_helpers._available_width(config)
        runtime = context.setdefault("formula_runtime", {})
        convert_enabled = bool(runtime.get("convert_enabled"))
        runtime["to_table_enabled"] = True
        parse_result = parse_document_formulas(doc)
        stats = FormulaRuleStats()

        candidates_by_para: dict[int, object] = {}
        mixed_skip_logged: set[int] = set()
        for occ in parse_result.occurrences:
            if convert_enabled and occ.source_type != "word_native":
                continue
            if block_only and not occ.is_block:
                continue
            stats.matched += 1
            stats.note_confidence(float(getattr(occ.node, "confidence", 0.0)))
            para_key = id(occ.paragraph)
            if not occ.is_formula_only:
                if para_key not in mixed_skip_logged:
                    mixed_skip_logged.add(para_key)
                    stats.skipped_dependency += 1
                    tracker.record(
                        rule_name=self.name,
                        target=occ.location,
                        section="formula",
                        change_type="skip",
                        before="formula paragraph",
                        after="MVP 仅处理纯公式段落",
                        paragraph_index=occ.paragraph_index,
                    )
                continue
            if para_key not in candidates_by_para:
                candidates_by_para[para_key] = occ

        candidates = list(candidates_by_para.values())
        candidates.sort(
            key=lambda item: item.paragraph_index if item.paragraph_index >= 0 else 10**9,
            reverse=True,
        )

        for occ in candidates:
            para = occ.paragraph
            if occ.in_table:
                stats.skipped_dependency += 1
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="skip",
                    before="formula paragraph",
                    after="暂不转换表格内公式",
                    paragraph_index=occ.paragraph_index,
                )
                continue
            if not _is_body_paragraph(para):
                stats.skipped_dependency += 1
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="skip",
                    before="formula paragraph",
                    after="暂不转换非正文段落",
                    paragraph_index=occ.paragraph_index,
                )
                continue

            table = doc.add_table(rows=1, cols=2)
            table.alignment = _TABLE_ALIGNMENT_MAP.get(
                table_alignment,
                WD_TABLE_ALIGNMENT.CENTER,
            )
            left_para = table.cell(0, 0).paragraphs[0]
            right_para = table.cell(0, 1).paragraphs[0]
            # Structure-only rule: keep original formula payload as-is and only containerize.
            _copy_paragraph_payload(para, left_para)
            right_para.text = ""
            if auto_shrink_number_column and avail_w and avail_w > 0:
                try:
                    table_helpers._shrink_equation_number_column(table._tbl, avail_w)
                except Exception:
                    # Keep conversion robust; width adjustment failure should not block conversion.
                    pass

            try:
                p = para._p
                parent = p.getparent()
                insert_pos = parent.index(p) + 1
                tbl_el = table._tbl
                # Move generated table from document tail to target position.
                if tbl_el.getparent() is not None:
                    tbl_el.getparent().remove(tbl_el)
                parent.insert(insert_pos, tbl_el)
                parent.remove(p)
            except Exception as exc:
                stats.errors += 1
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="text",
                    before="formula paragraph",
                    after="公式表格转换失败",
                    paragraph_index=occ.paragraph_index,
                    success=False,
                    failure_reason=str(exc),
                )
                continue

            stats.converted += 1
            tracker.record(
                rule_name=self.name,
                target=occ.location,
                section="formula",
                change_type="text",
                before="公式段落",
                after="已转换为公式表格（两列）",
                paragraph_index=occ.paragraph_index,
            )

        runtime.setdefault("stats", {})[self.name] = {
            "matched": stats.matched,
            "converted": stats.converted,
            "skipped_low_confidence": stats.skipped_low_confidence,
            "skipped_unsupported": stats.skipped_unsupported,
            "skipped_dependency": stats.skipped_dependency,
            "errors": stats.errors,
        }
        tracker.record(
            rule_name=self.name,
            target="summary",
            section="formula",
            change_type="format",
            before="统计",
            after=stats.to_summary(),
            paragraph_index=-1,
        )
        tracker.record(
            rule_name=self.name,
            target="confidence_confirmation",
            section="formula",
            change_type="format",
            before="自动置信分层确认",
            after=f"已在执行日志末尾标注：{stats.confidence_summary()}。",
            paragraph_index=-1,
        )

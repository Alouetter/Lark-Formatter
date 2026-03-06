"""Section-format rule: apply style config by detected document sections."""

from __future__ import annotations

import re
import unicodedata
from copy import deepcopy

from lxml import etree
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt

from src.engine.change_tracker import ChangeTracker
from src.engine.doc_tree import DocTree
from src.engine.rules.base import BaseRule
from src.scene.heading_model import (
    detect_level_by_style_name,
    get_front_matter_title_norms,
    get_non_numbered_heading_style_name,
    get_non_numbered_title_norms,
    get_non_numbered_title_sections,
    get_section_title_style_map,
)
from src.scene.schema import FormatScopeConfig, SceneConfig
from src.utils.line_spacing import apply_line_spacing

# section_type -> style key
SECTION_STYLE_MAP = {
    "body": "normal",
    "references": "references_body",
    "errata": "references_body",
    "acknowledgment": "acknowledgment_body",
    "appendix": "appendix_body",
    "abstract_cn": "abstract_body",
    "abstract_en": "abstract_body_en",
    "resume": "resume_body",
}

# section title paragraph -> style key
SECTION_TITLE_STYLE_MAP = {
    "abstract_cn": "abstract_title_cn",
    "abstract_en": "abstract_title_en",
    "references": "heading1",
    "errata": "heading1",
    "appendix": "heading1",
    "acknowledgment": "heading1",
    "resume": "heading1",
}

_NON_NUMBERED_TITLE_SECTIONS = {"references", "errata", "appendix", "acknowledgment", "resume"}
_NON_NUMBERED_TITLE_TEXTS = {
    "参考文献",
    "勘误页",
    "勘误",
    "附录",
    "致谢",
    "个人简历",
    "在学期间发表的学术论文与研究成果",
}
_NON_NUMBERED_HEADING_STYLE_NAME = "Heading 1 Unnumbered"

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

_RE_FIGURE_CAPTION = re.compile(r"^(图|表|Figure|Table|Fig\.?)\s*\d", re.IGNORECASE)
_HEADING_STYLE_PREFIXES = ("Heading", "heading", "标题")
_RE_TOC_STYLE = re.compile(r"^(toc|目录)\s*\d+", re.IGNORECASE)

# Body text subfigure references: "3.11　(a)" -> "3.11 (a)"
_RE_SUBFIG_REF_FULLWIDTH_SPACE = re.compile(
    r"(?P<num>\d+(?:\.\d+)+)\u3000(?P<label>[（(][A-Za-z][)）])"
)

_RE_REFERENCE_ENTRY_LINE = re.compile(r"^\s*(\[\d{1,4}\]|[（(]\d{1,4}[)）]|\d{1,4}\.)\s+\S")
_RE_TITLE_TAIL_MARKS = re.compile(r"[：:;；·•\-—_~\.。…]+$")
_RE_ABSTRACT_CN_TITLE = re.compile(r"^\u6458\u8981(?:[（(][^()（）]{0,8}[)）])?$")
_RE_ABSTRACT_EN_TITLE = re.compile(
    r"^abstract(?:[（(][^()（）]{0,16}[)）])?$",
    re.IGNORECASE,
)
_RE_TOC_TITLE_LINE = re.compile(
    r"^(?:\u76ee\u5f55|\u76ee\u9304|contents|tableofcontents)(?:[（(][^()（）]{0,8}[)）])?$",
    re.IGNORECASE,
)
_CHEM_TOKEN_RE = re.compile(
    r"[A-Za-z0-9Ａ-Ｚａ-ｚ０-９µμα-ωΑ-ΩΩω\(\)\[\]\{\}（）［］｛｝"
    r"\+\-−＋－‐‑‒–—―=\^\.·•∙⋅≡/%％／<>→←↔"
    r"⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₙ]+"
)
_CHARGE_SIGNS = "+-−＋－"
_DOT_SEPARATORS = ".·•"
_RADICAL_DOT_CHARS = {"·", "•"}
_BOND_CHARS = "-−－‐‑‒–—―=≡"
_SUPER_SUB_UNICODE_CHARS = set(
    "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ"
    "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₙ"
)
_SUPER_SUB_ASCII_MAP = {
    "⁰": "0",
    "¹": "1",
    "²": "2",
    "³": "3",
    "⁴": "4",
    "⁵": "5",
    "⁶": "6",
    "⁷": "7",
    "⁸": "8",
    "⁹": "9",
    "⁺": "+",
    "⁻": "-",
    "⁼": "=",
    "⁽": "(",
    "⁾": ")",
    "₀": "0",
    "₁": "1",
    "₂": "2",
    "₃": "3",
    "₄": "4",
    "₅": "5",
    "₆": "6",
    "₇": "7",
    "₈": "8",
    "₉": "9",
    "₊": "+",
    "₋": "-",
    "₌": "=",
    "₍": "(",
    "₎": ")",
}
_INTERNAL_CHEM_CONFUSABLE_CHAR_MAP = {
    # Small symbol variants seen in copied equations/units.
    "﹣": "-",
    "﹢": "+",
    # Common non-ASCII unit letters.
    "ɡ": "g",
    "ℊ": "g",
    "ℎ": "h",
    "ℓ": "l",
    "K": "K",
}
_BRACKET_OPEN_TO_CLOSE = {"(": ")", "[": "]", "{": "}"}
_RE_MASS_ADDUCT_CHARGE = re.compile(r"^\[[A-Za-z0-9+\-−＋－]+\](\d*[+\-−＋－])$")
_ELEMENT_SYMBOLS = set(
    (
        "H He Li Be B C N O F Ne Na Mg Al Si P S Cl Ar K Ca Sc Ti V Cr Mn Fe Co Ni Cu Zn "
        "Ga Ge As Se Br Kr Rb Sr Y Zr Nb Mo Tc Ru Rh Pd Ag Cd In Sn Sb Te I Xe Cs Ba La Ce "
        "Pr Nd Pm Sm Eu Gd Tb Dy Ho Er Tm Yb Lu Hf Ta W Re Os Ir Pt Au Hg Tl Pb Bi Po At Rn "
        "Fr Ra Ac Th Pa U Np Pu Am Cm Bk Cf Es Fm Md No Lr Rf Db Sg Bh Hs Mt Ds Rg Cn Nh Fl "
        "Mc Lv Ts Og"
    ).split()
)
# Common isotope symbols in chemistry writing.
_ELEMENT_SYMBOLS.update({"D", "T"})

# Internal hardening dictionaries for extreme/ambiguous cases.
# These are intentionally conservative and not exposed in UI.
_INTERNAL_CHEM_IGNORE_TOKENS = {
    "H1N1",
    "H5N1",
    "H3N2",
}
_INTERNAL_CHEM_ALLOW_TOKENS: set[str] = set()
_INTERNAL_CHEM_IGNORE_PATTERNS = [
    re.compile(r"^H\d{1,2}N\d{1,2}$"),  # Influenza subtype naming.
    # Instrument/model codes.
    re.compile(r"^CHI\d{2,}[A-Z]*$"),
    re.compile(r"^LCMS\d{2,}[A-Z]*$"),
]
_INTERNAL_CHEM_ALLOW_PATTERNS: list[re.Pattern[str]] = []
_INTERNAL_CHEM_MANUAL_OVERRIDES = {
    # Singlet oxygen: leading isotope/singlet marker + oxygen count.
    "1O2": "^._",
    # Radical forms often represented with a raised dot and sign.
    "O2·-": "._.^",
    "O2•-": "._.^",
    # Sulfate radical anion; keep leading dot baseline.
    "·SO42-": "..._^^",
    "•SO42-": "..._^^",
    "SO42-": ".._^^",
    "SO42+": ".._^^",
    "SO4^2-": ".._.^^",
    "SO4^2+": ".._.^^",
    "SO4-2": ".._^^",
    "SO4+2": ".._^^",
}
_INSTRUMENT_MODEL_PREFIXES = {
    "CHI",
    "LCMS",
    "HPLC",
    "UPLC",
    "GCMS",
    "ICP",
    "XRD",
    "XPS",
    "FTIR",
    "NMR",
    "SEM",
    "TEM",
}
_RE_BOND_TOKEN = re.compile(
    r"^(?:[A-Z][a-z]?)(?:[\-−－‐‑‒–—―=≡](?:[A-Z][a-z]?))+$"
)
_RE_ROMAN_OX_STATE = re.compile(
    r"^(?P<elem>[A-Z][a-z]?)\((?P<roman>I|II|III|IV|V|VI|VII)\)$"
)
_RE_UNIT_RATIO_TOKEN = re.compile(
    r"^(?:[A-Za-zµμΩω%]{1,10}(?:\([A-Za-zµμΩω%·\.]+\))?)(?:/(?:[A-Za-zµμΩω%]{1,10}(?:\([A-Za-zµμΩω%·\.]+\))?))+?$"
)
_RE_ISOTOPE_LABEL_TOKEN = re.compile(
    r"^\[(?P<mass>\d{1,3})(?P<elem>[A-Z][a-z]?)\]\d*(?:-[A-Za-z0-9]+)*$"
)
_RE_REACTION_ARROW_TOKEN = re.compile(
    r"^(?P<left>.+?)(?:<->|->|<-|→|←|↔)(?P<right>.+)$"
)
_RE_SPACED_FORMULA_SPAN = re.compile(
    r"[A-Za-z0-9Ａ-Ｚａ-ｚ０-９µμα-ωΑ-ΩΩω\(\)\[\]\{\}（）［］｛｝"
    r"\+\-−＋－‐‑‒–—―=\^\.·•∙⋅≡/%％／<>→←↔"
    r"⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₙ\s]{4,}"
)
_RE_GREEK_PREFIX_TOKEN = re.compile(r"^(?P<prefix>[α-ωΑ-Ω])(?P<sep>[\-−－‐‑‒–—―])(?P<rest>.+)$")
_RE_PLAIN_FORMULA_CHAIN = re.compile(r"^(?:\d{0,3})?(?:[A-Z][a-z]?\d*){2,}$")
_RE_SINGLE_ELEMENT_TOKEN = re.compile(r"^[A-Z][a-z]?$")
_RE_HYBRIDIZATION_TOKEN = re.compile(
    r"^(?:sp\d(?:d\d{0,2})?|dsp\d|d\d{1,2}sp\d(?:d\d{0,2})?)$",
    re.IGNORECASE,
)
_RE_BRACKET_COMPLEX_WITH_CHARGE = re.compile(
    r"^\[[A-Za-z0-9\(\)\+\-−＋－·•]+\]\d*[+\-−＋－]$"
)
_RE_SIMPLE_UNIT_CONNECT_TOKEN = re.compile(
    r"^[A-Za-zµμΩω%]{1,10}(?:[·•∙⋅/][A-Za-zµμΩω%]{1,10})+$"
)
_RE_PERCENT_SHORT_UNIT = re.compile(r"^(?:wt|vol|at)%$", re.IGNORECASE)
_RE_PERCENT_BRACKET_UNIT = re.compile(r"^%\((?:w|v)/(?:w|v)\)$", re.IGNORECASE)
_PLAIN_UNIT_WORDS = {"ppm", "ppb", "ppt", "ppmv", "ppbv", "pptr"}

ChemRuntime = tuple[
    set[str],  # ignore_tokens
    list[re.Pattern[str]],  # ignore_patterns
    dict[str, str],  # manual_overrides
    set[str],  # allow_tokens
    list[re.Pattern[str]],  # allow_patterns
]


def _run_has_no_proof(run) -> bool:
    """Runs carrying w:noProof should keep original formatting."""
    rpr = run._element.find(f"{{{_W_NS}}}rPr")
    return rpr is not None and rpr.find(f"{{{_W_NS}}}noProof") is not None


def _para_has_num_pr(para) -> bool:
    """List paragraphs should keep list indentation/formatting."""
    ppr = para._element.find(f"{{{_W_NS}}}pPr")
    return ppr is not None and ppr.find(f"{{{_W_NS}}}numPr") is not None


def _clear_para_num_pr(para) -> None:
    """Remove paragraph numbering properties to avoid list indentation carry-over."""
    ppr = para._element.find(f"{{{_W_NS}}}pPr")
    if ppr is None:
        return
    numpr = ppr.find(f"{{{_W_NS}}}numPr")
    if numpr is not None:
        ppr.remove(numpr)


def _para_is_special_block(para) -> bool:
    """Paragraphs tagged by md_cleanup as code/quote blocks."""
    ppr = para._element.find(f"{{{_W_NS}}}pPr")
    if ppr is not None and ppr.find(f"{{{_W_NS}}}pBdr") is not None:
        return True
    for run in para.runs:
        if _run_has_no_proof(run):
            return True
    return False


def _para_has_toc_pageref(para) -> bool:
    has_pageref = False
    for instr in para._element.iter(f"{{{_W_NS}}}instrText"):
        if "PAGEREF" in (instr.text or "").upper():
            has_pageref = True
            break
    if not has_pageref:
        return False
    return para._element.find(f".//{{{_W_NS}}}tab") is not None


def _para_is_heading(para, config: SceneConfig | None = None) -> bool:
    style_name = para.style.name if para.style else ""
    style_id = getattr(para.style, "style_id", "") if para.style else ""
    if config is not None:
        if detect_level_by_style_name(config, style_name):
            return True
        if detect_level_by_style_name(config, style_id):
            return True
    for prefix in _HEADING_STYLE_PREFIXES:
        if style_name.startswith(prefix):
            return True
    return False


def _para_is_caption(para) -> bool:
    text = (para.text or "").strip()
    return bool(text and _RE_FIGURE_CAPTION.match(text))


def _norm_no_space(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip()


def _norm_title_token(text: str) -> str:
    norm = _norm_no_space(text)
    if not norm:
        return ""
    return _RE_TITLE_TAIL_MARKS.sub("", norm)


def _contains_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in (text or ""))


def _para_is_resume_title(para) -> bool:
    norm = _norm_no_space(para.text)
    if not norm:
        return False
    if norm in ("个人简历", "在学期间发表的学术论文与研究成果"):
        return True
    return norm.startswith("在学期间发表的学术论文与研究成果")


def _para_is_appendix_title(para) -> bool:
    norm = _norm_no_space(para.text)
    if not norm:
        return False
    if len(norm) > 80:
        return False
    return norm.startswith("附录") or norm.lower().startswith("appendix")


def _is_non_numbered_post_title_text(
    text: str,
    non_numbered_title_norms: set[str] | None = None,
) -> bool:
    norm = _norm_no_space(text)
    if not norm:
        return False
    title_norms = non_numbered_title_norms
    if title_norms is None:
        title_norms = {_norm_no_space(v) for v in _NON_NUMBERED_TITLE_TEXTS if _norm_no_space(v)}
    if norm in title_norms:
        return True
    if norm.startswith("附录") or norm.lower().startswith("appendix"):
        return True
    return norm.startswith("在学期间发表的学术论文与研究成果")


def _looks_like_reference_entry(text: str) -> bool:
    return bool(_RE_REFERENCE_ENTRY_LINE.match((text or "").strip()))


def _looks_like_section_title_line(
    sec_type: str,
    text: str,
    *,
    front_matter_title_norms: set[str] | None = None,
    non_numbered_title_sections: set[str] | None = None,
    non_numbered_title_norms: set[str] | None = None,
) -> bool:
    norm = _norm_title_token(text)
    if not norm:
        return False
    low = norm.lower()

    fm_norms = front_matter_title_norms
    if fm_norms is None:
        fm_norms = {
            _norm_title_token("\u6458\u8981").lower(),
            _norm_title_token("abstract").lower(),
            _norm_title_token("\u76ee\u5f55").lower(),
            _norm_title_token("contents").lower(),
            _norm_title_token("tableofcontents").lower(),
        }
    else:
        fm_norms = {
            _norm_title_token(v).lower()
            for v in fm_norms
            if _norm_title_token(v)
        }

    if sec_type == "abstract_cn":
        if _RE_ABSTRACT_CN_TITLE.match(norm):
            return True
        return low in fm_norms and _contains_cjk(norm)

    if sec_type == "abstract_en":
        if _RE_ABSTRACT_EN_TITLE.match(norm):
            return True
        return low in fm_norms and "abstract" in low

    if sec_type == "toc":
        if _RE_TOC_TITLE_LINE.match(norm):
            return True
        return low in fm_norms or low in {"contents", "tableofcontents"}

    nn_sections = non_numbered_title_sections
    if nn_sections is None:
        nn_sections = set(_NON_NUMBERED_TITLE_SECTIONS)
    if sec_type in nn_sections:
        return _is_non_numbered_post_title_text(text, non_numbered_title_norms)

    return False


def _should_treat_first_para_as_section_title(
    section,
    para,
    *,
    config: SceneConfig | None = None,
    front_matter_title_norms: set[str] | None = None,
    non_numbered_title_sections: set[str] | None = None,
    non_numbered_title_norms: set[str] | None = None,
) -> bool:
    """First-para title guard: avoid aggressive title treatment on weak sections."""
    text = (para.text or "").strip()
    if not text:
        return False

    if section.section_type == "references" and _looks_like_reference_entry(text):
        return False

    if _looks_like_section_title_line(
        section.section_type,
        text,
        front_matter_title_norms=front_matter_title_norms,
        non_numbered_title_sections=non_numbered_title_sections,
        non_numbered_title_norms=non_numbered_title_norms,
    ):
        return True

    if getattr(section, "title_confident", True):
        return _para_is_heading(para, config=config) and len(text) <= 40
    return False


def _lstrip_para_leading_space(para) -> None:
    if not para.runs:
        if para.text:
            para.text = para.text.lstrip(" \t\u3000")
        return
    for run in para.runs:
        txt = run.text or ""
        if not txt:
            continue
        run.text = txt.lstrip(" \t\u3000")
        if run.text:
            break


def _normalize_subfigure_ref_space(para) -> bool:
    """Normalize full-width space before sub-figure marker in body text.

    Example:
    - 3.11　(a) -> 3.11 (a)
    """
    if not para.runs:
        text = para.text or ""
        new_text = _RE_SUBFIG_REF_FULLWIDTH_SPACE.sub(r"\g<num> \g<label>", text)
        if new_text != text:
            para.text = new_text
            return True
        return False

    spans = []
    full_text_parts = []
    cursor = 0
    for run in para.runs:
        text = run.text or ""
        start = cursor
        end = start + len(text)
        spans.append((run, start, end))
        full_text_parts.append(text)
        cursor = end

    full_text = "".join(full_text_parts)
    if not full_text:
        return False

    target_offsets = []
    for m in _RE_SUBFIG_REF_FULLWIDTH_SPACE.finditer(full_text):
        fw_offset = m.end("num")
        if fw_offset < len(full_text) and full_text[fw_offset] == "\u3000":
            target_offsets.append(fw_offset)

    if not target_offsets:
        return False

    changed = False
    for off in target_offsets:
        for run, start, end in spans:
            if start <= off < end:
                local = off - start
                txt = run.text or ""
                if local < len(txt) and txt[local] == "\u3000":
                    run.text = f"{txt[:local]} {txt[local + 1:]}"
                    changed = True
                break
    return changed


def _para_is_equation(para) -> bool:
    el = para._element
    if el.findall(f"{{{_M_NS}}}oMathPara") or el.findall(f"{{{_M_NS}}}oMath"):
        return True
    if el.findall(f"{{{_W_NS}}}r/{{{_W_NS}}}object"):
        return True
    return False


def _para_has_image(para) -> bool:
    el = para._element
    if el.findall(f"{{{_W_NS}}}r/{{{_W_NS}}}pict"):
        return True
    if el.findall(f"{{{_W_NS}}}r/{{{_W_NS}}}drawing"):
        return True
    if el.findall(f".//{{{_WP_NS}}}inline") or el.findall(f".//{{{_WP_NS}}}anchor"):
        return True
    return False


def _center_paragraph(para):
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = None
    pf.left_indent = None


def _normalize_reference_entry_paragraph(para) -> bool:
    """Force reference entry layout: left aligned, no first-line/list indentation."""
    text = (para.text or "").strip()
    if not text or not _looks_like_reference_entry(text):
        return False

    _clear_para_num_pr(para)
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    _sanitize_para_indent_ooxml(para, force_zero=True)
    return True


def _sanitize_para_indent_ooxml(para, *, force_zero: bool = False) -> None:
    """Clear legacy char-based indents; optionally force paragraph indent to zero."""
    ppr = para._element.find(f"{{{_W_NS}}}pPr")
    if ppr is None:
        if not force_zero:
            return
        ppr = etree.SubElement(para._element, f"{{{_W_NS}}}pPr")
    ind = ppr.find(f"{{{_W_NS}}}ind")
    if ind is None:
        if not force_zero:
            return
        ind = etree.SubElement(ppr, f"{{{_W_NS}}}ind")

    if force_zero:
        ind.set(f"{{{_W_NS}}}left", "0")
        ind.set(f"{{{_W_NS}}}right", "0")
        ind.set(f"{{{_W_NS}}}firstLine", "0")
        ind.set(f"{{{_W_NS}}}hanging", "0")

    # Character-based indent attrs can survive style rewriting and visually keep indentation.
    for char_attr in ("leftChars", "rightChars", "firstLineChars", "hangingChars"):
        ind.attrib.pop(f"{{{_W_NS}}}{char_attr}", None)


def _center_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _para_has_sectpr(para) -> bool:
    ppr = para._element.find(f"{{{_W_NS}}}pPr")
    return ppr is not None and ppr.find(f"{{{_W_NS}}}sectPr") is not None


def _has_table_between_paragraphs(doc, start_idx: int, end_idx: int) -> bool:
    """Return True when any body-level table exists between two paragraphs."""
    if start_idx < 0 or end_idx <= start_idx:
        return False
    if start_idx >= len(doc.paragraphs) or end_idx >= len(doc.paragraphs):
        return False

    start_el = doc.paragraphs[start_idx]._element
    end_el = doc.paragraphs[end_idx]._element
    curr = start_el.getnext()
    while curr is not None and curr is not end_el:
        tag = curr.tag.split("}")[-1] if "}" in curr.tag else curr.tag
        if tag == "tbl":
            return True
        curr = curr.getnext()
    return False


def _has_nonempty_content(doc, start_idx: int, end_idx: int) -> bool:
    if start_idx > end_idx:
        return False
    end = min(end_idx, len(doc.paragraphs) - 1)
    for i in range(max(0, start_idx), end + 1):
        if (doc.paragraphs[i].text or "").strip():
            return True
    return False


def _break_applies_to_boundary(doc, break_idx: int, boundary_idx: int) -> bool:
    if break_idx < 0 or break_idx >= boundary_idx:
        return False
    if _has_table_between_paragraphs(doc, break_idx, boundary_idx):
        return False
    for i in range(break_idx + 1, boundary_idx):
        if i >= len(doc.paragraphs):
            break
        if (doc.paragraphs[i].text or "").strip():
            return False
    return True


def _remove_caption_table_section_breaks(doc) -> int:
    """Drop sectPr on caption paragraphs directly followed by a table."""
    removed = 0
    for para in doc.paragraphs:
        if not _para_has_sectpr(para):
            continue
        if not _para_is_caption(para):
            continue
        nxt = para._element.getnext()
        if nxt is None:
            continue
        tag = nxt.tag.split("}")[-1] if "}" in nxt.tag else nxt.tag
        if tag != "tbl":
            continue
        ppr = para._element.find(f"{{{_W_NS}}}pPr")
        if ppr is None:
            continue
        sect_pr = ppr.find(f"{{{_W_NS}}}sectPr")
        if sect_pr is None:
            continue
        ppr.remove(sect_pr)
        removed += 1
    return removed


def _apply_run_font(run, font_cn: str, font_en: str, size_pt: float):
    if _run_has_no_proof(run):
        return
    run.font.name = font_en
    run.font.size = Pt(size_pt)
    rpr = run._element.find(f"{{{_W_NS}}}rPr")
    if rpr is None:
        rpr = etree.SubElement(run._element, f"{{{_W_NS}}}rPr")
    rfonts = rpr.find(f"{{{_W_NS}}}rFonts")
    if rfonts is None:
        rfonts = etree.SubElement(rpr, f"{{{_W_NS}}}rFonts")
    rfonts.set(f"{{{_W_NS}}}eastAsia", font_cn)
    rfonts.set(f"{{{_W_NS}}}ascii", font_en)
    rfonts.set(f"{{{_W_NS}}}hAnsi", font_en)
    rfonts.set(f"{{{_W_NS}}}cs", font_en)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme", "cstheme"):
        rfonts.attrib.pop(f"{{{_W_NS}}}{attr}", None)


def _run_vert_align(run) -> str | None:
    """Return run vertical alignment value if present."""
    if run.font.superscript:
        return "superscript"
    if run.font.subscript:
        return "subscript"
    rpr = run._element.find(f"{{{_W_NS}}}rPr")
    if rpr is None:
        return None
    va = rpr.find(f"{{{_W_NS}}}vertAlign")
    if va is None:
        return None
    val = (va.get(f"{{{_W_NS}}}val") or "").strip()
    if val in {"superscript", "subscript"}:
        return val
    return None


def _para_has_existing_vert_align(para) -> bool:
    """Skip paragraphs that already contain explicit superscript/subscript."""
    for run in para.runs:
        if _run_vert_align(run) in {"superscript", "subscript"}:
            return True
    return False


def _mark_style(marks: list[str | None], start: int, end: int, style: str) -> None:
    for i in range(max(0, start), min(len(marks), end)):
        marks[i] = style


def _is_ascii_upper(ch: str) -> bool:
    return "A" <= ch <= "Z"


def _is_ascii_lower(ch: str) -> bool:
    return "a" <= ch <= "z"


def _is_cjk(ch: str) -> bool:
    return "\u4e00" <= ch <= "\u9fff"


def _is_word_or_cjk(ch: str) -> bool:
    return ch.isalnum() or _is_cjk(ch)


def _normalize_math_alnum_char(ch: str) -> str:
    """Map mathematical alphanumeric symbols to plain ASCII when possible."""
    if not ch:
        return ""
    name = unicodedata.name(ch, "")
    if not name.startswith("MATHEMATICAL "):
        return ""

    m_alpha = re.search(r"(CAPITAL|SMALL) ([A-Z])$", name)
    if m_alpha:
        letter = m_alpha.group(2)
        return letter if m_alpha.group(1) == "CAPITAL" else letter.lower()

    m_digit = re.search(r"DIGIT ([0-9])$", name)
    if m_digit:
        return m_digit.group(1)

    if name.endswith("SMALL DOTLESS I"):
        return "i"
    if name.endswith("SMALL DOTLESS J"):
        return "j"
    return ""


def _normalize_formula_scan_char(ch: str) -> str:
    """Normalize single char for candidate scanning while keeping length 1:1."""
    if not ch:
        return ""
    # Fullwidth ASCII to halfwidth ASCII.
    if "\uff01" <= ch <= "\uff5e":
        ch = chr(ord(ch) - 0xFEE0)

    if ch in _INTERNAL_CHEM_CONFUSABLE_CHAR_MAP:
        return _INTERNAL_CHEM_CONFUSABLE_CHAR_MAP[ch]

    mapped = _normalize_math_alnum_char(ch)
    if mapped:
        return mapped
    return ch


def _normalize_formula_scan_text(text: str) -> str:
    if not text:
        return ""
    return "".join(_normalize_formula_scan_char(ch) for ch in str(text))


def _first_non_space_char(text: str) -> str:
    for ch in text or "":
        if not ch.isspace():
            return ch
    return ""


def _normalize_formula_token_chars(token: str) -> str:
    if not token:
        return ""
    out = []
    for ch in str(token):
        ch = _normalize_formula_scan_char(ch)

        if ch in {"µ"}:
            out.append("μ")
        elif ch in {"•"}:
            out.append("·")
        elif ch in {"∙", "⋅"}:
            out.append("·")
        elif ch in {"＋"}:
            out.append("+")
        elif ch in {"−", "－", "‐", "‑", "‒", "–", "—", "―"}:
            out.append("-")
        elif ch in {"（"}:
            out.append("(")
        elif ch in {"）"}:
            out.append(")")
        elif ch in {"［"}:
            out.append("[")
        elif ch in {"］"}:
            out.append("]")
        elif ch in {"｛"}:
            out.append("{")
        elif ch in {"｝"}:
            out.append("}")
        elif ch in {"／"}:
            out.append("/")
        elif ch in {"％"}:
            out.append("%")
        elif ch in _SUPER_SUB_ASCII_MAP:
            out.append(_SUPER_SUB_ASCII_MAP[ch])
        else:
            out.append(ch)
    return "".join(out)


def _normalize_unit_placeholder_token(token: str) -> str:
    """Normalize ad-hoc placeholders that may appear in imported corpora."""
    t = str(token or "")
    t = t.replace("{DOT}", "·").replace("{dot}", "·")
    return t


def _looks_like_roman_oxidation_token(token: str) -> bool:
    t = _normalize_formula_token_chars(token or "")
    m = _RE_ROMAN_OX_STATE.fullmatch(t)
    if not m:
        return False
    return m.group("elem") in _ELEMENT_SYMBOLS


def _looks_like_unit_ratio_token(token: str) -> bool:
    t = _normalize_unit_placeholder_token(_normalize_formula_token_chars(token or ""))
    if not t:
        return False
    if _RE_UNIT_RATIO_TOKEN.fullmatch(t):
        return True
    # Relaxed fallback: unit fragments connected by '/', no spaces.
    if "/" in t and " " not in t:
        parts = [p for p in t.split("/") if p]
        if len(parts) >= 2 and all(
            re.fullmatch(r"[A-Za-zµμΩω%\(\)·\.\-]+", p or "") for p in parts
        ):
            return True
    return False


def _looks_like_isotope_label_token(token: str) -> bool:
    t = _normalize_formula_token_chars(token or "")
    if not t:
        return False
    m = _RE_ISOTOPE_LABEL_TOKEN.fullmatch(t)
    if m and m.group("elem") in _ELEMENT_SYMBOLS:
        return True
    # Common isotope shorthand.
    if t in {"D2O", "T2O"}:
        return True
    return False


def _looks_like_plain_unit_token(token: str) -> bool:
    t = _normalize_unit_placeholder_token(_normalize_formula_token_chars(token or ""))
    if not t:
        return False
    if t.lower() in _PLAIN_UNIT_WORDS:
        return True
    if _RE_PERCENT_SHORT_UNIT.fullmatch(t):
        return True
    if _RE_PERCENT_BRACKET_UNIT.fullmatch(t):
        return True
    if _RE_SIMPLE_UNIT_CONNECT_TOKEN.fullmatch(t):
        return True
    if any(ch in _SUPER_SUB_UNICODE_CHARS for ch in token or "") and re.search(r"[A-Za-zµμΩω]", t):
        return True
    return False


def _is_formula_like_token_without_arrow(
    token: str,
    right_char: str = "",
    right_non_space_char: str = "",
) -> bool:
    t = _normalize_formula_token_chars(token or "")
    if not t:
        return False
    if _is_radical_token_candidate(t):
        return True
    if _is_bond_like_chem_token(t):
        return True
    if _looks_like_roman_oxidation_token(t):
        return True
    if _looks_like_unit_ratio_token(t):
        return True
    if _looks_like_plain_unit_token(token):
        return True
    if _looks_like_isotope_label_token(t):
        return True
    if _RE_BRACKET_COMPLEX_WITH_CHARGE.fullmatch(t):
        return True

    greek_m = _RE_GREEK_PREFIX_TOKEN.fullmatch(t)
    if greek_m and _is_formula_like_token_without_arrow(
        greek_m.group("rest"),
        right_char=right_char,
        right_non_space_char=right_non_space_char,
    ):
        return True

    if _RE_PLAIN_FORMULA_CHAIN.fullmatch(t):
        symbols = re.findall(r"[A-Z][a-z]?", t)
        if len(symbols) >= 2 and all(sym in _ELEMENT_SYMBOLS for sym in symbols):
            return True

    return any(
        _build_token_formula_marks(
            t,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        )
    )


def _is_formula_side_token(side: str) -> bool:
    t = _normalize_formula_token_chars((side or "").strip())
    if not t:
        return False
    if _RE_SINGLE_ELEMENT_TOKEN.fullmatch(t):
        return t in _ELEMENT_SYMBOLS
    return _is_formula_like_token_without_arrow(t)


def _looks_like_reaction_arrow_token(token: str) -> bool:
    t = _normalize_formula_token_chars(token or "")
    if not t:
        return False
    m = _RE_REACTION_ARROW_TOKEN.fullmatch(t)
    if not m:
        return False
    left = m.group("left").strip()
    right = m.group("right").strip()
    if not left or not right:
        return False
    return _is_formula_side_token(left) and _is_formula_side_token(right)


def _normalize_chem_token_key(token: str) -> str:
    raw = re.sub(r"\s+", "", str(token or "")).strip()
    if not raw:
        return ""
    return _normalize_unit_placeholder_token(_normalize_formula_token_chars(raw))


def _token_matches_patterns(
    patterns: list[re.Pattern[str]],
    token: str,
    normalized_token: str,
) -> bool:
    for pat in patterns:
        if pat.search(token):
            return True
        if normalized_token != token and pat.search(normalized_token):
            return True
    return False


def _resolve_chem_token_policy(
    token: str,
    normalized_token: str,
    chem_runtime: ChemRuntime,
) -> tuple[bool, str]:
    ignore_tokens, ignore_patterns, manual_overrides, allow_tokens, allow_patterns = chem_runtime
    if not normalized_token:
        return False, ""

    manual_mask = manual_overrides.get(normalized_token)
    if manual_mask:
        return True, manual_mask

    is_allowed = (
        normalized_token in allow_tokens
        or _token_matches_patterns(allow_patterns, token, normalized_token)
    )
    if not is_allowed:
        if normalized_token in ignore_tokens:
            return False, ""
        if _token_matches_patterns(ignore_patterns, token, normalized_token):
            return False, ""
        if _looks_like_instrument_model_token(token):
            return False, ""
    return True, ""


def _build_chem_dictionary_runtime(chem_cfg) -> ChemRuntime:
    ignore_tokens: set[str] = {
        _normalize_chem_token_key(v)
        for v in _INTERNAL_CHEM_IGNORE_TOKENS
        if _normalize_chem_token_key(v)
    }
    ignore_patterns: list[re.Pattern[str]] = list(_INTERNAL_CHEM_IGNORE_PATTERNS)
    manual_overrides: dict[str, str] = {
        _normalize_chem_token_key(k): str(v)
        for k, v in _INTERNAL_CHEM_MANUAL_OVERRIDES.items()
        if _normalize_chem_token_key(k) and str(v)
    }
    allow_tokens: set[str] = {
        _normalize_chem_token_key(v)
        for v in _INTERNAL_CHEM_ALLOW_TOKENS
        if _normalize_chem_token_key(v)
    }
    allow_patterns: list[re.Pattern[str]] = list(_INTERNAL_CHEM_ALLOW_PATTERNS)
    if chem_cfg is None:
        return ignore_tokens, ignore_patterns, manual_overrides, allow_tokens, allow_patterns

    raw_allow_tokens = getattr(chem_cfg, "allow_tokens", []) or []
    for raw in raw_allow_tokens:
        token = _normalize_chem_token_key(str(raw))
        if token:
            allow_tokens.add(token)

    raw_allow_patterns = getattr(chem_cfg, "allow_patterns", []) or []
    for raw in raw_allow_patterns:
        pat = str(raw).strip()
        if not pat:
            continue
        try:
            allow_patterns.append(re.compile(pat))
        except re.error:
            continue

    raw_tokens = getattr(chem_cfg, "ignore_tokens", []) or []
    for raw in raw_tokens:
        token = _normalize_chem_token_key(str(raw))
        if token:
            ignore_tokens.add(token)

    raw_patterns = getattr(chem_cfg, "ignore_patterns", []) or []
    for raw in raw_patterns:
        pat = str(raw).strip()
        if not pat:
            continue
        try:
            ignore_patterns.append(re.compile(pat))
        except re.error:
            continue

    raw_overrides = getattr(chem_cfg, "manual_overrides", {}) or {}
    if isinstance(raw_overrides, dict):
        for raw_token, raw_mask in raw_overrides.items():
            token = _normalize_chem_token_key(str(raw_token))
            mask = str(raw_mask or "")
            if token and mask:
                manual_overrides[token] = mask

    return ignore_tokens, ignore_patterns, manual_overrides, allow_tokens, allow_patterns


def _iter_chem_token_candidates(
    text: str,
    *,
    chem_cfg=None,
    chem_runtime: ChemRuntime | None = None,
):
    if not text:
        return
    trim_right_punct = ",.;:，。；："
    if chem_runtime is None:
        chem_runtime = _build_chem_dictionary_runtime(chem_cfg)
    scan_text = _normalize_formula_scan_text(text)
    for m in _CHEM_TOKEN_RE.finditer(scan_text):
        raw_token = text[m.start():m.end()]
        if not raw_token:
            continue

        token_start = 0
        token_end = len(raw_token)
        while token_end > token_start and raw_token[token_end - 1] in trim_right_punct:
            token_end -= 1
        token = raw_token[token_start:token_end]
        if not token:
            continue

        abs_start = m.start() + token_start
        abs_end = abs_start + len(token)
        normalized_token = _normalize_chem_token_key(token)
        if not normalized_token:
            continue
        right_char = text[abs_end] if abs_end < len(text) else ""
        right_non_space_char = _first_non_space_char(text[abs_end:])

        use_token, manual_mask = _resolve_chem_token_policy(token, normalized_token, chem_runtime)
        if not use_token:
            continue
        yield abs_start, abs_end, token, normalized_token, right_char, right_non_space_char, manual_mask


def _iter_compacted_formula_spans(text: str):
    """Yield spaced spans and compacted tokens for secondary matching.

    This recovers cases such as:
    - Fe - O
    - Fe (III)
    - [ Fe(CN)6 ]3-
    - mg / L
    """
    if not text:
        return
    trim_right_punct = ",.;:，。；："
    scan_text = _normalize_formula_scan_text(text)
    for m in _RE_SPACED_FORMULA_SPAN.finditer(scan_text):
        span = text[m.start():m.end()]
        if not span or not any(ch.isspace() for ch in span):
            continue
        if len(span) > 80:
            continue

        left_trim = len(span) - len(span.lstrip())
        right_trim = len(span) - len(span.rstrip())
        start = m.start() + left_trim
        end = m.end() - right_trim
        while end > start and text[end - 1] in trim_right_punct:
            end -= 1
        if end <= start:
            continue

        original = text[start:end]
        compact_raw = re.sub(r"\s+", "", original)
        if len(compact_raw) < 2:
            continue
        # Skip plain words/sentences.
        if not re.search(r"[A-Za-z0-9α-ωΑ-ΩµμΩω%％]", compact_raw):
            continue
        if not re.search(
            r"[\-−＋+＝=≡/／·•∙⋅\(\)\[\]\{\}<>→←↔^0-9－‐‑‒–—―%％"
            r"⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻₀₁₂₃₄₅₆₇₈₉₊₋]",
            compact_raw,
        ):
            continue

        compact_norm = _normalize_chem_token_key(compact_raw)
        if not compact_norm:
            continue
        yield start, end, original, compact_raw, compact_norm


def _apply_manual_override_mask(
    marks: list[str | None],
    token_start: int,
    token_text: str,
    mask: str,
) -> None:
    limit = min(len(token_text), len(mask))
    for i in range(limit):
        flag = mask[i]
        if flag == "^":
            marks[token_start + i] = "superscript"
        elif flag == "_":
            marks[token_start + i] = "subscript"


def _parse_formula_core_marks(core: str, marks: list[str | None]) -> tuple[bool, int, int]:
    """Parse chemistry core text and mark subscript digits."""
    if not core:
        return False, 0, 0

    i = 0
    n = len(core)
    element_count = 0
    subscript_count = 0
    close_stack: list[str] = []
    group_has_content: list[bool] = []
    component_start = True
    prev_was_dot = False
    saw_group = False

    while i < n:
        ch = core[i]
        if ch in _DOT_SEPARATORS:
            if prev_was_dot or i == 0 or i == n - 1:
                return False, 0, 0
            prev_was_dot = True
            component_start = True
            i += 1
            continue

        prev_was_dot = False
        if component_start:
            while i < n and core[i].isdigit():
                i += 1
            if i >= n:
                return False, 0, 0

        ch = core[i]
        if ch in _BRACKET_OPEN_TO_CLOSE:
            close_stack.append(_BRACKET_OPEN_TO_CLOSE[ch])
            group_has_content.append(False)
            i += 1
            component_start = True
            continue

        if ch in _BRACKET_OPEN_TO_CLOSE.values():
            if not close_stack or ch != close_stack[-1]:
                return False, 0, 0
            if not group_has_content or not group_has_content[-1]:
                return False, 0, 0
            close_stack.pop()
            group_has_content.pop()
            i += 1
            d0 = i
            while i < n and core[i].isdigit():
                i += 1
            if i > d0:
                _mark_style(marks, d0, i, "subscript")
                subscript_count += (i - d0)
            if group_has_content:
                group_has_content[-1] = True
            saw_group = True
            component_start = False
            continue

        if not _is_ascii_upper(ch):
            return False, 0, 0

        symbol = ch
        if i + 1 < n and _is_ascii_lower(core[i + 1]):
            candidate = core[i:i + 2]
            if candidate in _ELEMENT_SYMBOLS:
                symbol = candidate
                i += 2
            elif ch in _ELEMENT_SYMBOLS:
                i += 1
            else:
                return False, 0, 0
        else:
            if symbol not in _ELEMENT_SYMBOLS:
                return False, 0, 0
            i += 1

        element_count += 1
        d0 = i
        while i < n and core[i].isdigit():
            i += 1
        if i > d0:
            _mark_style(marks, d0, i, "subscript")
            subscript_count += (i - d0)
        if group_has_content:
            group_has_content[-1] = True
        saw_group = True
        component_start = False

    if close_stack or prev_was_dot or not saw_group:
        return False, 0, 0
    return True, element_count, subscript_count


def _choose_charge_digits_len(core_prefix: str, charge_digits: str, has_caret: bool) -> int:
    if not charge_digits:
        return 0
    if has_caret:
        return len(charge_digits)
    if core_prefix and core_prefix[-1] in ")]}":
        return len(charge_digits)
    if len(charge_digits) >= 2:
        return 1

    # One trailing digit before +/- is ambiguous:
    # - Fe3+ -> charge digit
    # - NH4+ / MnO4- -> formula digit
    test = core_prefix + charge_digits
    temp_marks: list[str | None] = [None] * len(test)
    ok, element_count, _ = _parse_formula_core_marks(test, temp_marks)
    if not ok:
        return len(charge_digits)
    return 1 if element_count <= 1 else 0


def _apply_isotope_prefix_marks(token: str, core_end: int, marks: list[str | None]) -> None:
    if core_end <= 1 or not token or not token[0].isdigit():
        return

    i = 0
    while i < core_end and token[i].isdigit():
        i += 1
    if i <= 0 or i >= core_end:
        return

    metastable = False
    elem_pos = i
    if token[elem_pos] in {"m", "M"} and (elem_pos + 1) < core_end and _is_ascii_upper(token[elem_pos + 1]):
        metastable = True
        elem_pos += 1
    if elem_pos >= core_end or not _is_ascii_upper(token[elem_pos]):
        return

    prefix = token[:i]
    # Avoid turning stoichiometric coefficients (e.g. 2H2O) into isotope marks.
    if metastable or len(prefix) >= 2 or token.startswith("1O"):
        _mark_style(marks, 0, i, "superscript")
        if metastable:
            _mark_style(marks, i, i + 1, "superscript")


def _build_token_chem_marks(
    token: str,
    right_char: str = "",
    right_non_space_char: str = "",
) -> list[str | None]:
    token = _normalize_formula_token_chars(token or "")
    n = len(token)
    token_marks: list[str | None] = [None] * n
    if n <= 0:
        return token_marks

    # Leading isotope notation with caret, e.g. ^14C, ^13CH4
    if token.startswith("^") and n > 2 and token[1].isdigit():
        nested = _build_token_chem_marks(
            token[1:],
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        )
        if any(nested):
            for i, style in enumerate(nested):
                if style:
                    token_marks[i + 1] = style
            return token_marks

    # Leading radical dot, e.g. ·SO42- / •SO42-.
    # Keep the leading dot in baseline; only format formula body.
    if token[0] in {"·", "•"} and n > 1:
        nested = _build_token_chem_marks(
            token[1:],
            # Charge parsing should be based on the inner token boundary.
            right_char="",
            right_non_space_char="",
        )
        if any(nested):
            for i, style in enumerate(nested):
                if style:
                    token_marks[i + 1] = style
            return token_marks

    # Metastable isotope notation without caret, e.g. 99mTc
    m0 = re.match(r"^(?P<mass>\d{1,3})(?P<meta>[mM])(?P<rest>[A-Z].*)$", token)
    if m0:
        mass = m0.group("mass")
        rest = m0.group("rest")
        nested = _build_token_chem_marks(
            mass + rest,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        )
        if any(nested):
            mass_len = len(mass)
            for i, style in enumerate(nested):
                if not style:
                    continue
                mapped_i = i if i < mass_len else i + 1
                if mapped_i < n:
                    token_marks[mapped_i] = style
            token_marks[mass_len] = "superscript"
            return token_marks

    # Wrapper-enclosed formula/ion, e.g. (Fe3+), [SO42-], (^14C)
    open_br = token[0] if token else ""
    close_br = _BRACKET_OPEN_TO_CLOSE.get(open_br, "")
    if close_br and n >= 3 and token[-1] == close_br:
        inner = token[1:-1]
        nested = _build_token_chem_marks(
            inner,
            # Inner formula right boundary is the closing bracket, not outer sentence char.
            right_char=close_br,
            right_non_space_char=close_br,
        )
        if any(nested):
            for i, style in enumerate(nested):
                if style:
                    token_marks[i + 1] = style
            return token_marks

    # Mass-spec adduct notation: [M+H]+, [M+2H]2+, [M-H]-
    m_adduct = _RE_MASS_ADDUCT_CHARGE.match(token)
    if m_adduct:
        body = token[1:token.rfind("]")]
        # Avoid misclassifying coordination complexes like [ZnCl4]2- as adducts.
        has_inner_adduct_sign = any(ch in _CHARGE_SIGNS for ch in body)
        if has_inner_adduct_sign and any(_is_ascii_upper(ch) for ch in body):
            charge = m_adduct.group(1)
            charge_start = n - len(charge)
            _mark_style(token_marks, charge_start, n, "superscript")
            return token_marks

    sign_start = n
    while sign_start > 0 and token[sign_start - 1] in _CHARGE_SIGNS:
        sign_start -= 1
    has_charge_tail = sign_start < n
    if has_charge_tail and right_char and _is_word_or_cjk(right_char):
        has_charge_tail = False
        sign_start = n

    charge_start = n
    charge_digits_start = sign_start
    if has_charge_tail:
        dstart = sign_start
        while dstart > 0 and token[dstart - 1].isdigit():
            dstart -= 1
        has_caret = dstart > 0 and token[dstart - 1] == "^"
        core_prefix = token[:dstart]
        charge_digits = token[dstart:sign_start]
        # Bond-like text split by whitespace, e.g. "Fe- O", should not treat '-' as ion charge.
        if (
            not charge_digits
            and not has_caret
            and right_char
            and right_char.isspace()
            and right_non_space_char
            and _is_ascii_upper(right_non_space_char)
            and core_prefix in _ELEMENT_SYMBOLS
        ):
            has_charge_tail = False
            sign_start = n
            charge_start = n
            charge_digits_start = n
        if not has_charge_tail:
            charge_start = n
            charge_digits_start = n
        else:
            charge_digits_len = _choose_charge_digits_len(core_prefix, charge_digits, has_caret)
            charge_digits_start = sign_start - charge_digits_len
            if has_caret:
                charge_start = dstart - 1
                charge_digits_start = dstart
            else:
                charge_start = charge_digits_start
    else:
        charge_start = n

    core = token[:charge_start]
    if not core:
        return token_marks

    core_marks: list[str | None] = [None] * len(core)
    ok, _, _ = _parse_formula_core_marks(core, core_marks)
    if not ok:
        return token_marks
    for i, style in enumerate(core_marks):
        if style:
            token_marks[i] = style

    _apply_isotope_prefix_marks(token, charge_start, token_marks)
    if has_charge_tail and charge_digits_start < n:
        _mark_style(token_marks, charge_digits_start, n, "superscript")

    return token_marks


def _is_radical_token_candidate(token: str) -> bool:
    if not token:
        return False
    open_br = token[0]
    close_br = _BRACKET_OPEN_TO_CLOSE.get(open_br, "")
    if close_br and len(token) >= 3 and token[-1] == close_br:
        inner = token[1:-1]
        if any(ch in _RADICAL_DOT_CHARS for ch in inner):
            return True
    if not any(ch in _RADICAL_DOT_CHARS for ch in token):
        return False

    # Prefix/suffix radical marks, e.g. ·OH, OH·
    if token[0] in _RADICAL_DOT_CHARS or token[-1] in _RADICAL_DOT_CHARS:
        return True

    # Dot next to charge patterns, e.g. O2·-, SO4^-·, SO4·2-
    for i, ch in enumerate(token):
        if ch not in _RADICAL_DOT_CHARS:
            continue
        prev_ch = token[i - 1] if i > 0 else ""
        next_ch = token[i + 1] if (i + 1) < len(token) else ""
        if prev_ch in _CHARGE_SIGNS + "^" or next_ch in _CHARGE_SIGNS + "^":
            return True
    return False


def _build_token_radical_marks(
    token: str,
    right_char: str = "",
    right_non_space_char: str = "",
) -> list[str | None]:
    """Recognize broader radical notations while keeping radical dot baseline."""
    n = len(token)
    marks: list[str | None] = [None] * n
    if n <= 0 or not _is_radical_token_candidate(token):
        return marks

    # Wrapper-enclosed radicals, e.g. (·OH), [SO4·-]
    open_br = token[0]
    close_br = _BRACKET_OPEN_TO_CLOSE.get(open_br, "")
    if close_br and n >= 3 and token[-1] == close_br:
        nested = _build_token_radical_marks(
            token[1:-1],
            # Inner radical right boundary is the closing bracket.
            right_char=close_br,
            right_non_space_char=close_br,
        )
        if any(nested):
            for i, style in enumerate(nested):
                if style:
                    marks[i + 1] = style
            return marks

    cleaned_chars: list[str] = []
    mapping: list[int] = []
    for idx, ch in enumerate(token):
        if ch in _RADICAL_DOT_CHARS:
            continue
        cleaned_chars.append(ch)
        mapping.append(idx)

    cleaned = "".join(cleaned_chars)
    if not cleaned:
        return marks

    nested = _build_token_chem_marks(
        cleaned,
        right_char=right_char,
        right_non_space_char=right_non_space_char,
    )
    # In radical contexts, O2- / N2- style is usually "subscript stoichiometry + charge".
    m_single = re.fullmatch(r"([A-Z][a-z]?)(\d)([+\-−＋－])", cleaned)
    if m_single:
        elem = m_single.group(1)
        digit_idx = len(elem)
        if (
            elem in {"O", "N", "S", "C", "H", "P"}
            and digit_idx < len(nested)
            and len(nested) >= 1
            and nested[digit_idx] == "superscript"
            and nested[-1] == "superscript"
        ):
            nested = list(nested)
            nested[digit_idx] = "subscript"

    if not any(nested):
        return marks

    for i, style in enumerate(nested):
        if not style:
            continue
        if i < len(mapping):
            marks[mapping[i]] = style
    return marks


def _is_unit_symbol_char(ch: str) -> bool:
    if _is_ascii_upper(ch) or _is_ascii_lower(ch):
        return True
    return ch in {"µ", "μ", "Ω", "ω", "°", "℃", "%", "‰"}


def _build_token_unit_formula_marks(token: str, right_char: str = "") -> list[str | None]:
    """Recognize scientific unit expressions and mark exponent part.

    Examples:
    - μmol·g-1
    - h-1
    - cm-2
    - mL-1
    """
    raw_token = token or ""
    token = _normalize_formula_token_chars(raw_token)
    n = len(token)
    token_marks: list[str | None] = [None] * n
    if n <= 0:
        return token_marks

    # Unit exponents are expected to be attached with sign/digit hints.
    if not any(ch.isdigit() for ch in token):
        return token_marks

    def _mark_unit_exp(start: int, end: int) -> None:
        for idx in range(max(0, start), min(n, end)):
            if idx < len(raw_token) and raw_token[idx] in _SUPER_SUB_UNICODE_CHARS:
                continue
            token_marks[idx] = "superscript"

    i = 0
    has_mark = False
    while i < n:
        while i < n and token[i].isdigit():
            i += 1
        if i >= n:
            return [None] * n

        seg_start = i
        if token[i] == "(":
            depth = 0
            while i < n:
                ch = token[i]
                if ch == "(":
                    depth += 1
                elif ch == ")":
                    depth -= 1
                    if depth == 0:
                        i += 1
                        break
                elif not (
                    _is_unit_symbol_char(ch)
                    or ch in _DOT_SEPARATORS
                    or ch == "/"
                    or ch.isspace()
                ):
                    return [None] * n
                i += 1
            if depth != 0:
                return [None] * n
        else:
            while i < n and _is_unit_symbol_char(token[i]):
                i += 1
            if i == seg_start:
                return [None] * n

        exp_start = i
        if i < n and token[i] == "^":
            i += 1
            exp_start = i

        has_sign = False
        if i < n and token[i] in _CHARGE_SIGNS:
            has_sign = True
            i += 1

        d0 = i
        while i < n and token[i].isdigit():
            i += 1

        # Guard number-range tokens like "200k-1.1M": this is usually a scale
        # interval, not a unit exponent, and should stay plain text.
        if (
            has_sign
            and i > d0
            and i < n
            and token[i] == "."
            and (i + 1) < n
            and token[i + 1].isdigit()
        ):
            return [None] * n

        # Signed exponent in unit expression, e.g. g-1 / mA-2.
        if has_sign and i > d0:
            _mark_unit_exp(exp_start, i)
            has_mark = True
        elif i > d0:
            # Bare positive exponent, e.g. cm2/m3. Keep this conservative.
            unit_text = token[seg_start:d0]
            has_lower = any(_is_ascii_lower(ch) for ch in unit_text)
            if (
                (has_lower and len(unit_text) <= 3)
                or any(ch in {"µ", "μ"} for ch in unit_text)
                or any(sep in token for sep in ("·", "•"))
                or unit_text.startswith("(")
            ):
                _mark_unit_exp(d0, i)
                has_mark = True
            else:
                return [None] * n

        if i >= n:
            break

        if token[i] in _DOT_SEPARATORS + "/":
            i += 1
            continue
        # Accept implicit product segments, e.g. g-1h-1.
        if _is_unit_symbol_char(token[i]) or token[i] == "(":
            continue
        return [None] * n

    # Avoid cases like CHI760 where trailing digits should stay plain.
    if right_char and _is_word_or_cjk(right_char):
        return [None] * n
    return token_marks if has_mark else [None] * n


def _build_token_formula_marks(
    token: str,
    right_char: str = "",
    right_non_space_char: str = "",
) -> list[str | None]:
    """Unified formula recognizer: chemistry first, then scientific units."""
    token = _normalize_formula_token_chars(token or "")
    n = len(token)
    if n <= 0:
        return []

    # Common tail qualifiers in corpora, e.g. μmol·g-1·h-1-dry
    m_suffix = re.fullmatch(r"(?P<core>.+?)(?P<suffix>[-_][A-Za-z]{2,12})", token)
    if m_suffix:
        core = m_suffix.group("core")
        suffix = m_suffix.group("suffix")
        if any(ch.isdigit() for ch in core):
            nested = _build_token_formula_marks(
                core,
                right_char=right_char,
                right_non_space_char=right_non_space_char,
            )
            if any(nested):
                out: list[str | None] = [None] * n
                for i, style in enumerate(nested):
                    if style:
                        out[i] = style
                return out

    # Unbalanced wrapper noise around units, e.g. "(μmol·g-1·h-1" / "μmol·g-1·h-1)"
    unmatched_tail = {")": ("(", ")"), "]": ("[", "]"), "}": ("{", "}")}
    if token[-1] in unmatched_tail:
        open_ch, close_ch = unmatched_tail[token[-1]]
        if token.count(open_ch) < token.count(close_ch):
            nested = _build_token_formula_marks(
                token[:-1],
                right_char=token[-1],
                right_non_space_char=right_non_space_char,
            )
            if any(nested):
                return list(nested) + [None]

    unmatched_head = {"(": ("(", ")"), "[": ("[", "]"), "{": ("{", "}")}
    if token[0] in unmatched_head:
        open_ch, close_ch = unmatched_head[token[0]]
        if token.count(open_ch) > token.count(close_ch):
            nested = _build_token_formula_marks(
                token[1:],
                right_char=right_char,
                right_non_space_char=right_non_space_char,
            )
            if any(nested):
                out: list[str | None] = [None] * n
                for i, style in enumerate(nested):
                    if style and (i + 1) < n:
                        out[i + 1] = style
                return out

    m_roman = _RE_ROMAN_OX_STATE.fullmatch(token)
    if m_roman and m_roman.group("elem") in _ELEMENT_SYMBOLS:
        roman_marks: list[str | None] = [None] * n
        left = token.find("(")
        right = token.rfind(")")
        if left >= 0 and right > left + 1:
            _mark_style(roman_marks, left + 1, right, "superscript")
        return roman_marks

    if _looks_like_isotope_label_token(token):
        m_iso = _RE_ISOTOPE_LABEL_TOKEN.fullmatch(token)
        if m_iso:
            isotope_marks: list[str | None] = [None] * n
            _mark_style(isotope_marks, m_iso.start("mass"), m_iso.end("mass"), "superscript")
            return isotope_marks

    m_greek = _RE_GREEK_PREFIX_TOKEN.fullmatch(token)
    if m_greek:
        rest = m_greek.group("rest")
        nested = _build_token_formula_marks(
            rest,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        )
        if any(nested):
            offset = len(m_greek.group("prefix") + m_greek.group("sep"))
            greek_marks: list[str | None] = [None] * n
            for i, style in enumerate(nested):
                if style and (offset + i) < n:
                    greek_marks[offset + i] = style
            return greek_marks

    is_hybrid = bool(_RE_HYBRIDIZATION_TOKEN.fullmatch(token))
    if not is_hybrid and "/" in token:
        parts = token.split("/")
        is_hybrid = bool(parts) and all(
            part and _RE_HYBRIDIZATION_TOKEN.fullmatch(part)
            for part in parts
        )
    if is_hybrid:
        hybrid_marks: list[str | None] = [None] * n
        for idx, ch in enumerate(token):
            if ch.isdigit():
                hybrid_marks[idx] = "superscript"
        return hybrid_marks

    radical_marks = _build_token_radical_marks(
        token,
        right_char=right_char,
        right_non_space_char=right_non_space_char,
    )
    if any(radical_marks):
        return radical_marks

    has_upper = any(_is_ascii_upper(ch) for ch in token)
    has_math_hint = any(ch.isdigit() or ch in _CHARGE_SIGNS + "^()[]{}" for ch in token)
    if has_upper and has_math_hint:
        chem_marks = _build_token_chem_marks(
            token,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        )
        if any(chem_marks):
            return chem_marks

    if _RE_BRACKET_COMPLEX_WITH_CHARGE.fullmatch(token):
        complex_marks: list[str | None] = [None] * n
        m_charge = re.search(r"\d*[+\-−＋－]$", token)
        if m_charge:
            _mark_style(complex_marks, m_charge.start(), m_charge.end(), "superscript")
        for m_count in re.finditer(r"\)\d+", token):
            _mark_style(complex_marks, m_count.start() + 1, m_count.end(), "subscript")
        if any(complex_marks):
            return complex_marks

    return _build_token_unit_formula_marks(token, right_char=right_char)


def _is_bond_like_chem_token(token: str) -> bool:
    t = _normalize_formula_token_chars((token or "").strip())
    if not t:
        return False
    if not _RE_BOND_TOKEN.fullmatch(t):
        return False
    symbols = [part for part in re.split(r"[\-−－‐‑‒–—―=≡]+", t) if part]
    return len(symbols) >= 2 and all(sym in _ELEMENT_SYMBOLS for sym in symbols)


def _is_formula_like_token_for_font(
    token: str,
    right_char: str = "",
    right_non_space_char: str = "",
) -> bool:
    t = _normalize_formula_token_chars((token or "").strip())
    if not t:
        return False
    if _is_formula_like_token_without_arrow(
        t,
        right_char=right_char,
        right_non_space_char=right_non_space_char,
    ):
        return True
    return _looks_like_reaction_arrow_token(t)


def _looks_like_instrument_model_token(token: str) -> bool:
    t = _normalize_formula_token_chars((token or "").strip())
    t = re.sub(r"[?？]+$", "", t)
    t = re.sub(r"-type$", "", t, flags=re.IGNORECASE)
    if len(t) < 5:
        return False
    if not t.isascii():
        return False
    t_upper = t.upper()
    if any(ch in t_upper for ch in (_CHARGE_SIGNS + "^()[]{}" + _DOT_SEPARATORS)):
        return False
    if not re.fullmatch(r"[A-Z0-9\-]+", t_upper):
        return False
    m = re.fullmatch(r"([A-Z]{2,})(\d{2,})([A-Z0-9\-]*)", t_upper)
    if not m:
        return False
    prefix = m.group(1)
    digits = m.group(2)
    if len(digits) >= 3:
        return True
    return prefix in _INSTRUMENT_MODEL_PREFIXES


def _build_chem_style_marks(
    text: str,
    chem_cfg=None,
    chem_runtime: ChemRuntime | None = None,
) -> list[str | None]:
    """Build per-char desired style marks for chemistry-like typography."""
    marks: list[str | None] = [None] * len(text)
    if not text:
        return marks

    runtime = chem_runtime if chem_runtime is not None else _build_chem_dictionary_runtime(chem_cfg)

    for abs_start, _, token, _, right_char, right_non_space_char, manual_mask in _iter_chem_token_candidates(
        text,
        chem_cfg=chem_cfg,
        chem_runtime=runtime,
    ):
        if manual_mask:
            _apply_manual_override_mask(marks, abs_start, token, manual_mask)
            continue

        token_marks = _build_token_formula_marks(
            token,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        )
        if not any(token_marks):
            continue
        for i, style in enumerate(token_marks):
            if style:
                marks[abs_start + i] = style

    for span_start, span_end, span_text, compact_raw, compact_norm in _iter_compacted_formula_spans(text):
        right_char = text[span_end] if span_end < len(text) else ""
        right_non_space_char = _first_non_space_char(text[span_end:])
        use_span, manual_mask = _resolve_chem_token_policy(compact_raw, compact_norm, runtime)
        if not use_span:
            continue

        non_space_offsets = [idx for idx, ch in enumerate(span_text) if not ch.isspace()]
        if not non_space_offsets:
            continue

        if manual_mask:
            limit = min(len(non_space_offsets), len(manual_mask))
            for i in range(limit):
                flag = manual_mask[i]
                if flag == "^":
                    marks[span_start + non_space_offsets[i]] = "superscript"
                elif flag == "_":
                    marks[span_start + non_space_offsets[i]] = "subscript"
            continue

        token_marks = _build_token_formula_marks(
            compact_raw,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        )
        if not any(token_marks):
            continue
        limit = min(len(non_space_offsets), len(token_marks))
        for i in range(limit):
            style = token_marks[i]
            if style:
                marks[span_start + non_space_offsets[i]] = style
    return marks


def _build_chem_font_mask(
    text: str,
    chem_cfg=None,
    chem_runtime: ChemRuntime | None = None,
) -> list[bool]:
    """Build per-char mask for formula-like tokens that should keep western font."""
    mask: list[bool] = [False] * len(text)
    if not text:
        return mask

    runtime = chem_runtime if chem_runtime is not None else _build_chem_dictionary_runtime(chem_cfg)

    for abs_start, abs_end, token, _, right_char, right_non_space_char, manual_mask in _iter_chem_token_candidates(
        text,
        chem_cfg=chem_cfg,
        chem_runtime=runtime,
    ):
        if manual_mask:
            for i in range(abs_start, abs_end):
                mask[i] = True
            continue
        if not _is_formula_like_token_for_font(
            token,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        ):
            continue
        for i in range(abs_start, abs_end):
            mask[i] = True

    for span_start, span_end, span_text, compact_raw, compact_norm in _iter_compacted_formula_spans(text):
        right_char = text[span_end] if span_end < len(text) else ""
        right_non_space_char = _first_non_space_char(text[span_end:])
        use_span, manual_mask = _resolve_chem_token_policy(compact_raw, compact_norm, runtime)
        if not use_span:
            continue

        non_space_offsets = [idx for idx, ch in enumerate(span_text) if not ch.isspace()]
        if not non_space_offsets:
            continue

        if manual_mask or _is_formula_like_token_for_font(
            compact_raw,
            right_char=right_char,
            right_non_space_char=right_non_space_char,
        ):
            for idx in non_space_offsets:
                mask[span_start + idx] = True
    return mask


def _set_rpr_vert_align(r_el, style: str) -> None:
    """Apply/override w:vertAlign on run element."""
    if style not in {"superscript", "subscript"}:
        return
    rpr = r_el.find(f"{{{_W_NS}}}rPr")
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r_el.insert(0, rpr)
    va = rpr.find(f"{{{_W_NS}}}vertAlign")
    if va is not None:
        rpr.remove(va)
    va = OxmlElement("w:vertAlign")
    va.set(f"{{{_W_NS}}}val", style)
    rpr.append(va)


def _new_text_run_element(text: str, base_rpr_el):
    r = OxmlElement("w:r")
    if base_rpr_el is not None:
        r.append(deepcopy(base_rpr_el))
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    return r


def _resolve_run_ascii_font(run_el) -> str | None:
    rpr = run_el.find(f"{{{_W_NS}}}rPr")
    if rpr is None:
        return None
    rfonts = rpr.find(f"{{{_W_NS}}}rFonts")
    if rfonts is None:
        return None
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        val = rfonts.get(f"{{{_W_NS}}}{key}")
        if val:
            return val
    return None


def _set_run_all_fonts(run_el, font_name: str) -> bool:
    if not font_name:
        return False
    rpr = run_el.find(f"{{{_W_NS}}}rPr")
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_el.insert(0, rpr)
    rfonts = rpr.find(f"{{{_W_NS}}}rFonts")
    if rfonts is None:
        rfonts = etree.SubElement(rpr, f"{{{_W_NS}}}rFonts")

    changed = False
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        q = f"{{{_W_NS}}}{key}"
        if rfonts.get(q) != font_name:
            rfonts.set(q, font_name)
            changed = True
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme", "cstheme"):
        q = f"{{{_W_NS}}}{attr}"
        if q in rfonts.attrib:
            rfonts.attrib.pop(q, None)
            changed = True
    return changed


def _normalize_super_sub_unicode_font_in_para(para) -> int:
    """Force western font on unicode super/subscript symbols (e.g. ⁻²)."""
    if not para.runs:
        return 0

    changed_chars = 0
    runs_snapshot = list(para.runs)
    for run in runs_snapshot:
        txt = run.text or ""
        if not txt:
            continue
        has_special = any(ch in _SUPER_SUB_UNICODE_CHARS for ch in txt)
        if not has_special:
            continue

        parent = run._element.getparent()
        if parent is None:
            continue
        target_font = _resolve_run_ascii_font(run._element) or (run.font.name or "Times New Roman")

        if all(ch in _SUPER_SUB_UNICODE_CHARS for ch in txt):
            if _set_run_all_fonts(run._element, target_font):
                changed_chars += len(txt)
            continue

        base_rpr = run._element.find(f"{{{_W_NS}}}rPr")
        insert_at = parent.index(run._element)
        seg_start = 0
        current_is_special = txt[0] in _SUPER_SUB_UNICODE_CHARS
        for i in range(1, len(txt) + 1):
            is_boundary = i == len(txt) or (txt[i] in _SUPER_SUB_UNICODE_CHARS) != current_is_special
            if not is_boundary:
                continue
            seg_text = txt[seg_start:i]
            if seg_text:
                new_r = _new_text_run_element(seg_text, base_rpr)
                if current_is_special and _set_run_all_fonts(new_r, target_font):
                    changed_chars += len(seg_text)
                parent.insert(insert_at, new_r)
                insert_at += 1
            if i < len(txt):
                seg_start = i
                current_is_special = txt[i] in _SUPER_SUB_UNICODE_CHARS
        parent.remove(run._element)

    return changed_chars


def _apply_chem_marks_to_para_runs(para, marks: list[str | None]) -> int:
    """Split runs minimally and apply per-char vertical alignment marks."""
    if not para.runs:
        return 0

    changed_chars = 0
    runs_snapshot = list(para.runs)
    cursor = 0

    for run in runs_snapshot:
        txt = run.text or ""
        n = len(txt)
        if n == 0:
            continue

        existing = _run_vert_align(run)
        if existing in {"superscript", "subscript"}:
            cursor += n
            continue

        local_marks = marks[cursor: cursor + n]
        cursor += n
        if not any(local_marks):
            continue

        segments: list[tuple[str, str | None]] = []
        seg_start = 0
        current = local_marks[0]
        for i in range(1, n + 1):
            if i == n or local_marks[i] != current:
                seg_text = txt[seg_start:i]
                if seg_text:
                    segments.append((seg_text, current))
                if i < n:
                    seg_start = i
                    current = local_marks[i]

        parent = run._element.getparent()
        if parent is None:
            continue

        insert_at = parent.index(run._element)
        base_rpr = run._element.find(f"{{{_W_NS}}}rPr")
        for seg_text, seg_style in segments:
            new_r = _new_text_run_element(seg_text, base_rpr)
            if seg_style in {"superscript", "subscript"}:
                _set_rpr_vert_align(new_r, seg_style)
                if not _contains_cjk(seg_text):
                    target_font = _resolve_run_ascii_font(new_r)
                    if target_font:
                        _set_run_all_fonts(new_r, target_font)
                changed_chars += len(seg_text)
            parent.insert(insert_at, new_r)
            insert_at += 1
        parent.remove(run._element)

    return changed_chars


def _apply_formula_font_mask_to_para_runs(para, mask: list[bool]) -> int:
    """Split runs and force western font on formula-like token spans."""
    if not para.runs:
        return 0

    changed_chars = 0
    runs_snapshot = list(para.runs)
    cursor = 0

    for run in runs_snapshot:
        txt = run.text or ""
        n = len(txt)
        if n == 0:
            continue

        local_mask = mask[cursor: cursor + n]
        cursor += n
        if not any(local_mask):
            continue

        segments: list[tuple[str, bool]] = []
        seg_start = 0
        current = bool(local_mask[0])
        for i in range(1, n + 1):
            if i == n or bool(local_mask[i]) != current:
                seg_text = txt[seg_start:i]
                if seg_text:
                    segments.append((seg_text, current))
                if i < n:
                    seg_start = i
                    current = bool(local_mask[i])

        parent = run._element.getparent()
        if parent is None:
            continue

        insert_at = parent.index(run._element)
        base_rpr = run._element.find(f"{{{_W_NS}}}rPr")
        for seg_text, need_western_font in segments:
            new_r = _new_text_run_element(seg_text, base_rpr)
            if need_western_font and not _contains_cjk(seg_text):
                target_font = _resolve_run_ascii_font(new_r) or (run.font.name or "Times New Roman")
                if _set_run_all_fonts(new_r, target_font):
                    changed_chars += len(seg_text)
            parent.insert(insert_at, new_r)
            insert_at += 1
        parent.remove(run._element)

    return changed_chars


def _restore_reference_chem_typography(
    para,
    chem_cfg=None,
    chem_runtime: ChemRuntime | None = None,
) -> tuple[int, int]:
    """Recover lost superscript/subscript for chemistry tokens in references."""
    text = "".join((run.text or "") for run in para.runs) if para.runs else (para.text or "")
    if not text:
        return 0, 0

    font_changed = _normalize_super_sub_unicode_font_in_para(para)
    font_mask = _build_chem_font_mask(text, chem_cfg=chem_cfg, chem_runtime=chem_runtime)
    if any(font_mask):
        font_changed += _apply_formula_font_mask_to_para_runs(para, font_mask)

    marks = _build_chem_style_marks(text, chem_cfg=chem_cfg, chem_runtime=chem_runtime)
    if not any(marks):
        return font_changed, 0
    mark_changed = _apply_chem_marks_to_para_runs(para, marks)
    return font_changed, mark_changed


class SectionFormatRule(BaseRule):
    name = "section_format"
    description = "按分区范围应用样式格式"

    def apply(self, doc: Document, config: SceneConfig, tracker: ChangeTracker, context: dict) -> None:
        self._active_config = config
        detected_heading1_indices: set[int] = set()
        for h in context.get("headings", []) or []:
            if getattr(h, "level", None) != "heading1":
                continue
            idx = getattr(h, "para_index", None)
            if isinstance(idx, int):
                detected_heading1_indices.add(idx)
        self._detected_heading1_indices = detected_heading1_indices

        doc_tree: DocTree = context.get("doc_tree")
        if not doc_tree:
            return

        scope: FormatScopeConfig = context.get("format_scope", config.format_scope)
        has_caption_index_context = "caption_indices" in context
        caption_indices = set(context.get("caption_indices", set()))

        chem_cfg = getattr(config, "chem_typography", None)
        if chem_cfg is None:
            chem_enabled = True
            chem_scopes = {"references"}
        else:
            chem_enabled = bool(getattr(chem_cfg, "enabled", True))
            raw_scopes = getattr(chem_cfg, "scopes", {}) or {}
            chem_scopes = {
                key
                for key in ("references", "body", "headings", "abstract_cn", "abstract_en")
                if bool(raw_scopes.get(key, False))
            }
            if bool(raw_scopes.get("abstract", False)):
                chem_scopes.update({"abstract_cn", "abstract_en"})
        if not chem_enabled:
            chem_scopes = set()
        chem_runtime = _build_chem_dictionary_runtime(chem_cfg)

        chem_para_count = {"references": 0, "body": 0, "headings": 0, "abstract": 0}
        chem_mark_char_count = {"references": 0, "body": 0, "headings": 0, "abstract": 0}
        chem_font_char_count = {"references": 0, "body": 0, "headings": 0, "abstract": 0}

        def _apply_chem_restore(scope_key: str, para_obj, *, metric_key: str | None = None) -> None:
            if scope_key not in chem_scopes:
                return
            font_chars, mark_chars = _restore_reference_chem_typography(
                para_obj,
                chem_cfg=chem_cfg,
                chem_runtime=chem_runtime,
            )
            if (font_chars + mark_chars) > 0:
                key = metric_key or scope_key
                if key not in chem_para_count:
                    return
                chem_para_count[key] += 1
                chem_mark_char_count[key] += mark_chars
                chem_font_char_count[key] += font_chars

        section_title_style_map = get_section_title_style_map(config)
        non_numbered_sections = get_non_numbered_title_sections(config)
        non_numbered_title_norms = get_non_numbered_title_norms(config)
        front_matter_title_norms = get_front_matter_title_norms(config)

        for section in doc_tree.sections:
            sec_type = section.section_type
            if not scope.is_section_enabled(sec_type):
                if sec_type in {"abstract_cn", "abstract_en"} and sec_type in chem_scopes:
                    for i in section.paragraph_range:
                        if i >= len(doc.paragraphs):
                            break
                        para = doc.paragraphs[i]
                        if not (para.text or "").strip():
                            continue
                        _apply_chem_restore(sec_type, para, metric_key="abstract")
                continue

            style_key = SECTION_STYLE_MAP.get(sec_type)
            sc = config.styles.get(style_key) if style_key else None
            if not sc:
                continue

            count = 0
            for i in section.paragraph_range:
                if i >= len(doc.paragraphs):
                    break

                para = doc.paragraphs[i]
                if not (para.text or "").strip():
                    continue

                # Back-matter heading misdetected in body: keep chapter-title look but unnumbered.
                if sec_type == "body" and _is_non_numbered_post_title_text(
                    para.text, non_numbered_title_norms
                ):
                    title_sc = config.styles.get("heading1")
                    if title_sc:
                        _lstrip_para_leading_space(para)
                        self._make_para_unnumbered(doc, para)
                        self._format_paragraph(
                            para,
                            title_sc,
                            force_plain=True,
                            alignment_override="center",
                        )
                        _apply_chem_restore("headings", para)
                        count += 1
                    continue

                if sec_type == "body" and _para_is_heading(para, config):
                    _apply_chem_restore("headings", para)
                    continue

                if sec_type == "body":
                    if has_caption_index_context:
                        if i in caption_indices:
                            continue
                        if _para_is_caption(para):
                            text = (para.text or "").strip()
                            if text[:1] in ("表",) or text.lower().startswith("table"):
                                cap_sc = config.styles.get("table_caption")
                            else:
                                cap_sc = config.styles.get("figure_caption")
                            if cap_sc:
                                _center_paragraph(para)
                                for run in para.runs:
                                    _apply_run_font(run, cap_sc.font_cn, cap_sc.font_en, cap_sc.size_pt)
                            continue
                    elif _para_is_caption(para):
                        continue

                if sec_type == "body" and _para_is_equation(para):
                    continue

                style_name = (para.style.name or "") if para.style else ""
                style_id = getattr(para.style, "style_id", "") if para.style else ""
                if _RE_TOC_STYLE.match(style_name) or _RE_TOC_STYLE.match(style_id):
                    continue
                if _para_has_toc_pageref(para):
                    continue

                if sec_type == "body":
                    _normalize_subfigure_ref_space(para)

                if sec_type != "body" and i == section.start_index:
                    if _should_treat_first_para_as_section_title(
                        section,
                        para,
                        config=config,
                        front_matter_title_norms=front_matter_title_norms,
                        non_numbered_title_sections=non_numbered_sections,
                        non_numbered_title_norms=non_numbered_title_norms,
                    ):
                        title_key = section_title_style_map.get(sec_type)
                        title_sc = config.styles.get(title_key) if title_key else None
                        if title_sc:
                            if sec_type in non_numbered_sections:
                                _lstrip_para_leading_space(para)
                                self._make_para_unnumbered(doc, para)
                            self._format_paragraph(
                                para,
                                title_sc,
                                force_plain=True,
                                alignment_override="center",
                            )
                        _apply_chem_restore("headings", para)
                    continue

                if sec_type == "resume" and _para_is_resume_title(para):
                    title_sc = config.styles.get("heading1")
                    if title_sc:
                        _lstrip_para_leading_space(para)
                        self._make_para_unnumbered(doc, para)
                        self._format_paragraph(
                            para,
                            title_sc,
                            force_plain=True,
                            alignment_override="center",
                        )
                        _apply_chem_restore("headings", para)
                        continue

                if sec_type == "appendix" and _para_is_appendix_title(para):
                    title_sc = config.styles.get("heading1")
                    if title_sc:
                        _lstrip_para_leading_space(para)
                        self._make_para_unnumbered(doc, para)
                        self._format_paragraph(
                            para,
                            title_sc,
                            force_plain=True,
                            alignment_override="center",
                        )
                        _apply_chem_restore("headings", para)
                        continue

                force_plain = sec_type == "resume"
                self._format_paragraph(
                    para,
                    sc,
                    force_plain=force_plain,
                    clear_indent_for_lists=(sec_type == "references"),
                )
                if sec_type == "references":
                    _apply_chem_restore("references", para)
                elif sec_type == "body":
                    _apply_chem_restore("body", para)
                elif sec_type in {"abstract_cn", "abstract_en"}:
                    _apply_chem_restore(sec_type, para, metric_key="abstract")
                count += 1

            tracker.record(
                rule_name=self.name,
                target=f"{count} 段落",
                section=sec_type,
                change_type="format",
                before="(mixed)",
                after=f"{sc.font_cn} {sc.size_pt}pt",
                paragraph_index=-1,
            )

        for scope_key, label, sec in (
            ("references", "参考文献", "references"),
            ("body", "正文", "body"),
            ("abstract", "摘要/Abstract", "global"),
            ("headings", "标题", "global"),
        ):
            if chem_para_count[scope_key] <= 0:
                continue
            tracker.record(
                rule_name=self.name,
                target=f"{chem_para_count[scope_key]} 个{label}段落",
                section=sec,
                change_type="format",
                before="化学式格式异常",
                after=(
                    f"已恢复上下标字符 {chem_mark_char_count[scope_key]} 个，"
                    f"字体归一化 {chem_font_char_count[scope_key]} 个"
                ),
                paragraph_index=-1,
            )

        self._enforce_reference_entries_no_indent(doc, doc_tree, tracker)
        self._format_front_matter_title_fallback(
            doc,
            doc_tree,
            config,
            scope,
            tracker,
            section_title_style_map=section_title_style_map,
            front_matter_title_norms=front_matter_title_norms,
            non_numbered_sections=non_numbered_sections,
            non_numbered_title_norms=non_numbered_title_norms,
        )
        self._ensure_front_title_section_breaks(doc, doc_tree, tracker)
        self._ensure_chapter_section_breaks(doc, doc_tree, tracker)
        self._ensure_post_title_section_breaks(doc, doc_tree, tracker)
        removed_caption_breaks = _remove_caption_table_section_breaks(doc)
        if removed_caption_breaks:
            tracker.record(
                rule_name=self.name,
                target=f"{removed_caption_breaks} 个题注段落",
                section="body",
                change_type="section_break",
                before="题注与表格之间存在分节符",
                after="已清理异常分节符",
                paragraph_index=-1,
            )
        self._center_figures_and_tables(
            doc,
            tracker,
            caption_indices=caption_indices,
            has_caption_index_context=has_caption_index_context,
        )

    def _center_figures_and_tables(
        self,
        doc,
        tracker,
        caption_indices: set[int] | None = None,
        has_caption_index_context: bool = False,
    ):
        fig_count = 0

        for para in doc.paragraphs:
            if _para_has_image(para):
                _center_paragraph(para)
                fig_count += 1

        if has_caption_index_context:
            for i in sorted(caption_indices or set()):
                if 0 <= i < len(doc.paragraphs):
                    _center_paragraph(doc.paragraphs[i])
                    fig_count += 1
        else:
            for para in doc.paragraphs:
                text = (para.text or "").strip()
                if text and _RE_FIGURE_CAPTION.match(text):
                    _center_paragraph(para)
                    fig_count += 1

        tbl_count = 0
        for table in doc.tables:
            _center_table(table)
            tbl_count += 1

        if fig_count or tbl_count:
            tracker.record(
                rule_name=self.name,
                target=f"{fig_count} 图片/标题, {tbl_count} 表格",
                section="global",
                change_type="format",
                before="(mixed)",
                after="居中+清除缩进",
                paragraph_index=-1,
            )

    def _enforce_reference_entries_no_indent(self, doc, doc_tree, tracker):
        """Best-effort safety net for reference entries (even if section detection drifts)."""
        start = None
        end = len(doc.paragraphs) - 1

        active_cfg = getattr(self, "_active_config", None)
        non_numbered_sections = (
            get_non_numbered_title_sections(active_cfg)
            if active_cfg is not None
            else set(_NON_NUMBERED_TITLE_SECTIONS)
        )
        non_numbered_title_norms = (
            get_non_numbered_title_norms(active_cfg)
            if active_cfg is not None
            else None
        )
        front_matter_title_norms = (
            get_front_matter_title_norms(active_cfg)
            if active_cfg is not None
            else None
        )

        title_start = None
        for i, para in enumerate(doc.paragraphs):
            text = (para.text or "").strip()
            if not text:
                continue
            if _looks_like_section_title_line(
                "references",
                text,
                front_matter_title_norms=front_matter_title_norms,
                non_numbered_title_sections=non_numbered_sections,
                non_numbered_title_norms=non_numbered_title_norms,
            ):
                title_start = i + 1
                break

        ref_sec = doc_tree.get_section("references") if doc_tree else None
        if ref_sec is not None:
            start = max(0, int(getattr(ref_sec, "start_index", 0)))
            if title_start is not None:
                start = min(start, title_start)
        else:
            start = title_start
            if start is None:
                return

        count = 0
        for i in range(max(0, start), min(len(doc.paragraphs) - 1, end) + 1):
            para = doc.paragraphs[i]
            if _normalize_reference_entry_paragraph(para):
                count += 1

        if count:
            tracker.record(
                rule_name=self.name,
                target=f"{count} 条参考文献",
                section="references",
                change_type="format",
                before="可能存在首行缩进/列表缩进",
                after="左对齐且无缩进",
                paragraph_index=-1,
            )

    def _resolve_section_title_index(
        self,
        doc,
        section,
        *,
        front_matter_title_norms: set[str] | None = None,
        non_numbered_title_sections: set[str] | None = None,
        non_numbered_title_norms: set[str] | None = None,
    ) -> int | None:
        if section is None or len(doc.paragraphs) <= 0:
            return None

        start = max(0, int(getattr(section, "start_index", -1)))
        end = min(len(doc.paragraphs) - 1, int(getattr(section, "end_index", start)))
        if end < start:
            return None

        scan_end = min(end, start + 16)
        for idx in range(start, scan_end + 1):
            para = doc.paragraphs[idx]
            if not (para.text or "").strip():
                continue
            if _should_treat_first_para_as_section_title(
                section,
                para,
                config=getattr(self, "_active_config", None),
                front_matter_title_norms=front_matter_title_norms,
                non_numbered_title_sections=non_numbered_title_sections,
                non_numbered_title_norms=non_numbered_title_norms,
            ):
                return idx
        return None

    def _resolve_front_title_without_section(
        self,
        doc,
        doc_tree,
        sec_type: str,
        *,
        front_matter_title_norms: set[str] | None = None,
        non_numbered_title_sections: set[str] | None = None,
        non_numbered_title_norms: set[str] | None = None,
    ) -> int | None:
        """Best-effort front-title recovery when doc_tree misses abstract/toc sections."""
        total = len(doc.paragraphs)
        if total <= 0:
            return None

        upper = min(total - 1, 120)
        if doc_tree is not None:
            toc_sec = doc_tree.get_section("toc")
            body_sec = doc_tree.get_section("body")
            for sec in (toc_sec, body_sec):
                if sec is None:
                    continue
                start = int(getattr(sec, "start_index", -1))
                if start > 0:
                    upper = min(upper, start - 1)
        if upper < 0:
            return None

        for idx in range(0, upper + 1):
            para = doc.paragraphs[idx]
            text = (para.text or "").strip()
            if not text:
                continue
            if _looks_like_section_title_line(
                sec_type,
                text,
                front_matter_title_norms=front_matter_title_norms,
                non_numbered_title_sections=non_numbered_title_sections,
                non_numbered_title_norms=non_numbered_title_norms,
            ):
                return idx
        return None

    def _format_front_matter_title_fallback(
        self,
        doc,
        doc_tree,
        config: SceneConfig,
        scope: FormatScopeConfig,
        tracker: ChangeTracker,
        *,
        section_title_style_map: dict[str, str],
        front_matter_title_norms: set[str],
        non_numbered_sections: set[str],
        non_numbered_title_norms: set[str],
    ) -> None:
        """Restore front-matter title style even if section bodies are disabled."""
        style_map = dict(section_title_style_map or {})
        style_map.setdefault("toc", "toc_title")

        count = 0
        used_title_indices: set[int] = set()
        for sec_type in ("abstract_cn", "abstract_en", "toc"):
            sec = doc_tree.get_section(sec_type)

            title_idx = None
            if sec is not None:
                title_idx = self._resolve_section_title_index(
                    doc,
                    sec,
                    front_matter_title_norms=front_matter_title_norms,
                    non_numbered_title_sections=non_numbered_sections,
                    non_numbered_title_norms=non_numbered_title_norms,
                )
            if title_idx is None:
                title_idx = self._resolve_front_title_without_section(
                    doc,
                    doc_tree,
                    sec_type,
                    front_matter_title_norms=front_matter_title_norms,
                    non_numbered_title_sections=non_numbered_sections,
                    non_numbered_title_norms=non_numbered_title_norms,
                )
            if title_idx is None:
                continue
            if title_idx in used_title_indices:
                continue

            style_key = style_map.get(sec_type)
            sc = config.styles.get(style_key) if style_key else None
            if not sc:
                continue

            para = doc.paragraphs[title_idx]
            # Front-matter titles should keep heading semantics (outline level),
            # but stay unnumbered and use scene-specific visual style.
            self._make_para_unnumbered(doc, para)
            _lstrip_para_leading_space(para)
            self._format_paragraph(
                para,
                sc,
                force_plain=True,
                alignment_override="center",
            )
            count += 1
            used_title_indices.add(title_idx)

        if count:
            tracker.record(
                rule_name=self.name,
                target=f"{count} front titles",
                section="global",
                change_type="format",
                before="scope disabled/mixed",
                after="title fallback formatted",
                paragraph_index=-1,
            )

    def _ensure_front_title_section_breaks(self, doc, doc_tree, tracker):
        """Ensure front-matter title blocks start with a next-page section break."""
        if not doc_tree:
            return

        active_cfg = getattr(self, "_active_config", None)
        front_matter_title_norms = (
            get_front_matter_title_norms(active_cfg)
            if active_cfg is not None
            else None
        )
        non_numbered_sections = (
            get_non_numbered_title_sections(active_cfg)
            if active_cfg is not None
            else set(_NON_NUMBERED_TITLE_SECTIONS)
        )
        non_numbered_title_norms = (
            get_non_numbered_title_norms(active_cfg)
            if active_cfg is not None
            else None
        )

        body_el = doc.element.body
        base_sectpr = body_el.find(f"{{{_W_NS}}}sectPr")
        boundaries: set[int] = set()

        for sec_type in ("abstract_cn", "abstract_en", "toc"):
            sec = doc_tree.get_section(sec_type)
            title_idx = None
            if sec is not None:
                title_idx = self._resolve_section_title_index(
                    doc,
                    sec,
                    front_matter_title_norms=front_matter_title_norms,
                    non_numbered_title_sections=non_numbered_sections,
                    non_numbered_title_norms=non_numbered_title_norms,
                )
            if title_idx is None:
                title_idx = self._resolve_front_title_without_section(
                    doc,
                    doc_tree,
                    sec_type,
                    front_matter_title_norms=front_matter_title_norms,
                    non_numbered_title_sections=non_numbered_sections,
                    non_numbered_title_norms=non_numbered_title_norms,
                )
            if title_idx is None or title_idx <= 0:
                continue
            if not _has_nonempty_content(doc, title_idx + 1, len(doc.paragraphs) - 1):
                continue
            boundaries.add(title_idx)

        count = 0
        for boundary in sorted(boundaries):
            prev_idx = boundary - 1
            if prev_idx < 0 or prev_idx >= len(doc.paragraphs):
                continue
            if _has_table_between_paragraphs(doc, prev_idx, boundary):
                continue

            prev = doc.paragraphs[prev_idx]
            prev_el = prev._element
            prev_ppr = prev_el.find(f"{{{_W_NS}}}pPr")

            effective_has_break = False
            for j in range(max(0, boundary - 8), boundary):
                if j >= len(doc.paragraphs):
                    continue
                if not _para_has_sectpr(doc.paragraphs[j]):
                    continue
                if _break_applies_to_boundary(doc, j, boundary):
                    effective_has_break = True
                    break
            if effective_has_break:
                continue

            if prev_ppr is not None and prev_ppr.find(f"{{{_W_NS}}}sectPr") is not None:
                continue

            for br in list(prev_el.iter(f"{{{_W_NS}}}br")):
                if br.get(f"{{{_W_NS}}}type") == "page":
                    br.getparent().remove(br)

            if prev_ppr is None:
                prev_ppr = etree.SubElement(prev_el, f"{{{_W_NS}}}pPr")
                prev_el.insert(0, prev_ppr)

            if base_sectpr is not None:
                new_sect = deepcopy(base_sectpr)
            else:
                new_sect = etree.SubElement(prev_ppr, f"{{{_W_NS}}}sectPr")
            prev_ppr.append(new_sect)
            count += 1

        if count:
            tracker.record(
                rule_name=self.name,
                target=f"{count} front titles",
                section="global",
                change_type="section_break",
                before="鏃犲垎鑺傜/鍒嗛〉绗?",
                after="涓嬩竴鑺傚垎鑺傜",
                paragraph_index=-1,
            )

    def _ensure_chapter_section_breaks(self, doc, doc_tree, tracker):
        """Ensure each chapter heading in body has a preceding next-page section break."""
        body_sec = doc_tree.get_section("body")
        if not body_sec:
            return

        body_el = doc.element.body
        base_sectpr = body_el.find(f"{{{_W_NS}}}sectPr")
        active_cfg = getattr(self, "_active_config", None)
        detected_heading1 = set(getattr(self, "_detected_heading1_indices", set()) or set())
        if detected_heading1:
            candidate_indices = [
                idx for idx in body_sec.paragraph_range
                if idx in detected_heading1
            ]
        else:
            candidate_indices = list(body_sec.paragraph_range)

        count = 0
        for i in candidate_indices:
            if i >= len(doc.paragraphs) or i == 0:
                continue

            para = doc.paragraphs[i]
            if not detected_heading1:
                style_name = para.style.name if para.style else ""
                style_id = getattr(para.style, "style_id", "") if para.style else ""
                level = None
                if active_cfg is not None:
                    level = detect_level_by_style_name(active_cfg, style_name)
                    if level is None:
                        level = detect_level_by_style_name(active_cfg, style_id)
                if level != "heading1" and style_name not in ("Heading 1", "一级标题"):
                    continue

            prev = doc.paragraphs[i - 1]
            if _has_table_between_paragraphs(doc, i - 1, i):
                continue
            prev_el = prev._element
            prev_ppr = prev_el.find(f"{{{_W_NS}}}pPr")

            near_has_break = False
            for j in range(max(body_sec.start_index, i - 3), i):
                if _para_has_sectpr(doc.paragraphs[j]):
                    near_has_break = True
                    break
            if near_has_break:
                continue

            if prev_ppr is not None and prev_ppr.find(f"{{{_W_NS}}}sectPr") is not None:
                continue

            for br in list(prev_el.iter(f"{{{_W_NS}}}br")):
                if br.get(f"{{{_W_NS}}}type") == "page":
                    br.getparent().remove(br)

            if prev_ppr is None:
                prev_ppr = etree.SubElement(prev_el, f"{{{_W_NS}}}pPr")
                prev_el.insert(0, prev_ppr)

            if base_sectpr is not None:
                new_sect = deepcopy(base_sectpr)
            else:
                new_sect = etree.SubElement(prev_ppr, f"{{{_W_NS}}}sectPr")
            prev_ppr.append(new_sect)
            count += 1

        if count:
            tracker.record(
                rule_name=self.name,
                target=f"{count} 个一级标题",
                section="body",
                change_type="section_break",
                before="无分节符/分页符",
                after="下一节分节符",
                paragraph_index=-1,
            )

    def _ensure_post_title_section_breaks(self, doc, doc_tree, tracker):
        """Ensure post-body title blocks start with a next-page section break."""
        if not doc_tree:
            return

        active_cfg = getattr(self, "_active_config", None)
        non_numbered_back_types = (
            get_non_numbered_title_sections(active_cfg)
            if active_cfg is not None
            else set(_NON_NUMBERED_TITLE_SECTIONS)
        )

        body_el = doc.element.body
        base_sectpr = body_el.find(f"{{{_W_NS}}}sectPr")
        boundaries: set[int] = set()

        for sec in doc_tree.sections:
            if sec.section_type not in non_numbered_back_types:
                continue
            if sec.start_index <= 0:
                continue
            if _has_nonempty_content(doc, sec.start_index + 1, sec.end_index):
                boundaries.add(sec.start_index)

        resume_sec = doc_tree.get_section("resume")
        if resume_sec:
            end = min(resume_sec.end_index, len(doc.paragraphs) - 1)
            for i in range(resume_sec.start_index + 1, end + 1):
                if not _para_is_resume_title(doc.paragraphs[i]):
                    continue
                if _has_nonempty_content(doc, i + 1, len(doc.paragraphs) - 1):
                    boundaries.add(i)

        count = 0
        for boundary in sorted(boundaries):
            prev_idx = boundary - 1
            if prev_idx < 0 or prev_idx >= len(doc.paragraphs):
                continue
            if _has_table_between_paragraphs(doc, prev_idx, boundary):
                continue

            prev = doc.paragraphs[prev_idx]
            prev_el = prev._element
            prev_ppr = prev_el.find(f"{{{_W_NS}}}pPr")

            effective_has_break = False
            for j in range(max(0, boundary - 8), boundary):
                if j >= len(doc.paragraphs):
                    continue
                if not _para_has_sectpr(doc.paragraphs[j]):
                    continue
                if _break_applies_to_boundary(doc, j, boundary):
                    effective_has_break = True
                    break
            if effective_has_break:
                continue

            if prev_ppr is not None and prev_ppr.find(f"{{{_W_NS}}}sectPr") is not None:
                continue

            for br in list(prev_el.iter(f"{{{_W_NS}}}br")):
                if br.get(f"{{{_W_NS}}}type") == "page":
                    br.getparent().remove(br)

            if prev_ppr is None:
                prev_ppr = etree.SubElement(prev_el, f"{{{_W_NS}}}pPr")
                prev_el.insert(0, prev_ppr)

            if base_sectpr is not None:
                new_sect = deepcopy(base_sectpr)
            else:
                new_sect = etree.SubElement(prev_ppr, f"{{{_W_NS}}}sectPr")
            prev_ppr.append(new_sect)
            count += 1

        if count:
            tracker.record(
                rule_name=self.name,
                target=f"{count} 个后置标题",
                section="global",
                change_type="section_break",
                before="无分节符/分页符",
                after="下一节分节符",
                paragraph_index=-1,
            )

    def _make_para_unnumbered(self, doc, para) -> None:
        """Keep chapter-heading semantics but remove auto-number participation."""
        heading_style = self._ensure_non_numbered_heading_style(doc)
        if heading_style is not None:
            para.style = heading_style

        ppr = para._element.find(f"{{{_W_NS}}}pPr")
        if ppr is None:
            ppr = etree.SubElement(para._element, f"{{{_W_NS}}}pPr")
            para._element.insert(0, ppr)

        numpr = ppr.find(f"{{{_W_NS}}}numPr")
        if numpr is not None:
            ppr.remove(numpr)
        outline = ppr.find(f"{{{_W_NS}}}outlineLvl")
        if outline is None:
            outline = etree.SubElement(ppr, f"{{{_W_NS}}}outlineLvl")
        outline.set(f"{{{_W_NS}}}val", "0")

    def _ensure_non_numbered_heading_style(self, doc):
        """Ensure the unnumbered chapter-heading style exists and return it."""
        active_cfg = getattr(self, "_active_config", None) or SceneConfig()
        style_name = get_non_numbered_heading_style_name(active_cfg) or _NON_NUMBERED_HEADING_STYLE_NAME

        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        try:
            style.base_style = doc.styles["Normal"]
        except KeyError:
            pass

        ppr = style.element.find(f"{{{_W_NS}}}pPr")
        if ppr is None:
            ppr = etree.SubElement(style.element, f"{{{_W_NS}}}pPr")
        numpr = ppr.find(f"{{{_W_NS}}}numPr")
        if numpr is not None:
            ppr.remove(numpr)
        outline = ppr.find(f"{{{_W_NS}}}outlineLvl")
        if outline is None:
            outline = etree.SubElement(ppr, f"{{{_W_NS}}}outlineLvl")
        outline.set(f"{{{_W_NS}}}val", "0")
        return style

    def _format_paragraph(
        self,
        para,
        sc,
        *,
        force_plain: bool = False,
        alignment_override: str | None = None,
        clear_indent_for_lists: bool = False,
    ):
        """Format one paragraph while preserving md_cleanup special blocks."""
        pf = para.paragraph_format
        has_list = _para_has_num_pr(para)
        is_special = _para_is_special_block(para)
        if clear_indent_for_lists and not is_special:
            style_name = (para.style.name or "") if para.style else ""
            style_id = (getattr(para.style, "style_id", "") or "") if para.style else ""
            style_name_low = style_name.lower()
            style_id_low = style_id.lower()

            # References should not keep list-number indentation from source docs.
            if has_list:
                _clear_para_num_pr(para)

            # Some docs carry numbering/indent via list styles (without local numPr).
            if "list" in style_name_low or "list" in style_id_low or "列表" in style_name:
                try:
                    para.style = "Normal"
                except Exception:
                    pass

            has_list = _para_has_num_pr(para)
        can_adjust_indent = not is_special and (not has_list or clear_indent_for_lists)

        if can_adjust_indent:
            if pf.left_indent is not None:
                pf.left_indent = None
            align_key = alignment_override or sc.alignment
            if align_key in ALIGNMENT_MAP:
                pf.alignment = ALIGNMENT_MAP[align_key]

        pf.space_before = Pt(sc.space_before_pt)
        pf.space_after = Pt(sc.space_after_pt)
        apply_line_spacing(pf, sc.line_spacing_type, sc.line_spacing_pt)

        if can_adjust_indent:
            if sc.first_line_indent_chars > 0:
                pf.first_line_indent = Pt(sc.size_pt * sc.first_line_indent_chars)
                _sanitize_para_indent_ooxml(para, force_zero=False)
            else:
                pf.first_line_indent = Pt(0)
                pf.left_indent = Pt(0)
                pf.right_indent = Pt(0)
                _sanitize_para_indent_ooxml(para, force_zero=True)

        for run in para.runs:
            _apply_run_font(run, sc.font_cn, sc.font_en, sc.size_pt)
            if force_plain:
                run.font.bold = bool(sc.bold)
                run.font.italic = bool(sc.italic)
                continue
            if sc.bold:
                run.font.bold = True
            if sc.italic:
                run.font.italic = True

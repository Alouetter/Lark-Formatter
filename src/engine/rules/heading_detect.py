"""Heading detection rule: identify heading paragraphs in body scope."""

from __future__ import annotations

import re

from docx import Document
from docx.text.paragraph import Paragraph

from src.engine.change_tracker import ChangeTracker
from src.engine.doc_tree import DocTree
from src.engine.rules.base import BaseRule
from src.scene.heading_model import (
    detect_level_by_style_name,
    get_front_matter_title_norms,
    get_non_numbered_title_norms,
    get_post_section_types,
)
from src.scene.schema import HeadingRiskGuardConfig, SceneConfig
from src.utils.toc_entry import (
    looks_like_toc_entry_line,
    looks_like_reference_entry_line,
    looks_like_date_placeholder_line,
    looks_like_numbered_toc_entry_with_page_suffix,
)


_HEADING_MIN_CONF_SCORE = 1
_TABULAR_NUMERIC_TOKEN_RE = re.compile(
    r"^[+\-]?\d+(?:\.\d+)?(?:[eE][+\-]?\d+)?%?$"
)
_CHAPTER_REF_RE = re.compile(r"第[一二三四五六七八九十百零\d]+章")
_SECTION_REF_RE = re.compile(r"第[一二三四五六七八九十百零\d]+节")
_NARRATIVE_PUNCT_RE = re.compile(r"[。；;？！?!]")
_BROKEN_TOC_BOOKMARK_HINTS = (
    "error! bookmark not defined",
    "bookmark not defined",
    "\u672a\u5b9a\u4e49\u4e66\u7b7e",
)

# Heading pattern sets (chapter -> level3), used as style-detection fallback.
CHAPTER_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百零\d]+章[\s\u3000\t]"),
    re.compile(r"^第\d+章[\s\u3000\t]"),
    re.compile(r"^第[一二三四五六七八九十百零\d]+章(?:[\s\u3000\t]*\S.*)?$"),
    # Keep plain "1 绪论" style headings, but avoid broad year/date false positives.
    re.compile(r"^\d{1,2}[\s\u3000\t]+(?!年)"),
]
LEVEL1_PATTERNS = [
    re.compile(r"^\d+\.\d+[\s\u3000\t]"),
    re.compile(r"^第[一二三四五六七八九十百零\d]+节[\s\u3000\t]"),
    re.compile(r"^第[一二三四五六七八九十百零\d]+节(?:[\s\u3000\t]*\S.*)?$"),
]
LEVEL2_PATTERNS = [
    re.compile(r"^\d+\.\d+\.\d+[\s\u3000\t]"),
    re.compile(r"^[一二三四五六七八九十百零]+、"),
]
LEVEL3_PATTERNS = [
    re.compile(r"^\d+\.\d+\.\d+\.\d+[\s\u3000\t]"),
    re.compile(r"^（[一二三四五六七八九十百零]+）"),
    re.compile(r"^\([一二三四五六七八九十百零]+\)"),
]
LEVEL4_PATTERNS = [re.compile(r"^\d+\.\d+\.\d+\.\d+\.\d+[\s\u3000\t]")]
LEVEL5_PATTERNS = [re.compile(r"^\d+\.\d+\.\d+\.\d+\.\d+\.\d+[\s\u3000\t]")]
LEVEL6_PATTERNS = []
LEVEL7_PATTERNS = []


class HeadingInfo:
    """Heading detection result."""

    def __init__(self, para_index: int, level: str, text: str, confidence: str = "high"):
        self.para_index = para_index
        self.level = level  # "heading1" | "heading2" | "heading3" | "heading4"
        self.text = text
        self.confidence = confidence  # "high" | "low"


def _norm_no_space(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip()


def _is_non_numbered_post_title(text: str, non_numbered_title_norms: set[str]) -> bool:
    norm = _norm_no_space(text)
    if not norm:
        return False
    if norm in non_numbered_title_norms:
        return True
    if norm.startswith("附录") or norm.lower().startswith("appendix"):
        return True
    # Keep historical compatibility for long resume subtitle variants.
    return norm.startswith("在学期间发表的学术论文与研究成果")


def _is_front_matter_title(text: str, front_title_norms: set[str]) -> bool:
    norm = _norm_no_space(text).lower()
    return norm in front_title_norms


def _is_reference_entry_line(text: str) -> bool:
    return looks_like_reference_entry_line(text)


def _looks_like_date_placeholder_line(text: str) -> bool:
    return looks_like_date_placeholder_line(text)


def _looks_like_tabular_numeric_line(text: str) -> bool:
    """Heuristic guard for table-like numeric rows accidentally styled as headings."""
    raw = (text or "").strip()
    if "\t" not in raw:
        return False

    tokens = [tok.strip() for tok in raw.split("\t") if tok.strip()]
    if len(tokens) < 2:
        return False

    numeric_like = 0
    for tok in tokens:
        if _TABULAR_NUMERIC_TOKEN_RE.fullmatch(tok):
            numeric_like += 1
            continue
        if re.fullmatch(r"\d+:\d+", tok):
            numeric_like += 1

    if numeric_like < 2:
        return False
    if re.search(r"[A-Za-z\u4e00-\u9fff]", "".join(tokens)):
        return False
    return (numeric_like / len(tokens)) >= 0.67


def _looks_like_broken_toc_entry_line(text: str) -> bool:
    raw = (text or "").strip()
    if not raw or "\t" not in raw:
        return False
    low = raw.lower()
    if any(h in low for h in _BROKEN_TOC_BOOKMARK_HINTS):
        return True
    # Chinese fallback without relying on punctuation variants.
    return ("\u4e66\u7b7e" in low and "\u672a\u5b9a\u4e49" in low)


def _looks_like_chapter_outline_sentence(text: str) -> bool:
    """Reject chapter-outline narrative sentences mis-detected as headings."""
    raw = (text or "").strip()
    if not raw:
        return False

    chapter_hits = len(_CHAPTER_REF_RE.findall(raw))
    section_hits = len(_SECTION_REF_RE.findall(raw))
    total_hits = chapter_hits + section_hits
    if total_hits >= 2:
        return True
    if total_hits == 0:
        return False
    if len(raw) < 24:
        return False
    if not _NARRATIVE_PUNCT_RE.search(raw):
        return False

    sep_count = (
        raw.count("；")
        + raw.count(";")
        + raw.count("、")
        + raw.count("，")
        + raw.count(",")
    )
    return sep_count >= 2


def _is_toc_style_para(para: Paragraph) -> bool:
    style_name = ((para.style.name if para.style else "") or "").strip()
    style_id = (getattr(para.style, "style_id", "") if para.style else "").strip()
    style_name_lower = style_name.lower()
    style_id_lower = style_id.lower()
    return bool(
        re.match(r"^(toc|目录)\s*\d+$", style_name_lower, re.IGNORECASE)
        or re.match(r"^(toc|目录)\s*\d+$", style_id_lower, re.IGNORECASE)
    )


def _para_has_pageref_field(para: Paragraph) -> bool:
    for instr in para._element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText"):
        if "PAGEREF" in ((instr.text or "").upper()):
            return True
    return False


def _detect_by_style(para: Paragraph, config: SceneConfig) -> str | None:
    style_name = para.style.name if para.style else ""
    return detect_level_by_style_name(config, style_name)


def _detect_by_pattern(text: str) -> str | None:
    raw = text.strip()
    for pat in LEVEL7_PATTERNS:
        if pat.match(raw): return "heading8"
    for pat in LEVEL6_PATTERNS:
        if pat.match(raw): return "heading7"
    for pat in LEVEL5_PATTERNS:
        if pat.match(raw): return "heading6"
    for pat in LEVEL4_PATTERNS:
        if pat.match(raw): return "heading5"
    for pat in LEVEL3_PATTERNS:
        if pat.match(raw):
            return "heading4"
    for pat in LEVEL2_PATTERNS:
        if pat.match(raw):
            return "heading3"
    for pat in LEVEL1_PATTERNS:
        if pat.match(raw):
            return "heading2"
    for pat in CHAPTER_PATTERNS:
        if pat.match(raw):
            return "heading1"
    return None


def _detect_heading_level(para: Paragraph, text: str, config: SceneConfig) -> str | None:
    level = _detect_by_style(para, config)
    if level:
        return level
    return _detect_by_pattern(text)


def _heading_confidence_score(
        para: Paragraph,
        text: str,
        style_level: str | None,
        pattern_level: str | None) -> int:
    """Whitelist/blacklist confidence score for heading candidate acceptance."""
    score = 0

    # Whitelist signals.
    if style_level:
        score += 3
    elif pattern_level:
        score += 2
    if len(text) <= 80:
        score += 1

    # Blacklist signals.
    if _is_toc_style_para(para):
        score -= 5
    if looks_like_toc_entry_line(text):
        score -= 5
    if _is_reference_entry_line(text):
        score -= 4
    if _looks_like_date_placeholder_line(text):
        score -= 4
    if _looks_like_tabular_numeric_line(text):
        score -= 6
    if _looks_like_broken_toc_entry_line(text):
        score -= 8
    if _looks_like_chapter_outline_sentence(text):
        score -= 8
    if len(text) >= 180:
        score -= 1

    return score


def _scan_heading_candidates(
        doc: Document,
        allowed_indices: set[int],
        config: SceneConfig,
        non_numbered_title_norms: set[str],
        front_title_norms: set[str],
        confidence: str = "high") -> list[HeadingInfo]:
    found: list[HeadingInfo] = []
    for i in sorted(allowed_indices):
        if i < 0 or i >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[i]
        text = (para.text or "").strip()
        if not text or len(text) > 200:
            continue
        if _is_non_numbered_post_title(text, non_numbered_title_norms):
            continue
        if _is_front_matter_title(text, front_title_norms):
            continue
        if _is_reference_entry_line(text):
            continue
        if _looks_like_date_placeholder_line(text):
            continue
        if _looks_like_tabular_numeric_line(text):
            continue
        if _looks_like_broken_toc_entry_line(text):
            continue
        if _looks_like_chapter_outline_sentence(text):
            continue
        if _is_toc_style_para(para):
            continue
        if looks_like_toc_entry_line(text):
            continue
        if looks_like_numbered_toc_entry_with_page_suffix(text):
            continue
        if _para_has_pageref_field(para):
            continue

        style_level = _detect_by_style(para, config)
        pattern_level = _detect_by_pattern(text) if not style_level else None
        level = style_level or pattern_level
        if level:
            conf_score = _heading_confidence_score(para, text, style_level, pattern_level)
            if conf_score >= _HEADING_MIN_CONF_SCORE:
                found.append(HeadingInfo(i, level, text, confidence))
    return found


def _dedupe_headings(headings: list[HeadingInfo]) -> list[HeadingInfo]:
    by_idx: dict[int, HeadingInfo] = {}
    rank = {"high": 2, "low": 1}
    for h in headings:
        prev = by_idx.get(h.para_index)
        if prev is None or rank.get(h.confidence, 0) > rank.get(prev.confidence, 0):
            by_idx[h.para_index] = h
    return [by_idx[i] for i in sorted(by_idx.keys())]


def _build_excluded_indices(
        doc_tree: DocTree | None, total: int, post_section_types: set[str]) -> set[int]:
    """Indices that should never be treated as body heading candidates."""
    excluded: set[int] = set()
    if not doc_tree:
        return excluded

    for sec in doc_tree.sections:
        if sec.section_type == "toc":
            excluded.update(range(max(0, sec.start_index), min(total - 1, sec.end_index) + 1))
            continue
        if sec.section_type in post_section_types and getattr(sec, "title_confident", True):
            excluded.update(range(max(0, sec.start_index), min(total - 1, sec.end_index) + 1))
    return excluded


def _need_risk_fallback(
        body,
        primary: list[HeadingInfo],
        fallback: list[HeadingInfo],
        total: int,
        guard: HeadingRiskGuardConfig) -> bool:
    if not getattr(guard, "enabled", True):
        return False
    if not body or not fallback:
        return False
    body_len = max(0, body.end_index - body.start_index + 1)
    p_count = len(primary)
    p_chapters = sum(1 for h in primary if h.level == "heading1")
    f_count = len(fallback)
    f_chapters = sum(1 for h in fallback if h.level == "heading1")

    # Risk-1: no headings in body, but enough candidates outside body.
    if p_count == 0 and (
        f_count >= max(1, int(getattr(guard, "no_body_min_candidates", 2)))
        or f_chapters >= max(1, int(getattr(guard, "no_body_min_chapters", 1)))
    ):
        return True
    # Risk-2: headings exist but no chapter in body, while chapter candidates exist outside.
    if p_count > 0 and p_chapters == 0 and (
        f_chapters >= max(1, int(getattr(guard, "no_chapter_min_outside_chapters", 1)))
    ):
        return True
    # Risk-3: body seems too tiny, outside candidates look significantly better.
    tiny_abs = max(1, int(getattr(guard, "tiny_body_max_abs_paras", 4)))
    tiny_ratio = max(0.0, float(getattr(guard, "tiny_body_max_ratio", 0.08)))
    tiny_limit = max(tiny_abs, int(total * tiny_ratio))
    tiny_min = max(1, int(getattr(guard, "tiny_body_min_candidates", 2)))
    tiny_margin = max(0, int(getattr(guard, "tiny_body_primary_margin", 1)))
    if body_len <= tiny_limit and f_count >= max(tiny_min, p_count + tiny_margin):
        return True
    return False


def _scan_toc_overreach_candidates(
        doc: Document,
        doc_tree: DocTree | None,
        config: SceneConfig,
        non_numbered_title_norms: set[str],
        front_title_norms: set[str]) -> list[HeadingInfo]:
    """Recover headings when doc_tree TOC range overreaches into body content."""
    if not doc_tree:
        return []
    toc = doc_tree.get_section("toc")
    if not toc:
        return []
    total = len(doc.paragraphs)
    if total <= 0:
        return []

    span = max(0, toc.end_index - toc.start_index + 1)
    if span <= 0:
        return []
    # Only for suspiciously large TOC ranges.
    if span < max(20, int(total * 0.55)):
        return []

    # Scan the whole overreach range instead of only the tail window;
    # otherwise early chapters may be missed (e.g., only chapter 5 gets recovered).
    start = max(0, toc.start_index + 1)
    end = min(total - 1, toc.end_index)
    if end < start:
        return []
    indices = set(range(start, end + 1))
    return _scan_heading_candidates(
        doc,
        indices,
        config,
        non_numbered_title_norms,
        front_title_norms,
        confidence="low",
    )


class HeadingDetectRule(BaseRule):
    name = "heading_detect"
    description = "识别文档中的标题段落并标记层级"

    def apply(self, doc: Document, config: SceneConfig, tracker: ChangeTracker, context: dict) -> None:
        doc_tree: DocTree = context.get("doc_tree")
        body = doc_tree.get_section("body") if doc_tree else None
        headings: list[HeadingInfo] = []
        total = len(doc.paragraphs)
        guard = config.heading_numbering.risk_guard
        non_numbered_title_norms = get_non_numbered_title_norms(config)
        front_title_norms = get_front_matter_title_norms(config)
        post_section_types = get_post_section_types(config)

        # No body section -> try TOC-tail recovery first, then give up.
        if not body:
            toc_overreach_candidates = _scan_toc_overreach_candidates(
                doc,
                doc_tree,
                config,
                non_numbered_title_norms,
                front_title_norms,
            )
            if toc_overreach_candidates:
                headings = _dedupe_headings(toc_overreach_candidates)
                first_chapter_idx = next((h.para_index for h in headings if h.level == "heading1"), None)
                if getattr(guard, "keep_after_first_chapter", True) and first_chapter_idx is not None:
                    headings = [h for h in headings if h.para_index >= first_chapter_idx]
                tracker.record(
                    rule_name=self.name,
                    target="toc_overreach_recovery",
                    section="body",
                    change_type="numbering",
                    before="body_headings=0 (no body section)",
                    after=f"recovered={len(headings)} from toc_overreach_range",
                    paragraph_index=-1,
                )
            context["headings"] = headings
            return

        body_indices = set(range(max(0, body.start_index), min(total - 1, body.end_index) + 1))
        excluded_indices = _build_excluded_indices(doc_tree, total, post_section_types)
        body_indices -= excluded_indices

        # Default path: only detect inside body section (high confidence).
        primary_headings = _scan_heading_candidates(
            doc, body_indices, config, non_numbered_title_norms, front_title_norms, confidence="high"
        )
        headings = list(primary_headings)

        # Risk fallback: scan outside body only when section split is suspicious.
        all_indices = set(range(total))
        outside_indices = all_indices - body_indices - excluded_indices
        fallback_candidates = _scan_heading_candidates(
            doc, outside_indices, config, non_numbered_title_norms, front_title_norms, confidence="low"
        )
        if _need_risk_fallback(body, primary_headings, fallback_candidates, total, guard):
            merged = _dedupe_headings(primary_headings + fallback_candidates)
            first_chapter_idx = next((h.para_index for h in merged if h.level == "heading1"), None)
            if getattr(guard, "keep_after_first_chapter", True) and first_chapter_idx is not None:
                merged = [h for h in merged if h.para_index >= first_chapter_idx]
            headings = merged
            fallback_used = sum(1 for h in headings if h.confidence == "low")
            tracker.record(
                rule_name=self.name,
                target="risk_fallback",
                section="body",
                change_type="numbering",
                before=f"body_headings={len(primary_headings)}",
                after=f"merged={len(headings)}, fallback_used={fallback_used}",
                paragraph_index=-1,
            )
        elif not headings:
            toc_overreach_candidates = _scan_toc_overreach_candidates(
                doc,
                doc_tree,
                config,
                non_numbered_title_norms,
                front_title_norms,
            )
            if toc_overreach_candidates:
                headings = _dedupe_headings(toc_overreach_candidates)
                first_chapter_idx = next((h.para_index for h in headings if h.level == "heading1"), None)
                if getattr(guard, "keep_after_first_chapter", True) and first_chapter_idx is not None:
                    headings = [h for h in headings if h.para_index >= first_chapter_idx]
                tracker.record(
                    rule_name=self.name,
                    target="toc_overreach_recovery",
                    section="body",
                    change_type="numbering",
                    before="body_headings=0",
                    after=f"recovered={len(headings)} from toc_overreach_range",
                    paragraph_index=-1,
                )

        context["headings"] = headings

        for h in headings:
            tracker.record(
                rule_name=self.name,
                target=f"段落 #{h.para_index}",
                section="body",
                change_type="numbering",
                before=h.text[:50],
                after=f"识别为 {h.level} (置信度: {h.confidence})",
                paragraph_index=h.para_index,
            )

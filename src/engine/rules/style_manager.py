"""样式创建/更新规则"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from src.engine.rules.base import BaseRule
from src.engine.rules.table_format import top_level_table_anchor_positions
from src.engine.change_tracker import ChangeTracker
from src.scene.schema import SceneConfig, StyleConfig
from src.scene.heading_model import get_level_to_style_key, get_level_to_word_style
from src.utils.indent import apply_style_config_indents
from src.utils.line_spacing import apply_line_spacing, sync_spacing_ooxml
from src.utils.ooxml import apply_explicit_rfonts

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# 不全局修改的样式（由 section_format / toc_format / page_setup 按段落直接应用）
SKIP_GLOBAL_STYLES = {
    "references_body", "acknowledgment_body",
    "abstract_body", "abstract_body_en",
    "abstract_title_cn", "abstract_title_en",
    "appendix_body", "resume_body", "symbol_table_body",
    "toc_title", "toc_chapter", "toc_level1", "toc_level2",
    "header_cn", "header_en", "page_number",
}

SPECIAL_WORD_STYLE_NAMES = {
    "normal": "Normal",
}


def _iter_cover_paragraphs(doc: Document, context: dict):
    doc_tree = context.get("doc_tree") if isinstance(context, dict) else None
    if doc_tree is None:
        return

    cover = doc_tree.get_section("cover")
    if cover is None:
        return

    for idx in cover.paragraph_range:
        if 0 <= idx < len(doc.paragraphs):
            yield doc.paragraphs[idx]

    anchor_positions = top_level_table_anchor_positions(doc)
    seen_cells: set[int] = set()
    for tbl_idx, table in enumerate(doc.tables):
        pos = anchor_positions[tbl_idx] if tbl_idx < len(anchor_positions) else -1
        if pos < 0:
            continue
        try:
            sec_type = doc_tree.get_section_for_paragraph(pos)
        except Exception:
            sec_type = None
        if sec_type != "cover":
            continue
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                for para in cell.paragraphs:
                    yield para


def _style_lineage_names(style) -> set[str]:
    names: set[str] = set()
    current = style
    safety = 0
    while current is not None and safety < 32:
        safety += 1
        name = str(getattr(current, "name", "") or "").strip()
        if name:
            names.add(name)
        current = getattr(current, "base_style", None)
    return names


def _collect_cover_used_style_names(doc: Document, context: dict) -> set[str]:
    protected: set[str] = set()
    for para in _iter_cover_paragraphs(doc, context) or []:
        protected.update(_style_lineage_names(getattr(para, "style", None)))
    return protected


def _apply_style_config(style, sc: StyleConfig) -> None:
    """将 StyleConfig 应用到 Word 样式对象"""
    from lxml import etree
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    def _set_on_off_rpr(tag: str, enabled: bool) -> None:
        old = rpr.find(f'{W}{tag}')
        if old is not None:
            rpr.remove(old)
        el = etree.SubElement(rpr, f'{W}{tag}')
        if not enabled:
            el.set(f'{W}val', '0')

    font = style.font
    font.name = sc.font_en
    font.size = Pt(sc.size_pt)
    # 显式控制粗斜体，避免内置样式残留的 bCs/iCs 继续生效。
    font.bold = None
    font.italic = None

    # 中文字体 + 清除主题字体/颜色引用
    rpr = style.element.find(f'{W}rPr')
    if rpr is None:
        rpr = etree.SubElement(style.element, f'{W}rPr')

    _set_on_off_rpr('b', bool(sc.bold))
    _set_on_off_rpr('bCs', bool(sc.bold))
    _set_on_off_rpr('i', bool(sc.italic))
    _set_on_off_rpr('iCs', bool(sc.italic))

    # 强制黑色字体，清除文档主题色（如蓝色 accent1）
    color_el = rpr.find(f'{W}color')
    if color_el is not None:
        rpr.remove(color_el)
    color_el = etree.SubElement(rpr, f'{W}color')
    color_el.set(f'{W}val', '000000')
    rfonts = rpr.find(f'{W}rFonts')
    if rfonts is None:
        rfonts = etree.SubElement(rpr, f'{W}rFonts')
    apply_explicit_rfonts(
        rpr,
        font_cn=sc.font_cn,
        font_en=sc.font_en,
    )
    # 设置显式字体
    rfonts.set(f'{W}eastAsia', sc.font_cn)
    rfonts.set(f'{W}ascii', sc.font_en)
    rfonts.set(f'{W}hAnsi', sc.font_en)
    # 移除主题字体引用（否则 theme 优先级高于显式字体名）
    for theme_attr in ('asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme'):
        rfonts.attrib.pop(f'{W}{theme_attr}', None)
    # 同步 szCs（font.size 只设置 sz，szCs 可能残留原值）
    sz_cs = rpr.find(f'{W}szCs')
    if sz_cs is not None:
        sz_cs.set(f'{W}val', str(int(sc.size_pt * 2)))
    else:
        sz_cs = etree.SubElement(rpr, f'{W}szCs')
        sz_cs.set(f'{W}val', str(int(sc.size_pt * 2)))

    # 段落格式
    pf = style.paragraph_format
    if sc.alignment in ALIGNMENT_MAP:
        pf.alignment = ALIGNMENT_MAP[sc.alignment]
    pf.space_before = Pt(sc.space_before_pt)
    pf.space_after = Pt(sc.space_after_pt)

    apply_line_spacing(pf, sc.line_spacing_type, sc.line_spacing_pt)
    sync_spacing_ooxml(
        style.element,
        space_before_pt=sc.space_before_pt,
        space_after_pt=sc.space_after_pt,
        line_spacing_type=sc.line_spacing_type,
        line_spacing_value=sc.line_spacing_pt,
    )

    apply_style_config_indents(pf, style.element, sc)


class StyleManagerRule(BaseRule):
    name = "style_manager"
    description = "创建/更新文档样式集"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        protected_word_styles: set[str] = set()
        scope = context.get("format_scope") if isinstance(context, dict) else None
        if scope is not None and hasattr(scope, "is_section_enabled"):
            try:
                cover_enabled = bool(scope.is_section_enabled("cover"))
            except Exception:
                cover_enabled = True
            if not cover_enabled:
                protected_word_styles.update(_collect_cover_used_style_names(doc, context))

        style_name_map = {}
        level_style_key_map = get_level_to_style_key(config)
        level_word_style_map = get_level_to_word_style(config)
        for level, style_key in level_style_key_map.items():
            word_style = level_word_style_map.get(level)
            if style_key and word_style:
                style_name_map[style_key] = word_style

        for style_key, sc in config.styles.items():
            # 跳过由 section_format 按段落直接应用的样式
            if style_key in SKIP_GLOBAL_STYLES:
                continue

            word_name = SPECIAL_WORD_STYLE_NAMES.get(
                style_key,
                style_name_map.get(style_key, style_key),
            )
            if word_name in protected_word_styles:
                continue
            try:
                style = doc.styles[word_name]
            except KeyError:
                style = doc.styles.add_style(
                    word_name, WD_STYLE_TYPE.PARAGRAPH)

            _apply_style_config(style, sc)

            tracker.record(
                rule_name=self.name,
                target=f"Style '{word_name}'",
                section="global",
                change_type="style",
                before="(existing)",
                after=f"font={sc.font_cn}/{sc.font_en} {sc.size_pt}pt",
                paragraph_index=-1,
            )


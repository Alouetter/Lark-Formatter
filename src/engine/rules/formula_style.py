"""Formula style normalization rule."""

from __future__ import annotations

import re
from copy import deepcopy

from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.base import BaseRule
from src.engine.rules import table_format as table_helpers
from src.formula_core.normalize import normalize_formula_node
from src.formula_core.parse import parse_document_formulas
from src.formula_core.runtime import FormulaRuleStats
from src.formula_core.semantics import (
    BOLD_ITALIC_MARKER_COMMANDS,
    BOLD_ROMAN_MARKER_COMMANDS,
    FUNCTION_NAMES,
    GREEK_COMMAND_NAMES,
    ITALIC_MARKER_COMMANDS,
    ROMAN_MARKER_COMMANDS,
    UPRIGHT_FAMILY_MARKER_COMMANDS,
)
from src.scene.schema import SceneConfig
from src.utils.line_spacing import sync_spacing_ooxml

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_RE_MATH_TOKEN = re.compile(
    r"\s+"
    r"|\\[A-Za-z]+"
    r"|[0-9]+(?:\.[0-9]+)?"
    r"|[A-Za-zΑ-Ωα-ω]+"
    r"|."
)
_RE_NUMBER_TOKEN = re.compile(r"^[0-9]+(?:\.[0-9]+)?$")
_RE_ASCII_WORD = re.compile(r"^[A-Za-z]+$")
_RE_GREEK_WORD = re.compile(r"^[Α-Ωα-ω]+$")
_RE_COMMAND = re.compile(r"^\\([A-Za-z]+)$")

_FUNCTION_NAMES = FUNCTION_NAMES
_GREEK_COMMANDS = GREEK_COMMAND_NAMES
_PARAGRAPH_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_OPERATOR_TOKENS = {
    "+", "-", "=", "×", "÷", "/", "^", "_", "(", ")", "[", "]", "{", "}", ",",
    ";", ":", ".", "·", "!", "%", "&", "|", "<", ">", "≤", "≥", "≠", "±", "∓",
    "∝", "∞", "∫", "∑", "∏", "√",
}


def _normalized_alignment(value: str | None, *, default: str = "center") -> str:
    raw = str(value or "").strip().lower()
    return raw if raw in _PARAGRAPH_ALIGNMENT_MAP else default


def _safe_positive_float(value, default: float) -> float:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        numeric = float(default)
    if numeric <= 0:
        return float(default)
    return numeric


def _apply_paragraph_layout(
        para,
        *,
        is_block: bool,
        unify_spacing: bool,
        line_spacing: float,
        space_before_pt: float,
        space_after_pt: float,
        block_alignment: str) -> bool:
    changed = False
    if unify_spacing:
        pf = para.paragraph_format
        before_value = Pt(space_before_pt)
        after_value = Pt(space_after_pt)
        if pf.space_before != before_value:
            pf.space_before = before_value
            changed = True
        if pf.space_after != after_value:
            pf.space_after = after_value
            changed = True
        if pf.line_spacing != line_spacing:
            pf.line_spacing = line_spacing
            changed = True
        sync_spacing_ooxml(
            para._element,
            space_before_pt=space_before_pt,
            space_after_pt=space_after_pt,
            line_spacing_type="multiple",
            line_spacing_value=line_spacing,
        )

    target_alignment = _PARAGRAPH_ALIGNMENT_MAP.get(
        _normalized_alignment(block_alignment, default="center"),
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    if is_block and para.alignment != target_alignment:
        para.alignment = target_alignment
        changed = True
    return changed


def _apply_run_style(
        para,
        *,
        unify_font: bool,
        unify_size: bool,
        font_name: str,
        font_size_pt: float) -> bool:
    changed = False
    size_value = Pt(font_size_pt)
    for run in para.runs:
        if unify_font and run.font.name != font_name:
            run.font.name = font_name
            changed = True
        if unify_size and run.font.size != size_value:
            run.font.size = size_value
            changed = True
    return changed


def _paragraph_has_omml(para) -> bool:
    try:
        return bool(
            para._p.xpath(
                ".//*[namespace-uri()='%s' and "
                "(local-name()='oMath' or local-name()='oMathPara')]" % _M_NS
            )
        )
    except Exception:
        return "oMath" in para._p.xml


def _paragraph_key(para) -> str:
    try:
        node = para._p
        return node.getroottree().getpath(node)
    except Exception:
        return str(id(getattr(para, "_p", para)))


def _paragraph_table_element_id(para) -> int | None:
    node = getattr(para, "_p", None)
    while node is not None:
        if _local_name(node) == "tbl":
            return id(node)
        node = node.getparent()
    return None


def _paragraph_is_in_allowed_equation_table(
        para,
        allowed_table_ids: set[int] | None) -> bool:
    if allowed_table_ids is None:
        return True
    table_id = _paragraph_table_element_id(para)
    return table_id in allowed_table_ids


def _first_child(parent, namespace: str, local_name: str):
    if parent is None:
        return None
    target = f"{{{namespace}}}{local_name}"
    for child in parent:
        if child.tag == target:
            return child
    return None


def _ensure_child(parent, namespace: str, local_name: str):
    node = _first_child(parent, namespace, local_name)
    if node is not None:
        return node
    node = etree.SubElement(parent, f"{{{namespace}}}{local_name}")
    return node


def _remove_child(parent, namespace: str, local_name: str) -> bool:
    node = _first_child(parent, namespace, local_name)
    if node is None:
        return False
    parent.remove(node)
    return True


def _local_name(node) -> str:
    tag = getattr(node, "tag", "")
    if not isinstance(tag, str):
        return ""
    return tag.split("}")[-1] if "}" in tag else tag


def _command_name(token: str) -> str | None:
    m = _RE_COMMAND.match(str(token or ""))
    if not m:
        return None
    return m.group(1).strip().lower()


def _is_whitespace_token(token: str) -> bool:
    return bool(token) and token.isspace()


def _is_number_token(token: str) -> bool:
    return bool(_RE_NUMBER_TOKEN.match(str(token or "")))


def _is_operator_token(token: str) -> bool:
    return str(token or "") in _OPERATOR_TOKENS


def _is_alpha_token(token: str) -> bool:
    raw = str(token or "")
    return bool(_RE_ASCII_WORD.match(raw) or _RE_GREEK_WORD.match(raw))


def _is_identifier_token(token: str) -> bool:
    cmd = _command_name(token)
    if cmd in _GREEK_COMMANDS:
        return True
    return _is_alpha_token(token)


def _tokenize_math_text(text: str, *, preserve_words: bool) -> list[str]:
    raw_tokens = [m.group(0) for m in _RE_MATH_TOKEN.finditer(str(text or ""))]
    out: list[str] = []
    for token in raw_tokens:
        if _is_whitespace_token(token) or _is_number_token(token) or _is_operator_token(token):
            out.append(token)
            continue
        cmd = _command_name(token)
        if cmd is not None:
            out.append(token)
            continue
        if _is_alpha_token(token):
            low = token.lower()
            if low in _FUNCTION_NAMES:
                out.append(token)
                continue
            if preserve_words:
                out.append(token)
                continue
            if len(token) <= 1:
                out.append(token)
                continue
            out.extend(list(token))
            continue
        out.append(token)
    return out


def _prev_sig_idx(tokens: list[str], idx: int) -> int | None:
    j = idx - 1
    while j >= 0:
        if not _is_whitespace_token(tokens[j]):
            return j
        j -= 1
    return None


def _next_sig_idx(tokens: list[str], idx: int) -> int | None:
    j = idx + 1
    while j < len(tokens):
        if not _is_whitespace_token(tokens[j]):
            return j
        j += 1
    return None


def _is_function_usage(tokens: list[str], idx: int) -> bool:
    token = str(tokens[idx] or "")
    if token.lower() not in _FUNCTION_NAMES:
        return False
    next_idx = _next_sig_idx(tokens, idx)
    if next_idx is None:
        return False
    nxt = tokens[next_idx]
    if nxt in {"(", "{", "["}:
        return True
    if _is_identifier_token(nxt):
        return True
    cmd = _command_name(nxt)
    if cmd in _GREEK_COMMANDS:
        return True
    return False


def _marker_kind_from_command(cmd: str | None) -> str | None:
    if cmd in BOLD_ITALIC_MARKER_COMMANDS:
        return "bold_italic"
    if cmd in BOLD_ROMAN_MARKER_COMMANDS:
        return "bold_roman"
    if cmd in UPRIGHT_FAMILY_MARKER_COMMANDS:
        return "upright_family"
    if cmd in ROMAN_MARKER_COMMANDS:
        return "roman"
    if cmd in ITALIC_MARKER_COMMANDS:
        return "italic"
    return None


def _explicit_marker_for_token(tokens: list[str], idx: int) -> str | None:
    token = tokens[idx]
    cmd = _command_name(token)
    if cmd is None and not _is_identifier_token(token):
        return None

    # Direct marker usage: \vec v / \mathbf A / \mathrm d
    prev_idx = _prev_sig_idx(tokens, idx)
    prev_cmd = _command_name(tokens[prev_idx]) if prev_idx is not None else None
    marker = _marker_kind_from_command(prev_cmd)
    if marker is not None:
        return marker

    # Marker scope usage: \vec{ab} / \mathbf{ABC}
    depth = 0
    open_idx = None
    j = idx - 1
    while j >= 0:
        tok = tokens[j]
        if tok == "}":
            depth += 1
        elif tok == "{":
            if depth == 0:
                open_idx = j
                break
            depth -= 1
        j -= 1

    if open_idx is None:
        return None
    marker_idx = _prev_sig_idx(tokens, open_idx)
    if marker_idx is None:
        return None
    return _marker_kind_from_command(_command_name(tokens[marker_idx]))


def _apply_marker_style(style: dict[str, bool | None], marker: str | None) -> bool:
    if marker == "bold_italic":
        style["italic"] = True
        style["bold"] = True
        return True
    if marker == "bold_roman":
        style["italic"] = False
        style["bold"] = True
        return True
    if marker == "roman":
        style["italic"] = False
        style["bold"] = False
        return True
    if marker == "upright_family":
        style["italic"] = False
        style["bold"] = False
        return True
    if marker == "italic":
        style["italic"] = True
        style["bold"] = False
        return True
    return False


def _is_differential_symbol(tokens: list[str], idx: int) -> bool:
    token = str(tokens[idx] or "").lower()
    if token != "d":
        return False
    next_idx = _next_sig_idx(tokens, idx)
    if next_idx is None or not _is_identifier_token(tokens[next_idx]):
        return False

    prev_idx = _prev_sig_idx(tokens, idx)
    next2_idx = _next_sig_idx(tokens, next_idx)
    prev = tokens[prev_idx] if prev_idx is not None else None
    next2 = tokens[next2_idx] if next2_idx is not None else None

    if prev in {"/", "∫"}:
        return True
    if prev is None and next2 == "/":
        return True
    if prev in {")", "]"}:
        return True
    if _is_identifier_token(str(prev or "")) or _is_number_token(str(prev or "")):
        return True
    return False


def _is_imaginary_unit(tokens: list[str], idx: int) -> bool:
    token = str(tokens[idx] or "").lower()
    if token not in {"i", "j"}:
        return False
    next_idx = _next_sig_idx(tokens, idx)
    prev_idx = _prev_sig_idx(tokens, idx)
    next_tok = tokens[next_idx] if next_idx is not None else ""
    prev_tok = tokens[prev_idx] if prev_idx is not None else ""

    if next_tok == "^":
        exp_idx = _next_sig_idx(tokens, next_idx)
        if exp_idx is not None and tokens[exp_idx] == "2":
            return True

    if prev_tok in {"=", "+", "-", "(", ","} and (
        next_idx is None or _is_operator_token(next_tok)
    ):
        return True
    return False


def _is_natural_constant_e(tokens: list[str], idx: int) -> bool:
    token = str(tokens[idx] or "").lower()
    if token != "e":
        return False
    next_idx = _next_sig_idx(tokens, idx)
    if next_idx is None or tokens[next_idx] != "^":
        return False
    prev_idx = _prev_sig_idx(tokens, idx)
    prev_tok = tokens[prev_idx] if prev_idx is not None else ""
    if _is_number_token(prev_tok) or _is_identifier_token(prev_tok):
        return False
    if prev_tok in {")", "]", "}"}:
        return False
    return True


def _merge_segments(parts: list[dict]) -> list[dict]:
    merged: list[dict] = []
    for part in parts:
        text = part.get("text", "")
        if not text:
            continue
        italic = part.get("italic")
        bold = part.get("bold")
        if (
            merged
            and merged[-1].get("italic") == italic
            and merged[-1].get("bold") == bold
        ):
            merged[-1]["text"] += text
            continue
        merged.append({"text": text, "italic": italic, "bold": bold})
    return merged


def _semantic_segments_for_text(
    text: str,
    *,
    in_sub_or_sup: bool,
    base_of_superscript: bool,
) -> list[dict]:
    tokens = _tokenize_math_text(text, preserve_words=in_sub_or_sup)
    parts: list[dict] = []

    for idx, token in enumerate(tokens):
        style = {"italic": None, "bold": None}
        cmd = _command_name(token)
        marker = _explicit_marker_for_token(tokens, idx)

        if _is_whitespace_token(token):
            parts.append({"text": token, **style})
            continue

        if _is_number_token(token) or _is_operator_token(token):
            style["italic"] = False
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        if cmd is not None:
            if cmd in _GREEK_COMMANDS:
                style["italic"] = True
                style["bold"] = False
            else:
                style["italic"] = False
                style["bold"] = False
            _apply_marker_style(style, marker)
            parts.append({"text": token, **style})
            continue

        if _apply_marker_style(style, marker):
            parts.append({"text": token, **style})
            continue

        low = token.lower()
        if _is_function_usage(tokens, idx):
            style["italic"] = False
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        if in_sub_or_sup:
            if _is_identifier_token(token):
                style["italic"] = len(token) == 1
            else:
                style["italic"] = False
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        if low == "d" and _is_differential_symbol(tokens, idx):
            style["italic"] = False
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        if low == "e" and _is_natural_constant_e(tokens, idx):
            style["italic"] = False
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        if low == "e" and base_of_superscript and len(token) == 1:
            # OMML superscript base run often contains only "e" (without "^" marker).
            style["italic"] = False
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        if low in {"i", "j"} and _is_imaginary_unit(tokens, idx):
            style["italic"] = False
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        if _is_identifier_token(token):
            style["italic"] = len(token) == 1
            style["bold"] = False
            parts.append({"text": token, **style})
            continue

        style["italic"] = False
        style["bold"] = False
        parts.append({"text": token, **style})

    return _merge_segments(parts)


def _run_has_ancestor_math_arg(m_run, arg_name: str, parent_names: set[str]) -> bool:
    node = m_run
    while node is not None:
        if _local_name(node) == arg_name:
            parent = node.getparent()
            if _local_name(parent) in parent_names:
                return True
        node = node.getparent()
    return False


def _run_in_sub_or_sup(m_run) -> bool:
    return (
        _run_has_ancestor_math_arg(m_run, "sub", {"sSub", "sSubSup", "sPre", "nary"})
        or _run_has_ancestor_math_arg(m_run, "sup", {"sSup", "sSubSup", "sPre", "nary"})
    )


def _run_is_base_of_superscript(m_run) -> bool:
    return _run_has_ancestor_math_arg(m_run, "e", {"sSup", "sSubSup"})


def _set_toggle_rpr(w_rpr, local_name: str, enabled: bool) -> bool:
    changed = False
    node = _first_child(w_rpr, _W_NS, local_name)
    if node is None:
        node = etree.SubElement(w_rpr, f"{{{_W_NS}}}{local_name}")
        changed = True
    desired = "1" if enabled else "0"
    key = f"{{{_W_NS}}}val"
    if (node.get(key) or "") != desired:
        node.set(key, desired)
        changed = True
    return changed


def _math_style_from_msty(m_run) -> tuple[bool, bool] | None:
    m_rpr = _first_child(m_run, _M_NS, "rPr")
    sty = _first_child(m_rpr, _M_NS, "sty")
    if sty is None:
        return None
    key = f"{{{_M_NS}}}val"
    value = str(sty.get(key) or sty.get("val") or "").strip().lower()
    mapping = {
        "p": (False, False),
        "b": (False, True),
        "i": (True, False),
        "bi": (True, True),
    }
    return mapping.get(value)


def _apply_semantic_to_math_run(m_run, *, italic: bool | None, bold: bool | None) -> bool:
    if italic is None and bold is None:
        return False
    msty_style = _math_style_from_msty(m_run)
    m_rpr = _ensure_child(m_run, _M_NS, "rPr")
    w_rpr = _first_child(m_rpr, _W_NS, "rPr")
    if msty_style is not None:
        if w_rpr is None:
            return False
        changed = False
        for local_name in ("i", "iCs", "b", "bCs"):
            changed |= _remove_child(w_rpr, _W_NS, local_name)
        return changed

    changed = False
    w_rpr = _ensure_child(m_rpr, _W_NS, "rPr")
    if italic is not None:
        changed |= _set_toggle_rpr(w_rpr, "i", bool(italic))
        changed |= _set_toggle_rpr(w_rpr, "iCs", bool(italic))
    if bold is not None:
        changed |= _set_toggle_rpr(w_rpr, "b", bool(bold))
        changed |= _set_toggle_rpr(w_rpr, "bCs", bool(bold))
    return changed


def _split_math_run_with_segments(m_run, segments: list[dict]) -> bool:
    parent = m_run.getparent()
    if parent is None:
        return False
    insert_idx = parent.index(m_run)
    template_children = [
        deepcopy(child)
        for child in list(m_run)
        if _local_name(child) != "t"
    ]
    changed = False

    for seg in segments:
        text = seg.get("text", "")
        if not text:
            continue
        new_run = etree.Element(m_run.tag)
        for child in template_children:
            new_run.append(deepcopy(child))
        text_el = etree.SubElement(new_run, f"{{{_M_NS}}}t")
        text_el.text = text
        changed |= _apply_semantic_to_math_run(
            new_run,
            italic=seg.get("italic"),
            bold=seg.get("bold"),
        )
        parent.insert(insert_idx, new_run)
        insert_idx += 1

    parent.remove(m_run)
    return True


def _apply_omml_semantic_style(para) -> bool:
    changed = False
    m_runs = list(
        para._p.xpath(
            ".//*[namespace-uri()='%s' and local-name()='r']" % _M_NS
        )
    )
    for m_run in m_runs:
        text_nodes = [
            child for child in list(m_run)
            if getattr(child, "tag", "") == f"{{{_M_NS}}}t"
        ]
        text = "".join((node.text or "") for node in text_nodes)
        if not text:
            continue
        segments = _semantic_segments_for_text(
            text,
            in_sub_or_sup=_run_in_sub_or_sup(m_run),
            base_of_superscript=_run_is_base_of_superscript(m_run),
        )
        if not segments:
            continue
        if len(segments) == 1 and segments[0].get("text", "") == text and len(text_nodes) == 1:
            changed |= _apply_semantic_to_math_run(
                m_run,
                italic=segments[0].get("italic"),
                bold=segments[0].get("bold"),
            )
            continue
        changed |= _split_math_run_with_segments(m_run, segments)
    return changed


def _apply_omml_run_style(
        para,
        *,
        unify_font: bool,
        unify_size: bool,
        font_name: str,
        font_size_pt: float) -> bool:
    if not (unify_font or unify_size):
        return False

    changed = False
    target_size = str(max(1, int(round(float(font_size_pt) * 2))))
    m_runs = para._p.xpath(
        ".//*[namespace-uri()='%s' and local-name()='r']" % _M_NS
    )
    for m_run in m_runs:
        m_rpr = _ensure_child(m_run, _M_NS, "rPr")
        w_rpr = _ensure_child(m_rpr, _W_NS, "rPr")
        if unify_font:
            w_fonts = _ensure_child(w_rpr, _W_NS, "rFonts")
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                key = f"{{{_W_NS}}}{attr}"
                if w_fonts.get(key) != font_name:
                    w_fonts.set(key, font_name)
                    changed = True
        if unify_size:
            w_sz = _ensure_child(w_rpr, _W_NS, "sz")
            w_szcs = _ensure_child(w_rpr, _W_NS, "szCs")
            key = f"{{{_W_NS}}}val"
            if w_sz.get(key) != target_size:
                w_sz.set(key, target_size)
                changed = True
            if w_szcs.get(key) != target_size:
                w_szcs.set(key, target_size)
                changed = True
    return changed


def _apply_formula_paragraph_style(
        para,
        *,
        is_block: bool,
        formula_only: bool,
        unify_font: bool,
        unify_size: bool,
        unify_spacing: bool,
        font_name: str,
        font_size_pt: float,
        line_spacing: float,
        space_before_pt: float,
        space_after_pt: float,
        block_alignment: str) -> bool:
    para_has_omml = _paragraph_has_omml(para)
    changed = False
    if formula_only:
        changed |= _apply_paragraph_layout(
            para,
            is_block=is_block,
            unify_spacing=unify_spacing,
            line_spacing=line_spacing,
            space_before_pt=space_before_pt,
            space_after_pt=space_after_pt,
            block_alignment=block_alignment,
        )
        changed |= _apply_run_style(
            para,
            unify_font=unify_font,
            unify_size=unify_size,
            font_name=font_name,
            font_size_pt=font_size_pt,
        )
    if para_has_omml:
        changed |= _apply_omml_run_style(
            para,
            unify_font=unify_font,
            unify_size=unify_size,
            font_name=font_name,
            font_size_pt=font_size_pt,
        )
        changed |= _apply_omml_semantic_style(para)
    return changed


def _collect_equation_table_formula_paragraphs(
        doc: Document,
        *,
        allowed_table_ids: set[int] | None = None,
        occurrence_paragraph_ids: set[str] | None = None,
        formula_paragraph_ids: set[str] | None = None) -> list[dict[str, object]]:
    allowed_table_ids = None if allowed_table_ids is None else set(allowed_table_ids)
    occurrence_paragraph_ids = set(occurrence_paragraph_ids or set())
    formula_paragraph_ids = set(formula_paragraph_ids or set())
    paragraph_items: list[dict[str, object]] = []
    seen: set[str] = set()

    for tbl in doc.tables:
        tbl_el = tbl._tbl
        if allowed_table_ids is not None and id(tbl_el) not in allowed_table_ids:
            continue
        if not table_helpers._is_equation_table(tbl_el):
            continue
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_key = _paragraph_key(para)
                    if para_key in seen:
                        continue
                    para_has_omml = _paragraph_has_omml(para)
                    has_formula_occurrence = para_key in occurrence_paragraph_ids
                    if not has_formula_occurrence and not para_has_omml:
                        continue
                    seen.add(para_key)
                    paragraph_items.append(
                        {
                            "paragraph": para,
                            "paragraph_key": para_key,
                            "formula_only": para_key in formula_paragraph_ids,
                            "has_omml": para_has_omml,
                        }
                    )
    return paragraph_items


def _style_formula_table_cells(
        paragraph_items: list[dict[str, object]],
        *,
        unify_font: bool,
        unify_size: bool,
        unify_spacing: bool,
        font_name: str,
        font_size_pt: float,
        line_spacing: float,
        space_before_pt: float,
        space_after_pt: float,
        block_alignment: str,
        already_styled_paragraph_ids: set[str] | None = None) -> dict[str, int]:
    """Apply style only to formula cells in equation tables."""
    already_styled_paragraph_ids = set(already_styled_paragraph_ids or set())
    styled_keys: set[str] = set()
    styled_native_keys: set[str] = set()
    styled_visual_keys: set[str] = set()
    changed_keys: set[str] = set()
    changed_native_keys: set[str] = set()
    changed_visual_keys: set[str] = set()

    for item in paragraph_items:
        para = item.get("paragraph")
        para_key = str(item.get("paragraph_key", ""))
        if not para_key:
            continue

        para_has_omml = bool(item.get("has_omml"))
        if para_key in already_styled_paragraph_ids:
            styled_keys.add(para_key)
            if para_has_omml:
                styled_native_keys.add(para_key)
            else:
                styled_visual_keys.add(para_key)
            continue

        cell_changed = _apply_formula_paragraph_style(
            para,
            is_block=True,
            formula_only=bool(item.get("formula_only")),
            unify_font=unify_font,
            unify_size=unify_size,
            unify_spacing=unify_spacing,
            font_name=font_name,
            font_size_pt=font_size_pt,
            line_spacing=line_spacing,
            space_before_pt=space_before_pt,
            space_after_pt=space_after_pt,
            block_alignment=block_alignment,
        )
        if not cell_changed:
            continue

        styled_keys.add(para_key)
        changed_keys.add(para_key)
        if para_has_omml:
            styled_native_keys.add(para_key)
            changed_native_keys.add(para_key)
        else:
            styled_visual_keys.add(para_key)
            changed_visual_keys.add(para_key)

    return {
        "styled_cells": len(styled_keys),
        "native_cell_styled": len(styled_native_keys),
        "visual_only_cell_styled": len(styled_visual_keys),
        "changed_cells": len(changed_keys),
        "changed_native_cell_styled": len(changed_native_keys),
        "changed_visual_only_cell_styled": len(changed_visual_keys),
    }


def _apply_equation_table_visual_style(
        doc: Document,
        *,
        total_width: int,
        table_alignment: str,
        formula_cell_alignment: str,
        number_alignment: str,
        number_font_name: str,
        number_font_size_pt: float,
        auto_shrink_number_column: bool,
        allowed_table_ids: set[int] | None = None) -> tuple[int, int]:
    """Apply table-level visual style for equation tables.

    Returns:
        (styled_table_count, formatted_number_cell_count)
    """
    table_count = 0
    formatted_number_cells = 0
    for tbl in doc.tables:
        tbl_el = tbl._tbl
        if allowed_table_ids is not None and id(tbl_el) not in allowed_table_ids:
            continue
        if not table_helpers._is_equation_table(tbl_el):
            continue
        table_count += 1
        formatted_number_cells += table_helpers._format_equation_table(
            tbl_el,
            total_width if total_width > 0 else None,
            table_alignment=table_alignment,
            formula_cell_alignment=formula_cell_alignment,
            number_alignment=number_alignment,
            number_font_name=number_font_name,
            number_font_size_pt=number_font_size_pt,
            auto_shrink_number_column=auto_shrink_number_column,
        )
    return table_count, formatted_number_cells


def _scoped_equation_table_ids(doc: Document, context: dict | None) -> set[int] | None:
    if not isinstance(context, dict):
        return None

    doc_tree = context.get("doc_tree")
    raw_target_indices = context.get("target_paragraph_indices")
    has_doc_tree_scope = bool(getattr(doc_tree, "sections", None))
    if raw_target_indices is None and not has_doc_tree_scope:
        return None

    target_indices = None
    if raw_target_indices is not None:
        target_indices = {
            idx for idx in raw_target_indices
            if isinstance(idx, int) and idx >= 0
        }

    anchor_positions = table_helpers.top_level_table_anchor_positions(doc)
    allowed: set[int] = set()
    for tbl_idx, tbl in enumerate(doc.tables):
        pos = anchor_positions[tbl_idx] if tbl_idx < len(anchor_positions) else -1
        if target_indices is not None:
            if pos < 0 or pos not in target_indices:
                continue
        if has_doc_tree_scope:
            if pos < 0:
                continue
            try:
                sec_type = doc_tree.get_section_for_paragraph(pos)
            except Exception:
                continue
            if sec_type != "body":
                continue
        allowed.add(id(tbl._tbl))
    return allowed


class FormulaStyleRule(BaseRule):
    name = "formula_style"
    description = "Formula style normalization"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        cfg = getattr(config, "formula_style", None)
        if cfg is not None and not bool(getattr(cfg, "enabled", False)):
            return

        unify_font = bool(getattr(cfg, "unify_font", True))
        unify_size = bool(getattr(cfg, "unify_size", True))
        unify_spacing = bool(getattr(cfg, "unify_spacing", True))
        formula_table_cfg = getattr(config, "formula_table", None)
        font_name = str(
            getattr(formula_table_cfg, "formula_font_name", "Cambria Math")
        ).strip() or "Cambria Math"
        font_size_pt = _safe_positive_float(
            getattr(formula_table_cfg, "formula_font_size_pt", 12.0),
            12.0,
        )
        line_spacing = _safe_positive_float(
            getattr(formula_table_cfg, "formula_line_spacing", 1.0),
            1.0,
        )
        try:
            space_before_pt = max(
                0.0,
                float(getattr(formula_table_cfg, "formula_space_before_pt", 0.0) or 0.0),
            )
        except (TypeError, ValueError):
            space_before_pt = 0.0
        try:
            space_after_pt = max(
                0.0,
                float(getattr(formula_table_cfg, "formula_space_after_pt", 0.0) or 0.0),
            )
        except (TypeError, ValueError):
            space_after_pt = 0.0
        block_alignment = _normalized_alignment(
            getattr(formula_table_cfg, "block_alignment", "center"),
            default="center",
        )
        table_alignment = _normalized_alignment(
            getattr(formula_table_cfg, "table_alignment", "center"),
            default="center",
        )
        formula_cell_alignment = _normalized_alignment(
            getattr(formula_table_cfg, "formula_cell_alignment", "center"),
            default="center",
        )
        number_alignment = _normalized_alignment(
            getattr(formula_table_cfg, "number_alignment", "right"),
            default="right",
        )
        number_font_name = str(
            getattr(formula_table_cfg, "number_font_name", "Times New Roman")
        ).strip() or "Times New Roman"
        number_font_size_pt = _safe_positive_float(
            getattr(formula_table_cfg, "number_font_size_pt", 10.5),
            10.5,
        )
        auto_shrink_number_column = bool(
            getattr(formula_table_cfg, "auto_shrink_number_column", True)
        )

        runtime = context.setdefault("formula_runtime", {})
        convert_enabled = bool(runtime.get("convert_enabled"))
        runtime["style_enabled"] = True
        parse_result = parse_document_formulas(doc)
        avail_w = table_helpers._available_width(config)
        stats = FormulaRuleStats()
        native_styled = 0
        visual_only = 0
        mixed_omml = 0
        styled_tables = 0
        styled_number_cells = 0
        styled_cells = 0
        styled_cell_native = 0
        styled_cell_visual_only = 0
        styled_cell_changed = 0
        allowed_equation_table_ids = _scoped_equation_table_ids(doc, context)
        table_formula_paragraph_ids = {
            _paragraph_key(occ.paragraph)
            for occ in parse_result.occurrences
            if occ.in_table and occ.is_formula_only
            and _paragraph_is_in_allowed_equation_table(
                occ.paragraph,
                allowed_equation_table_ids,
            )
        }
        table_occurrence_paragraph_ids = {
            _paragraph_key(occ.paragraph)
            for occ in parse_result.occurrences
            if occ.in_table
            and _paragraph_is_in_allowed_equation_table(
                occ.paragraph,
                allowed_equation_table_ids,
            )
        }
        equation_table_paragraph_items = _collect_equation_table_formula_paragraphs(
            doc,
            allowed_table_ids=allowed_equation_table_ids,
            occurrence_paragraph_ids=table_occurrence_paragraph_ids,
            formula_paragraph_ids=table_formula_paragraph_ids,
        )
        equation_table_paragraph_ids = {
            str(item.get("paragraph_key", ""))
            for item in equation_table_paragraph_items
        }
        pre_styled_equation_table_paragraph_ids: set[str] = set()

        seen: set[str] = set()
        for occ in parse_result.occurrences:
            if convert_enabled and occ.source_type != "word_native":
                continue
            if occ.in_table and not _paragraph_is_in_allowed_equation_table(
                occ.paragraph,
                allowed_equation_table_ids,
            ):
                continue
            normalized_node = normalize_formula_node(occ.node)
            stats.note_confidence(normalized_node.confidence)
            stats.matched += 1
            para_key = _paragraph_key(occ.paragraph)
            if para_key in seen:
                continue
            seen.add(para_key)

            if not occ.is_formula_only:
                para = occ.paragraph
                if _paragraph_has_omml(para):
                    changed = _apply_formula_paragraph_style(
                        para,
                        is_block=False,
                        formula_only=False,
                        unify_font=unify_font,
                        unify_size=unify_size,
                        unify_spacing=unify_spacing,
                        font_name=font_name,
                        font_size_pt=font_size_pt,
                        line_spacing=line_spacing,
                        space_before_pt=space_before_pt,
                        space_after_pt=space_after_pt,
                        block_alignment=block_alignment,
                    )
                    if changed:
                        native_styled += 1
                        mixed_omml += 1
                        if para_key in equation_table_paragraph_ids:
                            pre_styled_equation_table_paragraph_ids.add(para_key)
                        stats.converted += 1
                        tracker.record(
                            rule_name=self.name,
                            target=occ.location,
                            section="formula",
                            change_type="style",
                            before="混排段落中的原生公式样式未统一",
                            after="已统一混排段落中的原生公式样式",
                            paragraph_index=occ.paragraph_index,
                        )
                    else:
                        stats.skipped_unsupported += 1
                        tracker.record(
                            rule_name=self.name,
                            target=occ.location,
                            section="formula",
                            change_type="skip",
                            before="混排段落中的原生公式样式未统一",
                            after="未启用任何样式统一项",
                            paragraph_index=occ.paragraph_index,
                        )
                else:
                    stats.skipped_dependency += 1
                    tracker.record(
                        rule_name=self.name,
                        target=occ.location,
                        section="formula",
                        change_type="skip",
                        before="formula paragraph",
                        after="MVP 暂不改写混排公式段落",
                        paragraph_index=occ.paragraph_index,
                    )
                continue

            para = occ.paragraph
            changed = _apply_formula_paragraph_style(
                para,
                is_block=occ.is_block,
                formula_only=True,
                unify_font=unify_font,
                unify_size=unify_size,
                unify_spacing=unify_spacing,
                font_name=font_name,
                font_size_pt=font_size_pt,
                line_spacing=line_spacing,
                space_before_pt=space_before_pt,
                space_after_pt=space_after_pt,
                block_alignment=block_alignment,
            )

            if changed:
                if _paragraph_has_omml(para):
                    native_styled += 1
                    before_text = "原生公式样式未统一"
                    after_text = "已统一原生公式样式"
                    if para_key in equation_table_paragraph_ids:
                        pre_styled_equation_table_paragraph_ids.add(para_key)
                else:
                    visual_only += 1
                    before_text = "公式视觉样式未统一"
                    after_text = "已统一公式视觉样式（未转换为原生公式）"
                    if para_key in equation_table_paragraph_ids:
                        pre_styled_equation_table_paragraph_ids.add(para_key)
                stats.converted += 1
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="style",
                    before=before_text,
                    after=after_text,
                    paragraph_index=occ.paragraph_index,
                )
            else:
                stats.skipped_unsupported += 1
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="skip",
                    before="公式样式未统一",
                    after="未启用任何样式统一项",
                    paragraph_index=occ.paragraph_index,
                )

        styled_table_count, formatted_number_cells = _apply_equation_table_visual_style(
            doc,
            total_width=avail_w,
            table_alignment=table_alignment,
            formula_cell_alignment=formula_cell_alignment,
            number_alignment=number_alignment,
            number_font_name=number_font_name,
            number_font_size_pt=number_font_size_pt,
            allowed_table_ids=allowed_equation_table_ids,
            auto_shrink_number_column=auto_shrink_number_column,
        )
        if styled_table_count:
            styled_tables = styled_table_count
            styled_number_cells = formatted_number_cells
            stats.converted += styled_table_count
            tracker.record(
                rule_name=self.name,
                target="equation_tables",
                section="formula",
                change_type="style",
                before="公式表格容器样式未统一",
                after=(
                    f"已统一 {styled_table_count} 个公式表格容器样式"
                    + (
                        f"，并规范 {styled_number_cells} 个编号单元格样式"
                        if styled_number_cells
                        else ""
                    )
                ),
                paragraph_index=-1,
            )

        table_style_stats = _style_formula_table_cells(
            equation_table_paragraph_items,
            unify_font=unify_font,
            unify_size=unify_size,
            unify_spacing=unify_spacing,
            font_name=font_name,
            font_size_pt=font_size_pt,
            line_spacing=line_spacing,
            space_before_pt=space_before_pt,
            space_after_pt=space_after_pt,
            block_alignment=block_alignment,
            already_styled_paragraph_ids=pre_styled_equation_table_paragraph_ids,
        )
        styled_cells = int(table_style_stats.get("styled_cells", 0))
        styled_cell_native = int(table_style_stats.get("native_cell_styled", 0))
        styled_cell_visual_only = int(table_style_stats.get("visual_only_cell_styled", 0))
        styled_cell_changed = int(table_style_stats.get("changed_cells", 0))
        if styled_cells:
            stats.converted += styled_cell_changed
            detail_parts: list[str] = []
            if styled_cell_native:
                detail_parts.append(f"原生公式 {styled_cell_native} 个")
            if styled_cell_visual_only:
                detail_parts.append(f"仅视觉统一 {styled_cell_visual_only} 个")
            tracker.record(
                rule_name=self.name,
                target="equation_table_cells",
                section="formula",
                change_type="style",
                before="公式表格内样式未统一",
                after=(
                    f"已统一 {styled_cells} 个公式单元格样式"
                    + (f"（{'，'.join(detail_parts)}）" if detail_parts else "")
                ),
                paragraph_index=-1,
            )

        summary_text = (
            f"{stats.to_summary()}, "
            f"native_styled={native_styled}, visual_only={visual_only}, "
            f"mixed_omml={mixed_omml}, styled_tables={styled_tables}, "
            f"styled_number_cells={styled_number_cells}, "
            f"styled_cells={styled_cells}, styled_cell_native={styled_cell_native}, "
            f"styled_cell_visual_only={styled_cell_visual_only}"
        )
        runtime.setdefault("stats", {})[self.name] = {
            "matched": stats.matched,
            "converted": stats.converted,
            "skipped_low_confidence": stats.skipped_low_confidence,
            "skipped_unsupported": stats.skipped_unsupported,
            "skipped_dependency": stats.skipped_dependency,
            "errors": stats.errors,
            "native_styled": native_styled,
            "visual_only": visual_only,
            "mixed_omml": mixed_omml,
            "styled_tables": styled_tables,
            "styled_number_cells": styled_number_cells,
            "styled_cells": styled_cells,
            "styled_cell_native": styled_cell_native,
            "styled_cell_visual_only": styled_cell_visual_only,
        }
        tracker.record(
            rule_name=self.name,
            target="summary",
            section="formula",
            change_type="format",
            before="统计",
            after=summary_text,
            paragraph_index=-1,
        )
        tracker.record(
            rule_name=self.name,
            target="confidence_confirmation",
            section="formula",
            change_type="format",
            before="自动置信分层确认",
            after=f"已在执行日志末尾标注：{stats.confidence_summary()}。",
            paragraph_index=-1,
        )

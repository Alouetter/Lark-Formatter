"""Experimental rule: link body citation tokens to reference entries via fields."""

from __future__ import annotations

import re
from copy import deepcopy

from docx import Document
from lxml import etree

from src.engine.change_tracker import ChangeTracker
from src.engine.rules.base import BaseRule
from src.scene.schema import SceneConfig
from src.utils.ooxml import apply_explicit_rfonts

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


_CITATION_OPEN_TO_CLOSE = {
    "[": "]",
    "\uFF3B": "\uFF3D",
}
_CITATION_OPEN_CHARS = "".join(_CITATION_OPEN_TO_CLOSE.keys())
_CITATION_CLOSE_CHARS = "".join(_CITATION_OPEN_TO_CLOSE.values())

_RE_REFERENCE_LABELS = [
    re.compile(r"^\s*([\[\uFF3B]\s*(?P<num>\d{1,4})\s*[\]\uFF3D])"),
]
_RE_CITATION_TOKEN = re.compile(
    r"(?P<open>\[|\uFF3B)"
    r"(?P<inner>[^\[\]\uFF3B\uFF3D]+)"
    r"(?P<close>\]|\uFF3D)"
)
_RE_NUM_TOKEN = re.compile(r"\d{1,4}")
_RE_REF_FIELD_TARGET = re.compile(r"\bREF\s+(?P<target>[^\s\\]+)", re.IGNORECASE)
_RE_NUMERIC_CITATION_INNER = re.compile(
    r"^[0-9\s,;\uFF0C\uFF1B\u3001\-\u2013\u2014~]+$"
)
_RE_OUTER_PAGE_CANDIDATE = re.compile(
    r"(?P<pages>\d{1,4}(?:\s*[-\u2013\u2014~,\uFF0C\u3001]\s*\d{1,4}){0,5})"
)
_REFERENCE_HEADING_TOKENS = {"\u53C2\u8003\u6587\u732E", "references"}
_REF_ENTRY_BOOKMARK_PREFIX = "_RefEntry_"
_REF_NUM_BOOKMARK_PREFIX = "_RefNum_"
_REF_ENTRY_SEQ_ID = "RefEntry"
_RE_REF_ENTRY_BOOKMARK = re.compile(r"^_RefEntry_(?P<idx>\d+)$")
_OUTER_PAGE_TERMINATORS = set(
    ".,;:!?)]}>\"'"
    "\u3002\uFF0C\uFF1B\uFF1A\uFF01\uFF1F\u3001\uFF09\u3011\u300B\u3009\u201D\u2019"
)


def _extract_reference_number(text: str) -> int | None:
    raw = text or ""
    for pat in _RE_REFERENCE_LABELS:
        m = pat.match(raw)
        if not m:
            continue
        num_raw = m.group("num")
        if not num_raw:
            continue
        try:
            return int(num_raw)
        except (TypeError, ValueError):
            continue
    return None


def _extract_reference_label_end(text: str) -> int | None:
    raw = text or ""
    for pat in _RE_REFERENCE_LABELS:
        m = pat.match(raw)
        if not m:
            continue
        end = m.end(1)
        while end < len(raw) and raw[end].isspace():
            end += 1
        return end
    return None


def _norm_no_space_lower(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip().lower()


def _is_reference_heading_line(text: str) -> bool:
    norm = _norm_no_space_lower(text)
    if not norm:
        return False
    if norm in {token.lower() for token in _REFERENCE_HEADING_TOKENS}:
        return True
    return norm in {"参考文献：", "references:"}


def _iter_para_bookmark_names(para):
    for b in para._element.iter(_w("bookmarkStart")):
        name = b.get(_w("name"))
        if name:
            yield name


def _find_ref_entry_bookmark_name(para) -> str | None:
    for name in _iter_para_bookmark_names(para):
        if name.startswith(_REF_ENTRY_BOOKMARK_PREFIX):
            return name
    return None


def _find_ref_num_bookmark_name(para, entry_bookmark: str | None = None) -> str | None:
    preferred = _compose_ref_num_bookmark_name(entry_bookmark) if entry_bookmark else None
    for name in _iter_para_bookmark_names(para):
        if preferred and name == preferred:
            return name
        if name.startswith(_REF_NUM_BOOKMARK_PREFIX):
            return name
    return None


def _next_ref_entry_serial(existing_names: set[str]) -> int:
    max_idx = 0
    for name in existing_names:
        m = _RE_REF_ENTRY_BOOKMARK.match(name or "")
        if not m:
            continue
        try:
            max_idx = max(max_idx, int(m.group("idx")))
        except (TypeError, ValueError):
            continue
    return max_idx


def _ensure_unique_bookmark_name(base: str, existing_names: set[str]) -> str:
    if base not in existing_names:
        return base
    suffix = 2
    while True:
        candidate = f"{base}_{suffix}"
        if candidate not in existing_names:
            return candidate
        suffix += 1


def _compose_ref_num_bookmark_name(entry_bookmark: str | None) -> str:
    tail = (entry_bookmark or "").strip().strip("_")
    if not tail:
        tail = "entry"
    return f"{_REF_NUM_BOOKMARK_PREFIX}{tail}"


def _is_matching_citation_pair(open_bracket: str, close_bracket: str) -> bool:
    return _CITATION_OPEN_TO_CLOSE.get(open_bracket) == close_bracket


def _text_has_citation_brackets(text: str) -> bool:
    raw = text or ""
    return any(ch in raw for ch in _CITATION_OPEN_CHARS) and any(
        ch in raw for ch in _CITATION_CLOSE_CHARS
    )


def _paragraph_run_field_depths(para) -> dict[int, int]:
    """Map each direct run element id to its surrounding field nesting depth."""
    depth = 0
    run_depths: dict[int, int] = {}
    for child in para._element:
        if child.tag != _w("r"):
            continue
        run_depths[id(child)] = depth
        fld_char = child.find(_w("fldChar"))
        if fld_char is None:
            continue
        fld_type = (fld_char.get(_w("fldCharType")) or "").strip().lower()
        if fld_type == "begin":
            depth += 1
        elif fld_type == "end":
            depth = max(0, depth - 1)
    return run_depths


def _run_is_safe_plain_text(run_el, run_depths: dict[int, int]) -> bool:
    return run_depths.get(id(run_el), 0) == 0 and _run_is_plain_text_only(run_el)


def _paragraph_has_field(para) -> bool:
    p_el = para._element
    return (
        p_el.find(f".//{_w('instrText')}") is not None
        or p_el.find(f".//{_w('fldChar')}") is not None
    )


def _paragraph_has_blocking_field(para) -> bool:
    """Only block for existing citation-like fields (EndNote / this rule)."""
    if not _paragraph_has_field(para):
        return False
    for instr in para._element.iter(_w("instrText")):
        code = (instr.text or "").upper()
        if "ADDIN EN.CITE" in code:
            return True
        if "HYPERLINK" in code and "_REFENTRY_" in code:
            return True
        if "REF " in code and ("_REFNUM_" in code or "_REFENTRY_" in code):
            return True
    return False


def _max_bookmark_id(doc: Document) -> int:
    max_id = 0
    for b in doc.element.body.iter(_w("bookmarkStart")):
        try:
            max_id = max(max_id, int(b.get(_w("id"), "0")))
        except Exception:
            continue
    return max_id


def _collect_bookmark_names(doc: Document) -> set[str]:
    names: set[str] = set()
    for b in doc.element.body.iter(_w("bookmarkStart")):
        name = b.get(_w("name"))
        if name:
            names.add(name)
    return names


def _para_has_bookmark(para, bookmark_name: str) -> bool:
    for b in para._element.iter(_w("bookmarkStart")):
        if b.get(_w("name")) == bookmark_name:
            return True
    return False


def _ensure_paragraph_bookmark(para, bookmark_name: str, bookmark_id: int) -> bool:
    """Ensure a zero-width bookmark at paragraph start. Returns True when created."""
    p_el = para._element
    if _para_has_bookmark(para, bookmark_name):
        return False

    ppr = p_el.find(_w("pPr"))
    insert_idx = list(p_el).index(ppr) + 1 if ppr is not None else 0

    b_start = etree.Element(_w("bookmarkStart"))
    b_start.set(_w("id"), str(bookmark_id))
    b_start.set(_w("name"), bookmark_name)
    b_end = etree.Element(_w("bookmarkEnd"))
    b_end.set(_w("id"), str(bookmark_id))

    p_el.insert(insert_idx, b_start)
    p_el.insert(insert_idx + 1, b_end)
    return True


def _make_text_run(text: str, rpr_src) -> etree._Element:
    r = etree.Element(_w("r"))
    if rpr_src is not None:
        r.append(deepcopy(rpr_src))
    t = etree.SubElement(r, _w("t"))
    t.set(_XML_SPACE, "preserve")
    t.text = text
    return r


def _build_citation_rpr(rpr_src) -> etree._Element:
    if rpr_src is not None:
        rpr = deepcopy(rpr_src)
    else:
        rpr = etree.Element(_w("rPr"))

    rfonts = rpr.find(_w("rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, _w("rFonts"))
    apply_explicit_rfonts(
        rpr,
        font_cn="Times New Roman",
        font_en="Times New Roman",
        font_cs="Times New Roman",
    )
    rfonts.set(_w("ascii"), "Times New Roman")
    rfonts.set(_w("hAnsi"), "Times New Roman")
    rfonts.set(_w("eastAsia"), "Times New Roman")
    rfonts.set(_w("cs"), "Times New Roman")
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        if _w(attr) in rfonts.attrib:
            rfonts.attrib.pop(_w(attr), None)

    vert = rpr.find(_w("vertAlign"))
    if vert is None:
        vert = etree.SubElement(rpr, _w("vertAlign"))
    vert.set(_w("val"), "superscript")
    return rpr


def _make_field_run(
    fld_type: str,
    rpr_src,
    *,
    dirty: bool = False,
    citation: bool = False,
) -> etree._Element:
    r = etree.Element(_w("r"))
    if citation:
        r.append(_build_citation_rpr(rpr_src))
    elif rpr_src is not None:
        r.append(deepcopy(rpr_src))
    fc = etree.SubElement(r, _w("fldChar"))
    fc.set(_w("fldCharType"), fld_type)
    if dirty:
        fc.set(_w("dirty"), "true")
    return r


def _make_instr_run(instr_text: str, rpr_src, *, citation: bool = False) -> etree._Element:
    r = etree.Element(_w("r"))
    if citation:
        r.append(_build_citation_rpr(rpr_src))
    elif rpr_src is not None:
        r.append(deepcopy(rpr_src))
    it = etree.SubElement(r, _w("instrText"))
    it.set(_XML_SPACE, "preserve")
    it.text = instr_text
    return r


def _make_citation_text_run(text: str, rpr_src) -> etree._Element:
    r = etree.Element(_w("r"))
    r.append(_build_citation_rpr(rpr_src))
    t = etree.SubElement(r, _w("t"))
    t.set(_XML_SPACE, "preserve")
    t.text = text
    return r


def _build_hyperlink_field_runs(anchor: str, display_text: str, rpr_src) -> list[etree._Element]:
    return [
        _make_field_run("begin", rpr_src, dirty=True, citation=True),
        _make_instr_run(f' HYPERLINK \\l "{anchor}" ', rpr_src, citation=True),
        _make_field_run("separate", rpr_src, citation=True),
        _make_citation_text_run(display_text, rpr_src),
        _make_field_run("end", rpr_src, citation=True),
    ]


def _para_has_reference_seq_field(para) -> bool:
    for instr in para._element.iter(_w("instrText")):
        code = (instr.text or "").upper()
        if "SEQ" in code and _REF_ENTRY_SEQ_ID.upper() in code:
            return True
    return False


def _strip_prefix_chars_from_runs(para, strip_len: int) -> None:
    if strip_len <= 0:
        return
    remaining = strip_len
    for run in para.runs:
        txt = run.text or ""
        if not txt:
            continue
        if len(txt) <= remaining:
            remaining -= len(txt)
            run.text = ""
        else:
            run.text = txt[remaining:]
            remaining = 0
        if remaining <= 0:
            break

    for run in para.runs:
        txt = run.text or ""
        if not txt:
            continue
        run.text = txt.lstrip(" \t\u3000")
        break


def _reference_prefix_rpr(para):
    for run in para.runs:
        txt = run.text or ""
        if txt:
            return run._element.find(_w("rPr"))
    return None


def _make_bookmark_start(name: str, bookmark_id: int) -> etree._Element:
    el = etree.Element(_w("bookmarkStart"))
    el.set(_w("id"), str(bookmark_id))
    el.set(_w("name"), name)
    return el


def _make_bookmark_end(bookmark_id: int) -> etree._Element:
    el = etree.Element(_w("bookmarkEnd"))
    el.set(_w("id"), str(bookmark_id))
    return el


def _insert_reference_seq_prefix(
    para,
    *,
    number_bookmark: str,
    bookmark_id: int,
    display_hint: int,
) -> None:
    p_el = para._element
    ppr = p_el.find(_w("pPr"))
    insert_idx = list(p_el).index(ppr) + 1 if ppr is not None else 0
    rpr_src = _reference_prefix_rpr(para)

    prefix_nodes: list[etree._Element] = [
        _make_text_run("[", rpr_src),
        _make_bookmark_start(number_bookmark, bookmark_id),
        _make_field_run("begin", rpr_src, dirty=True),
        _make_instr_run(f" SEQ {_REF_ENTRY_SEQ_ID} \\* ARABIC ", rpr_src),
        _make_field_run("separate", rpr_src),
        _make_text_run(str(max(1, int(display_hint))), rpr_src),
        _make_field_run("end", rpr_src),
        _make_bookmark_end(bookmark_id),
        _make_text_run("] ", rpr_src),
    ]
    for node in prefix_nodes:
        p_el.insert(insert_idx, node)
        insert_idx += 1


def _build_ref_field_runs(bookmark_name: str, display_text: str, rpr_src) -> list[etree._Element]:
    return [
        _make_field_run("begin", rpr_src, dirty=True, citation=True),
        _make_instr_run(
            f" REF {bookmark_name} \\h \\* CHARFORMAT \\* MERGEFORMAT ",
            rpr_src,
            citation=True,
        ),
        _make_field_run("separate", rpr_src, citation=True),
        _make_citation_text_run(display_text, rpr_src),
        _make_field_run("end", rpr_src, citation=True),
    ]


def _find_reference_seq_field_bounds(para) -> tuple[int, int] | None:
    """Locate the first SEQ RefEntry field run range (begin-run idx, end-run idx)."""
    children = list(para._element)
    begin_idx: int | None = None
    seq_field_hit = False
    depth = 0

    for idx, child in enumerate(children):
        if child.tag != _w("r"):
            continue

        fld_char = child.find(_w("fldChar"))
        if fld_char is not None:
            fld_type = (fld_char.get(_w("fldCharType")) or "").strip().lower()
            if fld_type == "begin":
                if depth == 0:
                    begin_idx = idx
                    seq_field_hit = False
                depth += 1
            elif fld_type == "end" and depth > 0:
                depth -= 1
                if depth == 0 and begin_idx is not None:
                    if seq_field_hit:
                        return begin_idx, idx
                    begin_idx = None
                    seq_field_hit = False

        if begin_idx is not None:
            instr = child.find(_w("instrText"))
            if instr is not None:
                code = re.sub(r"\s+", " ", (instr.text or "")).strip().upper()
                if "SEQ" in code and _REF_ENTRY_SEQ_ID.upper() in code:
                    seq_field_hit = True

    return None


def _ensure_reference_seq_bookmark(
    para,
    *,
    bookmark_name: str,
    bookmark_id: int,
) -> bool:
    """Wrap existing SEQ RefEntry field with bookmark so REF targets survive field updates."""
    if _para_has_bookmark(para, bookmark_name):
        return False
    bounds = _find_reference_seq_field_bounds(para)
    if bounds is None:
        return False
    begin_idx, end_idx = bounds
    p_el = para._element
    p_el.insert(begin_idx, _make_bookmark_start(bookmark_name, bookmark_id))
    # +2: one offset from bookmarkStart insertion, then insert after original end-run.
    p_el.insert(end_idx + 2, _make_bookmark_end(bookmark_id))
    return True


def _resolve_reference_range(doc: Document, doc_tree) -> tuple[int, int] | None:
    total = len(doc.paragraphs)
    if total <= 0:
        return None

    if doc_tree is not None:
        sec = doc_tree.get_section("references")
        if sec is not None:
            start = max(0, int(sec.start_index))
            end = min(total - 1, int(sec.end_index))
            if end >= start:
                return start, end

    start = None
    for i, para in enumerate(doc.paragraphs):
        text = re.sub(r"\s+", "", (para.text or ""))
        if text.lower() in _REFERENCE_HEADING_TOKENS:
            start = i + 1
            break
    if start is None or start >= total:
        return None
    return start, total - 1


def _resolve_body_range(doc: Document, doc_tree, reference_start: int | None) -> tuple[int, int]:
    total = len(doc.paragraphs)
    if total <= 0:
        return 0, -1
    if doc_tree is not None:
        sec = doc_tree.get_section("body")
        if sec is not None:
            start = max(0, int(sec.start_index))
            end = min(total - 1, int(sec.end_index))
            if end >= start:
                return start, end
    end = total - 1
    if isinstance(reference_start, int) and reference_start > 0:
        end = min(end, reference_start - 1)
    return 0, end


def _first_non_space_char(text: str) -> str | None:
    for ch in text or "":
        if not ch.isspace():
            return ch
    return None


def _first_non_space_char_from_runs(
    runs_snapshot: list,
    start_idx: int,
    run_depths: dict[int, int],
) -> str | None:
    for run in runs_snapshot[start_idx:]:
        if not _run_is_safe_plain_text(run._element, run_depths):
            continue
        ch = _first_non_space_char(run.text or "")
        if ch is not None:
            return ch
    return None


def _normalize_field_code(code: str) -> str:
    return re.sub(r"\s+", " ", code or "").strip().upper()


def _ensure_run_citation_rpr(run_el) -> bool:
    """Apply citation superscript rPr to a run element in-place."""
    if run_el.tag != _w("r"):
        return False
    old_rpr = run_el.find(_w("rPr"))
    new_rpr = _build_citation_rpr(old_rpr)
    if old_rpr is None:
        run_el.insert(0, new_rpr)
        return True
    if etree.tostring(old_rpr) == etree.tostring(new_rpr):
        return False
    run_el.replace(old_rpr, new_rpr)
    return True


def _normalize_existing_citation_ref_fields(para) -> int:
    """Repair legacy REF citation fields (format switches + superscript styling)."""
    children = list(para._element)
    changed_fields = 0

    for idx, child in enumerate(children):
        if child.tag != _w("r"):
            continue
        instr = child.find(_w("instrText"))
        if instr is None:
            continue

        raw_code = instr.text or ""
        code_upper = _normalize_field_code(raw_code)
        if "REF " not in code_upper:
            continue
        if "_REFNUM_" not in code_upper and "_REFENTRY_" not in code_upper:
            continue

        target_match = _RE_REF_FIELD_TARGET.search(raw_code)
        if not target_match:
            continue
        target = (target_match.group("target") or "").strip()
        if not target:
            continue

        desired_code = f" REF {target} \\h \\* CHARFORMAT \\* MERGEFORMAT "
        field_changed = False
        if _normalize_field_code(raw_code) != _normalize_field_code(desired_code):
            instr.text = desired_code
            field_changed = True

        # Normalize field run formatting from begin ... end so current display also becomes superscript.
        begin_idx = idx
        while begin_idx >= 0:
            run_el = children[begin_idx]
            if run_el.tag == _w("r"):
                fld = run_el.find(_w("fldChar"))
                if fld is not None and (fld.get(_w("fldCharType")) or "").strip().lower() == "begin":
                    break
            begin_idx -= 1

        end_idx = idx
        while end_idx < len(children):
            run_el = children[end_idx]
            if run_el.tag == _w("r"):
                fld = run_el.find(_w("fldChar"))
                if fld is not None and (fld.get(_w("fldCharType")) or "").strip().lower() == "end":
                    break
            end_idx += 1

        if 0 <= begin_idx <= idx and idx <= end_idx < len(children):
            for run_idx in range(begin_idx, end_idx + 1):
                if _ensure_run_citation_rpr(children[run_idx]):
                    field_changed = True
        else:
            if _ensure_run_citation_rpr(child):
                field_changed = True

        if field_changed:
            changed_fields += 1

    return changed_fields


def _is_probable_year_token(compact_pages: str) -> bool:
    if not compact_pages.isdigit() or len(compact_pages) != 4:
        return False
    year = int(compact_pages)
    return 1900 <= year <= 2099


def _consume_outer_page_span(text: str, start_idx: int, next_char_hint: str | None) -> int:
    """Return end index for confident outer-page marker, else start_idx."""
    if start_idx >= len(text):
        return start_idx
    if not text[start_idx].isdigit():
        return start_idx

    tail = text[start_idx:]
    m = _RE_OUTER_PAGE_CANDIDATE.match(tail)
    if not m:
        return start_idx
    pages = m.group("pages") or ""
    if not pages:
        return start_idx

    end_idx = start_idx + m.end("pages")
    follow = _first_non_space_char(text[end_idx:]) or next_char_hint
    if follow is not None:
        if follow.isalnum():
            return start_idx
        if follow not in _OUTER_PAGE_TERMINATORS:
            return start_idx

    compact = re.sub(r"\s+", "", pages)
    if _is_probable_year_token(compact):
        return start_idx
    return end_idx


def _strip_reference_label_prefix(para) -> bool:
    cut_len = _extract_reference_label_end(para.text or "")
    if cut_len is None:
        return False
    _strip_prefix_chars_from_runs(para, cut_len)
    return True


def _run_is_plain_text_only(run_el) -> bool:
    if run_el is None or run_el.tag != _w("r"):
        return False
    for child in list(run_el):
        tag = child.tag
        if tag in {_w("rPr"), _w("t")}:
            continue
        return False
    return True


def _coalesce_split_bracket_runs(para) -> None:
    """Merge minimal safe run ranges so split citations become one plain-text run."""
    run_depths = _paragraph_run_field_depths(para)
    runs = list(para.runs)
    i = 0
    while i < len(runs):
        start_run = runs[i]
        if not _run_is_safe_plain_text(start_run._element, run_depths):
            i += 1
            continue
        start_text = start_run.text or ""
        if not any(ch in start_text for ch in _CITATION_OPEN_CHARS):
            i += 1
            continue
        if any(ch in start_text for ch in _CITATION_CLOSE_CHARS):
            i += 1
            continue

        j = i + 1
        while j < len(runs):
            if not _run_is_safe_plain_text(runs[j]._element, run_depths):
                break
            end_text = runs[j].text or ""
            if not any(ch in end_text for ch in _CITATION_CLOSE_CHARS):
                j += 1
                continue

            segment = runs[i:j + 1]
            merged_text = "".join(r.text or "" for r in segment)
            if not _RE_CITATION_TOKEN.search(merged_text):
                j += 1
                continue
            start_run.text = merged_text
            for k in range(j, i, -1):
                run_el = runs[k]._element
                parent = run_el.getparent()
                if parent is not None:
                    parent.remove(run_el)
            runs = list(para.runs)
            break

        i += 1


def _replace_citation_tokens_in_paragraph(
    para,
    num_to_target: dict[int, str],
    *,
    superscript_outer_page_numbers: bool,
    use_ref_field: bool,
) -> tuple[int, set[int]]:
    linked = 0
    unresolved: set[int] = set()
    if not para.runs:
        return linked, unresolved

    _coalesce_split_bracket_runs(para)
    run_depths = _paragraph_run_field_depths(para)
    runs_snapshot = list(para.runs)
    for run_idx, run in enumerate(runs_snapshot):
        if not _run_is_safe_plain_text(run._element, run_depths):
            continue
        run_text = run.text or ""
        if not _text_has_citation_brackets(run_text):
            continue
        matches = list(_RE_CITATION_TOKEN.finditer(run_text))
        if not matches:
            continue

        parent = run._element.getparent()
        if parent is None:
            continue
        insert_at = parent.index(run._element)
        rpr_src = run._element.find(_w("rPr"))
        cursor = 0
        next_char_hint = _first_non_space_char_from_runs(
            runs_snapshot,
            run_idx + 1,
            run_depths,
        )

        for m in matches:
            start, end = m.span()
            if start > cursor:
                parent.insert(insert_at, _make_text_run(run_text[cursor:start], rpr_src))
                insert_at += 1

            token = m.group(0)
            inner = m.group("inner") or ""
            open_bracket = m.group("open") or ""
            close_bracket = m.group("close") or ""
            if (
                not _is_matching_citation_pair(open_bracket, close_bracket)
                or not _RE_NUMERIC_CITATION_INNER.fullmatch(inner)
                or not _RE_NUM_TOKEN.search(inner)
            ):
                parent.insert(insert_at, _make_text_run(token, rpr_src))
                insert_at += 1
                cursor = end
                continue

            parent.insert(insert_at, _make_citation_text_run(open_bracket, rpr_src))
            insert_at += 1

            inner_cursor = 0
            for num_m in _RE_NUM_TOKEN.finditer(inner):
                nstart, nend = num_m.span()
                if nstart > inner_cursor:
                    sep_text = inner[inner_cursor:nstart]
                    parent.insert(insert_at, _make_citation_text_run(sep_text, rpr_src))
                    insert_at += 1
                raw_num = num_m.group(0)
                num = int(raw_num)
                target = num_to_target.get(num)
                if target:
                    if use_ref_field:
                        field_runs = _build_ref_field_runs(target, raw_num, rpr_src)
                    else:
                        field_runs = _build_hyperlink_field_runs(target, raw_num, rpr_src)
                    for field_run in field_runs:
                        parent.insert(insert_at, field_run)
                        insert_at += 1
                    linked += 1
                else:
                    parent.insert(insert_at, _make_citation_text_run(raw_num, rpr_src))
                    insert_at += 1
                    unresolved.add(num)
                inner_cursor = nend

            if inner_cursor < len(inner):
                tail_text = inner[inner_cursor:]
                parent.insert(insert_at, _make_citation_text_run(tail_text, rpr_src))
                insert_at += 1

            parent.insert(insert_at, _make_citation_text_run(close_bracket, rpr_src))
            insert_at += 1

            outer_end = end
            if superscript_outer_page_numbers:
                outer_end = _consume_outer_page_span(run_text, end, next_char_hint)
                if outer_end > end:
                    parent.insert(
                        insert_at,
                        _make_citation_text_run(run_text[end:outer_end], rpr_src),
                    )
                    insert_at += 1
            cursor = outer_end

        if cursor < len(run_text):
            parent.insert(insert_at, _make_text_run(run_text[cursor:], rpr_src))
            insert_at += 1

        parent.remove(run._element)

    return linked, unresolved


def _build_reference_targets(
    doc: Document,
    *,
    ref_start: int,
    ref_end: int,
    existing_names: set[str],
    next_bookmark_id: int,
    auto_number_reference_entries: bool,
) -> tuple[dict[int, str], int, int, int, set[int], int]:
    """Build mapping from reference number -> target bookmark.

    Returns:
      (num_to_target, next_bookmark_id, entry_bookmarks_added,
       number_fields_inserted, duplicated_numbers, reference_entries_count)
    """
    num_to_target: dict[int, str] = {}
    duplicated_numbers: set[int] = set()
    entry_bookmarks_added = 0
    number_fields_inserted = 0
    reference_entries_count = 0
    next_entry_serial = _next_ref_entry_serial(existing_names)

    for i in range(ref_start, ref_end + 1):
        if i < 0 or i >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[i]
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_reference_heading_line(text):
            continue
        if _paragraph_has_blocking_field(para):
            # Existing citation field in references entry is treated as locked content.
            continue

        parsed_num = _extract_reference_number(text)
        entry_bookmark = _find_ref_entry_bookmark_name(para)
        has_seq = _para_has_reference_seq_field(para)
        if parsed_num is None and not entry_bookmark and not has_seq:
            # Only [1]-style labels (including fullwidth square brackets) are
            # eligible as bibliography entries. Skip list items like "(1) 标题".
            continue

        reference_entries_count += 1

        if not entry_bookmark:
            next_entry_serial += 1
            entry_bookmark = _ensure_unique_bookmark_name(
                f"{_REF_ENTRY_BOOKMARK_PREFIX}{next_entry_serial}",
                existing_names,
            )
            if _ensure_paragraph_bookmark(para, entry_bookmark, next_bookmark_id):
                entry_bookmarks_added += 1
                next_bookmark_id += 1
            existing_names.add(entry_bookmark)

        target_bookmark = entry_bookmark
        if auto_number_reference_entries:
            number_bookmark = _find_ref_num_bookmark_name(para, entry_bookmark)
            if not number_bookmark:
                number_bookmark = _ensure_unique_bookmark_name(
                    _compose_ref_num_bookmark_name(entry_bookmark),
                    existing_names,
                )
                if not has_seq:
                    _strip_reference_label_prefix(para)
                    _insert_reference_seq_prefix(
                        para,
                        number_bookmark=number_bookmark,
                        bookmark_id=next_bookmark_id,
                        display_hint=parsed_num or reference_entries_count,
                    )
                    number_fields_inserted += 1
                    next_bookmark_id += 1
                    existing_names.add(number_bookmark)
                else:
                    # Existing SEQ field: re-wrap it with a stable bookmark if missing.
                    wrapped = _ensure_reference_seq_bookmark(
                        para,
                        bookmark_name=number_bookmark,
                        bookmark_id=next_bookmark_id,
                    )
                    if wrapped:
                        next_bookmark_id += 1
                        existing_names.add(number_bookmark)
                    else:
                        # Cannot safely locate SEQ field bounds, avoid broken REF target.
                        number_bookmark = None
            target_bookmark = number_bookmark if number_bookmark else ""

        if parsed_num is None:
            continue
        if not target_bookmark:
            continue
        if parsed_num in duplicated_numbers:
            continue
        if parsed_num in num_to_target:
            duplicated_numbers.add(parsed_num)
            num_to_target.pop(parsed_num, None)
            continue
        num_to_target[parsed_num] = target_bookmark

    return (
        num_to_target,
        next_bookmark_id,
        entry_bookmarks_added,
        number_fields_inserted,
        duplicated_numbers,
        reference_entries_count,
    )


class CitationLinkRule(BaseRule):
    name = "citation_link"
    description = "正文引用与参考文献域关联"

    def apply(
        self,
        doc: Document,
        config: SceneConfig,
        tracker: ChangeTracker,
        context: dict,
    ) -> None:
        citation_cfg = getattr(config, "citation_link", None)
        if citation_cfg is not None and not bool(getattr(citation_cfg, "enabled", True)):
            return
        superscript_outer_pages = bool(
            getattr(citation_cfg, "superscript_outer_page_numbers", False)
        )
        auto_number_reference_entries = bool(
            getattr(citation_cfg, "auto_number_reference_entries", True)
        )

        doc_tree = context.get("doc_tree")
        ref_range = _resolve_reference_range(doc, doc_tree)
        if ref_range is None:
            tracker.record(
                rule_name=self.name,
                target="references",
                section="global",
                change_type="field",
                before="references section not detected",
                after="skipped",
                paragraph_index=-1,
            )
            return

        ref_start, ref_end = ref_range
        next_bookmark_id = _max_bookmark_id(doc) + 1
        existing_names = _collect_bookmark_names(doc)

        (
            num_to_target,
            next_bookmark_id,
            entry_bookmarks_added,
            number_fields_inserted,
            duplicated_ref_numbers,
            reference_entries_count,
        ) = _build_reference_targets(
            doc,
            ref_start=ref_start,
            ref_end=ref_end,
            existing_names=existing_names,
            next_bookmark_id=next_bookmark_id,
            auto_number_reference_entries=auto_number_reference_entries,
        )

        if not num_to_target and number_fields_inserted <= 0 and entry_bookmarks_added <= 0:
            tracker.record(
                rule_name=self.name,
                target="references",
                section="references",
                change_type="field",
                before="no numbered reference entry matched",
                after="skipped",
                paragraph_index=-1,
            )
            return

        body_start, body_end = _resolve_body_range(doc, doc_tree, ref_start)
        linked_total = 0
        unresolved_numbers: set[int] = set()
        skipped_field_paras = 0
        repaired_field_paras = 0

        if num_to_target:
            for i in range(body_start, body_end + 1):
                if i < 0 or i >= len(doc.paragraphs):
                    continue
                para = doc.paragraphs[i]
                if not (para.text or "").strip():
                    continue
                has_blocking_field = _paragraph_has_blocking_field(para)
                if has_blocking_field:
                    normalized_fields = _normalize_existing_citation_ref_fields(para)
                    if normalized_fields > 0:
                        repaired_field_paras += 1
                        tracker.record(
                            rule_name=self.name,
                            target=f"paragraph #{i}",
                            section="body",
                            change_type="field",
                            before="existing citation field needs normalization",
                            after=f"normalized existing citation fields ({normalized_fields})",
                            paragraph_index=i,
                        )
                    skipped_field_paras += 1
                if not _text_has_citation_brackets(para.text or ""):
                    continue

                linked, unresolved = _replace_citation_tokens_in_paragraph(
                    para,
                    num_to_target,
                    superscript_outer_page_numbers=superscript_outer_pages,
                    use_ref_field=auto_number_reference_entries,
                )
                if linked > 0:
                    linked_total += linked
                    tracker.record(
                        rule_name=self.name,
                        target=f"paragraph #{i}",
                        section="body",
                        change_type="field",
                        before="citation token is plain text",
                        after=f"converted to linked field ({linked})",
                        paragraph_index=i,
                    )
                if unresolved:
                    unresolved_numbers.update(unresolved)

        tracker.record(
            rule_name=self.name,
            target="summary",
            section="global",
            change_type="field",
            before=(
                f"plain citations={linked_total}, references={len(num_to_target)}, "
                f"reference_entries={reference_entries_count}, "
                f"entry_bookmarks_added={entry_bookmarks_added}, "
                f"reference_seq_inserted={number_fields_inserted}, "
                f"auto_numbering={'on' if auto_number_reference_entries else 'off'}, "
                f"outer_page_superscript="
                f"{'on' if superscript_outer_pages else 'off'}"
            ),
            after=(
                f"linked={linked_total}, skipped_field_paragraphs={skipped_field_paras}, "
                f"repaired_field_paragraphs={repaired_field_paras}, "
                f"unresolved={len(unresolved_numbers)}"
            ),
            paragraph_index=-1,
        )

        if duplicated_ref_numbers:
            dup_text = ",".join(str(x) for x in sorted(duplicated_ref_numbers)[:12])
            if len(duplicated_ref_numbers) > 12:
                dup_text += ", ..."
            tracker.record(
                rule_name=self.name,
                target="references",
                section="references",
                change_type="field",
                before=f"duplicated numbering detected: {dup_text}",
                after="duplicate references left unlinked to avoid ambiguity",
                paragraph_index=ref_start,
            )

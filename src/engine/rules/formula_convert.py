"""Formula recognition and conversion rule."""

from __future__ import annotations

from copy import deepcopy

from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.rules import table_format as table_helpers
from src.engine.rules.base import BaseRule
from src.formula_core.ast import FormulaNode
from src.formula_core.convert import ConversionOutcome, convert_formula_node
from src.formula_core.normalize import normalize_formula_node
from src.formula_core.parse import parse_document_formulas, _looks_like_non_math_latex_text
from src.formula_core.runtime import FormulaRuleStats
from src.scene.schema import SceneConfig

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_CONFIDENCE_THRESHOLD = 0.60
_FORMULA_DIAGNOSTIC_MESSAGES = {
    "mathtype_binary_unparsed": "MathType 二进制未解析",
    "missing_equation_native_stream": "缺少 Equation Native 数据流",
    "mathtype_decode_failed": "MathType OLE 解码失败",
    "ole_relationship_missing": "OLE 关系缺失",
    "ole_target_missing": "OLE 内嵌对象缺失",
    "mathtype_ole_dependency_missing": "缺少 olefile 依赖",
    "ole_binary_unparsed": "OLE 对象未解析",
}


def _replace_paragraph_with_omml(paragraph, omml_element) -> bool:
    p = paragraph._p
    if p is None:
        return False
    ppr_tag = f"{{{_W_NS}}}pPr"
    run_tag = f"{{{_W_NS}}}r"
    for child in list(p):
        if child.tag != ppr_tag:
            p.remove(child)

    omml_copy = deepcopy(omml_element)
    local = omml_copy.tag.split("}")[-1] if "}" in omml_copy.tag else omml_copy.tag
    if local == "oMathPara":
        p.append(omml_copy)
        return True

    run = p.makeelement(run_tag)
    run.append(omml_copy)
    p.append(run)
    return True


def _paragraph_table_element(paragraph):
    node = getattr(paragraph, "_p", None)
    while node is not None:
        parent = node.getparent()
        if parent is None:
            return None
        tag = parent.tag.split("}")[-1] if "}" in parent.tag else parent.tag
        if tag == "tbl":
            return parent
        node = parent
    return None


def _source_end_index(occ) -> int:
    indices = [
        int(idx)
        for idx in list(getattr(occ, "source_paragraph_indices", []) or [])
        if isinstance(idx, int) and idx >= 0
    ]
    if indices:
        return max(indices)
    try:
        return int(getattr(occ, "paragraph_index", -1))
    except (TypeError, ValueError):
        return -1


def _remove_consumed_source_paragraphs(doc: Document, occ) -> None:
    anchor_idx = getattr(occ, "paragraph_index", -1)
    indices = sorted(
        {
            int(idx)
            for idx in list(getattr(occ, "source_paragraph_indices", []) or [])
            if isinstance(idx, int) and idx >= 0 and idx != anchor_idx
        },
        reverse=True,
    )
    for idx in indices:
        if idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        parent = para._p.getparent()
        if parent is not None:
            parent.remove(para._p)


def _remove_non_math_latex_wrapper_paragraphs(doc: Document) -> int:
    removed = 0
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[idx]
        text = str(getattr(para, "text", "") or "").strip()
        if not _looks_like_non_math_latex_text(text):
            continue
        parent = para._p.getparent()
        if parent is None:
            continue
        parent.remove(para._p)
        removed += 1
    return removed


def _dedupe(values: list[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for v in values:
        item = str(v or "").strip()
        if not item or item in seen:
            continue
        seen.add(item)
        out.append(item)
    return out


def _contains_escape_placeholder_noise(text: str) -> bool:
    value = str(text or "")
    return "ESC_" in value or "__ESC" in value


def _guard_untrusted_conversion(
        occ,
        normalized_node: FormulaNode,
        editable_outcome: ConversionOutcome,
        exchange_latex: str) -> str | None:
    warnings = [
        str(item or "").strip()
        for item in (
            list(getattr(normalized_node, "warnings", []) or [])
            + list(editable_outcome.warnings or [])
        )
    ]
    recovered_escape = (
        "escape_placeholder_command_recovered" in warnings
        or "escape_placeholder_formula_extracted" in warnings
        or "rendered_segment_extracted" in warnings
        or "command_context_extracted" in warnings
        or "big_operator_context_extracted" in warnings
    )
    if editable_outcome.reason == "fallback_literal_conversion":
        return "fallback_literal_conversion"
    if "fallback_literal_conversion" in warnings:
        return "fallback_literal_conversion"

    payload = normalized_node.payload if isinstance(normalized_node.payload, dict) else {}
    normalized_texts = [
        payload.get("latex", ""),
        payload.get("normalized_latex", ""),
        exchange_latex,
    ]
    if not recovered_escape:
        normalized_texts.append(payload.get("text", ""))
    if any(_contains_escape_placeholder_noise(text) for text in normalized_texts):
        return "escape_placeholder_noise"
    if _contains_escape_placeholder_noise(getattr(occ, "source_text", "")) and not recovered_escape:
        return "escape_placeholder_noise"
    return None


def _formula_diagnostic_code(node: FormulaNode) -> str | None:
    payload = node.payload if isinstance(node.payload, dict) else {}
    candidates = [payload.get("diagnostic_code")]
    candidates.extend(list(getattr(node, "warnings", []) or []))
    for item in candidates:
        code = str(item or "").strip()
        if not code:
            continue
        if code in _FORMULA_DIAGNOSTIC_MESSAGES:
            return code
        if code.startswith("unsupported_mtef_version"):
            return code
    return None


def _formula_diagnostic_message(code: str) -> str:
    value = str(code or "").strip()
    if value.startswith("unsupported_mtef_version"):
        _, _, version = value.partition(":")
        version = version.strip() or "unknown"
        return f"不支持的 MTEF 版本 ({version})"
    return _FORMULA_DIAGNOSTIC_MESSAGES.get(value, value or "公式诊断失败")


def _convert_to_editable_outcome(
        node: FormulaNode,
        output_mode: str,
        *,
        block: bool) -> tuple[ConversionOutcome, str]:
    """Return an outcome that is always intended for editable OMML replacement.

    - output_mode=word_native: direct conversion to OMML.
    - output_mode=latex: first normalize to LaTeX expression, then round-trip to OMML.
      This keeps LaTeX as the conversion logic while final document remains editable.
    """
    if output_mode == "word_native":
        return convert_formula_node(node, "word_native", block=block), ""

    source = str(getattr(node, "source_type", "") or "").strip().lower()
    if source == "word_native":
        # Safety guard: keep existing editable OMML untouched in latex logic mode.
        # Round-tripping complex native formulas can lose trailing terms.
        return ConversionOutcome(
            success=True,
            target_mode="word_native",
            confidence=max(0.85, float(getattr(node, "confidence", 0.0) or 0.0)),
            reason="already_word_native_latex_mode_noop",
            warnings=_dedupe(list(getattr(node, "warnings", []) or [])),
            transformed=False,
        ), ""

    latex_out = convert_formula_node(node, "latex", block=block)
    if not latex_out.success:
        return latex_out, ""

    latex_expr = str(latex_out.latex_text or "").strip()
    if not latex_expr:
        return ConversionOutcome(
            success=False,
            target_mode="word_native",
            confidence=max(0.0, float(latex_out.confidence)),
            reason="empty_latex_exchange",
            warnings=_dedupe(list(latex_out.warnings or []) + ["empty_latex_exchange"]),
        ), ""

    # Round-trip through LaTeX expression and finally output editable OMML.
    roundtrip_node = FormulaNode(
        kind=node.kind,
        payload={"latex": latex_expr},
        children=[],
        source_type="latex",
        confidence=latex_out.confidence,
        warnings=list(latex_out.warnings or []),
    )
    word_out = convert_formula_node(roundtrip_node, "word_native", block=block)
    if not word_out.success:
        return ConversionOutcome(
            success=False,
            target_mode="word_native",
            confidence=min(float(latex_out.confidence), float(word_out.confidence)),
            reason=f"latex_roundtrip_failed:{word_out.reason or 'unknown'}",
            warnings=_dedupe(
                list(latex_out.warnings or [])
                + list(word_out.warnings or [])
                + ["latex_roundtrip_failed"]
            ),
        ), latex_expr

    word_out.confidence = min(float(latex_out.confidence), float(word_out.confidence))
    word_out.reason = "latex_roundtrip_to_word_native"
    word_out.warnings = _dedupe(list(latex_out.warnings or []) + list(word_out.warnings or []))
    return word_out, latex_expr


class FormulaConvertRule(BaseRule):
    name = "formula_convert"
    description = "Formula recognition and conversion"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        cfg = getattr(config, "formula_convert", None)
        if cfg is not None and not bool(getattr(cfg, "enabled", False)):
            return

        output_mode = str(getattr(cfg, "output_mode", "word_native")).strip().lower()
        if output_mode not in {"word_native", "latex"}:
            output_mode = "word_native"
        low_policy = str(getattr(cfg, "low_confidence_policy", "skip_and_mark")).strip().lower()

        runtime = context.setdefault("formula_runtime", {})
        parse_result = parse_document_formulas(doc)
        runtime["parse_result"] = parse_result
        runtime["occurrences"] = parse_result.occurrences
        runtime["convert_enabled"] = True
        runtime["output_mode"] = output_mode
        if output_mode == "latex":
            runtime["latex_logic_editable_output"] = True
        low_conf_list = runtime.setdefault("low_confidence", [])
        diagnostic_list = runtime.setdefault("diagnostics", [])
        latex_exchange = runtime.setdefault("latex_exchange", [])
        rule_stats = FormulaRuleStats()

        occurrences = list(parse_result.occurrences)
        occurrences.sort(key=_source_end_index, reverse=True)
        saw_latex_source = any(getattr(occ, "source_type", "") == "latex" for occ in occurrences)

        for occ in occurrences:
            rule_stats.matched += 1
            normalized_node = normalize_formula_node(occ.node)
            diagnostic_code = _formula_diagnostic_code(normalized_node)
            if diagnostic_code is None:
                diagnostic_code = _formula_diagnostic_code(occ.node)
            if diagnostic_code is not None:
                rule_stats.skipped_unsupported += 1
                diagnostic_message = _formula_diagnostic_message(diagnostic_code)
                diagnostic_list.append({
                    "rule": self.name,
                    "location": occ.location,
                    "paragraph_index": occ.paragraph_index,
                    "source_type": occ.source_type,
                    "reason_code": diagnostic_code,
                    "reason": diagnostic_message,
                })
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="diagnostic",
                    before=f"{occ.source_type} -> {output_mode}",
                    after=f"诊断: {diagnostic_message}",
                    paragraph_index=occ.paragraph_index,
                )
                continue
            editable_outcome, exchange_latex = _convert_to_editable_outcome(
                normalized_node,
                output_mode,
                block=occ.is_block,
            )
            guard_reason = _guard_untrusted_conversion(
                occ,
                normalized_node,
                editable_outcome,
                exchange_latex,
            )

            confidence_for_stats = (
                editable_outcome.confidence if editable_outcome.success else normalized_node.confidence
            )
            if guard_reason:
                confidence_for_stats = min(float(confidence_for_stats), 0.45)
            rule_stats.note_confidence(confidence_for_stats)
            is_low_conf = confidence_for_stats < _CONFIDENCE_THRESHOLD

            if guard_reason:
                if is_low_conf and low_policy == "skip_and_mark":
                    rule_stats.skipped_low_confidence += 1
                    low_conf_list.append({
                        "rule": self.name,
                        "location": occ.location,
                        "paragraph_index": occ.paragraph_index,
                        "source_type": occ.source_type,
                        "target_mode": output_mode,
                        "confidence": round(confidence_for_stats, 3),
                        "reason": guard_reason,
                        "suggestion": "建议人工确认公式后处理",
                    })
                    tracker.record(
                        rule_name=self.name,
                        target=occ.location,
                        section="formula",
                        change_type="skip",
                        before=f"{occ.source_type} -> {output_mode}",
                        after=f"低置信跳过 ({confidence_for_stats:.2f})",
                        paragraph_index=occ.paragraph_index,
                    )
                else:
                    rule_stats.skipped_unsupported += 1
                    tracker.record(
                        rule_name=self.name,
                        target=occ.location,
                        section="formula",
                        change_type="skip",
                        before=f"{occ.source_type} -> {output_mode}",
                        after=f"跳过: {guard_reason}",
                        paragraph_index=occ.paragraph_index,
                    )
                continue

            if not editable_outcome.success:
                reason = editable_outcome.reason or "unsupported_conversion"
                if is_low_conf and low_policy == "skip_and_mark":
                    rule_stats.skipped_low_confidence += 1
                    low_conf_list.append({
                        "rule": self.name,
                        "location": occ.location,
                        "paragraph_index": occ.paragraph_index,
                        "source_type": occ.source_type,
                        "target_mode": output_mode,
                        "confidence": round(confidence_for_stats, 3),
                        "reason": reason,
                        "suggestion": "建议人工确认公式后处理",
                    })
                    tracker.record(
                        rule_name=self.name,
                        target=occ.location,
                        section="formula",
                        change_type="skip",
                        before=f"{occ.source_type} -> {output_mode}",
                        after=f"低置信跳过 ({confidence_for_stats:.2f})",
                        paragraph_index=occ.paragraph_index,
                    )
                else:
                    rule_stats.skipped_unsupported += 1
                    tracker.record(
                        rule_name=self.name,
                        target=occ.location,
                        section="formula",
                        change_type="skip",
                        before=f"{occ.source_type} -> {output_mode}",
                        after=f"跳过: {reason}",
                        paragraph_index=occ.paragraph_index,
                    )
                continue

            if not editable_outcome.transformed:
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="skip",
                    before=f"{occ.source_type} -> {output_mode}",
                    after="无需转换",
                    paragraph_index=occ.paragraph_index,
                )
                continue

            if is_low_conf and low_policy == "skip_and_mark":
                rule_stats.skipped_low_confidence += 1
                low_conf_list.append({
                    "rule": self.name,
                    "location": occ.location,
                    "paragraph_index": occ.paragraph_index,
                    "source_type": occ.source_type,
                    "target_mode": output_mode,
                    "confidence": round(confidence_for_stats, 3),
                    "reason": editable_outcome.reason or "low_confidence",
                    "suggestion": "建议人工确认公式后处理",
                })
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="skip",
                    before=f"{occ.source_type} -> {output_mode}",
                    after=f"低置信跳过 ({confidence_for_stats:.2f})",
                    paragraph_index=occ.paragraph_index,
                )
                continue

            if not occ.is_formula_only:
                rule_stats.skipped_dependency += 1
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="skip",
                    before=f"{occ.source_type} -> {output_mode}",
                    after="MVP 仅改写纯公式段落",
                    paragraph_index=occ.paragraph_index,
                )
                continue

            if occ.in_table:
                table_el = _paragraph_table_element(occ.paragraph)
                if table_el is None or not table_helpers._is_equation_table(table_el):
                    rule_stats.skipped_dependency += 1
                    tracker.record(
                        rule_name=self.name,
                        target=occ.location,
                        section="formula",
                        change_type="skip",
                        before=f"{occ.source_type} -> {output_mode}",
                        after="MVP 暂不改写普通表格内公式",
                        paragraph_index=occ.paragraph_index,
                    )
                    continue

            applied = False
            if editable_outcome.omml_element is not None:
                applied = _replace_paragraph_with_omml(occ.paragraph, editable_outcome.omml_element)

            if not applied:
                rule_stats.errors += 1
                tracker.record(
                    rule_name=self.name,
                    target=occ.location,
                    section="formula",
                    change_type="text",
                    before=f"{occ.source_type} -> {output_mode}",
                    after="改写失败",
                    paragraph_index=occ.paragraph_index,
                    success=False,
                    failure_reason="unable to apply editable formula result",
                )
                continue

            _remove_consumed_source_paragraphs(doc, occ)

            if output_mode == "latex" and exchange_latex:
                latex_exchange.append({
                    "location": occ.location,
                    "paragraph_index": occ.paragraph_index,
                    "latex": exchange_latex,
                })

            rule_stats.converted += 1
            tracker.record(
                rule_name=self.name,
                target=occ.location,
                section="formula",
                change_type="text",
                before=f"{occ.source_type} -> {output_mode}",
                after=(
                    "已按 LaTeX 逻辑转换为可编辑公式"
                    if output_mode == "latex"
                    else "已转换"
                ),
                paragraph_index=occ.paragraph_index,
            )

        removed_wrappers = 0
        if saw_latex_source and rule_stats.converted > 0:
            removed_wrappers = _remove_non_math_latex_wrapper_paragraphs(doc)
            if removed_wrappers > 0:
                tracker.record(
                    rule_name=self.name,
                    target="latex_document_wrappers",
                    section="formula",
                    change_type="text",
                    before=f"{removed_wrappers} wrapper paragraphs",
                    after="已移除非数学 LaTeX 包装行",
                    paragraph_index=-1,
                )

        runtime.setdefault("stats", {})[self.name] = {
            "matched": rule_stats.matched,
            "converted": rule_stats.converted,
            "skipped_low_confidence": rule_stats.skipped_low_confidence,
            "skipped_unsupported": rule_stats.skipped_unsupported,
            "skipped_dependency": rule_stats.skipped_dependency,
            "errors": rule_stats.errors,
            "removed_wrappers": removed_wrappers,
        }
        tracker.record(
            rule_name=self.name,
            target="summary",
            section="formula",
            change_type="format",
            before="统计",
            after=rule_stats.to_summary(),
            paragraph_index=-1,
        )
        tracker.record(
            rule_name=self.name,
            target="confidence_confirmation",
            section="formula",
            change_type="format",
            before="自动置信分层确认",
            after=(
                f"已在执行日志末尾标注：{rule_stats.confidence_summary()}；"
                f"无需人工交互确认，低置信按策略跳过并标注。"
            ),
            paragraph_index=-1,
        )

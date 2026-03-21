"""目录内容重建与格式化规则：清除旧 TOC、根据标题重建条目、包装域代码、应用格式"""

import re
from copy import deepcopy
from types import SimpleNamespace
from lxml import etree
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from src.engine.rules.base import BaseRule
from src.engine.change_tracker import ChangeTracker
from src.scene.heading_model import (
    detect_level_by_style_name,
    get_front_matter_title_norms,
    get_level_to_word_style,
    get_non_numbered_heading_style_name,
)
from src.scene.schema import SceneConfig, StyleConfig
from src.engine.doc_tree import DocTree
from src.utils.heading_numbering_template import (
    default_heading_numbering_template,
    render_heading_numbering,
)
from src.utils.heading_numbering_v2 import (
    advance_heading_counters,
    included_toc_levels,
    merged_level_binding,
)
from src.utils.indent import apply_style_config_indents
from src.utils.line_spacing import apply_line_spacing, sync_spacing_ooxml
from src.utils.ooxml import apply_explicit_rfonts
from src.utils.toc_entry import (
    looks_like_toc_entry_line,
    looks_like_date_placeholder_line,
    looks_like_reference_entry_line,
    is_toc_level_style_name,
    toc_whitelist_score,
    toc_blacklist_score,
    looks_like_numbered_toc_entry_with_page_suffix,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TOC_BOOKMARK_PREFIX = "_TocRef_"
_TOC_MODE_WORD_NATIVE = "word_native"
_TOC_MODE_PLAIN = "plain"
_FRONT_MATTER_HEADING_STYLE_NAME = "Front Matter Heading Unnumbered"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


def _resolve_toc_mode(config: SceneConfig) -> str:
    toc_cfg = getattr(config, "toc", None)
    raw_mode = str(getattr(toc_cfg, "mode", _TOC_MODE_WORD_NATIVE) or "").strip().lower()
    if raw_mode in {"plain", "legacy"}:
        return _TOC_MODE_PLAIN
    if raw_mode in {"word_native", "native", "auto", "field"}:
        return _TOC_MODE_WORD_NATIVE
    return _TOC_MODE_WORD_NATIVE


def _configured_toc_levels(config: SceneConfig) -> list[str]:
    levels = included_toc_levels(config.heading_numbering_v2.level_bindings)
    return levels or ["heading1", "heading2", "heading3"]


def _native_toc_field_instruction(config: SceneConfig) -> str:
    """Build TOC field instruction using outline-level switch (\\o).

    Built-in heading levels are collected via ``\\o``/``\\u`` to stay resilient
    to WPS / locale / case differences. Custom unnumbered heading styles used by
    front/back matter are appended via ``\\t`` so Word refresh keeps those
    entries in native TOC.

    Relying on ``\\t`` alone is brittle for built-in heading styles because style
    names can vary in case across office suites.
    """
    levels = _configured_toc_levels(config)
    if not levels:
        max_level = 3
    else:
        max_level = max(int(lv[7:]) for lv in levels)
    custom_styles = [
        _FRONT_MATTER_HEADING_STYLE_NAME,
        get_non_numbered_heading_style_name(config),
    ]
    seen: set[str] = set()
    style_pairs: list[str] = []
    for style_name in custom_styles:
        style_text = str(style_name or "").strip()
        if not style_text:
            continue
        lowered = style_text.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        style_pairs.append(f"{style_text},1")

    style_switch = ""
    if style_pairs:
        style_switch = f' \\t "{",".join(style_pairs)}"'
    return f' TOC \\o "1-{max_level}" \\h \\z \\u{style_switch} '


# ── 编号文本生成 ──

# 层级 → TOC 样式 key 映射
LEVEL_TO_TOC_STYLE = {
    "heading1": "toc_chapter",
    "heading2": "toc_level1",
    "heading3": "toc_level2",
}

# 层级 → Word TOC 样式名
LEVEL_TO_WORD_STYLE = {
    f"heading{level_idx}": f"TOC {level_idx}"
    for level_idx in range(1, 9)
}


def _toc_style_key_for_level(level: str) -> str:
    level_idx = int(str(level or "heading1")[7:]) if str(level or "").startswith("heading") else 1
    if level_idx <= 1:
        return "toc_chapter"
    if level_idx == 2:
        return "toc_level1"
    return "toc_level2"

_TOC_HEADING_STYLE_CANDIDATES = ["TOC Heading", "目录标题"]

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_ABSTRACT_CN_TITLES = {
    "摘要",
    "摘要。",
    "摘要：",
    "中文摘要",
    "内容摘要",
}

_RE_ABSTRACT_EN_TITLE = re.compile(r"^abstract([:：.]|\s|$)", re.IGNORECASE)
_RE_TOC_TITLE_CN = re.compile(r"^目[录錄]$")
_RE_TOC_TITLE_EN = re.compile(r"^(contents|tableofcontents)$", re.IGNORECASE)

_RE_CHAPTER_CN = re.compile(r"^第[一二三四五六七八九十百零\d]+[章节篇]")
_RE_TOC_CHAPTER_CN = re.compile(r"^第[一二三四五六七八九十百零〇两\d]+(?:章|篇)")
_RE_TOC_SECTION_CN = re.compile(r"^第[一二三四五六七八九十百零〇两\d]+节")
_RE_LEVEL2 = re.compile(r"^\d+\.\d+\.\d+(?:[\.、．])?\s*\S")
_RE_LEVEL1 = re.compile(r"^\d+\.\d+(?:[\.、．])?\s*\S")
_RE_CHAPTER_NUM = re.compile(r"^\d+(?:[\.、．]|\s+)\S")
_TOC_FALLBACK_MIN_CONF = 2
_TOC_GUARD_ENTRY_CONF = 2
_RE_PLAIN_TOC_PAGE_SUFFIX = re.compile(
    r"^(?P<label>.+?)(?:\t+|[.．。·…]{2,}|\s{1,})(?P<page>\d+|[IVXLCDMivxlcdm]+|[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]+)\s*$"
)
_RE_APPENDIX_TITLE = re.compile(
    r"^(?:附录|appendix)\s*[A-Za-zＡ-Ｚａ-ｚ0-9一二三四五六七八九十]*"
    r"(?:[\s\u3000:：、.．\-—_（）()]*\S.*)?$",
    re.IGNORECASE,
)

_BACK_MATTER_SECTION_TITLES = {
    "references": "参考文献",
    "errata": "勘误页",
    "appendix": "附录",
    "acknowledgment": "致谢",
    "resume": "个人简历",
}
_BACK_MATTER_RESUME_SUBTITLE = "在学期间发表的学术论文与研究成果"


def _strip_heading_number(text: str) -> str:
    """从标题文本中剥离编号前缀，只保留标题内容。

    heading_numbering 模式B已剥离编号，但模式A保留了编号。
    此函数兼容两种情况。
    """
    from src.engine.rules.heading_numbering import split_heading_text
    _, _, title = split_heading_text(text)
    return title.strip() if title.strip() else text.strip()


def _norm_no_space(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip()


def _looks_like_abstract_cn_title(text: str) -> bool:
    norm = _norm_no_space(text)
    if norm in _ABSTRACT_CN_TITLES:
        return True
    return norm.endswith("摘要") and len(norm) <= 6


def _looks_like_abstract_en_title(text: str) -> bool:
    return bool(_RE_ABSTRACT_EN_TITLE.match((text or "").strip()))


def _looks_like_appendix_title(text: str) -> bool:
    raw = (text or "").strip()
    if not raw or len(raw) > 80:
        return False
    if looks_like_toc_entry_line(raw):
        return False
    if looks_like_reference_entry_line(raw):
        return False
    if raw.endswith(("。", "；", ";", "，", ",")):
        return False
    return bool(_RE_APPENDIX_TITLE.match(raw))


def _looks_like_numbered_heading_text(text: str) -> bool:
    return bool(
        _RE_CHAPTER_CN.match(text)
        or _RE_LEVEL1.match(text)
        or _RE_LEVEL2.match(text)
        or _RE_CHAPTER_NUM.match(text)
    )


def _is_toc_entry_like_para(para, text: str, style_name: str, style_id: str) -> bool:
    conf = _toc_entry_confidence(para, text, style_name, style_id)
    return conf >= _TOC_GUARD_ENTRY_CONF


def _toc_entry_confidence(para, text: str, style_name: str, style_id: str) -> int:
    numbered_heading_like = _looks_like_numbered_heading_text(text)
    white = toc_whitelist_score(
        text,
        style_name=style_name,
        style_id=style_id,
        has_pageref=_para_has_pageref_field(para),
        numbered_heading_like=numbered_heading_like,
    )
    black = toc_blacklist_score(text)

    # Additional local blacklist hints.
    if looks_like_reference_entry_line(text):
        black += 1
    if looks_like_date_placeholder_line(text):
        black += 1

    return white - black


def _infer_toc_level_from_text(text: str) -> str:
    raw = (text or "").strip()
    # 兼容中文“第X章/第X节”样式，避免无法识别时全部回落到一级目录。
    if _RE_TOC_SECTION_CN.match(raw):
        return "heading2"
    if _RE_TOC_CHAPTER_CN.match(raw):
        return "heading1"
    if _RE_LEVEL2.match(raw):
        return "heading3"
    if _RE_LEVEL1.match(raw):
        return "heading2"
    return "heading1"


def _dedupe_entries(entries: list[dict]) -> list[dict]:
    deduped = []
    seen_idx = set()
    for entry in entries:
        idx = entry.get("para_index")
        if idx is None:
            continue
        if idx in seen_idx:
            continue
        seen_idx.add(idx)
        deduped.append(entry)
    return deduped


def _style_to_level(config: SceneConfig, style_name: str) -> str | None:
    return detect_level_by_style_name(config, style_name)


def _fallback_collect_headings(doc: Document, config: SceneConfig, toc_section=None):
    """当 heading_detect 未给出结果时，回退扫描全篇标题。"""
    toc_start = toc_section.start_index if toc_section else -1
    toc_end = toc_section.end_index if toc_section else -1
    front_title_norms = get_front_matter_title_norms(config)
    found = []
    for i, para in enumerate(doc.paragraphs):
        if toc_start <= i <= toc_end:
            continue
        text = (para.text or "").strip()
        if not text or len(text) > 200:
            continue
        if looks_like_toc_entry_line(text):
            continue
        if looks_like_numbered_toc_entry_with_page_suffix(text):
            continue
        if looks_like_reference_entry_line(text):
            continue
        if looks_like_date_placeholder_line(text):
            continue
        if _norm_no_space(text).lower() in front_title_norms:
            continue

        style_name = para.style.name if para.style else ""
        level = _style_to_level(config, style_name)

        if level is None:
            if _RE_LEVEL2.match(text):
                level = "heading3"
            elif _RE_LEVEL1.match(text):
                level = "heading2"
            elif _RE_CHAPTER_CN.match(text) or _RE_CHAPTER_NUM.match(text):
                level = "heading1"

        if level:
            found.append(SimpleNamespace(para_index=i, level=level, text=text))
    return found


def _extract_heading_title_runs(
    doc: Document,
    heading,
    fallback_title: str,
    *,
    prefix_len: int = 0,
) -> list[dict]:
    """提取标题正文 run 片段，尽量保留下/上角标信息。"""
    para_index = getattr(heading, "para_index", -1)
    if para_index < 0 or para_index >= len(doc.paragraphs):
        return [{"text": fallback_title, "vert_align": None}] if fallback_title else []

    para = doc.paragraphs[para_index]
    if not para.runs:
        return [{"text": fallback_title, "vert_align": None}] if fallback_title else []

    strip_len = max(0, int(prefix_len or 0))
    if strip_len <= 0:
        from src.engine.rules.heading_numbering import split_heading_text

        full_text = para.text or ""
        number_part, sep_part, _ = split_heading_text(full_text)
        strip_len = len(number_part) + len(sep_part) if number_part else 0

    segments = []
    remaining = strip_len
    is_first_non_empty = True
    for run in para.runs:
        txt = run.text or ""
        if not txt:
            continue
        if remaining >= len(txt):
            remaining -= len(txt)
            continue
        if remaining > 0:
            txt = txt[remaining:]
            remaining = 0

        if is_first_non_empty:
            txt = txt.lstrip(" \t\u3000")
        if not txt:
            continue

        is_first_non_empty = False
        vert_align = None
        if run.font.subscript:
            vert_align = "subscript"
        elif run.font.superscript:
            vert_align = "superscript"
        else:
            rpr = run._element.find(_w("rPr"))
            if rpr is not None:
                va = rpr.find(_w("vertAlign"))
                if va is not None:
                    val = (va.get(_w("val")) or "").strip()
                    if val in ("subscript", "superscript"):
                        vert_align = val

        segments.append({"text": txt, "vert_align": vert_align})

    if not segments and fallback_title:
        return [{"text": fallback_title, "vert_align": None}]
    return segments


def _configured_heading_prefix_len(text: str, numbering: str) -> int:
    raw = str(text or "")
    if not raw or not numbering:
        return 0
    trimmed = raw.lstrip(" \t\u3000")
    if not trimmed.startswith(numbering):
        return 0
    leading_len = len(raw) - len(trimmed)
    suffix = trimmed[len(numbering):]
    ws_suffix = len(suffix) - len(suffix.lstrip(" \t\u3000"))
    return leading_len + len(numbering) + ws_suffix


def _toc_entry_separator(separator: str | None) -> str:
    raw = str(separator or "")
    if "\t" not in raw:
        return raw
    return raw.replace("\t", "   ")


def _build_toc_entries(doc: Document, headings, levels_config: dict,
                       level_bindings: dict | None = None,
                       separator: str = "\u3000") -> list[dict]:
    """根据标题列表和编号配置，生成 TOC 条目列表。

    返回 [{"level": str, "numbering": str, "title": str, "sep": str}, ...]
    """
    entries = []
    counters = {f"heading{idx}": 0 for idx in range(1, 9)}
    # 层级计数器
    for h in headings:
        level = h.level

        if level not in LEVEL_TO_WORD_STYLE:
            continue

        counters = advance_heading_counters(level, counters, level_bindings)
        if not merged_level_binding(level, level_bindings).include_in_toc:
            continue

        numbering = _format_numbering(level, counters, levels_config)

        # 获取该层级的分隔符
        lc = levels_config.get(level)
        sep = _toc_entry_separator(lc.effective_separator if lc else separator)
        prefix_len = _configured_heading_prefix_len(h.text, numbering)
        if prefix_len > 0:
            title = (h.text or "")[prefix_len:].strip()
        else:
            title = _strip_heading_number(h.text)
        current_para_text = ""
        para_index = getattr(h, "para_index", -1)
        if 0 <= para_index < len(doc.paragraphs):
            current_para_text = doc.paragraphs[para_index].text or ""
        run_prefix_len = _configured_heading_prefix_len(current_para_text, numbering)

        title_runs = _extract_heading_title_runs(
            doc,
            h,
            title,
            prefix_len=run_prefix_len,
        )
        title_runs_text = "".join(str(part.get("text", "") or "") for part in title_runs).strip()
        if title and title_runs_text and _norm_no_space(title_runs_text) != _norm_no_space(title):
            title_runs = [{"text": title, "vert_align": None}]

        entries.append({
            "level": level,
            "numbering": numbering,
            "title": title,
            "title_runs": title_runs,
            "sep": sep,
            "para_index": h.para_index,
        })

    return entries


def _build_front_matter_toc_entries(doc: Document, doc_tree: DocTree) -> list[dict]:
    """构建前置分区目录条目（摘要 / Abstract）。"""
    entries = []
    seen_idx = set()

    def _push(level, numbering, title, sep, para_index):
        if para_index is None or para_index < 0 or para_index >= len(doc.paragraphs):
            return
        if para_index in seen_idx:
            return
        seen_idx.add(para_index)
        entries.append({
            "level": level,
            "numbering": numbering,
            "title": title,
            "sep": sep,
            "para_index": para_index,
        })

    for sec_type, fallback in (("abstract_cn", "摘要"), ("abstract_en", "Abstract")):
        sec = doc_tree.get_section(sec_type)
        if sec:
            para_index = sec.start_index
            txt = doc.paragraphs[para_index].text.strip()
            if looks_like_toc_entry_line(txt):
                continue
            title = txt or fallback
            _push("heading1", "", title, "", para_index)
            continue

        # 回退：按文本锚点搜索（摘要/Abstract），限定前文区域避免正文误命中
        scan_limit = max(10, int(len(doc.paragraphs) * 0.45))
        for i, para in enumerate(doc.paragraphs[:scan_limit]):
            txt = (para.text or "").strip()
            if looks_like_toc_entry_line(txt):
                continue
            if sec_type == "abstract_cn" and _looks_like_abstract_cn_title(txt):
                _push("heading1", "", "摘要", "", i)
                break
            if sec_type == "abstract_en" and _looks_like_abstract_en_title(txt):
                _push("heading1", "", "Abstract", "", i)
                break
    return entries


def _build_back_matter_toc_entries(doc: Document, doc_tree: DocTree) -> list[dict]:
    """构建后置分区目录条目（参考文献/勘误页/致谢/个人简历/成果标题）。"""
    entries = []
    seen_idx = set()

    def _push(para_index: int, title: str):
        if para_index is None or para_index < 0 or para_index >= len(doc.paragraphs):
            return
        if para_index in seen_idx:
            return
        seen_idx.add(para_index)
        entries.append({
            "level": "heading1",
            "numbering": "",
            "title": title,
            "sep": "",
            "para_index": para_index,
        })

    for sec_type, fallback_title in _BACK_MATTER_SECTION_TITLES.items():
        sec = doc_tree.get_section(sec_type) if doc_tree else None
        if sec:
            para_index = sec.start_index
            title_text = (doc.paragraphs[para_index].text or "").strip()
            if looks_like_toc_entry_line(title_text):
                continue
            title = title_text or fallback_title
            _push(para_index, title)

            # 简历分区内“成果标题”也按一级目录条目收录。
            if sec_type == "resume":
                end = min(sec.end_index, len(doc.paragraphs) - 1)
                for i in range(sec.start_index + 1, end + 1):
                    txt = (doc.paragraphs[i].text or "").strip()
                    if not txt:
                        continue
                    if looks_like_toc_entry_line(txt):
                        continue
                    norm = _norm_no_space(txt)
                    if norm.startswith(_BACK_MATTER_RESUME_SUBTITLE):
                        _push(i, txt)
                        break
            elif sec_type == "appendix":
                # Capture all appendix titles (e.g. 附录A/附录B), even when
                # doc_tree only anchors one appendix boundary.
                for i, para in enumerate(doc.paragraphs):
                    txt = (para.text or "").strip()
                    if _looks_like_appendix_title(txt):
                        _push(i, txt)
            continue

        # doc_tree 未识别时文本回退，避免后置标题漏收。
        fallback_norm = _norm_no_space(fallback_title)
        if sec_type == "appendix":
            for i, para in enumerate(doc.paragraphs):
                txt = (para.text or "").strip()
                if _looks_like_appendix_title(txt):
                    _push(i, txt)
            continue
        for i, para in enumerate(doc.paragraphs):
            txt = (para.text or "").strip()
            if not txt:
                continue
            if looks_like_toc_entry_line(txt):
                continue
            norm = _norm_no_space(txt)
            if norm == fallback_norm:
                _push(i, txt)
                break

    # “成果标题”独立回退
    for i, para in enumerate(doc.paragraphs):
        txt = (para.text or "").strip()
        if not txt:
            continue
        if looks_like_toc_entry_line(txt):
            continue
        norm = _norm_no_space(txt)
        if norm.startswith(_BACK_MATTER_RESUME_SUBTITLE):
            _push(i, txt)
            break

    return entries


def _format_numbering(level: str,
                      counters: dict[str, int],
                      levels_config: dict) -> str:
    lc = levels_config.get(level)
    template = lc.template if lc else default_heading_numbering_template(level)
    return render_heading_numbering(level, template, levels_config, counters)


# ── OOXML 段落构建 ──

def _make_run(parent, text=None, *, vert_align: str | None = None):
    """创建 w:r 元素"""
    r = etree.SubElement(parent, _w("r"))
    if text is not None:
        t = etree.SubElement(r, _w("t"))
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    if vert_align in ("subscript", "superscript"):
        rpr = r.find(_w("rPr"))
        if rpr is None:
            rpr = etree.SubElement(r, _w("rPr"))
        va = rpr.find(_w("vertAlign"))
        if va is None:
            va = etree.SubElement(rpr, _w("vertAlign"))
        va.set(_w("val"), vert_align)
    return r


def _make_tab_run(parent):
    """创建包含 tab 字符的 run"""
    r = etree.SubElement(parent, _w("r"))
    etree.SubElement(r, _w("tab"))
    return r


def _make_fld_char(parent, fld_type: str, *, dirty: bool = False):
    """创建 fldChar run（begin/separate/end）。"""
    r = etree.SubElement(parent, _w("r"))
    fc = etree.SubElement(r, _w("fldChar"))
    fc.set(_w("fldCharType"), fld_type)
    if dirty:
        fc.set(_w("dirty"), "true")
    return r


def _make_instr_text(parent, instr: str):
    """创建 instrText run"""
    r = etree.SubElement(parent, _w("r"))
    it = etree.SubElement(r, _w("instrText"))
    it.text = instr
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _build_toc_title_para(title_text: str = "目录") -> etree._Element:
    """构建 TOC 标题段落元素"""
    p = OxmlElement('w:p')
    ppr = etree.SubElement(p, _w("pPr"))
    jc = etree.SubElement(ppr, _w("jc"))
    jc.set(_w("val"), "center")
    outline = etree.SubElement(ppr, _w("outlineLvl"))
    outline.set(_w("val"), "9")
    _make_run(p, text=title_text)
    return p


def _build_toc_field_begin_para(field_instruction: str) -> etree._Element:
    """构建 TOC 域代码起始段落（begin + instrText + separate）"""
    p = OxmlElement('w:p')
    # dirty=true 提示 Word 打开时刷新目录域
    _make_fld_char(p, "begin", dirty=True)
    _make_instr_text(p, field_instruction)
    _make_fld_char(p, "separate")
    return p


def _build_toc_field_end_para() -> etree._Element:
    """构建 TOC 域代码结束段落（end）"""
    p = OxmlElement('w:p')
    _make_fld_char(p, "end")
    return p


def _max_bookmark_id(doc: Document) -> int:
    """获取文档中当前最大 bookmark id。"""
    max_id = 0
    for b in doc.element.body.iter(_w("bookmarkStart")):
        try:
            max_id = max(max_id, int(b.get(_w("id"), "0")))
        except Exception:
            continue
    return max_id


def _bookmark_name_for_para(para_index: int) -> str:
    return f"{_TOC_BOOKMARK_PREFIX}{para_index}"


def _purge_generated_toc_bookmarks(doc: Document) -> int:
    """Remove stale auto-generated TOC bookmarks to avoid wrong PAGEREF targets."""
    to_remove_ids: set[str] = set()
    removed = 0

    for b_start in list(doc.element.body.iter(_w("bookmarkStart"))):
        name = (b_start.get(_w("name")) or "").strip()
        if not name.startswith(_TOC_BOOKMARK_PREFIX):
            continue
        bid = (b_start.get(_w("id")) or "").strip()
        if bid:
            to_remove_ids.add(bid)
        parent = b_start.getparent()
        if parent is not None:
            parent.remove(b_start)
            removed += 1

    if to_remove_ids:
        for b_end in list(doc.element.body.iter(_w("bookmarkEnd"))):
            bid = (b_end.get(_w("id")) or "").strip()
            if bid not in to_remove_ids:
                continue
            parent = b_end.getparent()
            if parent is not None:
                parent.remove(b_end)
                removed += 1

    return removed


def _ensure_bookmark_for_para(doc: Document, para_index: int,
                              bookmark_name: str, bookmark_id: int) -> bool:
    """确保指定段落存在某书签；新建返回 True，已存在返回 False。"""
    if para_index < 0 or para_index >= len(doc.paragraphs):
        return False
    para_el = doc.paragraphs[para_index]._element

    # 已存在同名书签
    for b in para_el.iter(_w("bookmarkStart")):
        if b.get(_w("name")) == bookmark_name:
            return False

    ppr = para_el.find(_w("pPr"))
    insert_idx = list(para_el).index(ppr) + 1 if ppr is not None else 0

    b_start = etree.Element(_w("bookmarkStart"))
    b_start.set(_w("id"), str(bookmark_id))
    b_start.set(_w("name"), bookmark_name)
    b_end = etree.Element(_w("bookmarkEnd"))
    b_end.set(_w("id"), str(bookmark_id))

    para_el.insert(insert_idx, b_start)
    para_el.insert(insert_idx + 1, b_end)
    return True


def _attach_entry_bookmarks(doc: Document, entries: list[dict]) -> None:
    """给目录条目目标段落绑定书签，供 PAGEREF 取页码。"""
    _purge_generated_toc_bookmarks(doc)
    next_id = _max_bookmark_id(doc) + 1
    # 先收集全局已存在书签，避免重复
    existing_names = set()
    for b in doc.element.body.iter(_w("bookmarkStart")):
        name = b.get(_w("name"))
        if name:
            existing_names.add(name)

    for entry in entries:
        para_index = entry.get("para_index")
        if para_index is None:
            continue
        bname = _bookmark_name_for_para(para_index)
        entry["bookmark"] = bname
        if bname in existing_names:
            continue
        created = _ensure_bookmark_for_para(doc, para_index, bname, next_id)
        if created:
            existing_names.add(bname)
            next_id += 1


def _make_pageref_field(parent, bookmark_name: str):
    """创建 PAGEREF 域（页码字段），用于目录条目页码自动更新。"""
    _make_fld_char(parent, "begin", dirty=True)
    _make_instr_text(parent, f" PAGEREF {bookmark_name} \\h ")
    _make_fld_char(parent, "separate")
    _make_run(parent, text="1")  # 占位页码，更新域后刷新
    _make_fld_char(parent, "end")


def _make_hyperlink_field(parent, bookmark_name: str, parts: list[dict]) -> None:
    """Create an in-document hyperlink field for TOC labels."""
    _make_fld_char(parent, "begin", dirty=True)
    _make_instr_text(parent, f' HYPERLINK \\l "{bookmark_name}" ')
    _make_fld_char(parent, "separate")
    for part in parts:
        txt = part.get("text", "")
        if not txt:
            continue
        _make_run(parent, text=txt, vert_align=part.get("vert_align"))
    _make_fld_char(parent, "end")


def _build_toc_entry_para(entry: dict, tab_pos_twips: int,
                          word_style: str) -> etree._Element:
    """Build one TOC entry paragraph: hyperlink label + tab + page field."""
    p = OxmlElement('w:p')

    # Paragraph properties: right-aligned dot leader tab stop.
    ppr = etree.SubElement(p, _w("pPr"))
    style_id = re.sub(r"\s+", "", word_style or "")
    if style_id:
        p_style = etree.SubElement(ppr, _w("pStyle"))
        p_style.set(_w("val"), style_id)
    tabs = etree.SubElement(ppr, _w("tabs"))
    tab = etree.SubElement(tabs, _w("tab"))
    tab.set(_w("val"), "right")
    tab.set(_w("leader"), "dot")
    tab.set(_w("pos"), str(tab_pos_twips))

    label_parts: list[dict] = []
    numbering = entry.get("numbering", "")
    sep = entry.get("sep", "")
    if numbering:
        label_parts.append({"text": f"{numbering}{sep}", "vert_align": None})

    title_runs = entry.get("title_runs") or []
    if title_runs:
        for part in title_runs:
            txt = part.get("text", "")
            if not txt:
                continue
            label_parts.append({
                "text": txt,
                "vert_align": part.get("vert_align"),
            })
    else:
        title = entry.get("title", "")
        if title:
            label_parts.append({"text": title, "vert_align": None})

    if not label_parts:
        label_parts.append({"text": "", "vert_align": None})

    bname = entry.get("bookmark")
    if bname:
        _make_hyperlink_field(p, bname, label_parts)
    else:
        for part in label_parts:
            _make_run(p, text=part.get("text", ""), vert_align=part.get("vert_align"))

    # Append tab + PAGEREF page-number field.
    _make_tab_run(p)
    if bname:
        _make_pageref_field(p, bname)
    else:
        _make_run(p, text="1")  # Fallback placeholder.

    return p


def _format_toc_para(para, sc: StyleConfig, is_title: bool = False):
    """对 TOC 段落应用格式（字体、字号、行距、间距、缩进）"""
    pf = para.paragraph_format

    # 对齐：优先使用配置，缺省时标题居中/条目左对齐。
    if is_title:
        ppr = para._element.find(_w("pPr"))
        if ppr is None:
            ppr = etree.SubElement(para._element, _w("pPr"))
            para._element.insert(0, ppr)
        numpr = ppr.find(_w("numPr"))
        if numpr is not None:
            ppr.remove(numpr)
        outline = ppr.find(_w("outlineLvl"))
        if outline is None:
            outline = etree.SubElement(ppr, _w("outlineLvl"))
        outline.set(_w("val"), "9")

    align_key = str(getattr(sc, "alignment", "") or "").strip().lower()
    if align_key in ALIGNMENT_MAP:
        pf.alignment = ALIGNMENT_MAP[align_key]
    elif is_title:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 间距
    pf.space_before = Pt(sc.space_before_pt)
    pf.space_after = Pt(sc.space_after_pt)
    apply_line_spacing(pf, sc.line_spacing_type, sc.line_spacing_pt)
    sync_spacing_ooxml(
        para._element,
        space_before_pt=sc.space_before_pt,
        space_after_pt=sc.space_after_pt,
        line_spacing_type=sc.line_spacing_type,
        line_spacing_value=sc.line_spacing_pt,
    )

    # 缩进：按配置字符数应用（与正文样式一致的换算规则）。
    apply_style_config_indents(pf, para._element, sc)

    for run in para.runs:
        run.font.name = sc.font_en
        run.font.size = Pt(sc.size_pt)
        run.font.bold = bool(sc.bold)
        run.font.italic = bool(sc.italic)
        rpr = run._element.find(_w("rPr"))
        if rpr is None:
            rpr = etree.SubElement(run._element, _w("rPr"))
        rfonts = rpr.find(_w("rFonts"))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, _w("rFonts"))
        apply_explicit_rfonts(
            rpr,
            font_cn=sc.font_cn,
            font_en=sc.font_en,
        )
        rfonts.set(_w("eastAsia"), sc.font_cn)
        rfonts.set(_w("ascii"), sc.font_en)
        rfonts.set(_w("hAnsi"), sc.font_en)
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            rfonts.attrib.pop(_w(attr), None)


def _toc_style_candidates(level: str) -> list[str]:
    """返回某 TOC 层级可能存在的样式名候选（中英/大小写）。"""
    level_idx = max(1, min(8, int(level[7:]) if str(level).startswith("heading") else 1))
    return [f"TOC {level_idx}", f"toc {level_idx}", f"目录 {level_idx}"]


def _find_toc_style_by_level(doc: Document, level: str):
    """按层级查找已有 TOC 段落样式。"""
    for style_name in _toc_style_candidates(level):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            return style

    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        style_name = style.name or ""
        style_id = getattr(style, "style_id", "") or ""
        if (
            _toc_level_from_style_name(style_name) == level
            or _toc_level_from_style_name(style_id) == level
        ):
            return style
    return None


def _ensure_toc_style_for_level(doc: Document, level: str):
    """确保 TOC 层级样式存在，必要时创建标准样式名。"""
    style = _find_toc_style_by_level(doc, level)
    if style is not None:
        return style

    style_name = LEVEL_TO_WORD_STYLE.get(level, "TOC 1")
    try:
        return doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        # 样式名可能已存在于隐藏/延迟加载集合中，再次按名称取一次。
        try:
            style = doc.styles[style_name]
        except KeyError:
            return None
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            return style
        return None


def _is_toc_heading_style_name(style_name: str) -> bool:
    norm = _norm_no_space(style_name).lower()
    return norm in {"tocheading", "目录标题"}


def _find_toc_heading_style(doc: Document):
    for name in _TOC_HEADING_STYLE_CANDIDATES:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            return style

    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        style_name = style.name or ""
        style_id = getattr(style, "style_id", "") or ""
        if _is_toc_heading_style_name(style_name) or _is_toc_heading_style_name(style_id):
            return style
    return None


def _ensure_toc_heading_style(doc: Document):
    style = _find_toc_heading_style(doc)
    if style is not None:
        return style
    try:
        return doc.styles.add_style("TOC Heading", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        try:
            style = doc.styles["TOC Heading"]
        except KeyError:
            return None
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            return style
        return None


def _is_toc_level_style_name(style_name: str) -> bool:
    """判断是否为 TOC 条目样式名（TOC 1/2/3 或 目录 1/2/3）。"""
    return is_toc_level_style_name(style_name)


def _toc_level_from_style_name(style_name: str) -> str | None:
    """从 TOC 样式名推导层级键。"""
    s = (style_name or "").strip()
    # 仅识别 TOC 样式，避免把 "Heading 1/2/3" 误判成目录条目
    if not _is_toc_level_style_name(s):
        return None
    m = re.search(r'(\d+)\s*$', s)
    if not m:
        return None
    n = int(m.group(1))
    if 1 <= n <= 8:
        return f"heading{n}"
    return None


def _para_has_pageref_field(para) -> bool:
    """检测段落中是否包含 PAGEREF 域代码（目录页码项）。"""
    for instr in para._element.iter(_w("instrText")):
        if "PAGEREF" in (instr.text or "").upper():
            return True
    return False


def _normalize_toc_style(style, sc: StyleConfig | None = None) -> None:
    """将 TOC 样式定义同步到配置，避免 Word 更新域后格式回弹。"""
    sc = sc or StyleConfig()
    pf = style.paragraph_format

    align_key = str(getattr(sc, "alignment", "") or "").strip().lower()
    if align_key in ALIGNMENT_MAP:
        pf.alignment = ALIGNMENT_MAP[align_key]
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    apply_style_config_indents(pf, style.element, sc)
    pf.space_before = Pt(sc.space_before_pt)
    pf.space_after = Pt(sc.space_after_pt)
    apply_line_spacing(pf, sc.line_spacing_type, sc.line_spacing_pt)
    sync_spacing_ooxml(
        style.element,
        space_before_pt=sc.space_before_pt,
        space_after_pt=sc.space_after_pt,
        line_spacing_type=sc.line_spacing_type,
        line_spacing_value=sc.line_spacing_pt,
    )

    style.font.name = sc.font_en
    style.font.size = Pt(sc.size_pt)
    style.font.bold = bool(sc.bold)
    style.font.italic = bool(sc.italic)

    rpr = style.element.find(_w("rPr"))
    if rpr is None:
        rpr = etree.SubElement(style.element, _w("rPr"))
    rfonts = rpr.find(_w("rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, _w("rFonts"))
    apply_explicit_rfonts(
        rpr,
        font_cn=sc.font_cn,
        font_en=sc.font_en,
    )
    rfonts.set(_w("eastAsia"), sc.font_cn)
    rfonts.set(_w("ascii"), sc.font_en)
    rfonts.set(_w("hAnsi"), sc.font_en)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(_w(attr), None)


def _normalize_toc_heading_style(doc: Document, style, sc: StyleConfig | None = None) -> None:
    """Prevent TOC title style from inheriting heading numbering / outline level."""
    try:
        style.base_style = doc.styles["Normal"]
    except Exception:
        pass

    _normalize_toc_style(style, sc)

    ppr = style.element.find(_w("pPr"))
    if ppr is None:
        ppr = etree.SubElement(style.element, _w("pPr"))
    numpr = ppr.find(_w("numPr"))
    if numpr is not None:
        ppr.remove(numpr)
    outline = ppr.find(_w("outlineLvl"))
    if outline is None:
        outline = etree.SubElement(ppr, _w("outlineLvl"))
    outline.set(_w("val"), "9")


def _normalize_toc_styles_in_doc(doc: Document, config: SceneConfig) -> int:
    """统一修正文档内 TOC 样式定义，返回修正样式数。"""
    count = 0
    normalized_style_ids: set[str] = set()
    toc_levels = _configured_toc_levels(config)
    level_styles: dict[str, list] = {lv: [] for lv in toc_levels}

    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        style_name = style.name or ""
        style_id = getattr(style, "style_id", "") or ""
        level = _toc_level_from_style_name(style_name) or _toc_level_from_style_name(style_id)
        if level in level_styles:
            level_styles[level].append(style)

    for level in toc_levels:
        style_key = _toc_style_key_for_level(level)
        styles = level_styles.get(level) or []
        if not styles:
            ensured = _ensure_toc_style_for_level(doc, level)
            if ensured is not None:
                styles = [ensured]

        sc = config.styles.get(style_key) if style_key else None
        for style in styles:
            style_id = getattr(style, "style_id", "") or style.name or f"id({id(style)})"
            if style_id in normalized_style_ids:
                continue
            _normalize_toc_style(style, sc)
            normalized_style_ids.add(style_id)
            count += 1

    heading_style = _ensure_toc_heading_style(doc)
    if heading_style is not None:
        heading_style_id = (
            getattr(heading_style, "style_id", "")
            or heading_style.name
            or f"id({id(heading_style)})"
        )
        if heading_style_id not in normalized_style_ids:
            _normalize_toc_heading_style(doc, heading_style, config.styles.get("toc_title") or StyleConfig())
            normalized_style_ids.add(heading_style_id)
            count += 1

    return count


def _cm_to_twips(cm: float) -> int:
    return int(cm * 567)


def _enable_update_fields(doc):
    """在文档设置中启用 updateFields，使 Word 打开时自动更新域（含 TOC 页码）。

    对应 XML: <w:updateFields w:val="true"/> in word/settings.xml
    """
    settings_el = doc.settings.element
    existing = settings_el.find(_w("updateFields"))
    if existing is not None:
        existing.set(_w("val"), "true")
    else:
        uf = etree.SubElement(settings_el, _w("updateFields"))
        uf.set(_w("val"), "true")


def _has_page_break(elements: list) -> bool:
    """检测段落元素列表中是否包含分页符。

    检查两种常见形式：
    1. w:br type="page"（手动分页符）
    2. w:pPr/w:pageBreakBefore（段前分页）
    """
    for el in elements:
        # 检查 w:br type="page"
        for br in el.iter(_w("br")):
            if br.get(_w("type")) == "page":
                return True
        # 检查 pageBreakBefore
        ppr = el.find(_w("pPr"))
        if ppr is not None:
            pbb = ppr.find(_w("pageBreakBefore"))
            if pbb is not None:
                return True
    return False


def _extract_last_toc_section_break(elements: list):
    """Capture trailing sectPr inside old TOC block so rebuild won't drop section boundaries."""
    for el in reversed(elements):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag != "p":
            continue
        ppr = el.find(_w("pPr"))
        if ppr is None:
            continue
        sect = ppr.find(_w("sectPr"))
        if sect is not None:
            return deepcopy(sect)
    return None


def _apply_section_break_to_para(para_el, sect_pr) -> None:
    """Attach preserved sectPr to a rebuilt TOC paragraph."""
    if para_el is None or sect_pr is None:
        return
    ppr = para_el.find(_w("pPr"))
    if ppr is None:
        ppr = etree.SubElement(para_el, _w("pPr"))
        para_el.insert(0, ppr)
    old = ppr.find(_w("sectPr"))
    if old is not None:
        ppr.remove(old)
    ppr.append(deepcopy(sect_pr))


def _build_page_break_para() -> etree._Element:
    """构建包含分页符的空段落"""
    p = OxmlElement('w:p')
    r = etree.SubElement(p, _w("r"))
    br = etree.SubElement(r, _w("br"))
    br.set(_w("type"), "page")
    return p


def _split_plain_toc_entry_text(text: str) -> tuple[str, str] | None:
    raw = (text or "").strip()
    if not raw:
        return None
    m = _RE_PLAIN_TOC_PAGE_SUFFIX.match(raw)
    if not m:
        return None
    label = (m.group("label") or "").rstrip(" \t.．。·…")
    page = (m.group("page") or "").strip()
    if not label or not page:
        return None
    return label, page


def _ensure_right_tab_stop(para, tab_pos_twips: int) -> None:
    ppr = para._element.find(_w("pPr"))
    if ppr is None:
        ppr = etree.SubElement(para._element, _w("pPr"))
        para._element.insert(0, ppr)
    tabs = ppr.find(_w("tabs"))
    if tabs is None:
        tabs = etree.SubElement(ppr, _w("tabs"))
    else:
        for tab in list(tabs.findall(_w("tab"))):
            if (tab.get(_w("val")) or "").lower() == "right":
                tabs.remove(tab)
    tab = etree.SubElement(tabs, _w("tab"))
    tab.set(_w("val"), "right")
    tab.set(_w("leader"), "dot")
    tab.set(_w("pos"), str(tab_pos_twips))


def _rewrite_plain_toc_para(para, label: str, page: str, tab_pos_twips: int) -> None:
    p = para._element
    for child in list(p):
        if child.tag == _w("pPr"):
            continue
        p.remove(child)
    _ensure_right_tab_stop(para, tab_pos_twips)
    _make_run(p, text=label)
    _make_tab_run(p)
    _make_run(p, text=page)


class TocFormatRule(BaseRule):
    name = "toc_format"
    description = "目录内容重建与格式化"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        # 无论是否识别到 TOC 分区，都启用域更新并修正 TOC 样式
        _enable_update_fields(doc)
        _normalize_toc_styles_in_doc(doc, config)

        doc_tree: DocTree = context.get("doc_tree")
        if not doc_tree:
            return

        toc_section = doc_tree.get_section("toc")
        if not toc_section:
            # 回退：按“目录”标题与后续连续段落推断 TOC 区间
            toc_section = self._fallback_toc_section(doc, config)
            if not toc_section:
                # 若仍未识别，至少尝试格式化现有 TOC 样式段落
                self._format_existing_toc_paras(doc, config)
                return
        suspicious, reason = self._toc_section_is_suspicious(doc, config, toc_section)
        if suspicious:
            rescued = self._fallback_toc_section(doc, config)
            if (
                rescued
                and (
                    rescued.start_index != toc_section.start_index
                    or rescued.end_index != toc_section.end_index
                )
            ):
                rescue_suspicious, rescue_reason = self._toc_section_is_suspicious(
                    doc, config, rescued
                )
                if not rescue_suspicious:
                    tracker.record(
                        rule_name=self.name,
                        target="TOC range rescue",
                        section="toc",
                        change_type="skip",
                        before=(
                            f"range={toc_section.start_index}-{toc_section.end_index}"
                        ),
                        after=f"fallback_range={rescued.start_index}-{rescued.end_index}",
                        paragraph_index=rescued.start_index,
                    )
                    toc_section.start_index = rescued.start_index
                    toc_section.end_index = rescued.end_index
                    suspicious = False
                else:
                    reason = f"{reason}; fallback rejected: {rescue_reason}"

        if suspicious:
            tracker.record(
                rule_name=self.name,
                target="TOC range guard",
                section="toc",
                change_type="skip",
                before=(
                    f"range={toc_section.start_index}-{toc_section.end_index}, "
                    f"paras={len(doc.paragraphs)}"
                ),
                after=reason,
                paragraph_index=toc_section.start_index,
            )
            self._keep_existing_toc_with_normalize(
                doc,
                config,
                tracker,
                toc_section=toc_section,
                paragraph_index=toc_section.start_index,
            )
            return
        existing_toc_entries = self._count_existing_toc_entries(doc, toc_section)

        headings = context.get("headings", []) or []
        front_title_norms = get_front_matter_title_norms(config)
        if not headings:
            headings = _fallback_collect_headings(doc, config, toc_section)
        else:
            toc_start = toc_section.start_index if toc_section else -1
            toc_end = toc_section.end_index if toc_section else -1
            headings = [
                h for h in headings
                if not (toc_start <= getattr(h, "para_index", -1) <= toc_end)
                and not looks_like_toc_entry_line(getattr(h, "text", ""))
                and not looks_like_numbered_toc_entry_with_page_suffix(getattr(h, "text", ""))
                and _norm_no_space(getattr(h, "text", "")).lower() not in front_title_norms
            ]
        if not headings and existing_toc_entries >= 3:
            tracker.record(
                rule_name=self.name,
                target="TOC rebuild guard",
                section="toc",
                change_type="skip",
                before=f"existing_entries={existing_toc_entries}",
                after="no heading evidence, keep existing TOC entries",
                paragraph_index=toc_section.start_index,
            )
            self._keep_existing_toc_with_normalize(
                doc,
                config,
                tracker,
                toc_section=toc_section,
                paragraph_index=toc_section.start_index,
            )
            return

        toc_mode = _resolve_toc_mode(config)
        delta = 0
        if toc_mode == _TOC_MODE_PLAIN:
            # 确定编号方案
            levels_config = config.heading_numbering.levels

            # 生成普通目录条目
            entries = []
            entries.extend(_build_front_matter_toc_entries(doc, doc_tree))
            entries.extend(_build_back_matter_toc_entries(doc, doc_tree))
            if headings:
                headings = sorted(headings, key=lambda h: getattr(h, "para_index", 10**9))
                entries.extend(
                    _build_toc_entries(
                        doc,
                        headings,
                        levels_config,
                        config.heading_numbering_v2.level_bindings,
                    )
                )
            entries.sort(key=lambda e: e.get("para_index", 10**9))
            entries = _dedupe_entries(entries)
            if not entries:
                self._keep_existing_toc_with_normalize(
                    doc,
                    config,
                    tracker,
                    toc_section=toc_section,
                    paragraph_index=toc_section.start_index if toc_section else -1,
                )
                return
            if self._rebuild_coverage_too_low(existing_toc_entries, len(entries)):
                tracker.record(
                    rule_name=self.name,
                    target="TOC rebuild guard",
                    section="toc",
                    change_type="skip",
                    before=f"existing_entries={existing_toc_entries}",
                    after=f"rebuilt_entries={len(entries)} (coverage too low)",
                    paragraph_index=toc_section.start_index,
                )
                self._keep_existing_toc_with_normalize(
                    doc,
                    config,
                    tracker,
                    toc_section=toc_section,
                    paragraph_index=toc_section.start_index,
                )
                return

            # 为条目目标段落绑定书签，供 PAGEREF 自动取页码
            _attach_entry_bookmarks(doc, entries)

            # 计算右对齐 tab stop 位置（页面宽度 - 左右边距）
            ps = config.page_setup
            page_w_twips = 11906  # A4 宽度
            left_twips = _cm_to_twips(ps.margin.left_cm)
            right_twips = _cm_to_twips(ps.margin.right_cm)
            gutter_twips = _cm_to_twips(ps.gutter_cm)
            tab_pos = page_w_twips - left_twips - right_twips - gutter_twips

            # 清除旧 TOC 段落，插入新内容，返回段落数变化量
            delta = self._rebuild_toc(
                doc,
                doc_tree,
                toc_section,
                entries,
                tab_pos,
                config,
                tracker,
            )
        else:
            if not headings:
                tracker.record(
                    rule_name=self.name,
                    target="TOC rebuild guard",
                    section="toc",
                    change_type="skip",
                    before=f"existing_entries={existing_toc_entries}",
                    after="native mode needs heading evidence, keep existing TOC",
                    paragraph_index=toc_section.start_index,
                )
                self._keep_existing_toc_with_normalize(
                    doc,
                    config,
                    tracker,
                    toc_section=toc_section,
                    paragraph_index=toc_section.start_index,
                )
                return
            _purge_generated_toc_bookmarks(doc)
            levels_config = config.heading_numbering.levels
            native_seed_entries = []
            native_seed_entries.extend(_build_front_matter_toc_entries(doc, doc_tree))
            native_seed_entries.extend(_build_back_matter_toc_entries(doc, doc_tree))
            native_seed_entries.extend(
                _build_toc_entries(
                    doc,
                    sorted(headings, key=lambda h: getattr(h, "para_index", 10**9)),
                    levels_config,
                    config.heading_numbering_v2.level_bindings,
                )
            )
            native_seed_entries.sort(key=lambda e: e.get("para_index", 10**9))
            native_seed_entries = _dedupe_entries(native_seed_entries)

            ps = config.page_setup
            page_w_twips = 11906  # A4 宽度
            left_twips = _cm_to_twips(ps.margin.left_cm)
            right_twips = _cm_to_twips(ps.margin.right_cm)
            gutter_twips = _cm_to_twips(ps.gutter_cm)
            tab_pos = page_w_twips - left_twips - right_twips - gutter_twips

            delta = self._rebuild_native_toc(
                doc,
                toc_section,
                config,
                tracker,
                seed_entries=native_seed_entries,
                tab_pos_twips=tab_pos,
            )

        # 修正后续规则依赖的段落索引（TOC 重建改变了段落总数）
        if delta != 0:
            self._fix_indices_after_toc(context, doc_tree,
                                        toc_section, delta)

        # 再次应用目录段落格式，确保新旧目录段落都符合配置
        self._format_existing_toc_paras(doc, config, toc_section=toc_section)

    def _keep_existing_toc_with_normalize(
        self,
        doc: Document,
        config: SceneConfig,
        tracker: ChangeTracker,
        *,
        toc_section=None,
        paragraph_index: int = -1,
    ) -> None:
        normalized = self._normalize_plain_toc_entries(doc, config, toc_section=toc_section)
        self._format_existing_toc_paras(doc, config, toc_section=toc_section)
        if normalized > 0:
            tracker.record(
                rule_name=self.name,
                target="TOC normalize fallback",
                section="toc",
                change_type="format",
                before="plain text entries",
                after=f"normalized_entries={normalized}",
                paragraph_index=paragraph_index,
            )

    @staticmethod
    def _normalize_plain_toc_entries(doc: Document, config: SceneConfig, toc_section=None) -> int:
        if not doc.paragraphs:
            return 0

        start = 0
        end = len(doc.paragraphs) - 1
        if toc_section is not None:
            start = max(0, int(getattr(toc_section, "start_index", 0)))
            end = min(end, int(getattr(toc_section, "end_index", end)))
            if end < start:
                return 0

        ps = config.page_setup
        page_w_twips = 11906
        left_twips = _cm_to_twips(ps.margin.left_cm)
        right_twips = _cm_to_twips(ps.margin.right_cm)
        gutter_twips = _cm_to_twips(ps.gutter_cm)
        tab_pos = page_w_twips - left_twips - right_twips - gutter_twips

        changed = 0
        for idx in range(start, end + 1):
            para = doc.paragraphs[idx]
            text = (para.text or "").strip()
            if not text:
                continue
            norm = _norm_no_space(text).lower()
            if norm in {"目录", "目次", "contents", "tableofcontents"}:
                continue

            style_name = para.style.name if para.style else ""
            style_id = getattr(para.style, "style_id", "") if para.style else ""
            conf = _toc_entry_confidence(para, text, style_name, style_id)
            if conf < _TOC_GUARD_ENTRY_CONF:
                continue
            if _para_has_pageref_field(para):
                continue

            split = _split_plain_toc_entry_text(text)
            if not split:
                continue
            label, page = split

            level = _toc_level_from_style_name(style_name) or _toc_level_from_style_name(style_id)
            if level is None:
                level = _infer_toc_level_from_text(label)
            if level is None:
                level = "heading1"

            _rewrite_plain_toc_para(para, label, page, tab_pos)

            for style_name_candidate in _toc_style_candidates(level):
                try:
                    para.style = doc.styles[style_name_candidate]
                    break
                except KeyError:
                    continue

            style_key = _toc_style_key_for_level(level)
            sc = config.styles.get(style_key) if style_key else None
            _format_toc_para(para, sc or StyleConfig())
            changed += 1

        return changed

    def _rebuild_toc(self, doc, doc_tree, toc_section,
                     entries, tab_pos, config, tracker):
        """清除旧 TOC 段落并插入新内容。返回段落数变化量 delta。"""
        body_el = doc.element.body

        # 收集 TOC 分区内的所有段落元素
        toc_start = toc_section.start_index
        toc_end = toc_section.end_index
        toc_elements = []
        para_idx = 0
        for child in list(body_el):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                if toc_start <= para_idx <= toc_end:
                    toc_elements.append(child)
                para_idx += 1

        if not toc_elements:
            return 0

        # 检测旧 TOC 段落中是否包含分页符（w:br type="page"）
        had_page_break = _has_page_break(toc_elements)
        preserved_sect_pr = _extract_last_toc_section_break(toc_elements)

        # 记录插入位置（在删除前获取索引）
        insert_idx = list(body_el).index(toc_elements[0])

        # 删除所有旧 TOC 段落
        for el in toc_elements:
            body_el.remove(el)

        # 构建新 TOC 元素列表
        new_elements = []

        # 1. TOC 标题段落
        new_elements.append(_build_toc_title_para("目录"))

        # 2. TOC 条目段落
        for entry in entries:
            word_style = LEVEL_TO_WORD_STYLE.get(entry["level"], "TOC 1")
            new_elements.append(
                _build_toc_entry_para(entry, tab_pos, word_style))

        # 3. 恢复分页符（如果原 TOC 末尾有分页符）
        if had_page_break:
            new_elements.append(_build_page_break_para())

        # 插入新元素
        for i, el in enumerate(new_elements):
            body_el.insert(insert_idx + i, el)
        if new_elements:
            _apply_section_break_to_para(new_elements[-1], preserved_sect_pr)

        # 格式化新段落（通过 doc.paragraphs 找到对应段落）
        self._format_new_toc(doc, new_elements, entries, config)

        # 记录变更
        tracker.record(
            rule_name=self.name,
            target=f"TOC: {len(entries)} 个条目",
            section="toc",
            change_type="rebuild",
            before=f"{len(toc_elements)} 个旧段落",
            after=f"标题+{len(entries)}条目(PAGEREF页码域)",
            paragraph_index=toc_start,
        )

        # 返回段落数变化量
        old_count = len(toc_elements)
        new_count = len(new_elements)
        return new_count - old_count

    def _rebuild_native_toc(
        self,
        doc,
        toc_section,
        config,
        tracker,
        *,
        seed_entries: list[dict] | None = None,
        tab_pos_twips: int = 9000,
    ):
        """重建为 Word 原生自动目录（TOC 域）。返回段落数变化量 delta。"""
        body_el = doc.element.body

        toc_start = toc_section.start_index
        toc_end = toc_section.end_index
        toc_elements = []
        para_idx = 0
        for child in list(body_el):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                if toc_start <= para_idx <= toc_end:
                    toc_elements.append(child)
                para_idx += 1

        if not toc_elements:
            return 0

        had_page_break = _has_page_break(toc_elements)
        preserved_sect_pr = _extract_last_toc_section_break(toc_elements)
        insert_idx = list(body_el).index(toc_elements[0])

        for el in toc_elements:
            body_el.remove(el)

        new_elements = [
            _build_toc_title_para("目录"),
            _build_toc_field_begin_para(_native_toc_field_instruction(config)),
        ]
        for entry in seed_entries or []:
            word_style = LEVEL_TO_WORD_STYLE.get(entry.get("level", "heading1"), "TOC 1")
            new_elements.append(
                _build_toc_entry_para(entry, tab_pos_twips, word_style)
            )
        new_elements.append(_build_toc_field_end_para())
        if had_page_break:
            new_elements.append(_build_page_break_para())

        for i, el in enumerate(new_elements):
            body_el.insert(insert_idx + i, el)
        if new_elements:
            _apply_section_break_to_para(new_elements[-1], preserved_sect_pr)

        el_to_para = {p._element: p for p in doc.paragraphs}
        title_para = el_to_para.get(new_elements[0])
        if title_para:
            title_sc = config.styles.get("toc_title") or StyleConfig()
            _format_toc_para(title_para, title_sc, is_title=True)

        tracker.record(
            rule_name=self.name,
            target="TOC 原生自动目录",
            section="toc",
            change_type="rebuild",
            before=f"{len(toc_elements)} 个旧段落",
            after="标题+TOC域(Word原生自动目录)",
            paragraph_index=toc_start,
        )

        return len(new_elements) - len(toc_elements)

    @staticmethod
    def _fix_indices_after_toc(context, doc_tree, toc_section, delta):
        """修正 TOC 重建后所有后续段落索引的偏移。

        TOC 重建改变了段落总数，TOC 之后的所有段落索引都需要加上 delta。
        影响：context["headings"]、doc_tree 各分区边界。
        """
        toc_end_old = toc_section.end_index

        # 1. 修正 doc_tree 分区边界
        for sec in doc_tree.sections:
            if sec.section_type == "toc":
                # TOC 分区自身：end_index 调整为新 TOC 的末尾
                sec.end_index = sec.start_index + (sec.end_index - sec.start_index) + delta
                continue
            if sec.start_index > toc_end_old:
                sec.start_index += delta
            if sec.end_index > toc_end_old:
                sec.end_index += delta

        # 2. 修正 headings 段落索引
        headings = context.get("headings", [])
        for h in headings:
            if h.para_index > toc_end_old:
                h.para_index += delta

    def _format_new_toc(self, doc, new_elements, entries, config):
        """对新插入的 TOC 段落应用格式"""
        # 建立 element → paragraph 映射
        el_to_para = {p._element: p for p in doc.paragraphs}

        # 格式化标题段落（第一个元素）
        title_sc = config.styles.get("toc_title") or StyleConfig()
        title_para = el_to_para.get(new_elements[0])
        if title_para:
            _format_toc_para(title_para, title_sc, is_title=True)

        # 格式化条目段落（跳过标题，末尾可能有分页符）
        # new_elements: [title, entry0, entry1, ..., (page_break)]
        entry_elements = new_elements[1:1 + len(entries)]

        for i, el in enumerate(entry_elements):
            if i >= len(entries):
                break
            entry = entries[i]
            para = el_to_para.get(el)
            if not para:
                continue

            # 设置 Word TOC 样式（通过 python-docx API，自动解析 style_id）
            for style_name in _toc_style_candidates(entry["level"]):
                try:
                    para.style = doc.styles[style_name]
                    break
                except KeyError:
                    continue

            style_key = _toc_style_key_for_level(entry["level"])
            sc = config.styles.get(style_key) if style_key else None
            _format_toc_para(para, sc or StyleConfig())

    @staticmethod
    def _count_existing_toc_entries(doc: Document, toc_section) -> int:
        total = len(doc.paragraphs)
        start = max(0, int(getattr(toc_section, "start_index", 0)))
        end = min(total - 1, int(getattr(toc_section, "end_index", -1)))
        if end < start:
            return 0
        count = 0
        for i in range(start, end + 1):
            para = doc.paragraphs[i]
            text = (para.text or "").strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            style_id = getattr(para.style, "style_id", "") if para.style else ""
            if _toc_entry_confidence(para, text, style_name, style_id) >= _TOC_GUARD_ENTRY_CONF:
                count += 1
        return count

    @staticmethod
    def _rebuild_coverage_too_low(existing_entries: int, rebuilt_entries: int) -> bool:
        if existing_entries < 5:
            return False
        if rebuilt_entries <= 0:
            return True
        coverage = rebuilt_entries / max(existing_entries, 1)
        if coverage < 0.40 and (existing_entries - rebuilt_entries) >= 4:
            return True
        return False

    @staticmethod
    def _toc_section_is_suspicious(doc: Document, config: SceneConfig, toc_section) -> tuple[bool, str]:
        total = len(doc.paragraphs)
        if total <= 0:
            return True, "文档为空，跳过目录重建"

        start = max(0, int(getattr(toc_section, "start_index", 0)))
        end = min(total - 1, int(getattr(toc_section, "end_index", -1)))
        if end < start:
            return True, "目录区间无效，跳过目录重建"

        span = end - start + 1
        nonempty = 0
        entry_like = 0
        heading_like = 0
        strong_blacklist = 0

        for i in range(start, end + 1):
            para = doc.paragraphs[i]
            text = (para.text or "").strip()
            if not text:
                continue
            nonempty += 1
            style_name = para.style.name if para.style else ""
            style_id = getattr(para.style, "style_id", "") if para.style else ""
            conf = _toc_entry_confidence(para, text, style_name, style_id)
            if conf >= _TOC_GUARD_ENTRY_CONF:
                entry_like += 1
                continue
            if conf <= -2:
                strong_blacklist += 1
            if (
                detect_level_by_style_name(config, style_name) is not None
                or detect_level_by_style_name(config, style_id) is not None
            ):
                heading_like += 1

        span_ratio = span / max(total, 1)
        entry_ratio = entry_like / max(nonempty, 1)
        non_entry_count = max(0, nonempty - entry_like)

        # 防误删：区间过大且目录条目信号不足时，不执行重建。
        if span >= max(30, int(total * 0.60)) and entry_like < 3:
            return True, f"目录区间过大({span}/{total})且条目信号不足({entry_like})"
        if span_ratio >= 0.35 and entry_ratio < 0.15 and entry_like < 5:
            return True, (
                f"目录区间可疑(span={span}/{total}, entry={entry_like}/{nonempty})"
            )
        # Even with many TOC entries, a near-full-doc TOC range mixed with non-entry
        # paragraphs indicates overreach and can lead to destructive rebuild.
        if span_ratio >= 0.80 and non_entry_count >= 2:
            return True, (
                f"目录区间过大且混入非目录段落(span={span}/{total}, "
                f"entry={entry_like}, non_entry={non_entry_count})"
            )
        if strong_blacklist >= max(3, int(nonempty * 0.40)) and entry_like < 4:
            return True, (
                f"目录区间黑名单命中偏多(black={strong_blacklist}/{nonempty}, entry={entry_like})"
            )
        if heading_like >= 3 and entry_like == 0:
            return True, "目录区间包含大量正文标题信号，疑似误识别"

        return False, ""

    @staticmethod
    def _fallback_toc_section(doc: Document, config: SceneConfig):
        """When doc_tree misses TOC, infer TOC range from title + fallback boundary rules."""
        toc_title_idx = None
        for i, para in enumerate(doc.paragraphs):
            norm = _norm_no_space(para.text)
            if _RE_TOC_TITLE_CN.match(norm) or _RE_TOC_TITLE_EN.match(norm.lower()):
                toc_title_idx = i
                break
        if toc_title_idx is None:
            return None

        end = toc_title_idx
        found_entry = False
        blank_streak = 0
        for j in range(toc_title_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[j]
            txt = (para.text or "").strip()
            style_name = para.style.name if para.style else ""
            style_id = getattr(para.style, "style_id", "") if para.style else ""

            if not txt:
                # Allow sparse blank lines inside TOC area, but do not drift too far.
                if found_entry and blank_streak < 1:
                    blank_streak += 1
                    end = j
                    continue
                if not found_entry and j == toc_title_idx + 1:
                    end = j
                    continue
                break
            blank_streak = 0

            if (
                detect_level_by_style_name(config, style_name) is not None
                or detect_level_by_style_name(config, style_id) is not None
            ):
                break

            conf = _toc_entry_confidence(para, txt, style_name, style_id)
            if conf >= _TOC_FALLBACK_MIN_CONF:
                found_entry = True
                end = j
                continue

            numbered_heading_like = _looks_like_numbered_heading_text(txt)
            if numbered_heading_like:
                break

            # Once an entry cluster has started, first unrelated content ends TOC.
            if found_entry:
                break
            # Before first entry, unrelated content means fallback is unreliable.
            break

        if not found_entry:
            return None
        return SimpleNamespace(start_index=toc_title_idx, end_index=end)

    @staticmethod
    def _format_existing_toc_paras(doc: Document, config: SceneConfig, toc_section=None) -> None:
        """对文档中现有 TOC 条目段落应用配置化格式。"""
        start = None
        end = None
        if toc_section is not None:
            total = len(doc.paragraphs)
            start = max(0, int(getattr(toc_section, "start_index", 0)))
            end = min(total - 1, int(getattr(toc_section, "end_index", -1)))

        for idx, para in enumerate(doc.paragraphs):
            if start is not None and end is not None and not (start <= idx <= end):
                continue
            style_name = para.style.name if para.style else ""
            style_id = getattr(para.style, "style_id", "") if para.style else ""
            level = _toc_level_from_style_name(style_name) or _toc_level_from_style_name(style_id)
            text = (para.text or "").strip()
            if level is None:
                if _norm_no_space(text) in {"目录", "目次", "contents", "tableofcontents"}:
                    toc_heading_style = _find_toc_heading_style(doc)
                    if toc_heading_style is not None:
                        para.style = toc_heading_style
                    sc_title = config.styles.get("toc_title")
                    if sc_title:
                        _format_toc_para(para, sc_title, is_title=True)
                    continue
                if (
                    _para_has_pageref_field(para)
                    or looks_like_toc_entry_line(text)
                    or looks_like_numbered_toc_entry_with_page_suffix(text)
                ):
                    level = _infer_toc_level_from_text(text)
                else:
                    continue
            if level is None:
                level = "heading1"
            for style_name_candidate in _toc_style_candidates(level):
                try:
                    para.style = doc.styles[style_name_candidate]
                    break
                except KeyError:
                    continue
            style_key = _toc_style_key_for_level(level)
            sc = config.styles.get(style_key) if style_key else None
            _format_toc_para(para, sc or StyleConfig())

"""公式表格编号规则（实验功能）。"""

from __future__ import annotations

from docx import Document

from src.engine.change_tracker import ChangeTracker
from src.engine.rules import table_format as table_helpers
from src.engine.rules.base import BaseRule
from src.scene.schema import SceneConfig

_EXPLICIT_FRONT_MATTER_TYPES = {"cover", "toc", "abstract_cn", "abstract_en"}


def _target_paragraph_indices(context: dict | None) -> set[int]:
    if not isinstance(context, dict):
        return set()
    raw = context.get("target_paragraph_indices")
    if raw is None:
        return set()
    return {idx for idx in raw if isinstance(idx, int) and idx >= 0}


def _preserved_body_start_index(old_tree, total_paragraphs: int) -> int | None:
    """Only preserve manual pre_body/body scopes; keep explicit front matter auto-detected."""
    if old_tree is None or total_paragraphs <= 0:
        return None

    try:
        if any(
            old_tree.get_section(sec_type) is not None
            for sec_type in _EXPLICIT_FRONT_MATTER_TYPES
        ):
            return None
        pre_body = old_tree.get_section("pre_body")
        body_sec = old_tree.get_section("body")
    except Exception:
        return None

    if pre_body is None or body_sec is None or not isinstance(body_sec.start_index, int):
        return None
    return max(0, min(body_sec.start_index, total_paragraphs - 1))


def _refresh_structure_context(doc: Document, config: SceneConfig, context: dict | None) -> dict:
    """在公式规则改写文档结构后，刷新 doc_tree 与 headings。"""
    if not isinstance(context, dict):
        return {}

    try:
        from src.engine.doc_tree import DocTree
        from src.engine.rules.heading_detect import HeadingDetectRule
    except Exception:
        return context

    total_paragraphs = len(doc.paragraphs)
    old_tree = context.get("doc_tree")
    body_start_index = _preserved_body_start_index(old_tree, total_paragraphs)

    doc_tree = DocTree()
    if total_paragraphs > 0:
        if body_start_index is None:
            doc_tree.build(doc)
        else:
            doc_tree.build(doc, body_start_index=body_start_index)
    context["doc_tree"] = doc_tree

    heading_context = {"doc_tree": doc_tree}
    target_indices = _target_paragraph_indices(context)
    if target_indices:
        heading_context["target_paragraph_indices"] = set(target_indices)
    HeadingDetectRule().apply(doc, config, ChangeTracker(), heading_context)
    context["headings"] = heading_context.get("headings", [])
    if target_indices:
        context["target_paragraph_indices"] = set(target_indices)
    return context


def _parse_equation_numbering_format(numbering_format: str | None) -> tuple[bool, str]:
    raw = str(numbering_format or "chapter.seq").strip().lower()
    if raw == "seq":
        return False, ""
    if raw == "chapter-seq":
        return True, "-"
    return True, "."


def _compose_equation_number_id(
    *,
    chapter_num: int,
    seq_num: int,
    numbering_format: str | None,
) -> str:
    include_chapter, chapter_sep = _parse_equation_numbering_format(numbering_format)
    if include_chapter:
        return f"{chapter_num}{chapter_sep}{seq_num}"
    return str(seq_num)


def _existing_equation_number_matches(
    number_id: str | None,
    *,
    chapter_num: int,
    seq_num: int,
    numbering_format: str | None,
) -> bool:
    existing = str(number_id or "").strip()
    if not existing:
        return False
    expected = _compose_equation_number_id(
        chapter_num=chapter_num,
        seq_num=seq_num,
        numbering_format=numbering_format,
    )
    include_chapter, _ = _parse_equation_numbering_format(numbering_format)
    if include_chapter:
        parsed = table_helpers._parse_chapter_seq(existing)
        return parsed == (chapter_num, seq_num) and existing == expected
    return existing.isdigit() and int(existing) == seq_num and existing == expected


def _infer_table_existing_chapter(tbl_el) -> int | None:
    counts: dict[int, int] = {}
    for tr in tbl_el.findall(table_helpers._w("tr")):
        number_id, _, _ = table_helpers._find_existing_equation_number_in_row(tr)
        parsed = table_helpers._parse_chapter_seq(number_id or "")
        if parsed is None:
            continue
        chapter_num, _ = parsed
        counts[chapter_num] = counts.get(chapter_num, 0) + 1
    if not counts:
        return None
    return sorted(counts.items(), key=lambda item: (-item[1], item[0]))[0][0]


def _infer_dominant_existing_chapter(doc: Document, tbl_para_pos: list[int], scope_predicate) -> int | None:
    counts: dict[int, int] = {}
    for tbl_idx, tbl in enumerate(doc.tables):
        pos = tbl_para_pos[tbl_idx] if tbl_idx < len(tbl_para_pos) else -1
        if not scope_predicate(pos):
            continue
        tbl_el = tbl._tbl
        if not table_helpers._is_equation_table(tbl_el):
            continue
        chapter_num = _infer_table_existing_chapter(tbl_el)
        if chapter_num is None:
            continue
        counts[chapter_num] = counts.get(chapter_num, 0) + 1
    if not counts:
        return None
    return sorted(counts.items(), key=lambda item: (-item[1], item[0]))[0][0]


class EquationTableFormatRule(BaseRule):
    name = "equation_table_format"
    description = "公式表格编号"

    def apply(
        self,
        doc: Document,
        config: SceneConfig,
        tracker: ChangeTracker,
        context: dict,
    ) -> None:
        cfg = getattr(config, "equation_table_format", None)
        if cfg is not None and not bool(getattr(cfg, "enabled", False)):
            return
        numbering_format = str(
            getattr(cfg, "numbering_format", "chapter.seq")
        ).strip().lower() or "chapter.seq"
        if numbering_format not in {"seq", "chapter-seq", "chapter.seq"}:
            numbering_format = "chapter.seq"
        include_chapter_number, _ = _parse_equation_numbering_format(numbering_format)

        context = _refresh_structure_context(doc, config, context)
        avail_w = table_helpers._available_width(config)
        visual_cfg = getattr(config, "formula_table", None)
        number_alignment = str(
            getattr(visual_cfg, "number_alignment", "right")
        ).strip().lower() or "right"
        number_font_name = str(
            getattr(visual_cfg, "number_font_name", "Times New Roman")
        ).strip() or "Times New Roman"
        try:
            number_font_size_pt = float(
                getattr(visual_cfg, "number_font_size_pt", 10.5)
            )
        except (TypeError, ValueError):
            number_font_size_pt = 10.5
        if number_font_size_pt <= 0:
            number_font_size_pt = 10.5
        auto_shrink_number_column = bool(
            getattr(visual_cfg, "auto_shrink_number_column", True)
        )

        body_range = None
        target_indices = context.get("target_paragraph_indices")
        if target_indices is not None:
            target_indices = set(target_indices)

        doc_tree = context.get("doc_tree")
        has_doc_tree_scope = bool(getattr(doc_tree, "sections", None))
        if doc_tree:
            body_sec = doc_tree.get_section("body")
            if body_sec:
                start, end = body_sec.start_index, body_sec.end_index
                if end >= start:
                    body_range = (start, end)

        def _table_section_type(pos: int) -> str | None:
            if not has_doc_tree_scope or pos < 0:
                return None
            try:
                return doc_tree.get_section_for_paragraph(pos)
            except Exception:
                return None

        def _is_table_in_scope(pos: int) -> bool:
            if target_indices is not None:
                if pos < 0 or pos not in target_indices:
                    return False
            sec_type = _table_section_type(pos)
            if sec_type is not None:
                return sec_type == "body"
            if body_range is not None:
                return pos >= 0 and body_range[0] <= pos <= body_range[1]
            # Leading body tables can legitimately have no paragraph anchor yet.
            # Fall back to treating top-level document tables as in-scope.
            return True

        headings = context.get("headings", [])
        body_end = body_range[1] if body_range else (len(doc.paragraphs) - 1)
        chapter_ranges = table_helpers._build_chapter_ranges(headings, body_end)
        chapter_eq_counts: dict[int, int] = {}
        tbl_para_pos = table_helpers.top_level_table_anchor_positions(doc)
        dominant_existing_chapter = (
            _infer_dominant_existing_chapter(doc, tbl_para_pos, _is_table_in_scope)
            if include_chapter_number
            else None
        )
        last_effective_chapter: int | None = None
        global_eq_count = 0
        eq_table_count = 0
        eq_autonum_count = 0
        eq_renumber_count = 0

        for tbl_idx, tbl in enumerate(doc.tables):
            tbl_el = tbl._tbl
            pos = tbl_para_pos[tbl_idx] if tbl_idx < len(tbl_para_pos) else -1

            if not _is_table_in_scope(pos):
                continue
            if not table_helpers._is_equation_table(tbl_el):
                continue

            chap = table_helpers._get_chapter_num(pos, chapter_ranges)
            if include_chapter_number and chap <= 0:
                chap = (
                    _infer_table_existing_chapter(tbl_el)
                    or last_effective_chapter
                    or dominant_existing_chapter
                    or 1
                )
            if chap <= 0:
                chap = 1
            if include_chapter_number:
                last_effective_chapter = chap

            eq_rows_in_table = 0
            for row_idx, tr in enumerate(tbl_el.findall(table_helpers._w("tr"))):
                if not table_helpers._is_equation_row(tr):
                    continue

                eq_rows_in_table += 1
                current_seq = (
                    chapter_eq_counts.get(chap, 0)
                    if include_chapter_number
                    else global_eq_count
                )
                expected_seq = current_seq + 1
                expected_id = _compose_equation_number_id(
                    chapter_num=chap,
                    seq_num=expected_seq,
                    numbering_format=numbering_format,
                )
                expected_number = f"({expected_id})"
                existing_id, pure_cell, number_cell = table_helpers._find_existing_equation_number_in_row(tr)

                if existing_id:
                    if _existing_equation_number_matches(
                        existing_id,
                        chapter_num=chap,
                        seq_num=expected_seq,
                        numbering_format=numbering_format,
                    ):
                        if include_chapter_number:
                            chapter_eq_counts[chap] = expected_seq
                        else:
                            global_eq_count = expected_seq
                        continue

                    before_number = f"({existing_id})"
                    updated_cell = table_helpers._rewrite_existing_equation_number_in_row(
                        tr,
                        expected_number,
                        pure_cell=pure_cell,
                        number_cell=number_cell,
                    )
                    if updated_cell is None:
                        updated_cell = table_helpers._insert_equation_number_cell_in_row(
                            tr,
                            expected_number,
                        )
                        if (
                            updated_cell is not None
                            and number_cell is not None
                            and number_cell is not pure_cell
                        ):
                            table_helpers._strip_equation_number_tail_from_cell(
                                number_cell
                            )
                    if updated_cell is None:
                        continue
                    table_helpers._format_equation_number_cell(
                        updated_cell,
                        alignment=number_alignment,
                        font_name=number_font_name,
                        font_size_pt=number_font_size_pt,
                    )

                    if include_chapter_number:
                        chapter_eq_counts[chap] = expected_seq
                    else:
                        global_eq_count = expected_seq
                    eq_renumber_count += 1
                    tracker.record(
                        rule_name=self.name,
                        target=f"公式表格 #{tbl_idx} 第{row_idx + 1}行",
                        section="body",
                        change_type="numbering",
                        before=before_number,
                        after=expected_number,
                        paragraph_index=pos,
                    )
                    continue

                inserted_cell = table_helpers._insert_equation_number_cell_in_row(
                    tr,
                    expected_number,
                )
                if inserted_cell is None:
                    continue
                table_helpers._format_equation_number_cell(
                    inserted_cell,
                    alignment=number_alignment,
                    font_name=number_font_name,
                    font_size_pt=number_font_size_pt,
                )

                if include_chapter_number:
                    chapter_eq_counts[chap] = expected_seq
                else:
                    global_eq_count = expected_seq
                eq_autonum_count += 1
                tracker.record(
                    rule_name=self.name,
                    target=f"公式表格 #{tbl_idx} 第{row_idx + 1}行",
                    section="body",
                    change_type="numbering",
                    before="(无序号)",
                    after=expected_number,
                    paragraph_index=pos,
                )

            if eq_rows_in_table <= 0:
                continue

            if auto_shrink_number_column and avail_w > 0:
                try:
                    table_helpers._shrink_equation_number_column(tbl_el, avail_w)
                except Exception:
                    pass

            eq_table_count += 1

        if eq_table_count:
            tracker.record(
                rule_name=self.name,
                target=f"{eq_table_count} 个公式表格",
                section="body",
                change_type="format",
                before="未处理编号",
                after="已完成公式表格编号扫描与续编",
                paragraph_index=-1,
            )
        if eq_autonum_count:
            tracker.record(
                rule_name=self.name,
                target=f"{eq_autonum_count} 个公式行",
                section="body",
                change_type="numbering",
                before="(缺失序号)",
                after="已自动补全序号",
                paragraph_index=-1,
            )
        if eq_renumber_count:
            tracker.record(
                rule_name=self.name,
                target=f"{eq_renumber_count} 个公式行",
                section="body",
                change_type="numbering",
                before="(序号不连续/错误)",
                after="已按顺序自动纠正并递推后续序号",
                paragraph_index=-1,
            )

        stats_summary = (
            f"tables={eq_table_count}, auto_number={eq_autonum_count}, renumber={eq_renumber_count}"
        )
        context.setdefault("formula_runtime", {}).setdefault("stats", {})[self.name] = {
            "tables": eq_table_count,
            "auto_number": eq_autonum_count,
            "renumber": eq_renumber_count,
        }
        tracker.record(
            rule_name=self.name,
            target="summary",
            section="formula",
            change_type="format",
            before="统计",
            after=stats_summary,
            paragraph_index=-1,
        )

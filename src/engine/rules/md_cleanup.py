"""Markdown 粘贴修复规则：检测并转换段落中的 Markdown 语法"""

from copy import deepcopy
from collections import Counter
from functools import lru_cache
import re
import unicodedata
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

from src.engine.rules.base import BaseRule
from src.engine.change_tracker import ChangeTracker
from src.engine.page_scope import (
    parse_page_ranges_text,
    page_ranges_to_paragraph_ranges,
    page_number_to_start_paragraph_index,
    paragraph_ranges_to_index_set,
)
from src.scene.schema import SceneConfig, StyleConfig
from src.scene.heading_model import get_level_to_word_style, detect_level_by_style_name
from src.engine.doc_tree import DocTree
from src.formula_core.ast import FormulaNode
from src.formula_core.convert import convert_formula_node
from src.formula_core.parse import parse_document_formulas
from src.formula_core.repair import repair_formula_text
from src.markdown.block_parser import parse_docx_paragraphs
from src.markdown.inline_parser import parse_inline
from src.markdown.ir import MarkdownBlock, BlockType, InlineSpan, InlineType
from src.markdown.word_render import (
    write_spans, set_no_proof, add_hyperlink, write_table_cell,
    register_list_numbering, apply_num_pr, apply_blockquote_border,
    normalize_ordered_list_style, normalize_unordered_list_style,
    get_ordered_level_format, get_unordered_level_definition,
    _apply_inline_format, normalize_markdown_table, normalize_markdown_table_visual_only,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_RE_ESC_PLACEHOLDER_ARTIFACT = re.compile(r"(?:__)?ESC_?[0-9A-Fa-f]{8}(?:__)?")
_RE_LATEX_COMMAND = re.compile(
    r"(?:\\\[|\\\]|"
    r"\\(?:frac|sqrt|sum|int|prod|left|right|begin|end|lim|sin|cos|tan|log|ln|exp|"
    r"alpha|beta|gamma|delta|theta|lambda|mu|pi|sigma|omega|mathcal|mathbf|mathrm|"
    r"text|cdot|times|pm|mp|infty|nabla|partial|leq|geq|neq|approx|to|rightarrow|"
    r"leftarrow)\b)"
)

# Markdown 标题级别 → 场景 heading level key
HEADING_LEVEL_KEY_MAP = {
    1: "heading1", 2: "heading2",
    3: "heading3", 4: "heading4",
    5: "heading4", 6: "heading4",
}


def _replace_paragraph_with_omml(paragraph, omml_element) -> bool:
    """Replace paragraph content with OMML and keep paragraph properties."""
    p = paragraph._p
    if p is None:
        return False
    ppr_tag = f"{{{W_NS}}}pPr"
    run_tag = f"{{{W_NS}}}r"
    for child in list(p):
        if child.tag != ppr_tag:
            p.remove(child)

    omml_copy = deepcopy(omml_element)
    local = omml_copy.tag.split("}")[-1] if "}" in omml_copy.tag else omml_copy.tag
    if local == "oMathPara":
        p.append(omml_copy)
        return True

    run = p.makeelement(run_tag)
    run.append(omml_copy)
    p.append(run)
    return True


class MdCleanupRule(BaseRule):
    name = "md_cleanup"
    description = "检测并转换 Markdown 语法"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        doc_tree: DocTree = context.get("doc_tree")
        body = doc_tree.get_section("body") if doc_tree else None
        target_indices = context.get("target_paragraph_indices")
        if target_indices is not None:
            target_indices = set(target_indices)

        # 无 body 分区时不处理，防止跨分区误操作
        if not body:
            return

        manual_break_signature_count = self._count_manual_break_signatures(
            doc,
            start_index=body.start_index,
            end_index=body.end_index,
            target_indices=target_indices,
        )

        # 先把段落中的手动换行（Shift+Enter）拆成真实段落（Enter）。
        # 这一步可避免后续规则把“下箭头换行”继续传递到最终文档。
        split_count, target_indices = self._split_manual_break_paragraphs(
            doc,
            body.start_index,
            body.end_index,
            target_indices=target_indices,
        )
        body_end_index = min(len(doc.paragraphs) - 1, body.end_index + split_count)
        marker_separator = self._resolve_list_marker_separator(config)
        ordered_list_style = self._resolve_ordered_list_style(config)
        unordered_list_style = self._resolve_unordered_list_style(config)

        normalized_num_defs, normalized_prefix_runs = self._normalize_existing_word_lists(
            doc,
            marker_separator=marker_separator,
            ordered_list_style=ordered_list_style,
            unordered_list_style=unordered_list_style,
            start_index=body.start_index,
            end_index=body_end_index,
            target_indices=target_indices,
        )
        if normalized_num_defs or normalized_prefix_runs:
            tracker.record(
                rule_name=self.name,
                target="现有 Word 列表分隔符统一",
                section="body",
                change_type="format",
                before=(
                    f"num_defs={normalized_num_defs}, "
                    f"prefix_runs={normalized_prefix_runs}"
                ),
                after=(
                    f"→ marker_separator={marker_separator}, "
                    f"ordered={ordered_list_style}, unordered={unordered_list_style}"
                ),
                paragraph_index=body.start_index,
            )
        if split_count:
            tracker.record(
                rule_name=self.name,
                target=f"body 段落手动换行拆分",
                section="body",
                change_type="format",
                before=f"{split_count} 处段内换行",
                after="→ 拆分为真实段落（Enter）",
                paragraph_index=body.start_index,
            )

        formula_noise_cleaned, fake_formula_markers_cleared, repaired_para_indices = (
            self._normalize_formula_copy_noise_paragraphs(
                doc,
                config,
                tracker,
                start_index=body.start_index,
                end_index=body_end_index,
                target_indices=target_indices,
            )
        )
        if formula_noise_cleaned or fake_formula_markers_cleared:
            tracker.record(
                rule_name=self.name,
                target="formula_copy_noise_cleanup",
                section="body",
                change_type="text",
                before=(
                    f"formula_lines={formula_noise_cleaned}, "
                    f"fake_list_markers={fake_formula_markers_cleared}"
                ),
                after="→ markdown formula copy noise normalized",
                paragraph_index=body.start_index,
            )

        # 保护已识别出的公式段落，避免 Markdown inline 解析把 LaTeX 片段、
        # 显示公式边界行（如 \[ / \]）等再次改写坏掉。
        protected_formula_paragraphs: set[int] = set()
        try:
            formula_parse_result = parse_document_formulas(doc)
        except Exception:
            formula_parse_result = None
        if formula_parse_result is not None:
            for occ in formula_parse_result.occurrences:
                indices = list(getattr(occ, "source_paragraph_indices", []) or [])
                if not indices and getattr(occ, "paragraph_index", -1) >= 0:
                    indices = [int(occ.paragraph_index)]
                for raw_idx in indices:
                    try:
                        para_idx = int(raw_idx)
                    except (TypeError, ValueError):
                        continue
                    if para_idx < body.start_index or para_idx > body_end_index:
                        continue
                    if target_indices and para_idx not in target_indices:
                        continue
                    if self._should_protect_formula_occurrence(occ):
                        protected_formula_paragraphs.add(para_idx)

        # 收集作用范围内的段落文本（跳过已有 Heading 样式的段落和已修复公式的段落）
        para_texts: list[tuple[int, str]] = []
        for i, para in enumerate(doc.paragraphs):
            if not (body.start_index <= i <= body_end_index):
                continue
            if target_indices and i not in target_indices:
                continue
            if i in protected_formula_paragraphs:
                continue
            # 跳过已修复公式的段落，避免LaTeX命令被当作转义字符重新破坏
            if i in repaired_para_indices:
                continue
            para_text = para.text or ""
            # 已有 Heading 样式的段落不参与 Markdown 解析
            style_name = para.style.name if para.style else ""
            if detect_level_by_style_name(config, style_name):
                continue
            para_texts.append((i, para_text))

        # 解析为 IR
        markdown_paste_hint = bool(manual_break_signature_count or split_count)
        blocks = parse_docx_paragraphs(
            para_texts,
            markdown_paste_hint=markdown_paste_hint,
        )

        # 预注册列表编号
        bullet_num_id = register_list_numbering(
            doc,
            "bullet",
            marker_separator=marker_separator,
            unordered_style=unordered_list_style,
        )
        decimal_num_id = register_list_numbering(
            doc,
            "decimal",
            marker_separator=marker_separator,
            ordered_style=ordered_list_style,
        )

        # 收集脚注定义
        footnote_defs: dict[str, str] = {}
        for block in blocks:
            if block.type == BlockType.FOOTNOTE_DEF:
                footnote_defs[block.list_marker] = block.raw_text

        # 从场景配置获取正文字体（供 write_spans 使用）
        normal_sc = config.styles.get("normal")
        base_font_name = normal_sc.font_en if normal_sc else None
        base_font_size = Pt(normal_sc.size_pt) if normal_sc and normal_sc.size_pt else None
        source_doc_path = context.get("source_doc_path")
        md_base_path = context.get("source_doc_dir")
        if source_doc_path and not md_base_path:
            try:
                md_base_path = str(Path(str(source_doc_path)).resolve().parent)
            except Exception:
                md_base_path = None

        render_ctx = {
            "bullet_num_id": bullet_num_id,
            "decimal_num_id": decimal_num_id,
            "footnote_defs": footnote_defs,
            "base_font_name": base_font_name,
            "base_font_size": base_font_size,
            "heading_style_map": get_level_to_word_style(config),
            "md_base_path": md_base_path,
            "list_marker_separator": marker_separator,
            "ordered_list_style": ordered_list_style,
            "unordered_list_style": unordered_list_style,
        }
        md_table_element_ids = context.get("md_cleanup_table_element_ids")
        if not isinstance(md_table_element_ids, set):
            md_table_element_ids = set()
            context["md_cleanup_table_element_ids"] = md_table_element_ids
        render_ctx["md_cleanup_table_element_ids"] = md_table_element_ids

        # 过滤有效块（包含 BLANK / FOOTNOTE_DEF），倒序处理（避免索引偏移）
        md_blocks = [b for b in blocks if b.source_para_indices]
        md_blocks.sort(
            key=lambda b: b.source_para_indices[0], reverse=True)

        for block in md_blocks:
            self._apply_block(doc, block, config, tracker, render_ctx)

        normalized_table_count = self._normalize_existing_markdown_tables(
            doc,
            context,
            render_ctx,
        )
        if normalized_table_count > 0:
            tracker.record(
                rule_name=self.name,
                target=f"{normalized_table_count} 个表格",
                section="body",
                change_type="format",
                before="粘贴表格样式(虚线/缩进/未居中)",
                after="→ Markdown 表格规范化(去虚框+清缩进+强制居中)",
                paragraph_index=body.start_index,
            )

        promoted_list_markers, removed_list_markers = self._normalize_orphan_list_markers(
            doc,
            config,
            start_index=body.start_index,
            end_index=body_end_index,
            target_indices=target_indices,
        )
        if removed_list_markers > 0:
            tracker.record(
                rule_name=self.name,
                target="列表空标记段落清理",
                section="body",
                change_type="format",
                before=f"orphan_list_markers={removed_list_markers}",
                after=(
                    f"removed={removed_list_markers}, promoted={promoted_list_markers}"
                ),
                paragraph_index=body.start_index,
            )

        # md_cleanup 删除/合并了段落，doc_tree 索引已失效
        # 重建 doc_tree 供后续规则（section_format 等）使用
        if doc_tree:
            scope = context.get("format_scope", config.format_scope)
            body_start = None
            target_page_ranges: list[tuple[int, int]] = []
            if scope.mode == "manual":
                page_ranges_text = str(getattr(scope, "page_ranges_text", "") or "").strip()
                if page_ranges_text:
                    try:
                        target_page_ranges = parse_page_ranges_text(page_ranges_text)
                    except ValueError:
                        target_page_ranges = []
                if not target_page_ranges:
                    if scope.body_start_index is not None:
                        body_start = scope.body_start_index
                    elif scope.body_start_page is not None:
                        body_start = page_number_to_start_paragraph_index(
                            doc,
                            scope.body_start_page,
                            require_word=True,
                        )
                    elif scope.body_start_keyword:
                        body_start = self._find_keyword_index(
                            doc, scope.body_start_keyword)
            if target_page_ranges:
                doc_tree.build(doc)
                target_para_ranges = page_ranges_to_paragraph_ranges(
                    doc,
                    target_page_ranges,
                    require_word=True,
                )
                context["target_page_ranges"] = target_page_ranges
                context["target_paragraph_ranges"] = target_para_ranges
                context["target_paragraph_indices"] = paragraph_ranges_to_index_set(target_para_ranges)
            else:
                doc_tree.build(doc, body_start_index=body_start)
                context.pop("target_page_ranges", None)
                context.pop("target_paragraph_ranges", None)
                context.pop("target_paragraph_indices", None)
            context["doc_tree"] = doc_tree

    @staticmethod
    def _table_has_merge(tbl_el) -> bool:
        for tc in tbl_el.iter(qn("w:tc")):
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                continue
            grid_span = tc_pr.find(qn("w:gridSpan"))
            if grid_span is not None:
                val = grid_span.get(qn("w:val"))
                try:
                    if int(str(val).strip()) > 1:
                        return True
                except (TypeError, ValueError):
                    return True
            v_merge = tc_pr.find(qn("w:vMerge"))
            if v_merge is not None:
                return True
        return False

    @staticmethod
    def _table_cell_counts(tbl) -> list[int]:
        counts: list[int] = []
        for row in tbl.rows:
            counts.append(len(row.cells))
        return counts

    @staticmethod
    def _table_is_simple_rectangular(tbl) -> bool:
        counts = MdCleanupRule._table_cell_counts(tbl)
        if not counts:
            return False
        if min(counts) < 2:
            return False
        if min(counts) != max(counts):
            return False
        return not MdCleanupRule._table_has_merge(tbl._tbl)

    @staticmethod
    def _table_has_non_solid_border(tbl_el) -> bool:
        solid_values = {"single", "none", "nil"}
        for border in tbl_el.iter():
            tag = getattr(border, "tag", "")
            if tag not in {
                qn("w:top"),
                qn("w:left"),
                qn("w:bottom"),
                qn("w:right"),
                qn("w:insideH"),
                qn("w:insideV"),
            }:
                continue
            val = str(border.get(qn("w:val")) or "").strip().lower()
            if not val:
                continue
            if val not in solid_values:
                return True
        return False

    @staticmethod
    def _table_has_tc_borders(tbl_el) -> bool:
        for tc in tbl_el.iter(qn("w:tc")):
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                continue
            if tc_pr.find(qn("w:tcBorders")) is not None:
                return True
        return False

    @staticmethod
    def _table_has_borderless_gridline_signature(tbl_el) -> bool:
        tbl_pr = tbl_el.find(qn("w:tblPr"))
        if tbl_pr is None:
            return False
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            return False
        border_sides = ("top", "left", "bottom", "right", "insideH", "insideV")
        found_side = False
        for side in border_sides:
            border = tbl_borders.find(qn(f"w:{side}"))
            if border is None:
                continue
            found_side = True
            val = str(border.get(qn("w:val")) or "").strip().lower()
            if val not in {"none", "nil"}:
                return False
        return found_side

    @staticmethod
    def _table_has_any_visible_border(tbl_el) -> bool:
        border_tags = {
            qn("w:top"),
            qn("w:left"),
            qn("w:bottom"),
            qn("w:right"),
            qn("w:insideH"),
            qn("w:insideV"),
        }
        for node in tbl_el.iter():
            if getattr(node, "tag", "") not in border_tags:
                continue
            val = str(node.get(qn("w:val")) or "").strip().lower()
            if val and val not in {"none", "nil"}:
                return True
        return False

    @staticmethod
    def _table_has_style_or_look(tbl_el) -> bool:
        tbl_pr = tbl_el.find(qn("w:tblPr"))
        if tbl_pr is None:
            return False
        if tbl_pr.find(qn("w:tblStyle")) is not None:
            return True
        if tbl_pr.find(qn("w:tblLook")) is not None:
            return True
        return False

    @staticmethod
    def _table_has_indent_or_not_center(tbl_el) -> bool:
        for p in tbl_el.iter(qn("w:p")):
            p_pr = p.find(qn("w:pPr"))
            if p_pr is None:
                continue
            jc = p_pr.find(qn("w:jc"))
            if jc is not None and str(jc.get(qn("w:val")) or "").strip().lower() not in {"", "center"}:
                return True
            ind = p_pr.find(qn("w:ind"))
            if ind is not None:
                for attr in (
                    "left", "leftChars", "right", "rightChars",
                    "firstLine", "firstLineChars", "hanging", "hangingChars",
                ):
                    val = str(ind.get(qn(f"w:{attr}")) or "").strip()
                    if val and val != "0":
                        return True
        return False

    @staticmethod
    def _table_is_markdown_like(tbl) -> bool:
        tbl_el = tbl._tbl
        if not MdCleanupRule._table_has_any_visible_border(tbl_el):
            return True
        if MdCleanupRule._table_has_non_solid_border(tbl_el):
            return True
        if MdCleanupRule._table_has_tc_borders(tbl_el):
            return True
        if MdCleanupRule._table_has_borderless_gridline_signature(tbl_el):
            return True
        if (
            MdCleanupRule._table_has_style_or_look(tbl_el)
            and MdCleanupRule._table_has_indent_or_not_center(tbl_el)
        ):
            return True
        return False

    @staticmethod
    def _table_text_matrix(tbl) -> tuple[list[str], list[list[str]]]:
        rows = tbl.rows
        if not rows:
            return [], []
        headers = [str(cell.text or "").strip() for cell in rows[0].cells]
        body_rows: list[list[str]] = []
        for row in rows[1:]:
            body_rows.append([str(cell.text or "").strip() for cell in row.cells])
        return headers, body_rows

    @staticmethod
    def _collect_table_positions(doc: Document) -> list[int]:
        positions: list[int] = []
        para_idx = -1
        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                para_idx += 1
            elif tag == "tbl":
                positions.append(para_idx if para_idx >= 0 else 0)
        return positions

    def _normalize_existing_markdown_tables(self, doc: Document, context: dict, ctx: dict) -> int:
        doc_tree: DocTree | None = context.get("doc_tree")
        body_range: tuple[int, int] | None = None
        if doc_tree is not None:
            body_sec = doc_tree.get_section("body")
            if body_sec and body_sec.end_index >= body_sec.start_index:
                body_range = (int(body_sec.start_index), int(body_sec.end_index))

        target_indices = context.get("target_paragraph_indices")
        if target_indices is not None:
            target_indices = set(target_indices)

        def in_scope(pos: int) -> bool:
            if body_range is not None and (pos < body_range[0] or pos > body_range[1]):
                return False
            if target_indices is not None and pos not in target_indices:
                return False
            return True

        md_table_ids = ctx.get("md_cleanup_table_element_ids")
        if not isinstance(md_table_ids, set):
            md_table_ids = set()
            ctx["md_cleanup_table_element_ids"] = md_table_ids

        table_positions = self._collect_table_positions(doc)
        normalized = 0
        for idx, table in enumerate(doc.tables):
            tbl_el = table._tbl
            if id(tbl_el) in md_table_ids:
                continue
            pos = table_positions[idx] if idx < len(table_positions) else -1
            if pos >= 0 and not in_scope(pos):
                continue
            if not self._table_is_markdown_like(table):
                continue
            headers, rows = self._table_text_matrix(table)
            if self._table_is_simple_rectangular(table) and len(headers) >= 2:
                normalize_markdown_table(doc, table, headers, rows)
            else:
                normalize_markdown_table_visual_only(table)
            md_table_ids.add(id(tbl_el))
            normalized += 1
        return normalized

    def _apply_block(self, doc, block, config, tracker, ctx):
        """按类型分发处理"""
        handlers = {
            BlockType.HEADING: self._apply_heading,
            BlockType.CODE_BLOCK: self._apply_code_block,
            BlockType.TABLE: self._apply_table,
            BlockType.BLOCKQUOTE: self._apply_blockquote,
            BlockType.HORIZONTAL_RULE: self._apply_hr,
            BlockType.LATEX_BLOCK: self._apply_latex_block,
            BlockType.LIST_ITEM: self._apply_list_item,
            BlockType.TASK_LIST_ITEM: self._apply_task_list,
            BlockType.PARAGRAPH: self._apply_inline,
            BlockType.BLANK: self._apply_blank,
            BlockType.FOOTNOTE_DEF: self._apply_footnote_def,
        }
        handler = handlers.get(block.type)
        if handler:
            handler(doc, block, config, tracker, ctx)

    def _apply_heading(self, doc, block, config, tracker, ctx):
        """剥离 # 标记，应用 Heading 样式"""
        idx = block.source_para_indices[0]
        para = doc.paragraphs[idx]
        old_text = para.text
        style_map = ctx.get("heading_style_map") or {}
        level_key = HEADING_LEVEL_KEY_MAP.get(block.level, "heading4")
        style_name = style_map.get(level_key, f"Heading {min(block.level, 4)}")

        write_spans(para, block.spans,
                    doc=doc,
                    footnote_defs=ctx.get("footnote_defs"),
                    base_font_name=ctx.get("base_font_name"),
                    base_font_size=ctx.get("base_font_size"),
                    image_base_path=ctx.get("md_base_path"),
                    prefer_real_image=True)
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            pass

        tracker.record(
            rule_name=self.name, target=f"段落 #{idx}",
            section="body", change_type="text",
            before=old_text[:80],
            after=f"→ {style_name}: {block.raw_text[:50]}",
            paragraph_index=idx)

    def _apply_code_block(self, doc, block, config, tracker, ctx):
        """围栏代码块：删除 ``` 行，内容段落应用等宽字体 + noProof"""
        indices = block.source_para_indices
        if len(indices) < 2:
            return

        code_sc = config.styles.get("code_block", StyleConfig(
            font_cn="Consolas", font_en="Consolas", size_pt=10,
            alignment="left", first_line_indent_chars=0,
            line_spacing_type="exact", line_spacing_pt=16))

        content_indices = indices[1:-1] if len(indices) > 2 else []
        fence_indices = [indices[0], indices[-1]]

        for ci in content_indices:
            if ci >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[ci]
            pf = para.paragraph_format
            pf.first_line_indent = Pt(0)
            pf.left_indent = Cm(1.0)
            for run in para.runs:
                run.font.name = code_sc.font_en
                run.font.size = Pt(code_sc.size_pt)
                set_no_proof(run)

        for fi in sorted(fence_indices, reverse=True):
            if fi >= len(doc.paragraphs):
                continue
            p = doc.paragraphs[fi]
            p._element.getparent().remove(p._element)

        tracker.record(
            rule_name=self.name,
            target=f"段落 #{indices[0]}-#{indices[-1]}",
            section="body", change_type="format",
            before=f"```{block.language}...```",
            after=f"→ code_block ({len(content_indices)} 行)",
            paragraph_index=indices[0])

    def _apply_table(self, doc, block, config, tracker, ctx):
        """Markdown 表格 → Word 原生表格（支持单元格换行和行内格式）"""
        indices = block.source_para_indices
        if not indices:
            return

        num_cols = len(block.table_headers)
        num_rows = 1 + len(block.table_rows)

        table = doc.add_table(rows=num_rows, cols=num_cols)

        for j, header in enumerate(block.table_headers):
            if j < num_cols:
                cell = table.rows[0].cells[j]
                write_table_cell(
                    cell,
                    header,
                    image_base_path=ctx.get("md_base_path"),
                    prefer_real_image=True,
                )
                for run in cell.paragraphs[0].runs:
                    run.bold = True

        for i, row_data in enumerate(block.table_rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    write_table_cell(
                        table.rows[i + 1].cells[j],
                        cell_text,
                        image_base_path=ctx.get("md_base_path"),
                        prefer_real_image=True,
                    )

        # 应用列对齐
        align_map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
        }
        if block.table_alignments:
            for row in table.rows:
                for j, cell in enumerate(row.cells):
                    if j < len(block.table_alignments):
                        align = align_map.get(block.table_alignments[j])
                        if align and cell.paragraphs:
                            cell.paragraphs[0].alignment = align

        normalize_markdown_table(doc, table, block.table_headers, block.table_rows)

        # 移动表格到原位置
        first_para = doc.paragraphs[indices[0]]
        first_para._element.addprevious(table._tbl)
        md_table_ids = ctx.get("md_cleanup_table_element_ids")
        if isinstance(md_table_ids, set):
            md_table_ids.add(id(table._tbl))

        # 删除原始 | 段落（倒序）
        for pi in sorted(indices, reverse=True):
            if pi >= len(doc.paragraphs):
                continue
            p = doc.paragraphs[pi]
            p._element.getparent().remove(p._element)

        tracker.record(
            rule_name=self.name,
            target=f"段落 #{indices[0]}-#{indices[-1]}",
            section="body", change_type="text",
            before=f"{len(indices)} 行 Markdown 表格",
            after=f"→ Word 表格 {num_rows}x{num_cols}",
            paragraph_index=indices[0])

    def _apply_blockquote(self, doc, block, config, tracker, ctx):
        """引用块：剥离 > 前缀，添加左边框+底色"""
        base_name, base_size = ctx.get("base_font_name"), ctx.get("base_font_size")
        for idx in block.source_para_indices:
            if idx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[idx]
            old_text = para.text
            clean = re.sub(r'^(?:>\s*)+', '', para.text)
            write_spans(
                para,
                self._resolve_inline_spans(clean),
                doc=doc,
                footnote_defs=ctx.get("footnote_defs"),
                base_font_name=base_name,
                base_font_size=base_size,
                image_base_path=ctx.get("md_base_path"),
                prefer_real_image=True,
            )
            apply_blockquote_border(para, block.quote_level)
            tracker.record(
                rule_name=self.name, target=f"段落 #{idx}",
                section="body", change_type="text",
                before=old_text[:60],
                after=f"→ blockquote (level {block.quote_level})",
                paragraph_index=idx)

    def _apply_blank(self, doc, block, config, tracker, ctx):
        """删除 Markdown 空行产生的空段落（保留含图片的段落）"""
        md_cfg = getattr(config, "md_cleanup", None)
        keep_list_markers = bool(
            getattr(md_cfg, "preserve_existing_word_lists", True)
        )
        for idx in sorted(block.source_para_indices, reverse=True):
            if idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]
                if keep_list_markers and self._para_has_list_marker(p):
                    continue
                if not p.text.strip() and not self._para_has_content(p):
                    p._element.getparent().remove(p._element)

    def _apply_footnote_def(self, doc, block, config, tracker, ctx):
        """删除脚注定义源行，正文引用已在 spans 渲染阶段处理。"""
        indices = block.source_para_indices
        for idx in sorted(indices, reverse=True):
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                para._element.getparent().remove(para._element)

        if indices:
            tracker.record(
                rule_name=self.name,
                target=f"段落 #{indices[0]}",
                section="body",
                change_type="text",
                before=f"[^{block.list_marker}]: {block.raw_text[:40]}",
                after="(脚注定义已转为引用并删除定义行)",
                paragraph_index=indices[0],
            )

    def _normalize_formula_copy_noise_paragraphs(
        self,
        doc: Document,
        config: SceneConfig,
        tracker: ChangeTracker,
        *,
        start_index: int,
        end_index: int,
        target_indices: set[int] | None,
    ) -> tuple[int, int, set[int]]:
        md_cfg = getattr(config, "md_cleanup", None)
        if md_cfg is None:
            return 0, 0, set()

        formula_cleanup_enabled = bool(
            getattr(md_cfg, "formula_copy_noise_cleanup", True)
        )
        fake_list_cleanup_enabled = bool(
            getattr(md_cfg, "suppress_formula_fake_lists", True)
        )
        if not formula_cleanup_enabled and not fake_list_cleanup_enabled:
            return 0, 0, set()

        if start_index < 0:
            start_index = 0
        if end_index < start_index or not doc.paragraphs:
            return 0, 0, set()
        end_index = min(end_index, len(doc.paragraphs) - 1)
        if end_index < start_index:
            return 0, 0, set()

        target_set = set(target_indices) if target_indices else None
        normalized_count = 0
        fake_list_removed_count = 0
        repaired_para_indices = set()  # 记录被修复的段落索引

        if formula_cleanup_enabled:
            for idx in range(start_index, end_index + 1):
                if target_set is not None and idx not in target_set:
                    continue
                para = doc.paragraphs[idx]
                if self._para_has_content(para):
                    continue
                raw_text = para.text or ""
                stripped = raw_text.strip()
                if not stripped:
                    continue
                cleaned = self._normalize_formula_copy_noise_text(stripped)
                if not cleaned or cleaned == stripped:
                    continue
                self._replace_text(para, cleaned)
                normalized_count += 1
                repaired_para_indices.add(idx)  # 记录被修复的段落
                tracker.record(
                    rule_name=self.name,
                    target=f"段落 #{idx}",
                    section="body",
                    change_type="text",
                    before=raw_text[:120],
                    after=f"→ formula copy noise cleanup: {cleaned[:120]}",
                    paragraph_index=idx,
                )

        if fake_list_cleanup_enabled:
            for idx in range(min(end_index, len(doc.paragraphs) - 1), start_index - 1, -1):
                if idx >= len(doc.paragraphs):
                    continue
                if target_set is not None and idx not in target_set:
                    continue
                para = doc.paragraphs[idx]
                if self._para_has_content(para):
                    continue
                if self._para_has_list_marker(para):
                    continue

                stripped = (para.text or "").strip()
                if not self._is_fake_formula_list_marker_line(stripped):
                    continue

                prev_text = self._find_neighbor_formula_context_text(
                    doc,
                    config,
                    start_index=start_index,
                    end_index=end_index,
                    target_set=target_set,
                    from_index=idx,
                    direction=-1,
                )
                next_text = self._find_neighbor_formula_context_text(
                    doc,
                    config,
                    start_index=start_index,
                    end_index=end_index,
                    target_set=target_set,
                    from_index=idx,
                    direction=1,
                )
                if not (prev_text and next_text):
                    continue
                if not (
                    self._looks_like_formula_line(prev_text)
                    and self._looks_like_formula_line(next_text)
                ):
                    continue

                para._element.getparent().remove(para._element)
                fake_list_removed_count += 1
                tracker.record(
                    rule_name=self.name,
                    target=f"段落 #{idx}",
                    section="body",
                    change_type="text",
                    before=stripped[:120],
                    after="→ removed fake list marker between formula lines",
                    paragraph_index=idx,
                )

        return normalized_count, fake_list_removed_count, repaired_para_indices

    def _normalize_formula_copy_noise_text(self, text: str) -> str | None:
        raw = str(text or "").strip()
        if not raw:
            return None

        repaired = repair_formula_text(raw)
        repaired_text = str(repaired.text or "").strip()
        if "ESC_" in raw or "__ESC" in raw:
            if self._is_valid_escape_formula_repair(
                raw_text=raw,
                repaired_text=repaired_text,
                warnings=repaired.warnings,
            ):
                return repaired_text
        elif self._is_valid_general_formula_repair(
            raw_text=raw,
            repaired_text=repaired_text,
            source=str(repaired.source or "").strip(),
            confidence=float(repaired.confidence or 0.0),
        ):
            return repaired_text

        compact = self._compact_formula_text(raw)
        if not compact:
            return None

        segments = self._split_formula_segments(compact)
        if len(segments) < 2:
            return None

        skeletons = [
            self._formula_skeleton_key(seg)
            for seg in segments
            if self._formula_skeleton_key(seg)
        ]
        if not skeletons:
            return None

        skeleton_counter = Counter(skeletons)
        dominant_key, dominant_count = skeleton_counter.most_common(1)[0]
        if dominant_count < 2:
            return None

        dominant_segments = [
            seg for seg in segments if self._formula_skeleton_key(seg) == dominant_key
        ]
        if not dominant_segments:
            return None

        seg_counter = Counter(dominant_segments)
        best = max(
            set(dominant_segments),
            key=lambda seg: (
                self._formula_detail_score(seg),
                seg_counter[seg],
                len(seg),
            ),
        )
        if best == compact:
            return None
        return best

    @staticmethod
    def _is_valid_escape_formula_repair(
        *,
        raw_text: str,
        repaired_text: str,
        warnings: list[str],
    ) -> bool:
        """Validate if ESC formula repair produced a valid result.

        Uses a more lenient approach that allows partial repairs if they
        show significant improvement over the original text.
        """
        if not repaired_text or repaired_text == raw_text:
            return False

        # Count unresolved ESC placeholders in the result
        unresolved_count = repaired_text.count("ESC_") + repaired_text.count("__ESC")

        # Allow up to 2 unresolved placeholders if there's significant improvement
        if unresolved_count > 2:
            return False

        # Must have LaTeX commands or show significant length reduction
        has_latex = "\\" in repaired_text
        significantly_shorter = len(repaired_text) < len(raw_text) * 0.6

        return has_latex or significantly_shorter

    @staticmethod
    def _is_valid_general_formula_repair(
        *,
        raw_text: str,
        repaired_text: str,
        source: str,
        confidence: float,
    ) -> bool:
        if not repaired_text or repaired_text == raw_text:
            return False

        if source in {"rendered_fraction", "fraction_identity", "quadratic_formula_copy"}:
            return True
        if source == "repeated_segment":
            return len(repaired_text) <= len(raw_text) * 0.75
        if "\\" in repaired_text and len(repaired_text) < len(raw_text):
            return True
        return confidence >= 0.9 and len(repaired_text) < len(raw_text)

    def _split_formula_segments(self, compact_text: str) -> list[str]:
        s = compact_text
        if not s:
            return []
        if len(s) > 240:
            return []
        if s.count("=") < 2:
            return []

        @lru_cache(maxsize=None)
        def _dfs(start: int) -> tuple[str, ...] | None:
            if start == len(s):
                return ()
            if start < 0 or start >= len(s):
                return None
            if not (s[start].isalpha() or s[start] in {"("}):
                return None

            best: tuple[str, ...] | None = None
            best_score = (-1, -1, -1)
            max_end = min(len(s), start + 120)
            for end in range(start + 5, max_end + 1):
                seg = s[start:end]
                eq_count = seg.count("=")
                if eq_count == 0:
                    continue
                if eq_count > 1:
                    break
                if not self._is_formula_segment_candidate(seg):
                    continue

                rest = _dfs(end)
                if rest is None:
                    continue

                candidate = (seg,) + rest
                # Prefer more segments, then uniform lengths (lower variance
                # favours correct repetition-boundary alignment), then higher
                # detail score sum.
                lengths = [len(item) for item in candidate]
                avg_len = sum(lengths) / len(lengths)
                variance = sum((ln - avg_len) ** 2 for ln in lengths)
                score = (
                    len(candidate),
                    -variance,
                    sum(self._formula_detail_score(item) for item in candidate),
                )
                if best is None or score > best_score:
                    best = candidate
                    best_score = score

            return best

        result = _dfs(0)
        if result is None or len(result) < 2:
            return []
        return list(result)

    @staticmethod
    def _is_formula_segment_candidate(seg: str) -> bool:
        if not seg:
            return False
        if len(seg) < 5 or len(seg) > 120:
            return False
        if seg.count("=") != 1:
            return False
        if re.search(r"[\u4e00-\u9fff]", seg):
            return False

        left, right = seg.split("=", 1)
        if not left or not right:
            return False
        if not (left[0].isalpha() or left[0] in {"("}):
            return False
        if not (right[0].isalpha() or right[0].isdigit() or right[0] in {"("}):
            return False
        if not (left[-1].isalnum() or left[-1] in {")", "]", "}"}):
            return False
        if not (right[-1].isalnum() or right[-1] in {")", "]", "}"}):
            return False

        for ch in seg:
            if ch.isalnum():
                continue
            if ch in {"+", "-", "*", "/", "^", "_", "=", "(", ")", "[", "]", "{", "}", ".", ",", "\\", "<", ">", "!", "~"}:
                continue
            return False

        if not any(ch.isalpha() for ch in seg):
            return False
        # Must contain at least one of:
        # 1. Operators: + - * / ^ _
        # 2. Digits
        # 3. Brackets: [ ] (for subscripts/array notation like E[X])
        # 4. Function call pattern: letter followed by (
        if not re.search(r"[+\-*/^_\d\[\]]", seg):
            # Check for function call pattern
            if not re.search(r"[a-zA-Z]\(", seg):
                return False
        return True

    @staticmethod
    def _formula_skeleton_key(text: str) -> str:
        return re.sub(r"[\^_{}\(\)\s]", "", text.lower())

    @staticmethod
    def _formula_detail_score(text: str) -> int:
        return (
            text.count("^") * 4
            + text.count("_") * 4
            + text.count("{")
            + text.count("}")
            + text.count("\\")
            + len(text) // 20
        )

    @staticmethod
    def _compact_formula_text(text: str) -> str:
        normalized = unicodedata.normalize("NFKC", str(text or ""))
        normalized = normalized.replace("\u2212", "-")  # MINUS SIGN → hyphen-minus
        normalized = normalized.replace("\u00d7", "*")  # × → *
        normalized = normalized.replace("\u00b7", "*")  # · → *
        # Common math Unicode symbols
        normalized = normalized.replace("\u2192", "->")  # → RIGHTWARDS ARROW
        normalized = normalized.replace("\u2190", "<-")  # ← LEFTWARDS ARROW
        normalized = normalized.replace("\u221e", "inf")  # ∞ INFINITY
        normalized = normalized.replace("\u2211", "sum")  # ∑ N-ARY SUMMATION
        normalized = normalized.replace("\u222b", "int")  # ∫ INTEGRAL
        normalized = normalized.replace("\u220f", "prod")  # ∏ N-ARY PRODUCT
        normalized = normalized.replace("\u221a", "sqrt")  # √ SQUARE ROOT
        normalized = normalized.replace("\u2248", "~=")  # ≈ ALMOST EQUAL TO
        normalized = normalized.replace("\u2260", "!=")  # ≠ NOT EQUAL TO
        normalized = normalized.replace("\u2264", "<=")  # ≤ LESS-THAN OR EQUAL TO
        normalized = normalized.replace("\u2265", ">=")  # ≥ GREATER-THAN OR EQUAL TO
        return re.sub(r"\s+", "", normalized)

    def _looks_like_formula_line(self, text: str) -> bool:
        compact = self._compact_formula_text(text)
        if self._is_formula_segment_candidate(compact):
            return True
        normalized = self._normalize_formula_copy_noise_text(compact)
        return bool(
            normalized and self._is_formula_segment_candidate(normalized)
        )

    @staticmethod
    def _is_fake_formula_list_marker_line(text: str) -> bool:
        s = str(text or "").strip()
        if not s or len(s) > 10:
            return False
        return bool(
            re.match(r"^(?:\d+[.)、]|[（(]\d+[)）])$", s)
        )

    def _find_neighbor_formula_context_text(
        self,
        doc: Document,
        config: SceneConfig,
        *,
        start_index: int,
        end_index: int,
        target_set: set[int] | None,
        from_index: int,
        direction: int,
    ) -> str | None:
        step = -1 if direction < 0 else 1
        idx = from_index + step
        while start_index <= idx <= end_index:
            if target_set is not None and idx not in target_set:
                idx += step
                continue

            para = doc.paragraphs[idx]
            if self._para_has_content(para):
                return None

            style_name = para.style.name if para.style else ""
            if detect_level_by_style_name(config, style_name):
                return None

            text = (para.text or "").strip()
            if not text:
                idx += step
                continue
            if self._is_fake_formula_list_marker_line(text):
                idx += step
                continue
            return text

        return None

    @staticmethod
    def _para_has_content(para) -> bool:
        """检查段落是否包含非文本内容（图片、公式等）"""
        el = para._element
        W = f"{{{W_NS}}}"
        if el.findall(f"{W}r/{W}drawing"):
            return True
        if el.findall(f"{W}r/{W}pict"):
            return True
        # Keep non-text elements (e.g. formulas/embedded objects).
        for node in el.iter():
            tag = getattr(node, "tag", "")
            if not isinstance(tag, str):
                continue
            local = etree.QName(tag).localname
            if local in {"oMath", "oMathPara", "object", "oleObject"}:
                return True
        return False

    @staticmethod
    def _count_manual_break_signatures(
        doc: Document,
        *,
        start_index: int,
        end_index: int,
        target_indices: set[int] | None,
    ) -> int:
        """统计范围内含手动换行（w:br -> '\\n'）的段落数量。"""
        if start_index < 0:
            start_index = 0
        if end_index < start_index:
            return 0
        target_set = set(target_indices) if target_indices else None
        count = 0
        for idx, para in enumerate(doc.paragraphs):
            if idx < start_index or idx > end_index:
                continue
            if target_set is not None and idx not in target_set:
                continue
            text = para.text or ""
            if "\n" in text or "\r" in text:
                count += 1
        return count

    @staticmethod
    def _para_has_list_marker(para) -> bool:
        p_pr = para._element.find(qn("w:pPr"))
        if p_pr is None:
            return False
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            return False
        num_id = num_pr.find(qn("w:numId"))
        ilvl = num_pr.find(qn("w:ilvl"))
        return bool(
            (num_id is not None and num_id.get(qn("w:val")))
            or (ilvl is not None and ilvl.get(qn("w:val")))
        )

    @staticmethod
    def _resolve_list_marker_separator(config: SceneConfig) -> str:
        md_cfg = getattr(config, "md_cleanup", None)
        raw = str(getattr(md_cfg, "list_marker_separator", "tab") or "").strip().lower()
        if raw in {"half_space", "space", "halfwidth_space"}:
            return "half_space"
        if raw in {"full_space", "fullwidth_space"}:
            return "full_space"
        return "tab"

    @staticmethod
    def _resolve_ordered_list_style(config: SceneConfig) -> str:
        md_cfg = getattr(config, "md_cleanup", None)
        raw = str(getattr(md_cfg, "ordered_list_style", "mixed") or "").strip().lower()
        return normalize_ordered_list_style(raw)

    @staticmethod
    def _resolve_unordered_list_style(config: SceneConfig) -> str:
        md_cfg = getattr(config, "md_cleanup", None)
        raw = str(getattr(md_cfg, "unordered_list_style", "word_default") or "").strip().lower()
        return normalize_unordered_list_style(raw)

    @staticmethod
    def _list_separator_text(separator: str) -> str:
        if separator == "full_space":
            return "\u3000"
        if separator == "half_space":
            return " "
        return "\t"

    @staticmethod
    def _is_ordered_list_marker(marker: str) -> bool:
        s = str(marker or "").strip()
        if not s:
            return False
        return bool(re.match(r"^(?:\d+[.)、]|[（(]\d+[)）])$", s))

    @staticmethod
    def _separator_to_suff_value(separator: str) -> str:
        if separator == "half_space":
            return "space"
        if separator == "full_space":
            return "nothing"
        return "tab"

    @staticmethod
    def _strip_leading_full_space_from_para(para) -> bool:
        for run in para.runs:
            text = run.text or ""
            if not text:
                continue
            if text.startswith("\u3000"):
                run.text = text[1:]
                return True
            break
        return False

    def _set_numbering_suff_for_num_id(
        self,
        doc: Document,
        *,
        num_id: int,
        separator: str,
        ordered_list_style: str,
        unordered_list_style: str,
    ) -> bool:
        numbering_part = getattr(doc.part, "numbering_part", None)
        if numbering_part is None:
            return False
        numbering_el = getattr(numbering_part, "element", None)
        if numbering_el is None:
            return False

        target_num = None
        num_id_str = str(num_id)
        for num_el in numbering_el.findall(qn("w:num")):
            if str(num_el.get(qn("w:numId"), "")).strip() == num_id_str:
                target_num = num_el
                break
        if target_num is None:
            return False

        abstract_ref = target_num.find(qn("w:abstractNumId"))
        if abstract_ref is None:
            return False
        abstract_id = str(abstract_ref.get(qn("w:val"), "")).strip()
        if not abstract_id:
            return False

        abstract_num = None
        for abs_el in numbering_el.findall(qn("w:abstractNum")):
            if str(abs_el.get(qn("w:abstractNumId"), "")).strip() == abstract_id:
                abstract_num = abs_el
                break
        if abstract_num is None:
            return False

        desired = self._separator_to_suff_value(separator)
        changed = False
        levels = abstract_num.findall(qn("w:lvl"))
        if not levels:
            return False

        level0 = None
        for lvl in levels:
            if str(lvl.get(qn("w:ilvl")) or "").strip() == "0":
                level0 = lvl
                break
        if level0 is None:
            level0 = levels[0]
        fmt0 = level0.find(qn("w:numFmt"))
        fmt0_val = str(fmt0.get(qn("w:val")) or "").strip().lower() if fmt0 is not None else ""
        is_bullet = fmt0_val == "bullet"

        for lvl in levels:
            ilvl_raw = str(lvl.get(qn("w:ilvl")) or "").strip()
            try:
                ilvl = int(ilvl_raw) if ilvl_raw else 0
            except (TypeError, ValueError):
                ilvl = 0
            ilvl = max(0, min(ilvl, 8))

            fmt = lvl.find(qn("w:numFmt"))
            if fmt is None:
                fmt = etree.SubElement(lvl, qn("w:numFmt"))
                changed = True
            lvl_text = lvl.find(qn("w:lvlText"))
            if lvl_text is None:
                lvl_text = etree.SubElement(lvl, qn("w:lvlText"))
                changed = True

            if is_bullet:
                bullet_char, ascii_font, hansi_font = get_unordered_level_definition(
                    unordered_list_style, ilvl
                )
                if str(fmt.get(qn("w:val")) or "").strip().lower() != "bullet":
                    fmt.set(qn("w:val"), "bullet")
                    changed = True
                if str(lvl_text.get(qn("w:val")) or "") != bullet_char:
                    lvl_text.set(qn("w:val"), bullet_char)
                    changed = True
                rpr = lvl.find(qn("w:rPr"))
                if rpr is None:
                    rpr = etree.SubElement(lvl, qn("w:rPr"))
                    changed = True
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = etree.SubElement(rpr, qn("w:rFonts"))
                    changed = True
                if str(rfonts.get(qn("w:ascii")) or "") != ascii_font:
                    rfonts.set(qn("w:ascii"), ascii_font)
                    changed = True
                if str(rfonts.get(qn("w:hAnsi")) or "") != hansi_font:
                    rfonts.set(qn("w:hAnsi"), hansi_font)
                    changed = True
                if str(rfonts.get(qn("w:hint")) or "").strip().lower() != "default":
                    rfonts.set(qn("w:hint"), "default")
                    changed = True
            else:
                num_fmt_val, lvl_text_val = get_ordered_level_format(
                    ordered_list_style, ilvl
                )
                if str(fmt.get(qn("w:val")) or "").strip().lower() != num_fmt_val.lower():
                    fmt.set(qn("w:val"), num_fmt_val)
                    changed = True
                if str(lvl_text.get(qn("w:val")) or "") != lvl_text_val:
                    lvl_text.set(qn("w:val"), lvl_text_val)
                    changed = True

            suff = lvl.find(qn("w:suff"))
            if suff is None:
                suff = etree.SubElement(lvl, qn("w:suff"))
                changed = True
            current = str(suff.get(qn("w:val"), "")).strip().lower()
            if current != desired:
                suff.set(qn("w:val"), desired)
                changed = True
        return changed

    def _normalize_existing_word_lists(
        self,
        doc: Document,
        *,
        marker_separator: str,
        ordered_list_style: str,
        unordered_list_style: str,
        start_index: int,
        end_index: int,
        target_indices: set[int] | None,
    ) -> tuple[int, int]:
        if start_index < 0:
            start_index = 0
        if end_index < start_index:
            return 0, 0

        changed_num_ids: set[int] = set()
        changed_prefix_runs = 0
        target_set = set(target_indices) if target_indices else None

        for idx, para in enumerate(doc.paragraphs):
            if idx < start_index or idx > end_index:
                continue
            if target_set is not None and idx not in target_set:
                continue
            marker_values = self._get_list_marker_values(para)
            if not marker_values:
                continue
            num_id, _ = marker_values
            if self._set_numbering_suff_for_num_id(
                doc,
                num_id=num_id,
                separator=marker_separator,
                ordered_list_style=ordered_list_style,
                unordered_list_style=unordered_list_style,
            ):
                changed_num_ids.add(num_id)

            if marker_separator == "full_space":
                if not (para.text or "").startswith("\u3000"):
                    self._prepend_text_to_para(
                        para,
                        self._list_separator_text(marker_separator),
                    )
                    changed_prefix_runs += 1
            else:
                if self._strip_leading_full_space_from_para(para):
                    changed_prefix_runs += 1

        return len(changed_num_ids), changed_prefix_runs

    @staticmethod
    def _prepend_text_to_para(para, prefix: str) -> None:
        if not prefix:
            return
        if para.runs:
            para.runs[0].text = f"{prefix}{para.runs[0].text or ''}"
            return
        para.add_run(prefix)

    @staticmethod
    def _get_list_marker_values(para) -> tuple[int, int] | None:
        p_pr = para._element.find(qn("w:pPr"))
        if p_pr is None:
            return None
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            return None
        num_id = num_pr.find(qn("w:numId"))
        if num_id is None:
            return None
        num_id_val = num_id.get(qn("w:val"))
        if not num_id_val:
            return None
        ilvl = num_pr.find(qn("w:ilvl"))
        ilvl_val = ilvl.get(qn("w:val")) if ilvl is not None else "0"
        try:
            parsed_num_id = int(str(num_id_val).strip())
            parsed_ilvl = int(str(ilvl_val).strip()) if ilvl_val is not None else 0
        except (TypeError, ValueError):
            return None
        if parsed_num_id < 0:
            return None
        return parsed_num_id, max(0, parsed_ilvl)

    @staticmethod
    def _looks_like_non_list_block_start(text: str) -> bool:
        s = (text or "").strip()
        if not s:
            return False
        if s.startswith(("#", "```", "~~~", "$$", ">", "|")):
            return True
        if re.match(r"^(?:---+|\*\*\*+|___+)\s*$", s):
            return True
        if re.match(r"^(\s*)([-*+])\s+\S", s):
            return True
        if re.match(r"^(\s*)(?:\d+[.)、]|[（(]\d+[)）])\s*\S", s):
            return True
        if re.match(r"^\[\^([A-Za-z0-9_-]+)\]:\s+\S", s):
            return True
        return False

    def _normalize_orphan_list_markers(
        self,
        doc: Document,
        config: SceneConfig,
        *,
        start_index: int,
        end_index: int,
        target_indices: set[int] | None,
    ) -> tuple[int, int]:
        """修复“空列表符号段落 + 正文段落”错位，避免出现孤立圆点行。"""
        if start_index < 0:
            start_index = 0
        if end_index < start_index:
            return 0, 0

        marker_separator = self._resolve_list_marker_separator(config)
        target_set = set(target_indices) if target_indices else None
        promoted = 0
        removed = 0
        i = 0
        while i < len(doc.paragraphs):
            if i < start_index or i > end_index:
                i += 1
                continue
            if target_set is not None and i not in target_set:
                i += 1
                continue
            para = doc.paragraphs[i]
            is_empty_list_marker = (
                self._para_has_list_marker(para)
                and not (para.text or "").strip()
                and not self._para_has_content(para)
            )
            if not is_empty_list_marker:
                i += 1
                continue

            marker_values = self._get_list_marker_values(para)
            target_idx = i + 1
            while target_idx < len(doc.paragraphs):
                target_para = doc.paragraphs[target_idx]
                if (target_para.text or "").strip() or self._para_has_content(target_para):
                    break
                target_idx += 1

            if target_idx > end_index:
                i += 1
                continue
            if target_set is not None and target_idx < len(doc.paragraphs) and target_idx not in target_set:
                i += 1
                continue

            if marker_values and target_idx < len(doc.paragraphs):
                target_para = doc.paragraphs[target_idx]
                target_text = (target_para.text or "").strip()
                target_style_name = target_para.style.name if target_para.style else ""
                if (
                    target_text
                    and not self._para_has_list_marker(target_para)
                    and not detect_level_by_style_name(config, target_style_name)
                    and not self._looks_like_non_list_block_start(target_text)
                ):
                    apply_num_pr(target_para, marker_values[0], marker_values[1])
                    if marker_separator == "full_space" and not target_text.startswith("\u3000"):
                        self._prepend_text_to_para(
                            target_para,
                            self._list_separator_text(marker_separator),
                        )
                    promoted += 1

            para._element.getparent().remove(para._element)
            removed += 1
            if i <= end_index:
                end_index -= 1
            if target_set is not None:
                shifted: set[int] = set()
                for idx in target_set:
                    if idx < i:
                        shifted.add(idx)
                    elif idx > i:
                        shifted.add(idx - 1)
                target_set = shifted

        return promoted, removed

    def _apply_hr(self, doc, block, config, tracker, ctx):
        """水平线：删除 --- 段落"""
        idx = block.source_para_indices[0]
        if idx >= len(doc.paragraphs):
            return
        para = doc.paragraphs[idx]
        para._element.getparent().remove(para._element)
        tracker.record(
            rule_name=self.name, target=f"段落 #{idx}",
            section="body", change_type="text",
            before="---", after="(已删除)",
            paragraph_index=idx)

    def _apply_latex_block(self, doc, block, config, tracker, ctx):
        """LaTeX 块：转换为可编辑公式对象（优先），失败时保留纯文本。"""
        indices = block.source_para_indices
        if not indices:
            return

        first_idx = indices[0]
        if first_idx >= len(doc.paragraphs):
            return

        first_para = doc.paragraphs[first_idx]
        raw_expr = (block.raw_text or "").strip()
        converted = False

        if raw_expr:
            node = FormulaNode(
                kind="text",
                payload={"latex": raw_expr},
                source_type="latex",
                confidence=0.95,
                warnings=[],
            )
            outcome = convert_formula_node(node, "word_native", block=True)
            if outcome.success and outcome.omml_element is not None:
                converted = _replace_paragraph_with_omml(first_para, outcome.omml_element)

        if not converted:
            # Keep visible source text when conversion fails, avoid silent data loss.
            self._replace_text(first_para, raw_expr)

        for pi in sorted(indices[1:], reverse=True):
            if pi >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[pi]
            para._element.getparent().remove(para._element)

        tracker.record(
            rule_name=self.name,
            target=f"段落 #{first_idx}",
            section="body",
            change_type="text",
            before=f"$$...$$ ({len(indices)} 行)",
            after=(
                "→ 可编辑公式对象"
                if converted else
                "→ LaTeX 文本（公式对象转换失败，已保留表达式）"
            ),
            paragraph_index=first_idx,
        )

    def _apply_list_item(self, doc, block, config, tracker, ctx):
        """列表项：剥离标记，应用 Word 原生编号"""
        idx = block.source_para_indices[0]
        if idx >= len(doc.paragraphs):
            return
        para = doc.paragraphs[idx]
        old_text = para.text

        base_name, base_size = ctx.get("base_font_name"), ctx.get("base_font_size")
        marker_separator = str(ctx.get("list_marker_separator", "tab") or "tab")
        spans = list(block.spans or [])
        if marker_separator == "full_space":
            needs_prefix = True
            if spans and spans[0].type == InlineType.TEXT and str(spans[0].text or "").startswith("\u3000"):
                needs_prefix = False
            if needs_prefix:
                spans = [InlineSpan(InlineType.TEXT, self._list_separator_text(marker_separator))] + spans

        write_spans(para, spans,
                    doc=doc,
                    footnote_defs=ctx.get("footnote_defs"),
                    base_font_name=base_name, base_font_size=base_size,
                    image_base_path=ctx.get("md_base_path"),
                    prefer_real_image=True)
        is_ordered = self._is_ordered_list_marker(block.list_marker)
        num_id = ctx["decimal_num_id"] if is_ordered else ctx["bullet_num_id"]
        apply_num_pr(para, num_id, block.list_level)

        tracker.record(
            rule_name=self.name, target=f"段落 #{idx}",
            section="body", change_type="text",
            before=old_text[:60],
            after=f"→ list (level {block.list_level}): {block.raw_text[:40]}",
            paragraph_index=idx)

    def _apply_task_list(self, doc, block, config, tracker, ctx):
        """任务清单：剥离 - [x] 标记，添加勾选符号"""
        idx = block.source_para_indices[0]
        if idx >= len(doc.paragraphs):
            return
        para = doc.paragraphs[idx]
        old_text = para.text
        marker_separator = str(ctx.get("list_marker_separator", "tab") or "tab")
        sep = self._list_separator_text(marker_separator)
        prefix = f"{'☑' if block.checked else '☐'}{sep}"
        # 清除现有 runs
        for run in para.runs:
            run._element.getparent().remove(run._element)
        para.add_run(prefix)
        write_spans(para, block.spans,
                    doc=doc,
                    footnote_defs=ctx.get("footnote_defs"),
                    clear_existing=False,
                    image_base_path=ctx.get("md_base_path"),
                    prefer_real_image=True)
        # 任务列表使用显式缩进体现嵌套层级。
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.left_indent = Cm(0.75 * max(0, block.list_level))
        tracker.record(
            rule_name=self.name, target=f"段落 #{idx}",
            section="body", change_type="text",
            before=old_text[:60],
            after=f"→ task: {prefix}{block.raw_text[:30]}",
            paragraph_index=idx)

    def _apply_inline(self, doc, block, config, tracker, ctx):
        """普通段落：处理行内 Markdown 格式 + 多段落合并"""
        indices = block.source_para_indices
        idx = indices[0]
        if idx >= len(doc.paragraphs):
            return

        base_name, base_size = ctx.get("base_font_name"), ctx.get("base_font_size")
        # 关键修正：
        # 多个 source 段落时不再压缩成“单段 + 手动换行(w:br)”，
        # 直接逐段渲染，保留真正段落边界（Enter）。
        if len(indices) > 1:
            for para_idx in indices:
                if para_idx >= len(doc.paragraphs):
                    continue
                para = doc.paragraphs[para_idx]
                spans = self._resolve_inline_spans(para.text or "")
                write_spans(
                    para, spans,
                    doc=doc,
                    footnote_defs=ctx.get("footnote_defs"),
                    base_font_name=base_name, base_font_size=base_size,
                    image_base_path=ctx.get("md_base_path"),
                    prefer_real_image=True)
            after_desc = f"→ inline formatted ({len(indices)} para preserved)"
            before_text = doc.paragraphs[idx].text[:60]
        else:
            para = doc.paragraphs[idx]
            old_text = para.text
            spans = self._resolve_inline_spans(
                para.text or "",
                fallback_spans=block.spans,
            )
            write_spans(para, spans,
                        doc=doc,
                        footnote_defs=ctx.get("footnote_defs"),
                        base_font_name=base_name, base_font_size=base_size,
                        image_base_path=ctx.get("md_base_path"),
                        prefer_real_image=True)
            before_text = old_text[:60]
            after_desc = "→ inline formatted (1 para)"

        tracker.record(
            rule_name=self.name, target=f"段落 #{idx}",
            section="body", change_type="format",
            before=before_text,
            after=after_desc,
            paragraph_index=idx)

    # ── 辅助方法 ──

    @staticmethod
    def _looks_like_markdown_sensitive_formula_text(text: str) -> bool:
        value = str(text or "").strip()
        if not value:
            return False
        if not re.search(r"[*_`]", value):
            return False
        if _RE_LATEX_COMMAND.search(value):
            return True
        if re.search(r"[A-Za-zΑ-Ωα-ω0-9\)\]}]\s*[*_]\s*[A-Za-zΑ-Ωα-ω0-9\(\[{\\]", value):
            return True
        if re.search(r"[A-Za-zΑ-Ωα-ω][A-Za-z0-9]*\([^)]*[*_`][^)]*\)", value):
            return True
        return False

    @staticmethod
    def _should_protect_formula_occurrence(occ) -> bool:
        source_type = str(getattr(occ, "source_type", "") or "").strip().lower()
        if source_type in {"word_native", "mathtype", "old_equation", "ole_equation"}:
            return True
        if source_type != "latex":
            if source_type in {"plain_text", "ocr_fragment", "unicode_text"}:
                node = getattr(occ, "node", None)
                payload = getattr(node, "payload", {}) if node is not None else {}
                source_text = (
                    getattr(occ, "source_text", "")
                    or payload.get("text")
                    or payload.get("latex")
                    or ""
                )
                return MdCleanupRule._looks_like_markdown_sensitive_formula_text(source_text)
            return False
        if bool(getattr(occ, "is_formula_only", False)):
            return True
        indices = list(getattr(occ, "source_paragraph_indices", []) or [])
        return len(indices) > 1

    @staticmethod
    def _replace_text(para, new_text: str):
        """替换段落文本，保留第一个 run 的格式"""
        if not para.runs:
            para.text = new_text
            return
        first_run = para.runs[0]
        for run in para.runs[1:]:
            run.text = ""
        first_run.text = new_text

    @staticmethod
    def _get_base_font(para):
        """从段落第一个 run 提取基础字体信息"""
        if para.runs:
            r = para.runs[0]
            return r.font.name, r.font.size
        return None, None

    @staticmethod
    def _find_keyword_index(doc: Document, keyword: str) -> int | None:
        """在文档中查找包含关键字的段落索引"""
        kw = re.sub(r'\s+', '', keyword or "")
        if not kw:
            return None
        for i, para in enumerate(doc.paragraphs):
            text = re.sub(r'\s+', '', (para.text or "").strip())
            if kw in text:
                return i
        return None

    @staticmethod
    def _find_page_start_index(doc: Document, target_page: int) -> int | None:
        """根据目标页码找到该页出现的第一个段落索引。"""
        return page_number_to_start_paragraph_index(doc, target_page)

    def _split_manual_break_paragraphs(
            self,
            doc: Document,
            start_idx: int,
            end_idx: int,
            *,
            target_indices: set[int] | None) -> tuple[int, set[int] | None]:
        """将段落中的 '\\n' 手动换行拆分为多个真实段落。"""
        if start_idx is None or end_idx is None:
            return 0, target_indices
        if start_idx < 0 or end_idx < start_idx:
            return 0, target_indices

        inserted = 0
        target_filter = set(target_indices) if target_indices is not None else None
        updated_target_indices = set(target_indices) if target_indices is not None else None
        end_idx = min(end_idx, len(doc.paragraphs) - 1)
        for idx in range(end_idx, start_idx - 1, -1):
            if idx >= len(doc.paragraphs):
                continue
            if target_filter is not None and idx not in target_filter:
                continue
            para = doc.paragraphs[idx]
            if self._para_has_content(para):
                continue
            text = para.text or ""
            if "\n" not in text and "\r" not in text:
                continue

            parts = re.split(r'\r?\n', text)
            if len(parts) <= 1:
                continue

            src_style = para.style
            src_pf = para.paragraph_format
            self._replace_text(para, parts[0])

            anchor = para
            inserted_here = 0
            for part in parts[1:]:
                new_para = self._insert_paragraph_after(anchor, part)
                try:
                    new_para.style = src_style
                except Exception:
                    pass
                new_pf = new_para.paragraph_format
                new_pf.alignment = src_pf.alignment
                new_pf.first_line_indent = src_pf.first_line_indent
                new_pf.left_indent = src_pf.left_indent
                new_pf.right_indent = src_pf.right_indent
                new_pf.space_before = src_pf.space_before
                new_pf.space_after = src_pf.space_after
                new_pf.line_spacing = src_pf.line_spacing
                new_pf.line_spacing_rule = src_pf.line_spacing_rule
                anchor = new_para
                inserted += 1
                inserted_here += 1

            if updated_target_indices is not None and inserted_here > 0:
                shifted: set[int] = set()
                for target_idx in updated_target_indices:
                    if target_idx > idx:
                        shifted.add(target_idx + inserted_here)
                    elif target_idx == idx:
                        shifted.add(target_idx)
                        shifted.update(
                            range(target_idx + 1, target_idx + inserted_here + 1)
                        )
                    else:
                        shifted.add(target_idx)
                updated_target_indices = shifted

        return inserted, updated_target_indices

    @staticmethod
    def _contains_latex_commands(text: str) -> bool:
        """检查文本是否包含LaTeX命令（避免被parse_inline破坏）"""
        if not text:
            return False
        # 检查是否包含LaTeX命令模式：\command 或 \begin{...}
        return bool(_RE_LATEX_COMMAND.search(text))

    @staticmethod
    def _resolve_inline_spans(text: str, fallback_spans: list[InlineSpan] | None = None) -> list[InlineSpan]:
        """解析内联 Markdown 格式，如果失败则使用 fallback_spans"""
        if not text:
            return fallback_spans or []
        spans = parse_inline(text)
        if not spans and fallback_spans:
            return fallback_spans
        return spans

    @staticmethod
    def _insert_paragraph_after(para, text: str = ""):
        """在指定段落后插入新段落。"""
        new_p = OxmlElement("w:p")
        para._element.addnext(new_p)
        new_para = Paragraph(new_p, para._parent)
        if text:
            new_para.add_run(text)
        return new_para

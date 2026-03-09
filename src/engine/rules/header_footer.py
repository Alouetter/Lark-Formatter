"""页眉页脚自动更新规则

规则：
- 页眉：STYLEREF 域代码自动跟随章节标题；Abstract 部分用 Arial 字体
- 页码：正文用阿拉伯数字，摘要/目录用罗马数字
- 封面不加页眉页码
"""

import re
from lxml import etree
from docx import Document
from docx.oxml import OxmlElement
from src.engine.rules.base import BaseRule
from src.engine.change_tracker import ChangeTracker
from src.scene.heading_model import (
    get_header_back_types,
    get_header_front_text,
    get_non_numbered_heading_style_name,
    get_non_numbered_title_sections,
)
from src.scene.schema import SceneConfig, StyleConfig

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
_RE_CAPTION_LINE = re.compile(r"^(图|表|Figure|Table|Fig\.?)\s*\d", re.IGNORECASE)
_COVER_LIKE_SECTION_TYPES = {"cover", "unknown"}


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


# ── 基础工具 ──

def _apply_run_style(r_el, sc: StyleConfig):
    """给 run 元素应用字体和字号"""
    rPr = r_el.find(_w("rPr"))
    if rPr is None:
        rPr = etree.SubElement(r_el, _w("rPr"))
        r_el.insert(0, rPr)
    rf = rPr.find(_w("rFonts"))
    if rf is None:
        rf = etree.SubElement(rPr, _w("rFonts"))
    rf.set(_w("ascii"), sc.font_en)
    rf.set(_w("hAnsi"), sc.font_en)
    rf.set(_w("eastAsia"), sc.font_cn)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rf.attrib.pop(_w(attr), None)
    half_pt = str(int(sc.size_pt * 2))
    for tag in ("sz", "szCs"):
        el = rPr.find(_w(tag))
        if el is None:
            el = etree.SubElement(rPr, _w(tag))
        el.set(_w("val"), half_pt)


def _set_para_center(p_el):
    """居中、缩进0、间距0、单倍行距"""
    pPr = p_el.find(_w("pPr"))
    if pPr is None:
        pPr = etree.SubElement(p_el, _w("pPr"))
        p_el.insert(0, pPr)
    jc = pPr.find(_w("jc"))
    if jc is None:
        jc = etree.SubElement(pPr, _w("jc"))
    jc.set(_w("val"), "center")
    # 缩进 0
    ind = pPr.find(_w("ind"))
    if ind is None:
        ind = etree.SubElement(pPr, _w("ind"))
    ind.set(_w("firstLine"), "0")
    ind.set(_w("left"), "0")
    # 间距 0，单倍行距
    spacing = pPr.find(_w("spacing"))
    if spacing is None:
        spacing = etree.SubElement(pPr, _w("spacing"))
    spacing.set(_w("before"), "0")
    spacing.set(_w("after"), "0")
    spacing.set(_w("line"), "240")
    spacing.set(_w("lineRule"), "auto")


def _set_header_border(header_el, enable: bool):
    """设置或移除页眉段落的底部横线"""
    for p in header_el.findall(_w("p")):
        pPr = p.find(_w("pPr"))
        if pPr is None:
            pPr = etree.SubElement(p, _w("pPr"))
            p.insert(0, pPr)
        pBdr = pPr.find(_w("pBdr"))
        if enable:
            if pBdr is None:
                pBdr = etree.SubElement(pPr, _w("pBdr"))
            bottom = pBdr.find(_w("bottom"))
            if bottom is None:
                bottom = etree.SubElement(pBdr, _w("bottom"))
            bottom.set(_w("val"), "single")
            bottom.set(_w("sz"), "4")
            bottom.set(_w("space"), "1")
            bottom.set(_w("color"), "auto")
        else:
            if pBdr is not None:
                bottom = pBdr.find(_w("bottom"))
                if bottom is not None:
                    pBdr.remove(bottom)


def _clear_element(el):
    for child in list(el):
        el.remove(child)


def _make_styled_run(parent, sc, text=None):
    """创建带样式的 run，可选文本"""
    r = etree.SubElement(parent, _w("r"))
    _apply_run_style(r, sc)
    if text is not None:
        t = etree.SubElement(r, _w("t"))
        t.text = text
        t.set(_XML_SPACE, "preserve")
    return r


def _make_field_run(parent, sc, fld_type):
    """创建 fldChar run"""
    r = _make_styled_run(parent, sc)
    fc = etree.SubElement(r, _w("fldChar"))
    fc.set(_w("fldCharType"), fld_type)
    return r


def _make_instr_run(parent, sc, instr_text):
    """创建 instrText run"""
    r = _make_styled_run(parent, sc)
    it = etree.SubElement(r, _w("instrText"))
    it.text = instr_text
    it.set(_XML_SPACE, "preserve")
    return r


# ── 分节符插入 ──

def _remove_page_breaks(p_el):
    """移除段落中的分页符（w:br type="page" 和 pageBreakBefore）"""
    for br in list(p_el.iter(_w("br"))):
        if br.get(_w("type")) == "page":
            br.getparent().remove(br)
    pPr = p_el.find(_w("pPr"))
    if pPr is not None:
        pbb = pPr.find(_w("pageBreakBefore"))
        if pbb is not None:
            pPr.remove(pbb)


def _get_existing_break_indices(doc):
    """返回已有分节符所在的段落索引集合"""
    breaks = set()
    para_idx = 0
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            pPr = child.find(_w("pPr"))
            if pPr is not None and pPr.find(_w("sectPr")) is not None:
                breaks.add(para_idx)
            para_idx += 1
    return breaks


def _has_sectpr(para) -> bool:
    ppr = para._element.find(_w("pPr"))
    return ppr is not None and ppr.find(_w("sectPr")) is not None


def _has_table_between_paragraphs(doc, start_idx: int, end_idx: int) -> bool:
    """Return True when any body-level table exists between two paragraphs."""
    if start_idx < 0 or end_idx <= start_idx:
        return False
    if start_idx >= len(doc.paragraphs) or end_idx >= len(doc.paragraphs):
        return False

    start_el = doc.paragraphs[start_idx]._element
    end_el = doc.paragraphs[end_idx]._element
    curr = start_el.getnext()
    while curr is not None and curr is not end_el:
        tag = curr.tag.split("}")[-1] if "}" in curr.tag else curr.tag
        if tag == "tbl":
            return True
        curr = curr.getnext()
    return False


def _remove_caption_table_breaks(doc) -> int:
    """Clean legacy sectPr accidentally attached to table captions."""
    removed = 0
    for para in doc.paragraphs:
        if not _has_sectpr(para):
            continue
        text = (para.text or "").strip()
        if not text or not _RE_CAPTION_LINE.match(text):
            continue
        nxt = para._element.getnext()
        if nxt is None:
            continue
        tag = nxt.tag.split("}")[-1] if "}" in nxt.tag else nxt.tag
        if tag != "tbl":
            continue
        ppr = para._element.find(_w("pPr"))
        if ppr is None:
            continue
        sect_pr = ppr.find(_w("sectPr"))
        if sect_pr is None:
            continue
        ppr.remove(sect_pr)
        removed += 1
    return removed


def _nearby_break_indices(doc, boundary_para: int, lookback: int = 3) -> list[int]:
    """返回 boundary 前 lookback 个段落中已有分节符的索引。"""
    start = max(0, boundary_para - lookback)
    hits = []
    for idx in range(start, boundary_para):
        if idx >= len(doc.paragraphs):
            break
        if _has_sectpr(doc.paragraphs[idx]):
            hits.append(idx)
    return hits


def _prune_redundant_blank_breaks(doc, break_indices: list[int], keep_idx: int) -> None:
    """清理同一边界附近冗余分节符（仅清理空白段落上的冗余项）。"""
    for idx in break_indices:
        if idx == keep_idx:
            continue
        if idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        # 只清理空白段落，避免误删正文中有意分节。
        if (para.text or "").strip():
            continue
        ppr = para._element.find(_w("pPr"))
        if ppr is None:
            continue
        sp = ppr.find(_w("sectPr"))
        if sp is not None:
            ppr.remove(sp)


def _copy_sect_pr(src_sect):
    """从 docx section 复制页面尺寸/边距到新 sectPr"""
    import copy
    new_sp = etree.Element(_w("sectPr"))
    # 复制 pgSz
    old_sp = src_sect._sectPr
    for tag in ("pgSz", "pgMar"):
        el = old_sp.find(_w(tag))
        if el is not None:
            new_sp.append(copy.deepcopy(el))
    # 设为 nextPage 分节
    typ = etree.SubElement(new_sp, _w("type"))
    typ.set(_w("val"), "nextPage")
    return new_sp


def _is_heading_like_para(para) -> bool:
    """启发式判断段落是否像章节/小节标题。"""
    style_name = para.style.name if para.style else ""
    if (style_name.startswith("Heading")
            or style_name.startswith("heading")
            or style_name.startswith("标题")
            or style_name == "一级标题"):
        return True

    txt = (para.text or "").strip()
    if not txt:
        return False
    if re.match(r'^第[一二三四五六七八九十百零\d]+[章节篇]', txt):
        return True
    # 数字层级标题（如 1.2 / 3.1.4）才视为子标题；避免把“2024年…”等正文误判
    if (re.match(r'^\d+\.\d+(?:\.\d+){0,3}(?:[\.、．])?\s*\S', txt)
            and not re.search(r'[。！？；;，,]', txt)):
        return True
    return False


def _break_applies_to_boundary(doc, break_idx: int, boundary_para: int) -> bool:
    """判断 break_idx 处分节符是否真正作用于 boundary 边界。

    仅当 break 与 boundary 之间全为空白段落时，认为该分节符已满足该边界。
    """
    if break_idx < 0 or break_idx >= boundary_para:
        return False
    if _has_table_between_paragraphs(doc, break_idx, boundary_para):
        return False
    for idx in range(break_idx + 1, boundary_para):
        if idx >= len(doc.paragraphs):
            break
        if (doc.paragraphs[idx].text or "").strip():
            return False
    return True


def _effective_break_indices_for_boundary(
        doc, boundary_para: int, lookback: int = 8) -> list[int]:
    """返回真正作用于该 boundary 的已有分节符索引列表。"""
    start = max(0, boundary_para - lookback)
    hits = []
    for idx in range(start, boundary_para):
        if idx >= len(doc.paragraphs):
            break
        if not _has_sectpr(doc.paragraphs[idx]):
            continue
        if _break_applies_to_boundary(doc, idx, boundary_para):
            hits.append(idx)
    return hits


def _remove_unexpected_internal_breaks(doc, doc_tree, expected_prev_indices: set[int]) -> int:
    """清理致谢/简历分区内的错位旧分节符。

    仅删除以下分节符：
    1) 不在本轮目标边界集合 expected_prev_indices 中；
    2) 分节前后仍属于同一逻辑分区（非真正分区边界）；
    3) 分区类型属于 acknowledgment/resume。
    """
    if not doc_tree:
        return 0

    removed = 0
    for idx in sorted(_get_existing_break_indices(doc)):
        if idx in expected_prev_indices:
            continue
        if idx < 0 or idx + 1 >= len(doc.paragraphs):
            continue

        sec_a = doc_tree.get_section_for_paragraph(idx)
        sec_b = doc_tree.get_section_for_paragraph(idx + 1)
        if sec_a != sec_b:
            continue
        if sec_a not in {"acknowledgment", "resume"}:
            continue

        ppr = doc.paragraphs[idx]._element.find(_w("pPr"))
        if ppr is None:
            continue
        sp = ppr.find(_w("sectPr"))
        if sp is None:
            continue
        ppr.remove(sp)
        removed += 1

    return removed


def _norm_no_space(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip()


def _find_ack_range_fallback(doc) -> tuple[int, int] | None:
    """当 doc_tree 未识别 acknowledgment 时，按文本锚点回退识别致谢区间。"""
    ack_idx = None
    for i, para in enumerate(doc.paragraphs):
        norm = _norm_no_space(para.text)
        if norm in ("致谢", "致谢。", "致谢："):
            ack_idx = i
            break
    if ack_idx is None:
        return None

    # 结束边界：遇到后置分区显式锚点或文末
    end = len(doc.paragraphs) - 1
    for j in range(ack_idx + 1, len(doc.paragraphs)):
        norm = _norm_no_space(doc.paragraphs[j].text)
        if norm in ("参考文献", "附录", "简历", "个人简历", "在学期间"):
            end = j - 1
            break
    if end < ack_idx:
        end = ack_idx
    return ack_idx, end


def _collect_ack_internal_boundaries(doc, doc_tree) -> list[int]:
    """致谢分区内部：若正文后出现子章节标题，在标题前建立分节边界。"""
    if not doc_tree:
        return []
    sec = doc_tree.get_section("acknowledgment") if doc_tree else None
    if sec:
        start_idx, end_idx = sec.start_index, sec.end_index
    else:
        fallback = _find_ack_range_fallback(doc)
        if not fallback:
            return []
        start_idx, end_idx = fallback

    boundaries = []
    seen_ack_text = False

    # start_index 通常是“致谢”标题本身，从下一段开始扫描
    for idx in range(start_idx + 1, min(end_idx, len(doc.paragraphs) - 1) + 1):
        para = doc.paragraphs[idx]
        txt = (para.text or "").strip()
        if not txt:
            continue

        if _is_heading_like_para(para):
            # 只有在“已出现致谢正文”后才切分，避免把标题行后直接小标题误切
            if seen_ack_text:
                boundaries.append(idx)
                break
            continue

        seen_ack_text = True

    return boundaries


def _collect_resume_internal_boundaries(doc, doc_tree) -> list[int]:
    """简历分区内部：成果标题前建立分节边界。"""
    if not doc_tree:
        return []
    sec = doc_tree.get_section("resume")
    if not sec:
        return []

    start = max(0, sec.start_index + 1)
    end = min(sec.end_index, len(doc.paragraphs) - 1)
    boundaries = []
    for idx in range(start, end + 1):
        txt = (doc.paragraphs[idx].text or "").strip()
        if not txt:
            continue
        norm = _norm_no_space(txt)
        if not norm.startswith("在学期间发表的学术论文与研究成果"):
            continue
        # 仅在后续仍有内容时切分，避免末尾空标题产生无意义分节
        has_following = False
        for j in range(idx + 1, len(doc.paragraphs)):
            if (doc.paragraphs[j].text or "").strip():
                has_following = True
                break
        if has_following:
            boundaries.append(idx)
        break
    return boundaries


def _ensure_section_breaks(doc, doc_tree):
    """在逻辑分区边界插入分节符（如果尚不存在）。

    确保 cover、每个前置分区、body 各自拥有独立的 DOCX section。
    """
    if not doc_tree or not doc_tree.sections:
        return

    existing = _get_existing_break_indices(doc)
    # 收集需要分节的边界：每个逻辑分区的 start_index
    # 在 start_index 的前一个段落末尾插入 sectPr
    boundaries = []
    for sec in doc_tree.sections:
        if sec.section_type == "cover":
            continue  # cover 从 0 开始，不需要在前面插入
        if sec.start_index > 0:
            boundaries.append(sec.start_index)

    # 特殊场景：致谢正文后跟随子章节标题，需要在标题前分节（下一页）
    ack_internal_boundaries = set(_collect_ack_internal_boundaries(doc, doc_tree))
    boundaries.extend(ack_internal_boundaries)
    resume_internal_boundaries = set(_collect_resume_internal_boundaries(doc, doc_tree))
    boundaries.extend(resume_internal_boundaries)

    # 去重 + 排序，从后往前插入避免索引偏移
    boundaries = sorted(set(boundaries), reverse=True)
    expected_prev_indices = {b - 1 for b in boundaries if b > 0}
    base_sect = doc.sections[0]

    for boundary_para in boundaries:
        if boundary_para <= 0 or boundary_para >= len(doc.paragraphs):
            continue
        prev_idx = boundary_para - 1
        if prev_idx < 0:
            continue
        boundary_el = doc.paragraphs[boundary_para]._element
        has_table_gap = _has_table_between_paragraphs(doc, prev_idx, boundary_para)
        # 所有边界统一使用“有效覆盖该边界”的判断，避免误用更早分节符。
        near = _effective_break_indices_for_boundary(
            doc, boundary_para, lookback=8
        )
        if near:
            keep_idx = max(near)
            _prune_redundant_blank_breaks(doc, near, keep_idx)
            continue
        # 已有分节符则跳过
        if not has_table_gap and prev_idx in existing:
            continue
        # 在前一段落的 pPr 中插入 sectPr
        if has_table_gap:
            prev_el = OxmlElement("w:p")
            boundary_el.addprevious(prev_el)
            pPr = etree.SubElement(prev_el, _w("pPr"))
        else:
            prev_el = doc.paragraphs[prev_idx]._element
            pPr = prev_el.find(_w("pPr"))
            if pPr is None:
                pPr = etree.SubElement(prev_el, _w("pPr"))
                prev_el.insert(0, pPr)
        # 已有 sectPr 则跳过
        if pPr.find(_w("sectPr")) is not None:
            continue
        pPr.append(_copy_sect_pr(base_sect))
        if not has_table_gap:
            existing.add(prev_idx)
        # 移除该段落已有的分页符（sectPr nextPage 已包含分页）
        _remove_page_breaks(prev_el)
        # 也移除下一段落的 pageBreakBefore
        _remove_page_breaks(boundary_el)

    # 清理旧版本遗留的“错位分节符”（常见于致谢/简历内部小标题后）
    _remove_unexpected_internal_breaks(doc, doc_tree, expected_prev_indices)
    _remove_caption_table_breaks(doc)


# ── DOCX section ↔ 逻辑分区映射 ──

def _map_docx_sections(doc, doc_tree):
    """将每个 DOCX section 映射到具体分区类型。"""
    if not doc_tree:
        return [(s, "body") for s in doc.sections]

    body_el = doc.element.body
    sec_breaks = []
    para_idx = 0
    docx_sec_idx = 0

    for child in body_el:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            pPr = child.find(_w("pPr"))
            if pPr is not None and pPr.find(_w("sectPr")) is not None:
                sec_breaks.append((para_idx, docx_sec_idx))
                docx_sec_idx += 1
            para_idx += 1

    total_paras = len(doc.paragraphs)
    ranges = []
    prev_start = 0
    for end_idx, _ in sec_breaks:
        ranges.append((prev_start, end_idx))
        prev_start = end_idx + 1
    ranges.append((prev_start, total_paras - 1))

    result = []
    for i, docx_sec in enumerate(doc.sections):
        if i >= len(ranges):
            result.append((docx_sec, "body"))
            continue
        start, end = ranges[i]
        mid = (start + end) // 2
        sec_type = doc_tree.get_section_for_paragraph(mid)
        result.append((docx_sec, sec_type))

    return result


# ── 页眉构建 ──

def _build_empty_header(header_el):
    """空页眉（封面用）"""
    _clear_element(header_el)
    etree.SubElement(header_el, _w("p"))


def _build_styleref_header(
        header_el, sc, *,
        style_ref: str = "标题 1",
        include_number: bool = True):
    """STYLEREF 域代码页眉，可配置引用样式与是否包含编号。"""
    _clear_element(header_el)
    p = etree.SubElement(header_el, _w("p"))
    _set_para_center(p)

    if include_number:
        # STYLEREF \n -> 章节编号（从 numPr 取）
        _make_field_run(p, sc, "begin")
        _make_instr_run(p, sc, f' STYLEREF "{style_ref}" \\n ')
        _make_field_run(p, sc, "separate")
        _make_styled_run(p, sc, text="1")
        _make_field_run(p, sc, "end")

    # STYLEREF -> 章节标题文本
    _make_field_run(p, sc, "begin")
    _make_instr_run(p, sc, f' STYLEREF "{style_ref}" ')
    _make_field_run(p, sc, "separate")
    _make_styled_run(p, sc, text="章节标题")
    _make_field_run(p, sc, "end")


def _build_text_header(header_el, text, sc):
    """静态文本页眉"""
    _clear_element(header_el)
    p = etree.SubElement(header_el, _w("p"))
    _set_para_center(p)
    _make_styled_run(p, sc, text=text)


# ── 页脚构建 ──

def _build_empty_footer(footer_el):
    """空页脚（封面用）"""
    _clear_element(footer_el)
    etree.SubElement(footer_el, _w("p"))


def _build_page_footer(footer_el, sc, num_format="decimal"):
    """PAGE 域代码页脚，支持 decimal / upperRoman / lowerRoman"""
    _clear_element(footer_el)
    p = etree.SubElement(footer_el, _w("p"))
    _set_para_center(p)

    _make_field_run(p, sc, "begin")
    if num_format == "decimal":
        _make_instr_run(p, sc, " PAGE ")
    else:
        _make_instr_run(p, sc, " PAGE \\* ROMAN ")
    _make_field_run(p, sc, "separate")
    _make_styled_run(p, sc, text="1")
    _make_field_run(p, sc, "end")


def _set_page_num_format(sect_pr, fmt="decimal", start=None):
    """设置 sectPr 中的页码格式和起始页号。

    fmt: "decimal" / "upperRoman" / "lowerRoman"
    """
    pgNumType = sect_pr.find(_w("pgNumType"))
    if pgNumType is None:
        pgNumType = etree.SubElement(sect_pr, _w("pgNumType"))
    pgNumType.set(_w("fmt"), fmt)
    if start is not None:
        pgNumType.set(_w("start"), str(start))
    else:
        # 显式移除 start 属性，确保页码从上一节继续而非重新开始
        pgNumType.attrib.pop(_w("start"), None)


class HeaderFooterRule(BaseRule):
    name = "header_footer"
    description = "自动更新页眉页脚"

    def apply(self, doc: Document, config: SceneConfig,
              tracker: ChangeTracker, context: dict) -> None:
        if not config.update_header and not config.update_page_number and not config.update_header_line:
            return

        doc_tree = context.get("doc_tree")
        _ensure_section_breaks(doc, doc_tree)
        sec_map = _map_docx_sections(doc, doc_tree)

        header_cn = config.styles.get("header_cn")
        header_en = config.styles.get("header_en")
        page_sc = config.styles.get("page_number")

        h_count = 0
        f_count = 0
        seen_front = False
        seen_body = False

        show_line = config.update_header_line
        front_header_text = get_header_front_text(config)
        front_types = set(front_header_text.keys())
        back_types = get_header_back_types(config)
        non_numbered_back_types = get_non_numbered_title_sections(config)
        non_numbered_heading_style = get_non_numbered_heading_style_name(config)

        for docx_sec, sec_type in sec_map:
            if config.update_header and header_cn:
                h_count += self._apply_header(
                    docx_sec, sec_type, header_cn, header_en,
                    show_line=show_line,
                    front_header_text=front_header_text,
                    non_numbered_back_types=non_numbered_back_types,
                    non_numbered_heading_style=non_numbered_heading_style,
                )
            elif show_line:
                # 仅设置横线，不重建页眉内容
                header = docx_sec.header
                header.is_linked_to_previous = False
                _set_header_border(header._element,
                                   sec_type not in _COVER_LIKE_SECTION_TYPES)

            if config.update_page_number and page_sc:
                is_front = sec_type in front_types
                is_body_or_back = (sec_type == "body" or sec_type in back_types)

                if is_front:
                    start = 1 if not seen_front else None
                    seen_front = True
                    f_count += self._apply_footer_front(
                        docx_sec, page_sc, start)
                elif is_body_or_back:
                    start = 1 if not seen_body else None
                    seen_body = True
                    f_count += self._apply_footer_body(
                        docx_sec, page_sc, start)
                else:
                    # cover
                    footer = docx_sec.footer
                    footer.is_linked_to_previous = False
                    _build_empty_footer(footer._element)
                    f_count += 1

        if h_count:
            tracker.record(
                rule_name=self.name, target=f"{h_count} 个页眉",
                section="global", change_type="format",
                before="(原页眉)", after="STYLEREF/静态 章节标题",
                paragraph_index=-1,
            )
        if f_count:
            tracker.record(
                rule_name=self.name, target=f"{f_count} 个页脚",
                section="global", change_type="format",
                before="(原页脚)", after="PAGE 域代码",
                paragraph_index=-1,
            )

    def _apply_header(
            self,
            docx_sec,
            sec_type,
            header_cn,
            header_en,
            *,
            show_line=False,
            front_header_text: dict[str, str] | None = None,
            non_numbered_back_types: set[str] | None = None,
            non_numbered_heading_style: str = "Heading 1 Unnumbered"):
        header = docx_sec.header
        header.is_linked_to_previous = False
        front_header_text = front_header_text or {}
        non_numbered_back_types = non_numbered_back_types or set()

        if sec_type in _COVER_LIKE_SECTION_TYPES:
            _build_empty_header(header._element)
            _set_header_border(header._element, False)
        elif sec_type in front_header_text:
            text = front_header_text[sec_type]
            sc = header_en if sec_type == "abstract_en" else header_cn
            _build_text_header(header._element, text, sc or header_cn)
            _set_header_border(header._element, show_line)
        elif sec_type in non_numbered_back_types:
            # 后置无编号一级标题：仅显示标题文本，不显示编号字段
            _build_styleref_header(
                header._element, header_cn,
                style_ref=non_numbered_heading_style,
                include_number=False,
            )
            _set_header_border(header._element, show_line)
        else:
            _build_styleref_header(header._element, header_cn)
            _set_header_border(header._element, show_line)
        return 1

    def _apply_footer_front(self, docx_sec, page_sc, start):
        footer = docx_sec.footer
        footer.is_linked_to_previous = False
        _build_page_footer(footer._element, page_sc, "upperRoman")
        _set_page_num_format(docx_sec._sectPr, "upperRoman", start=start)
        return 1

    def _apply_footer_body(self, docx_sec, page_sc, start):
        footer = docx_sec.footer
        footer.is_linked_to_previous = False
        _build_page_footer(footer._element, page_sc, "decimal")
        _set_page_num_format(docx_sec._sectPr, "decimal", start=start)
        return 1

"""Formula parsing utilities for Word OOXML, LaTeX and text-like formulas."""

from __future__ import annotations

import re
from typing import Iterable

from docx import Document
from lxml import etree

from .ast import FormulaNode, FormulaOccurrence, FormulaParseResult
from .mathtype_ole import extract_ole_formula_candidates
from .normalize import (
    looks_like_bibliographic_reference_text,
    looks_like_caption_text,
    looks_like_formula_text,
    text_formula_to_latex,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

_RE_LATEX_BLOCK = re.compile(r"^\s*\$\$(?P<body>[\s\S]+?)\$\$\s*$")
_RE_LATEX_INLINE_ONLY = re.compile(r"^\s*\$(?!\$)(?P<body>[\s\S]+?)(?<!\\)\$\s*$")
_RE_LATEX_INLINE = re.compile(r"(?<!\\)\$(?!\$)(?P<body>[^$\n]+?)(?<!\\)\$")
_RE_LATEX_COMMAND = re.compile(r"\\[A-Za-z]+\b")
_RE_FORMULA_TOKEN = re.compile(r"(\\[A-Za-z]+|[=+\-*/^_∑∫∏√→←↔])")
_RE_CHINESE = re.compile(r"[\u4e00-\u9fff]")
_RE_PROSE_PUNCT = re.compile(r"[，。；：！？、]")
_RE_STRONG_MATH_ANCHOR = re.compile(r"(=|≈|≠|≤|≥|∑|∫|∏|√|→|←|↔|⇌|±|∂)")
_RE_TRAILING_EQ_NUM = re.compile(
    r"(?:\s|[\u00a0])*(?:[\.．…·•]{2,}\s*)?[\(（]\s*(\d+(?:[\-\.]\d+)+)\s*[\)）]\s*$"
)
_RE_LATEX_PREAMBLE_ONLY = re.compile(
    r"^\s*\\(?:documentclass|usepackage|title|author|date)"
    r"(?:\[[^\]]*\])?(?:\{[\s\S]*\})?\s*$"
)
_LATEX_DISPLAY_BLOCK_PAIRS = {
    r"\[": r"\]",
    "$$": "$$",
    r"\begin{equation}": r"\end{equation}",
    r"\begin{equation*}": r"\end{equation*}",
    r"\begin{align}": r"\end{align}",
    r"\begin{align*}": r"\end{align*}",
    r"\begin{gather}": r"\end{gather}",
    r"\begin{gather*}": r"\end{gather*}",
    r"\begin{multline}": r"\end{multline}",
    r"\begin{multline*}": r"\end{multline*}",
}


def _iter_all_paragraphs(doc: Document) -> Iterable[tuple[object, int, str, bool]]:
    body_indexes: dict[str, int] = {}
    tree = doc.element.body.getroottree()
    idx = -1
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            idx += 1
            body_indexes[tree.getpath(child)] = idx

    visited: set[str] = set()
    for para in doc.paragraphs:
        para_key = para._p.getroottree().getpath(para._p)
        if para_key in visited:
            continue
        visited.add(para_key)
        body_idx = body_indexes.get(para_key, -1)
        yield para, body_idx, f"paragraph[{body_idx}]", False

    table_pos = 0
    for tbl in doc.tables:
        table_pos += 1
        for row_idx, row in enumerate(tbl.rows):
            for col_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    para_key = para._p.getroottree().getpath(para._p)
                    if para_key in visited:
                        continue
                    visited.add(para_key)
                    body_idx = body_indexes.get(para_key, -1)
                    location = f"table[{table_pos}] r{row_idx + 1}c{col_idx + 1} p{p_idx + 1}"
                    yield para, body_idx, location, True


def _is_formula_only_text(text: str) -> bool:
    value = str(text or "").strip()
    if not value:
        return False
    if looks_like_caption_text(value):
        return False
    if looks_like_bibliographic_reference_text(value):
        return False
    value, _ = _strip_trailing_equation_number(value)
    value = value.strip()
    if not value:
        return False
    if len(value) > 240:
        return False
    chinese_count = len(_RE_CHINESE.findall(value))
    if chinese_count and chinese_count > max(2, int(len(value) * 0.25)):
        return False
    return bool(_RE_FORMULA_TOKEN.search(value))


def _strip_trailing_equation_number(text: str) -> tuple[str, str | None]:
    value = str(text or "")
    if not value:
        return "", None
    match = _RE_TRAILING_EQ_NUM.search(value)
    if not match:
        return value, None
    return value[:match.start()].rstrip(), match.group(1)


def _looks_like_non_math_latex_text(text: str) -> bool:
    value = str(text or "").strip()
    if not value:
        return False
    if value in {
        r"\[",
        r"\]",
        "$$",
        r"\begin{document}",
        r"\end{document}",
        r"\maketitle",
        r"\tableofcontents",
    }:
        return True
    return bool(_RE_LATEX_PREAMBLE_ONLY.match(value))


def _looks_like_bare_latex(text: str) -> bool:
    value = str(text or "").strip()
    if not value or "$" in value:
        return False
    if _looks_like_non_math_latex_text(value):
        return False
    commands = re.findall(r"\\[A-Za-z]+", value)
    if not commands:
        return False
    if len(commands) >= 2:
        return True
    if value.startswith("\\") and re.search(r"[{}^_=]", value):
        return True
    if value.startswith("\\") and len(value.split()) <= 16:
        return True
    return False


def _looks_like_prose_omml_text(text: str, *, shape_kind: str = "text") -> bool:
    value = str(text or "").strip()
    if not value:
        return False
    if shape_kind not in {"text", "unknown"}:
        return False
    if looks_like_caption_text(value):
        return True
    if looks_like_bibliographic_reference_text(value):
        return True
    if len(value) > 240:
        return True
    if _RE_STRONG_MATH_ANCHOR.search(value):
        return False
    if value.startswith("\\") and _RE_LATEX_COMMAND.search(value):
        return False

    chinese_count = len(_RE_CHINESE.findall(value))
    prose_punct_count = len(_RE_PROSE_PUNCT.findall(value))
    matched, conf, _ = looks_like_formula_text(value)
    if matched and float(conf) >= 0.78 and chinese_count <= max(2, int(len(value) * 0.12)):
        return False

    if chinese_count >= max(6, int(len(value) * 0.16)) and prose_punct_count >= 1:
        return True
    if chinese_count >= max(12, int(len(value) * 0.28)):
        return True
    if any(
        marker in value
        for marker in ("体积比为", "混合液中", "搅拌", "透析", "分别得到", "分别表示", "其中，")
    ):
        return True
    return False


def _extract_word_formula(paragraph) -> tuple[FormulaNode | None, bool, bool, str]:
    p = paragraph._p
    has_formula = bool(
        p.xpath(
            (
                ".//*[namespace-uri()='%s' and "
                "(local-name()='oMath' or local-name()='oMathPara')]"
            ) % _M_NS
        )
    )
    if not has_formula:
        return None, False, False, ""

    linear = "".join(
        p.xpath(
            ".//*[namespace-uri()='%s' and local-name()='t']/text()" % _M_NS
        )
    ).strip()
    is_block = bool(
        p.xpath(
            ".//*[namespace-uri()='%s' and local-name()='oMathPara']" % _M_NS
        )
    )

    shape_kind, shape_payload = _detect_omml_shape(p)
    if _looks_like_prose_omml_text(linear, shape_kind=shape_kind):
        return None, False, False, ""

    allowed = {
        "oMath", "oMathPara", "r", "t", "rPr", "sty", "scr",
        "f", "num", "den", "sSup", "sSub", "sSubSup", "sPre",
        "e", "sup", "sub", "rad", "deg", "degHide",
        "nary", "naryPr", "chr", "limLoc", "subHide", "supHide",
        "m", "mPr", "mr", "func", "fName",
        "d", "dPr", "begChr", "endChr", "sepChr",
    }
    complex_nodes = []
    for node in p.xpath(".//*[namespace-uri()='%s']" % _M_NS):
        local = node.tag.split("}")[-1]
        if local not in allowed:
            complex_nodes.append(local)
    warnings: list[str] = []
    confidence = 0.95
    if complex_nodes:
        confidence = 0.78 if shape_kind in {"big_operator", "matrix", "delimited", "function"} else 0.58
        warnings.append(
            "complex_omml:" + ",".join(sorted(set(complex_nodes))[:6])
        )
    elif shape_kind not in {"text", "unknown"}:
        confidence = 0.91

    normal_text = "".join(
        p.xpath(
            (
                ".//*[namespace-uri()='%s' and local-name()='t' and "
                "not(ancestor::*[namespace-uri()='%s' and "
                "(local-name()='oMath' or local-name()='oMathPara')])]/text()"
            ) % (_W_NS, _M_NS),
        )
    ).strip()
    normal_text_without_number, _ = _strip_trailing_equation_number(normal_text)
    is_formula_only = not normal_text_without_number.strip()

    node = FormulaNode(
        kind="equation",
        payload={
            "linear_text": linear,
            "is_simple_omml": not complex_nodes,
            "shape": shape_kind,
            "shape_payload": shape_payload,
        },
        children=[],
        source_type="word_native",
        confidence=confidence,
        warnings=warnings,
    )
    return node, is_formula_only, is_block, linear


def _extract_ole_formula(paragraph) -> list[tuple[FormulaNode, bool, bool, tuple[int, int] | None, str]]:
    p = paragraph._p
    has_ole = bool(
        p.xpath(".//*[namespace-uri()='%s' and local-name()='object']" % _W_NS)
    )
    if not has_ole:
        return []

    typed_hits = extract_ole_formula_candidates(paragraph)
    if typed_hits:
        results: list[tuple[FormulaNode, bool, bool, tuple[int, int] | None, str]] = []
        for hit in typed_hits:
            source_type = str(hit.source_type or "ole_equation").strip().lower()
            payload = dict(hit.payload or {})
            candidate_raw = str(
                payload.get("text")
                or paragraph.text
                or ""
            ).strip()
            candidate, _ = _strip_trailing_equation_number(candidate_raw)
            candidate = candidate.strip()
            warnings = list(hit.warnings or [])
            confidence = float(hit.confidence or 0.0)

            if candidate and "latex" not in payload and "normalized_latex" not in payload:
                latex, norm_conf, norm_warnings = text_formula_to_latex(candidate, source_hint=source_type)
                warnings.extend(norm_warnings)
                if latex:
                    payload["latex"] = latex
                    payload["normalized_latex"] = latex
                    confidence = max(confidence, norm_conf)
                else:
                    warnings.append("ole_formula_candidate_not_structured")
            elif not candidate and "latex" not in payload and "diagnostic_code" not in payload:
                warnings.append("ole_formula_text_blank")

            node = FormulaNode(
                kind="equation",
                payload=payload,
                children=[],
                source_type=source_type,
                confidence=confidence,
                warnings=list(dict.fromkeys(warnings)),
            )
            is_formula_only = bool(getattr(hit, "object_only", False)) or _is_formula_only_text(candidate)
            body = str(
                candidate
                or payload.get("latex")
                or payload.get("normalized_latex")
                or ""
            ).strip()
            results.append((node, is_formula_only, is_formula_only, None, body))
        return results

    raw_xml = etree.tostring(p, encoding="unicode").lower()
    if "mathtype" in raw_xml or "equation.dsmt" in raw_xml or "dsmt" in raw_xml:
        source_type = "mathtype"
    elif "equation.3" in raw_xml or "microsoft equation" in raw_xml:
        source_type = "old_equation"
    else:
        source_type = "ole_equation"

    candidate_raw = (paragraph.text or "").strip()
    candidate, _ = _strip_trailing_equation_number(candidate_raw)
    candidate = candidate.strip()
    latex, norm_conf, norm_warnings = text_formula_to_latex(candidate, source_hint=source_type)

    warnings = list(norm_warnings)
    confidence = 0.58
    payload = {"text": candidate}
    if not candidate:
        warnings.append("ole_formula_text_blank")
    if latex:
        payload["latex"] = latex
        payload["normalized_latex"] = latex
        confidence = max(confidence, norm_conf)
    else:
        warnings.append("ole_formula_candidate_not_structured")

    node = FormulaNode(
        kind="equation",
        payload=payload,
        children=[],
        source_type=source_type,
        confidence=confidence,
        warnings=warnings,
    )
    is_formula_only = _is_formula_only_text(candidate)
    return [(node, is_formula_only, is_formula_only, None, candidate)]


def _extract_latex_formulas(paragraph) -> list[tuple[FormulaNode, bool, bool, tuple[int, int] | None, str]]:
    raw_text = paragraph.text or ""
    text, _ = _strip_trailing_equation_number(raw_text)
    stripped = text.strip()
    result: list[tuple[FormulaNode, bool, bool, tuple[int, int] | None, str]] = []

    if _looks_like_non_math_latex_text(stripped):
        return result

    block_match = _RE_LATEX_BLOCK.match(text)
    if block_match:
        body = block_match.group("body").strip()
        node = FormulaNode(
            kind="equation",
            payload={"latex": body},
            children=[],
            source_type="latex",
            confidence=0.92 if body else 0.35,
            warnings=[] if body else ["empty_latex_block"],
        )
        result.append((node, True, True, None, body))
        return result

    inline_only = _RE_LATEX_INLINE_ONLY.match(text)
    if inline_only:
        body = inline_only.group("body").strip()
        node = FormulaNode(
            kind="equation",
            payload={"latex": body},
            children=[],
            source_type="latex",
            confidence=0.90 if body else 0.35,
            warnings=[] if body else ["empty_latex_inline"],
        )
        result.append((node, False, True, None, body))
        return result

    for m in _RE_LATEX_INLINE.finditer(text):
        body = (m.group("body") or "").strip()
        span = m.span()
        warnings: list[str] = []
        confidence = 0.76
        if not body:
            confidence = 0.35
            warnings.append("empty_latex_inline")
        node = FormulaNode(
            kind="equation",
            payload={"latex": body},
            children=[],
            source_type="latex",
            confidence=confidence,
            warnings=warnings,
        )
        is_formula_only = stripped == m.group(0).strip()
        result.append((node, False, is_formula_only, span, body))

    if result:
        return result

    if _looks_like_bare_latex(stripped):
        node = FormulaNode(
            kind="equation",
            payload={"latex": stripped, "latex_source": "bare_text"},
            children=[],
            source_type="latex",
            confidence=0.86,
            warnings=["bare_latex_detected"],
        )
        return [(node, True, True, None, stripped)]

    return result


def _make_multiline_latex_occurrence(
    entries: list[tuple[object, int, str, bool]],
    start_idx: int,
    end_idx: int,
) -> FormulaOccurrence | None:
    start_para, start_body_idx, _, start_in_table = entries[start_idx]
    if start_in_table:
        return None

    body_lines: list[str] = []
    source_para_indices: list[int] = []
    for idx in range(start_idx, end_idx + 1):
        para, body_idx, _, in_table = entries[idx]
        if in_table:
            return None
        if body_idx >= 0:
            source_para_indices.append(body_idx)
        if idx in {start_idx, end_idx}:
            continue
        body_lines.append(para.text or "")

    body = "\n".join(body_lines).strip()
    if not body:
        return None

    end_body_idx = entries[end_idx][1]
    location = (
        f"paragraph[{start_body_idx}]"
        if end_body_idx < 0 or end_body_idx == start_body_idx
        else f"paragraph[{start_body_idx}-{end_body_idx}]"
    )
    node = FormulaNode(
        kind="equation",
        payload={"latex": body, "latex_source": "multiline_display_block"},
        children=[],
        source_type="latex",
        confidence=0.94,
        warnings=["multiline_display_block_detected"],
    )
    return FormulaOccurrence(
        node=node,
        paragraph=start_para,
        paragraph_index=start_body_idx,
        location=location,
        source_type="latex",
        is_block=True,
        is_formula_only=True,
        source_text=body,
        source_range=None,
        in_table=False,
        source_paragraph_indices=source_para_indices,
    )


def _extract_multiline_latex_block(
    entries: list[tuple[object, int, str, bool]],
    start_idx: int,
) -> tuple[FormulaOccurrence, int] | None:
    para, _, _, in_table = entries[start_idx]
    if in_table:
        return None

    start_text = str(getattr(para, "text", "") or "").strip()
    end_marker = _LATEX_DISPLAY_BLOCK_PAIRS.get(start_text)
    if not end_marker:
        return None

    for idx in range(start_idx + 1, len(entries)):
        next_para, _, _, next_in_table = entries[idx]
        if next_in_table:
            break
        if str(getattr(next_para, "text", "") or "").strip() != end_marker:
            continue
        occurrence = _make_multiline_latex_occurrence(entries, start_idx, idx)
        if occurrence is not None:
            return occurrence, idx
        break
    return None


def _extract_text_formula(paragraph) -> list[tuple[FormulaNode, bool, bool, tuple[int, int] | None, str]]:
    text = (paragraph.text or "").strip()
    if not text:
        return []
    if looks_like_caption_text(text):
        return []
    if looks_like_bibliographic_reference_text(text):
        return []
    text, _ = _strip_trailing_equation_number(text)
    text = text.strip()
    if not text:
        return []
    if _looks_like_non_math_latex_text(text):
        return []
    matched, base_conf, source_kind = looks_like_formula_text(text)
    if not matched:
        return []

    latex, norm_conf, warnings = text_formula_to_latex(text, source_hint=source_kind)
    confidence = max(base_conf, norm_conf)
    payload = {"text": text}
    if latex:
        payload["latex"] = latex
        payload["normalized_latex"] = latex
    else:
        warnings = list(warnings) + ["unable_to_generate_latex_candidate"]

    node = FormulaNode(
        kind="equation",
        payload=payload,
        children=[],
        source_type=source_kind,
        confidence=confidence,
        warnings=warnings,
    )
    is_formula_only = _is_formula_only_text(text)
    is_block = is_formula_only
    return [(node, is_block, is_formula_only, None, text)]


def _m_first(parent, local: str):
    if parent is None:
        return None
    matches = parent.xpath(
        "./*[namespace-uri()='%s' and local-name()='%s']" % (_M_NS, local)
    )
    return matches[0] if matches else None


def _m_text(node) -> str:
    if node is None:
        return ""
    return "".join(
        node.xpath(".//*[namespace-uri()='%s' and local-name()='t']/text()" % _M_NS)
    ).strip()


def _m_attr(node, attr: str) -> str:
    if node is None:
        return ""
    key = f"{{{_M_NS}}}{attr}"
    return (node.get(key) or node.get(attr) or "").strip()


def _detect_omml_shape(paragraph_el) -> tuple[str, dict]:
    """Best-effort shape extraction from first math object."""
    math_nodes = paragraph_el.xpath(
        ".//*[namespace-uri()='%s' and local-name()='oMath']" % _M_NS
    )
    if not math_nodes:
        return "text", {}
    math = math_nodes[0]

    candidates = [
        child for child in math
        if child.tag.startswith("{%s}" % _M_NS)
    ]
    if not candidates:
        return "text", {"text": _m_text(math)}
    head = candidates[0]
    local = head.tag.split("}")[-1]

    if local == "f":
        num = _m_text(_m_first(head, "num"))
        den = _m_text(_m_first(head, "den"))
        return "fraction", {"num": num, "den": den}
    if local == "sSup":
        base = _m_text(_m_first(head, "e"))
        sup = _m_text(_m_first(head, "sup"))
        return "superscript", {"base": base, "sup": sup}
    if local == "sSub":
        base = _m_text(_m_first(head, "e"))
        sub = _m_text(_m_first(head, "sub"))
        return "subscript", {"base": base, "sub": sub}
    if local == "sSubSup":
        base = _m_text(_m_first(head, "e"))
        sub = _m_text(_m_first(head, "sub"))
        sup = _m_text(_m_first(head, "sup"))
        return "sub_sup", {"base": base, "sub": sub, "sup": sup}
    if local == "sPre":
        base = _m_text(_m_first(head, "e"))
        sub = _m_text(_m_first(head, "sub"))
        sup = _m_text(_m_first(head, "sup"))
        return "prescript", {"base": base, "sub": sub, "sup": sup}
    if local == "rad":
        body = _m_text(_m_first(head, "e"))
        deg = _m_text(_m_first(head, "deg"))
        return "sqrt", {"body": body, "deg": deg}
    if local == "nary":
        nary_pr = _m_first(head, "naryPr")
        chr_val = _m_attr(_m_first(nary_pr, "chr"), "val")
        operator = {
            "∑": "sum",
            "∫": "int",
            "∬": "iint",
            "∭": "iiint",
            "∮": "oint",
            "∏": "prod",
        }.get(chr_val, "nary")
        sub = _m_text(_m_first(head, "sub"))
        sup = _m_text(_m_first(head, "sup"))
        body = _m_text(_m_first(head, "e"))
        return "big_operator", {
            "operator": operator,
            "symbol": chr_val,
            "sub": sub,
            "sup": sup,
            "body": body,
        }
    if local == "d":
        dpr = _m_first(head, "dPr")
        left = _m_attr(_m_first(dpr, "begChr"), "val") or "("
        right = _m_attr(_m_first(dpr, "endChr"), "val") or ")"
        body = _m_text(_m_first(head, "e"))
        return "delimited", {"left": left, "right": right, "body": body}
    if local == "m":
        rows = []
        for mr in head.xpath("./*[namespace-uri()='%s' and local-name()='mr']" % _M_NS):
            cols = []
            for e in mr.xpath("./*[namespace-uri()='%s' and local-name()='e']" % _M_NS):
                cols.append(_m_text(e))
            if cols:
                rows.append(cols)
        if rows:
            return "matrix", {"rows": rows, "env": "matrix"}
    if local == "func":
        name = _m_text(_m_first(head, "fName"))
        arg = _m_text(_m_first(head, "e"))
        return "function", {"name": name, "arg": arg}

    return "text", {"text": _m_text(math)}


def parse_document_formulas(doc: Document) -> FormulaParseResult:
    """Parse formulas from document body and table cells."""

    occurrences: list[FormulaOccurrence] = []
    entries = list(_iter_all_paragraphs(doc))
    idx = 0
    while idx < len(entries):
        para, body_idx, location, in_table = entries[idx]

        multiline_block = _extract_multiline_latex_block(entries, idx)
        if multiline_block is not None:
            occurrence, end_idx = multiline_block
            occurrences.append(occurrence)
            idx = end_idx + 1
            continue

        word_node, is_formula_only, is_block, linear_text = _extract_word_formula(para)
        if word_node is not None:
            occurrences.append(
                FormulaOccurrence(
                    node=word_node,
                    paragraph=para,
                    paragraph_index=body_idx,
                    location=location,
                    source_type="word_native",
                    is_block=is_block,
                    is_formula_only=is_formula_only,
                    source_text=linear_text,
                    source_range=None,
                    in_table=in_table,
                    source_paragraph_indices=[body_idx] if body_idx >= 0 else [],
                )
            )

        ole_hits = _extract_ole_formula(para)
        for node, ole_block, ole_only, span, body in ole_hits:
            occurrences.append(
                FormulaOccurrence(
                    node=node,
                    paragraph=para,
                    paragraph_index=body_idx,
                    location=location,
                    source_type=node.source_type,
                    is_block=ole_block,
                    is_formula_only=ole_only,
                    source_text=body,
                    source_range=span,
                    in_table=in_table,
                    source_paragraph_indices=[body_idx] if body_idx >= 0 else [],
                )
            )

        latex_hits = _extract_latex_formulas(para)
        for node, latex_block, latex_only, span, body in latex_hits:
            occurrences.append(
                FormulaOccurrence(
                    node=node,
                    paragraph=para,
                    paragraph_index=body_idx,
                    location=location,
                    source_type="latex",
                    is_block=latex_block,
                    is_formula_only=latex_only,
                    source_text=body,
                    source_range=span,
                    in_table=in_table,
                    source_paragraph_indices=[body_idx] if body_idx >= 0 else [],
                )
            )

        if word_node is None and not ole_hits and not latex_hits:
            text_hits = _extract_text_formula(para)
            for node, txt_block, txt_only, span, body in text_hits:
                occurrences.append(
                    FormulaOccurrence(
                        node=node,
                        paragraph=para,
                        paragraph_index=body_idx,
                        location=location,
                        source_type=node.source_type,
                        is_block=txt_block,
                        is_formula_only=txt_only,
                        source_text=body,
                        source_range=span,
                        in_table=in_table,
                        source_paragraph_indices=[body_idx] if body_idx >= 0 else [],
                    )
                )
        idx += 1

    return FormulaParseResult(occurrences=occurrences)
